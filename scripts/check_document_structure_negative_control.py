#!/usr/bin/env python3
"""Negative control for scripts/check_document_structure.py.

A check nobody has seen fail is not evidence. This builds a throwaway repository layout,
reinjects each failure condition the gate claims to catch — including the two conditions
that actually shipped, AMOC's truncated Appendix C and ARCC's missing front and back
matter — and asserts the gate goes RED on every one. It also runs CLEAN cases that must
stay green, because a gate that fires on legitimate work teaches people to ignore it, and
three of these are variations this book genuinely uses.
"""
import os
import re
import shutil
import subprocess
import sys
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GATE = os.path.join('scripts', 'check_document_structure.py')

FULL = ['Headline', 'Valuation summary — every read at a glance',
        'Company overview — the company at a glance',
        '1  Fundamental valuation'] + \
       ['1.%d  Section %d' % (i, i) for i in range(1, 10)] + \
       ['%d  Section' % i for i in range(2, 8)] + \
       ['Appendix A  Financial statements'] + \
       ['A.%d  Statement' % i for i in (1, 2, 3)] + \
       ['Appendix B  Peer frame, risk register and the research register'] + \
       ['B.%d  Part' % i for i in (1, 2, 3)] + \
       ['Appendix C  The expert valuation panel'] + \
       ['C.%d  Part' % i for i in range(1, 7)] + \
       ['About this series', 'Disclosure and disclaimer']


def write_doc(path, heads, style='bold'):
    import docx
    from docx.shared import Pt
    d = docx.Document()
    for h in heads:
        if style == 'heading':
            d.add_paragraph(h, style='Heading 1')
        elif style == 'plain':
            # Deliberately NOT a heading under either convention. The first draft of this
            # file fell through to the bold branch here, so case 6 injected nothing and its
            # green proved only that the document was unchanged — the same defect this
            # repository has already paid for in another negative control.
            d.add_paragraph(h)
        else:
            p = d.add_paragraph()
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(12)
        d.add_paragraph('Body text for %s.' % h)
    d.save(path)


def build(tmp, studies, ratchet=None):
    """studies: {TICKER: (heads|None, style)}"""
    eng = os.path.join(tmp, 'engine')
    os.makedirs(os.path.join(eng, 'build_depth_audit'), exist_ok=True)
    shutil.copy(os.path.join(ROOT, 'engine', 'research_protocol.py'),
                os.path.join(eng, 'research_protocol.py'))
    mr = os.path.join(eng, 'model_report')
    os.makedirs(mr, exist_ok=True)
    shutil.copy(os.path.join(ROOT, 'engine', 'model_report',
                             'MODEL_REPORT_09-08-2026.docx'),
                os.path.join(mr, 'MODEL_REPORT_09-08-2026.docx'))
    os.makedirs(os.path.join(tmp, 'scripts'), exist_ok=True)
    shutil.copy(os.path.join(ROOT, GATE), os.path.join(tmp, GATE))
    for tk, (heads, style) in studies.items():
        sd = os.path.join(eng, '%s_study' % tk.lower())
        os.makedirs(sd, exist_ok=True)
        if heads is not None:
            write_doc(os.path.join(sd, '%s_Valuation_Study_03-09-2026.docx' % tk),
                      heads, style)
    import json
    json.dump({'entries': ratchet or {}},
              open(os.path.join(eng, 'build_depth_audit',
                                'document_outstanding.json'), 'w'))
    return tmp


def run(tmp):
    r = subprocess.run([sys.executable, GATE], cwd=tmp, capture_output=True, text=True)
    return r.returncode, (r.stdout or '') + (r.stderr or '')


CASES = []


def case(name, must_fail, studies, ratchet=None, expect=None, mutate=None):
    CASES.append((name, must_fail, studies, ratchet, expect, mutate))


# ---- conditions that MUST go red -------------------------------------------------
case('1. AMOC as it shipped — Appendix C stops after C.3, no Company overview',
     True, {'AMOC': ([h for h in FULL
                      if not h.startswith(('Company overview', 'C.1', 'C.2', 'C.3',
                                           'C.4', 'C.5', 'C.6', 'About this'))], 'bold')},
     expect='C.4')
case('2. ARCC as it shipped — no Headline, summary, overview, About or Disclosure, '
     'and A/B/C sub-parts unnumbered',
     True, {'ARCC': ([h for h in FULL
                      if not re.match(r'^(Headline|Valuation summary|Company overview|'
                                      r'[ABC]\.\d|About this|Disclosure)', h)], 'bold')},
     expect='Headline')
case('3. one section missing — the About section alone',
     True, {'X': ([h for h in FULL if not h.startswith('About this')], 'bold')},
     expect='About this series')
case('4. every section present but out of the model order',
     True, {'X': (FULL[:3] + FULL[19:] + FULL[3:19], 'bold')},
     expect='out of the model order')
case('5. no valuation study document in the directory at all',
     True, {'X': (None, 'bold')}, expect='no valuation study document')
case('6. the document opens but yields no headings — must REFUSE, never report '
     'every section missing',
     True, {'X': (FULL, 'plain')}, expect='headings unreadable')
case('7. a listed study no longer resolves on disk [R-ENF-04]',
     True, {'X': (FULL, 'bold')}, ratchet={'GONE': 'was failing'},
     expect='no longer resolve on disk')
case('8. the population is empty — no study directories at all [R-ENF-04]',
     True, {}, expect='examined zero study directories')
case('9. a NEW study with no ratchet entry either way',
     True, {'NEW': ([h for h in FULL if not h.startswith('C.5')], 'bold'),
            'X': (FULL, 'bold')}, expect='C.5')
case('10. the two records of the standard disagree — the skeleton stops naming a '
     'section the model report carries',
     True, {'X': (FULL, 'bold')}, expect='REFUSING',
     mutate=('research_protocol', 'B.1 peers and the sector frame; ', ''))

# ---- clean cases that MUST stay green ---------------------------------------------
case('11. CLEAN — the model report skeleton, bold-run headings',
     False, {'X': (FULL, 'bold')})
case('12. CLEAN — the same sections under Word Heading styles (PHDC/TMGH convention)',
     False, {'X': (FULL, 'heading')})
case('13. CLEAN — numbers with a trailing period and "About this study" '
     '(the ADNOCDIST/TMGH convention)',
     False, {'X': ([re.sub(r'^(\d(?:\.\d)?)\s', r'\1. ', h).replace(
         'About this series', 'About this study') for h in FULL], 'bold')})
case('14. CLEAN — a study carrying EXTRA sections beyond the model (ARCC 1.10, 1.11)',
     False, {'X': (FULL[:13] + ['1.10  Extra', '1.11  Extra'] + FULL[13:], 'bold')})
case('15. CLEAN — a known breach that is on the ratchet stays green',
     False, {'X': ([h for h in FULL if not h.startswith('C.6')], 'bold')},
     ratchet={'X': 'missing 1 section(s): C.6'})
case('16. CLEAN — a superseded off-standard edition beside a good current one',
     False, {'X': (FULL, 'bold')}, mutate=('extra_old', None, None))


def main():
    bad = 0
    for i, (name, must_fail, studies, ratchet, expect, mutate) in enumerate(CASES, 1):
        tmp = tempfile.mkdtemp(prefix='ncdoc')
        try:
            build(tmp, studies, ratchet)
            if mutate and mutate[0] == 'research_protocol':
                p = os.path.join(tmp, 'engine', 'research_protocol.py')
                s = open(p, encoding='utf-8').read()
                assert mutate[1] in s, 'mutation anchor missing: %r' % mutate[1]
                open(p, 'w', encoding='utf-8').write(s.replace(mutate[1], mutate[2]))
            if mutate and mutate[0] == 'extra_old':
                write_doc(os.path.join(tmp, 'engine', 'x_study',
                                       'X_Valuation_Study_01-08-2026.docx'),
                          FULL[:4], 'bold')
            rc, out = run(tmp)
            red = rc != 0
            ok = (red == must_fail) and (expect is None or expect in out)
            print('%-4s %s' % ('PASS' if ok else 'FAIL', name))
            if not ok:
                bad += 1
                print('      rc=%d  wanted %s%s' % (
                    rc, 'RED' if must_fail else 'GREEN',
                    (' containing %r' % expect) if expect else ''))
                print('      ' + '\n      '.join(out.strip().splitlines()[-12:]))
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    print('\n%d/%d cases behaved as specified' % (len(CASES) - bad, len(CASES)))
    return 1 if bad else 0


if __name__ == '__main__':
    sys.exit(main())
