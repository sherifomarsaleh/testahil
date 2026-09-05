#!/usr/bin/env python3
"""A NEW STUDY MUST CLEAR EVERY GATE, AND THIS PROVES IT RATHER THAN ASSERTING IT.

[R-ENF-07]

WHY THIS EXISTS
    The point of a ratcheted gate is that knowingly-outstanding work is listed and allowed
    to fail while the build breaks on a NEW violation. Every gate in this repository says
    so in its own docstring, and every one is negative-controlled on its own conditions.
    What none of them tests is the claim the whole design rests on:

        A STUDY DIRECTORY CREATED TOMORROW, WITH NOTHING IN IT, GOES RED EVERYWHERE.

    That is a property of the SYSTEM rather than of any gate, so no gate can check it. It
    is also the exact claim the programme was asked to deliver — that a study is produced
    correctly and passes several checks without anyone intervening — and it is the kind of
    claim that is true by construction right up to the day a ratchet is seeded one entry
    too generously, or a gate globs a pattern a new directory happens not to match, or a
    check skips a study whose numbers file will not parse.

    So this walks the whole set: it copies the repository into a sandbox, plants an empty
    study directory called ZZTEST_study, and runs every ratcheted gate. Each one must go
    RED and must NAME the new study. A gate that stays green on a study with no numbers, no
    documents, no workbook, no sweep and no records is a gate a new name can walk past.

WHAT A PASS MEANS, AND WHAT IT DOES NOT
    A pass means every gate in the set REFUSES an unknown study directory. It does not mean
    the gates are individually right — each has its own negative control for that — and it
    does not mean a study cannot be wrong in a way no gate models. It means the guideline
    BINDS on a new name rather than depending on whoever builds it remembering to look.

USAGE
    python3 scripts/check_new_study_gauntlet.py           # gate
    python3 scripts/check_new_study_gauntlet.py --verbose # each gate's own last line
"""
import glob
import json
import os
import shutil
import subprocess
import re
import sys
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TICKER = 'ZZTEST'

# THE SET IS SPLIT, AND THE SPLIT IS THE FINDING [measured 03-Sep-2026 on the first run].
# A single list of "study gates" turned out to conflate two different kinds of check, and
# four of the seventeen stayed green on an empty directory for a perfectly good reason:
# they bite on ARTEFACTS, not on the directory. A study with no delivered documents cannot
# leak internal vocabulary and cannot publish a retired blend. Demanding that those gates
# refuse an empty directory would be a false claim about what they check — so they are
# tested the way they actually work, by planting a MINIMAL OFFENDING ARTEFACT and asserting
# they catch it.
#
# a builder that READS the artefact, which is what makes the artefact gate's subject exist
BUILDER_STUB = "import json\nd = json.load(open('diagnostics.json'))\n"

# DIRECTORY GATES: must refuse a study directory that exists and holds nothing. These are
# the checks a new name cannot walk past by simply not producing something.
DIRECTORY_GATES = [
    'check_study_provenance.py',
    'check_rebuild_ledger.py',
    'check_workbook_structure.py',
    'check_document_structure.py',
    'check_sweep_module.py',
    'check_prose_figures.py',
    'check_valuation_gap.py',
    'check_macro_coherence.py',
    'check_bridge.py',
    'check_lens_design.py',
    'check_cost_of_capital.py',
    'check_output_records.py',
    'check_forecast_anchor.py',
    'check_delivered_pdf_currency.py',
    'check_table_footing.py',
    'check_source_integrity.py',
    # ADDED 05-Sep-2026, and how it got here is the finding beside it. This gate has
    # always run over every study directory — through engine/valuation_calibration/
    # terminal_census.census(), which does the glob — so its own source carried no
    # `_study` and the detector below could not see it. It became visible only when an
    # unrelated comment in it happened to mention a study path. An empty directory has no
    # numbers file, the census reports it unreadable, and the gate refuses it by name.
    'check_terminal_floor.py',
    # ADDED 05-Sep-2026, BY THIS FILE'S OWN REFUSAL rather than by anyone remembering. Both
    # gates were adopted the same day and neither was listed here, so the gauntlet reported
    # 29 of 29 refusing while two study-scoped gates on disk had been tested by nothing — the
    # failure shape this file exists to close, occurring inside it. Both are directory-scoped:
    # an empty study directory declares no walk-forward scope and carries no recalculation
    # instrument, neither has a ratchet entry, and each refuses BY NAME.
    'check_walkforward_scope.py',
    'check_workbook_values.py',
]

# ARTEFACT GATES: bite once the study produces the artefact they read, and are tested by
# planting one that should trip them. `plant` returns the files to create.
ARTEFACT_GATES = {
    'check_delivered_vocabulary.py': (
        'a delivered document naming a standing rule',
        lambda: {'%s_Valuation_Study_03-09-2026.docx' % TICKER: ('docx', 'Adopted under '
                                                                 '[R-GAP-01] in September.')}),
    'check_waterfall_assertions.py': (
        'a delivered table instructing a reader whose builder never checks it',
        lambda: {'%s_Valuation_Study_03-09-2026.docx' % TICKER: ('docx', [
            ['EGP million', 'Value'],
            ['Enterprise value', '6,617'],
            ['Plus net cash', '4,930'],
            ['Equity value', '11,426'],
        ]),
            'build_it.py': ('py', BUILDER_STUB)}),
    'check_band_vocabulary.py': (
        'a delivered document carrying the retired verdict vocabulary',
        lambda: {'%s_Valuation_Study_03-09-2026.docx' % TICKER: (
            'docx', 'The cone FAILED CALIBRATION TEST over the resolved windows.')}),
    'check_column_widths.py': (
        'a delivered table whose column is too narrow for the figure it prints',
        lambda: {'%s_Valuation_Study_03-09-2026.docx' % TICKER: ('docx', (
            [['EGP million', 'FY2025'], ['Revenue', '(1,234,567.89)']], [2.0, 0.18]))}),
    'check_edition_date.py': (
        'a delivered document whose masthead disagrees with its own filename',
        lambda: {'%s_Valuation_Study_03-09-2026.docx' % TICKER: (
            'docx', 'Valuation study issued 1 January 2026')}),
    'check_site_data_reader.py': (
        'a study script reading assets/data.js with a regex instead of a real parse',
        lambda: {'read_it.py': ('py', "import re\n"
                                      "src = open('assets/data.js').read()\n"
                                      "m = re.search(r'levels', src)\n")}),
    'check_figure_axes.py': (
        'a figure script drawing a reference line outside its own axes',
        lambda: {'figures.py': ('py', "import matplotlib\n"
                                      "matplotlib.use('Agg')\n"
                                      "import matplotlib.pyplot as plt\n"
                                      "fig, ax = plt.subplots()\n"
                                      "ax.plot([1, 2, 3], [1, 2, 3])\n"
                                      "ax.set_ylim(0, 3)\n"
                                      "ax.axhline(99.0, color='r')\n"
                                      "fig.savefig('fig_out.png')\n")}),
    'check_sign_convention.py': (
        'a delivered table printing deductions in two sign conventions at once',
        lambda: {'%s_Valuation_Study_03-09-2026.docx' % TICKER: ('docx', [
            ['AED million', '2026E'],
            ['Less cash operating expenses', '(2,650)'],
            ['Less capital expenditure', '(1,012)'],
            ['Less increase in working capital', '440'],
            ['Free cash flow to the firm', '4,368'],
        ])}),
    'check_artefact_currency.py': (
        'a builder-read JSON carrying a central and declaring no vintage',
        lambda: {'diagnostics.json': ('json', {'central': 12.34, 'note': 'no declaration'}),
                 'build_it.py': ('py', BUILDER_STUB)}),
    'check_eps_reconciliation.py': (
        'a committed numbers file whose profit over shares does not reproduce the '
        'reported earnings per share, with nothing naming the difference',
        lambda: {'study_numbers.json': ('json', {
            'meta': {'shares_mn': 1000.0},
            'inputs': {
                'npa_fy25': {'value': 8000.0, 'source': 'audited FY2025 statements',
                             'date': '2026-02-01', 'ring': 'Company'},
                'eps_fy25': {'value': 7.00, 'source': 'audited FY2025 statements',
                             'date': '2026-02-01', 'ring': 'Company'}}})}),
    'check_harness_outputs.py': (
        'a pricing harness that can write the committed numbers file on the override path',
        lambda: {'compute.py': ('py',
                                'import json, os\n'
                                "BETA_OVERRIDE = os.environ.get('BETA_OVERRIDE')\n"
                                'beta = float(BETA_OVERRIDE) if BETA_OVERRIDE else 0.488\n'
                                "json.dump({'beta': beta}, "
                                "open('study_numbers.json', 'w'))\n")}),
    'check_source_rebinding.py': (
        'a source constant rebound after inputs were registered against it',
        lambda: {'compute.py': ('py',
                                'FS25 = "Consolidated Financial Statements FY2025"\n'
                                "inp('a', 1.0, FS25, '2025-12-31', 'COMPANY')\n"
                                'FS25 = "Annual Report 2025, note 15"\n'
                                "inp('b', 2.0, FS25, '2025-12-31', 'COMPANY')\n")}),
}

# NOT IN EITHER SET, and each with the reason, because a name in a list that resolves to
# the wrong subject is worse than an absence [R-ENF-04]:
#   check_valuation_inputs.py       anchors on WALK-FORWARD run directories, not on study
#                                   directories — a new study is not its subject
#   check_calibration_deliverables.py  anchors on the campaign queue's calibrated names
#   check_lens_vocabulary.py        reads delivered PDFs; with no PDF there is nothing to
#                                   read, and its own population anchoring covers the case
#                                   where the whole book has none
EXCLUDED = {
    'check_valuation_inputs.py': 'anchors on walk-forward run directories, not study '
                                 'directories',
    'check_calibration_deliverables.py': "anchors on the campaign queue's calibrated names",
    'check_lens_vocabulary.py': 'reads delivered PDFs; an empty study has none',
    'check_page_integrity.py': "its subject is the site's ticker pages; the only mention "
                               'of a study directory in it is a comment',
    'check_published_coverage.py': 'holds PUBLISHED fair values against the studies behind '
                                   'them, so a study directory is its reference set rather '
                                   'than its subject — a new directory publishes nothing '
                                   'and there is correctly nothing to refuse',
    'check_publish_block.py': '[R-GAP-02] deliberately carries NO ratchet because it blocks '
                              'a FUTURE act rather than condemning a past one, and a held '
                              'study is not a red build. It reads the new study as '
                              'unreadable and HOLDS it, which is the rule working; '
                              'demanding a nonzero exit would contradict the rule',
    'check_new_study_gauntlet.py': 'this file',
    'check_exemplar_debt.py': 'its subject is the model report and the ratchets naming '
                              'it; a new empty study directory is not the exemplar and '
                              'there is correctly nothing for it to refuse',
    'check_lessons_register.py': 'anchors on the WALK-FORWARD run directories and the '
                                 'lessons behind them, not on study directories',
    'check_protocol_text.py': 'reads the two governing documents; it names study '
                              'directories only to check that what they claim exists '
                              'does exist',
    'check_walkforward_actuation.py': 'anchors on walk-forward runs, not on study '
                                      'directories',
}


# THE LISTS ARE HAND-MAINTAINED, WHICH MEANS A GATE ADDED TOMORROW IS SILENTLY UNTESTED
# AND THIS FILE STILL REPORTS CLEAN — the [R-ENF-04] species inside the very check written
# to close it. The completeness clause below reads the gates on disk rather than trusting
# the lists: any script that resolves study directories for itself is study-scoped, and
# must be named in EXACTLY ONE of the three lists. A new gate then fails this run until
# somebody says which kind it is, which is the whole point of the file.
STUDY_SCOPED = re.compile(r"engine['\"/,\s]{0,4}\*_study|_study['\"]?\s*\)|glob\("
                          r"[^)]*_study")


def study_scoped_gates(repo):
    """Every gate on disk that resolves study directories for itself."""
    out = set()
    for path in sorted(glob.glob(os.path.join(repo, 'scripts', 'check_*.py'))):
        name = os.path.basename(path)
        if name.endswith('_negative_control.py'):
            continue                     # a control is evidence about a gate, not a gate
        try:
            src = open(path, encoding='utf-8').read()
        except OSError:
            continue
        # ANY script that RESOLVES a study directory, not only one that globs the literal
        # pattern. Five gates construct the path instead — "engine/%s_study" % ticker —
        # and the first draft of this clause could not see them, which is the same
        # under-detection it exists to close, one level down. A comment about "the study"
        # carries no underscore, so this stays exact rather than becoming a word search.
        #
        # AND A GATE MAY DELEGATE THE RESOLUTION ENTIRELY [widened 05-Sep-2026].
        # check_terminal_floor.py runs over every study directory through
        # terminal_census.census(), which does the glob, so its own source carried no
        # `_study` at all and this detector could not see it — for as long as it has
        # existed. It surfaced only because an unrelated comment added to that gate
        # happened to name a study path, which is luck rather than a check. So the
        # first-party modules a gate imports are read too, one level down: a gate that
        # hands its population to a shared instrument is still a gate over that
        # population.
        if '_study' in src or _imports_a_study_resolver(repo, src):
            out.add(name)
    return out


def _imports_a_study_resolver(repo, src):
    """Does this script import a first-party module that resolves study directories?"""
    # NO REGEX HERE, DELIBERATELY. [R-ENF-03]'s gate refuses a regular-expression call in
    # any file that reads assets/data.js, and this file does — it plants a study script
    # reading data.js as one of its own fixtures. A pattern-scan call added here for an
    # unrelated purpose tripped that gate the first time it ran, which is the right
    # outcome: the rule is about the FILE, not about which line the call sits on. Splitting
    # the line is enough for an import statement and needs no pattern.
    #
    # AND THE COMMENT EXPLAINING THAT TRIPPED IT TOO, because it named the call. Worth
    # leaving recorded rather than tidied away: a shape-matching check cannot tell a call
    # from a description of one, which is the cost of shape-matching and is why it is only
    # used where the shape cannot occur innocently. Here it can, in a comment — so the
    # comment says what happened without writing the call.
    mods = set()
    for line in src.splitlines():
        head = line.strip()
        if head.startswith('import ') or head.startswith('from '):
            parts = head.split()
            if len(parts) >= 2:
                mods.add(parts[1].rstrip(','))
    for m in mods:
        rel = m.replace('.', os.sep)
        for cand in (os.path.join(repo, rel + '.py'),
                     os.path.join(repo, rel, '__init__.py'),
                     os.path.join(repo, 'engine', rel + '.py'),
                     os.path.join(repo, 'scripts', rel + '.py')):
            if os.path.exists(cand):
                try:
                    if '_study' in open(cand, encoding='utf-8').read():
                        return True
                except OSError:
                    continue
    return False


def sandbox():
    """A copy of the repository with one empty study directory planted in it.

    Copied rather than mutated: a gate that rebuilds a ratchet, or a --prune run reached by
    accident, must not be able to touch the real tree. This repository has already paid for
    running repo-mutating steps against a live checkout once.
    """
    tmp = tempfile.mkdtemp(prefix='gauntlet_')
    def ignore(d, names):
        # raw_indices was excluded in the first draft and check_study_provenance CRASHED on
        # its absence — going red for the wrong reason, which reads exactly like going red
        # for the right one. An excluded directory a gate needs is a sandbox defect
        # masquerading as a finding [R-ENF-04].
        # raw_ohlc (16MB) and panels (2.6MB) are COPIED, not excluded, and the cost is
        # accepted deliberately. The first draft excluded raw_indices and a gate crashed
        # on the absence, going red for the WRONG reason — which reads exactly like going
        # red for the right one [R-ENF-07]. Excluding these two reproduced that failure in
        # three more gates the day the lists were completed: page_integrity refused for a
        # missing OHLC population, band_vocabulary for a legacy page whose panel was not
        # there, figure_axes for figure scripts that could not run. A SANDBOX DEFECT
        # MASQUERADING AS A FINDING IS THE MORE DANGEROUS OUTCOME, because it is green-
        # looking evidence that a gate works.
        return [n for n in names
                if n in ('.git', '__pycache__', 'node_modules', 'filings')]
    shutil.copytree(ROOT, os.path.join(tmp, 'repo'), ignore=ignore, symlinks=True)
    repo = os.path.join(tmp, 'repo')
    os.makedirs(os.path.join(repo, 'engine', '%s_study' % TICKER.lower()), exist_ok=True)
    return tmp, repo


# THE TOOL CONTRACT. A gate that cannot run because a TOOL is absent exits 2, never
# 1: both are failures and neither is clean [R-ENF-04], but they send a reader to
# different repairs — one to the gate or the study, the other to the environment.
TOOL_EXIT = 2


def run(repo, gate):
    r = subprocess.run([sys.executable, os.path.join('scripts', gate)],
                       cwd=repo, capture_output=True, text=True, timeout=1800)
    out = (r.stdout or '') + (r.stderr or '')
    lines = [l for l in out.strip().splitlines() if l.strip()]
    return r.returncode, out, (lines[-1] if lines else '')


def plant(sdir, files):
    """Create the artefacts an artefact gate needs as its subject."""
    for name, (kind, payload) in files.items():
        path = os.path.join(sdir, name)
        if kind == 'docx':
            import docx
            d = docx.Document()
            if isinstance(payload, tuple):
                # (rows, widths_in_inches) — a gate whose subject is a COLUMN WIDTH cannot
                # be tested by a table that does not carry one
                rows_, widths = payload
                t = d.add_table(rows=len(rows_), cols=len(rows_[0]))
                t.autofit = False
                for j, w in enumerate(widths):
                    t.columns[j].width = docx.shared.Inches(w)
                for i, row in enumerate(rows_):
                    for j, cell in enumerate(row):
                        t.cell(i, j).text = str(cell)
                        t.cell(i, j).width = docx.shared.Inches(widths[j])
            elif isinstance(payload, list):
                # a TABLE fixture: [[cell, ...], ...]. A gate whose subject is a table
                # cannot be given a paragraph and be said to have been tested.
                t = d.add_table(rows=len(payload), cols=len(payload[0]))
                for i, row in enumerate(payload):
                    for j, cell in enumerate(row):
                        t.cell(i, j).text = str(cell)
            else:
                d.add_paragraph(payload)
            d.save(path)
        elif kind == 'json':
            json.dump(payload, open(path, 'w', encoding='utf-8'))
        else:
            open(path, 'w', encoding='utf-8').write(payload)


def main(argv):
    verbose = '--verbose' in argv
    tmp, repo = sandbox()
    sdir = os.path.join(repo, 'engine', '%s_study' % TICKER.lower())
    try:
        print('NEW-STUDY GAUNTLET — an empty engine/%s_study/ planted in a sandbox copy'
              % TICKER.lower())
        print('%d directory gates, %d artefact gates, %d excluded with a stated reason\n'
              % (len(DIRECTORY_GATES), len(ARTEFACT_GATES), len(EXCLUDED)))

        red, wrong, missing, notool = [], [], [], []

        print('DIRECTORY GATES — must refuse a study directory that holds nothing')
        for gate in DIRECTORY_GATES:
            if not os.path.exists(os.path.join(repo, 'scripts', gate)):
                missing.append(gate)
                continue
            try:
                rc, out, last = run(repo, gate)
            except Exception as e:                                      # noqa: BLE001
                rc, out, last = 1, '', '%s: %s' % (type(e).__name__, e)
            named = TICKER.lower() in out.lower()
            # EXIT 2 IS THE TOOL CONTRACT: the gate could not run, so it has refused
            # nothing and its silence says nothing about the study. Counting that as a
            # permissive gate is the gauntlet's own first-run finding arriving a second
            # time — on 04-Sep-2026 three gates hit it in CI (poppler installed 320 lines
            # below this step, matplotlib never installed at all) and this file reported
            # "3 gate(s) did not refuse the new study". They had refused nothing because
            # they had run nothing, and those are different repairs. It still FAILS the
            # run [R-ENF-04]; what changes is which way it sends the reader.
            if rc == TOOL_EXIT:
                notool.append((gate, rc, named, last))
                print('   %-4s %-40s exit %d   (could not run — broken tool)'
                      % ('TOOL', gate, rc))
                if verbose:
                    print('        %s' % last[:150])
                continue
            ok = rc != 0 and named
            (red if ok else wrong).append((gate, rc, named, last))
            print('   %-4s %-40s exit %d%s' % ('RED ' if ok else 'MISS', gate, rc,
                                               '' if named else '   (does not name it)'))
            if verbose:
                print('        %s' % last[:150])

        print('\nARTEFACT GATES — bite once the study produces what they read')
        for gate, (what, mk) in ARTEFACT_GATES.items():
            if not os.path.exists(os.path.join(repo, 'scripts', gate)):
                missing.append(gate)
                continue
            for f in os.listdir(sdir):
                os.remove(os.path.join(sdir, f))
            try:
                plant(sdir, mk())
                rc, out, last = run(repo, gate)
            except Exception as e:                                      # noqa: BLE001
                rc, out, last = 0, '', '%s: %s' % (type(e).__name__, e)
            named = TICKER.lower() in out.lower()
            if rc == TOOL_EXIT:                       # see the note above
                notool.append((gate, rc, named, last))
                print('   %-4s %-40s exit %d   (could not run — broken tool)'
                      % ('TOOL', gate, rc))
                if verbose:
                    print('        %s' % last[:150])
                continue
            ok = rc != 0 and named
            (red if ok else wrong).append((gate, rc, named, last))
            print('   %-4s %-40s exit %d   %s' % ('RED ' if ok else 'MISS', gate, rc, what))
            if verbose:
                print('        %s' % last[:150])
        for f in os.listdir(sdir):
            os.remove(os.path.join(sdir, f))

        print('\nEXCLUDED, each with its reason')
        for gate, why in EXCLUDED.items():
            print('   %-45s %s' % (gate, why))
            if not os.path.exists(os.path.join(repo, 'scripts', gate)):
                missing.append(gate)

        print('\n%d of %d gates refuse a new study'
              % (len(red), len(red) + len(wrong)))
        rc = 0
        if missing:
            print('\nFAIL — %d gate(s) named in this file do not exist: %s'
                  % (len(missing), ', '.join(sorted(set(missing)))))
            print('The lists ARE the claim; a name resolving to nothing means the claim is '
                  'about a gate that is not there [R-ENF-04].')
            rc = 1
        unlisted = sorted(study_scoped_gates(repo)
                          - set(DIRECTORY_GATES) - set(ARTEFACT_GATES) - set(EXCLUDED))
        if unlisted:
            print('\nFAIL — %d study-scoped gate(s) on disk are named in none of the three '
                  'lists: %s' % (len(unlisted), ', '.join(unlisted)))
            print('A gate nobody listed is a gate this run never tested, and the run still '
                  'reports clean — which is the failure shape this whole file exists to '
                  'close, occurring inside it. Say which kind it is.')
            rc = 1
        if notool:
            print('\nFAIL — %d gate(s) COULD NOT RUN; a tool they need is absent, so they '
                  'refused nothing and their silence is not evidence [R-ENF-04]:'
                  % len(notool))
            for gate, code, _named, last in notool:
                print('   %-40s exit %d' % (gate, code))
                print('        %s' % last[:170])
            print('Fix the ENVIRONMENT, not the gate: every tool a gate needs is installed '
                  'before any gate runs. This is not a finding about the study.')
            rc = 1
        if wrong:
            print('\nFAIL — %d gate(s) did not refuse the new study, or went red without '
                  'naming it:' % len(wrong))
            for gate, code, named, last in wrong:
                print('   %-40s exit %d  names it: %s' % (gate, code, named))
                print('   %-40s   %s' % ('', last[:130]))
            print('\nA gate that does not refuse a study with no numbers, no documents, no '
                  'workbook, no sweep and no records is one a new name can walk past. '
                  'Either it should cover a new study and does not, or it belongs in '
                  'ARTEFACT_GATES or EXCLUDED — and saying WHICH is the whole point of '
                  'this file.')
            rc = 1
        if rc == 0:
            print('\nOK — every directory gate refuses an empty study by name, and every '
                  'artefact gate catches a planted offender. The guideline binds on a new '
                  'study rather than on whoever builds it remembering to look.')
        return rc
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == '__main__':
    sys.exit(main(sys.argv[1:]))
