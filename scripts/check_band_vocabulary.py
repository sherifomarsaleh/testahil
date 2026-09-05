#!/usr/bin/env python3
"""[R-CAL-02] / [R-ENF-01] Gate: no skill-verdict vocabulary on a public surface,
and no published band record that disagrees with the live panels.

Checked from OUTSIDE the pages it governs, and it FAILS rather than warns. The
rule this enforces was written down before today; what was missing was anything
looking at the pages from outside, which is how riyadhcable.html sat for weeks
claiming 13 resolved windows against a panel holding 10.

The vocabulary table itself lives in engine/band_record.py and is shared with
assert_no_verdict_tokens(), so the gate and the generators cannot reach opposite
conclusions about the same string.

Run:  python3 scripts/check_band_vocabulary.py [--root DIR]
"""
import argparse
import glob

# THE READERS ARE IMPORTED AT MODULE SCOPE, DELIBERATELY. Both were originally imported
# inside the loops that use them, and on 5 September 2026 that turned a missing openpyxl
# in one workflow into TWENTY-TWO lines reading "could not be read ... an unreadable
# workbook is not a clean one" — every word of which is this gate's own correct rule, and
# none of which said the real thing, which was that a dependency was absent. A missing
# dependency should fail once, loudly, at import; twenty-two of them dressed as findings
# is red for the wrong reason, and it reads exactly like red for the right one.
import openpyxl
from docx import Document
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from engine import band_record as br  # noqa: E402

DEFAULT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Surfaces whose verdict text is BAKED INTO AN IMAGE and so cannot be read here.
# engine/metal_backtest.py renders every assets/calibration_*.png with the skill
# verdict in its subtitle. That is a real public surface this gate is blind to,
# so the SOURCE of the caption is scanned instead, and the already-rendered
# figures are carried as a ratchet: listed, reported, allowed to fail, and the
# list may only ever shrink ([R-ENF-02] — a permanently red check is one everyone
# learns to ignore, and a gate with no release is a stall).
FIGURE_CAPTION_SOURCE = os.path.join("engine", "metal_backtest.py")
OUTSTANDING = os.path.join("engine", "build_depth_audit", "band_outstanding.json")



def _ratchet_json(root):
    try:
        return json.load(open(os.path.join(root, OUTSTANDING), encoding="utf-8"))
    except Exception:
        return {}


def _workbook_key(fp):
    """Newest edition of a workbook. The dates are DDMMYYYY, run together, so they are
    PARSED rather than sorted lexically — 03092026 sorts below 09082026 as a string and
    would pick the wrong file, which is the mistake the document resolver beside this one
    records having already made once."""
    m = re.search(r"(\d{2})(\d{2})(\d{4})", os.path.basename(fp))
    if m:
        return (m.group(3), m.group(2), m.group(1))
    m = re.search(r"(\d{2})-(\d{2})-(\d{4})", os.path.basename(fp))
    return (m.group(3), m.group(2), m.group(1)) if m else ("", "", "")


def surfaces(root):
    for f in sorted(glob.glob(os.path.join(root, "*.html")) + glob.glob(os.path.join(root, "legacy", "*.html"))):
        yield f
    # Every reader-facing asset script, not a hand-listed two: markets.js and
    # fv_overlay.js are equally rendered and were going unscanned.
    for f in sorted(glob.glob(os.path.join(root, "assets", "*.js"))):
        yield f


# Fields that exist to record the INTERNAL verdict and are rendered by nothing.
# The verdict is still the Step 0 gate and the ledger note is its audit trail —
# scanning them would fail this gate on the very record the protocol says to keep.
# Verified: no reader of LEDGER `.note` exists in any .html or .js on the site.
INTERNAL_FIELDS = {"note"}


def js_reader_text(path):
    """Every string a .js file can put in front of a reader, comments excluded.

    By LOADING the file rather than regex-stripping comments: `//` also opens
    every https:// URL in the file, and a hand-rolled stripper on a 6,700-line
    data file is the kind of parser-substitute this repo has been bitten by
    before. Loading removes comments by construction.
    """
    import subprocess
    # Top-level `const` in a vm script creates a LEXICAL binding, not an own
    # property of the context object — so walking Object.keys(ctx) reaches
    # nothing a data file declares. The names are collected and re-exported
    # explicitly. (Caught by the negative control: without this the coverage.js
    # scan passed while seeing an empty object.)
    names = re.findall(r"^(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=",
                       open(path, encoding="utf-8").read(), re.M)
    if not names:
        raise RuntimeError("no top-level declarations found")
    export = ",".join(dict.fromkeys(names))
    js = """
      const fs=require('fs'), vm=require('vm'), ctx={window:{}};
      vm.createContext(ctx);
      vm.runInContext(fs.readFileSync(process.argv[1],'utf8')
                      + ';globalThis.__ALL={%s};', ctx);
      const drop = new Set(%s);
      const out = [];
      const seen = new Set();
      (function walk(v){
        if (v === null || typeof v !== 'object') { if (typeof v === 'string') out.push(v); return; }
        if (seen.has(v)) return; seen.add(v);
        for (const k of Object.keys(v)) { if (!drop.has(k)) walk(v[k]); }
      })(ctx.__ALL);
      console.log(JSON.stringify(out));
    """ % (export, json.dumps(sorted(INTERNAL_FIELDS)))
    r = subprocess.run(["node", "-e", js, path], capture_output=True, text=True)
    if r.returncode != 0:
        raise RuntimeError(r.stderr.strip()[:300])
    return "\n".join(json.loads(r.stdout))


def strip_html_comments(src):
    src = re.sub(r"^\s*//.*$", "", src, flags=re.M)
    src = re.sub(r"/\*.*?\*/", "", src, flags=re.S)
    src = re.sub(r"<!--.*?-->", "", src, flags=re.S)
    return src


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", default=DEFAULT_ROOT,
                    help="repository root to check (the negative control passes a copy)")
    a = ap.parse_args()
    root, fails = a.root, []
    panels = os.path.join(root, "engine", "panels")
    records = br.by_key(panels)

    # ---- 1. vocabulary a reader can see ------------------------------------
    for path in surfaces(root):
        rel = os.path.relpath(path, root)
        raw = open(path, encoding="utf-8").read()
        if path.endswith(".js"):
            # DATA files (data.js, coverage.js, markets.js) are walked after
            # loading, so comments and internal fields drop out by construction.
            # BEHAVIOUR files (app.js) need a DOM and cannot be loaded here —
            # their reader-facing strings live inside function bodies, which an
            # object walk would not reach anyway, so those fall back to stripping
            # comments. Both paths are covered; neither guesses.
            try:
                text = js_reader_text(path)
            except RuntimeError:
                text = strip_html_comments(raw)
        else:
            text = strip_html_comments(raw)
        fails += br.scan_text(text, rel)

    # ---- 1b. THE DELIVERED STUDY DOCUMENTS, WHICH THIS GATE NEVER SCANNED ----
    # [added 03-Sep-2026]. This gate has always read the SITE -- html and js -- and
    # a delivered study is as public a surface as a web page. ARCC's shipped study
    # published the retired skill verdict twice and a band flag its own record does
    # not earn, and nothing here could see either, because the study is a .docx and
    # this loop reads text files. Each study's own scrub checks INTERNAL-PROCEDURE
    # vocabulary and knows nothing about verdicts; this one knows about verdicts and
    # never looked at the study. Between them the surface was uncovered.
    # ONLY THE DELIVERED EDITION, resolved by the date in its name -- L-067, the
    # same discipline the workbook-structure and deliverables gates already use. A
    # superseded edition is a historical artefact nobody receives, and failing on
    # one is the permanently-red check [R-ENF-02] forbids. The filename date is
    # DD-MM-YYYY, so it is parsed rather than sorted lexically: a lexical sort puts
    # 08-08 after 03-09 and would pick the wrong file, which is a mistake this
    # session has already made once today.
    def _edition_key(fp):
        m = re.search(r"(\d{2})-(\d{2})-(\d{4})", os.path.basename(fp))
        return (m.group(3), m.group(2), m.group(1)) if m else ("", "", "")

    # RATCHETED per [R-ENF-02]: three current editions predate this extension and are
    # cleared at their own next re-issue. The list may only ever SHORTEN.
    _DOC_RATCHET = set(_ratchet_json(root).get("documents", []))
    _doc_outstanding = 0
    _delivered = []
    for _sdir in sorted(glob.glob(os.path.join(root, "engine", "*_study"))):
        _docs = [f for f in glob.glob(os.path.join(_sdir, "*.docx"))
                 if re.search(r"\d{2}-\d{2}-\d{4}", os.path.basename(f))]
        if not _docs:
            continue
        _latest = max(_edition_key(f) for f in _docs)
        _delivered += [f for f in _docs if _edition_key(f) == _latest]
    for path in sorted(_delivered):
        rel = os.path.relpath(path, root)
        try:
            _d = Document(path)
            _parts = [p_.text for p_ in _d.paragraphs]
            for _t in _d.tables:
                for _r in _t.rows:
                    _parts += [c_.text for c_ in _r.cells]
            _hits = br.scan_text("\n".join(_parts), rel)
            if rel.replace(os.sep, "/") in _DOC_RATCHET:
                _doc_outstanding += len(_hits)          # reported, not failing
            else:
                fails += _hits
        except Exception as e:                       # [R-ENF-04]: an unreadable
            fails.append("%s: could not be read (%s). An unreadable document is not "
                         "a clean one." % (rel, e))


    # ---- 1b. the delivered WORKBOOKS, which no gate had ever read ------------
    # [R-CAL-03] retires the verdict from "no page, figure, document or deck", and
    # this check grew one surface at a time: the pages first, the figures' caption
    # template next, the Word documents on 03-Sep-2026. THE WORKBOOK IS DELIVERED
    # BESIDE THE DOCUMENT and nothing read it, so the verdict survived there while
    # every other surface was swept — including in two workbooks built THIS WEEK,
    # which is what makes this a live hole rather than a legacy one. Same population
    # discipline as the documents: latest edition per study, [R-ENF-04] anchoring, and
    # a ratchet that may only SHORTEN.
    _WB_RATCHET = set(_ratchet_json(root).get("workbooks", []))
    _wb_outstanding = 0
    _books = []
    for _sdir in sorted(glob.glob(os.path.join(root, "engine", "*_study"))):
        _x = [f for f in glob.glob(os.path.join(_sdir, "*.xlsx"))
              if re.search(r"\d{2}", os.path.basename(f))]
        if not _x:
            continue
        _top = max(_workbook_key(f) for f in _x)
        _books += [f for f in _x if _workbook_key(f) == _top]
    if not _books:
        # [R-ENF-04] an empty population is not a clean one: this repository holds
        # delivered workbooks, so reading none means the resolver broke.
        fails.append("no delivered workbook was read at all — the population resolver "
                     "found none, which is a broken run rather than a clean one")
    for path in sorted(_books):
        rel = os.path.relpath(path, root).replace(os.sep, "/")
        try:
            _wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
            _cells = []
            for _ws in _wb.worksheets:
                for _row in _ws.iter_rows():
                    for _c in _row:
                        if isinstance(_c.value, str):
                            _cells.append(_c.value)
            _wb.close()
            _hits = br.scan_text("\n".join(_cells), rel)
            if rel in _WB_RATCHET:
                _wb_outstanding += len(_hits)
            else:
                fails += _hits
        except Exception as e:                       # [R-ENF-04]: unreadable is not clean
            fails.append("%s: could not be read (%s). An unreadable workbook is not "
                         "a clean one." % (rel, e))

    # ---- 2. the published record still matches its panel --------------------
    # Compare the block the generator PRODUCES against the block in the file,
    # rather than re-parsing data.js with a regex that mirrors the emitter's
    # exact formatting: that regex checked three of ten fields and would have
    # degraded to a silent no-op the moment the format moved.
    sys.path.insert(0, os.path.join(root, "scripts"))
    import build_band_records as bbr
    data_js = os.path.join(root, "assets", "data.js")
    src = open(data_js, encoding="utf-8").read()
    try:
        block, _ = bbr.build(src)
    except Exception as e:                                   # noqa: BLE001
        fails.append(f"assets/data.js: band records cannot be rebuilt — {e}")
    else:
        if block not in src:
            fails.append("assets/data.js: the BANDS block is stale or hand-edited — "
                         "re-run scripts/build_band_records.py --write")

    # ---- 3. every data-band-record span names a real record -----------------
    for path in sorted(glob.glob(os.path.join(root, "*.html")) + glob.glob(os.path.join(root, "legacy", "*.html"))):
        rel = os.path.relpath(path, root)
        for m in re.finditer(r'data-band-record="([^"]+)"',
                             open(path, encoding="utf-8").read()):
            try:
                br.resolve(m.group(1), records)
            except KeyError:
                fails.append(f'{rel}: data-band-record="{m.group(1)}" has no panel')

    # ---- 4. the caption behind the figures this gate cannot read ------------
    # The figures are images, so their text is unreadable here. What IS readable
    # is the caption TEMPLATE that produces it, so that is what gets scanned —
    # and ONLY that. Scanning the whole module would flag the CRPS diagnostic
    # [R-CAL-03] deliberately keeps in the codebase, which is not a public
    # surface and never reaches a reader.
    cap = os.path.join(root, FIGURE_CAPTION_SOURCE)
    baked = []
    if os.path.exists(cap):
        src_cap = open(cap, encoding="utf-8").read()
        # the header templates assigned in build(): h2 = (...) / h3 = (...)
        for m in re.finditer(r'^\s{4}(h[23]) = \((.*?)\)\n', src_cap, re.S | re.M):
            baked += br.scan_text(m.group(2),
                                  f"{FIGURE_CAPTION_SOURCE} ({m.group(1)} caption)")
    out_path = os.path.join(root, OUTSTANDING)
    known = json.load(open(out_path)) if os.path.exists(out_path) else {"figures": []}
    if baked and not known.get("figures"):
        fails += baked

    if fails:
        print(f"[R-CAL-02] FAIL — {len(fails)} problem(s):")
        for f in fails[:60]:
            print("  " + f)
        if len(fails) > 60:
            print(f"  ... and {len(fails) - 60} more")
        return 1
    print("[R-CAL-02] OK — no verdict vocabulary in any page text; every published "
          "band record agrees with its panel.")
    if baked:
        n = len(known.get("figures", []))
        if _doc_outstanding:
            print(f"  OUTSTANDING (reported, not failing): {_doc_outstanding} verdict "
                  f"hit(s) in {len(_DOC_RATCHET)} delivered study document(s) that predate "
                  f"this check — see {OUTSTANDING}.")
        if _wb_outstanding:
            print(f"  OUTSTANDING (reported, not failing): {_wb_outstanding} verdict "
                  f"hit(s) in {len(_WB_RATCHET)} delivered WORKBOOK(s) that predate the "
                  f"05-Sep-2026 extension of this check to workbooks — see {OUTSTANDING}.")
        print(f"  read {len(_books)} delivered workbook(s) and {len(_delivered)} delivered "
              f"document(s).")
        print(f"  OUTSTANDING (reported, not failing): the skill verdict is still baked into "
              f"{n} calibration figure(s) this gate cannot read — see {OUTSTANDING}.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
