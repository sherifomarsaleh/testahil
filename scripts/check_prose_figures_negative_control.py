#!/usr/bin/env python3
"""Negative control for scripts/check_prose_figures.py.

Reinjects each condition the gate claims to catch — including the shape that matters most,
a study that CARRIES the instrument while its own check is RED, which a gate testing only
for the file's existence would have reported clean.
"""
import json
import os
import shutil
import subprocess
import sys
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GATE = os.path.join('scripts', 'check_prose_figures.py')

PASSING = ("import sys\n"
           "print('prose figures checked: 10; unmatched: 0')\n"
           "print('reads X.docx via Document()')\n"
           "sys.exit(0)\n")
FAILING = ("import sys\n"
           "print('prose figures checked: 10; unmatched: 3')\n"
           "print('reads X.docx via Document()')\n"
           "sys.exit(1)\n")
NOT_A_CHECK = ("print('a report about figures, reading nothing')\n")
CRASHES = ("import sys\n"
           "raise SystemExit('reads X.docx via Document() then dies: no numbers file')\n")



# ---------------------------------------------------------------------------------------
# THE SIGN-WORD CHECK. A direction word is a claim and is checked against the sign beside
# it: ARCC shipped "+1.8% below the simple annualisation" with the ratio computed correctly
# and the word saying the opposite, and MODON ships "AED 2.50, -12% above the market". Two
# innocent constructions fired against the first draft and are clean cases here — a
# temporal "over" and a range dash — which is why bare "over" is not in the upward set and
# a sign preceded by a digit is not a minus.
_SIGN_FAIL = [
    ("ARCC as it shipped — a computed +1.8% beside a typed 'below'",
     ['This model forecasts EGP 3,842mn for FY2026, +1.8% below the simple annualisation.']),
    ("MODON as it shipped — a negative sign beside 'above'",
     ['Result: AED 2.50, -12% above the market.']),
    ("a minus sign beside 'more than'",
     ['The lens lands -8.4% more than the traded price.']),
    ("a plus sign two words before 'short of'",
     ['It comes in +3.0% well short of the benchmark.']),
]
_SIGN_CLEAN = [
    ("a temporal 'over' — fired against the first draft",
     ["each leg's OWN price fell about -2.4% over the same span"]),
    ("a RANGE dash, not a minus — fired against the first draft",
     ['Portland prices ran 80-85% above the 2024 average during 2025']),
    ("the corrected ARCC sentence",
     ['for FY2026, +1.8% above the simple annualisation']),
    ("a correct downward statement",
     ['the lens sits -13.6% below the latest known price']),
    ("a direction word far from any sign",
     ['margin of 42.9% against 40.6%, well above FY2023']),
]


def sign_word_control():
    """Returns the list of failures; empty means all nine conditions held."""
    import prose_figures as _PF
    bad = []
    for nm, txt in _SIGN_FAIL:
        n = len(_PF.sign_word_conflicts(txt))
        print('  [%s] %s: %d' % ('RED ' if n else 'MISS', nm, n))
        if not n:
            bad.append('SIGN FAIL CASE DID NOT FIRE: ' + nm)
    for nm, txt in _SIGN_CLEAN:
        n = len(_PF.sign_word_conflicts(txt))
        print('  [%s] %s: %d' % ('ok  ' if not n else 'FIRE', nm, n))
        if n:
            bad.append('SIGN CLEAN CASE FIRED: ' + nm)
    return bad


def _plant_docx(path, text):
    """A minimal delivered document, so the sandbox's population is not empty.

    The gate refuses on zero documents per [R-ENF-04], which is correct in production and
    would otherwise turn every case in this sandbox red for the WRONG reason — a sandbox
    that omits something a gate needs makes the gate crash on the absence, which reads
    exactly like going red for the right reason. That is [R-ENF-07]'s own first-run finding
    and it recurred here.
    """
    import docx
    d = docx.Document()
    d.add_paragraph(text)
    d.save(path)


def build(tmp, studies, ratchet=None, docs=True, doc_text='The lens sits -13.6% below.'):
    eng = os.path.join(tmp, 'engine')
    os.makedirs(os.path.join(eng, 'build_depth_audit'), exist_ok=True)
    os.makedirs(os.path.join(tmp, 'scripts'), exist_ok=True)
    shutil.copy(os.path.join(ROOT, GATE), os.path.join(tmp, GATE))
    # THE GATE IMPORTS THE SHARED INSTRUMENT, so a sandbox without it is testing a
    # different program. The sign-word check lives in engine/prose_figures.py and is run
    # book-wide by the gate.
    shutil.copy(os.path.join(ROOT, 'engine', 'prose_figures.py'),
                os.path.join(eng, 'prose_figures.py'))
    for tk, script in studies.items():
        sd = os.path.join(eng, '%s_study' % tk.lower())
        os.makedirs(sd, exist_ok=True)
        if script is not None:
            open(os.path.join(sd, 'prose_check.py'), 'w').write(script)
        if docs:
            _plant_docx(os.path.join(sd, '%s_Valuation_Study_01-01-2026.docx' % tk.upper()),
                        doc_text)
    json.dump({'entries': ratchet or {}},
              open(os.path.join(eng, 'build_depth_audit',
                                'prose_outstanding.json'), 'w'))
    return tmp


def run(tmp):
    r = subprocess.run([sys.executable, GATE], cwd=tmp, capture_output=True, text=True)
    return r.returncode, (r.stdout or '') + (r.stderr or '')


CASES = [
    ('1. no prose check at all', True, {'X': None}, None, 'no script reconciles'),
    ('2. THE ONE THAT MATTERS — carries the instrument and its own check is RED',
     True, {'X': FAILING}, None, 'unmatched: 3'),
    ('3. a file that mentions figures but reads no document', True,
     {'X': NOT_A_CHECK}, None, 'no script reconciles'),
    ('4. a check that crashes', True, {'X': CRASHES}, None, 'X'),
    ('5. a listed study no longer on disk [R-ENF-04]', True, {'X': PASSING},
     {'GONE': 'had none'}, 'no longer resolve on disk'),
    ('6. no study directories at all [R-ENF-04]', True, {}, None,
     'examined zero study directories'),
    ('7. CLEAN — carries it and passes', False, {'X': PASSING}, None, None),
    ('8. CLEAN — a study with no check that is on the ratchet', False, {'X': None},
     {'X': 'no prose check'}, None),
    ('9. CLEAN — a study whose check is RED but is on the ratchet: knowingly '
     'outstanding, allowed to fail', False, {'X': FAILING}, {'X': 'unmatched: 3'}, None),
]


def main():
    bad = 0
    for name, must_fail, studies, ratchet, expect in CASES:
        tmp = tempfile.mkdtemp(prefix='ncprose')
        try:
            build(tmp, studies, ratchet)
            rc, out = run(tmp)
            red = rc != 0
            ok = (red == must_fail) and (expect is None or expect in out)
            print('%-4s %s' % ('PASS' if ok else 'FAIL', name))
            if not ok:
                bad += 1
                print('      rc=%d wanted %s%s' % (
                    rc, 'RED' if must_fail else 'GREEN',
                    (' containing %r' % expect) if expect else ''))
                print('      ' + '\n      '.join(out.strip().splitlines()[-8:]))
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    print('\nSIGN-WORD CHECK:')
    sys.path.insert(0, os.path.join(ROOT, 'engine'))
    sbad = sign_word_control()
    for b in sbad:
        print('  ' + b)
    total = len(CASES) + len(_SIGN_FAIL) + len(_SIGN_CLEAN)
    passed = total - bad - len(sbad)
    print('\n%d/%d conditions behaved as specified' % (passed, total))
    return 1 if (bad or sbad) else 0


if __name__ == '__main__':
    sys.exit(main())
