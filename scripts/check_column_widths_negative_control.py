#!/usr/bin/env python3
"""The column-width gate must fire on every shape it claims to catch, and on nothing else.

Every failing case is a construction that actually shipped. The clean cases include the two
that matter most: a column holding a URL, which cannot fit any page and must NOT fire, and
a column short by less than the experiment can resolve — which MUST fire, because the
tolerance that would have excused it was measured to hide a defect verified on the page.
"""
import json
import os
import shutil
import subprocess
import sys
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GATE = os.path.join('scripts', 'check_column_widths.py')
sys.path.insert(0, os.path.join(ROOT, 'engine'))
import col_width as C                                                  # noqa: E402


def plant(root, study, name, headers, rows, widths, size=8.5):
    import docx
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    d = docx.Document()
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(lay)
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(round(w * 567))))
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.text = ''
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(size)
    for row in rows:
        tr = t.add_row()
        for i, v in enumerate(row):
            cell = tr.cells[i]
            cell.width = Cm(widths[i])
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(v))
            r.font.size = Pt(size)
    sd = os.path.join(root, 'engine', '%s_study' % study)
    os.makedirs(sd, exist_ok=True)
    d.save(os.path.join(sd, name))


def build(tmp, ratchet=None):
    os.makedirs(os.path.join(tmp, 'scripts'), exist_ok=True)
    os.makedirs(os.path.join(tmp, 'engine', 'build_depth_audit'), exist_ok=True)
    shutil.copy(os.path.join(ROOT, GATE), os.path.join(tmp, GATE))
    # the gate imports the real module; the sandbox gets a copy so the constants and the
    # committed measurements travel with it
    shutil.copy(os.path.join(ROOT, 'engine', 'col_width.py'),
                os.path.join(tmp, 'engine', 'col_width.py'))
    shutil.copytree(os.path.join(ROOT, 'engine', 'lab', 'col_width'),
                    os.path.join(tmp, 'engine', 'lab', 'col_width'))
    json.dump({'outstanding': ratchet or {}},
              open(os.path.join(tmp, 'engine', 'build_depth_audit',
                                'colwidth_outstanding.json'), 'w'), indent=1)


def run(tmp):
    r = subprocess.run([sys.executable, GATE], cwd=tmp, capture_output=True, text=True)
    return r.returncode, (r.stdout or '') + (r.stderr or '')


# (name, planter, must_fail, expected fragment)
CASES = [
    ("PHDC's register as it shipped — a 7.5pt ISO date in 1.90cm, which prints "
     "'2025-12-' with a bare '31' beneath",
     lambda t: plant(t, 'zzz', 'ZZZ_Bibliography_03-09-2026.docx',
                     ['Input', 'Value', 'Date'],
                     [['accounts receivable', '28,118.1', '2025-12-31']],
                     [4.0, 2.0, 1.90], size=7.5), True, 'Date'),
    ("TMGH's ladder as it shipped — ten rate cells at 1.28cm printing '32.4' with an "
     "orphaned '%'",
     lambda t: plant(t, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
                     ['Year', '1', '2'], [['Cost of capital', '32.4%', '28.6%']],
                     [3.6, 1.28, 1.28]), True, "docx::1"),
    ("PHDC's appendix as it shipped — a bracketed eight-character negative in 1.66cm",
     lambda t: plant(t, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
                     ['EGP million', '2035'], [['Cost of revenue', '-110,168']],
                     [4.6, 1.66]), True, '2035'),
    ("a BOLD header wider than its own column — the face, not the string",
     lambda t: plant(t, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
                     ['Case', '102,747'], [['Rating-based', '35.79%']],
                     [3.6, 1.70]), True, '102,747'),
    # ---- and the ones that must NOT fire ----------------------------------------------
    ("CLEAN — a URL column, which cannot fit any page and is a WORD, not a figure",
     lambda t: plant(t, 'zzz', 'ZZZ_Bibliography_03-09-2026.docx',
                     ['Source'],
                     [['https://ir.example.com/en/financial-information/'
                       'annual-reports/consolidated-statements-fy2025.pdf']],
                     [6.0]), False, None),
    ("CLEAN — a long prose column that wraps at word boundaries",
     lambda t: plant(t, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
                     ['Why it is a comparator'],
                     [['integrated-community developer, similar recognition basis']],
                     [4.0]), False, None),
    ("CLEAN — every column sized from the cells by fit_widths",
     lambda t: plant(t, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
                     ['EGP million', '2030'], [['Development revenue', '102,747']],
                     C.fit_widths(['EGP million', '2030'],
                                  [['Development revenue', '102,747']], 16.0)),
     False, None),
    ("CLEAN — the same table at 7.5pt, where the ink is smaller and the column fits",
     lambda t: plant(t, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
                     ['EGP million', '2030'], [['Development revenue', '102,747']],
                     [4.0, 1.55], size=7.5), False, None),
]


def main():
    bad = 0
    for name, planter, must_fail, expect in CASES:
        tmp = tempfile.mkdtemp(prefix='nccw')
        try:
            build(tmp)
            planter(tmp)
            rc, out = run(tmp)
            ok = ((rc != 0) == must_fail) and (expect is None or expect in out)
            print('%-4s %s' % ('PASS' if ok else 'FAIL', name[:100]))
            if not ok:
                bad += 1
                print('      rc=%d wanted %s' % (rc, 'RED' if must_fail else 'GREEN'))
                print('      ' + '\n      '.join(out.strip().splitlines()[-4:]))
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    # THE CASE THE TOLERANCE ARGUMENT TURNED ON. A column short by less than the
    # experiment's own 0.05cm step was once excused on the ground that the model cannot
    # resolve it. PHDC's sits 0.039cm under AND ITS PAGE WRAPS, so the excuse hid a
    # verified defect and the tolerance is zero. This case is that measurement.
    tmp = tempfile.mkdtemp(prefix='nccwt')
    try:
        build(tmp)
        plant(tmp, 'zzz', 'ZZZ_Bibliography_03-09-2026.docx',
              ['Input', 'Date'], [['x', '2025-12-31']], [4.0, 1.899], size=7.5)
        rc, out = run(tmp)
        ok = rc != 0
        print('%-4s a column short by LESS than one grid step still FAILS — the tolerance '
              'that would excuse it was measured to hide a defect' % ('PASS' if ok else 'FAIL'))
        bad += 0 if ok else 1
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    # the ratchet excuses, and only what it names
    tmp = tempfile.mkdtemp(prefix='nccwr')
    try:
        build(tmp, {'engine/zzz_study/ZZZ_Bibliography_03-09-2026.docx::Date': 'known'})
        plant(tmp, 'zzz', 'ZZZ_Bibliography_03-09-2026.docx',
              ['Input', 'Date'], [['x', '2025-12-31']], [4.0, 1.80], size=7.5)
        rc, out = run(tmp)
        ok = rc == 0 and 'ratcheted' in out
        print('%-4s a ratcheted narrow column stays GREEN' % ('PASS' if ok else 'FAIL'))
        bad += 0 if ok else 1
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    # an emptied population must FAIL, not report clean
    tmp = tempfile.mkdtemp(prefix='nccwz')
    try:
        build(tmp)
        rc, out = run(tmp)
        ok = rc != 0 and 'zero delivered documents' in out
        print('%-4s a run with no documents FAILS [R-ENF-04]' % ('PASS' if ok else 'FAIL'))
        bad += 0 if ok else 1
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    # a document with no TABLES is the absence the count would hide behind
    tmp = tempfile.mkdtemp(prefix='nccwn')
    try:
        build(tmp)
        import docx
        d = docx.Document()
        d.add_paragraph('a study with no tables at all')
        sd = os.path.join(tmp, 'engine', 'zzz_study')
        os.makedirs(sd, exist_ok=True)
        d.save(os.path.join(sd, 'ZZZ_Valuation_Study_03-09-2026.docx'))
        rc, out = run(tmp)
        ok = rc != 0 and 'zero TABLES' in out
        print('%-4s a run that read documents but zero TABLES FAILS [R-ENF-04]'
              % ('PASS' if ok else 'FAIL'))
        bad += 0 if ok else 1
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    # a ratchet naming a vanished document must FAIL
    tmp = tempfile.mkdtemp(prefix='nccwv')
    try:
        build(tmp, {'engine/gone_study/GONE_Valuation_Study_01-01-2026.docx::Date': 'x'})
        plant(tmp, 'zzz', 'ZZZ_Valuation_Study_03-09-2026.docx',
              ['EGP million', '2030'], [['Development revenue', '102,747']],
              C.fit_widths(['EGP million', '2030'],
                           [['Development revenue', '102,747']], 16.0))
        rc, out = run(tmp)
        ok = rc != 0 and 'no longer exist' in out
        print('%-4s a ratchet naming a vanished document FAILS [R-ENF-04]'
              % ('PASS' if ok else 'FAIL'))
        bad += 0 if ok else 1
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    total = len(CASES) + 5
    print('\n%d/%d conditions behaved as specified' % (total - bad, total))
    return 1 if bad else 0


if __name__ == '__main__':
    sys.exit(main())
