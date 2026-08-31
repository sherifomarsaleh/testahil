#!/usr/bin/env python3
"""Check the technical calibration from OUTSIDE the things it produced.  [R-TCAL-01]

The technical lens now carries the same three artefacts the other lenses carry —
a per-name record (engine/tech_records.json), a generated register of what the
calibration taught (engine/lab/ta_calibration/Technical_Lessons_Register.docx),
and the results files both are built from. Each validates itself while it is
being built, and a self-attested boolean is never a check [R-ENF-01], so this
job runs over the committed artefacts rather than inside their builders and
FAILS rather than warns.

What it checks, and why each one exists:

1. THE MODULES IMPORT. Not parse — import. lessons_source.py resolves every
   lesson figure from the RESULTS files at import, so a lesson whose evidence
   has gone missing dies here instead of printing a stale number.

2. THE RECORD NAMES THE READ IT GRADED. tech_records.json stores the sha256 of
   engine/technicals.py as harvested. If the read has moved since, the record
   is certifying a module that no longer exists — the same defect as the frozen
   technical read the 29-Jul-2026 rule closed, one layer up.

3. THE POPULATION IS ANCHORED OFF THE LIBRARIES ON DISK [R-ENF-04]. Every
   raw_ohlc/{MARKET}/{TICKER}.csv is either recorded or listed as skipped with
   a reason, and nothing is recorded that has no library behind it. Counted
   on (market, ticker), never on ticker alone — ADIB is two banks, and that
   collision has already defeated one count in this repository.

4. THE PAYLOAD IS THE GENERATED FORM. register_payload.json is regenerated and
   compared byte for byte; on a mismatch the committed file is restored before
   failing, so the check never leaves the tree dirty.

5. THE WORD FILE CARRIES EVERY LESSON, AND NO ID THAT RESOLVES NOWHERE. Every
   T-nn in the register document must exist in the source. The orphan half of
   this check caught a real defect the day it was written: the renumbering
   pass of 31-Aug-2026 mapped every id in the sources it knew about and missed
   one in build_register.py, so the delivered document told readers volume was
   "scored in T-013" — an id that no longer existed.
"""
import hashlib
import json
import os
import re
import subprocess
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENGINE = os.path.join(ROOT, "engine")
LAB = os.path.join(ENGINE, "lab", "ta_calibration")
RECORDS = os.path.join(ENGINE, "tech_records.json")
PAYLOAD = os.path.join(LAB, "register_payload.json")
DOCX = os.path.join(LAB, "Technical_Lessons_Register.docx")
RAW = os.path.join(ENGINE, "raw_ohlc")
sys.path.insert(0, ENGINE)
sys.path.insert(0, LAB)


def libraries_on_disk():
    libs = set()
    for market in sorted(os.listdir(RAW)):
        d = os.path.join(RAW, market)
        if not os.path.isdir(d):
            continue
        for f in sorted(os.listdir(d)):
            if f.endswith(".csv"):
                libs.add((market, f[:-4]))
    return libs


def main():
    fails = []

    # 1 — verify by import, not by parse
    import tech_record as TR
    import lessons_source as LS
    print("  [ok]   modules import — %d lessons resolve from the results files"
          % len(LS.LESSONS))

    doc = json.load(open(RECORDS))

    # 2 — the record names the read it graded
    want = hashlib.sha256(
        open(os.path.join(ENGINE, "technicals.py"), "rb").read()).hexdigest()
    got = doc.get("read_sha256")
    if got != want:
        fails.append("tech_records.json grades a technicals.py that is not the "
                     "one on disk (stored %s…, module %s…) — the read moved "
                     "without its record; rerun: python3 engine/tech_record.py"
                     % (str(got)[:12], want[:12]))
    else:
        print("  [ok]   the record grades the technicals.py on disk (%s…)"
              % want[:12])

    # 2b — and it was built on the lens's own clock
    if tuple(doc.get("horizons", ())) != TR.HORIZON_SESSIONS:
        fails.append("tech_records.json horizons %s != the module's %s — the "
                     "record is on the wrong clock"
                     % (doc.get("horizons"), list(TR.HORIZON_SESSIONS)))
    else:
        print("  [ok]   record horizons are the module's own: %s sessions"
              % (list(TR.HORIZON_SESSIONS),))

    # 3 — the population, anchored off the libraries on disk
    libs = libraries_on_disk()
    if not libs:
        fails.append("no OHLC libraries found at all — either the population "
                     "is empty or this check is looking in the wrong place; "
                     "an empty result is not a clean result")
    recorded = {(r["market"], r["ticker"]) for r in doc["records"].values()}
    skipped = {(s["market"], s["ticker"]) for s in doc["skipped"]}
    unexplained = sorted(libs - recorded - skipped)
    phantom = sorted((recorded | skipped) - libs)
    if unexplained:
        fails.append("%d library(ies) neither recorded nor skipped-with-reason:"
                     " %s — rerun: python3 engine/tech_record.py"
                     % (len(unexplained), unexplained[:6]))
    if phantom:
        fails.append("%d record(s) with no library on disk behind them: %s"
                     % (len(phantom), phantom[:6]))
    if doc.get("libraries") != len(libs):
        fails.append("tech_records.json says %s libraries; the disk holds %d"
                     % (doc.get("libraries"), len(libs)))
    if not unexplained and not phantom and doc.get("libraries") == len(libs):
        print("  [ok]   %d libraries on disk — %d recorded, %d skipped with a "
              "reason, every one accounted for"
              % (len(libs), len(recorded), len(skipped)))

    # 3b — no record half-lost. Every name carries one record per horizon, so a
    # name-level count stays whole while a single horizon quietly disappears —
    # the negative control caught this checker doing exactly that on its first
    # run. The stored total and the per-name horizon set close it.
    if doc.get("recorded") != len(doc["records"]):
        fails.append("tech_records.json says %s records but holds %d"
                     % (doc.get("recorded"), len(doc["records"])))
    want_h = set(doc.get("horizons", []))
    partial = sorted(
        (m, t) for (m, t) in recorded
        if {r["h"] for r in doc["records"].values()
            if (r["market"], r["ticker"]) == (m, t)} != want_h)
    if partial:
        fails.append("%d recorded name(s) missing a horizon: %s — a name that "
                     "cannot support a horizon belongs in skipped, with the "
                     "reason" % (len(partial), partial[:6]))
    if doc.get("recorded") == len(doc["records"]) and not partial:
        print("  [ok]   %d records — stored count matches, every recorded "
              "name whole at %d horizons"
              % (len(doc["records"]), len(want_h)))

    # 4 — the payload is the generated form (regenerate, compare, restore)
    committed = open(PAYLOAD, "rb").read()
    r = subprocess.run([sys.executable, os.path.join(LAB, "build_register.py")],
                       capture_output=True, text=True)
    if r.returncode != 0:
        open(PAYLOAD, "wb").write(committed)
        fails.append("build_register.py failed to run: %s"
                     % (r.stderr.strip().splitlines() or ["no output"])[-1])
    else:
        regenerated = open(PAYLOAD, "rb").read()
        if regenerated != committed:
            open(PAYLOAD, "wb").write(committed)
            fails.append("register_payload.json is not the generated form — "
                         "run: python3 engine/lab/ta_calibration/build_register.py")
        else:
            print("  [ok]   register_payload.json matches its generator exactly")

    # 5 — the Word file carries every lesson, and no orphan id
    ids = {json.loads(x)["id"] if isinstance(x, str) else x["id"]
           for x in json.load(open(PAYLOAD))["lessons"]}
    if not os.path.exists(DOCX):
        fails.append("Technical_Lessons_Register.docx does not exist — run: "
                     "node engine/lab/ta_calibration/build_register.js")
    else:
        try:
            from docx import Document
            d = Document(DOCX)
            text = "\n".join(p.text for p in d.paragraphs)
            for t in d.tables:
                for row in t.rows:
                    text += "\n" + "\n".join(c.text for c in row.cells)
            absent = sorted(i for i in ids if i not in text)
            orphan = sorted(set(re.findall(r"\bT-\d{2,3}\b", text)) - ids)
            if absent:
                fails.append("register document is missing %d lesson(s) %s — "
                             "regenerate: node engine/lab/ta_calibration/"
                             "build_register.js" % (len(absent), absent[:6]))
            if orphan:
                fails.append("register document cites id(s) that resolve to no "
                             "lesson: %s" % orphan)
            if not absent and not orphan:
                print("  [ok]   register document carries all %d lessons, no "
                      "orphan ids" % len(ids))
        except ImportError:
            print("  [--]   python-docx not installed; Word file not checked")

    # the read-live surface [R-TCAL-01]: the populations, printed from the
    # committed record rather than from any governing document
    per_h = {}
    for r_ in doc["records"].values():
        per_h[r_["h"]] = per_h.get(r_["h"], 0) + 1
    print("\n  population : %d libraries — %d recorded, %d skipped"
          % (doc.get("libraries", 0), len(recorded), len(skipped)))
    print("  records    : " + " · ".join("h=%d: %d" % (h, n)
                                         for h, n in sorted(per_h.items())))
    print("  register   : %d lessons (%s)" % (
        len(ids), " · ".join("%s %d" % (s, n) for s, n in sorted(
            __import__("collections").Counter(
                x["status"] for x in json.load(open(PAYLOAD))["lessons"]
            ).items()))))

    if fails:
        print("\nFAILED — %d problem(s):" % len(fails))
        for f in fails:
            print("  - %s" % f)
        return 1
    print("\ntechnical calibration OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
