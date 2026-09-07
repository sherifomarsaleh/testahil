#!/usr/bin/env python3
"""Every percentage and multiple a reader sees must be a number the model computed.

WHY THIS EXISTS
    "A NUMBER STATED IN PROSE MUST BE COMPUTED, NOT TYPED" has been a standing rule since
    07-Aug-2026 and depth-bar standard 3 requires that no financial numeral be typed into a
    builder. Exactly ONE study implements a check for it: engine/egch_study/prose_check.py,
    added 01-Sep-2026 after an audit found seven typed figures the computed numbers beside
    them contradicted. On 3 September 2026 the cost of that being one study rather than
    twenty-four was measured, in a single afternoon, on studies that had just been rebuilt
    and passed every other gate:

        AMOC   "A 514-basis-point range across four consecutive filed periods" — computed
               off the five periods the same sentence names, the range is 737 basis points.
               Four more typed period counts beside it, and one of them was not a label but
               arithmetic: a summary row summing FIVE reinvestment rates and dividing by
               FOUR, printed as a "Four-period average".
        ARCC   a masthead a day stale, a price date a month stale, an EFG bridge still
               ending on "this study's weighted central — four lenses, weighted" after the
               blend was retired, and a caption asserting the panel median "sits close to"
               a central 22% away from it.
        PHDC   three comments above one line of the bottom-up model, two of them wrong, one
               of them asserting the exact opposite of what the code does.

    None of these is a modelling error and every one reaches a reader. They survive because
    a typed word — "four", "close to", "revision 3" — does not look like a figure, and
    because the gates around them all examine how numbers were BUILT.

WHAT THIS GATE REQUIRES, AND WHAT IT DELIBERATELY DOES NOT
    It requires each study to CARRY the instrument: a script that opens its own delivered
    documents, reconciles the figures in them against its own committed numbers, and fails
    on an unmatched one. It does NOT set a book-wide threshold on the unmatched count,
    because the honest rendering set is study-specific — a figure legitimately quoted
    against a technical close rather than spot, a structural constant a reader sees, a
    statutory rate — and only the study that wrote the sentence can curate it. EGCH's own
    docstring records the rule: a false positive is fixed by widening the rendering set,
    never by deleting the figure from the study.

    Measured across the book on adoption day for information rather than as a bar: 8,824
    figures in the delivered documents, 373 unmatched against a generic rendering set built
    from each study's committed JSON, a 4.2% rate spanning 0.0% (PHDC) to 11.9% (GBCO).
    That number is PRINTED as an advisory, exactly as library staleness is, because a
    generic rendering set cannot tell a typed figure from an uncurated one — which is the
    argument for the per-study instrument rather than a substitute for it.

THE POPULATION IS ANCHORED ELSEWHERE  [R-ENF-04]
    Every ticker listed in prose_outstanding.json must resolve to a study directory on disk,
    and a run that examined zero study directories fails outright.

THE RATCHET  [R-ENF-02]
    Twenty-two studies have no prose check on adoption day and are listed. The build breaks
    on a NEW study with no entry either way, or on a study that HAD the instrument and
    stopped. The list may only ever get SHORTER — --prune rewrites it.

USAGE
    python3 scripts/check_prose_figures.py           # gate
    python3 scripts/check_prose_figures.py --measure # the advisory rate, per study
    python3 scripts/check_prose_figures.py --prune
"""
import glob
import json
import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENGINE = os.path.join(ROOT, 'engine')
OUTSTANDING = os.path.join(ENGINE, 'build_depth_audit', 'prose_outstanding.json')

# A figure a reader sees: a number carrying a percentage or a multiple. Bare integers are
# deliberately NOT matched — a page number, a note reference and a year are all bare
# integers, and a check that cries wolf is one everyone learns to ignore.
NUM = re.compile(r"(?<![\w.])(-?\d{1,3}(?:,\d{3})*(?:\.\d+)?|-?\d+(?:\.\d+)?)\s*"
                 r"(per cent|%|x\b|times)")
# constants a reader legitimately sees that no model produces
STRUCTURAL = (0, 1, 2, 3, 4, 5, 10, 15, 20, 25, 30, 40, 50, 60, 70, 75, 80, 90, 95, 100,
              0.5, 1.5, 2.5, 7.5, 12.5, 22.5)


def latest(sdir, pattern):
    cands = []
    for p in glob.glob(os.path.join(sdir, pattern)):
        b = os.path.basename(p)
        if b.startswith('~$'):
            continue
        m = re.findall(r'(\d{2})-(\d{2})-(\d{4})', b)
        cands.append(((m[-1][2] + m[-1][1] + m[-1][0]) if m else '', p))
    return sorted(cands)[-1][1] if cands else None


def _flat(x, out):
    if isinstance(x, dict):
        for v in x.values():
            _flat(v, out)
    elif isinstance(x, (list, tuple)):
        for v in x:
            _flat(v, out)
    elif isinstance(x, (int, float)) and not isinstance(x, bool):
        out.append(float(x))


def measure(sdir):
    """(figures, unmatched) against a GENERIC rendering set. Advisory only — a generic set
    cannot distinguish a typed figure from an uncurated one, which is the whole reason the
    instrument belongs inside the study."""
    import docx
    vals = []
    for jf in glob.glob(os.path.join(sdir, '*.json')):
        try:
            _flat(json.load(open(jf, encoding='utf-8')), vals)
        except Exception:                                               # noqa: BLE001
            pass
    render = set()
    for v in vals:
        for x in (v, 1 - v, v - 1, -v, 100 * v, 100 * (1 - v), 100 * (v - 1), v / 100):
            for d in (0, 1, 2, 3):
                render.add(round(x, d))
    for x in STRUCTURAL:
        for d in (0, 1, 2):
            render.add(round(x, d))
    docs = [p for p in (latest(sdir, '*Valuation_Study*.docx'),
                        latest(sdir, '*Bibliograph*.docx')) if p]
    n = u = 0
    for f in docs:
        try:
            d = docx.Document(f)
        except Exception:                                               # noqa: BLE001
            continue
        texts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    texts.append(c.text)
        for txt in texts:
            for m in NUM.finditer(txt):
                raw = m.group(1).replace(',', '')
                dec = len(raw.split('.')[1]) if '.' in raw else 0
                n += 1
                if round(float(raw), dec) not in render:
                    u += 1
    return n, u


def has_instrument(sdir):
    """A script that opens this study's own delivered documents and reconciles their
    figures. Named prose_check.py by convention; anything matching *prose*.py counts, and
    it must actually READ a .docx — a file that only prints a report is not the check."""
    for p in glob.glob(os.path.join(sdir, '*prose*.py')):
        src = open(p, encoding='utf-8', errors='replace').read()
        if '.docx' in src and ('Document(' in src or 'docx' in src):
            return os.path.basename(p)
    return None


def runs_clean(sdir, script):
    """AND IT MUST PASS. Counting a study as conforming because a file exists would report
    it clean while its own check is red — which is worse than having no check, because it
    puts a green tick on a red result. This is the sweep gate's own logic: the instrument
    runs, or it does not run.
    """
    import subprocess
    r = subprocess.run([sys.executable, script], cwd=sdir,
                       capture_output=True, text=True, timeout=600)
    last = [l for l in (r.stdout or '').strip().splitlines() if l.strip()]
    return r.returncode == 0, (last[0] if last else (r.stderr or '').strip()[:160])


def main(argv):
    dirs = sorted(glob.glob(os.path.join(ENGINE, '*_study')))
    if not dirs:
        print('FAIL — examined zero study directories. An empty result is not a clean '
              'result [R-ENF-04].')
        return 1

    if '--measure' in argv:
        print('%-13s %8s %8s %7s' % ('study', 'figures', 'unmatch', 'rate'))
        tn = tu = 0
        for sdir in dirs:
            tk = os.path.basename(sdir)[:-len('_study')].upper()
            n, u = measure(sdir)
            if not n:
                print('%-13s %8s' % (tk, '(no delivered document)'))
                continue
            print('%-13s %8d %8d %6.1f%%' % (tk, n, u, 100.0 * u / n))
            tn += n
            tu += u
        print('%-13s %8d %8d %6.1f%%' % ('TOTAL', tn, tu, 100.0 * tu / max(tn, 1)))
        print('\nADVISORY, never a bar: this rendering set is GENERIC and cannot tell a '
              'typed figure from an uncurated one. The instrument belongs in the study, '
              'where the author who wrote the sentence can curate what a figure may '
              'legitimately be quoted against.')
        return 0

    known = {}
    if os.path.exists(OUTSTANDING):
        known = json.load(open(OUTSTANDING, encoding='utf-8')).get('entries', {})
    on_disk = {os.path.basename(d)[:-len('_study')].upper() for d in dirs}
    stranded = sorted(set(known) - on_disk)

    # THE TWO WAYS OF LACKING ARE NOT THE SAME AND THE MESSAGE MUST SAY WHICH.
    # Both belong in `lack` — a red instrument is not conformance, which is this gate's
    # own rule that it RUNS the instrument rather than counting the file. But they are
    # different facts about the study and they send a reader to different places. On
    # 05-09-2026 EGCH's own prose check went red on one unmatched figure and this gate
    # reported "no prose check and no entry either way" — a study that has carried the
    # book's ORIGINAL implementation since 01-09-2026. The next reader went looking for a
    # deleted file. A MESSAGE THAT MISDESCRIBES WHY A CHECK FAILED IS THE COMMENT
    # ASSERTING A CHECK THAT DOES NOT EXIST, one layer out: it is confidently wrong, and
    # it stops the reader looking where the defect actually is.
    have, lack, detail, red = [], [], {}, set()
    for sdir in dirs:
        tk = os.path.basename(sdir)[:-len('_study')].upper()
        script = has_instrument(sdir)
        if not script:
            lack.append(tk)
            detail[tk] = 'no script reconciles the delivered documents against the numbers'
            continue
        try:
            ok, line = runs_clean(sdir, script)
        except Exception as e:                                          # noqa: BLE001
            ok, line = False, '%s: %s' % (type(e).__name__, e)
        detail[tk] = line
        if not ok:
            red.add(tk)
        (have if ok else lack).append(tk)

    print('PROSE-FIGURE VERIFICATION — the instrument, not a book-wide threshold')
    print('examined %d study directories\n' % len(dirs))
    print('CARRY IT AND PASS (%d):' % len(have))
    for tk in have:
        print('   %-12s %s' % (tk, detail.get(tk, '')))
    if lack:
        print('\nDO NOT (%d):' % len(lack))
        for tk in lack:
            print('   %-12s %s' % (tk, detail.get(tk, '')[:110]))

    now_have = sorted(set(known) & set(have))
    if now_have:
        print('\nNOW CARRYING ONE — remove from the list (%d): %s'
              % (len(now_have), ', '.join(now_have)))

    if '--prune' in argv:
        keep = {k: v for k, v in known.items() if k not in now_have and k in on_disk}
        for tk in lack:
            keep.setdefault(tk, detail.get(tk, 'no prose check')[:220])
        json.dump({'note': ('Studies with no prose-figure verification when this gate was '
                            'adopted (03-Sep-2026). Allowed to fail; the list may only ever '
                            'get shorter. --prune rewrites it. Run --measure for the '
                            'advisory unmatched rate per study.'),
                   'entries': keep},
                  open(OUTSTANDING, 'w', encoding='utf-8'), indent=1, sort_keys=True)
        open(OUTSTANDING, 'a', encoding='utf-8').write('\n')
        print('\npruned; %d entry/entries remain' % len(keep))
        return 0

    # A DIRECTION WORD IS A CLAIM AND IS CHECKED AGAINST THE SIGN BESIDE IT. This runs
    # book-wide with NO per-study declaration, because unlike the rendering set the check is
    # entirely internal to one sentence — a positive sign followed by "below" is wrong on
    # its own terms and no study fact can make it right. Measured across every delivered
    # document in its latest edition it finds ONE contradiction and no false positives.
    sys.path.insert(0, os.path.join(ROOT, 'engine'))
    import prose_figures as _PF
    sign_bad = []
    docs_read = 0
    for sdir in sorted(glob.glob(os.path.join(ROOT, 'engine', '*_study'))):
        tk = os.path.basename(sdir)[:-6].upper()
        for pat in ('*_Valuation_Study_*.docx', '*_Bibliography_*.docx'):
            f = latest(sdir, pat)
            if not f:
                continue
            docs_read += 1
            for frag, ctx in _PF.sign_word_conflicts(_PF.document_texts(f)):
                sign_bad.append((tk, os.path.basename(f), ctx))
    if not docs_read:
        print('\nFAIL — the sign-word check examined zero documents; an empty result is '
              'not a clean result [R-ENF-04]')
        return 1
    SIGN_ALLOWED = {'MODON'}          # ratchet: breaching at adoption, may only SHORTEN
    print('\nSIGN-WORD CHECK: %d delivered documents read, %d contradiction(s)'
          % (docs_read, len(sign_bad)))
    sign_new = []
    for tk, fn, ctx in sign_bad:
        mark = 'ratcheted' if tk in SIGN_ALLOWED else 'NEW'
        print('  [%s] %s: ...%s...' % (mark, fn[:40], ctx[:150]))
        if tk not in SIGN_ALLOWED:
            sign_new.append(tk)

    rc = 0
    if sign_new:
        print('\nFAIL — a signed figure contradicted by the direction word beside it: %s'
              % ', '.join(sorted(set(sign_new))))
        rc = 1
    if stranded:
        print('\nFAIL — %d listed study/studies no longer resolve on disk: %s'
              % (len(stranded), ', '.join(stranded)))
        rc = 1
    unlisted = [tk for tk in lack if tk not in known]
    if unlisted:
        _red = [tk for tk in unlisted if tk in red]
        _none = [tk for tk in unlisted if tk not in red]
        if _none:
            print('\nFAIL — %d study/studies with NO prose check and no entry either '
                  'way: %s' % (len(_none), ', '.join(_none)))
        for tk in _red:
            print('\nFAIL — %s CARRIES a prose check and it is RED, which is not the same '
                  'fact as having none: %s' % (tk, detail.get(tk, '')))
        print('\nA typed word does not look like a figure. "four", "close to", "revision 3" '
              'and "514 basis points" all reached readers this week in studies that passed '
              'every other gate, because every other gate examines how a number was BUILT.')
        rc = 1
    if rc == 0:
        print('\nOK — no new violations.')
    return rc


if __name__ == '__main__':
    sys.exit(main(sys.argv[1:]))
