#!/usr/bin/env python3
"""No delivered document names this repository's own machinery.  [R-ENF-01]

WHY THIS EXISTS
    Depth-bar standard 4 requires a programmatic scan of every delivered document for
    internal-procedure vocabulary, with ZERO hits. Every study implements it, and every
    study implements it as ITS OWN hand-maintained list of forbidden words — 39 terms in
    ARCC, 68 in EGCH, a different set again in AMOC. On 3 September 2026 EGCH's delivered
    bibliography was found shipping three standing-rule identifiers and a repository path

        "[R-MACRO-01]"   "[R-LENS-03]"   "engine/macro_paths/EG.json"

    straight out of an input register's source field, while its own scrub reported 0 hits
    across 68 patterns. AMOC's scrub, which happens to carry both shapes, caught the
    IDENTICAL sentence in its own bibliography the same hour. A sweep of the book then found
    three more delivered documents with the same defect, each through a different hole:
    ADNOCDRILL naming two repository files, PHDC a rule identifier, SCEM an engine module.

    A LIST OF FORBIDDEN WORDS CANNOT BE COMPLETE, and per [R-ENF-01] the fix closes the
    class rather than lengthening five lists. Two things are matched by SHAPE instead, from
    outside the studies, over all of them:

        1. a standing-rule identifier — [R-AREA-NN]
        2. a repository path with a file extension — engine/... or scripts/....py|json|md|csv|js

    Neither can occur innocently in a document written for an outside reader, which is what
    makes shape-matching safe here where a word list is not. This gate does NOT replace the
    per-study scrubs: they catch procedure NOUNS, which need judgement about ordinary senses
    ("register", "gate", "step") and are better decided inside the study that wrote them.

WHAT IT EXAMINES
    Every delivered .docx in every engine/*_study/ — the study, the bibliography or source
    register, and anything else a reader receives — at its LATEST edition, chosen by the
    date in the filename rather than by modification time (L-067: a check that opens a
    delivered file by name moves with the re-issue).

THE POPULATION IS ANCHORED ELSEWHERE  [R-ENF-04]
    Every ticker listed in vocabulary_outstanding.json must resolve to a study directory on
    disk, and a run that examined ZERO documents fails outright. A document that cannot be
    read is a failure, never a skip.

THE RATCHET  [R-ENF-02]
    Documents already leaking on adoption day are listed and allowed to fail. The build
    breaks on a NEW leak. The list may only ever get SHORTER — --prune rewrites it.

USAGE
    python3 scripts/check_delivered_vocabulary.py          # gate; exit 1 on any new leak
    python3 scripts/check_delivered_vocabulary.py --prune  # drop the now-clean entries
"""
import glob
import json
import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENGINE = os.path.join(ROOT, 'engine')
OUTSTANDING = os.path.join(ENGINE, 'build_depth_audit', 'vocabulary_outstanding.json')

# Matched by SHAPE, not by word. A rule identifier and a repository path are unambiguous:
# there is no ordinary sense in which a valuation study written for an investor mentions
# either, which is exactly why these two can be shape-matched where "register" cannot.
SHAPES = [
    ('a standing-rule identifier', re.compile(r'\[R-[A-Z]+-\d+', re.I)),
    ('a repository path',
     re.compile(r'\b(?:engine|scripts|assets)/[A-Za-z0-9_./-]+\.(?:py|json|md|csv|js|yml)\b',
                re.I)),
]


def latest_edition(sdir, pattern):
    """The newest file matching pattern, by the DD-MM-YYYY in its name."""
    cands = []
    for p in glob.glob(os.path.join(sdir, pattern)):
        b = os.path.basename(p)
        if b.startswith('~$'):
            continue
        m = re.findall(r'(\d{2})-(\d{2})-(\d{4})', b)
        cands.append(((m[-1][2] + m[-1][1] + m[-1][0]) if m else '', b, p))
    return sorted(cands)[-1][2] if cands else None


def delivered(sdir):
    """One file per KIND, at its latest edition — the set a reader receives."""
    out = []
    for pat in ('*Valuation_Study*.docx', '*Bibliograph*.docx', '*Source_Register*.docx',
                '*Sources*.docx', '*Source*.docx'):
        p = latest_edition(sdir, pat)
        if p and p not in out:
            out.append(p)
    return out


def text_of(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(c.text)
    return '\n'.join(parts)


def examine(sdir):
    tk = os.path.basename(sdir)[:-len('_study')].upper()
    docs = delivered(sdir)
    if not docs:
        return tk, [], []
    problems, seen = [], []
    for p in docs:
        seen.append(os.path.basename(p))
        try:
            t = text_of(p)
        except Exception as e:                                          # noqa: BLE001
            problems.append((os.path.basename(p), 'unreadable',
                             '%s: %s. An unreadable document is not a clean one.'
                             % (type(e).__name__, e)))
            continue
        for label, rx in SHAPES:
            hits = sorted(set(m.group(0) for m in rx.finditer(t)))
            if hits:
                problems.append((os.path.basename(p), label,
                                 ', '.join(hits[:4])
                                 + (' (+%d more)' % (len(hits) - 4) if len(hits) > 4 else '')))
    return tk, problems, seen


def main(argv):
    dirs = sorted(glob.glob(os.path.join(ENGINE, '*_study')))
    if not dirs:
        print('FAIL — examined zero study directories. An empty result is not a clean '
              'result [R-ENF-04].')
        return 1
    known = {}
    if os.path.exists(OUTSTANDING):
        known = json.load(open(OUTSTANDING, encoding='utf-8')).get('entries', {})
    on_disk = {os.path.basename(d)[:-len('_study')].upper() for d in dirs}
    stranded = sorted(set(known) - on_disk)

    results = [examine(d) for d in dirs]
    n_docs = sum(len(seen) for _, _, seen in results)
    if n_docs == 0:
        print('FAIL — examined zero delivered documents across %d study directories. An '
              'empty result is not a clean result [R-ENF-04].' % len(results))
        return 1

    print('DELIVERED-DOCUMENT VOCABULARY — %d shapes, matched by shape rather than by word'
          % len(SHAPES))
    print('examined %d document(s) across %d study directories\n' % (n_docs, len(results)))

    clean = [tk for tk, pr, seen in results if seen and not pr]
    leaking = [(tk, pr) for tk, pr, seen in results if pr]
    print('CLEAN (%d): %s' % (len(clean), ', '.join(clean) or 'none'))
    if leaking:
        print('\nNAMING THIS REPOSITORY (%d):' % len(leaking))
        for tk, pr in leaking:
            for fn, label, detail in pr:
                print('   %-12s %s' % (tk, fn))
                print('   %-12s   %s: %s' % ('', label, detail))

    now_clean = sorted(set(known) & set(clean))
    if now_clean:
        print('\nNOW CLEAN — remove from the list (%d): %s'
              % (len(now_clean), ', '.join(now_clean)))

    if '--prune' in argv:
        keep = {k: v for k, v in known.items() if k not in now_clean and k in on_disk}
        for tk, pr in leaking:
            keep.setdefault(tk, '; '.join('%s: %s' % (f, d) for f, _, d in pr)[:300])
        json.dump({'note': ('Delivered documents already naming this repository when this '
                            'gate was adopted (03-Sep-2026). Allowed to fail; the list may '
                            'only ever get shorter. --prune rewrites it.'),
                   'entries': keep},
                  open(OUTSTANDING, 'w', encoding='utf-8'), indent=1, sort_keys=True)
        open(OUTSTANDING, 'a', encoding='utf-8').write('\n')
        print('\npruned; %d entry/entries remain' % len(keep))
        return 0

    bad = 0
    if stranded:
        print('\nFAIL — %d listed study/studies no longer resolve on disk: %s'
              % (len(stranded), ', '.join(stranded)))
        bad += 1
    unlisted = [(tk, pr) for tk, pr in leaking if tk not in known]
    if unlisted:
        print('\nFAIL — %d new leak(s):' % len(unlisted))
        for tk, pr in unlisted:
            for fn, label, detail in pr:
                print('   %s / %s — %s: %s' % (tk, fn, label, detail))
        print('\nA rule identifier or a repository path in a document a reader receives is '
              'internal machinery on the page. Rewrite the sentence for the reader; do not '
              'add a word to a study-local list, because the next hole will be a different '
              'shape.')
        bad += 1
    if not bad:
        print('\nOK — no new leaks.')
    return 1 if bad else 0


if __name__ == '__main__':
    sys.exit(main(sys.argv[1:]))
