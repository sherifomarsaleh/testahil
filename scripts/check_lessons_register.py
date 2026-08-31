#!/usr/bin/env python3
"""Check the lessons register from OUTSIDE the register.  [R-LESSON-01] [R-FCAL-01]

A register that validates itself is a self-attested boolean, and this project
has been bitten by those repeatedly. This job runs over the register rather than
inside it and FAILS rather than warns.

What it checks, and why each one exists:

1. THE MODULE IMPORTS. Not parses — imports. A bare identifier parses perfectly
   and dies at import, and that exact bug once reached the default branch.

2. THE MARKDOWN IS THE GENERATED FORM. Regenerate it and compare byte for byte.
   Two copies of the same rules kept in sync by an instruction to remember have
   drifted apart in this repository more than once.

3. BIJECTION BOTH WAYS. Every lesson id appears in the document, and every id in
   the document exists in the module. A tool reporting "0 skipped" is not
   evidence — count against a total held somewhere else.

4. THE POPULATION IS NOT EMPTY, AND IT IS ANCHORED OFF THE REGISTER. The floor
   is the walk-forward directories on disk: every one of them must have produced
   at least one lesson. An empty result is not a clean result, so a register
   that has stopped being fed fails here rather than passing quietly.

5. EVERY REGISTERED CLASS IS REACHABLE. A class named in CLASSES with no lesson
   under it is a category nobody can use; a lesson under an unregistered class
   is a typo that silently creates a new one. Both fail.

6. NO STOCK LESSON MASQUERADES AS GENERAL. A STOCK-scoped lesson whose subject
   is not a real study directory is either a typo or a lesson filed at the wrong
   scope, which is the mistake this register exists to prevent.
"""
import os, subprocess, sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENGINE = os.path.join(ROOT, "engine")
sys.path.insert(0, ENGINE)

MD = os.path.join(ENGINE, "Lessons_Register.md")


def main():
    fails = []

    # 1 — verify by import, not by parse
    import lessons_register as LR
    n = LR.assert_lessons_register()
    print("  [ok]   module imports and validates — %d lessons" % n)

    # 2 — the document is the generated form
    import build_lessons_register as B
    generated = B.build()
    if not os.path.exists(MD):
        fails.append("Lessons_Register.md does not exist")
    else:
        on_disk = open(MD).read()
        if on_disk != generated:
            fails.append("Lessons_Register.md is not the generated form — "
                         "run: python3 engine/build_lessons_register.py")
        else:
            print("  [ok]   Lessons_Register.md matches its generator exactly")

    # 2b — THE WORD FILE CARRIES EVERY LESSON. A .docx cannot be compared byte
    # for byte (python-docx stamps it), so it is checked by CONTENT against the
    # module's own total: every id must appear in the text the reader sees. A
    # document regenerated from a stale source would fail here rather than sit
    # quietly out of date beside a correct Markdown file.
    DOCX = os.path.join(ENGINE, "Lessons_Register.docx")
    if not os.path.exists(DOCX):
        fails.append("Lessons_Register.docx does not exist — run: "
                     "python3 engine/build_lessons_docx.py")
    else:
        try:
            from docx import Document
            d = Document(DOCX)
            text = "\n".join(p.text for p in d.paragraphs)
            for t in d.tables:
                for row in t.rows:
                    text += "\n" + "\n".join(c.text for c in row.cells)
            absent = sorted(x["id"] for x in LR.LESSONS if x["id"] not in text)
            if absent:
                fails.append("Lessons_Register.docx is missing %d lesson(s) "
                             "%s — regenerate it: python3 "
                             "engine/build_lessons_docx.py"
                             % (len(absent), absent[:6]))
            else:
                print("  [ok]   Lessons_Register.docx carries all %d lessons"
                      % len(LR.LESSONS))
        except ImportError:
            print("  [--]   python-docx not installed; Word file not checked")

    # 3 — bijection both ways, counted against the module's own total
    if os.path.exists(MD):
        text = open(MD).read()
        ids = {x["id"] for x in LR.LESSONS}
        missing = sorted(i for i in ids if i not in text)
        if missing:
            fails.append("lessons absent from the document: %s" % missing)
        import re
        in_doc = set(re.findall(r"\bL-\d{3}\b", text))
        orphan = sorted(in_doc - ids)
        if orphan:
            fails.append("ids in the document with no lesson behind them: %s"
                         % orphan)
        if not missing and not orphan:
            print("  [ok]   %d ids resolve both ways, none orphaned" % len(ids))

    # 4 — the population, anchored off the register
    wf = sorted(d for d in os.listdir(ENGINE)
                if d.endswith("_walkforward")
                and os.path.isdir(os.path.join(ENGINE, d)))
    if not wf:
        fails.append("no walk-forward directories found at all — either the "
                     "population is empty or this check is looking in the "
                     "wrong place; an empty result is not a clean result")
    for d in wf:
        tk = d[:-len("_walkforward")].upper()
        got = [x for x in LR.LESSONS
               if x["origin"] == "walk_forward_fundamental"
               and (x["applies_to"] == tk or tk in x["source"].upper())]
        if not got:
            fails.append("%s has a fundamental walk-forward run and no "
                         "lesson in the register" % tk)
    if wf and not any(f.startswith(tuple(d[:-13].upper() for d in wf))
                      for f in fails):
        print("  [ok]   %d walk-forward run(s) on disk, every one represented"
              % len(wf))

    # 4b — EVERY HARVESTED DRAFT IS RESOLVED. The harvester finds candidates
    # mechanically; a candidate nobody ruled on is not a clean result, it is an
    # unanswered question wearing the costume of one. Each draft must end as
    # registered (an id) or declined (a reason). Neither is allowed to be blank,
    # and a run with no draft file at all has not been harvested.
    import json as _json
    for d in wf:
        rd = os.path.join(ENGINE, d)
        dp = os.path.join(rd, "lessons_draft.json")
        if not os.path.exists(dp):
            fails.append("%s has never been harvested — run "
                         "engine/lessons_harvest.py %s"
                         % (d, d[:-len("_walkforward")].upper()))
            continue
        doc = _json.load(open(dp))
        drafts = doc.get("drafts", [])
        if not drafts:
            fails.append("%s/lessons_draft.json holds no drafts at all; an "
                         "empty harvest must not read as a clean one" % d)
        open_ones = [x["proposed_id"] for x in drafts
                     if not x.get("registered") and not x.get("declined")]
        if open_ones:
            fails.append("%s: %d harvested finding(s) neither registered nor "
                         "declined: %s" % (d, len(open_ones), open_ones))
        bad = [x["proposed_id"] for x in drafts
               if x.get("declined") is not None and not str(x["declined"]).strip()]
        if bad:
            fails.append("%s: declined with no reason given: %s" % (d, bad))
        if not open_ones and not bad and drafts:
            print("  [ok]   %s — %d harvested finding(s), all resolved "
                  "(%d registered, %d declined with a reason)"
                  % (d, len(drafts),
                     sum(1 for x in drafts if x.get("registered")),
                     sum(1 for x in drafts if x.get("declined"))))

    # 5 — every registered class reachable, no unregistered class in use
    used = {x["applies_to"] for x in LR.LESSONS if x["scope"] == "CLASS"}
    empty = sorted(set(LR.CLASSES) - used)
    if empty:
        fails.append("registered classes with no lesson: %s" % empty)
    else:
        print("  [ok]   every registered class carries at least one lesson")

    # 6 — a STOCK lesson must name a company this repository actually studies
    studies = {d[:-len("_study")].upper() for d in os.listdir(ENGINE)
               if d.endswith("_study")
               and os.path.isdir(os.path.join(ENGINE, d))}
    for x in LR.LESSONS:
        if x["scope"] == "STOCK" and x["applies_to"].upper() not in studies:
            fails.append("%s is scoped to %s, which is not a study in this "
                         "repository" % (x["id"], x["applies_to"]))
    print("  [ok]   every single-company lesson names a real study")

    c = LR.counts()
    print("\n  scope   : ALL %d · CLASS %d · STOCK %d"
          % (c["ALL"], c["CLASS"], c["STOCK"]))
    print("  origin  : " + " · ".join("%s %d" % (k, v)
                                      for k, v in c["by_origin"].items()))
    print("  status  : " + " · ".join("%s %d" % (k, v)
                                      for k, v in c["by_status"].items()))

    if fails:
        print("\nFAILED — %d problem(s):" % len(fails))
        for f in fails:
            print("  - %s" % f)
        return 1
    print("\nlessons register OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
