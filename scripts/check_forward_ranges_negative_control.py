#!/usr/bin/env python3
"""Negative control for check_forward_ranges.  [R-ENF-01]

Reinjects every condition the gate refuses and every clean case it must not fire on.
EVERY MUTATION ASSERTS THAT IT LANDED before the gate runs, and the case COUNT is
asserted against a declared constant — because three times in one session a control was
caught passing a fixture that never injected its condition.

TWO HALVES, DELIBERATELY. The GATE half sandboxes the repository and plants runs,
documents and ratchet entries. The DETECTOR half feeds engine/range_disclosure's shape
reader the REAL TABLES OUT OF THE DELIVERED DOCUMENTS, pulled at run time rather than
transcribed, in both directions: the three shapes exactly as they ship must be found,
and the same three with their range evidence stripped must not be.

THE CLEAN HALF IS WHERE THIS TURNS. The detector's first draft fired on six studies with
no walk-forward at all, and all six were work that is right — a period written with a
dash, a multiple range in a lens table, a charter-rate spread. Those three kinds are
carried here as cases that must stay silent, because a control proving only that the
gate catches breaches says nothing about the studies it would have wrongly condemned.
"""
import glob
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, ROOT)
TARGET = os.path.join(HERE, "check_forward_ranges.py")
SRC_ENGINE = os.path.join(ROOT, "engine")

from engine import range_disclosure as RD          # noqa: E402

CASES_EXPECTED = 19
RED_EXPECTED = 10
CLEAN_EXPECTED = 9

STUDY_DOC = re.compile(r'valuation[_ ]study.*\.docx$', re.I)
DATE = re.compile(r'(\d{2})-(\d{2})-(\d{4})')


# --------------------------------------------------------------- shared helpers

def _latest(directory):
    docs = [n for n in os.listdir(directory)
            if STUDY_DOC.search(n) and not n.startswith('~')]

    def key(n):
        m = DATE.search(n)
        return (m.group(3), m.group(2), m.group(1)) if m else ('0000', '00', '00')

    return sorted(docs, key=key)[-1] if docs else None


def _grids(path):
    import docx
    return [[[c.text.strip().replace('\n', ' ') for c in r.cells] for r in t.rows]
            for t in docx.Document(path).tables]


def _shipped(ticker):
    """The delivered tables of a run's study, and the year it is dated in."""
    d = os.path.join(SRC_ENGINE, '%s_study' % ticker.lower())
    f = _latest(d)
    assert f, "FIXTURE DID NOT LOAD: no delivered study for %s" % ticker
    return _grids(os.path.join(d, f)), int(DATE.search(f).group(3))


def _points_docx(path, years=(2028, 2029, 2030)):
    """A minimal delivered study that publishes its far years as POINTS."""
    import docx
    doc = docx.Document()
    t = doc.add_table(rows=2, cols=1 + len(years))
    t.cell(0, 0).text = 'EGP mn'
    for j, y in enumerate(years, 1):
        t.cell(0, j).text = 'FY%dE' % y
    t.cell(1, 0).text = 'Revenue'
    for j in range(1, len(years) + 1):
        t.cell(1, j).text = '16,440'
    doc.save(path)


def _sandbox():
    tmp = tempfile.mkdtemp(prefix="fwdrange-nc-")
    eng = os.path.join(tmp, "engine")
    os.makedirs(os.path.join(eng, "build_depth_audit"))
    rat = os.path.join(SRC_ENGINE, "build_depth_audit",
                       "forward_ranges_outstanding.json")
    if os.path.exists(rat):
        shutil.copy(rat, os.path.join(eng, "build_depth_audit"))
    shutil.copy(os.path.join(SRC_ENGINE, "range_disclosure.py"), eng)
    open(os.path.join(eng, "__init__.py"), "w").close()
    for run in sorted(glob.glob(os.path.join(SRC_ENGINE, "*_walkforward"))):
        if not os.path.exists(os.path.join(run, "forward_ranges.json")):
            continue
        tk = os.path.basename(run)[:-len("_walkforward")]
        os.makedirs(os.path.join(eng, os.path.basename(run)))
        open(os.path.join(eng, os.path.basename(run), "forward_ranges.json"),
             "w").write("{}")
        src = os.path.join(SRC_ENGINE, "%s_study" % tk)
        dst = os.path.join(eng, "%s_study" % tk)
        os.makedirs(dst)
        f = _latest(src)
        if f:
            shutil.copy(os.path.join(src, f), os.path.join(dst, f))
    os.makedirs(os.path.join(tmp, "scripts"))
    shutil.copy(TARGET, os.path.join(tmp, "scripts", "check_forward_ranges.py"))
    return tmp


def _run(tmp):
    r = subprocess.run([sys.executable,
                        os.path.join(tmp, "scripts", "check_forward_ranges.py")],
                       capture_output=True, text=True, timeout=600)
    return r.returncode, r.stdout + r.stderr


def _rat_path(tmp):
    return os.path.join(tmp, "engine", "build_depth_audit",
                        "forward_ranges_outstanding.json")


def _unratchet(tmp, tk):
    p = _rat_path(tmp)
    o = json.load(open(p))
    assert tk in o["outstanding"], "MUTATION DID NOT LAND: %s not on the list" % tk
    del o["outstanding"][tk]
    json.dump(o, open(p, "w"))
    assert tk not in json.load(open(p))["outstanding"], "MUTATION DID NOT LAND"


# ------------------------------------------------------------------- red cases

def amoc_unratcheted(tmp):
    _unratchet(tmp, "AMOC")
    return "AMOC"


def arcc_unratcheted(tmp):
    _unratchet(tmp, "ARCC")
    return "ARCC"


def new_run_prints_points(tmp):
    """A run adopted tomorrow whose study publishes its far years as points."""
    eng = os.path.join(tmp, "engine")
    os.makedirs(os.path.join(eng, "newco_walkforward"))
    open(os.path.join(eng, "newco_walkforward", "forward_ranges.json"), "w").write("{}")
    d = os.path.join(eng, "newco_study")
    os.makedirs(d)
    _points_docx(os.path.join(d, "NEWCO_Valuation_Study_07-09-2026.docx"))
    assert os.path.exists(os.path.join(d, "NEWCO_Valuation_Study_07-09-2026.docx")), \
        "MUTATION DID NOT LAND"
    return "NEWCO"


def phantom_ratchet(tmp):
    p = _rat_path(tmp)
    o = json.load(open(p))
    o["outstanding"]["NOSUCHNAME"] = "a name with no run behind it"
    json.dump(o, open(p, "w"))
    assert "NOSUCHNAME" in json.load(open(p))["outstanding"], "MUTATION DID NOT LAND"
    return "anchored on nothing"


def emptied_population(tmp):
    runs = glob.glob(os.path.join(tmp, "engine", "*_walkforward"))
    assert runs, "MUTATION DID NOT LAND"
    for r in runs:
        shutil.rmtree(r)
    return "ZERO walk-forward runs"


def no_documents_read(tmp):
    hit = 0
    for d in glob.glob(os.path.join(tmp, "engine", "*_study")):
        for n in os.listdir(d):
            if STUDY_DOC.search(n):
                os.rename(os.path.join(d, n),
                          os.path.join(d, n.replace("Valuation_Study", "Renamed")))
                hit += 1
    assert hit > 0, "MUTATION DID NOT LAND: no documents renamed"
    return "READ ZERO delivered documents"


def unreadable_document(tmp):
    d = os.path.join(tmp, "engine", "phdc_study")
    f = _latest(d)
    assert f, "MUTATION DID NOT LAND: no PHDC document in the sandbox"
    open(os.path.join(d, f), "wb").write(b"not a zip archive")
    return "unreadable"


def _strip(grids, kill):
    out = []
    for tb in grids:
        out.append([[kill(c) for c in row] for row in tb])
    return out


def detector_shape_a_stripped():
    """PHDC's own table with the Low and High headers renamed. Must find nothing."""
    g, y = _shipped("PHDC")
    assert RD.far_year_range_shapes(g, y) == ["A"], "FIXTURE: PHDC is not shape A"
    m = _strip(g, lambda c: "Band" if c.strip().lower() in ("low", "high") else c)
    assert RD.far_year_range_shapes(m, y) != ["A"], "MUTATION DID NOT LAND"
    return RD.far_year_range_shapes(m, y)


def detector_shape_b_stripped():
    """EGCH's own rows with 'of the range' renamed. Must find nothing."""
    g, y = _shipped("EGCH")
    assert RD.far_year_range_shapes(g, y) == ["B"], "FIXTURE: EGCH is not shape B"
    m = _strip(g, lambda c: c.replace("of the range", "of the forecast"))
    assert RD.far_year_range_shapes(m, y) != ["B"], "MUTATION DID NOT LAND"
    return RD.far_year_range_shapes(m, y)


def detector_shape_c_stripped():
    """TMGH's own cells with the dash removed, leaving the low alone. Must find
    nothing — a single figure is the point this rule forbids."""
    g, y = _shipped("TMGH")
    assert RD.far_year_range_shapes(g, y) == ["C"], "FIXTURE: TMGH is not shape C"
    m = _strip(g, lambda c: re.sub(r"\s*[–—]\s*[0-9][0-9,\.]*$", "", c))
    assert RD.far_year_range_shapes(m, y) != ["C"], "MUTATION DID NOT LAND"
    return RD.far_year_range_shapes(m, y)


# ----------------------------------------------------------------- clean cases

def as_shipped(tmp):
    return "the book as it stands"


def ratchet_shortens(tmp):
    """A ratcheted study starts printing a range. The list may SHORTEN and the
    gate must stay GREEN — a ratchet that went red on its own clearing would be
    the permanently-red check [R-ENF-02] forbids."""
    src = os.path.join(tmp, "engine", "phdc_study")
    dst = os.path.join(tmp, "engine", "arcc_study")
    f = _latest(src)
    old = _latest(dst)
    assert f and old, "MUTATION DID NOT LAND"
    os.remove(os.path.join(dst, old))
    shutil.copy(os.path.join(src, f),
                os.path.join(dst, "ARCC_Valuation_Study_03-09-2026_public.docx"))
    assert _latest(dst) == "ARCC_Valuation_Study_03-09-2026_public.docx", \
        "MUTATION DID NOT LAND"
    return "ARCC now prints a range"


def non_run_study_prints_points(tmp):
    """A study with NO walk-forward publishing far-year points is not in scope and
    must not fire — most of the book is in exactly this state."""
    d = os.path.join(tmp, "engine", "nowf_study")
    os.makedirs(d)
    _points_docx(os.path.join(d, "NOWF_Valuation_Study_07-09-2026.docx"))
    assert os.path.exists(os.path.join(d, "NOWF_Valuation_Study_07-09-2026.docx")), \
        "MUTATION DID NOT LAND"
    return "outside the population"


def detector_shape_a_as_shipped():
    g, y = _shipped("PHDC")
    got = RD.far_year_range_shapes(g, y)
    assert got == ["A"], "PHDC's shipped table no longer reads as a range: %s" % got
    return got


def detector_shape_b_as_shipped():
    g, y = _shipped("EGCH")
    got = RD.far_year_range_shapes(g, y)
    assert got == ["B"], "EGCH's shipped table no longer reads as a range: %s" % got
    return got


def detector_shape_c_as_shipped():
    g, y = _shipped("TMGH")
    got = RD.far_year_range_shapes(g, y)
    assert got == ["C"], "TMGH's shipped table no longer reads as a range: %s" % got
    return got


def detector_period_not_a_range():
    """A far-year row carrying a PERIOD written with a dash. AMR and MODON both
    ship one and the first draft called each a published range."""
    tb = [["What to watch", "When", "Why"],
          ["2028", "2026–2027", "the Kuwait launch"]]
    got = RD.far_year_range_shapes([tb], 2026)
    assert got == [], "a period was read as a range: %s" % got
    return got


def detector_multiple_range_not_a_range():
    """A multiple range in a lens table. AIRARABIA, DU and SWDY all ship one."""
    tb = [["Step", "Value"],
          ["Range at 10× / 16×", "3.10 – 5.14"],
          ["Runs to", "2029"]]
    got = RD.far_year_range_shapes([tb], 2026)
    assert got == [], "a multiple range was read as a far-year range: %s" % got
    return got


def detector_history_range_not_a_range():
    """A range over REPORTED history. The rule is about the far FORECAST years."""
    tb = [["EGP mn", "2023", "2024", "2025"],
          ["Multiple", "8.9 – 15.2", "", ""]]
    got = RD.far_year_range_shapes([tb], 2026)
    assert got == [], "a historical range satisfied a forecast rule: %s" % got
    return got


# ------------------------------------------------------------------------ main

GATE_RED = [
    (amoc_unratcheted, "AMOC unratcheted — points while its own run commits a band"),
    (arcc_unratcheted, "ARCC unratcheted — points, band committed two days earlier"),
    (new_run_prints_points, "a NEW run whose study publishes far-year points"),
    (phantom_ratchet, "the ratchet names a run that does not exist"),
    (emptied_population, "ZERO runs — an empty result is not a clean result"),
    (no_documents_read, "ZERO documents read across present runs"),
    (unreadable_document, "a delivered study that cannot be opened"),
]
DET_RED = [
    (detector_shape_a_stripped, "PHDC's Low/High headers renamed"),
    (detector_shape_b_stripped, "EGCH's 'of the range' rows renamed"),
    (detector_shape_c_stripped, "TMGH's range cells reduced to one figure"),
]
GATE_CLEAN = [
    (as_shipped, "the book exactly as it stands"),
    (ratchet_shortens, "a ratcheted study begins printing a range"),
    (non_run_study_prints_points, "a study with no run publishing points"),
]
DET_CLEAN = [
    (detector_shape_a_as_shipped, "PHDC's table exactly as it ships"),
    (detector_shape_b_as_shipped, "EGCH's table exactly as it ships"),
    (detector_shape_c_as_shipped, "TMGH's table exactly as it ships"),
    (detector_period_not_a_range, "a period written with a dash"),
    (detector_multiple_range_not_a_range, "a multiple range in a lens table"),
    (detector_history_range_not_a_range, "a range over reported history"),
]


def main():
    assert (len(GATE_RED) + len(DET_RED) == RED_EXPECTED
            and len(GATE_CLEAN) + len(DET_CLEAN) == CLEAN_EXPECTED
            and RED_EXPECTED + CLEAN_EXPECTED == CASES_EXPECTED), "the case count moved"
    print("NEGATIVE CONTROL — check_forward_ranges  [R-ENF-01]")
    print("   %d cases: %d must go RED, %d must stay CLEAN"
          % (CASES_EXPECTED, RED_EXPECTED, CLEAN_EXPECTED))
    print("   every mutation asserts that it LANDED before the gate is asked\n")
    bad = 0

    for fn, label in GATE_RED:
        tmp = _sandbox()
        try:
            fn(tmp)
            rc, out = _run(tmp)
            ok = rc != 0
            print("  RED   %-4s %s" % ("ok" if ok else "MISS", label))
            if not ok:
                bad += 1
                print("        the gate reported clean:\n" + out[-400:])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    for fn, label in DET_RED:
        got = fn()
        print("  RED   ok   %-52s -> %s" % (label, got or "no shape"))

    for fn, label in GATE_CLEAN:
        tmp = _sandbox()
        try:
            fn(tmp)
            rc, out = _run(tmp)
            ok = rc == 0
            print("  CLEAN %-4s %s" % ("ok" if ok else "FIRED", label))
            if not ok:
                bad += 1
                print("        the gate went red on work that is right:\n" + out[-500:])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    for fn, label in DET_CLEAN:
        got = fn()
        print("  CLEAN ok   %-52s -> %s" % (label, got or "no shape"))

    if bad:
        print("\nFAIL — %d of %d cases behaved wrongly." % (bad, CASES_EXPECTED))
        return 1
    print("\nOK — all %d cases behaved as the rule requires." % CASES_EXPECTED)
    return 0


if __name__ == "__main__":
    sys.exit(main())
