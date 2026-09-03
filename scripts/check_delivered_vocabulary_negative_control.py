#!/usr/bin/env python3
"""Negative control for scripts/check_delivered_vocabulary.py.

Reinjects every condition the gate claims to catch — including the four leaks that actually
shipped, EGCH's rule identifiers and repository path, ADNOCDRILL's two file paths, PHDC's
identifier and SCEM's engine module — and asserts the gate goes RED on each. Then runs
clean cases that must stay GREEN, because a gate that fires on legitimate prose is one
everyone learns to route around, and two of these are sentences a study should be free to
write.
"""
import json
import os
import shutil
import subprocess
import sys
import tempfile

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GATE = os.path.join('scripts', 'check_delivered_vocabulary.py')


def write_doc(path, paragraphs):
    import docx
    d = docx.Document()
    for t in paragraphs:
        d.add_paragraph(t)
    d.save(path)


def write_doc_table(path, cells):
    """The leaks that shipped were in TABLE cells — an input register's source column."""
    import docx
    d = docx.Document()
    d.add_paragraph('Input register')
    t = d.add_table(rows=len(cells), cols=2)
    for i, (k, v) in enumerate(cells):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
    d.save(path)


CLEAN_PROSE = [
    'Arabian Cement is worth EGP 53.21 a share on the cash-flow lens.',
    'The forecast opens at an EBITDA margin of 39.03% against a filed peak of 39.25%.',
    'Every input carries a value, a source, a date and a research layer in the companion '
    'source register, and the commercial register records the company at Suez.',
    'Section 1.4 sets out the cost of capital. The 2026 step in the tariff is disclosed.',
]


def build(tmp, studies, ratchet=None, table_mode=None):
    eng = os.path.join(tmp, 'engine')
    os.makedirs(os.path.join(eng, 'build_depth_audit'), exist_ok=True)
    os.makedirs(os.path.join(tmp, 'scripts'), exist_ok=True)
    shutil.copy(os.path.join(ROOT, GATE), os.path.join(tmp, GATE))
    for tk, paras in studies.items():
        sd = os.path.join(eng, '%s_study' % tk.lower())
        os.makedirs(sd, exist_ok=True)
        if paras is None:
            continue
        f = os.path.join(sd, '%s_Valuation_Study_03-09-2026.docx' % tk)
        if table_mode:
            write_doc_table(f, [('spot', p) for p in paras])
        else:
            write_doc(f, paras)
    json.dump({'entries': ratchet or {}},
              open(os.path.join(eng, 'build_depth_audit',
                                'vocabulary_outstanding.json'), 'w'))
    return tmp


def run(tmp):
    r = subprocess.run([sys.executable, GATE], cwd=tmp, capture_output=True, text=True)
    return r.returncode, (r.stdout or '') + (r.stderr or '')


CASES = []


def case(name, must_fail, studies, ratchet=None, expect=None, table_mode=False):
    CASES.append((name, must_fail, studies, ratchet, expect, table_mode))


# ---- the four leaks that actually shipped --------------------------------------
case("1. EGCH's bibliography as it shipped — two rule ids and a repository path",
     True, {'EGCH': CLEAN_PROSE + [
         'DERIVED from the house macro path for Egypt (engine/macro_paths/EG.json), and '
         'under [R-MACRO-01] a study may not carry an inflation number of its own.',
         'RETIRED under [R-LENS-03] and consumed by nothing.']},
     expect='[R-MACRO-01')
case("2. ADNOCDRILL's bibliography as it shipped — two file paths",
     True, {'ADNOCDRILL': CLEAN_PROSE + [
         'Regressed by engine/adnocdrill_study/beta_reg.py against '
         'engine/raw_indices/AE/ADXGENERAL.csv.']},
     expect='engine/adnocdrill_study/beta_reg.py')
case("3. PHDC's bibliography as it shipped — one rule id, IN A TABLE CELL",
     True, {'PHDC': ['Egyptian Exchange closing price, 3 September 2026; '
                     '[R-GAP-01 AMENDED] delivers no study against a stale price']},
     expect='[R-GAP-01', table_mode=True)
case("4. SCEM's bibliography as it shipped — an engine module",
     True, {'SCEM': CLEAN_PROSE + ['The fitted parameters live in engine/market_profiles.py.']},
     expect='engine/market_profiles.py')

# ---- other shapes of the same defect ------------------------------------------
case('5. a scripts/ path', True,
     {'X': CLEAN_PROSE + ['Checked by scripts/check_valuation_gap.py before delivery.']},
     expect='scripts/check_valuation_gap.py')
case('6. a lower-case rule id', True,
     {'X': CLEAN_PROSE + ['Adopted under [r-bridge-01] in September.']},
     expect='r-bridge-01')
case('7. an assets/ path', True,
     {'X': CLEAN_PROSE + ['The band record is generated into assets/data.js.']},
     expect='assets/data.js')
case('8. a listed study no longer on disk [R-ENF-04]', True,
     {'X': CLEAN_PROSE}, ratchet={'GONE': 'was leaking'},
     expect='no longer resolve on disk')
case('9. no study directories at all [R-ENF-04]', True, {},
     expect='examined zero study directories')
case('10. study directories but NO delivered documents [R-ENF-04]', True,
     {'X': None}, expect='examined zero delivered documents')

# ---- clean cases that must stay green -----------------------------------------
case('11. CLEAN — ordinary prose using register, gate, step and engine in their '
     'everyday senses', False, {'X': CLEAN_PROSE})
case('12. CLEAN — a known leak that is on the ratchet stays green', False,
     {'X': CLEAN_PROSE + ['See engine/market_profiles.py.']},
     ratchet={'X': 'engine/market_profiles.py'})
case('13. CLEAN — a bare filename with no directory, which a study may legitimately name',
     False, {'X': CLEAN_PROSE + ['The audited statements are filed as FS2025.pdf.']})
case('14. CLEAN — a bracketed reference that is not a rule id', False,
     {'X': CLEAN_PROSE + ['As set out in [Note 14] to the accounts, and in [Table 3].']})


def main():
    bad = 0
    for name, must_fail, studies, ratchet, expect, table_mode in CASES:
        tmp = tempfile.mkdtemp(prefix='ncvocab')
        try:
            build(tmp, studies, ratchet, table_mode)
            rc, out = run(tmp)
            red = rc != 0
            ok = (red == must_fail) and (expect is None or expect.lower() in out.lower())
            print('%-4s %s' % ('PASS' if ok else 'FAIL', name))
            if not ok:
                bad += 1
                print('      rc=%d  wanted %s%s' % (
                    rc, 'RED' if must_fail else 'GREEN',
                    (' containing %r' % expect) if expect else ''))
                print('      ' + '\n      '.join(out.strip().splitlines()[-10:]))
        finally:
            shutil.rmtree(tmp, ignore_errors=True)
    print('\n%d/%d cases behaved as specified' % (len(CASES) - bad, len(CASES)))
    return 1 if bad else 0


if __name__ == '__main__':
    sys.exit(main())
