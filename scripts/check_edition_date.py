#!/usr/bin/env python3
"""[R-ENF-01] A DELIVERED DOCUMENT STATES ITS OWN EDITION DATE, AND STATES THE RIGHT ONE.

WHY THIS EXISTS, AND IT WAS FOUND BY READING RATHER THAN BY ANY GATE. ARCC shipped a
masthead a day stale and it was caught by a person reading the rendered pages. That was
recorded as a defect in one study. Reading PHDC's pages the same way found the identical
thing — a 3 September edition whose masthead reads "edition of 2 September 2026" — which
made it a CLASS rather than an incident, and closing the class rather than the instance is
what [R-ENF-01] requires.

Measured across every delivered valuation study: SEVEN OF THIRTY-ONE do not carry their own
edition date in their masthead, in three distinct shapes.

  STATED AND WRONG        PHDC (3 Sep edition reading 2 September), TMGH (2 Sep reading
                          1 September). Both a day stale, both otherwise clean, and ARCC
                          made three.
  STATED NOWHERE AT ALL   ADNOCLS and SAVOLA. The date in the filename appears in NO
                          paragraph of either document. A reader receives a valuation of a
                          listed company with nothing on it saying when it was struck —
                          and ADNOCLS IS THE MODEL REPORT, the document every other study
                          is written against.
  BURIED IN THE BODY      DU (paragraph 74, inside a sentence about a licence expiry),
                          GBCO (167, in an expert-log note) and RIYADHCABLE (119, in the
                          disclaimer). Present, but not where a reader looks for it, and
                          incidental rather than declared.

WHAT IS CHECKED AND WHAT IS NOT. The masthead must CARRY the filename's date in some
ordinary rendering of it. It is NOT required to be the only date there, and that matters:
AMOC's masthead reads "Valuation study as of 6 August 2026, issued 3 September 2026", which
is the better form — a study can be anchored on one date and issued on another, and saying
both is more honest than saying one. What may not happen is a document whose issue date
disagrees with the file it ships as, or is absent.

Ratcheted [R-ENF-02] — the seven are listed and allowed, the list may only ever SHORTEN.
Population-anchored [R-ENF-04]: a run that examined zero documents FAILS.
"""
import argparse
import datetime
import glob
import json
import os
import re
import subprocess
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RATCHET = os.path.join(ROOT, 'engine', 'build_depth_audit', 'edition_outstanding.json')
DATE = re.compile(r'(\d{2})-(\d{2})-(\d{4})')
# the masthead is the block before the document gets going. Ten paragraphs is generous:
# every study that states its date correctly does so within the first four.
MASTHEAD_PARAS = 10
_MONTHS = ('January', 'February', 'March', 'April', 'May', 'June', 'July',
           'August', 'September', 'October', 'November', 'December')


def renderings(d):
    """Every ordinary way a human writes one date. Shape, not one house format.

    A study may render its date as it likes — the point of this check is that the date is
    THERE and RIGHT, not that every study writes it the same way. Requiring one format
    would fail documents that are perfectly correct, which is the check firing on work that
    is right [R-COC-01].
    """
    out = set()
    for f in ('%d %B %Y', '%d %b %Y', '%B %d, %Y', '%b %d, %Y', '%Y-%m-%d', '%d-%m-%Y',
              '%d/%m/%Y', '%d.%m.%Y'):
        s = d.strftime(f)
        out.add(s)
        out.add(re.sub(r'\b0(\d)', r'\1', s))      # 06 August -> 6 August
    return out


def documents():
    out = []
    for f in sorted(glob.glob(os.path.join(ROOT, 'engine', '*_study', '*.docx'))):
        b = os.path.basename(f)
        if b.startswith('~$') or 'Valuation_Study' not in b:
            continue
        m = DATE.search(b)
        if m:
            out.append((f, datetime.date(int(m.group(3)), int(m.group(2)), int(m.group(1)))))
    return out


def audit():
    from docx import Document
    examined, bad = 0, {}
    for f, dt in documents():
        rel = os.path.relpath(f, ROOT)
        try:
            doc = Document(f)
        except Exception as e:                                    # noqa: BLE001
            bad[rel] = 'will not open: %s' % e
            examined += 1
            continue
        examined += 1
        forms = renderings(dt)
        head = ' | '.join(p.text for p in doc.paragraphs[:MASTHEAD_PARAS])
        if any(x in head for x in forms):
            continue
        whole = ' | '.join(p.text for p in doc.paragraphs)
        where = next((i for i, p in enumerate(doc.paragraphs)
                      if any(x in p.text for x in forms)), None)
        if where is None:
            # is some OTHER date stated in the masthead? that is the stale-masthead shape,
            # and it is worse than an absent date because it reads as a fact.
            # NAME A REAL DATE OR NAME NONE. A loose \w+ for the month matched any word
            # and reported ADNOCLS's masthead as stating '16 on 2026', which is not a date
            # and not what that document says — a diagnostic that misleads is worse than
            # one that admits it found nothing, because the next reader chases it.
            other = re.search(r'\b(\d{1,2}\s+(?:%s)[a-z]*\.?\s+20\d\d)\b'
                              % '|'.join(m[:3] for m in _MONTHS), head)
            bad[rel] = ('the masthead states %r and the file is dated %s'
                        % (other.group(1), dt) if other
                        else 'the edition date %s appears NOWHERE in the document' % dt)
        else:
            bad[rel] = ('the edition date %s appears only at paragraph %d, not in the '
                        'masthead' % (dt, where))
    return examined, bad


# ---------------------------------------------------------------------------------------
# THE SECOND CLAUSE: A DATE TYPED BESIDE A COMPUTED PRICE, INSIDE A FIGURE.
# Reading PHDC's page 2 found the valuation-summary table saying "Market price, 3 September
# 2026" and the figure directly beneath it labelling the SAME 14.40 as "close 14.40 (23 Aug
# 2026)" — one price, two dates, eleven days apart, in one document a reader receives.
# EGCH's price chart said "6 August 2026 close" against a committed spot date of 3
# September, twenty-eight days stale. In both the PRICE was computed and the DATE beside it
# was typed, which is why no gate saw it: everything that reconciles figures against a model
# inspects the number, and a date does not look like one.
#
# SCOPED TO FIGURE BUILDERS ON PURPOSE. A study's PROSE legitimately names many dates — an
# earlier edition, a filing, a licence expiry — and a check over prose would fire on work
# that is right. A figure label is different: it annotates a computed quantity, and a date
# there is describing THAT quantity. BOROUGE types one too and it is CORRECT (7 August 2026
# against a committed 2026-08-07), which is the clean case the control keeps.
_DATE_IN_TEXT = re.compile(
    r'(\d{1,2})\s+(' + '|'.join(m[:3] for m in _MONTHS) + r')[a-z]*\.?\s+(20\d\d)')
_PRICE_WORD = re.compile(r'(?i)\bclose\b|\bspot\b|\bprice\b|\blast\b')


def _committed_spot_date(study_dir):
    f = os.path.join(study_dir, 'study_numbers.json')
    if not os.path.exists(f):
        return None
    try:
        d = json.load(open(f))
    except Exception:                                             # noqa: BLE001
        return None
    raw = str((d.get('meta') or {}).get('spot_date') or d.get('spot_date') or '')
    raw = raw.replace('close ', '').strip()
    for f_ in ('%Y-%m-%d', '%d %b %Y', '%d %B %Y', '%d-%m-%Y'):
        try:
            return datetime.datetime.strptime(raw, f_).date()
        except ValueError:
            pass
    return None


def audit_figure_dates():
    """(examined, offenders) — a date typed in a figure label against the committed spot."""
    examined, bad = 0, {}
    for f in sorted(glob.glob(os.path.join(ROOT, 'engine', '*_study', 'figures.py'))):
        rel = os.path.relpath(f, ROOT)
        spot = _committed_spot_date(os.path.dirname(f))
        try:
            src = open(f, encoding='utf-8').read()
        except Exception:                                         # noqa: BLE001
            continue
        for ln_no, line in enumerate(src.splitlines(), 1):
            if line.lstrip().startswith('#'):
                continue          # a comment recording a fixed defect is not the defect
            m = _DATE_IN_TEXT.search(line)
            if not m or not _PRICE_WORD.search(line):
                continue
            examined += 1
            if spot is None:
                bad['%s:%d' % (rel, ln_no)] = ('a date is typed beside a price and the '
                                               'study commits no spot date to check it '
                                               'against')
                continue
            try:
                got = datetime.datetime.strptime(
                    '%s %s %s' % (m.group(1), m.group(2), m.group(3)), '%d %b %Y').date()
            except ValueError:
                continue
            if got != spot:
                bad['%s:%d' % (rel, ln_no)] = (
                    'a figure labels a price %s while the study commits %s' % (got, spot))
    return examined, bad


# THE MASTHEAD CLAUSE ASKS WHETHER THE EDITION DATE IS PRESENT, AND PRESENCE IS NOT
# AGREEMENT [ADDED 04-Sep-2026]. EGCH's 03-09-2026 edition carried "Valuation study — 1
# September 2026" beside "Anchor price EGP 14.41 at the close of 2026-09-03", and this gate
# passed it: the ISO date is a perfectly good rendering of the edition date, so the gate saw
# the right date sitting next to the wrong one. The study date was TYPED into the builder
# and had gone two days stale.
#
# WHAT IS CHECKED IS AGREEMENT WHERE THE DOCUMENT ITSELF MAKES A CLAIM. Where a date
# immediately follows an edition label, that date IS the document's statement of its own
# edition, and it must be the edition. Where no date follows the label the clause is silent,
# because the document has claimed nothing and the presence clause above already covers it.
_LABEL = re.compile(r'(valuation study|edition of)', re.I)
_ANYDATE = re.compile(
    r'(\d{4}-\d{2}-\d{2})'
    r'|(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|'
    r'November|December)\s+\d{4})'
    r'|((?:January|February|March|April|May|June|July|August|September|October|November|'
    r'December)\s+\d{1,2},\s*\d{4})', re.I)


def _parse(tok):
    for fmt in ('%Y-%m-%d', '%d %B %Y', '%B %d, %Y'):
        try:
            return datetime.datetime.strptime(tok.strip(), fmt).date()
        except ValueError:
            continue
    return None


def audit_masthead_agreement():
    """(examined, offenders) — a date the masthead states as its edition that is not one."""
    examined, bad = 0, {}
    for f, dt in documents():
        rel = os.path.relpath(f, ROOT)
        try:
            import docx
            head = ' '.join(p.text for p in docx.Document(f).paragraphs[:6])
        except Exception:                                             # noqa: BLE001
            continue
        m = _LABEL.search(head)
        if not m:
            continue
        seg = head[m.end():m.end() + 220]
        # AMOC'S FORM IS BETTER THAN THE RULE REQUIRES AND MUST NOT BE PUNISHED FOR IT.
        # "valuation study as of 6 August 2026, issued 3 September 2026" states two dates
        # on purpose — the price it was struck against and the day it was issued — and the
        # EDITION is the second. Where the label is qualified by "as of", the date that
        # must agree is the one after "issued".
        if re.match(r'\s*as of', seg, re.I):
            i = seg.lower().find('issued')
            if i < 0:
                bad[rel] = ('the masthead says "as of" a date and never says when it was '
                            'issued, so it states no edition date at all')
                examined += 1
                continue
            seg = seg[i:]
        d = _ANYDATE.search(seg)
        if not d:
            continue
        examined += 1
        got = _parse(d.group(0))
        if got and got != dt:
            bad[rel] = ('the masthead states this edition is of %s and the edition is %s; '
                        'presence of the right date elsewhere in the masthead is not '
                        'agreement' % (got, dt))
    return examined, bad


def audit_delivered_pdfs():
    """The same date rule, applied to THE FILE A READER ACTUALLY RECEIVES.

    THE PDF IS THE DELIVERABLE AND THE WORD FILE IS THE BUILD ARTEFACT — check_delivered_
    pdf_currency opens with that sentence, and it is why this clause exists. The masthead
    clause above reads .docx, so a study whose document is rebuilt and whose PDF is not
    passes with the old date still on the page a reader opens. That is not hypothetical:
    on 03-Sep-2026 TMGH's masthead was corrected in the .docx and its delivered PDF sat
    NINETEEN HOURS behind it, showing the date that had just been fixed.

    IT FINDS NO OFFENDER TODAY BEYOND THE FOUR ALREADY RATCHETED, and that is the point of
    running it: the four .docx offenders are the same four PDFs, so the two artefacts agree,
    which is what one wants to be able to SAY rather than assume.

    WHITESPACE IS NORMALISED BEFORE THE SEARCH AND THE FIRST DRAFT DID NOT DO IT. A PDF
    wraps text wherever the column ends, so AMOC's "issued 3 September 2026" renders with
    the "3" ending one line and "September 2026" beginning the next — and a substring search
    reported four perfectly correct documents as undated. Two pages are read rather than
    one, because a study with a COVER PAGE carries its masthead on page 2 and reading one
    page called those undated too. Both were the instrument being wrong about the object,
    caught by looking at what it flagged.
    """
    examined, bad = 0, {}
    for f, dt in documents():
        pdf = os.path.splitext(f)[0] + '.pdf'
        if not os.path.exists(pdf):
            continue
        rel = os.path.relpath(pdf, ROOT)
        try:
            t = subprocess.run(['pdftotext', '-f', '1', '-l', '2', '-layout', pdf, '-'],
                               capture_output=True, text=True, timeout=180).stdout
        except Exception:                                         # noqa: BLE001
            continue
        examined += 1
        flat = re.sub(r'\s+', ' ', t)
        if not any(x in flat for x in renderings(dt)):
            bad[rel] = ('the delivered PDF does not carry its own edition date %s in its '
                        'first two pages' % dt)
    return examined, bad


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--prune', action='store_true')
    a = ap.parse_args()
    rat = json.load(open(RATCHET)) if os.path.exists(RATCHET) else {}
    allowed = set(rat.get('outstanding', {}))

    examined, bad = audit()
    fx_examined, fx_bad = audit_figure_dates()
    bad.update(fx_bad)
    ma_examined, ma_bad = audit_masthead_agreement()
    bad.update({k: v for k, v in ma_bad.items() if k not in bad})
    pdf_examined, pdf_bad = audit_delivered_pdfs()
    # a study already ratcheted on its .docx is not counted twice for the same date
    _docx_bad = {os.path.splitext(k)[0] for k in bad}
    bad.update({k: v for k, v in pdf_bad.items()
                if os.path.splitext(k)[0] not in _docx_bad})
    if not examined:
        print('FAIL — examined zero delivered studies; an empty result is not a clean '
              'result [R-ENF-04]')
        return 1
    # THE PDF POPULATION IS ANCHORED SEPARATELY AND FOR A REASON THAT IS NOT PEDANTRY.
    # audit_delivered_pdfs() swallows an extraction failure per file, so a missing
    # pdftotext binary makes EVERY pdf skip and the clause reports nothing — which reads
    # exactly like every PDF carrying its date. That is [R-ENF-04]'s own failure: an
    # ABSENT answer wearing the costume of a clean one. The book delivers PDFs, so zero
    # of them is a broken probe, never a clean result.
    if not pdf_examined:
        print('FAIL — examined zero delivered PDFs while %d .docx editions exist; the '
              'extractor did not run [R-ENF-04]' % examined)
        return 1
    stranded = sorted(p for p in allowed
                      if not os.path.exists(os.path.join(ROOT, p)))
    if stranded:
        print('FAIL — the ratchet names documents that no longer exist: %s [R-ENF-04]'
              % stranded)
        return 1

    print('delivered studies examined: %d (.docx) + %d (.pdf);  mastheads stating their '
          'own edition: %d;  figure labels dating a price: %d;  not carrying the right '
          'date: %d' % (examined, pdf_examined, ma_examined, fx_examined, len(bad)))
    for p in sorted(bad):
        print('  [%s] %-52s %s' % ('ratcheted' if p in allowed else 'NEW',
                                   os.path.basename(p), bad[p]))
    fixed = sorted(allowed - set(bad))
    if fixed:
        print('\nNOW CARRYING IT — remove from the list (%d): %s'
              % (len(fixed), ', '.join(os.path.basename(x) for x in fixed)))

    if a.prune:
        grown = sorted(set(bad) - allowed)
        if grown:
            print('\nREFUSING TO PRUNE — the list would GROW by %s. A ratchet may only '
                  'ever SHORTEN [R-ENF-02].' % grown)
            return 1
        rat['outstanding'] = {k: bad[k] for k in sorted(set(bad) & allowed)}
        json.dump(rat, open(RATCHET, 'w'), indent=1)
        print('\npruned: %d -> %d' % (len(allowed), len(rat['outstanding'])))
        return 0

    new = sorted(set(bad) - allowed)
    if new:
        print('\nFAIL — a delivered document must state its own edition date in its '
              'masthead, and state the right one:\n  %s' % '\n  '.join(new))
        return 1
    print('\nOK — no new violations. %d document(s) on the ratchet, which may only SHORTEN.'
          % len(allowed))
    return 0


if __name__ == '__main__':
    sys.exit(main())
