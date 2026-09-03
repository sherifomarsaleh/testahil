#!/usr/bin/env python3
"""[R-ENF-01] A COLUMN A READER SEES IS WIDE ENOUGH FOR THE FIGURES IT PRINTS.

A table cell one character too narrow wraps, and Word breaks a line after a hyphen — so a
negative number renders as a bare dash with its digits on the line beneath, a date as
"2025-12-" with a bare "31" under it, and a rate as "32.4" with an orphaned "%". THE SIGN,
THE DAY OR THE UNIT OF A PRINTED FIGURE CHANGES AS IT IS READ, FROM A PURELY TYPOGRAPHIC
CAUSE, and every other instrument in this repository examines how a number was BUILT.

WHY THE MODEL AND NOT THE PAGE. A page-side scan was tried first and is the obvious
instrument: render the PDF, find lines that are nothing but a short fragment continuing a
number from the line above. It reports 2,946 hits across the book and almost all of them
are PROSE — a paragraph ending a line on "31-Dec-" and beginning the next with "2024" is
ordinary typesetting, not a defect. pdftotext cannot tell a table cell from a sentence, so
the detector cannot either. The defect is a property of a COLUMN, and only the document
knows where its columns are.

engine/col_width.py holds per-character widths MEASURED in the delivered font
(engine/lab/col_width/*.py build a document of single-token cells across a 0.05cm grid,
render it, and read back the width at which each token stops splitting). Three sets are
committed — plain at 8.5pt, bold at 8.5pt, plain at 7.5pt — and the module asserts its
constants still clear all three at import, so a constant that stopped reproducing its own
experiment fails the import rather than rounding a column down.

WHAT IS DELIBERATELY NOT CHECKED: a WORD breaking across two lines. "Where it was read"
columns hold URLs needing 26.89cm on a 16.79cm page — that token cannot fit any column, and
demanding it would be a false claim about what a table can do, which is the permanently-red
check [R-ENF-02] forbids. A broken word is a nuisance a reader reassembles; a broken figure
is read wrong. The line is drawn where the consequence changes.

Ratcheted [R-ENF-02]: what was already narrow when this was adopted is listed and allowed,
the build breaks on a NEW one, and the list may only ever SHORTEN. Population-anchored
[R-ENF-04] both ways: a run that examined zero DOCUMENTS fails, and so does one that
examined zero TABLES across present documents — the distinction an absent answer hides
behind.
"""
import argparse
import collections
import glob
import json
import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, 'engine'))
import col_width as C                                                  # noqa: E402

RATCHET = os.path.join(ROOT, 'engine', 'build_depth_audit', 'colwidth_outstanding.json')
DATE = re.compile(r'(\d{2})-(\d{2})-(\d{4})')


def latest(paths):
    """The LATEST edition of each document by DATE, never by string sort."""
    keep = {}
    for p in paths:
        b = os.path.basename(p)
        if b.startswith('~$'):
            continue
        m = DATE.search(b)
        key = (os.path.dirname(p), DATE.sub('', b))
        k = (m.group(3), m.group(2), m.group(1)) if m else ('0', '0', '0')
        if key not in keep or k > keep[key][0]:
            keep[key] = (k, p)
    return sorted(v[1] for v in keep.values())


def table_size(t):
    """The point size the cells are ACTUALLY set at, read off the runs.

    The size is not a detail: the ink scales with it, and an input register set at 7.5pt
    judged by 8.5pt figures is judged a tenth of a centimetre wrong on a ten-character
    date — which is exactly the margin these columns are decided by.
    """
    sizes = []
    for r in t.rows[:3]:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    if run.font.size is not None:
                        sizes.append(run.font.size.pt)
    return collections.Counter(sizes).most_common(1)[0][0] if sizes else C.BASE_PT


def audit():
    """(documents, tables, offenders) over every delivered document's latest edition."""
    import docx
    from docx.oxml.ns import qn
    docs = tables = 0
    bad = {}
    for f in latest(glob.glob(os.path.join(ROOT, 'engine', '*_study', '*.docx'))):
        rel = os.path.relpath(f, ROOT)
        try:
            d = docx.Document(f)
        except Exception as e:                                          # noqa: BLE001
            bad[rel] = 'will not open: %s' % e
            continue
        docs += 1
        for ti, t in enumerate(d.tables):
            g = t._tbl.find(qn('w:tblGrid'))
            if g is None:
                continue
            w = [int(gc.get(qn('w:w')) or 0) / 567.0
                 for gc in g.findall(qn('w:gridCol'))]
            rows = [[c.text for c in r.cells] for r in t.rows]
            if not rows or not w:
                continue
            tables += 1
            for col, declared, needed in C.audit(rows[0], rows[1:], w,
                                                 size=table_size(t)):
                key = '%s::%s' % (rel, str(col).strip()[:40] or '(unnamed)')
                bad.setdefault(key, 'declared %.2fcm, needs %.2fcm for its widest '
                                    'figure' % (declared, needed))
    return docs, tables, bad


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--prune', action='store_true')
    a = ap.parse_args()
    rat = json.load(open(RATCHET)) if os.path.exists(RATCHET) else {}
    allowed = set(rat.get('outstanding', {}))

    docs, tables, bad = audit()
    if not docs:
        print('FAIL — examined zero delivered documents; an empty result is not a clean '
              'result [R-ENF-04]')
        return 1
    if not tables:
        print('FAIL — read %d document(s) and zero TABLES; the reader did not run '
              '[R-ENF-04]' % docs)
        return 1
    stranded = sorted(p for p in allowed
                      if not os.path.exists(os.path.join(ROOT, p.split('::')[0])))
    if stranded:
        print('FAIL — the ratchet names documents that no longer exist: %s [R-ENF-04]'
              % stranded[:5])
        return 1

    print('documents examined: %d;  tables: %d;  columns too narrow for a figure they '
          'print: %d' % (docs, tables, len(bad)))
    for k in sorted(bad):
        print('  [%s] %-64s %s' % ('ratcheted' if k in allowed else 'NEW', k[-64:], bad[k]))
    fixed = sorted(allowed - set(bad))
    if fixed:
        print('\nNOW WIDE ENOUGH — remove from the list (%d): %s'
              % (len(fixed), ', '.join(x.split('::')[-1] for x in fixed[:8])))

    if a.prune:
        grown = sorted(set(bad) - allowed)
        if grown:
            print('\nREFUSING TO PRUNE — the list would GROW by %d. A ratchet may only '
                  'ever SHORTEN [R-ENF-02]: %s' % (len(grown), grown[:3]))
            return 1
        rat['outstanding'] = {k: bad[k] for k in sorted(set(bad) & allowed)}
        json.dump(rat, open(RATCHET, 'w'), indent=1)
        print('\npruned: %d -> %d' % (len(allowed), len(rat['outstanding'])))
        return 0

    new = sorted(set(bad) - allowed)
    if new:
        print('\nFAIL — %d column(s) too narrow for a figure they print:\n%s\n'
              'Size the table with engine/col_width.fit_widths() rather than by eye; a '
              'wrapped figure changes what a reader reads.' % (len(new), '\n'.join(new)))
        return 1
    print('\nOK — no new violations. %d column(s) on the ratchet, which may only SHORTEN.'
          % len(allowed))
    return 0


if __name__ == '__main__':
    sys.exit(main())
