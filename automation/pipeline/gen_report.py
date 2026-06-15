"""Build the canonical Word study from one config file.
READ FIRST -> Headline -> 1 Fundamental -> 2 Technical -> 3 Monte Carlo ->
4 Comparison -> 5 Catalysts -> 6 Probability zones -> 7 Caveats + RSI -> Disclosure.
Usage: python pipeline/gen_report.py data/TMGH.yaml output/TMGH_report.docx"""
import sys, datetime, yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import montecarlo

ROOT = Path(__file__).resolve().parent.parent
TEAL = RGBColor(0x1B, 0x5E, 0x5E)
GREY = RGBColor(0x66, 0x66, 0x66)


def h1(doc, txt):
    p = doc.add_heading(level=1)
    r = p.add_run(txt); r.font.name = "Arial"; r.font.color.rgb = TEAL; r.font.size = Pt(15)


def para(doc, txt, italic=False, color=None, size=10, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.font.name = "Arial"; r.font.size = Pt(size)
    r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i].paragraphs[0].add_run(h)
        c.font.bold = True; c.font.name = "Arial"; c.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            r = cells[i].paragraphs[0].add_run(str(val)); r.font.name = "Arial"; r.font.size = Pt(9)
    return t


def main(cfg_path, out_path):
    cfg = yaml.safe_load(open(cfg_path))
    mc = montecarlo.run(cfg)
    v, m, s = cfg["valuation"], cfg["market"], cfg["site"]
    pg = cfg["page"]; tech = s["tech"]; ccy = cfg["currency"]; anchor = m["anchor"]
    base = round(sum(c["value"] for c in v["components"]), 2)
    today = datetime.date.today()
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)

    # title
    p = doc.add_paragraph(); r = p.add_run(f"{cfg['name']} ({cfg['exchange']}:{cfg['ticker']}) — Independent Valuation Study")
    r.bold = True; r.font.size = Pt(16); r.font.name = "Arial"; r.font.color.rgb = TEAL
    para(doc, "Educational Analysis · Fundamental · Technical · Monte Carlo — one integrated read", italic=True, color=GREY)
    para(doc, f"Anchor: {ccy} {anchor} ({m['close_date']} close) · {m['shares_bn']}bn shares · prepared {today.isoformat()}", size=9, color=GREY)

    para(doc, "READ FIRST — what this document is, and is not.", bold=True)
    para(doc, "This study is a valuation exercise and an expression of personal analytical opinion, published free of charge for "
              "educational purposes. It is NOT investment advice, NOT a recommendation or solicitation to buy, sell or hold any "
              "security, and NOT directed at the circumstances of any reader. The preparer is not licensed by the Egyptian Financial "
              "Regulatory Authority (FRA) or any other regulator, manages no money, and accepts no fees, funds or clients. All values "
              "are model outputs presented as ranges and distributions because no single number should be relied on. Consult a "
              "licensed financial advisor before any investment decision.")

    h1(doc, "Headline")
    para(doc, f"Independent per-segment sum-of-the-parts puts risk-adjusted base value at {ccy} {base} "
              f"({(base/anchor-1)*100:+.0f}% vs the {ccy} {anchor} market), with a bear floor of {ccy} {v['bear']} and full-execution "
              f"ceiling of {ccy} {v['full_execution']}. The Monte Carlo T+20 median is {mc['percentiles'][20][50]} and the T+60 median "
              f"{mc['percentiles'][60][50]}, right-skewed on launch and cross-border-inflow events. This study estimates values and "
              f"probabilities — it does not recommend actions.")

    h1(doc, "1.  Fundamental analysis")
    para(doc, "Holding-company structure: each operating segment is valued on its own basis and summed, net cash is added at the "
              "parent-attributable share, and minority interests are deducted segment-by-segment. Consolidated DCF is a labelled "
              "cross-check only, never the primary number.")
    rows = [[c["label"], f"{c['value']:+.2f}" if c["value"] < 0 else f"{c['value']:.2f}"] for c in v["components"]]
    rows.append(["Base-case equity value", f"{base:.2f}"])
    table(doc, ["Component (risk-adjusted)", f"{ccy}/sh"], rows)
    para(doc, f"Scenarios: Bear {ccy} {v['bear']} · Base {ccy} {base} ({(base/anchor-1)*100:+.0f}%) · "
              f"Full execution {ccy} {v['full_execution']} ({(v['full_execution']/anchor-1)*100:+.0f}%).")

    h1(doc, "2.  Technical analysis")
    para(doc, tech["trend"], bold=True)
    para(doc, tech["summary"])
    table(doc, ["Resistance", "Support"], [[pg["resDetail"], pg["supDetail"]]])
    para(doc, f"Typical daily move (ATR): {pg['atr']}. Where the case breaks: {pg['breakBelow']}.", size=9, color=GREY)
    para(doc, f"Bull trigger: {tech['bull']}  Bear trigger: {tech['bear']}", size=9)

    h1(doc, "3.  Monte Carlo simulation — one model, ten external factors")
    para(doc, f"{cfg['montecarlo']['paths']:,} daily paths to T+60, seed {cfg['montecarlo']['seed']}, driven by five continuous factors "
              "(compounding daily) and five discrete event factors (each firing with a stated probability on a random session). Every "
              "drift, probability and impact is an explicit, editable judgment — the table IS the model.")
    frows = [[f["name"], "continuous", f"{f['drift']*100:+.1f}%", f"{f['vol']*100:.0f}%"] for f in cfg["montecarlo"]["continuous"]]
    frows += [[e["name"], "event", f"p={e['prob']*100:.0f}%", f"{e['impact']*100:+.0f}% (±{e['spread']*100:.0f}%)"] for e in cfg["montecarlo"]["events"]]
    table(doc, ["Factor", "Type", "Drift/Prob", "Vol/Impact"], frows)
    para(doc, "Percentile estimates", bold=True)
    pub = cfg["published_dist"]
    table(doc, ["Horizon", "5%", "25%", "Median", "75%", "95%"],
          [["T+20 (~1 month)", pub["t20"]["p5"], pub["t20"]["p25"], pub["t20"]["p50"], pub["t20"]["p75"], pub["t20"]["p95"]],
           ["T+60 (~3 months)", pub["t60"]["p5"], pub["t60"]["p25"], pub["t60"]["p50"], pub["t60"]["p75"], pub["t60"]["p95"]]])
    para(doc, "Level-touch probabilities (T+20 / T+60)", bold=True)
    table(doc, ["Level", "T+20", "T+60"], [[f"{l}", f"{a}%", f"{b}%"] for l, a, b in cfg["published_touch"]])

    h1(doc, "4.  Comparison — fundamental vs technical vs Monte Carlo")
    table(doc, ["Lens", "Core output", "What it says"],
          [["Fundamental", f"Base {base}; full {v['full_execution']}; bear {v['bear']}",
            f"Risk-adjusted value {(base/anchor-1)*100:+.0f}% vs market; upside is segment + launch execution"],
           ["Technical", tech["trend"], "Read the trigger levels above; volume is the confirmation tell"],
           ["Monte Carlo", f"T+20 median {mc['percentiles'][20][50]}; T+60 median {mc['percentiles'][60][50]}",
            "Three-month distribution skews right on the launch and inflow events"]])
    para(doc, "Verdict: the three lenses operate on different clocks — the fundamental upside and any technical caution are not in "
              "conflict. What, if anything, follows for any individual depends on circumstances this study does not know and does not address.")

    h1(doc, "5.  Catalyst calendar")
    table(doc, ["Catalyst", "Read"],
          [["Quarterly results (backlog → revenue)", "Revenue-visibility story; a T+60 item, not near-term"],
           ["Project launch & sales execution", "The pivotal event variable — the largest single factor in the simulation"],
           ["CBE MPC policy-rate path", "Factor 1 in the simulation; affects discounting and mortgage cost"]])

    h1(doc, "6.  Reading the probabilities — significant zones (analysis, not guidance)")
    para(doc, f"This translates the simulation into price zones. It is a map of the distribution — not a plan, not guidance, and not "
              f"calibrated to any reader. Resistance sits at {pg['resDetail']}; support at {pg['supDetail']}. Below {pg['breakBelow']} "
              f"the base framing would itself need rebuilding.")

    h1(doc, "7.  Caveats")
    para(doc, "Bottom-up valuation is assumption-driven (each segment basis, the net-cash parent share, the minority deduction and "
              "the normalized terminal value) — illustrative ranges, not precision. Monte Carlo factors are independent (no "
              "cross-correlation matrix yet); each drift, probability and impact is judgment, not estimation; event timing within the "
              "window is uniform-random.")
    para(doc, "Momentum-reading note (educational). Worry when price keeps rising but RSI stalls or falls under its moving average "
              "(bearish divergence — momentum dying); be comfortable holding when price and RSI rise together. Get interested when "
              "price falls but RSI rises (selling exhausting itself); stay away when price and RSI fall together (intact downtrend). "
              "These are descriptive patterns for reading momentum, not signals to act on.")

    para(doc, "Disclosure & Disclaimer — read in full.", bold=True)
    para(doc, "This study is an educational valuation exercise and an expression of the preparer's personal analytical opinion, based "
              "exclusively on publicly available information and on assumptions stated in the text and the accompanying model. It is "
              "published free of charge, on a standing periodic schedule, to demonstrate methodology and to invite scrutiny of that "
              "methodology. Nothing herein constitutes investment advice, financial consultancy, securities analysis services, a research "
              "recommendation, a rating, a price target, an offer, or a solicitation to deal in any security. The preparer is not licensed "
              "or registered with the FRA or any other securities regulator, manages no money, and accepts no clients, fees or funds. This "
              "document was prepared and published independently of any trading activity, on a pre-set schedule, and must not be used to "
              "influence the market price of any security. All valuations, scenarios, probabilities and levels are model outputs resting on "
              "explicit, subjective assumptions; they are illustrative, highly uncertain, and likely to prove wrong in material respects. "
              "Public and vendor-sourced data are believed reliable but are not guaranteed; vendor-derived figures are flagged where used. "
              "To the maximum extent permitted by law, the preparer accepts no liability for any decision taken or loss incurred by any "
              "person in reliance on this document. Any person considering an investment decision should conduct their own independent "
              "assessment and consult a financial advisor licensed in their jurisdiction.", size=9, color=GREY)
    para(doc, f"{today.strftime('%d %B %Y')}", size=9, color=GREY)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    # python-docx omits the required w:zoom percent attribute; patch it so the file validates
    import zipfile, os, re
    tmp = out_path + ".tmp"
    with zipfile.ZipFile(out_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for it in zin.namelist():
            data = zin.read(it)
            if it == "word/settings.xml":
                t = data.decode("utf-8")
                t = re.sub(r'<w:zoom\s*/>', '<w:zoom w:percent="100"/>', t)
                t = re.sub(r'(<w:zoom)(?![^>]*w:percent)([^>]*)/>', r'\1 w:percent="100"\2/>', t)
                data = t.encode("utf-8")
            zout.writestr(it, data)
    os.replace(tmp, out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
