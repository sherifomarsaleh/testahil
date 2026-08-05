"""ELEC_Source_Register_05-08-2026.docx — companion bibliography document.
Every input: value, source, date, ring — emitted from study_numbers.json (the
compute script's own INPUTS block), plus the research-sweep bibliography."""
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers.json'))
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def borders(table, color='C9D4D1', sz='4'):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    tblPr.append(b)

def P(text='', size=10, bold=False, italic=False, color=INK, space_after=5, space_before=0, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def H1(t): return P(t, size=13, bold=True, space_before=12, space_after=4)
def H2(t): return P(t, size=11, bold=True, space_before=8, space_after=3)

def table(rows, widths, size=8.6, header=True):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    borders(t); t.autofit = False
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# masthead
t = doc.add_table(rows=1, cols=1); borders(t, color=F_DARK)
c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.1)
r = c.paragraphs[0].add_run('TESTAHIL')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE
P('Electro Cable Egypt (EGX: ELEC) — Source Register', size=15, bold=True, space_before=8, space_after=2)
P('Companion to the Valuation Study dated 5 August 2026. Every figure that enters the model, with its source, '
  'its date, and the research ring it came from.', italic=True, color=GREY, space_after=8)

# READ FIRST
t = doc.add_table(rows=1, cols=1); borders(t, color='C0A45F', sz='6')
c = t.cell(0, 0); shade(c, F_CREAM); c.width = Inches(7.1)
p = c.paragraphs[0]
r = p.add_run('READ FIRST. ')
r.bold = True; r.font.size = Pt(9.6)
r = p.add_run(
    'This register exists so that no number in the study is unfalsifiable. Every input below is a four-field '
    'record — value, source, date, ring — emitted directly by the compute script that built the study '
    '(compute.py), not transcribed by hand afterwards. A bare numeral cannot enter the model: the build fails. '
    'Where a figure is a house judgment rather than an observable, the source field says so in those words. '
    'SPECIFIC TO THIS COMPANY: Electro Cable Egypt’s audited financial-statement PDFs (Mubasher file store, '
    'the company’s own site, and several data aggregators) were unreachable through the channels available '
    'for this study — headline figures are therefore multiply-sourced from bourse-disclosure reporting '
    'services, and several line items are DERIVED, each labelled as such. The one fully-triangulated year '
    '(FY2024) closes to the reported net profit within 0.8% using the derived lines. Testahil is not licensed '
    'by the FRA and publishes no ratings and no price targets.')
r.font.size = Pt(9.6)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

# A — rings
H1('A.  How to read the ring column')
P('Research runs outward-in — world, country, industry, company — before any forecast driver is set, and every '
  'driver must trace to a ring finding or be declared a house judgment.')
rings = [
 ['Ring', 'What it means', 'Count in this register'],
 ['Market', 'Prices and quotes observable on a screen', '—'],
 ['Global', 'World-level facts that reach the company through prices (copper, rates)', '—'],
 ['Country', 'Egypt macro: policy rate, inflation, FX, sovereign risk, tax', '—'],
 ['Industry', 'Cable-sector structure, demand programs, pass-through economics', '—'],
 ['Company', 'ELEC’s own disclosed figures and events', '—'],
 ['House', 'The preparer’s argued judgments — the rows a reader should attack first', '—'],
]
cnt = {}
for k, rec in D['inputs'].items():
    cnt[rec['ring']] = cnt.get(rec['ring'], 0) + 1
for row in rings[1:]:
    row[2] = str(cnt.get(row[0], 0)) + ' records' if row[0] in cnt else '—'
table(rings, [1.0, 4.4, 1.6], size=9)

# B — primary documents
H1('B.  Primary documents and channels relied on')
P('Two channels deliberately NOT used: no figure is taken from the single InvestingPro "fair value" flag '
  '(model output, not disclosure), and the July-2026 "Damodaran 14.87%" figure circulating in secondary '
  'reproductions was not adopted (the identical figure was previously caught as a misquote of the original '
  'file; the house-cached January-2026 original-file figures are used and the difference is sensitized).')
rows = [
 ['Document / channel', 'Publisher / provenance', 'Date', 'Used for'],
 ['FY2024 consolidated results reporting', 'Mubasher / MarketScreener (of the EGX filing)', 'Mar-2025', 'FY24 revenue, net profit, EPS; FY23 comparatives'],
 ['FY2025 consolidated results reporting', 'Arab Finance / Zawya / Decypha / Reuters flash', 'Mar-2026', 'FY25 revenue, NP, EPS, total assets; standalone NP'],
 ['1Q-2026 results reporting', 'Zawya / Arab Finance (of the EGX filing)', 'May-2026', 'Q1-26 loss, sales; Q1-25 restated comparatives'],
 ['9M/H1-2025 results reporting', 'Zawya', 'Aug/Nov-2025', 'Interim trajectory; 9M-25 total assets'],
 ['Company profile (capital, shares, ISIN)', 'Mubasher ELEC page; EGX listing PDFs', 'mid-2026', 'Shares 3,313,540,373; paid-in capital; par 0.20'],
 ['Balance-sheet health snapshot', 'Simply Wall St (FY24 vintage, 22-May-2025)', 'May-2025', 'Debt ~9,000 · cash 828 · equity 3,600 · EBIT 3,400 · coverage 2.0×'],
 ['Ownership & block-trade disclosures', 'Zawya / Arab Finance / Amwal Al Ghad / Alborsa', '2023–Jul-2026', 'Gadwa group ~81%; 2026 block sales at 2.00–2.21'],
 ['CBE MPC & CPI releases', 'Central Bank of Egypt (via FocusEconomics / DNE)', 'Jul-2026', 'Corridor 19.00/20.00%; inflation 14.3%; targets 7%/5% ±2pp'],
 ['Egypt 10Y yield / 5Y CDS', 'investing.com; MoF via DNE; MacroMicro', 'May–Jul-2026', 'rf 22.31% (21-Jul-26 print); CDS 270–330bp'],
 ['Damodaran country-risk file (house cache)', 'ctryprem, "Last updated January 5, 2026"', 'Jan-2026', 'ERP 9.41% (CDS) / 13.94% (rating); default spreads 3.40% / 6.37%'],
 ['Copper / aluminium market data', 'TradingEconomics; INN; ZAMAK; TradingKey', 'Aug-2026', 'Copper ~$6.63/lb COMEX, +51.6% y/y; LME ~$12.8k avg'],
 ['Grid-capex programs', 'Daily News Egypt; EEAS; Zawya', 'Nov-25–Jun-26', 'EETC EGP 45bn plan; EU €690mn; Saudi interconnector 95%'],
 ['Peer financials', 'Arab Finance (SWDY FY25); stockanalysis.com; multiples.vc', 'Mar–Jul-2026', 'SWDY P/E 10.4× / EV-EBITDA ~6×; Riyadh Cables 18×/15×'],
 ['Uploaded EGX daily price history (3,749 rows)', 'user-supplied OHLC export, cleaned by the data-quality gate', '2011–05-Aug-2026', 'Spot, technicals, calibration, simulation'],
]
table(rows, [1.9, 1.9, 0.85, 2.45], size=8.2)

# C — full register
H1('C.  The full input register')
P('Every record below was emitted by compute.py. “House” in the ring column means the number is an argued '
  'assumption, not an observable; each such row is sensitized in the study (§1.9).', size=9.4)
by_ring = {}
for k, rec in D['inputs'].items():
    by_ring.setdefault(rec['ring'], []).append((k, rec))
for ring in ['Market', 'Country', 'Industry', 'Company', 'House']:
    if ring not in by_ring: continue
    H2(f'C.{["Market","Country","Industry","Company","House"].index(ring)+1}  Ring: {ring}  ({len(by_ring[ring])} records)')
    rows = [['Input', 'Value', 'Date', 'Source / provenance']]
    for k, rec in by_ring[ring]:
        v = rec['value']
        if isinstance(v, float): vs = f'{v:,.0f}' if abs(v) >= 1000 else f'{v:,.4g}'
        elif isinstance(v, list): vs = ', '.join(f'{x:,.3g}' for x in v)
        elif isinstance(v, dict): vs = ', '.join(f'{a}={b:g}' for a, b in v.items())
        else: vs = str(v)[:60]
        rows.append([k, vs, rec['date'], rec['source']])
    table(rows, [1.25, 1.15, 0.85, 3.85], size=7.8)

# D — judgments
H1('D.  Numbers that are judgments, not observations')
P('Stated separately because they are the ones a reader should attack first. Each is sensitized in the study.')
rows = [
 ['Judgment', 'Value used', 'Why it is a judgment', 'Where it is sensitised'],
 ['EBITDA margin path', '13% → 19%', 'The true mid-cycle margin is unknowable from outside; set between the 1Q26 trough and the windfall prints', '§1.9 margin × NWC grid, ±3pp'],
 ['NWC intensity path', '108% → 88% of revenue', 'Collection of state-linked receivables is unproven', '§1.9 grid, 76–100% endpoints'],
 ['Net debt FY25 anchor', 'EGP 8,800 mn', 'Derived from twice-sourced totals, not a disclosed print', '±1,000 in bear/bull; SOTP Bridge note'],
 ['Kd 23.5% and the forward path', '23.5% → 15.5%', 'Facility disclosure unreachable; effective-rate checks are themselves partly derived', 'Explicit × terminal WACC grid'],
 ['Terminal rf / ERP / Kd', '10.5% / 7.0% / 15.0%', 'Norm-built from the CBE target + house conventions — a macro view, disclosed as such', 'Terminal WACC × g grid'],
 ['Terminal growth 5%', 'center of 3–7% grid', 'Standing convention for established Egyptian names post-disinflation', '§1.9 grid'],
 ['Justified multiples', '5.5× EV/EBITDA · 6.5× P/E', 'Peer-anchored with a leverage/concentration discount — a judgment of degree', 'Bear/bull columns'],
 ['Capex 1.6% of revenue', 'no disclosed capex any year', 'Derived through the D&A ratio; a capacity program would change it', 'Immaterial at ±50bp; noted §7'],
]
table(rows, [1.5, 1.25, 2.85, 1.5], size=8.2)

# E — falsification
H1('E.  What would falsify the study')
P('The single most useful thing a source register can do is name the evidence that would overturn the work.')
rows = [
 ['Test', 'What to watch', 'When'],
 ['The collection question (the crux)', 'H1-2026 receivables and operating cash flow; any receivables-ageing disclosure', '~mid-Aug-2026 (H1 results)'],
 ['The margin question', 'Two consecutive halves of EBITDA margin ≥ 20% without a devaluation', 'H1-26, FY26 results'],
 ['The rate path', 'CBE MPC decisions vs the modelled 23.5% → 15.5% forward Kd path', 'MPC 20-Aug / 01-Oct / 12-Nov-2026'],
 ['The net-debt anchor', 'Any audited balance sheet: debt, cash, facility schedule', 'On publication'],
 ['The demand thesis', 'EETC tender awards actually reaching ELEC’s order book (vs Elsewedy taking all of it)', 'FY26/27'],
 ['The flow overhang', 'Further controlling-group block sales — or their cessation', 'Ongoing'],
]
table(rows, [1.9, 3.6, 1.6], size=8.6)

P('Testahil · Independent valuation studies · Distributions, not tips. Educational analysis, not investment '
  'advice. Testahil is not licensed by the FRA or any securities regulator and publishes no ratings and no '
  'price targets.', size=8.4, color=GREY, align='center', space_before=10)
doc.save('ELEC_Source_Register_05-08-2026.docx')
print('register saved:', len(doc.tables), 'tables')
