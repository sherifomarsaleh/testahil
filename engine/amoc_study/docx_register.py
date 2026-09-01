"""AMOC_Bibliography_01-09-2026.docx — the companion bibliography document.

Every input in the model: value, source, date and research layer — emitted directly from
study_numbers.json (the compute script's own INPUTS block), plus the document bibliography,
the triangulations, and the negative results.
"""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(HERE, 'study_numbers.json')))
INP = D['inputs']
BASE, BETA = D['base'], D['wacc']['beta']
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_margins(table, top=36, bottom=36, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    table._tbl.tblPr.append(m)


def borders(table, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    table._tbl.tblPr.append(b)


def P(text='', size=10, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def H1(t): return P(t, size=13, bold=True, space_before=12, space_after=4)
def H2(t): return P(t, size=11, bold=True, space_before=8, space_after=3)


def table(rows, widths, size=8.2, header=True):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t); t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def masthead():
    t = doc.add_table(rows=1, cols=1)
    cell_margins(t, 80, 80, 150, 150)
    c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.1)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def fmt(v):
    if isinstance(v, bool):
        return str(v)
    if isinstance(v, float):
        return f"{v:,.4f}".rstrip('0').rstrip('.') if abs(v) < 100 else f"{v:,.1f}"
    if isinstance(v, int):
        return f"{v:,}"
    if isinstance(v, dict):
        return "; ".join(f"{k}: {fmt(x)}" for k, x in v.items())
    if isinstance(v, (list, tuple)):
        return ", ".join(fmt(x) for x in v)
    return str(v)


# ============================================================================
masthead()
H1('Alexandria Mineral Oils Company S.A.E. (EGX: AMOC) — Bibliography and Source Register')
P('Companion document to the valuation study dated 6 August 2026. It records where every number '
  'in that study came from.', size=9.5, color=GREY)

H2('READ FIRST')
P('This document exists so that a reader can check the study rather than trust it. It lists every '
  'input the valuation model uses, together with the value, the source, the date of that source '
  'and the research layer it belongs to. Nothing in the study is computed from a figure that does '
  'not appear here.')
P('Three things are worth knowing before reading the tables. First, inputs marked "House" are '
  'judgements or derivations made by the analyst, not disclosures by the company; each carries '
  'the reasoning that produced it, and a reader is free to disagree and re-run the model. Second, '
  'several figures are TRIANGULATED — derived independently by more than one route and averaged, '
  'with all the routes shown, because the underlying figure was disclosed only through a growth '
  'rate. Third, where a source could not be reached, that is recorded as a negative result at the '
  'end of this document rather than quietly filled in.')

H2('What changed in this edition')
P('The first edition of this study was built WITHOUT the audited financial statements. The '
  'company\'s investor-relations site, the exchange\'s disclosure pages and the commercial '
  'financial-data terminals all refused connections under the network policy in force, so every '
  'company figure came from published reporting of the company\'s releases and from aggregators. '
  'That edition triangulated each material figure across independent sources, reconstructed the '
  'base year from two reported halves, and marked every reconstructed line as reconstructed.')
P('THE FILINGS ARE NOW IN HAND AND THIS EDITION IS BUILT ON THEM. The audited consolidated '
  'statements for the transition period 1 July 2025 to 31 December 2025 (Crowe — Dr A. M. Hegazy '
  '& Co, UNQUALIFIED opinion, signed at Giza on 18 February 2026), the limited-review statements '
  'for the six months to 31 December 2024, and the reviewed statements for the three months to '
  '31 March 2026. Every figure in the company layer below carries one of those three as its '
  'source, with the note number where the filing gives one. No company figure in this edition is '
  'triangulated, reconstructed or inferred.')
P('The change was not cosmetic. Twelve published assumptions were overturned by the filings, '
  'including a capital-expenditure line modelled at roughly five times the actual cash spend, a '
  'depreciation charge modelled at three times the actual, an operating-expense base understated '
  'by a factor of three, a property-plant-and-equipment balance reconstructed at nearly twice the '
  'filed figure, a minority interest inferred at 3.0% against a disclosed 4.645%, and a '
  'tax-disputes provision of EGP 904.6mn that was never carried at all.')
P('THIS EDITION GOES ONE STEP FURTHER. The half-year results for 1 January to 30 June 2026, '
  'disclosed to the Egyptian Exchange on 29-30 July 2026 — one week before the anchor date — '
  'are restored to the record after the previous edition deleted them on a rule that treated '
  '"is it inside one of the four PDFs?" as the test for use. With them, a clean CONTIGUOUS '
  'twelve-month base year to 30 June 2026 exists and is the headline base. The disclosure is a '
  'press release, not a filing, and is flagged REPORTED wherever it is used; its gross-profit '
  'line is REJECTED on a coherence test recorded below and SOLVED from its own profit line '
  f"instead. The fair-value estimate is EGP {D['central']:.2f} in this edition, against "
  'EGP 7.16 in the audited nine-month edition and EGP 9.38 in the pre-filings edition — each '
  'restatement is itself part of the record and none is silently overwritten.')

H2('The research layers')
table([['Layer', 'What it covers'],
       ['Company', 'Alexandria Mineral Oils\' own results releases, general-assembly resolutions '
        'and disclosed operating statistics, as reported by the trade and financial press'],
       ['Country', 'Egyptian macroeconomic and sovereign data: policy rates, government bond '
        'yields, sovereign spreads, the exchange rate, the tax regime, the country equity risk '
        'premium'],
       ['Industry', 'Lubricant base oils, paraffin wax and refined-product markets; the crude '
        'and product price deck'],
       ['Global', 'World interest rates and growth used in the currency-of-discounting '
        'alternative and the terminal ceiling'],
       ['House', 'Analyst judgement — forecast drivers, normalisation choices and cost-of-capital '
        'construction. Each is argued in place rather than asserted']],
      [1.05, 5.95], size=8.6)

H2('Primary sources relied upon')
table([['Source', 'Publisher', 'Date', 'What was taken from it'],
       ['"AMOC Doubles FY 2025/26 Profit as Revenue Expands 35%"', 'Egypt Oil & Gas',
        'Aug 2026',
        'Revenue and profit after tax for the six months to 30 June 2026, on both a consolidated '
        'and a standalone basis, together with operating cash flow. Carried as corroboration '
        'only; the period it covers was IDENTIFIED by reconciling its two reported growth rates '
        'against the constructed comparative'],
       ['"AMOC Net Profit Rises 37% in Q1 2026"', 'Egypt Oil & Gas', 'Apr–May 2026',
        'Consolidated net sales of EGP 10.51bn and consolidated net profit of EGP 635.12mn for '
        'the quarter to 31 March 2026, with the prior-year comparative. Also the confirmation '
        'that the company now reports on CALENDAR quarters'],
       ['"AMOC Reports Sales of EGP 20 Bn in H2 2025" and the general-assembly report',
        'Egypt Oil & Gas; Egyptian Ministry of Petroleum', 'Mar 2026',
        'The July–December 2025 transition period: 808,000 tonnes sold, sales of about EGP 20bn, '
        'a 14.5% growth rate, exports of about 42,000 tonnes up 40%, the EGP 0.40 per share '
        'dividend approved on 28 March 2026, and the four new storage tanks'],
       ['"AMOC\'s consolidated profits rise 2% YoY in H1 FY 2025/26"', 'Arab Finance',
        'Feb 2026',
        'Consolidated sales of EGP 20.735bn against EGP 18.246bn, and consolidated profit after '
        'tax of EGP 656.428mn, +2%. The prior-year comparative half is what makes the base-year '
        'construction possible'],
       ['"AMOC Reports 27% Sales Growth for Q1 FY2025/26"', 'Egypt Oil & Gas', 'Nov 2025',
        'Standalone and consolidated net sales and profit for the quarter to 30 September 2025'],
       ['"AMOC Reports 17.3% Increase in Standalone Net Profit for FY 2024/25" and the FY2024/25 '
        'results commentary', 'Egypt Oil & Gas; MarketScreener', 'Sep–Oct 2025',
        'Standalone profit of EGP 1.49bn (+17.3%) and consolidated profit of EGP 1.55bn (+3%) '
        'for the year to 30 June 2025; sales of 1.26mn tonnes valued at EGP 36.9bn, +10.8% on '
        '2023/24; output of oils and waxes of 172,000 tonnes at 108% of target'],
       ['"AMOC approves FY2025/26 planning budget"; "AMOC Approves EGP 580mm Capital Budget"',
        'Zawya / Reuters; Egypt Oil & Gas', 'Jun–Jul 2025',
        'The approved capital budget of EGP 580.19mn and the planning-budget net sales of EGP '
        '37.332bn — the anchor for the capital-expenditure driver'],
       ['"AMOC\'s Exceptional Fiscal Year 22/23 Results Gain Approval at the Annual General '
        'Assembly"', 'MoneyController', '2023',
        'Cost of sales of EGP 21,218.64mn and gross profit of EGP 1,297.01mn for the year to 30 '
        'June 2023 — the ONLY period for which both the cost line and the margin line are '
        'separately available, and therefore the anchor for the whole margin discussion'],
       ['Company financial summary pages', 'stockanalysis.com; Investing.com; TradingView',
        'Aug 2026',
        'Shares outstanding, market capitalisation, total assets, total liabilities, cash and '
        'equivalents, total debt, dividend per share and payout ratio'],
       ['Monetary Policy Report Q1 2026 and the August 2026 rate decision',
        'Central Bank of Egypt', '2026',
        'Main operation rate 19.50% (corridor 19.00/20.00), held for a second consecutive '
        'meeting; annual headline inflation of 14.30% in June 2026; the 7% (±2pp) fourth-quarter '
        '2026 and 5% (±2pp) fourth-quarter 2028 inflation targets used to build the terminal '
        'risk-free rate'],
       ['Country default spreads and risk premiums, Egypt row',
        'A. Damodaran, Stern School of Business, New York University', 'January 2026',
        'The credit-default-swap-basis equity risk premium and sovereign default spread used as '
        'primary, and the rating-basis pair disclosed as the computed alternative'],
       ['Egyptian pound closing rates and the 10-year government bond yield',
        'Amwal Al Ghad; house cost-of-capital reference', 'Jul–Aug 2026',
        'USD/EGP at 50.25 (close of 50.30/50.40 on 4 August 2026) and the 10-year local-currency '
        'yield of 22.31%'],
       ['AMOC daily price history — open, high, low, close and volume',
        'Vendor export, screened through the house data-quality gate', '2 Jan 2011 – 6 Aug 2026',
        'The full price history: the closing price the study is anchored on, the beta regression, '
        'the volatility estimate and the simulated price distribution']],
      [1.85, 1.72, 0.80, 2.68], size=7.9)

# ---- what the filings replaced ----------------------------------------------
H1('Triangulated figures — RETIRED')
P('The first edition carried a table here of figures disclosed only through growth rates, each '
  'derived by two or three independent routes and averaged, with every route shown. That table is '
  'REMOVED rather than updated: the figures it triangulated are now read straight off the audited '
  'statements. Keeping it would imply a corroboration exercise that no longer has anything to '
  'corroborate.')
P('Two derivations remain, and the study states both on its face. FIRST, the base year: the '
  'twelve contiguous months to 30 June 2026 — the audited transition half plus the half '
  'disclosed to the exchange on 29-30 July 2026 — with no annualisation scalar; the audited '
  'nine months annualised by four thirds is published beside it as the fully-audited '
  'alternative. SECOND, the gross profit of the reported half: the released figure of EGP '
  f"{D['ttm']['gp_h1_released']:,.0f}mn, run through the company's own first-quarter expense "
  'run rates, implies a profit after tax '
  f"{D['ttm']['ct3']:+.1%} above the profit printed in the same release, so it is rejected and "
  f"the figure used, EGP {D['ttm']['gp_h1']:,.0f}mn, is SOLVED from the release's own profit "
  'line. Two coherence tests support the profit line it is solved from: the reported majority '
  f"profit ties to the AUDITED statement of changes in equity within {D['ttm']['ct1']:.1%}, "
  f"and the reported revenue ties to an independent triangulation within {D['ttm']['ct2']:.1%}.")
table([['What the first edition triangulated', 'What the filing says'],
       ['Base-year revenue, reconstructed from two reported halves at EGP 39,996mn',
        f"Twelve months to 30-Jun-2026, no scalar: EGP {D['ttm']['rev']:,.0f}mn "
        f"(audited nine months annualised: {D['audited']['base_rev']:,.0f}mn)"],
       ['Gross margin, built from a per-tonne cost stack at 6.06%',
        f"Twelve months to 30-Jun-2026: {D['ttm']['gm']:.2%} "
        f"(nine audited months: {D['audited']['base_gm']:.2%})"],
       ['Cost of sales, built from house yields, energy intensity and a solved feedstock '
        'differential',
        'Note 15-A as filed: raw materials ' + f"{D['unit']['cost_share']['raw']:.1%}" +
        ', salaries ' + f"{D['unit']['cost_share']['salaries']:.1%}" + ', other ' +
        f"{D['unit']['cost_share']['other']:.1%}" + ', supporting materials ' +
        f"{D['unit']['cost_share']['support']:.1%}" + ', depreciation ' +
        f"{D['unit']['cost_share']['dep']:.1%}"],
       ['Three product lines from a reviewer-sourced table',
        f"Note 14-A: eight lines, {D['unit']['tot_t']:,.0f} tonnes and EGP "
        f"{D['unit']['tot_v']:,.0f}mn, tonnage and value both disclosed"],
       ['Property, plant and equipment reconstructed as a residual at EGP 2,403mn',
        f"Note 6 as filed: EGP {D['audited']['ppe']:,.0f}mn including projects under construction"],
       ['Depreciation modelled at 1.1% of revenue',
        f"Actual, annualised from two filings: EGP {D['audited']['dep_ann']:,.0f}mn"],
       ['Capital expenditure modelled at 1.45% of revenue',
        f"Actual cash paid, annualised: EGP {D['audited']['capex_ann']:,.0f}mn"],
       ['Operating expense modelled at 1.25% of revenue',
        f"Actual, annualised: EGP {D['audited']['opex_ann']:,.0f}mn, "
        f"{D['audited']['opex_ann']/D['audited']['base_rev']:.2%} of revenue"],
       ['Minority interest inferred at 3.0% of group profit',
        f"Disclosed: {D['audited']['nci_share']:.3%}; AMOC owns 86.45% of Alexandria Wax Products"],
       ['Effective tax rate assumed at 23.5%',
        f"Computed from both filed periods: {D['audited']['tax_eff']:.2%}"],
       ['No tax-disputes provision carried',
        f"Note 10-1: EGP {D['audited']['provisions']:,.0f}mn"]],
      [3.05, 4.00], size=7.9)

H1('Triangulated figures — every route shown')
P('Where a figure was disclosed only through a growth rate, it is derived by more than one '
  'independent route and the AVERAGE is carried. The routes are on the face of the companion '
  'model as live formulas, not asserted here.')
P('The residual triangulation table and the period-identification check that followed it have '
  'both been removed. The first edition had to identify which period a press release covered by '
  'matching two reported growth rates against a constructed comparative; the filings state the '
  'period on their face, so the exercise is redundant. The audited transition period runs 1 July '
  '2025 to 31 December 2025 and says so in the auditor\'s report.')

H1('The full input register')
P('Every hardcoded figure the compute script consumes, in the order it is declared. A bare '
  'numeral anywhere in the inputs block fails the build, so this table is complete by '
  'construction.')
by_ring = {}
for k, v in INP.items():
    by_ring.setdefault(v['ring'], []).append((k, v))
for ring in ('Company', 'Industry', 'Country', 'Global', 'House'):
    if ring not in by_ring:
        continue
    H2(f'{ring} layer — {len(by_ring[ring])} inputs')
    rows = [['Input', 'Value', 'Date', 'Source and reasoning']]
    for k, v in by_ring[ring]:
        rows.append([k, fmt(v['value']), v['date'], v['source']])
    table(rows, [1.90, 1.55, 0.95, 2.67], size=7.5)

# ---- negative results -------------------------------------------------------
H1('Negative results — what could not be sourced')
P('Recorded rather than filled in. Each of these is a place where a reader with better access '
  'can improve the study.')
table([['What was sought', 'Why it matters', 'Outcome'],
       ['The audited consolidated financial statements for the year to 30 June 2025 and the '
        'transition period to 31 December 2025',
        'Would replace the entire reconstructed balance sheet and the closed income statement '
        'with disclosed lines',
        'NOT REACHED. The company\'s investor-relations site and its published statement PDFs '
        'refused connection under the network policy in force'],
       ['The exchange\'s own disclosure pages for the company',
        'Would confirm the transition-period filing and the year-end change directly rather than '
        'through reporting of them',
        'NOT REACHED. Same cause'],
       ['THE AUDITED FINANCIAL STATEMENTS THEMSELVES',
        'The whole of the historical record: statement of financial position, profit or loss, '
        'cash flows, changes in equity, and every explanatory note',
        'REACHED. This edition is built on the audited consolidated statements for the '
        'transition period 1-Jul-2025 to 31-Dec-2025 (Crowe, Dr A. M. Hegazy & Co, UNQUALIFIED '
        'opinion, signed Giza 18-Feb-2026), the limited-review statements for the six months to '
        '31-Dec-2024, and the reviewed statements for the three months to 31-Mar-2026. The '
        'previous edition of this study was built on triangulated press reporting because the '
        'filings could not be reached from the build environment. Everything that triangulation '
        'stood in for has been replaced by the filing'],
       ['A segmental or product-line revenue note',
        'The revenue build',
        'REACHED — note 14-A gives EIGHT product lines with tonnes AND value for the transition '
        f"half: {D['unit']['tot_t']:,.0f} tonnes for EGP {D['unit']['tot_v']:,.0f}mn. Realisations "
        'per tonne are that note divided by itself; nothing is reconstructed and no crack '
        'multiple, crude parity or feedstock differential is needed or used. The specialty slate '
        f"is {D['unit']['spec_share_t']*100:.2f}% of tonnage and {D['unit']['spec_share_v']*100:.2f}% "
        'of value'],
       ['A cost-of-sales breakdown',
        'The margin, which on a processor at this gross margin IS the valuation',
        'REACHED — note 15-A splits cost of sales five ways: raw materials '
        f"{D['unit']['cost_share']['raw']*100:.1f}%, salaries {D['unit']['cost_share']['salaries']*100:.1f}%, "
        f"other (natural gas, electricity, water, spare parts, maintenance and the EPROM "
        f"operating contract) {D['unit']['cost_share']['other']*100:.1f}%, supporting materials "
        f"{D['unit']['cost_share']['support']*100:.1f}% and depreciation "
        f"{D['unit']['cost_share']['dep']*100:.1f}%. The previous edition BUILT this stack from "
        'house estimates of yields, energy intensity and a solved feedstock differential, carried '
        'no salaries line inside cost of sales at all, and estimated chemicals at roughly five '
        'times the disclosed figure'],
       ['A disclosed non-controlling-interest balance or profit share',
        'The minority deduction in the bridge is inferred from the gap between consolidated and '
        'standalone profit',
        f"NOT FOUND. Held at {D['dcf']['nci_share']*100:.1f}% and sensitised to "
        f"{D['dcf']['nci_alt']*100:.0f}%, which moves the answer by "
        f"{(D['dcf']['ps_nci_alt']/D['dcf']['ps']-1)*100:+.1f}%"],
       ['A disclosed depreciation charge, capital-expenditure actual, or property, plant and '
        'equipment balance',
        'All three are drivers of the free-cash-flow waterfall',
        'NOT FOUND as disclosed lines. The capital budget of EGP 580.19mn is disclosed and is '
        'used as the anchor; depreciation is set as a share of revenue and property, plant and '
        'equipment is the residual against disclosed total assets'],
       ['A traded credit-default-swap quote for Egypt as at August 2026',
        'The sovereign spread netted out of the risk-free rate',
        'NOT REACHED live. The January-2026 published country-premium file is used, and the '
        'rating-basis alternative is computed and published as a value'],
       ['Listed peer multiples for Egyptian or regional lubricant base-oil processors',
        'The relative lens',
        'NO DIRECT LISTED COMPARATOR EXISTS — the company is the only listed refinery on the '
        'exchange. The multiple range used is taken from the international Group I base-oil and '
        'independent-processor band and is disclosed as such rather than presented as a peer set']],
      [2.30, 2.00, 2.65], size=7.9)

H1('Method note')
P('All financial arithmetic in the study originates in an executed, asserting compute script. '
  'Every hardcoded figure enters through the four-field register reproduced above; the script '
  'refuses to emit its results unless the enterprise-to-equity bridge closes exactly, terminal '
  'value as a percentage of enterprise value is computed and printed, the implied fair value '
  'sits inside a stated plausibility band, and net debt and minority interests carry the correct '
  'signs into the bridge.')
P('The companion workbook was verified on the DELIVERED file rather than on the script that '
  'wrote it: every formula cell was independently re-evaluated and required to reproduce the '
  'model\'s own value with none unresolvable and none unchecked, and every input was perturbed '
  'in place with the whole workbook re-evaluated to confirm it moves the headline in the '
  'asserted direction.')

doc.save(os.path.join(HERE, 'AMOC_Bibliography_01-09-2026.docx'))
print('wrote AMOC_Bibliography_01-09-2026.docx')
