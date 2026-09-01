"""External-reader scrub — the check that was attested but never written.

`compute.py` set external_reader_scrub=True and nothing performed a scrub. That is
precisely the defect [R-ENF-01] names: a self-attested boolean is never a check. This
runs over the DELIVERED files — the Word study, the bibliography and the workbook's
text — and fails on internal-procedure vocabulary a reader should never meet.

TWO KINDS OF PATTERN, DELIBERATELY SEPARATED. `FORBIDDEN` is house machinery: rule
identifiers, gate names, verdict tokens, the names of internal modules and registers.
A single hit fails. `AMBIGUOUS` is vocabulary that is ordinary English in one sense and
house jargon in another — "register" is a commercial register, "step" is a step-change —
so those are REPORTED with their surrounding words and adjudicated, never auto-failed.
A check that cries wolf is one everyone learns to ignore.
"""
import os, re, sys, json, zipfile

HERE = os.path.dirname(os.path.abspath(__file__))

FORBIDDEN = [
    (r"\bR-[A-Z]{3,6}-\d{2}\b", "standing-rule identifier"),
    (r"\bSIGCM\b", "internal mandate name"),
    (r"\bassert_(sigcm|ground_up|model_study|beta_provenance)\b", "gate function name"),
    (r"\bPENDING_REVIEW\b", "internal review directory"),
    (r"\bStep 0\.0\b|\bStep 2A\b", "internal step name"),
    (r"\b(Global|Country|Industry|Company) ring\b", "sweep ring"),
    (r"\bCRPS\b", "proper-score name, methodology page only"),
    # PARITY is matched CASE-SENSITIVELY in caps. Lowercase "parity" is an ordinary
    # word in this document — crude parity, a currency at parity — and a check that
    # cries wolf is one everyone learns to ignore.
    (r"\bPARITY\b", "retired verdict token", True),
    (r"\bwidth_cal\b|\bnu=|\bmc_v3\b|\bmarket_profiles\b", "engine internals"),
    (r"\bLONO\b|\bblock bootstrap\b", "calibration internals"),
    (r"\bwalk-forward\b(?!\s+(test|of|on))", "unqualified internal test name"),
    (r"\bDriverLine\b|\bfitted_configs\b|\bpanel_refresh\b", "module or class name"),
    (r"\bthe model report\b|\bADNOCLS\b", "internal exemplar"),
    (r"\bstudy_numbers\.json\b|\bcompute\.py\b|\bpanel\.py\b", "build artefact"),
    # A repository path is not a source a reader can follow. This pattern was added after
    # the table-width check surfaced "engine/raw_ohlc/EG/AMOC.csv" sitting in the delivered
    # bibliography as though it were a citation — found by the wrong check, which is how
    # these things usually surface.
    (r"\bengine/[A-Za-z0-9_/]+\.(csv|py|json|md)\b|\bscripts/[A-Za-z0-9_]+\.py\b",
     "repository path presented as a source"),
]
AMBIGUOUS = [
    (r"\bregister\b", "ordinary sense: commercial register, input register"),
    (r"\bgate\b", "ordinary sense: a test the study describes to its reader"),
    (r"\bstep\b", "ordinary sense: a step-change"),
    (r"\bengine\b", "ordinary sense: the cash-flow engine of the workbook"),
]


def docx_text(path):
    with zipfile.ZipFile(path) as z:
        parts = [n for n in z.namelist()
                 if n.startswith("word/") and n.endswith(".xml")]
        out = []
        for n in parts:
            xml = z.read(n).decode("utf-8", "replace")
            out.append(re.sub(r"<[^>]+>", " ", xml))
    return re.sub(r"\s+", " ", " ".join(out))


def xlsx_text(path):
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist() if n.endswith("sharedStrings.xml")]
        if not names:
            return ""
        xml = z.read(names[0]).decode("utf-8", "replace")
    return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", xml))


def scan(name, text):
    hard, soft = [], []
    for entry in FORBIDDEN:
        pat, why = entry[0], entry[1]
        case_sensitive = len(entry) > 2 and entry[2]
        for m in re.finditer(pat, text, 0 if case_sensitive else re.I):
            hard.append({"file": name, "term": m.group(0), "why": why,
                         "context": text[max(0, m.start() - 60):m.start() + 60].strip()})
    for pat, why in AMBIGUOUS:
        n = len(re.findall(pat, text, re.I))
        if n:
            soft.append({"file": name, "pattern": pat, "hits": n, "note": why})
    return hard, soft


def main():
    files = []
    for fn in sorted(os.listdir(HERE)):
        if fn.endswith("_public.docx") and "01-09-2026" in fn:
            files.append((fn, docx_text(os.path.join(HERE, fn))))
        if fn.startswith("AMOC_Bibliography_01-09-2026") and fn.endswith(".docx"):
            files.append((fn, docx_text(os.path.join(HERE, fn))))
        if fn.endswith("01092026_public.xlsx"):
            files.append((fn, xlsx_text(os.path.join(HERE, fn))))
    assert files, "SCRUB FAIL — no delivered files matched. An empty scan is not a clean scan."
    hard, soft = [], []
    for name, text in files:
        h, s = scan(name, text)
        hard += h
        soft += s
    json.dump({"files": [f for f, _ in files], "hard": hard, "soft": soft},
              open(os.path.join(HERE, "scrub_result.json"), "w"), indent=1)
    print("scrubbed %d delivered files: %s" % (len(files), ", ".join(f for f, _ in files)))
    print("hard hits (must be zero): %d" % len(hard))
    for h in hard[:25]:
        print("   %-28s %-22s %s" % (h["file"][:28], h["term"], h["context"][:70]))
    print("ambiguous, reported not failed:")
    for s in soft:
        print("   %-28s %-14s %3d  %s" % (s["file"][:28], s["pattern"], s["hits"], s["note"]))
    if hard:
        sys.exit(1)
    print("SCRUB PASS")


if __name__ == "__main__":
    main()


# ---------------------------------------------------------------------------
# TABLE DISCIPLINE — the other half of the depth bar, and also programmatic.
# Fixed layout with explicit widths, and no starved or bloated column. A column
# is STARVED when its width leaves less room than its longest cell needs; it is
# BLOATED when it is more than three times the width its content requires. Both
# are read off the DELIVERED file, not off the builder's intentions.
# ---------------------------------------------------------------------------
def table_discipline(path):
    """Fixed layout, and no starved or bloated column, read off the DELIVERED file.

    THE FIRST VERSION OF THIS CHECK WAS WRONG and it is worth saying how. It compared a
    column's width against the LONGEST CELL, so every prose callout — a single 7-inch
    column holding two thousand characters of perfectly well-set text that wraps over
    fifteen lines — came back "starved". Eight of its nine hits were of that kind. A
    check that cries wolf is one everyone learns to ignore, so the test now asks the
    question that actually matters: is the column too narrow for something that CANNOT
    wrap? That is the longest unbreakable token — a nine-digit figure, a long
    hyphen-free word — not the longest paragraph.
    """
    import docx
    doc = docx.Document(path)
    bad = []
    for ti, t in enumerate(doc.tables):
        ncol = len(t.columns)
        for ci in range(ncol):
            w = t.columns[ci].width
            if w is None:
                bad.append({"table": ti, "col": ci,
                            "problem": "no explicit width — layout is not fixed"})
                continue
            win = w.inches
            cells = [c.text.strip() for c in t.columns[ci].cells]
            tokens = [tok for c in cells for tok in c.split()]
            longest_token = max((len(tok) for tok in tokens), default=0)
            longest_cell = max((len(c) for c in cells), default=0)
            # ~11 characters an inch at the sizes these tables use
            if longest_token / 11.0 > win * 1.05:
                bad.append({"table": ti, "col": ci, "problem": "starved",
                            "width_in": round(win, 2), "longest_token": longest_token})
            elif longest_cell <= 8 and win > 1.8:
                bad.append({"table": ti, "col": ci, "problem": "bloated",
                            "width_in": round(win, 2), "longest_cell": longest_cell})
    return len(doc.tables), bad


if __name__ == "__main__" and os.environ.get("SCRUB_TABLES", "1") == "1":
    for fn in sorted(os.listdir(HERE)):
        if fn.endswith(".docx") and "01-09-2026" in fn:
            n, bad = table_discipline(os.path.join(HERE, fn))
            print("table discipline %-40s %2d tables, %d problem(s)" % (fn[:40], n, len(bad)))
            for b in bad[:12]:
                print("     ", b)
