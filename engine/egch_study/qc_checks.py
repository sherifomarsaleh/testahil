"""EGCH — the model study's programmatic depth-bar checks.

Four standards are machine-checkable and are checked here rather than asserted:
  (4) EXTERNAL-READER SCRUB   — internal-procedure vocabulary in any delivered document
  (5) FIGURE DISCIPLINE       — zero transparency on every figure
  (6) TABLE DISCIPLINE        — no table exceeds its text block, no starved or bloated column
  (3) NUMERIC TRACEABILITY    — no financial numeral typed into any builder

A scrub false positive is fixed in THIS CHECKER, never by rewording legitimate finance
prose in the study.
"""
import json, os, re, sys, glob
from docx import Document
from docx.shared import Inches
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
DOCS = ['EGCH_Valuation_Study_03-09-2026.docx', 'EGCH_Bibliography_03-09-2026.docx']
fails = []

# ---------------------------------------------- (4) external-reader scrub -----
# Internal-procedure vocabulary. Each pattern is a word-boundary regex so ordinary
# finance prose does not trip it: "step" alone is a normal English word, "Step 0" and
# "Step 2A" are ours; "gate" appears in no legitimate sentence here but "gate" inside
# "investigate" must not match, hence the boundaries.
PATTERNS = [
    r"\bstep\s*0\b", r"\bstep\s*2a\b", r"\bstep\s*0\.0\b", r"\bsigcm\b",
    r"\bsweep register\b", r"\binformation sweep\b", r"\bfour[- ]ring\b",
    r"\bglobal ring\b", r"\bcountry ring\b", r"\bindustry ring\b", r"\bcompany ring\b",
    r"\bqc gate\b", r"\bquality gate\b", r"\bcalibration gate\b", r"\bdata-quality gate\b",
    r"\bmateriality gate\b", r"\bpromotion rule\b", r"\bstanding research protocol\b",
    r"\bmodel study\b", r"\bdepth bar\b", r"\bdriver ledger\b", r"\bcalibration ledger\b",
    r"\bledger row\b", r"\broll[- ]forward\b", r"\bcohort\b", r"\bstrike date\b",
    r"\bgrade date\b", r"\bcrps\b", r"\bprobability[- ]integral\b", r"\bpit histogram\b",
    r"\bbacktest\b", r"\bwalk[- ]forward\b", r"\bmonte carlo\b", r"\bcone\b",
    r"\bwidth_cal\b", r"\bmc_v3\b", r"\bmarket_profiles\b", r"\bresearch_protocol\b",
    r"\bwacc_builder\b", r"\bdata_quality\b", r"\badaptive_width\b", r"\bhorizons\.py\b",
    r"\bstudy_numbers\b", r"\binput_register\b", r"\bxlsx_expected\b",
    r"\bpanel verdict\b", r"\bparity[- ]flagged\b", r"\bboundary case\b",
    r"\bpass/fail verdict\b",
    # added 1 September 2026 after an audit found the retired skill figure and engine names in prose
    r"\bskill\b", r"\bproper scoring rule\b", r"\bscoring rule\b", r"\bskill score\b",
    r"\b(non-overlapping|resolved|three-month|3-month|back-?tested|rolling) windows?\b",
    r"\bwindows since\b", r"\bbeta_result\b", r"\bown_stock_beta\b", r"\bbeta_regression\b",
    r"\bwacc_result\b", r"\bforward_ranges\b", r"\bscores\.json\b", r"\bdiagnostics\.json\b",
    r"\bcorrections_log\b", r"\blessons_register\b", r"\bpanel\.py\b", r"\bcompute\.py\b",
    r"\buniformity test\b", r"\bp-value\b",
    # ADDED 03-Sep-2026: THIS LIST NAMED FIFTY-EIGHT THINGS AND NOT THE TWO MOST OBVIOUS.
    # The delivered bibliography was shipping three standing-rule identifiers and a
    # repository path — "[R-MACRO-01]", "[R-LENS-03]", "engine/macro_paths/EG.json" —
    # straight out of an input register's source field, and this scrub reported 0 hits
    # across 68 patterns because it was enumerating procedure NOUNS and no pattern
    # matched the shape of an identifier or a path. AMOC's scrub carries both and caught
    # the identical sentence in its own bibliography the same hour.
    #
    # A list of forbidden words cannot be complete, so these two are matched by SHAPE:
    # any [R-AREA-NN] identifier, and any engine/ or scripts/ path with a file extension.
    # Neither can occur innocently in a document written for an outside reader.
    r"\[r-[a-z]+-\d+", r"\b(?:engine|scripts)/[a-z0-9_./-]+\.(?:py|json|md|csv|js)\b",
    r"\bmacro_paths\b", r"\blessons\.py\b", r"\bmacro_path\b",
]
scrub_hits = []
for f in DOCS:
    d = Document(f)
    text = " ".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                text += " " + c.text
    low = text.lower()
    for pat in PATTERNS:
        for m in re.finditer(pat, low):
            ctx = low[max(0, m.start() - 45):m.end() + 45].replace("\n", " ")
            scrub_hits.append(f"{f}: /{pat}/ -> ...{ctx}...")
print(f"(4) external-reader scrub : {len(PATTERNS)} patterns, {len(scrub_hits)} hits")
for h in scrub_hits[:12]:
    print("   !", h)
if scrub_hits:
    fails.append("external_reader_scrub")

# ------------------------------------------------------ (5) figure discipline -
figs = sorted(glob.glob('fig*.png'))
transparent = []
for f in figs:
    im = Image.open(f)
    if im.mode in ('RGBA', 'LA'):
        a = im.getchannel('A')
        if a.getextrema()[0] < 255:
            transparent.append(f)
print(f"(5) figure discipline     : {len(figs)} figures, {len(transparent)} with transparency")
if transparent:
    fails.append("figure_discipline")
    for f in transparent:
        print("   !", f)

# ------------------------------------------------------- (6) table discipline -
STARVED = 0.55     # inches — below this an ordinary label wraps mid-word
problems = []
for f in DOCS:
    d = Document(f)
    sec = d.sections[0]
    block = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0
    for i, t in enumerate(d.tables):
        widths = [c.width / 914400.0 if c.width else 0 for c in t.columns]
        total = sum(widths)
        if total > block + 0.02:
            problems.append(f"{f} table {i+1}: total {total:.2f}in exceeds the {block:.2f}in text block")
        for j, w in enumerate(widths):
            if 0 < w < STARVED:
                problems.append(f"{f} table {i+1} column {j+1}: {w:.2f}in is starved")
            if w > block * 0.75 and len(widths) > 1:
                problems.append(f"{f} table {i+1} column {j+1}: {w:.2f}in is bloated")
n_tables = sum(len(Document(f).tables) for f in DOCS)
print(f"(6) table discipline      : {n_tables} tables checked, {len(problems)} problems")
for p in problems[:12]:
    print("   !", p)
if problems:
    fails.append("table_discipline")

# --------------------------------------------------- (3) numeric traceability -
# No financial numeral may be typed into a builder. A builder may carry structural
# integers (row and column indices, font sizes, column widths, loop bounds) — those are
# layout, not finance. The test looks for numerals with a thousands separator or more
# than four significant digits, which is what a financial figure looks like.
BUILDERS = ['docx_egch.py', 'docx_biblio.py', 'build_xlsx.py', 'figures.py']
# Three shapes of financial numeral, all of which have reached a builder at some point:
#   1,234,567 or 1_234_567   — separated thousands
#   1101.6e6                 — scientific notation, which the separated-thousands test
#                              cannot see and which hid a stale materials figure until a
#                              recalculation caught it downstream
#   1292.0 / 8602.606        — any bare number of a thousand or more
FIN = re.compile(r"(?<![\w.])("
                 r"\d{1,3}(?:[_,]\d{3})+(?:\.\d+)?"
                 r"|\d+(?:\.\d+)?[eE]\d+"
                 r"|\d{4,}(?:\.\d+)?"
                 r")(?![\w])")
SCALE_TOKENS = {'1000', '10000', '100000', '1000000', '914400', '365', '1e6', '1e3',
                '1e5', '1e9'}
typed = []
for f in BUILDERS:
    for n, line in enumerate(open(f), 1):
        s = line.split('#')[0]
        if re.search(r'color\s*=|set_[xy]lim|set_[xy]ticks|figsize|dpi=', s):
            continue              # palette hex and axis geometry are layout, not finance
        for m in FIN.finditer(s):
            tok = m.group(0)
            plain = tok.replace('_', '').replace(',', '')
            if plain.lower() in SCALE_TOKENS:
                continue          # unit conversions and scale factors, not finance
            try:                  # a bare calendar year in prose is not a financial figure
                if float(plain) == int(float(plain)) and 1900 <= int(float(plain)) <= 2100 \
                        and '.' not in plain and 'e' not in plain.lower():
                    continue
            except ValueError:
                pass
            typed.append(f"{f}:{n}: {tok}  |  {line.strip()[:80]}")
print(f"(3) numeric traceability  : {len(BUILDERS)} builders, {len(typed)} typed financial numerals")
for t in typed[:12]:
    print("   !", t)
if typed:
    fails.append("numeric_traceability")

json.dump(dict(scrub_patterns=len(PATTERNS), scrub_hits=scrub_hits,
               figures=len(figs), transparent=transparent,
               tables=n_tables, table_problems=problems,
               builders=BUILDERS, typed_numerals=typed, fails=fails),
          open('qc_checks.json', 'w'), indent=1)
print()
print(("PASS — all four programmatic standards clean" if not fails
       else "FAIL: " + ", ".join(fails)))
sys.exit(0 if not fails else 1)
