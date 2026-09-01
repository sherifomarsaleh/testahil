"""EGCH_Bibliography_01-09-2026.docx — a standalone source register.

Every figure that reaches the study or the model traces to a row here: what it is, where
it came from, what kind of source that is, and the date the source itself carries.
Reads the sweep register and the filings source log — nothing is typed twice.
"""
import json, os, re, sys
HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SW = json.load(open('sweep_register.json'))
IRJ = json.load(open('input_register.json'))
from inputs import V
from docprops import strip_stub_counts
LN = json.load(open('lenses.json'))
EXJ = json.load(open('experts.json'))
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(11), Inches(8.5)
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(9.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.05


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_margins(t, top=40, bottom=40, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    t._tbl.tblPr.append(m)


def borders(t, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    t._tbl.tblPr.append(b)


def P(text='', size=9.5, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def H1(t):
    return P(t, size=14, bold=True, space_before=12, space_after=5)


def table(rows, widths, size=8.5):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; shade(c, F_PANEL)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ----------------------------------------------------------------- masthead --
t = doc.add_table(rows=1, cols=1); cell_margins(t, 90, 90, 160, 160)
c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(9.8)
p = c.paragraphs[0]
r = p.add_run('Testahil · Bibliography and source register')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE
r2 = p.add_run('   Egyptian Chemical Industries (KIMA), EGX: EGCH — 1 September 2026')
r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
doc.add_paragraph().paragraph_format.space_after = Pt(0)

P("This register lists every source behind the valuation study and the accompanying "
  "workbook. It separates the company's own issued documents, which are the only "
  "permitted source for anything the company itself reports, from external material used "
  "for market context and forecast drivers, and from cross-checks that inform no figure.",
  size=10)


H1("1.  The company's own issued documents")
P("Obtained from the company's investor-relations page at kimaegypt.com, which serves its "
  "filings through the Mist investor-relations portal. Every document was downloaded, "
  "opened and read; all are scanned images, so each figure was read from the statement "
  "itself and crossfooted against its own subtotals.", size=9.2, color=GREY)
rows = [["Document", "Period", "Report date", "Auditor", "Retrieved", "What it sources"]]
DOCS = [
 ("Audited financial statements, ten annuals", "Years to 30 June 2009, 2010, 2011, 2013, 2014, 2016, 2018, 2019, 2020 and 2021 (with each prior year as comparative)",
  "2009 to Nov 2021", "Central Auditing Organization", "1 Sep 2026",
  "The eighteen-year reported history on which the forecasting method was tested before use; "
  "not a source of any forecast driver"),
 ("Audited financial statements", "Year to 30 June 2022 (as comparatives)", "8 Oct 2023",
  "Central Auditing Organization; PKF Rashed Badr & Co", "8 Aug 2026",
  "Opening income statement and balance sheet of the four-year history"),
 ("Audited financial statements", "Year to 30 June 2023", "8 Oct 2023",
  "Central Auditing Organization; PKF Rashed Badr & Co", "8 Aug 2026",
  "Peak-utilisation reference year; cost-of-sales and selling-cost notes"),
 ("Audited financial statements", "Year to 30 June 2024", "23 Oct 2024",
  "Central Auditing Organization", "8 Aug 2026",
  "Trough-utilisation year; the one-off investment-property revaluation gain; "
  "the repayment of the pound tranche of the project loan"),
 ("Audited financial statements", "Year to 30 June 2025", "23 Sep 2025",
  "Central Auditing Organization", "8 Aug 2026",
  "The anchor year: note 20 revenue split, the auditor's production and unit-cost table, "
  "note 21 cost of sales, note 22 selling cost, note 18 loans, note 14 share capital, "
  "note 28 gas price and capacity plates, the appropriation statement"),
 ("Interim statements, limited review", "Three months to 30 September 2025", "13 Nov 2025",
  "Central Auditing Organization; Nasr Abou El Abbas & Co (Morison Global)", "8 Aug 2026",
  "Export volume and price movement; the gas-loss quantification and its unit valuation"),
 ("Interim statements, limited review", "Six months to 31 December 2025", "10 Feb 2026",
  "Central Auditing Organization; Nasr Abou El Abbas & Co (Morison Global)", "8 Aug 2026",
  "Margin confirmation; the partial disposal of the listed stake"),
 ("Interim statements, limited review", "Nine months to 31 March 2026", "20 May 2026",
  "Central Auditing Organization; Nasr Abou El Abbas & Co (Morison Global)", "8 Aug 2026",
  "The balance sheet used in the equity bridge; the company's own budget column; the "
  "translation loss on the dollar debt"),
]
for d in DOCS:
    rows.append(list(d))
table(rows, [1.5, 1.5, 0.85, 1.9, 0.75, 3.3], size=8.2)
P("The ten annual statements for 2009 to 2021 in the first row were retrieved on 1 September "
  "2026 and used only to test the forecasting method on the company's own history; no "
  "forecast driver is taken from them. Also indexed on the same portal but not used: the "
  "quarterly board and shareholder-structure disclosure reports through 30 June 2026. The "
  "portal's audit-reports section holds only audit-committee reports and its newest item "
  "covers the period to 31 December 2021; the statutory auditor's reports for the years "
  "used here are bound inside the annual statements above.", size=9, color=GREY)

# ----------------------------------------------- 2. market and macro data ----
H1("2.  Market data, sovereign data and macroeconomic series")
rows = [["Item", "Value used", "As of", "Source", "Where it is used"]]
rows += [
 ["Share price", f"EGP {V('spot_price'):.2f}", "6 Aug 2026", "Exchange close, from the study's own price library",
  "Market capitalisation, the equity weight in the cost of capital, and every comparison"],
 ["Share count", f"{V('shares_outstanding'):,}", "30 Jun 2025", "Note 14 of the audited statements",
  "Per-share values. NOT taken from an exchange page or an aggregator"],
 ["Ten-year government bond yield", f"{V('rf_observed')*100:.2f}%", "6 Aug 2026", "Market quote",
  "The observed risk-free rate before normalisation"],
 ["Treasury bond coupon cross-check", f"{V('sovereign_bond_coupon')*100:.3f}%", "2026", "New EGP 120.9bn issue to May 2029",
  "Corroborates the sovereign yield above"],
 ["Sovereign rating and default spread", f"{V('moodys_rating')}, "
  f"{V('sov_spread_rating')*100:.2f}%", "Jan 2026",
  "Country-premium workbook, Egypt row, read from the original file",
  "Normalisation of the risk-free rate, rating basis"],
 ["Country equity risk premium", f"{V('country_risk_premium_rating')*100:.2f}%",
  "Jan 2026", "Same workbook: the total rating-basis premium less the mature-market premium",
  "The country component of the equity premium"],
 ["Total equity risk premium", f"{V('erp_rating')*100:.2f}% "
  f"(rating) / {V('erp_cds_damodaran')*100:.2f}% (CDS)", "Jan 2026",
  "Same workbook, both columns", "Cost of equity, published on both bases"],
 ["Exchange rate", f"EGP {V('usd_egp_spot'):.2f} per US dollar", "7 Aug 2026", "Market quote",
  "The starting point of the currency path"],
 ["Urea, granular, free on board Egypt", f"US${V('urea_fob_egypt'):.0f} per tonne", "7 Aug 2026",
  "Listed futures contract", "The export-price anchor and the crux sensitivity"],
 ["Egyptian headline inflation", f"{V('cpi_latest')*100:.1f}% year on year", "June 2026", "Official statistics",
  "The domestic cost escalator"],
 ["Policy rate", f"{V('policy_rate')*100:.2f}% main operation rate, mid-corridor", "9 July 2026", "Central bank decision",
  "Context for the rate path; the terminal build uses the inflation target, not this rate"],
 ["Treasury-bill yields", f"{V('tbill_yield_range')[0]*100:.2f}% to {V('tbill_yield_range')[1]*100:.2f}% by tenor", "6 Aug 2026",
  "Secondary market quotes — the central bank's own auction pages were unreachable",
  "Context only; no valuation figure depends on them"],
]
table(rows, [1.85, 1.55, 0.85, 2.35, 3.2], size=8.2)

# ------------------------------------------------- 3. context and drivers ----
H1("3.  External context and forecast drivers")
P("These sources inform forward-looking judgements. None of them is a source for any "
  "figure the company reports about itself.", size=9.2, color=GREY)
rows = [["Subject", "What it establishes", "Source type", "Date"]]
CTX = [
 ("Egyptian gas curtailment", "Industrial gas rationing after the interruption of Israeli "
  "export flows; producers curtailed by up to half", "Energy and commodity press", "Mar 2026"),
 ("Regional urea supply shock", "Middle East conflict lifting global urea above US$700 a "
  "tonne; the region carries about 35% of seaborne trade", "Commodity research", "Mar 2026"),
 ("Fertilizer quota regime", "Cabinet decision 170 of 24 November 2021 and trade-ministry "
  "decree 241 of 2021; the September 2025 revision; the 2026 switch to a 10% ad-valorem "
  "export duty", "Government decisions, as recorded in the auditor's reports and in press "
  "coverage", "2021 to 2026"),
 ("Egyptian nitrogen capacity", "About 7.2 to 7.3 million tonnes a year across the named "
  "producers", "Industry analysis", "Mar 2026"),
 ("Local fertilizer prices", f"Subsidised supply at about EGP {V('subsidised_price'):,.0f} a "
  f"tonne against open-market sacks at EGP {V('local_sack_price_low'):,.0f} to "
  f"{V('local_sack_price_high'):,.0f} per 50 kilograms", "Egyptian press", "2025-2026"),
 ("Project contractor and scope", "The engineering consortium, the contract value and the "
  "settlement history of the earlier project", "Contractor announcements, corroborated "
  "against the notes to the audited statements", "2022 to 2025"),
 ("Carbon border adjustment", "European reporting obligations reaching Egyptian nitrogen "
  "exports from 2026", "Auditor's report; European policy", "2025-2026"),
]
for c_ in CTX:
    rows.append(list(c_))
table(rows, [1.7, 4.2, 2.3, 1.0], size=8.2)

# ---------------------------------------------------------- 4. cross-checks --
H1("4.  Cross-checks, used to test and never to build")
rows = [["Cross-check", "What it was used for", "Why it is not a build source"]]
rows += [
 ["Financial data aggregators", "Confirming the ticker, the fiscal-year convention and the "
  "order of magnitude of market capitalisation", "Aggregators restate and normalise; the "
  "company's own statements are the only permitted source for its reported figures"],
 ["Press reporting of half-year profit", "A first sanity check on the interim result before "
  "the filing itself was read", "Superseded entirely by the reviewed statements"],
 ["Peer capacity and location", "Framing the freight disadvantage of an inland plant",
  "Peer data informs comparison only, never this company's own history"],
]
table(rows, [2.0, 3.7, 4.1], size=8.2)

# --------------------------------------------------------- 5. access notes ---
H1("5.  Access notes, including what could not be obtained")
for a in SW['primary_access']:
    P(("REACHED — " if a['reachable'] else "NOT REACHED — ") + a['url'], bold=True, size=9)
    P(a['note'], size=8.8, color=GREY, space_after=7)
P("Two consequences follow and both are stated in the study itself. The share count comes "
  "from the audited statements rather than the exchange, and treasury-bill yields are "
  "secondary quotes rather than auction results. Neither substitution touches a figure the "
  "company reports about itself.", size=9)

# ----------------------------------------------------------- 6. how to check -

# ------------------------------------------------- 6. THE FULL INPUT REGISTER -
doc.add_page_break()
H1("6.  The full input register — every input, four fields, grouped by layer")
P("Value, source-and-construction, date and layer for every input the study uses. Nothing "
  "in the deliverables is an orphan number.", size=9.2, color=GREY)
for L in sorted(IRJ['layers']):
    rowsL = [r for r in IRJ['inputs'].values() if r['layer'] == L]
    if not rowsL:
        continue
    P(f"{L} — {IRJ['layers'][L]}  ({len(rowsL)} inputs)", bold=True, size=10.5, space_before=8)
    rows = [["Input", "Value", "Unit", "Date", "Source and construction"]]
    for rr in sorted(rowsL, key=lambda x: x['key']):
        v = rr['value']
        if isinstance(v, list):
            sval = ", ".join(f"{x:,.4g}" for x in v)
        elif isinstance(v, float):
            sval = f"{v:,.4f}".rstrip('0').rstrip('.') if abs(v) < 1000 else f"{v:,.0f}"
        else:
            sval = f"{v:,}" if isinstance(v, int) else str(v)
        note = rr['source'] + ((" — " + rr['note']) if rr.get('note') else "")
        rows.append([rr['key'], sval, rr['unit'], rr['date'], note])
    table(rows, [1.7, 1.35, 0.85, 0.8, 5.1], size=7.4)

# ----------------------------------------------------- judgements and negatives
doc.add_page_break()
H1("The judgements, and what would overturn each")
rows = [["Judgement", "What was decided", "What would overturn it"]]
JD = [("The valuation lens",
       "Free cash flow to the firm with a volume-times-price driver tree, because the "
       "revenue note and the balance sheet show a single operating business with no lending, "
       "rental or fee leg",
       "A disclosed second segment of any size, or a change in the asset mix that made the "
       "balance sheet something other than a plant"),
      ("The contested judgement — the capital programme",
       "Computed BOTH ways and published side by side rather than averaged, because the two "
       f"differ by EGP {LN['contested']['gap']:,.2f} a share",
       "A dated commissioning schedule with a disclosed nameplate, which would collapse the "
       "two cases toward one"),
      ("The gas share of the materials line",
       f"Three quarters of a single disclosed materials line, implying "
       f"{V('gas_m3_per_t_ammonia_modelled'):,.0f} cubic metres a "
       "tonne of ammonia — inside the auditor's own disclosed range",
       "Any disclosure splitting that line, or a stated consumption rate"),
      ("The new complex's nameplate",
       "Derived from the ammonia design plate less urea's draw at its own plate",
       "A filing stating the capacity, which no filing currently does"),
      ("The cost of capital basis",
       "The sovereign's own yield less its own default spread, with a country-loaded premium "
       "added back once; both premium bases published",
       "Evidence that Egyptian equity risk clears below its sovereign, which would revalue "
       "every Egyptian equity and not only this one"),
      ("The dividend assumption",
       "Zero, sourced from two consecutive appropriation statements rather than assumed",
       "A distribution proposed in any future appropriation statement"),
      ("The terminal growth rate",
       "The central bank's medium-term inflation target — nominal maintenance growth with no "
       "real growth assumed",
       "A demonstrated ability to grow volumes, which needs gas the country does not "
       "currently have")]
for j in JD:
    rows.append(list(j))
table(rows, [1.7, 4.0, 4.1], size=8.0)

H1("Negative results — what was searched for and not found")
rows = [["What was sought", "Where", "Outcome"]]
NEG = [("An investor presentation or earnings-call transcript",
        "The company's investor-relations portal and open search",
        "None exists. The company publishes filings and disclosure reports only, so the "
        "volume, price and utilisation data such material normally carries was taken "
        "instead from the statutory auditor's own tables"),
       ("A stated nameplate capacity for the new complex",
        "All four audited years and all three interim filings",
        "Not disclosed. The capacity used is derived and flagged as derived throughout"),
       ("A maintenance capital-expenditure guide or investment plan beyond the project",
        "Board reports inside the filings, and open search",
        "None found. Maintenance capital expenditure is set at a mature-plant standard and "
        "sensitised"),
       ("Separate audit reports for FY2023, FY2024 and FY2025",
        "The portal's audit-reports section",
        "That section holds only audit-committee reports and its newest item covers the "
        "period to 31 December 2021. The statutory auditor's reports for the years used are "
        "bound inside the annual statements themselves"),
       ("Official treasury-bill auction results",
        "The central bank's auction pages",
        "Blocked to automated access. Secondary market quotes are carried instead, used for "
        "context only, and no valuation figure depends on them"),
       ("The exchange's own stated share count",
        "The exchange company page for the listing",
        "Behind a bot challenge that refused every automated read. The share count is taken "
        "from note 14 of the audited statements instead")]
for n in NEG:
    rows.append(list(n))
table(rows, [2.5, 2.6, 4.7], size=8.0)

H1("Where a widely quoted third-party figure disagreed with the filings")
rows = [["Figure", "What third parties showed", "What the filings show", "Which was used"],
        ["Shares outstanding",
         f"One widely visible market page reported {V('aggregator_share_count')/1e6:,.2f} million "
         f"shares, inconsistent with its own market-capitalisation figure by a factor of ten; "
         f"a stale exchange listing document showed {V('stale_listing_share_count'):,} shares",
         f"{V('shares_outstanding'):,} shares of EGP {V('par_value'):.0f} par, paid-in capital EGP "
         f"{V('paid_capital')*1e6:,.0f}, note 14",
         "The filing. The stale document predates the capital increase"],
        ["Fiscal-year convention",
         "Several sources present results on a calendar-year basis",
         "The financial year ends 30 June; every statement is dated accordingly",
         "The filing"],
        ["FY2023/24 profitability",
         f"Reported net profit of EGP {V('is_net_FY2324'):,.1f} million is widely quoted without "
         f"qualification",
         f"That figure includes EGP {V('oneoff_reval_FY2324'):,.1f} million of one-off "
         f"investment-property revaluation gain",
         "The underlying figure, used for every margin and return in the study"]]
table(rows, [1.5, 3.4, 3.4, 1.5], size=7.8)

H1("7.  How a reader can check this study")
P("Every historical figure in the study appears in the accompanying workbook on the sheets "
  "marked historical, carried exactly as issued. Every forecast figure is a formula on "
  "those sheets, driven by the assumptions sheet, and each assumption carries its source "
  "in the adjacent column. A reader who disagrees with an assumption can change it and "
  "watch the valuation move; the two grids that do not redraw are labelled on their own "
  "sheets, because each of their cells is a separate run of the whole model.", size=9.5)

doc.save('EGCH_Bibliography_01-09-2026.docx')
strip_stub_counts('EGCH_Bibliography_01-09-2026.docx')
print("wrote EGCH_Bibliography_01-09-2026.docx")
