"""Every percentage and multiple that appears in PROSE in the delivered documents must be a
number the model computed. [Added 1 September 2026 after an audit found seven typed figures
the computed numbers beside them contradicted.]

Method: the numbers files (study_numbers, lenses, alternatives, experts, sensitivity grid,
the input register, the band record, the walk-forward bands) are flattened into every
plausible rendering of every value — 0/1/2 decimals as a percentage, as a plain number, as
a multiple — and each figure found in a paragraph or caption of either document must match
one of them. A figure with no computed counterpart FAILS the build; a false positive is
fixed HERE by widening the rendering set, never by deleting the figure from the study.
"""
import json, os, re, sys
from docx import Document
HERE = os.path.dirname(os.path.abspath(__file__)); os.chdir(HERE)
DOCS = ['EGCH_Valuation_Study_05-09-2026.docx', 'EGCH_Bibliography_05-09-2026.docx']


def walk(x, out):
    if isinstance(x, dict):
        for v in x.values(): walk(v, out)
    elif isinstance(x, (list, tuple)):
        for v in x: walk(v, out)
    elif isinstance(x, (int, float)) and not isinstance(x, bool):
        out.append(float(x))


vals = []
for f in ('study_numbers.json', 'lenses.json', 'alternatives.json', 'experts.json',
          'sensitivity_grid.json', 'band_record.json', 'strike_result.json',
          'technicals.json', 'beta_result.json', 'flat_rate_ladder.json'):
    if os.path.exists(f):
        walk(json.load(open(f)), vals)
# the input register contributes ONLY the values a builder actually consumes: a registered
# input nothing reads (L-018) cannot license a figure, and a correction note's own "the study
# had X" literal is never harvested — the numbers file, not the register's prose, is the
# counterpart. Consumption is read from the builders' own V('key') / _V('key') / src('key') calls.
_CONSUMED, _PREFIXES = set(), set()
for _f in sorted(os.listdir('.')):
    if _f.endswith('.py') and _f != 'prose_check.py':
        _src = open(_f).read()
        _CONSUMED |= set(re.findall(r"\b(?:_?V|src)\(\s*['\"]([A-Za-z0-9_]+)['\"]\s*\)", _src))
        # keys read through an f-string (V(f'bs_{key}_{t}')) are consumed by PREFIX
        _PREFIXES.update(re.findall(r"\b(?:_?V|src)\(\s*f['\"]([A-Za-z0-9_]+)\{", _src))
_IR = json.load(open('input_register.json'))['inputs']
_CONSUMED |= {k for k in _IR if any(k.startswith(p) for p in _PREFIXES if p)}
_UNCONSUMED = sorted(k for k in _IR if k not in _CONSUMED)
json.dump({"consumed": sorted(_CONSUMED & set(_IR)), "not_consumed": _UNCONSUMED,
           "prefixes": sorted(_PREFIXES)}, open('prose_check_result.json', 'w'), indent=1)
for k in _CONSUMED:
    if k in _IR:
        walk(_IR[k]['value'], vals)
walk(json.load(open(os.path.join('..', 'egch_walkforward', 'forward_ranges.json'))), vals)
walk(json.load(open(os.path.join('..', 'egch_walkforward', 'scores.json')))['drivers'], vals)
# derived renderings: a value v, its percentage forms, 1-v and v-1 (shares and changes), and its inverse
# THE TECHNICAL LADDER IS A DISTANCE FROM THE TECHNICAL READ'S OWN CLOSE, NOT FROM
# SPOT [corrected 03-Sep-2026]. These are two clocks and the protocol says so: a
# mid-cycle library arrival moves the technical read without re-striking the study,
# so the read is computed on the last session in the price library while the study
# is struck on the latest known price. They were the same number until the
# principal supplied closes for 3 September against a library ending 6 August; this
# checker then divided by 14.41 where the document divides by 13.98 and reported
# the document's own correct figure as unmatched. A checker that models what a
# document ought to do rather than reading what it does is checking a different
# document -- the [R-ENF-03] species, in its own small way. Spot is kept as a
# second divisor because some figures genuinely are quoted against it.
_tech = json.load(open('technicals.json')) if os.path.exists('technicals.json') else {}
_sn = json.load(open('study_numbers.json'))
_spot = _sn.get('spot')
_tclose = (_tech.get('close') if isinstance(_tech, dict) else None) or _spot
def _levels(x, out):
    if isinstance(x, dict):
        for v in x.values(): _levels(v, out)
    elif isinstance(x, (list, tuple)):
        for v in x: _levels(v, out)
    elif isinstance(x, (int, float)) and not isinstance(x, bool):
        for _den in (_tclose, _spot):
            if _den:
                out.append(x / _den - 1)
_lv = []; _levels(_tech, _lv); vals.extend(_lv)

# THE REVENUE-MIX SHARES ARE MODEL OUTPUTS AND THE SET DID NOT GENERATE THEM. Section 1.6
# prints each channel's share of that year's revenue; those are ratios of two committed
# figures and this checker had no route to them, so they matched only by COINCIDENCE
# against unrelated values in the set's own 1-v and 100*v expansions. Re-striking the price
# map on the live fit moved the coincidence away and the document's own correct 71.5%
# reported as unmatched.
#
# Widened rather than the figure deleted, which is what the rule says to do and is the
# right answer here twice over: the figures are real, and a figure that passes by
# coincidence is a figure this checker was not actually checking.
for _row in (_sn.get('cases', {}).get('base', {}).get('rows') or []):
    _tot = _row.get('revenue')
    if _tot:
        for _k in ('rev_exp', 'rev_sub', 'rev_free', 'rev_an', 'rev_other'):
            if _row.get(_k) is not None:
                vals.append(_row[_k] / _tot)
RENDER = set()
for v in vals:
    for x in (v, 1 - v, v - 1, -v, 100 * v, 100 * (1 - v), 100 * (v - 1), v / 100):
        for d in (0, 1, 2, 3):
            RENDER.add(round(x, d))
# structural constants a reader sees that are not model outputs
for x in (0, 5, 10, 15, 20, 25, 50, 80, 90, 95, 100, 22.5, 2.4, 3.95, 4.75, 9.5, 12.5, 6.5, 7.5, 0.5, 1.0, 1.5):
    for d in (0, 1, 2):
        RENDER.add(round(x, d))

NUM = re.compile(r"(?<![\w.])(-?\d{1,3}(?:,\d{3})*(?:\.\d+)?|-?\d+(?:\.\d+)?)\s*(per cent|%|x\b|times)")
problems, checked = [], 0
for f in DOCS:
    d = Document(f)
    texts = [p.text for p in d.paragraphs]
    for t in d.tables:
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if hdr == ["Input", "Value", "Unit", "Date", "Source and construction"]:
            continue      # the input register printed verbatim: the source of truth, not a claim about it
        for row in t.rows:
            for c in row.cells:
                texts.append(c.text)
    for txt in texts:
        for m in NUM.finditer(txt):
            raw = m.group(1).replace(',', '')
            v = float(raw)
            dec = len(raw.split('.')[1]) if '.' in raw else 0
            checked += 1
            if round(v, dec) not in RENDER and round(v, dec) not in {round(y, dec) for y in RENDER}:
                problems.append(f"{f}: '{m.group(0)}' in: …{txt[max(0, m.start()-60):m.end()+40]}…")
print(f"prose figures checked: {checked}; unmatched: {len(problems)}; register inputs consumed by a builder: "
      f"{len(_CONSUMED & set(_IR))} of {len(_IR)}; not consumed by any builder: {len(_UNCONSUMED)} (listed in prose_check_result.json)")
for p in problems:
    print("  !", p)
sys.exit(1 if problems else 0)
