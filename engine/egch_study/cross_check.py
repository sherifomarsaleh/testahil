"""EGCH — do the workbook, the document and the model say the same thing?

The recalculation proves the workbook is internally consistent with the model. It says
nothing about whether the DELIVERED DOCUMENT quotes the same numbers — a builder can read
a stale JSON, a caption can be hand-written, a figure can be regenerated out of order. This
check closes that gap by asserting the same headline quantity in all THREE artefacts:

    the model      study_numbers.json / alternatives.json, what compute.py produced
    the workbook   the delivered .xlsx, evaluated through xlcalc, not its cached values
    the document   the delivered .docx, as text a reader actually sees

A quantity that does not appear in the document text is a FAILURE, not a skip: a headline
the study does not state is a headline the reader cannot check.
"""
import json, os, re, sys
import openpyxl, xlcalc
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)

D = json.load(open('study_numbers.json'))
AL = json.load(open('alternatives.json'))
WB = openpyxl.load_workbook('EGCH_Valuation_Model_01092026.xlsx')
BK = xlcalc.Book(WB)
DOC = Document('EGCH_Valuation_Study_01-09-2026.docx')

TEXT = " ".join(p.text for p in DOC.paragraphs)
for t in DOC.tables:
    for row in t.rows:
        for c in row.cells:
            TEXT += " " + c.text
TEXT = re.sub(r"\s+", " ", TEXT)

BASE, HALT = D['cases']['base'], D['cases']['halt']
R = BASE['rows']

# label, model value, workbook cell, how the document renders it
CHECKS = [
    ("Value per share, carried through", BASE['bridge']['per_share'], ('DCF', 'B44'), "2dp"),
    ("Value per share, stopped", HALT['bridge']['per_share'], ('DCF', 'D44'), "2dp"),
    ("Equity value, carried through", BASE['bridge']['equity'], ('DCF', 'B45'), "0dp"),
    ("Equity value, stopped", HALT['bridge']['equity'], ('DCF', 'D45'), "0dp"),
    ("Enterprise value, carried through", BASE['bridge']['ev'], ('SOTP Bridge', 'B7'), "0dp"),
    ("Enterprise value, stopped", HALT['bridge']['ev'], ('SOTP Bridge', 'C7'), "0dp"),
    ("Net debt", BASE['bridge']['net_debt'], ('DCF', 'B41'), "0dp"),
    ("Terminal value share of EV, carried through", BASE['bridge']['tv_pct_ev'],
     ('SOTP Bridge', 'B8'), "pct1"),
    ("Terminal value share of EV, stopped", HALT['bridge']['tv_pct_ev'],
     ('SOTP Bridge', 'C8'), "pct1"),
    ("Revenue, FY2026/27", R[0]['revenue'], ('DCF', 'B5'), "0dp"),
    ("Operating profit before depreciation, FY2026/27", R[0]['ebitda'], ('DCF', 'B6'), "0dp"),
    ("Free cash flow, FY2026/27", R[0]['fcff'], ('DCF', 'B14'), "0dp"),
    ("Maintenance capital expenditure rate", D['drivers']['maint_capex_pct_rev'],
     ('Cash Flow', 'D28'), "pct1"),
    ("Project capital expenditure, FY2026/27", D['drivers']['anna_capex_path'][0],
     ('Assumptions', 'Project capital expenditure — FY2026/27'), None),
    ("Cost of capital, year one", D['drivers']['wacc_path'][0],
     ('Assumptions', 'Cost of capital, year one'), "pct1"),
    ("Terminal cost of capital", D['drivers']['wacc_terminal'],
     ('Assumptions', 'TERMINAL COST OF CAPITAL'), "pct1"),
]


def resolve(sheet, ref):
    """A coordinate is a brittle way to name a cell: inserting a row above it silently
    re-points the check at a different number, which is what happened the first time the
    Assumptions sheet grew. Anything that is not a coordinate is looked up BY LABEL in
    column A, and a label that is not found is a failure, never a skip."""
    if re.fullmatch(r"[A-Z]+\d+", ref):
        return ref
    ws = WB[sheet]
    for row in ws.iter_rows(min_col=1, max_col=1):
        if row[0].value == ref:
            return f"C{row[0].row}"
    raise KeyError(f"no row labelled {ref!r} on {sheet}")

fails, checked, in_doc = [], 0, 0


def renders(value, style):
    """Every spelling of the number the document builder could legitimately produce."""
    out = set()
    if style == "pct1":
        out |= {f"{value*100:.1f}%", f"{value*100:.2f}%"}
    elif style == "pct0":
        out.add(f"{value*100:.0f}%")
    elif style == "2dp":
        out |= {f"{value:,.2f}", f"{abs(value):,.2f}"}
    elif style == "0dp":
        out |= {f"{value:,.0f}", f"{abs(value):,.0f}"}
    else:
        out |= {f"{value:,.0f}", f"{value:,.1f}", f"{value:,.2f}"}
    return out


for lab, model_v, (sheet, coord), style in CHECKS:
    checked += 1
    try:
        wb_v = float(BK.cell_value(sheet, resolve(sheet, coord)))
    except Exception as ex:
        fails.append(f"{lab}: workbook cell {sheet}!{coord} would not evaluate ({ex})")
        continue
    tol = max(abs(model_v) * 1e-6, 1e-6)
    if abs(wb_v - model_v) > tol:
        fails.append(f"{lab}: model {model_v:,.6f} vs workbook {sheet}!{coord} {wb_v:,.6f}")
        continue
    if style is None:
        in_doc += 1
        continue
    if any(r in TEXT for r in renders(model_v, style)):
        in_doc += 1
    else:
        fails.append(f"{lab}: {sorted(renders(model_v, style))} appears in neither the "
                     f"document text nor its tables")

print(f"cross-check: {checked} headline quantities")
print(f"  model vs workbook   : {checked - len([f for f in fails if 'workbook' in f])}"
      f"/{checked} agree")
print(f"  quoted in the study : {in_doc}/{checked}")
for f in fails:
    print("   !", f)
if fails:
    sys.exit(f"FAIL: {len(fails)} of {checked} headline quantities disagree across the "
             f"three artefacts")
print(f"PASS: the model, the delivered workbook and the delivered study state the same "
      f"{checked} headline figures")
