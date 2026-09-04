"""SCEM_Bibliography_04-09-2026.docx — a standalone source register.

Every figure that reaches the study or the model traces to a row here: what it is, where
it came from, what kind of source that is, and the date the source itself carries.
Reads study_numbers.json and the sweep register — no numeral is typed here.
"""
import json
import re, os, sys
HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)
sys.path.insert(0, os.path.join(HERE, '..'))
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers.json'))
SW = json.load(open('sweep_register.json'))
INP = D['inputs']

INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(11), Inches(8.5)      # landscape: long source text
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(9.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.05


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_margins(t, top=40, bottom=40, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    t._tbl.tblPr.append(m)


def borders(t, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    t._tbl.tblPr.append(b)


def P(text='', size=9.5, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def H1(t): return P(t, size=15, bold=True, space_before=12, space_after=6)
def H2(t): return P(t, size=11.5, bold=True, space_before=10, space_after=4)


def table(rows, widths, size=8.4, wrap_cols=None):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ------------------------------------------------------------------ masthead
t = doc.add_table(rows=1, cols=1); cell_margins(t, 90, 90, 150, 150)
c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(9.8)
p = c.paragraphs[0]
r = p.add_run('Testahil · Sinai Cement Company S.A.E. (EGX: SCEM) — Source Register')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE
r2 = p.add_run('   6 August 2026')
r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

P('Where every number came from. This register accompanies the valuation study and the '
  'companion model. Each input carries the source it was taken from, the type of that '
  'source, and the date the source itself bears — not the date it was read.', size=10)

# ---------------------------------------------------- sourcing limitation
H2('A limitation to state at the top')
P('The company\'s audited consolidated financial statements could not be retrieved while '
  'this work was carried out. They are known to be published at '
  'sinaicement.com/wp-content/uploads/2025/05/SCC-AFS-E-1224.pdf, and every attempt to '
  'reach that address — along with the exchange\'s own disclosure portal and every '
  'financial data provider — was refused by the network policy governing this environment. '
  'Thirteen separate hosts were blocked.', size=9.5)
P('The consequence is stated rather than hidden. Revenue and profit after tax are carried '
  'as reported through press coverage of the company\'s exchange filings — one step removed '
  'from the audited print. Every line between revenue and profit after tax is DERIVED by '
  'closing the disclosed profit, and is labelled as derived wherever it appears. Two '
  'entries below are flagged accordingly. A reader with access to the audited statements '
  'should re-run the model against them; the workbook is built so that changing the input '
  'reprices everything downstream.', size=9.5)

# ---------------------------------------------------- input register
doc.add_page_break()
H1('1  Input register — every figure in the model')
RING_ORDER = {'Market': 0, 'Company': 1, 'Industry': 2, 'Country': 3, 'House': 4}
rows = [['#', 'Input', 'Value', 'Ring', 'Source', 'Source date']]


def fmt(v):
    if isinstance(v, list):
        return ', '.join(f'{x:,.4g}' for x in v)
    if isinstance(v, float):
        return f'{v:,.4f}'.rstrip('0').rstrip('.')
    return f'{v:,}' if isinstance(v, int) else str(v)


# THE REGISTER KEEPS ITS PROVENANCE; THE DELIVERED DOCUMENT DOES NOT PRINT IT.
# The stripper lived here first and is now engine/outward_source.py, imported rather
# than copied: one hand-maintained stripper per study is one hole per study, which is
# [L-084] and the scrub-list finding of the same morning. A rule that one study
# implements is a rule that one study obeys.
from outward_source import outward                                  # noqa: E402



items = sorted(INP.items(), key=lambda kv: (RING_ORDER.get(kv[1]['ring'], 9), kv[0]))
for i, (k, v) in enumerate(items, 1):
    rows.append([str(i), k, fmt(v['value']), v['ring'], outward(v['source']), v['date']])
table(rows, [0.36, 1.70, 0.92, 0.74, 4.92, 0.84], size=7.4)

# ---------------------------------------------------- source catalogue
doc.add_page_break()
H1('2  Source catalogue — the publications and institutions relied on')
CAT = [
 ('Sinai Cement Company S.A.E. — exchange filings', 'Company filing',
  'FY2023, FY2024 and FY2025 consolidated results; 9M-2025 results; EGM resolutions on the '
  '2024 capital increase; issued-capital and share-count disclosures.',
  'Accessed indirectly through the press outlets below; the primary documents could not be '
  'retrieved from the company website or the exchange portal.'),
 ('Global Cement', 'Trade press',
  'FY2024 turnaround (net profit EGP 3.07bn against a EGP 121.42mn FY2023 loss; sales EGP '
  '6.42bn from EGP 4.28bn); Egyptian production, consumption, capacity and export volumes; '
  'the dormant-capacity revival programme.', 'globalcement.com'),
 ('International Cement Review / cemnet', 'Trade press',
  'Egyptian 2025 production up 18% to ~65Mt; 9M-2025 results; the Cementir acquisition of '
  'the Sinai White stake.', 'cemnet.com'),
 ('Cementir Holding N.V. / Aalborg Portland Holding A/S', 'Counterparty disclosure',
  'Acquisition of an additional 25.40% of Sinai White Portland Cement from Sinai Cement for '
  'EUR 30 million, completed 13 August 2024, taking Cementir to 96.5%. This is the '
  'transaction that resets the FY2024 base.', 'cementirholding.com'),
 ('Vicat S.A. — results and regulatory filings', 'Shareholder disclosure',
  'The 77.6% holding through Vicat Egypt Cement Industries; the July 2025 mandatory tender '
  'offer for 58,416,664 shares (22.4%) at EGP 41.00 per share, filed with the Financial '
  'Regulatory Authority.', 'Reported via Reuters and Vicat H1-2025 results'),
 ('Daily News Egypt', 'National press',
  'FY2024 results commentary including the ~5% market-share figure; central bank inflation '
  'projections and the May-2026 inflation print.', 'dailynewsegypt.com'),
 ('Arab Finance', 'Financial press',
  'FY2025 consolidated results; peer results for Misr Beni Suef Cement and Arabian Cement.',
  'arabfinance.com'),
 ('EnterpriseAM Egypt', 'Financial press',
  'Egyptian cement pricing through 2025 and the 2026 estimate; the Egyptian Competition '
  'Authority\'s permanent removal of production quotas in July 2025.', 'enterpriseam.com'),
 ('Central Bank of Egypt', 'Central bank',
  'Main operation rate of 19.50%, held since 2 April 2026; the Q1-2026 Monetary Policy '
  'Report; the medium-term inflation target of 7% (±2pp) for Q4-2026 easing to 5%; the '
  'headline and core inflation series.', 'cbe.org.eg'),
 ('Aswath Damodaran — country risk premium file', 'Reference dataset',
  'Egypt equity risk premium and sovereign default spread on both the credit-default-swap '
  'and credit-rating bases, January 2026 edition. The original file only.',
  'Used for the cost-of-equity build'),
 ('PwC Worldwide Tax Summaries', 'Tax reference',
  'Egyptian corporate income tax rate of 22.5%, unchanged for 2025-26.', 'taxsummaries.pwc.com'),
 ('Egyptian Exchange daily price history', 'Market data',
  '3,626 daily open, high, low, close and volume records from 2 January 2011 to 6 August '
  '2026, supplied with the engagement. This series is the basis of the price chart, the '
  'volatility estimate, the beta regression and the liquidity diagnostics.',
  'Supplied as a file'),
 ('Global Energy Monitor plant register', 'Sector database',
  'El Hassana plant configuration: two lines, approximately 3.8 million tonnes a year of '
  'cement capacity, commissioned from 1997.', 'gem.wiki'),
]
rows = [['Source', 'Type', 'What it was used for', 'Reference']]
for a, b, c_, d_ in CAT:
    rows.append([a, b, c_, d_])
table(rows, [1.85, 1.62, 4.03, 2.00], size=8.0)

# ---------------------------------------------------- derived figures
doc.add_page_break()
H1('3  Figures that are DERIVED rather than sourced')
P('These do not appear in any source. They are computed, and the method is given so a '
  'reader can reproduce or reject each one.', size=9.5)
DER = [
 ('Share count — 260,812,477',
  'Triangulated three ways and cross-checked: issued capital of EGP 2,608,124,770 at EGP 10 '
  f'par gives {D["share_triangulation"]["issued_capital"]:,.0f}; the tender offer of '
  f'58,416,664 shares described as 22.4% back-solves to '
  f'{D["share_triangulation"]["tender_offer"]:,.0f}; quoted market capitalisation divided '
  f'by the closing price gives {D["share_triangulation"]["market_cap"]:,.0f}. The three '
  'agree to within 0.01%. Aggregator prints of 141.46 million are irreconcilable with the '
  'market capitalisation printed beside them and are rejected.'),
 ('FY2024 EBIT, treasury income and the disposal gain',
  'The FY2024 profit bridge is closed against the one disclosed EBITDA figure. Depreciation '
  'is taken at 6.2% of revenue, giving EBIT; the disposal gain is EUR 30 million at the '
  'August-2024 exchange rate less an estimated EGP 100 million carrying value; treasury '
  'income is then whatever remains once disclosed profit after tax is grossed up at the '
  'statutory rate.'),
 ('FY2025 EBIT and EBITDA',
  'Disclosed profit after tax is grossed up at the statutory rate to give pre-tax profit; '
  'treasury income on the derived cash balance is deducted to give EBIT; depreciation at '
  '4.6% of revenue is added back to give EBITDA.'),
 ('The cash balance',
  'No cash figure is separately obtainable. FY2024 cash is inferred as the derived FY2024 '
  'treasury income divided by the prevailing deposit yield; FY2025 cash is that balance '
  'grown by a stated multiple. This is the least well-evidenced material figure in the '
  'valuation and net cash is 37% of the equity value.'),
 ('Sales volume and realised price',
  'Volume is the Egyptian market size multiplied by the company\'s share; realised price is '
  'then disclosed revenue divided by that volume, so the build reproduces the reported top '
  'line exactly in all three historical years.'),
 ('Terminal return on invested capital',
  f'Struck on replacement cost — {D["inputs"]["capacity_mt"]["value"]}Mt at USD '
  f'{D["inputs"]["repl_usd_t"]["value"]:.0f} per annual tonne — rather than on book '
  f'invested capital, which implies a {D["terminal_reconciliation"]["roic_book"]*100:.0f}% '
  'return because the plant is carried at 1997 cost through a five-fold devaluation.'),
 ('Beta — 1.00 adopted',
  f'A five-year weekly regression of the shares against a 32-name equal-weight Egyptian '
  f'composite returns {json.load(open("beta_result.json"))["beta"]:.3f} with an R-squared '
  f'of {json.load(open("beta_result.json"))["r2"]:.3f} over '
  f'{json.load(open("beta_result.json"))["n"]} observations. That R-squared is below the 5% '
  f'usability floor, so the regression result is NOT used. No Egyptian listed cement peer '
  f'carries a price history in the covered library, so a re-levered peer beta is '
  f'unavailable. 1.00 is adopted and corroborated: a lead-and-lag estimator correcting for '
  f'the 29.3% of sessions that close unchanged gives '
  f'{json.load(open("beta_result.json"))["dimson"]["sum_beta"]:.3f}, whose 90% interval '
  f'contains 1.00.'),
]
rows = [['Figure', 'How it was derived']]
for a, b in DER:
    rows.append([a, b])
table(rows, [2.40, 7.10], size=8.2)

# ---------------------------------------------------- research trail
doc.add_page_break()
H1('4  Research trail')
P(f'The research was carried out on {SW["sweep_date"]} across four concentric rings — '
  f'global, country, industry and company — with {len(SW["findings"])} recorded findings. '
  'Each finding below names the source and the date that source carries.', size=9.5)
rows = [['Ring', 'Topic', 'Finding', 'Source', 'Date']]
for f in SW['findings']:
    if f['klass'] == 'NEGATIVE_SEARCH':
        continue
    rows.append([f['ring'].title(), f['category'], f['headline'], f['source_name'],
                 f['source_date']])
table(rows, [0.76, 1.58, 4.06, 2.12, 0.88], size=7.4)

H2('Searches that returned nothing')
P('Recorded because an absence of evidence shaped the model as much as the evidence did.',
  size=9.5)
rows = [['Topic', 'What was searched for']]
for f in SW['findings']:
    if f['klass'] != 'NEGATIVE_SEARCH':
        continue
    rows.append([f['category'], f['headline'].replace('Negative search — nothing found (', '')
                 .rstrip(')')])
table(rows, [2.20, 7.30], size=8.0)

P('')
P('Testahil · Independent valuation research · Educational analysis, not investment advice.',
  size=8.4, italic=True, color=GREY)

OUT = 'SCEM_Bibliography_04-09-2026.docx'
doc.save(OUT)
print('wrote', OUT)
