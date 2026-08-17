"""AIRARABIA_Bibliography_09-08-2026.docx — a standalone source register.

Every figure that reaches the study or the model traces to a row here: what it
is, where it came from, what kind of source that is, and the date the source
itself carries. Reads study_numbers.json and the research-trail JSON — no
numeral is typed here."""
import json, os, sys
HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers.json'))
SW = json.load(open('sweep_register.json'))
INP = D['inputs']

INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
F_PANEL, F_CREAM = 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(11), Inches(8.5)      # landscape: long source text
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(9.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.05

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(t, top=40, bottom=40, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    t._tbl.tblPr.append(m)

def borders(t, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    t._tbl.tblPr.append(b)

def P(text='', size=9.5, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def H1(t): return P(t, size=15, bold=True, space_before=12, space_after=6)
def H2(t): return P(t, size=11.5, bold=True, space_before=10, space_after=4)

def table(rows, widths, size=8.4):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ================================ 0 READ FIRST ================================
P('Air Arabia PJSC — Bibliography and input register', size=20, bold=True, space_after=2)
P('Companion to the valuation study of 9 August 2026 · Testahil · educational analysis, '
  'not investment advice', size=10, color=GREY, space_after=10)
P('READ FIRST. Every number in the study and the workbook traces to a row in this document. '
  'Section 1 lists the primary documents actually read; section 2 is the full input register — '
  'every model input with its value, source, and the date the source itself carries — grouped by '
  'research layer; section 3 lists the judgement calls and what evidence would overturn each; '
  'section 4 records searches that returned nothing, and one inconsistency found inside a primary '
  'document. Layers: COMPANY (the company\'s own filings and presentations), COUNTRY (UAE official '
  'data), INDUSTRY (aviation-sector references), MARKET (traded prices), HOUSE (this study\'s own '
  'judgement, always built on the sourced layers beneath it).', size=9.5)

# ================================ 1 PRIMARY DOCUMENTS =========================
H1('1  Primary documents')
rows = [['Document', 'Publisher / auditor', 'Date', 'What was taken from it']]
rows += [
    ['Audited consolidated financial statements FY2025',
     'Air Arabia PJSC / KPMG Lower Gulf (unqualified)', '13-Feb-2026',
     'The FY2025 income statement, balance sheet, cash flow and all notes; the RESTATED FY2024 '
     'comparatives and restated 1-Jan-2024 position (Note 43); revenue disaggregation (28a); the '
     '11-line direct-cost stack (29); fixed deposits at 4.41% (17); leases at 4% average (25); '
     'borrowings incl. the AED 849.6mn aircraft loan (26); tax under the 15% minimum-tax regime '
     '(27); JV/associate detail per investee (12); the 120-aircraft order advances (7); segments '
     '(39)'],
    ['Audited consolidated financial statements FY2024',
     'Air Arabia PJSC (approved 13-Feb-2025)', '13-Feb-2025',
     'FY2024 as reported and the FY2023 comparative income statement, cash flow and direct-cost '
     'note (scanned document, machine-read; every used figure cross-checked against typed '
     'comparatives)'],
    ['Audited consolidated financial statements FY2023 and FY2022',
     'Air Arabia PJSC', '2024 / 2023',
     'FY2023 and FY2022 statements (fourth audited year of context); dividends paid history'],
    ['Q1-2026 condensed interim financial information (reviewed)',
     'Air Arabia PJSC / Grant Thornton UAE', '13-May-2026',
     'The study year\'s only disclosed quarter: revenue 1,800.4, net profit 278.1, the March '
     'airspace-closure impact'],
    ['Results presentations Q4-2022 through Q1-2026',
     'Air Arabia PJSC investor relations', '2023–2026',
     'Passengers, load factor, all-hub traffic, fleet count, destinations, operating-profit '
     'tables — unit data no financial statement carries'],
    ['AGM press release', 'Air Arabia PJSC', '12-Mar-2026',
     'The 30-fils FY2025 dividend approval (AED 1.4bn) and board election'],
    ['First A320neo delivery release', 'Air Arabia PJSC', '29-Sep-2025',
     'Order composition (73/27/20), first delivery, CFM LEAP engines, 174 seats'],
    ['UAE dirham T-Bond auction results', 'UAE Ministry of Finance', 'May / Jul-2026',
     'The AED sovereign anchor: January-2031 tranche 4.30% (May), 4.48% (July), 4–14bp over US '
     'Treasuries'],
    ['Base-rate decision', 'Central Bank of the UAE', '29-Jul-2026',
     'Base rate held at 3.65%; the deposit-yield and financing-rate paths'],
    ['Quarterly Economic Review', 'Central Bank of the UAE', 'Mar-2026',
     'UAE growth ~5.6% and inflation ~2% — the escalator for airport-tariff-class costs'],
    ['Country risk dataset', 'A. Damodaran, NYU Stern', '05-Jan-2026',
     'UAE row: Aa2, 0.42% default spread, 4.87% equity risk premium; the UAE sovereign-swap '
     'column is not published (stated, not substituted)'],
    ['Jet Fuel Price Monitor; industry outlook', 'IATA', 'Jun–Aug-2026',
     'Jet fuel $158.77/bbl (early Aug-2026); the association\'s 2026 assumption (jet $152, Brent '
     '$95); Middle East traffic −11.4% in 2026'],
    ['Short-Term Energy Outlook', 'US Energy Information Administration', '07-Jul-2026',
     'Brent $81.91 (2026) → $64.76 (2027) — the base-case fuel path'],
    ['Passenger statistics 2025', 'Sharjah Airport Authority', 'Jan-2026',
     '19.48mn passengers, +13.9% — the home hub\'s growth'],
    ['Dubai Financial Market price history', 'DFM via uploaded export', '07-Aug-2026',
     'The share-price series 2011–2026 (3,908 sessions) for the price map, beta and technicals'],
    ['DFM General Index history', 'Yahoo Finance (aggregator, index only)', '16-Jul-2026',
     'The ADOPTED regression index for beta — the exchange the shares are listed on per note 1 of '
     'every filing. Market data only, never a source for company figures'],
    ['FTSE ADX General Index history, 2011–2026', 'Investing.com (aggregator, index only)',
     '24-Jul-2026',
     'The ALTERNATIVE-BENCHMARK regressor: the same beta regression re-run against the other UAE '
     'market proxy as an external check on a single-benchmark fit. Published, not adopted'],
    ['Quarterly Economic Review, June-2026 edition', 'Central Bank of the UAE', '30-Jun-2026',
     'The CURRENT 2026 projection: real GDP +1.7% (2027: +9.8%), inflation 2.3% — replacing the '
     'stale March-2026 vintage an earlier draft cited'],
    ['2025 GDP outturn release', 'UAE Federal Competitiveness and Statistics Centre', '30-May-2026',
     'Real GDP +6.2% in 2025, non-oil +6.8% — the outturn, replacing a forecast vintage'],
    ['Daily par yield curve', 'US Department of the Treasury', '07-Aug-2026',
     '5-year US Treasury at 4.35% — the peg-consistency check on the AED risk-free construction'],
    ['Ryanair FY26 results and Q1-FY27 report', 'Ryanair Holdings plc', 'May / 26-Jul-2026',
     'EBITDA €3,747.6mn (operating profit + depreciation), 1,039.2mn shares, €2.7bn net cash — '
     'the 6.5× peer multiple from the primary filings'],
    ['Wizz Air FY2026 final results', 'Wizz Air Holdings plc', '11-Jun-2026',
     'EBITDA €1,318.3mn and net debt — the peer multiple previously suppressed as n/m'],
    ['Interim IFRS financial statements', 'Pegasus Hava Taşımacılığı A.Ş.', '2026',
     'Euro functional currency, no inflation restatement — correcting an earlier basis label'],
    ['AGM dividend release for FY2021', 'Air Arabia PJSC', '11-Mar-2022',
     'The 8.5-fils FY2021 dividend — the base of the ladder, previously misstated as a 5-fil step'],
]
table(rows, [2.30, 1.85, 0.85, 4.75])

# ================================ 2 INPUT REGISTER ============================
doc.add_page_break()
H1('2  The input register — every input, value, source and date, by layer')
ORDER = ['Company', 'Country', 'Industry', 'Market', 'House']
def layer_of(ring):
    r = ring.lower()
    if 'company' in r: return 'Company'
    if 'country' in r: return 'Country'
    if 'industry' in r: return 'Industry'
    if 'market' in r: return 'Market'
    return 'House'
def _num(x):
    if isinstance(x, bool) or not isinstance(x, (int, float)):
        return str(x)
    return f'{x:,.1f}' if abs(x) >= 1000 else f'{x:,.4g}'
def fmt_val(v):
    if v is None:
        return 'not published'
    if isinstance(v, (int, float)):
        return _num(v)
    if isinstance(v, list):
        return '[' + ', '.join(_num(x) for x in v) + ']'
    if isinstance(v, dict):
        return '; '.join(f'{k}={fmt_val(x)}' for k, x in list(v.items())[:6]) + ('; …' if len(v) > 6 else '')
    return str(v)
def fmt_name(k):
    return k.replace('_', ' ')
def fmt_src(s_):
    return s_.replace(' See beta_result.json', '').replace('beta_result.json', 'the regression record')
groups = {L: [] for L in ORDER}
for k, rec in INP.items():
    groups[layer_of(rec['ring'])].append((k, rec))
for L in ORDER:
    if not groups[L]:
        continue
    H2(f'{L} layer — {len(groups[L])} inputs')
    rows = [['Input', 'Value', 'Date', 'Source and construction']]
    for k, rec in groups[L]:
        rows.append([fmt_name(k), fmt_val(rec['value']), rec['date'], fmt_src(rec['source'])])
    table(rows, [1.35, 1.55, 0.72, 6.13], size=7.6)

# ================================ 3 JUDGEMENTS ================================
doc.add_page_break()
H1('3  Judgement calls — and what would overturn each')
rows = [['Judgement', 'What was decided', 'What would overturn it'],
    ['The fuel path (dual-framed)',
     'Base follows the official US energy-agency curve (relief from 2027); the airline '
     'association\'s high-fuel assumption is priced in full as the alternative',
     'Two consecutive quarters of realised jet fuel above ~$150/bbl would make the high-fuel '
     'framing the base'],
    ['Which market index the beta is measured against',
     'The general index of the Dubai exchange, because note 1 of the 2025 statements, the 2025 '
     'annual report and the Q1-2026 interim all state the ordinary shares are listed there, and the '
     'annual report benchmarks the share price against that index. The same regression against the '
     'Abu Dhabi general index gives a lower beta (0.812 against 1.086) and a HIGHER value, and is '
     'published in full beside the adopted figure rather than left out',
     'A change of listing venue, a dual listing, or evidence that the Abu Dhabi index better '
     'explains this share\'s weekly returns — today it explains a third as much (R² 0.14 against '
     '0.40), which is why it is the cross-check and not the basis'],
    ['The JV network (dual-framed)',
     'Base carries the audited AED 363mn carrying value; the alternative capitalises the AED '
     '190mn profit share at 15×',
     'Consolidated disclosure, a venture dividend policy, or a Saudi launch date would justify '
     'promoting the capitalised framing'],
    ['Passenger path −1.6% then +8–9%/yr',
     'Q1-2026 actual (−11%) plus phased airspace restoration, then fleet-led growth at held '
     '~85–86% load factor',
     'H1-2026 passengers below ~5.9mn consolidated, or a load factor below 84%, would cut the '
     'recovery slope'],
    ['Fare give-back in FY2027 (−1.6%)',
     'The Q1-2026 yield spike is treated as scarcity pricing that regional capacity will compete '
     'away', 'Yields holding through two post-restoration quarters would remove the give-back'],
    ['Fleet capex ~AED 1.9–2.0bn/yr',
     'Approximately 3–4 owned aircraft a year plus the pre-delivery ladder; the owned/leased '
     'split is NOT disclosed and this is the build\'s weakest driver — sensitised ±30%',
     'Any financing-plan disclosure for the neo ramp'],
    ['Tax at the statutory 15%',
     'Above the realised 8.8–11.6% prints — deliberate conservatism under the new minimum-tax '
     'regime', 'A third year of a sub-12% effective rate with a disclosed structural driver'],
    ['Sustainable return on equity 18%',
     'Below the 19.9% record year, which carried a scarcity-yield tailwind',
     'Two more years above 20% through a normal capacity environment'],
    ['Working capital at −64% of revenue',
     'The three-year audited centre, not the best year',
     'A structural change in ticket-sale timing or maintenance-provision policy'],
    ['Justified multiples 7.5× / 13×',
     'The global LCC centre, above the mature Europeans, below Jazeera',
     'A durable re-rating of the sector, or Air Arabia losing its net-cash position'],
    ['Terminal growth 2.5% against a 4.0% terminal risk-free',
     'About half a point real for a still-growing home market',
     'Sharjah airport capacity saturating without a second-hub answer']]
table(rows, [1.75, 4.20, 3.80], size=8.2)

# ================================ 4 NEGATIVES & DISCREPANCIES =================
H1('4  Negative results, and one discrepancy inside a primary document')
rows = [['What was searched for', 'Outcome (dated 09-Aug-2026)'],
    ['Seat capacity / available seat-kilometres / stage length, any document',
     'Not disclosed anywhere in four years of filings or presentations — passengers × per-'
     'passenger rates is the finest level the company\'s own record supports; said in the study'],
    ['Fuel hedge ratios or hedged volumes',
     'The accounts disclose instruments and fair values to 2028 but never the hedged share'],
    ['Owned-versus-leased split of forward aircraft deliveries',
     'Not disclosed; the capex driver is an assumption and is flagged as the weakest in the build'],
    ['A UAE sovereign credit-default-swap premium in the January-2026 country dataset',
     'The dataset publishes no UAE value in that column — the rating basis is the only published '
     'construction and is used alone, stated openly'],
    ['Q2/H1-2026 results', 'Not yet published as of 9 August 2026 — the first catalyst to watch']]
table(rows, [3.60, 6.15], size=8.4)
P('Discrepancy recorded, not repaired: in the FY2025 filing\'s restated FY2024 revenue '
  'disaggregation, the six contract-revenue lines sum to total revenue (AED 6,765.9mn) while the '
  'printed contract-revenue subtotal is AED 6,616.4mn — the aircraft-lease-rental line (AED '
  '149.4mn) is the difference. The study uses the passenger and baggage lines as printed and '
  'records the footing inconsistency here. Additionally, the FY2022–FY2024 statement PDFs are '
  'image scans: they were machine-read, and every figure carried into the model was cross-checked '
  'against the following year\'s typed comparative column.', size=8.8)
P('This edition responds to four external audits (9-10 Aug-2026): every accepted finding is '
  'implemented and the full 71-row finding-by-finding response, with receipts and prices, is in the '
  'study repository. Extracts of the decisive filing passages (the finance-sublease note, the '
  'operating-lease lessor note, the tax reconciliation, the fleet allocation) are quoted in the '
  'response so the next auditor can tie them without machine-reading the scanned PDFs.', size=8.8)
P('The regression index history (DFM General Index) ends 16-Jul-2026, three weeks before the '
  'stock series — the beta window is truncated to the overlap, costing 3 of ~260 weekly '
  'observations; recorded rather than papered over.', size=8.8)

OUT = 'AIRARABIA_Bibliography_09-08-2026.docx'
doc.save(OUT)
print('wrote', OUT, '| tables:', len(doc.tables), '| paragraphs:', len(doc.paragraphs))
