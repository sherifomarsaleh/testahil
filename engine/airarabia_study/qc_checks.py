"""Automated QC checks over the delivered documents:
  (k)/(m) no internal-procedure vocabulary reaches the external reader
  (o)      table column widths fit the text block and no column is starved or bloated
  (h)/(n)  figure canvases are opaque and light (numbers legible on any page)
  general  no placeholder or unformatted values leaked into the documents
"""
import os, re, sys, json
from docx import Document
from docx.shared import Inches
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
STUDY = os.path.join(HERE, 'AIRARABIA_Valuation_Study_09-08-2026_public.docx')
BIB = os.path.join(HERE, 'AIRARABIA_Bibliography_09-08-2026.docx')


def _latest_xlsx():
    """The LATEST delivered workbook by PARSED DDMMYYYY date, never by string sort —
    '03092026' sorts below '09082026' as text and a text sort silently picks a superseded
    edition [L-067]. THE WORKBOOK IS A DELIVERED DOCUMENT AND EVERY SCRUB IN THE BOOK
    EXCLUDED IT [L-350]: a reader receives three files and this scrub named two."""
    c = []
    for f in os.listdir(HERE):
        m = re.search(r'Valuation_Model_(\d{2})(\d{2})(\d{4})_.*\.xlsx$', f)
        if m and not f.startswith('~$'):
            c.append((m.group(3) + m.group(2) + m.group(1), f))
    return os.path.join(HERE, sorted(c)[-1][1]) if c else None


XLSX = _latest_xlsx()
TEXT_WIDTH = 7.0            # 8.5in page less 0.75in margins each side
BIB_WIDTH = 9.8             # landscape bibliography uses 0.6in margins

BANNED = [
    r'\bstep\s*0(\.0)?\b', r'\bstep\s*2a\b', r'\bstep\s*[1-9]\b',
    r'\bfour[- ]ring\b', r'\bring\s+(classification|guide)\b', r'\bB/S/D/C\b',
    r'\binformation sweep\b', r'\bsweep register\b', r'\bQC gate\b', r'\bgate item\b',
    r'\bcalibration gate\b', r'\bmateriality gate\b', r'\bpromotion rule\b',
    r'\bstanding research protocol\b', r'\bresearch protocol\b',
    r'\bmc_v3\b', r'\bmarket_profiles\b', r'\bfitted_configs\b', r'\bstudy_numbers\b',
    r'\bcompute\.py\b', r'\bdata_quality\b', r'\bclean_ohlc\b', r'\bbacktest_v3\b',
    r'\bapply_breaks\b', r'\brobust_verdict\b', r'\bpanel_refresh\b', r'\bwacc_builder\b',
    r'\bLONO\b', r'\bCRPS\b', r'\bPIT\b', r'\bwidth_cal\b', r'\bkd_path\b',
    r'\bcalibration ledger\b', r'\bledger cohort\b',
    r'\bdevice A-\d\b', r'\bexpert persona library\b', r'\bpersona library\b',
    r'\bscale-normalis?zed\b', r'\bbootstrap block\b', r'\bengine reconciliation\b',
]

CASE_BANNED = [r'\bPARITY\b', r'\bBOUNDARY\b', r'\bFAIL\b(?! )']

def doc_text(path):
    # .xlsx branch: every STRING cell a reader sees. A numeric cell is a model output the
    # recalculation gate reconciles; a formula is skipped because data_only=False hands back
    # its text and a formula is not a sentence. Returns (None, text) — the width and figure
    # checks below are .docx-only and never receive a workbook.
    if path.lower().endswith(('.xlsx', '.xlsm')):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for v in row:
                    if isinstance(v, str) and not v.startswith('='):
                        parts.append(v)
        wb.close()
        return None, '\n'.join(parts)
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return d, '\n'.join(parts)

fails = []

# ---- (k)/(m) procedure-reference scrub --------------------------------------
print('=' * 74)
print('(k)+(m)  internal-procedure vocabulary scrub')
for label, path in [('study', STUDY), ('bibliography', BIB), ('workbook', XLSX)]:
    assert path and os.path.exists(path), f'{label}: delivered document not found'
    _, txt = doc_text(path)
    low = txt.lower()
    hits = []
    for pat in BANNED + CASE_BANNED:
        flags = 0 if pat in CASE_BANNED else re.I
        hay = txt if pat in CASE_BANNED else low
        for m in re.finditer(pat, hay, flags):
            ctx = txt[max(0, m.start() - 45):m.end() + 45].replace('\n', ' ')
            hits.append(f'{pat} -> "...{ctx}..."')
    print(f'  {label}: {len(hits)} hits')
    for h in hits:
        print('     ', h)
    if hits:
        fails.append(f'procedure vocabulary in {label}')

# ---- (o) table column widths --------------------------------------------------
print('=' * 74)
print('(o)  table column widths')
for label, path, maxw in [('study', STUDY, TEXT_WIDTH), ('bibliography', BIB, BIB_WIDTH)]:
    d = Document(path)
    over, narrow, wide = [], [], []
    for i, t in enumerate(d.tables):
        ws = [c.width.inches if c.width else None for c in t.columns]
        if any(w is None for w in ws):
            over.append(f'table {i}: unset width')
            continue
        tot = sum(ws)
        if tot > maxw + 0.02:
            over.append(f'table {i}: total {tot:.2f}in > {maxw:.2f}in')
        # a text column narrower than 0.55in will wrap mid-word for ordinary labels
        for j, w in enumerate(ws):
            if w < 0.55:
                narrow.append(f'table {i} col {j}: {w:.2f}in')
        # a single column taking more than 75% of a multi-column table is bloated
        if len(ws) > 2 and max(ws) / tot > 0.75:
            wide.append(f'table {i}: col takes {max(ws)/tot:.0%} of width')
    print(f'  {label}: {len(d.tables)} tables | over-wide {len(over)} | starved {len(narrow)} | '
          f'bloated {len(wide)}')
    for x in over + narrow + wide:
        print('     ', x)
    if over or narrow or wide:
        fails.append(f'column widths in {label}')

# ---- (h)/(n) figure canvases ----------------------------------------------------
print('=' * 74)
print('(h)+(n)  figure canvases opaque and light')
for fn in sorted(f for f in os.listdir(HERE) if f.endswith('.png')):
    im = Image.open(os.path.join(HERE, fn)).convert('RGBA')
    alpha = im.getchannel('A')
    transparent = alpha.getextrema()[0] < 255
    rgb = im.convert('RGB')
    px = list(rgb.getdata())
    corners = [rgb.getpixel(p) for p in [(1, 1), (rgb.width - 2, 1), (1, rgb.height - 2),
                                         (rgb.width - 2, rgb.height - 2)]]
    light = all(sum(c) / 3 > 200 for c in corners)
    ok = (not transparent) and light
    print(f'  {fn}: {"OK " if ok else "FAIL"} '
          f'(transparent={transparent}, corner luminance={[round(sum(c)/3) for c in corners]})')
    if not ok:
        fails.append(f'figure canvas {fn}')

# ---- general: no leaked placeholders --------------------------------------------
print('=' * 74)
print('general  leaked placeholders / unformatted values')
for label, path in [('study', STUDY), ('bibliography', BIB)]:
    d, txt = doc_text(path)
    bad = re.findall(r'\{[a-z_]+\}|\bnan\b|\d+e[+-]\d\d|\bTODO\b|\bXXX\b', txt)
    # a whole cell whose entire content is a Python repr is the real leakage mode
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                if c.text.strip() in ('None', 'nan', 'inf', '-inf', '[]', '{}', '0.0%'):
                    bad.append(f'cell="{c.text.strip()}"')
    print(f'  {label}: {len(bad)} hits {sorted(set(bad))[:8]}')
    if bad:
        fails.append(f'placeholders in {label}')

print('=' * 74)
if fails:
    print('QC CHECKS FAILED:', '; '.join(fails))
    sys.exit(1)
print('ALL AUTOMATED QC CHECKS PASSED')
