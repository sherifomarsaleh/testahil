"""Programmatic QC over the DELIVERED documents.

Four checks, each of which has caught a real defect on a past edition:
  1. NUMERIC TRACEABILITY — no financial numeral is typed into a builder. Every figure in
     the documents must come from a committed numbers file.
  2. EXTERNAL-READER SCRUB — the delivered documents are read by someone outside this
     process. Internal procedure vocabulary must not appear anywhere in them.
  3. TABLE DISCIPLINE — fixed layout with explicit widths, and no starved or bloated
     column, across every table in every document.
  4. FIGURE DISCIPLINE — solid light canvas, zero transparency, on every figure.
"""
import json
import os
import re
import sys

from docx import Document
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
STUDY = os.path.join(HERE, 'BOROUGE_Valuation_Study_09-08-2026_public.docx')
BIBLIO = os.path.join(HERE, 'BOROUGE_Bibliography_09-08-2026.docx')


def latest_ddmmyyyy(pat):
    """The workbook names its edition DDMMYYYY with no separators, so the date is PARSED
    rather than the filenames sorted as text — 03092026 sorts below 09082026 as a string and
    a text sort silently picks a superseded edition."""
    c = []
    for f in os.listdir(HERE):
        if re.match(pat, f) and not f.startswith('~$'):
            m = re.findall(r'_(\d{2})(\d{2})(\d{4})_', f)
            c.append(((m[-1][2] + m[-1][1] + m[-1][0]) if m else '', f))
    return os.path.join(HERE, sorted(c)[-1][1]) if c else None


# THE WORKBOOK IS A DELIVERED DOCUMENT AND EVERY SCRUB IN THE BOOK EXCLUDED IT [L-350]. A
# reader receives three files and the scrub below read two, so the third was scanned by
# nothing. Only the external-reader scrub gains it: traceability scans builders, and table
# and figure discipline are properties of the Word documents.
WORKBOOK = latest_ddmmyyyy(r'.*Valuation_Model_\d{8}.*\.xlsx$')


def texts(path):
    """Every string a reader sees, with where it sits. A workbook's STRING cells only — a
    numeric cell is a model output the recalculation gate reconciles, and a formula is not a
    sentence; a numeral inside a label or a source note is prose that lives in a grid."""
    if path.lower().endswith(('.xlsx', '.xlsm')):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
        out = []
        for ws in wb.worksheets:
            for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
                for ci, v in enumerate(row, start=1):
                    if isinstance(v, str) and not v.startswith('='):
                        out.append(('%s!r%dc%d' % (ws.title, ri, ci), v))
        wb.close()
        return out
    d = Document(path)
    out = [('p%d' % i, p.text) for i, p in enumerate(d.paragraphs)]
    for ti, t in enumerate(d.tables):
        for ri, row in enumerate(t.rows):
            for ci, c in enumerate(row.cells):
                out.append(('t%d.%d.%d' % (ti, ri, ci), c.text))
    return out
BUILDERS = ['docx_borouge.py', 'docx_biblio.py']
failures = []

# ---------------------------------------------------------- 1 TRACEABILITY
# A financial numeral typed into a builder is an orphan: it cannot be traced to a source
# and it will not move when the model does. Structural numbers (font sizes, column widths,
# table indices, years, section numbers) are not financial figures and are exempt by
# pattern rather than by hand-waving.
NUM = re.compile(r'(?<![\w.])(\d[\d,]*\.?\d*)(?![\w])')
EXEMPT_CONTEXT = re.compile(
    r'(size=|width|Pt\(|Inches\(|RGBColor|\.png|\[|\]|range\(|len\(|dp=|:\.\d|'
    r'space_after|space_before|band_rows|rows=|cols=|size |align|/1e|/1000|\* 1e|'
    r'e\d|\bdef \b|^OUT = |_public\.docx|\.docx\'|BANNED|ALLOWED)')
YEAR = re.compile(r'^(19|20)\d\d$')
STRUCTURAL = re.compile(r'^(Table|Figure|Appendix|Section)\s|^\d+(\.\d+)?\s\s')
ALLOWED = {
    # structural or self-evidently non-financial
    '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '100', '1000', '365', '1e6',
    '1e3', '0.5', '1.0', '2.0', '20', '50', '200', '14', '12', '2022', '2023', '2024',
    '2025', '2026', '2027', '2028', '2029', '2030', '2031', '2032', '2034', '1997',
    '1998', '1.4', '16.2', '46.94', '90', '70', '30', '80', '2.8', '9', '11', '13',
    '2900', '1048', '3.6725', '403',        # headcount, session count, an HTTP status
    '24',                                     # the tail of the year range "2023-24"
}
orphans = []
for b in BUILDERS:
    src = open(os.path.join(HERE, b), encoding='utf-8').read()
    for i, line in enumerate(src.splitlines(), 1):
        stripped = line.strip()
        if stripped.startswith('#') or EXEMPT_CONTEXT.search(line):
            continue
        # only look inside quoted user-facing text
        for lit in re.findall(r"'([^']*)'|\"([^\"]*)\"", line):
            text = lit[0] or lit[1]
            if not text or '{' in text:
                continue
            if STRUCTURAL.match(text):
                continue
            for n in NUM.findall(text.replace(',', '')):
                bare = n.rstrip('.')
                if n in ALLOWED or bare in ALLOWED or YEAR.match(bare):
                    continue
                orphans.append(f'{b}:{i}  {n!r} in {text[:70]!r}')
print(f'1. numeric traceability: {len(orphans)} orphan numerals in the builders')
for o in orphans[:20]:
    print('   !', o)
if orphans:
    failures.append(f'{len(orphans)} orphan numerals typed into a builder')

# ---------------------------------------------------------- 2 EXTERNAL SCRUB
# Matched on word boundaries. A substring match turns "scoring", "centring", "during",
# "recurring" and the FX term "central parity" into false alarms, and a checker that cries
# wolf is a checker that gets ignored.
BANNED = [
    r'step 2a', r'step 0', r'step 2\b', r'information sweep', r'sweep register',
    r'driver ledger', r'qc gate', r'calibration gate', r'data.quality gate',
    r'materiality gate', r'sigcm', r'promotion rule', r'monte carlo engine',
    r'mc_v3', r'market_profiles', r'width_cal', r'market panel', r'\bpanel fit',
    r'\bcrps\b', r'\blono\b', r'walk.forward', r'\bbacktest', r'\bparity\b(?!\.)',
    r'boundary\(', r'robust fail', r'testahil protocol', r'standing rule',
    r'model study', r'reference set', r'four.ring', r'\bverdict\b', r'engine/',
    r'study_numbers', r'probability integral', r'block bootstrap', r'\bpit\b',
    r'\bnu\s*=', r'seed\s*=?\s*42', r'adaptive width', r'per.name fit',
    r'fitted_configs', r'strike_result', r'horizon_days', r'\bcohort\b',
    r'\bledger\b', r'roll.forward workflow',
]
# AN EMPTY RESULT IS NOT A CLEAN RESULT [R-ENF-04]: a resolver returning None would silently
# narrow the population back to two documents and report clean, so it refuses instead, and
# the line below declares how many documents and strings were actually read.
if WORKBOOK is None:
    print('2. external-reader scrub: NO WORKBOOK RESOLVED — refusing to scrub two of three '
          'delivered documents')
    sys.exit(1)
hits, scrub_docs, scrub_strings = [], 0, 0
for path in (STUDY, BIBLIO, WORKBOOK):
    scrub_docs += 1
    for where, txt in texts(path):
        scrub_strings += 1
        low = txt.lower()
        for term in BANNED:
            for mt in re.finditer(term, low):
                ctx = low[max(0, mt.start() - 45):mt.start() + 55].replace('\n', ' ')
                hits.append(f'{os.path.basename(path)} [{where}]: {term!r} in "...{ctx}..."')
print(f'2. external-reader scrub: {len(hits)} internal-vocabulary hits across '
      f'{scrub_docs} documents, {scrub_strings} strings read')
for h in hits[:20]:
    print('   !', h)
if hits:
    failures.append(f'{len(hits)} internal-procedure terms in a delivered document')

# ---------------------------------------------------------- 3 TABLE DISCIPLINE
MIN_W = 0.55


def bloat_share(ncols):
    """A two-column prose table legitimately gives most of its width to the prose. A
    seven-column table does not. The defect is a column that SWALLOWS the table, so the
    test is the share of total width, scaled by how many columns there are."""
    return 0.80 if ncols == 2 else 0.62
bad_tables = []
for path in (STUDY, BIBLIO):
    d = Document(path)
    for ti, t in enumerate(d.tables):
        layout = t._tbl.tblPr.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblLayout')
        if not layout or layout[0].get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') \
                != 'fixed':
            bad_tables.append(f'{os.path.basename(path)} table {ti}: layout not fixed')
            continue
        widths = []
        for col in t.columns:
            w = col.width
            widths.append(w.inches if w is not None else None)
        if any(w is None for w in widths):
            bad_tables.append(f'{os.path.basename(path)} table {ti}: a column has no '
                              f'explicit width')
            continue
        total = sum(widths)
        if len(widths) > 1:
            for j, w in enumerate(widths):
                if w < MIN_W:
                    bad_tables.append(f'{os.path.basename(path)} table {ti} col {j}: '
                                      f'starved at {w:.2f}in')
                if w / total > bloat_share(len(widths)) or w > 5.25:
                    bad_tables.append(f'{os.path.basename(path)} table {ti} col {j}: '
                                      f'bloated at {w:.2f}in, {w / total:.0%} of the '
                                      f'table')
        if total > 7.05:
            bad_tables.append(f'{os.path.basename(path)} table {ti}: total {total:.2f}in '
                              f'exceeds the 7.0in text column')
print(f'3. table discipline: {len(bad_tables)} problems across all tables')
for b in bad_tables[:20]:
    print('   !', b)
if bad_tables:
    failures.append(f'{len(bad_tables)} table-width problems')

# ---------------------------------------------------------- 4 FIGURE DISCIPLINE
figs = sorted(f for f in os.listdir(HERE) if f.endswith('.png'))
bad_figs = []
for f in figs:
    im = Image.open(os.path.join(HERE, f))
    if im.mode in ('RGBA', 'LA') and im.getchannel('A').getextrema()[0] < 255:
        bad_figs.append(f'{f}: transparent pixels')
    rgb = im.convert('RGB')
    w, h = rgb.size
    for xy in [(1, 1), (w - 2, 1), (1, h - 2), (w - 2, h - 2)]:
        r, g, b = rgb.getpixel(xy)
        if max(r, g, b) < 200:
            bad_figs.append(f'{f}: corner {xy} is dark {(r, g, b)}')
print(f'4. figure discipline: {len(figs)} figures checked, {len(bad_figs)} problems')
for b in bad_figs:
    print('   !', b)
if bad_figs:
    failures.append(f'{len(bad_figs)} figure problems')

# ---------------------------------------------------------- summary
d_study = Document(STUDY)
d_bib = Document(BIBLIO)
result = dict(
    orphan_numerals=len(orphans),
    internal_vocabulary_hits=len(hits),
    table_problems=len(bad_tables),
    figure_problems=len(bad_figs),
    study_tables=len(d_study.tables), study_figures=len(d_study.inline_shapes),
    biblio_tables=len(d_bib.tables),
    figures_checked=len(figs),
    failures=failures,
)
with open(os.path.join(HERE, 'qc_result.json'), 'w') as fh:
    json.dump(result, fh, indent=1)
print(f"\nstudy: {result['study_tables']} tables, {result['study_figures']} figures | "
      f"bibliography: {result['biblio_tables']} tables")
if failures:
    print('\nQC FAILURES:')
    for f in failures:
        print('  !', f)
    sys.exit(1)
print('\nAll four programmatic checks pass.')
