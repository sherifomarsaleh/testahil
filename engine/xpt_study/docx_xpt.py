"""XPTUSD_Valuation_Study_20-07-2026_public.docx — python-docx builder, TMPV house
style (adapted from the STC/silver builders). Metals-class study: silver Combined
1-3-12M study is the structural reference; engine language updated to mc_v3."""
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers_xpt.json'))
S0 = json.load(open('step0_results.json'))
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
BRASS = RGBColor(0x89, 0x6F, 0x36); GOLD = RGBColor(0xC0, 0xA4, 0x5F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_PANEL2, F_CREAM = '1C3A36', 'EAF0EE', 'EFF3F1', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.top_margin, sec.bottom_margin = Inches(0.65), Inches(0.65)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.06

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(table, top=40, bottom=40, left=90, right=90):
    tblPr = table._tbl.tblPr
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tblPr.append(m)

def borders(table, color='C9D4D1', sz='4'):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    tblPr.append(b)

def P(text='', size=10.5, bold=False, italic=False, color=INK, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def rich(parts, size=10.5, space_after=6, space_before=0):
    p = doc.add_paragraph()
    for txt, kw in parts:
        r = p.add_run(txt)
        r.font.size = Pt(kw.get('size', size)); r.bold = kw.get('bold', False)
        r.italic = kw.get('italic', False); r.font.color.rgb = kw.get('color', INK)
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    return p

def H1(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = INK; r.font.size = Pt(15 if level == 1 else 12); r.bold = True
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    return p

def H2(text): return H1(text, level=2)

def caption(text):
    return P(text, size=8.7, italic=True, color=GREY, space_after=10)

def table(rows, widths, header=True, first_col_bold=False, size=9.3, header_fill=F_PANEL,
          align_right_from=1, band_rows=None):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, header_fill)
            if band_rows and i in band_rows:
                r.bold = True; shade(c, F_CREAM)
            if first_col_bold and j == 0 and i > 0:
                r.bold = True
            if j >= align_right_from and i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, width, caption_text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    p.paragraph_format.space_after = Pt(2)
    caption(caption_text)

def masthead():
    t = doc.add_table(rows=1, cols=1)
    cell_margins(t, 90, 90, 160, 160)
    c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.0)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

spot = D['meta']['spot']; MC = D['mc']; Z = D['zone']; A = D['anchors']; T = D['tech']; PR = D['price']; PA = D['path']; RA = D['ratio']

# ============================ COVER ============================
masthead()
P('Independent Valuation Study — Educational Analysis', size=10, color=GREY, space_after=2)
P('Platinum  (XPT/USD)', size=24, bold=True, space_after=2)
P('Fundamental analysis · Technical analysis · Monte Carlo simulation — one integrated read', size=11, color=BRASS, space_after=8)
P(f"Anchor: ${spot:,.2f}/oz (20 Jul 2026 close) · gold ${RA['gold_spot']:,.0f} → Pt/Au ratio ~{RA['now']:.2f}× (Au/Pt ~{1/RA['now']:.2f}×) · palladium ${RA['pd_spot']:,.0f} → Pt/Pd ~{RA['pt_pd']:.2f}× · "
  f"the industrial-and-precious metal, priced in US dollars per troy ounce · prices and probabilities computed 20 Jul 2026 from the attached daily history (04 Jan 2011 – 20 Jul 2026, {D['meta']['rows']:,} sessions after the data-quality gate) · "
  "primary fundamental lenses: the Pt/Au ratio and the analyst-consensus forward anchor, cross-checked against the structural supply/demand deficit, the mining cost curve and real-rate carry · "
  "this study combines the 1-month, 3-month and 12-month horizons in one document. The swing factors are the Fed real-rate path (10-year real yield ~2.31%, an ~18-year high), reverse Pt-for-Pd substitution now that platinum trades at a premium to palladium, and whether a fourth consecutive deficit outlasts the demand destruction the 2025–26 price shock set off.", size=9.3, color=GREY, space_after=10)

# ============================ HEADLINE ============================
H1('Headline')
P(f"The model's read: a wide, two-sided distribution centred on spot over one and three months and tilting modestly higher over twelve, as platinum consolidates a violent boom-and-bust inside a fair-value zone that spans roughly ${Z['lo']:,.0f}–{Z['hi']:,.0f}. "
  f"Platinum closed ${spot:,.2f} after the most dramatic sequence in its modern history: a ~128% run in 2025 from ${PA['end2024']:,.0f} at end-2024, an all-time closing high of ${PR['ath_close']:,.2f} on 23 January 2026 (attached-series basis; other vendors print within ~1.5%; intraday records near $2,920 followed days later), then a −17.6% single-day crash on Friday 30 January 2026 — the session of the surprise Warsh Fed-chair nomination, which markets read as a vote for Fed independence and higher-for-longer rates: the dollar rallied, gold fell ~10% and silver ~30% the same day, and the haven mania unwound — then a grinding slide to a 30 June low of ${PA['lo2026']:,.0f} as the Hormuz war premium built in the spring faded and the Fed repricing hardened. "
  f"Spot now sits {abs(PR['off_ath'])*100:.0f}% below the January closing record, −{abs(PR['ytd'])*100:.0f}% year-to-date, yet still +{PR['from_2024_end']*100:.0f}% above end-2024. "
  f"The fundamental lens — for a commodity, fair-value anchors rather than cash flows — puts a fair-value ZONE (what an ounce is worth now, undated) of ${Z['lo']:,.0f}–{Z['hi']:,.0f} (centre ≈ ${Z['centre']:,.0f}), with spot {abs(Z['spot_vs_centre'])*100:.1f}% below the centre — essentially at fair value. "
  f"The Pt/Au ratio, platinum's signature relative lens, sits at ~{RA['now']:.3f}× against gold's ${RA['gold_spot']:,.0f}: at the 5-year average ratio ({RA['mean_5y']:.3f}×) platinum is worth ~${A['ratio']['base']:,.0f} — above spot — while the 2-year average ({RA['mean_2y']:.3f}×) implies ~${RA['mean_2y']*RA['gold_spot']:,.0f}, about 4% below spot. "
  f"The published consensus splits by vintage: the LBMA January survey averaged $2,222 — set at the record and stale — while UBS's post-crash 29 June targets sit at $1,700 (end-2026) and $1,800 (mid-2027), just above spot. "
  f"The structural picture is genuinely tight — a fourth consecutive annual deficit (297 koz 2026f) has run above-ground stocks down to 1,747 koz, about eleven weeks of demand cover, the thinnest on record — but the 2025–26 price shock is destroying demand in real time (China jewellery fabrication −42% YoY in Q1-2026), and platinum's premium to palladium has, for the first time since 2017, flipped the substitution incentive that quietly added ~700 koz a year to automotive demand.", space_after=6)
P(f"A 50,000-path Monte Carlo sized by the carry-anchored YZ-HAR engine (Step-0-tested, PARITY) carries a forward volatility of ~{MC['t63']['ann_fwd_vol']*100:.0f}% — mean-reverting down from a trailing one-year {PR['vol252']*100:.0f}% that the crash inflated, wider than gold, ~two-thirds of silver's. Read the 50% band first: over one month the interquartile range is ${MC['t20']['p25']:,.0f}–{MC['t20']['p75']:,.0f}, over three months ${MC['t60']['p25']:,.0f}–{MC['t60']['p75']:,.0f}, over twelve ${MC['t252']['p25']:,.0f}–{MC['t252']['p75']:,.0f}; the predicted median sits at spot at 1 and 3 months (${MC['t20']['p50']:,.0f} / ${MC['t60']['p50']:,.0f}) and drifts to ~${MC['t252']['p50']:,.0f} at 12 months on pure cost-of-carry — the engine adds no factor drift, by design. "
  f"Technically the tape is stabilising, not reversed: price sits on the 20-day average but {abs(T['vs50'])*100:.0f}% below the 50-day and {abs(T['vs200'])*100:.0f}% below the 200-day, RSI {T['rsi14']:.0f} is neutral, MACD is negative with a positive but fading histogram, and the 50/200 stack is in a death-cross configuration. "
  f"Net: platinum has round-tripped a squeeze and now trades essentially at a fair value that a structurally tight physical market underwrites and a hawkish Fed and demand destruction cap. The right tail belongs to the stock-depletion squeeze re-igniting (eleven weeks of cover leaves no buffer), a 2027 Fed-easing re-pricing and the gold anchor; the left tail to reverse substitution, jewellery destruction compounding, and the stale half of the consensus being cut to price. This study estimates a distribution, not an action.", space_after=8)

# ============================ VALUATION SUMMARY ============================
H1('Valuation summary — every read at a glance')
P('One table for the four reads that follow — what the metal is worth (fundamental anchors), what the tape is doing (technical), where price could travel over one, three and twelve months (Monte Carlo), and how three independent expert methods land. Every row is developed in the sections and appendices below.')
P('Two kinds of number, kept distinct throughout. A fair value is what an ounce is worth today (the fundamental anchors and the expert estimates — a "fair value now"), and it does not carry a date. A prediction is where the price could actually be at a future moment (the Monte Carlo medians and bands — labelled "predicted at 1 / 3 / 12 months"). The two answer different questions: fair value asks "is it cheap or dear right now?"; the prediction asks "where might it trade by then?" Wherever a dollar figure appears below it is tagged as one or the other.')
table([
    ['Read', 'Level (US$/oz)', 'Kind', 'One-line verdict'],
    ['Fair-value zone (weighted anchors)', f"{Z['lo']:,.0f} – {Z['hi']:,.0f} (centre {Z['centre']:,.0f})", 'fair value now', 'spot essentially at the centre'],
    ['Pt/Au ratio lens (primary)', f"{A['ratio']['bear']:,.0f} – {A['ratio']['bull']:,.0f} (base {A['ratio']['base']:,.0f})", 'fair value now', 'platinum still cheap vs gold on the 5y window'],
    ['Analyst consensus (forward anchor)', f"{A['consensus']['bear']:,.0f} – {A['consensus']['bull']:,.0f} (base {A['consensus']['base']:,.0f})", 'fair value now', 'fresh UBS ~1,700–1,800; Jan survey 2,222 stale'],
    ['Supply/demand balance', f"{A['balance']['bear']:,.0f} – {A['balance']['bull']:,.0f} (base {A['balance']['base']:,.0f})", 'fair value now', '4th deficit, 11 weeks of cover — floor, not moonshot'],
    ['Cost curve (floor & incentive)', f"{A['cost']['bear']:,.0f} – {A['cost']['bull']:,.0f} (base {A['cost']['base']:,.0f})", 'fair value now', 'AISC ~1,006; greenfield incentive 2,300–2,500'],
    ['Technical tape', f"sup {T['sup'][0]:,.0f}/{T['sup'][1]:,.0f}/{T['sup'][2]:,.0f} · res {T['res'][0]:,.0f}/{T['res'][2]:,.0f}", 'tape', 'stabilising below trend; death-cross stack'],
    ['MC predicted median — T+20 / T+60 / T+252 (ledger cohorts + 12m)', f"{MC['t20']['p50']:,.0f} / {MC['t60']['p50']:,.0f} / {MC['t252']['p50']:,.0f}", 'prediction', 'carry-only drift; §3 tables quote the T+21/T+63 study horizons'],
    ['MC 90% band — 3 months (T+60 cohort)', f"{MC['t60']['p5']:,.0f} – {MC['t60']['p95']:,.0f}", 'prediction', f"width {MC['t60']['w90_pct_spot']*100:.0f}% of spot; T+63 study variant {MC['t63']['p5']:,.0f}–{MC['t63']['p95']:,.0f}"],
    ['Expert panel (three methods)', '1,500 / 1,800 / 1,670', 'fair value now', 'spread = the reverse-substitution question'],
], [2.5, 1.75, 0.95, 1.8], band_rows=[1])
P(f"Bottom line. The anchors agree that platinum is neither cheap nor dear at ${spot:,.0f}: the ratio lens says it is still inexpensive against gold, the balance says eleven weeks of cover puts a structural floor nearby, and the fresh half of the consensus sits just above spot while the stale half sits far above it. The three experts span $1,500–1,800 and each turns on one question — whether the deficit outlasts reverse substitution and demand destruction. The technical tape is the near-term dissenter, below trend and in a death-cross stack, which is why the one- and three-month distributions centre on spot with wide two-sided tails rather than already at the consensus, and only the twelve-month view drifts up — and then only by carry.", space_after=8)

# ============================ MARKET OVERVIEW ============================
H1('Market overview — platinum at a glance')
table([
    ['Item', 'Value', 'Item ', 'Value '],
    ['Spot (20 Jul 2026 close)', f"${spot:,.2f}/oz", '2026 supply (WPIC f)', '7,377 koz (+2%)'],
    ['All-time closing high', f"${PR['ath_close']:,.2f} (23 Jan 2026)", '2026 demand (WPIC f)', '7,674 koz (−9%)'],
    ['52-week range', f"${PR['lo52']:,.0f} – {PR['hi52']:,.0f}", '2026 balance', '−297 koz (4th consecutive deficit)'],
    ['Return 2025 / YTD 2026', f"+{PA['ret2025']*100:.0f}% / {PR['ytd']*100:+.0f}%", 'Above-ground stocks, end-26f', '1,747 koz ≈ 11 weeks of cover'],
    ['Realized vol 90d / 1y', f"{PR['vol90']*100:.0f}% / {PR['vol252']*100:.0f}%", 'Mine supply concentration', 'South Africa ~72% of 5,551 koz'],
    ['Pt/Au · Pt/Pd', f"{RA['now']:.3f}× · {RA['pt_pd']:.2f}×", 'Demand mix 2026f', 'auto 39% · jewellery 26% · industrial 29% · investment 7%'],
    ['NYMEX managed money (14 Jul)', '+8.3k contracts net long', 'Recycling 2026f', '1,826 koz (+9%)'],
], [1.9, 1.75, 1.9, 1.75], first_col_bold=False)
caption('Source: WPIC Platinum Quarterly Q1 2026 (18 May 2026, with Metals Focus), CFTC COT via metalcharts.org, exchange data, attached daily history. Values rounded. Demand mix shares computed on the 7,674 koz total.')

# ============================ 1 FUNDAMENTAL ============================
H1('1  Fundamental valuation')
P("How a commodity is valued. Platinum pays nothing and has no cash flows, so there is no sum-of-the-parts, no discounted-cash-flow of a business and no earnings multiple; the standard template's company items are replaced by their metal analogues (Appendix A). Every figure in this section is a fair value NOW — an estimate of what an ounce is worth today, undated — not a price forecast; the dated 1 / 3 / 12-month price predictions live in §3. "
  "But platinum is neither gold nor silver: it is one-third precious and two-thirds industrial, its mine supply is the most concentrated of any major metal (~72% South Africa), and it has a sibling metal — palladium — it substitutes for and against inside the same catalyst. Five anchors bound a fair-value ZONE: (1) the Pt/Au ratio, the primary relative lens; (2) the published analyst-consensus forward anchor; (3) the structural supply/demand balance; (4) the mining cost curve, which for platinum is unusually informative — spot spent 2018–2024 pinned against it; and (5) real-rate carry. The synthesis and football field are in §1.5; the swing factor — reverse substitution against the deficit — is dissected in §1.7 and the sensitivity grid in §1.9.")

H2('1.1  The Pt/Au ratio — the primary lens')
P(f"Platinum's most defensible relative reference is its price against gold, because the two share the precious-metal bid and their ratio defines regimes. The ratio sits at ~{RA['now']:.3f}× today (platinum ${spot:,.0f} ÷ gold ${RA['gold_spot']:,.0f}) — i.e. an ounce of gold buys ~{1/RA['now']:.2f} ounces of platinum. The full 2011–26 history averages {RA['mean_full']:.2f}× — but that average is a fossil: platinum traded at a premium to gold (ratio >1) until 2015, and the diesel scandal, the death of diesel demand and gold's own official-sector re-rating broke the old relationship for good. The modern, post-2016 regime averages {RA['mean_post2016']:.2f}×, the last five years {RA['mean_5y']:.3f}×, the last two {RA['mean_2y']:.3f}× — and today's {RA['now']:.3f}× sits at the {RA['pctile_now_post2016']*100:.0f}th percentile of the post-2016 distribution: even after a 128% year, platinum remains historically cheap against a gold price that quadrupled. The record extremes bound the lens: {RA['min_full']:.2f}× at the April 2025 trough (platinum's cheapest ever against gold, days before the squeeze began) and {RA['max_full']:.2f}× in 2011.")
table([
    ['Ratio scenario', 'Ratio', f"Implied Pt @ gold ${RA['gold_spot']:,.0f}", 'Reading'],
    ['2-year average (post-re-rating)', f"{RA['mean_2y']:.3f}×", f"${RA['mean_2y']*RA['gold_spot']:,.0f}", 'essentially spot — the recent regime'],
    ['Today', f"{RA['now']:.3f}×", f"${spot:,.0f}", '12th percentile of post-2016 — still cheap'],
    ['5-year average', f"{RA['mean_5y']:.3f}×", f"${RA['mean_5y']*RA['gold_spot']:,.0f}", 'the mean-reversion case, +14% vs spot'],
    ['Post-2016 average', f"{RA['mean_post2016']:.2f}×", f"${RA['mean_post2016']*RA['gold_spot']:,.0f}", 'the full modern regime — the bull anchor'],
], [2.15, 0.95, 1.7, 2.2], band_rows=[2])
P(f"Read. At the 2-year-average ratio platinum is worth ~${RA['mean_2y']*RA['gold_spot']:,.0f} — within rounding of spot — so on the recent window platinum is fairly valued, while the 5-year window says ~${A['ratio']['base']:,.0f} and the full post-2016 window ~${RA['mean_post2016']*RA['gold_spot']:,.0f}. The lens has the second dimension the others lack: it moves with gold. If gold reverts toward its own consensus zone (LBMA 2026 average forecast $4,742) and the ratio merely holds, platinum is dragged toward ~$1,920 with no change in platinum's own story. The honest caveat cuts the other way too: the post-2016 mean embeds years when platinum was structurally over-supplied; if the 2-year window is the true new regime, the ratio lens collapses onto spot and stops being an upside argument. A second relative check — the Pt/Pd cross at {RA['pt_pd']:.2f}× — is no longer a valuation lens but a demand driver, and it is doing the opposite of helping: see §1.7.")

H2('1.2  Analyst-consensus forward anchor — the market-implied fair value')
P("With no cash flows, the published professional consensus is a second market-implied fair value — but for platinum in July 2026 it splits into two vintages that disagree by $500, and the split itself is information. Targets:")
table([
    ['Source', 'Level (US$/oz)', 'Set when', 'Vintage problem'],
    ['LBMA 2026 survey average', '2,222', '20 Jan 2026 — at the record', 'set pre-crash; stale-high'],
    ['UBS — Sep & Dec 2026', '1,700', '29 Jun 2026 — spot $1,618', 'fresh, post-crash'],
    ['UBS — Mar & Jun 2027', '1,800', '29 Jun 2026', 'fresh, post-crash'],
    ['Valterra long-term planning', '2,300–2,500', 'Feb 2026 — for greenfield sanction', 'an incentive price, not a forecast'],
    ['WPIC 2026–30 average deficit', '~331 koz p.a.', 'May 2026', 'structure, not a price'],
], [2.05, 1.25, 1.95, 1.9])
P("Read. The fresh half of the consensus (UBS, $1,700–1,800) sits just above spot; the stale half ($2,222, written at the top) sits 38% above it. Platinum has fallen below the January sheet exactly as silver fell below its own — the same staleness crux: either price mean-reverts up toward targets written in the mania, or the targets are cut down to price. UBS has already cut once and flags the bear mechanism explicitly: weak investment demand could tip the 2026 balance into surplus. The base consensus anchor used here is $1,750 — the fresh vintage — with the stale $2,222 kept as the bull edge and a $1,500 'consensus-capitulation' level as the bear edge. The miners' own $2,300–2,500 planning range is deliberately excluded from the base: it is the price needed to build new mines, which is precisely what nobody is building — an incentive ceiling the market only pays when it must.")

H2('1.3  The supply/demand balance — the structural support')
P("Platinum has run a physical deficit for four consecutive years — 799 koz (2023), 1,033 koz (2024), 1,191 koz (2025), 297 koz forecast for 2026 — and the cumulative draw has taken above-ground stocks from 4,268 koz at end-2023 to a forecast 1,747 koz at end-2026: from 34 weeks of demand cover to about 11. That is the tightest platinum stock position on record, and it is the structural case beneath the price. Mine supply cannot respond: 72% comes from South African ore bodies whose real costs compound (electricity tariffs up >900% since 2008), the industry warns publicly of terminal decline, and the producers themselves are returning cash rather than sanctioning projects — only 2 of the 20 projects launched in the 2005–10 boom still operate. Recycling (+9% to 1,826 koz) is the only elastic leg. The 2026 twist is on the demand side: WPIC forecasts total demand −9% to 7,674 koz, because the price shock is doing what price shocks do — jewellery −12%, autocatalyst −2%, ETF and exchange stocks flipping to net sales — while bars and coins (+27%, led by China) and industrial (+9%) absorb part of the blow. The full balance is Appendix A.")
P("Read. The deficit is real but narrowing: 297 koz is ~4% of demand, a floor-builder, not a moonshot — the squeeze already happened, in 2025, and the price paid for it in the 2026 halving. What is NOT narrowing is the stock depletion: even a 297 koz deficit takes cover below three months, into territory where the 2025 lease-rate spikes showed the physical market seizing on any surprise. The honest read: the balance underwrites a structurally higher floor — the $1,350–1,500 region that held on the June retest, well above the cost curve — and it hands the price squeeze optionality (eleven weeks of cover leaves no buffer against a South African outage or an investment-demand surge), but it does not by itself justify the stale consensus's $2,222.")

H2('1.4  Cost floor and real-rate carry')
P("Two bounding anchors. The mining cost floor: S&P Global projects 2026 all-in sustaining costs of ~$1,006/oz (+7.7% YoY). Unlike silver (spot ~2× its cost curve) platinum spent 2018–2024 trading AT the curve — the industry restructured against it — so the floor is unusually well-tested: spot at ~1.6× AISC is the first sustained margin the industry has enjoyed in a decade, and a return to the curve (~$1,000) is the deep-bear boundary at which shafts close and supply self-corrects. The upper cost anchor is the incentive price: Valterra plans long-term around $2,300–2,500 — below that, no new greenfield ounces get sanctioned, which is how the market stays structurally short supply. Real-rate carry: platinum yields nothing, so the opportunity cost of holding it is the real rate. At ~2.31% real 10-year (18 Jul close — the highest in roughly 18 years, up from 1.99% in mid-May and 2.07% in early June) and a firm dollar the carry is a genuine, and strengthening, headwind — the hawkish 2026 repricing is a large part of why the squeeze deflated — and each 2025-style lease-rate spike is the physical market pricing that scarcity premium back in. Over twelve months it flips: easing is expected to resume in 2027, and every easing leg historically pulls investment flows back into the complex. Neither anchor pins a level; cost bounds the floor (~$1,000 deep-bear, $1,350–1,500 practical) and carry sets the near-term tilt.")

H2('1.5  Synthesis — five anchors and the fair-value zone')
P("No single anchor is decisive, so we weight them: the Pt/Au ratio and the consensus 30% each (market-implied and platinum-specific), the balance 20% and the cost curve 20% (the structural bounds), with real-rate carry as the timing tilt rather than a level. The weights are a stated house judgment, logged to the Fundamental Driver Ledger, and the silver-study precedent (35/25/20/12/8) lands within a few dollars of the same centre.")
figure('fig1_football.png', 6.9, f"Figure 1. Valuation football field. Bars span the bear–bull range per anchor; brass ticks mark the base; the ink line is spot ${spot:,.0f}; the shaded band is the ${Z['lo']:,.0f}–{Z['hi']:,.0f} weighted zone.")
P(f"The zone is ${Z['lo']:,.0f}–{Z['hi']:,.0f}, centre ~${Z['centre']:,.0f}, with spot ${spot:,.0f} sitting {abs(Z['spot_vs_centre'])*100:.1f}% below the centre — essentially at fair value. That is the honest fundamental statement: platinum has round-tripped a squeeze and now trades almost exactly where the anchors cross, held up by an eleven-week stock position and a still-cheap ratio to gold, held down by a hawkish Fed, demand destruction and a consensus that is half-stale. The width of the zone is not vagueness — the bear edge is a cost-curve world where reverse substitution wins, the bull edge is the stale consensus being right after all, and $800 of honest disagreement separates them.")

H2('1.6  Demand segments — a deeper look')
P("Platinum's demand is split across four legs with four different drivers — more balanced than silver (58% industrial) or gold (mostly monetary), which is why single-factor platinum stories are usually wrong. The scorecard, on WPIC's 2026 forecast:")
table([
    ['Segment', '2026f (koz)', 'Share', 'YoY', 'Driver / swing'],
    ['Automotive (catalysts)', '2,959', '39%', '−2%', 'ICE+hybrid volumes; Pt-for-Pd substitution now reversing — §1.7'],
    ['Jewellery', '1,958', '26%', '−12%', 'China −42% Q1; price shock + VAT-rebate removal (1 Nov 2025)'],
    ['Industrial (chem/glass/medical/H2…)', '2,238', '29%', '+9%', 'chemical 612, glass 377, medical 332, hydrogen 69 — the quiet grower'],
    ['Investment (bars/coins/ETF/exch)', '519', '7%', '−54%', 'bars +27% (China 13t vs <1t in 2019) vs ETF −100 and exchange −100'],
], [2.3, 0.95, 0.65, 0.65, 2.6])
P("Read. The investment leg is what took platinum to $2,772 and back — ETF and exchange-stock flows are the price-sensitive marginal buyer, and WPIC forecasts both negative for 2026 while physical bars boom in China. The industrial leg is the quiet structural story (+9%, with hydrogen at 69 koz still an option on the 2030s, not a 2026 driver — WPIC's >500 koz p.a. electrolyser scenario is a within-10-years case). Jewellery is the shock absorber that failed: the −42% Q1 China print is the single loudest demand-destruction datum in the complex. And automotive — the largest leg — is where the crux lives.")

H2('1.7  The crux — reverse substitution versus the structural deficit')
P(f"The swing factor, stated in real units. Roughly 700 koz a year of platinum's automotive demand exists because carmakers substituted platinum FOR palladium between 2021 and 2025, while palladium traded at a large premium. That premium is gone: platinum now trades at ~{RA['pt_pd']:.2f}× palladium (${spot:,.0f} vs ${RA['pd_spot']:,.0f}) — the first sustained platinum premium since 2017 — so the economic incentive that built the 700 koz has flipped sign. The mechanics are slow in both directions: substitution happens almost exclusively on new vehicle platforms (~15% of the market each year) and a platform locks its catalyst spec for ~7 years, so reversal is a 2028–30 erosion, not a 2026 cliff — WPIC itself forecast only 'gradual' reversal once palladium entered surplus. But the direction is now against platinum, and it compounds with jewellery destruction and BEV share growth. The arithmetic that matters: the 2026 deficit is 297 koz; reversing 150 koz a year of substitution — half the annual build rate — erases half the deficit; reversing the stock takes the market to surplus around the same time above-ground cover would otherwise hit zero. The whole platinum question is which side runs out of road first.")
table([
    ['Reverse-substitution pace (koz/yr)', 'Balance effect', 'AGS trajectory', 'Reading'],
    ['0 (lock-in holds through 2028)', 'deficit ~300 koz persists', 'cover < 8 weeks by end-27', 'squeeze redux — bull'],
    ['~150 (half the build rate)', 'deficit ~150 koz', 'cover ~9–10 weeks', 'tight but stable — base'],
    ['~300+ (full-rate reversal + jewellery slide)', 'balanced/surplus 2027', 'depletion stops', 'premium deflates to cost-plus — bear'],
], [2.35, 1.5, 1.5, 1.8], band_rows=[2])
P("The number to watch is not the price: it is palladium's discount (every dollar of Pt premium over Pd pays engineers to switch back), the WPIC quarterly substitution estimate, and Chinese jewellery fabrication. This is platinum's grams-per-solar-watt: not an abstract margin, but ounces per vehicle platform, decided in procurement meetings on a 7-year clock.")

H2('1.8  Macro — the Fed, the dollar, and geopolitics')
P("Three macro forces set the near-term tilt. The Fed and real rates: the dominant swing. The January mania was, at root, a Fed-independence trade — and the surprise Warsh nomination on 30 January reversed it in one session (dollar up, gold ~−10%, silver ~−30%, platinum −17.6%). The regime then turned durably hawkish: a hot CPI print (the Hormuz oil shock feeding through) collapsed cut odds and pushed markets to price hikes; the funds range has been held at 3.50–3.75% (midpoint 3.63%) since the 17 June statement, and the 10-year real yield reached ~2.31% by the 18 July close — an ~18-year high, a direct tax on a zero-yield metal. The consensus assumes easing resumes in 2027; if it does, the whole complex re-rates, platinum included; if it does not, the carry headwind persists. The dollar: platinum is priced in dollars and the dollar firmed all year on the rate repricing — a steady drag. Geopolitics and policy: the Strait of Hormuz war that began on 28 February (strikes, strait closure, Brent to $126) is what held platinum near $2,000 through Q1 — WPIC's own observation — and its spring de-escalation into the April toll-and-blockade stalemate is part of what the Q2 slide priced out; South Africa (72% of mine supply) remains the standing supply-shock risk; and China keeps hardening the demand floor — WPIC reports PGMs as fundamentally key to the 15th Five-Year Plan (AI, hydrogen; an industry-council characterization — the official strategic-minerals catalogue is undisclosed and its reconstructed lists carry no PGMs), the GFEX launch brought published stockpiles and sponge delivery, and the bar-investment market grew from under a tonne in 2019 to ~13 tonnes in 2025. The structural bid is migrating east, and it now has an exchange of its own.")

H2('1.9  Sensitivity — fair value versus gold and the ratio')
P("Because platinum's primary lens is two-dimensional, the sensitivity grid is platinum fair value as a function of the gold price and the Pt/Au ratio — the two variables that between them carry the relative-value case.")
figure('fig2_grid.png', 6.4, f"Figure 2. Platinum fair value = gold × Pt/Au ratio (US$/oz). Columns are gold scenarios; rows are ratio scenarios. Today's cell (gold ~${RA['gold_spot']:,.0f}, ratio {RA['now']:.3f}×) sits at spot. The upside is a joint move: gold higher and the ratio re-rating toward its 5-year mean.")
P(f"Read. The grid makes the cases concrete. Bull: gold at $4,800 with the ratio back at ~0.45× → ~$2,160. Bear: gold at $3,600 with the ratio stuck at the 2-year 0.36× → ~$1,296. The base — gold near ${RA['gold_spot']:,.0f} and the ratio between its 2- and 5-year means — brackets spot from ${RA['mean_2y']*RA['gold_spot']:,.0f} to ${RA['mean_5y']*RA['gold_spot']:,.0f}. Note the asymmetry the grid hides: the ratio's reversion path depends on the crux in §1.7 — reverse substitution is precisely the mechanism that would keep the ratio pinned at the 2-year lows — so the two dimensions of the grid are not independent, and reading the top-right corner as a free option overstates it.")

# ============================ 2 TECHNICAL ============================
H1('2  Technical and price structure')
P(f"Computed from the attached XPT/USD daily history (04 Jan 2011 – 20 Jul 2026, {D['meta']['rows']:,} sessions after cleaning).")
table([
    ['Indicator', 'Reading', 'Signal'],
    ['Close', f"${spot:,.2f}", '—'],
    ['SMA 20 / 50 / 200', f"{T['sma20']:,.0f} / {T['sma50']:,.0f} / {T['sma200']:,.0f}", f"price {T['vs20']*100:+.1f}% / {T['vs50']*100:+.1f}% / {T['vs200']*100:+.1f}% vs each"],
    ['50/200 stack', f"{T['sma50']:,.0f} < {T['sma200']:,.0f}", 'death-cross configuration, both falling'],
    ['RSI(14)', f"{T['rsi14']:.1f}", 'neutral — recovered from the June oversold'],
    ['MACD (12,26,9)', f"{T['macd']:,.0f} / signal {T['macd_sig']:,.0f}", f"negative, histogram +{T['macd_hist']:,.0f} but fading (was +17 five sessions ago)"],
    ['30-day average daily range', f"{PR['avg_true_range_pct_30d']*100:.1f}%", 'still a high-volatility tape'],
    ['52-week range', f"{PR['lo52']:,.0f} – {PR['hi52']:,.0f}", f"now {abs(PR['off_ath'])*100:.0f}% below the January closing record"],
    ['Support', f"{T['sup'][0]:,.0f} → {T['sup'][1]:,.0f} → {T['sup'][2]:,.0f}", '1 Jul swing low → round level → Feb-2021 breakout shelf'],
    ['Resistance', f"{T['res'][0]:,.0f} → {T['res'][2]:,.0f}", '50-day average → May-2026 recovery high'],
], [1.85, 1.9, 3.3])
P(f"Read. Near-term momentum is stabilising rather than reversing: price has reclaimed the 20-day average after holding the ${PA['lo2026']:,.0f} June low, RSI at {T['rsi14']:.0f} is neutral, and the MACD histogram is positive but shrinking — a bounce inside a downtrend, not yet a trend change. The larger structure is unambiguous: price is {abs(T['vs50'])*100:.0f}% below the 50-day and {abs(T['vs200'])*100:.0f}% below the 200-day, and the 50-day has crossed below the 200-day with both falling — the death-cross configuration that typically follows a halving. For a metal this volatile the tape says 'consolidating after a crash', which is exactly why the near-term distributions centre on spot with wide two-sided tails. The levels that matter: ${T['sup'][0]:,.0f} (the June–July floor), then ${T['sup'][1]:,.0f}, then the ${T['sup'][2]:,.0f} shelf — the 2021 breakout level whose loss would reopen the pre-squeeze range; above, the falling 50-day at ${T['res'][0]:,.0f}, then the ${T['res'][2]:,.0f} May recovery high, beyond which the squeeze narrative re-arms.")
figure('fig3_ma.png', 6.9, "Figure 3. Price and the 20 / 50 / 200-day moving-average stack (last ~2 years). The 2025 re-rating, the January 2026 record, the crash, and the consolidation below a rolling-over stack.")
figure('fig4_mom.png', 6.9, "Figure 4. Momentum — RSI(14) and MACD(12,26,9). RSI neutral after the June oversold; MACD negative with a positive but fading histogram.")

# ============================ 3 MONTE CARLO ============================
H1('3  Monte Carlo — a probabilistic price map')
P("Engine (v3, 'carry-anchored YZ-HAR-t'). Volatility width is sized by a pooled log-HAR cascade (variance lags 1/5/22) on a gap-aware Yang-Zhang variance proxy (overnight² + Rogers-Satchell), with the lognormal bias correction and a 0.8/0.2 log-space shrink toward the trailing-252d proxy — projecting the average daily variance over each forecast window, not trailing close-to-close vol. Drift is the cost-of-carry anchor ln(1+rf) − ln(1+q) with rf = 3.63% (the held Fed-funds midpoint) and q = 0 (a zero-yield store of value; the lease rate is a user's borrow cost, not a holder's yield) — no secular drift, no trend term, no signal (metals run carry-only), and the §1 fair-value gap is deliberately kept OUT of the drift so the fundamental and probabilistic lenses stay independent. One framing sensitivity, stated: if a positive convenience yield were imputed from the lease market instead of q = 0, the forward median would sit LOWER (each 100bp of q ≈ −1% on the 12-month median); q = 0 is the documented house null for a holder — the lease rate is a user's borrow cost, not a yield the vaulted ounce pays — and the 2025 squeeze rates it reflects have normalized. Shape and width come from platinum's own fitted configuration: a PROVISIONAL single-instrument fit (Gaussian tail; width_cal 0.853 — the MLE wanted an even narrower 0.79 and the house shrink-clip stopped it at 0.853) on 62 non-overlapping 60-day windows, 2012–2026 — flagged circular exactly as gold's first fit was, cross-checked out-of-sample in Appendix B. 50,000 paths, seed 42, anchored at spot; horizons T+21 (1 month), T+63 (3 months) and T+252 (12 months), with the ledger cohorts anchored at T+20/T+60 per house convention. Everything in this section is a dated prediction, as distinct from the undated fair-value-now anchors in §1.")
P(f"Read the 50% band first. The engine's forward width is ~{MC['t63']['ann_fwd_vol']*100:.0f}% annualised at three months — mean-reverting down from the crash-inflated trailing year ({PR['vol252']*100:.0f}%) but far above the pre-2025 norm — roughly two-thirds of silver's width and well above gold's. Drift is pure carry (+{MC['t63']['drift_log']*100:.1f}% log over 3 months, +{MC['t252']['drift_log']*100:.1f}% over 12), so the median barely moves: the distribution is genuinely two-sided, with {MC['t63']['p_below_spot']*100:.0f}% of three-month paths finishing below spot and {MC['t252']['p_below_spot']*100:.0f}% of twelve-month paths. Lead with the interquartile band: at T+63 the middle 50% of outcomes spans ${MC['t63']['p25']:,.0f}–{MC['t63']['p75']:,.0f}, and the 5–95% cone spans ${MC['t63']['p5']:,.0f}–{MC['t63']['p95']:,.0f}.")
table([
    ['Horizon', 'p5', 'p25', 'p50 (median)', 'p75', 'p95', 'mean'],
    ['T+21 (1 month)', f"{MC['t21']['p5']:,.2f}", f"{MC['t21']['p25']:,.2f}", f"{MC['t21']['p50']:,.2f}", f"{MC['t21']['p75']:,.2f}", f"{MC['t21']['p95']:,.2f}", f"{MC['t21']['mean']:,.2f}"],
    ['T+63 (3 months)', f"{MC['t63']['p5']:,.2f}", f"{MC['t63']['p25']:,.2f}", f"{MC['t63']['p50']:,.2f}", f"{MC['t63']['p75']:,.2f}", f"{MC['t63']['p95']:,.2f}", f"{MC['t63']['mean']:,.2f}"],
    ['T+252 (12 months)', f"{MC['t252']['p5']:,.2f}", f"{MC['t252']['p25']:,.2f}", f"{MC['t252']['p50']:,.2f}", f"{MC['t252']['p75']:,.2f}", f"{MC['t252']['p95']:,.2f}", f"{MC['t252']['mean']:,.2f}"],
], [1.5, 1.0, 1.0, 1.15, 1.0, 1.0, 1.0], size=8.8)
caption("Percentile map (US$/oz), 50,000 paths. The mean sits above the median at every horizon — the lognormal right skew. Ledger cohorts (T+20/T+60, same engine, same seed) are anchored separately for grading: T+20 median 1,612.84, T+60 median 1,623.02.")
P("The three horizons side by side — the combined read this document is built around: over one month the cone is $1,382–1,882 (90%), too wide for the fundamental anchors to bind; over three months $1,229–2,152, with the fair-value centre inside the interquartile band; over twelve months $961–2,897 — the 90% cone runs from the cost curve to the old record, which is an honest statement of how little a year of platinum can be predicted and why this study insists on distributions.")
figure('fig5_fan.png', 6.9, "Figure 5. Forward cone to T+252 (50k paths, carry-only drift, Gaussian tail, width_cal 0.853). Median (ink) drifts by carry only; the fan is the honest uncertainty. Dashed brass = fair-value centre.")
figure('fig6_dist63.png', 5.6, "Figure 6. Terminal price distribution at T+63 (3 months). Median at spot; right-skewed.")
figure('fig7_dist252.png', 5.6, "Figure 7. Terminal price distribution at T+252 (12 months) — the same shape stretched; p5 sits at the cost curve, p95 at the January record.")
P('Outcome zones (probability the terminal price lands in each band vs spot):')
z63, z252 = MC['t63']['zones'], MC['t252']['zones']
table([
    ['Zone vs spot', 'T+63', 'T+252'],
    ['below −20%', f"{z63['below_m20']*100:.0f}%", f"{z252['below_m20']*100:.0f}%"],
    ['−20% … −10%', f"{z63['m10_m20']*100:.0f}%", f"{z252['m10_m20']*100:.0f}%"],
    ['−10% … −5%', f"{z63['m5_m10']*100:.0f}%", f"{z252['m5_m10']*100:.0f}%"],
    ['±5%', f"{z63['pm5']*100:.0f}%", f"{z252['pm5']*100:.0f}%"],
    ['+5% … +10%', f"{z63['p5_p10']*100:.0f}%", f"{z252['p5_p10']*100:.0f}%"],
    ['+10% … +20%', f"{z63['p10_p20']*100:.0f}%", f"{z252['p10_p20']*100:.0f}%"],
    ['above +20%', f"{z63['above_p20']*100:.0f}%", f"{z252['above_p20']*100:.0f}%"],
], [1.7, 1.0, 1.0])
P(f"The modal three-month outcome is 'within ±5% of spot' but it holds only {z63['pm5']*100:.0f}% of the mass — a high-dispersion instrument where the tails carry the weight; by twelve months the largest single zone is 'above +20%' at {z252['above_p20']*100:.0f}% against {z252['below_m20']*100:.0f}% below −20%, the lognormal skew, not a directional call.")
P('Level-touch probabilities (path touches the level at least once within the window; running max for upside, running min for downside):')
t20, t60, t252 = MC['t20'], MC['t60'], MC['t252']
table([
    ['Level', 'By T+20 (1m)', 'By T+60 (3m)', 'By T+252 (12m)', 'Meaning'],
    ['$1,750 (June shelf)', f"{t20['touch_abs']['1750']*100:.0f}%", f"{t60['touch_abs']['1750']*100:.0f}%", f"{t252['touch_abs']['1750']*100:.0f}%", 'first resistance'],
    ['$1,990 (May recovery high)', f"{t20['touch_abs']['1990.5']*100:.0f}%", f"{t60['touch_abs']['1990.5']*100:.0f}%", f"{t252['touch_abs']['1990.5']*100:.0f}%", 'squeeze re-arms'],
    ['$2,222 (stale consensus)', f"{t20['touch_abs']['2222']*100:.0f}%", f"{t60['touch_abs']['2222']*100:.0f}%", f"{t252['touch_abs']['2222']*100:.0f}%", 'stale-vintage vindicated'],
    ['$2,925 (January record)', f"{t20['touch_abs']['2925']*100:.0f}%", f"{t60['touch_abs']['2925']*100:.0f}%", f"{t252['touch_abs']['2925']*100:.0f}%", 'full round trip'],
    ['$1,540 (June–July floor)', f"{t20['touch_abs']['1539.6']*100:.0f}%", f"{t60['touch_abs']['1539.6']*100:.0f}%", f"{t252['touch_abs']['1539.6']*100:.0f}%", 'the floor retested'],
    ['$1,348 (Feb-2021 shelf)', f"{t20['touch_abs']['1348.2']*100:.0f}%", f"{t60['touch_abs']['1348.2']*100:.0f}%", f"{t252['touch_abs']['1348.2']*100:.0f}%", 'breakout base gives way'],
    ['$1,006 (AISC floor)', f"{t20['touch_abs']['1006']*100:.0f}%", f"{t60['touch_abs']['1006']*100:.0f}%", f"{t252['touch_abs']['1006']*100:.0f}%", 'back to the cost curve'],
], [1.85, 1.05, 1.05, 1.1, 1.7], size=8.8)

# ============================ 4 COMPARISON ============================
H1('4  Comparison of the lenses, and a verdict')
P(f"The fundamental anchors, the technical tape and the probabilistic map tell a consistent story with different emphases. The fundamental zone (${Z['lo']:,.0f}–{Z['hi']:,.0f}, centre ${Z['centre']:,.0f}) says platinum is fairly valued at ${spot:,.0f}, underwritten by the tightest stock position on record and a still-cheap ratio to gold. The technical tape says 'consolidating a crash' — stabilising but below a rolling-over trend stack. The Monte Carlo reconciles them: a near-term median at spot (the tape's caution), a ~{MC['t63']['ann_fwd_vol']*100:.0f}%-wide cone (the metal's volatility), and a twelve-month drift that is carry only — the engine deliberately refuses to monetise the fair-value gap, because Step 0 shows it has no licence to (drift terms degrade skill; Appendix B). The three experts span $1,500–1,800 and disagree only on how fast reverse substitution erodes the deficit. Verdict: this is a distribution centred on a fair-value zone, not a directional call — the honest statement is a ${Z['lo']:,.0f}–{Z['hi']:,.0f} zone with spot essentially at its centre, a wide quarter-ahead cone, and the twelve-month right tail owned by stock-depletion squeeze mechanics and a 2027 easing cycle, the left tail by reverse substitution and consensus capitulation.")

# ============================ 5 CATALYSTS ============================
H1('5  Catalysts to watch')
table([
    ['Catalyst', 'When', 'Direction', 'Why it matters'],
    ['FOMC + CPI prints', '29–30 Jul FOMC; monthly CPI', 'both', 'the real-rate carry is the dominant macro driver; 2027 easing is the re-rating option'],
    ['WPIC Platinum Quarterly (Q2)', '~Sep 2026', 'both', 'the substitution estimate, the jewellery print, and any deficit revision'],
    ['Palladium discount', 'continuous', 'bearish if wide', f"Pt/Pd at {RA['pt_pd']:.2f}× pays for reverse substitution; the crux in real units"],
    ['Lease rates / EFP spreads', 'continuous', 'bullish if spiking', 'eleven weeks of cover means any physical surprise reprices instantly'],
    ['SA producer H2 prints (Valterra, Implats…)', 'Aug–Sep 2026', 'both', '72% of mine supply; guidance changes move the balance directly'],
    ['China: GFEX volumes, stockpile data, jewellery', 'monthly', 'both', 'the marginal physical buyer moved east; VAT-shock base effects lap in Nov'],
    ['ETF/COT positioning extremes', 'weekly', 'contrarian', 'net long +8.3k is modest; an extreme either way telegraphs the unwind'],
], [1.95, 1.35, 1.05, 2.75], size=8.8)

# ============================ 6 ZONES ============================
H1('6  Reading the probability zones')
P(f"The distribution is best read as zones, not a point. The ${MC['t60']['p25']:,.0f}–{MC['t60']['p75']:,.0f} region is the 'fair-value core' — the three-month interquartile band, containing the anchors' centre. Below ~${T['sup'][0]:,.0f} is the 'reverse-substitution / capitulation' zone — a {t60['touch_abs']['1539.6']*100:.0f}% touch by three months and {t252['touch_abs']['1539.6']*100:.0f}% by twelve — where the ${T['sup'][2]:,.0f} shelf and, in extremis, the ~$1,000 cost curve arrest the fall ({t252['touch_abs']['1006']*100:.0f}% twelve-month touch). Above ~${T['res'][2]:,.0f} is the 'squeeze-redux / consensus-reversion' zone — a {t60['touch_abs']['1990.5']*100:.0f}% touch by three months and {t252['touch_abs']['1990.5']*100:.0f}% by twelve — which needs a physical trigger against eleven weeks of cover or the Fed turning. The stale consensus at $2,222 carries a {t252['touch_abs']['2222']*100:.0f}% twelve-month touch; the January record a {t252['touch_abs']['2925']*100:.0f}% one. Lead with the interquartile band: the cone's width is platinum's genuine quarter-ahead uncertainty at ~{MC['t63']['ann_fwd_vol']*100:.0f}% volatility, not padding.")

# ============================ 7 CAVEATS ============================
H1('7  Caveats and what would change our mind')
P('This is a distribution, not a target, and several things would move it materially:', space_after=4)
for head, body in [
    ('The fit is provisional and circular. ', 'Platinum\'s (Gaussian, 0.853) configuration is a single-instrument self-fit, scored on its own history like gold\'s first fit. The de-circularized cross-check (Appendix B) is PARITY, and metals remain the weakest calibration in the system. A first live grade that lands badly triggers an immediate re-fit.'),
    ('Reverse substitution could be faster than the 7-year lock-in implies. ', 'The lock-in evidence comes from the forward direction; a sustained 30%+ platinum premium has no modern precedent, and procurement can accelerate what engineering permits. Watch the WPIC substitution line each quarter.'),
    ('The consensus staleness cuts both ways. ', 'We anchored on the fresh UBS vintage; if gold runs and the ratio re-rates, the January $2,222 sheet stops being stale and becomes right, and the zone centre moves up several hundred dollars.'),
    ('Demand destruction is measured with a lag. ', 'The −42% China jewellery print is Q1; if Q2 shows stabilisation at lower prices, the bear leg of the balance anchor is overdone — and vice versa.'),
    ('A South African supply event breaks every bear case instantly. ', 'With eleven weeks of cover, a single major outage (flood, power, closure) at 72% supply concentration is the fastest path to the right tail; the 2025 lease-rate seizures are the template.'),
    ('Data and vendor caveats. ', 'Spot cross-references (gold, palladium) are the Friday 17 Jul closes against the Monday 20 Jul platinum anchor — a one-session mismatch, immaterial at these widths but stated; on 20 Jul vendors printed gold ~$4,000–4,030, at which the ratio base is ~$1,845 and the zone centre ~$1,639 (+0.3%, inside rounding). All record levels (the $2,772.45 closing high, end-2024 $904, +128% in 2025) are attached-series figures; other vendors print within ~1–1.5%, and the anchor itself is the series close the ledger will grade against. The vendor series carries ~12% of sessions with Open marginally outside the High–Low range (a spot-quote splicing artifact) plus a single partial Sunday row (19 Jul); the engine\'s variance proxy guards degenerate bars, and the artifact is shared by the gold/silver library files.'),
]:
    rich([(head, dict(bold=True)), (body, dict())], space_after=4)

# ============================ APPENDIX A ============================
H1("Appendix A  The platinum balance — a commodity's financial statements")
P("A commodity has no income statement, balance sheet, cash-flow statement or per-share metrics, so the standard template's Appendix A items (A.1 income statement, A.2 balance sheet, A.3 cash flow) and the company DCF are a justified N/A here — the WACC build likewise does not apply (house Cost-of-Capital rule: metals are valued through the real-rates/flows/cost-curve lens, not a discount rate on a capital structure). Their metal analogue — the object that plays the role financial statements play for a firm — is the physical supply/demand balance: three years of history plus the current-year forecast, the ledger that decides whether the market draws down or rebuilds inventory. Figures from the WPIC Platinum Quarterly Q1 2026 (18 May 2026, research by Metals Focus).")
H2('A.1  Supply and demand balance (koz)')
table([
    ['Line (koz)', '2023', '2024', '2025', '2026f'],
    ['South Africa mine supply', '3,957', '4,133', '3,957', '4,005'],
    ['Zimbabwe', '507', '512', '516', '508'],
    ['Russia', '674', '677', '677', '646'],
    ['North America', '278', '265', '212', '201'],
    ['Other mining', '190', '191', '196', '192'],
    ['Total mine supply', '5,620', '5,787', '5,561', '5,551'],
    ['Recycling (auto / jewellery / industrial)', '1,515', '1,536', '1,679', '1,826'],
    ['TOTAL SUPPLY', '7,135', '7,323', '7,240', '7,377'],
    ['Automotive', '3,204', '3,108', '3,031', '2,959'],
    ['Jewellery', '1,849', '2,008', '2,214', '1,958'],
    ['Industrial (chem, glass, medical, H2…)', '2,491', '2,526', '2,049', '2,238'],
    ['  of which hydrogen', '22', '40', '65', '69'],
    ['Investment (bars/coins/ETF/exchange)', '388', '713', '1,136', '519'],
    ['TOTAL DEMAND', '7,933', '8,355', '8,431', '7,674'],
    ['BALANCE', '−799', '−1,033', '−1,191', '−297'],
    ['Above-ground stocks (year-end)', '4,268', '3,235', '2,044', '1,747'],
    ['Weeks of demand cover', '~34', '~20', '~12', '~11'],
], [2.75, 0.95, 0.95, 0.95, 0.95], band_rows=[8, 14, 15], size=8.8)
P("The balance is the metal's bottom line: four consecutive deficits totalling ~3.3 Moz have removed the equivalent of seven months of demand from above-ground stocks, taking cover from 34 weeks to 11 (WPIC-stated basis; a simple AGS ÷ same-year demand × 52 gives ~28/20/13/12, and summing the printed lines leaves a 1 koz print-rounding residual in 2023–24 — both stated, dual-framing rule). Mine supply barely moves (capital-starved, cost-inflated, geographically concentrated); recycling is the elastic leg; the demand swing in 2026 is the price shock itself — jewellery and ETF/exchange stocks give back what bars, coins and industrial gain. Supply minus demand equals the balance on every column. This table is the single most important 'financial statement' platinum has.")

# ============================ APPENDIX B ============================
H1('Appendix B  Step 0 — calibration backtest')
sf = S0['scores']['self_fit']; lg = S0['scores']['lono_gold_silver']; bw = S0['scores']['borrowed_live_metals']
sf5 = S0['scores']['self_fit_5y']; dg = S0['diag_self']
P("Before valuing the metal we test the forecasting engine on platinum's own history — under the CURRENT house gate: walk-forward, horizon 60 sessions, non-overlapping windows, each realised close scored against the full distribution the carry-anchored YZ-HAR-t engine would have produced using only prior data, graded on scale-normalized CRPS (crps/spot) against a CARRY-ANCHORED lognormal random-walk benchmark (same carry anchor both sides, so skill can never harvest the time-value of money). The sample is the full cleaned history: 62 windows, origins 05 Jan 2012 – 13 Feb 2026. Because platinum is a NEW single-instrument market, its (ν, width_cal) is fitted on its own residuals — a PROVISIONAL SELF-FIT, circular by construction, exactly as gold's first fit was flagged — so we also score two de-circularized configurations: the metal-family fit trained on gold+silver only (platinum fully out-of-sample), and the live production metals config.")
table([
    ['Configuration', 'ν / width_cal', 'CRPS skill (norm.)', '90% CI (block 2)', 'Robust verdict {2,3,4}'],
    ['XPT self-fit (ADOPTED, provisional)', 'Gaussian / 0.853', f"{sf['skill']:+.4f}", f"[{sf['ci_block2'][0]:+.3f}, {sf['ci_block2'][1]:+.3f}]", sf['verdict']],
    ['LONO — trained on gold+silver (OOS)', '20 / 1.035', f"{lg['skill']:+.4f}", f"[{lg['ci_block2'][0]:+.3f}, {lg['ci_block2'][1]:+.3f}]", lg['verdict']],
    ['Borrowed live METALS config', 'Gaussian / 1.0', f"{bw['skill']:+.4f}", f"[{bw['ci_block2'][0]:+.3f}, {bw['ci_block2'][1]:+.3f}]", bw['verdict']],
    ['Self-fit, last-5y windows only (n=20)', 'Gaussian / 0.853', f"{sf5['skill']:+.4f}", f"[{sf5['ci_block2'][0]:+.3f}, {sf5['ci_block2'][1]:+.3f}]", sf5['verdict']],
], [2.5, 1.3, 1.3, 1.35, 1.55], size=8.8)
P(f"Passes and blemishes, honestly stated. The verdict is PARITY under every configuration — the engine neither beats nor loses to the carry-anchored benchmark on this history (the market-panel CI is the standing gate; platinum does NOT arrive failing). Coverage under the adopted fit: 50% band {dg['cov50']*100:.0f}%, 80% band {dg['cov80']*100:.0f}%, 90% band {dg['cov90']*100:.1f}% against targets 50/80/90 — mildly over-covered at the tails even after the width narrowed to 0.853 (the MLE wanted 0.79; the house clip floor binds from BELOW — it prevented a narrower cone, a conservatism guard, not a risk-truncation device), the same over-coverage signature the system's other young fits show. PIT mean {dg['pit_mean']:.3f} — essentially centred; unlike silver there is no bull-skew blemish, because the carry-anchored benchmark absorbs the store-of-value drift. The reproduction check: rebuilding gold's live panel with this session's chain reproduced the production registry EXACTLY (67 windows, skill +0.0035, CI [−0.005, +0.013], PARITY) — the numbers above come from the production engine, not an approximation. Standing caveats: METALS REMAIN THE WEAKEST CALIBRATION IN THE SYSTEM — gold is a self-fit, silver borrows gold's, and platinum's own fit is provisional until the metals panel pools; the pooled 3-metal fit (ν≈20, width 0.965 on 148 windows) is the likely future configuration and its cone at T+63 is ~8% wider on the left tail (${MC['t63_lono']['p5']:,.0f} vs ${MC['t63']['p5']:,.0f} at p5 under the gold+silver-trained variant) — a stated sensitivity, not an adopted config.")
figure('figB1_calibration.png', 6.9, "Figure B1. Left: all 62 non-overlapping 60-day cones (90%) replayed over 2012–2026 with realised closes. Middle: PIT histogram under the adopted fit — roughly uniform, no U-shape. Right: interval coverage vs nominal (diagnostic only; band-containment is never the grade).")

# ============================ APPENDIX C ============================
H1('Appendix C  Peer set, sector structure, and risks')
P("Platinum has no 'comparable companies' — it is the underlying — but the investable platinum complex spans several vehicles with different exposures, useful for context and for reading what the equity market implies about the metal.")
table([
    ['Cohort', 'Names', 'Character'],
    ['SA primary producers', 'Valterra (ex-Amplats), Impala, Sibanye-Stillwater, Northam', 'levered to spot & the rand; AISC ~$1,000; returning cash, not building'],
    ['Zimbabwe / other', 'Zimplats, Tharisa', 'expansion optionality, political discount'],
    ['Russian by-product', 'Norilsk Nickel', '~12% of supply; sanctions-shadowed'],
    ['Physical ETFs', 'PPLT (~$1.8bn), PLTM', 'direct bullion; the 2026f −100 koz outflow line'],
    ['New venue', 'GFEX platinum futures (Guangzhou)', "China's first PGM derivatives; published stockpiles"],
    ['Platinum (this study)', '—', 'the metal itself — no operating leverage, no balance sheet, no counterparty'],
], [1.5, 2.6, 3.2], size=8.8)
P("Risks, in order of materiality: the Fed real-rate path and the dollar (the dominant macro driver of a zero-yield asset); reverse Pt-for-Pd substitution now that the price relationship has flipped (the demand-side structural risk); Chinese jewellery and bar demand (the marginal physical buyer, policy-sensitive); South African supply concentration (72% — both the biggest downside insurance and the biggest upside trigger); positioning and liquidity (platinum futures are ~10× less liquid than gold's — moves overshoot); and consensus staleness (half the published sheet was written at the record). The mitigants are structural: four consecutive deficits, eleven weeks of above-ground cover, a capital-starved supply base with a $2,300+ incentive price, and a tested cost floor at ~$1,000. Unlike the miners, the metal carries no operating leverage and no balance-sheet risk — the trade-off is that it pays nothing, which is precisely the carry cost Expert 1 presses.")

# ============================ APPENDIX D ============================
H1('Appendix D  The expert valuation panel')
P("Three independent methods, cast by instrument class from the house persona library: for a precious/industrial metal the panel pairs a real-rates macro valuer, an industrial supply/demand (cost-curve) valuer, and an investment-flows/ratio valuer — the same cast as the silver study, and for the same declared reason: platinum's industrial leg (two-thirds of demand) is its defining difference from gold, so the industrial seat replaces gold's official-sector seat. Each expert runs a genuinely different method, shows workings, and states a falsification condition. They are labelled Expert 1 / 2 / 3. Every expert fair value below is a fair value NOW, not a dated prediction.")

H2('D.1  Expert 1 — the real-rate / opportunity-cost valuer')
P("Worldview and tradition. Values platinum as a zero-yield asset whose fair price is set by the real interest rate and the dollar: holding an ounce forgoes the real return on cash, so the metal is 'worth it' only if expected appreciation compensates the carry. Platinum is only partly a monetary asset, so this expert applies the framework to the investment-driven premium over the industrial floor rather than to the whole price.")
P("When it works / fails. Works at rate-cycle turning points — the 2026 hawkish repricing crushed the squeeze premium on schedule. Fails when the physical market takes over (it is blind to an eleven-week stockpile) and says nothing about substitution.")
P("Standard workings. (1) Take the real 10-year (2.31% at the 18 Jul close — an ~18-year high) and the dollar as the discount-rate proxies. (2) Anchor a 'carry-fair' level consistent with a neutral real rate on the investment premium. (3) Adjust by platinum's real-rate beta (in the recent regime, roughly −$120–180/oz per 100bp of real yield on the premium component). (4) Overlay the dollar.")
P(f"Worked example. At 2.31% real and a firm dollar the carry-fair level is ~$1,500 — about 7% below spot: on this lens the market is still paying a residual squeeze premium the rate regime no longer funds. Easing discounted early (real back toward ~1.8%) lifts the claim toward ~$1,700; a push toward 2.7–2.8% real cuts it to ~$1,300. (The first edition quoted a stale 1.99% May print here — corrected; the section's direction is unchanged, the headwind is simply stronger.)")
P("Sensitivity and falsification. Each 50bp of real yield is worth roughly $70–90/oz on the premium. Falsification: platinum rising materially while real yields also rise — then the physical or substitution story is driving and this model is off-duty.")
P("Cross-examination. To Expert 2: 'Your deficit ran four years and the price still halved in a week when the Fed turned — the carry sets the mood the physical story trades inside.' To Expert 3: 'The Pt/Au ratio is just the relative rate-sensitivity of two metals; you price my variable twice.'")
P(f"Verdict. Fair value ~$1,500 (range $1,300–1,700). What the price implies: at ${spot:,.0f} the market is paying ~$100 over carry-fair — a residual squeeze/deficit premium this lens does not credit. Bull $1,700: 2027 easing discounted early, dollar softens. Base $1,500: the ~2.3% regime persists. Bear $1,300: sticky inflation, hikes delivered, real 10Y toward 2.7–2.8%.")

H2('D.2  Expert 2 — the industrial supply/demand (cost-curve) valuer')
P("Worldview and tradition. Values platinum as an industrial commodity in structural deficit: fair value is the marginal cost of clearing supply plus a scarcity premium that scales with the stock drawdown. Two-thirds of demand is industrial/automotive, mine supply is inelastic and 72% concentrated — this is the metal this method was built for.")
P("When it works / fails. Works over multi-year horizons and in genuine shortages (the 3.3 Moz cumulative draw is exactly that); 2018–24, when platinum pinned to its cost curve in surplus, is this model's success in the other direction. Fails on short horizons and cannot see demand destruction coming — the −42% China jewellery print and the substitution flip are precisely its blind spots.")
P("Standard workings. (1) Build the balance (Appendix A). (2) Marginal cost: AISC ~$1,006, +7.7%/yr. (3) Scarcity premium scaled to weeks-of-cover (11 and falling — the thinnest ever). (4) Cross-check against the incentive price ($2,300–2,500) at which new supply would eventually cap the premium.")
P("Worked example. At a 297 koz deficit and 11 weeks of cover, cost ($1,006) plus a ~$800 scarcity premium gives ~$1,800 — cover this thin has no precedent, and the 2025 lease-rate seizures showed what the physical market pays under stress. If reverse substitution and jewellery losses take the deficit to zero, the premium compresses toward ~$1,350; if cover keeps falling toward 8 weeks, $2,200 is defensible.")
P("Sensitivity and falsification. The swing is the substitution line: 150 koz/yr of reversal halves the deficit. Falsification: two consecutive annual surpluses, or a sustained rebuild in visible stocks (WPIC AGS line, NYMEX/GFEX inventories).")
P("Cross-examination. To Expert 1: 'Your carry model had platinum “fair” at $900 for six years while the industry starved — the cost curve and the stock line, not the real rate, decide where a physical shortage clears.' To Expert 3: 'The ratio is a heuristic; a catalyst plant does not care what gold costs.'")
P(f"Verdict. Fair value ~$1,800 (range $1,350–2,200), the top of the cluster. What the price implies: at ${spot:,.0f} the market pays cost plus ~$600 — it believes the deficit narrows but not that it ends. Bull $2,200: substitution lock-in holds, cover thins toward 8 weeks, any SA event ignites it. Base $1,800: deficit ~150–300 koz persists. Bear $1,350: full-rate reverse substitution plus jewellery destruction reach surplus by 2027.")

H2('D.3  Expert 3 — the investment-flows / Pt-Au ratio valuer')
P("Worldview and tradition. Values platinum relative to gold and through the flows that move a small market — the Pt/Au ratio, ETF holdings, futures positioning, and now the GFEX/China bid. Platinum is a high-beta satellite of the precious complex: it lags early, overshoots late, and its marginal buyer changed nationality in 2025.")
P("When it works / fails. Works in precious-complex phases: the 2025 squeeze (ratio 0.28→0.55 intraday extremes) and the January unwind both tracked flows, not fundamentals. Fails when the industrial leg decouples the metal from gold, and positioning signals are noisy.")
P(f"Standard workings. (1) Take gold (${RA['gold_spot']:,.0f} spot; LBMA 2026 average forecast $4,742). (2) Take the ratio the flows justify (2-year mean {RA['mean_2y']:.3f}×; 5-year {RA['mean_5y']:.3f}×; the post-2016 regime {RA['mean_post2016']:.2f}×). (3) Platinum = gold × ratio. (4) Overlay positioning (+8.3k net long — modest) and the China bar bid (13t and growing).")
P(f"Worked example, inputs named. At gold ${RA['gold_spot']:,.0f} and a neutral 0.42× (between the 2- and 5-year means), platinum is ~$1,670 — modestly above spot. The upside is not platinum-specific: gold at its LBMA-survey $4,742 with the ratio merely held (0.405×) maps to ~$1,921; the same gold at the 5-year-mean ratio (0.461×) is ~$2,186; the bull case below pairs a more conservative gold $4,600 with 0.45× → ~$2,070. The downside is a gold reversal dragging the complex: $3,600 gold at the 2-year ratio (0.388×) is ~$1,300.")
P("Sensitivity and falsification. A $400 move in gold ≈ $160/oz; a 0.04 move in the ratio ≈ $160/oz — the two dimensions carry equal weight. Falsification: a sustained Pt/gold decoupling (platinum moving hard while gold and the ratio do not), which would mean the industrial leg has taken over.")
P("Cross-examination. To Expert 1: 'Real rates drive gold and I already price gold — plus the flows you ignore.' To Expert 2: 'Your balance was in deficit for four straight years while the price went $900 → $2,772 → $1,608; flows set the path, your ledger sets the destination.'")
P(f"Verdict. Fair value ~$1,670 (range $1,300–2,070), the swing vote. What the price implies: at ${spot:,.0f} the market prices the recent-regime ratio with no re-rating — neither the gold consensus nor ratio reversion. Bull $2,070: gold to its consensus zone plus 5y-mean ratio reversion. Base $1,670: gold holds, ratio splits the recent windows. Bear $1,300: gold reverses, ratio pins at the 2-year mean, spec length liquidates.")

H2('D.5  The three in one room')
P(f"A short exchange on the single question they most disagree on — whether platinum is cheap at ${spot:,.0f}.")
P("Expert 2: 'Four deficits, 3.3 million ounces out of the vault, eleven weeks of cover — the thinnest platinum market ever recorded — and you call $1,608 fair? The cost curve is $1,006 and rising 8% a year, nobody is building, and the incentive price is $2,300. Scarcity this deep is worth $1,800 before any squeeze.'")
P("Expert 1: 'It ran to $2,772 on that story and gave half back in five months — because the Fed turned and the real yield did what it always does to a zero-coupon asset. Your stockpile did not defend the record; carry broke it. At two-point-three percent real — an eighteen-year high — the premium over your own cost floor is an option on 2027 easing, and options decay.'")
P(f"Expert 3: 'You are both pricing the wrong marginal buyer. The flow that built the squeeze was Chinese bars and NYMEX shorts covering; the flow that broke it was the same book unwinding. Gold at ${RA['gold_spot']:,.0f} and a recent-regime ratio puts platinum near $1,670, and the interesting fact is the ratio is still in its 12th percentile — cheap to gold even now. The swing is whether reverse substitution stops the ratio ever re-rating — which is Expert 2's variable, priced my way.'")

H2('D.6  Reading the divergence')
figure('figD1_experts.png', 6.9, "Figure D1. Expert and anchor fair values on one field. The spread ($1,500–1,800) is narrow relative to the Monte Carlo cone — the experts agree on the zone and disagree on where in it platinum sits.")
P(f"What the spread measures. The $1,500–1,800 range is tight for an asset this volatile — the experts agree platinum is roughly fairly valued and disagree only on the durability of the deficit. The spread is, precisely, the reverse-substitution question (with the rate regime setting its floor): Expert 2's $1,800 is the deficit persisting against it, Expert 1's $1,500 is a 2.3%-real-yield world capping everything, Expert 3's $1,670 is the ratio splitting the difference. Every one of them turns on the same real-unit variable the crux (§1.7) isolates — ounces per vehicle platform, on a 7-year procurement clock, against eleven weeks of inventory. That the panel is tight while the Monte Carlo cone (${MC['t63']['p5']:,.0f}–{MC['t63']['p95']:,.0f} at three months) is enormous is the honest picture: the fair value is knowable within a zone; the quarter-ahead price is not.")

# ============================ ABOUT / DISCLAIMER ============================
H1('About this series')
P("Each instrument in the Testahil series follows the same standing format — Headline → Valuation summary → 1. Fundamental → 2. Technical → 3. Monte Carlo → 4. Comparison → 5. Catalysts → 6. Probability map → 7. Caveats → Appendices (A financial statements or their metal analogue, B Step 0 calibration, C peer set, D expert panel) → Disclosure. A Step 0 calibration backtest gates every study: the forecasting engine is tested against a carry-anchored random-walk benchmark on scale-normalized CRPS before the instrument is valued, and the verdict is printed, not asserted — platinum's is PARITY under a provisional self-fit, stated plainly in Appendix B. For a commodity the fundamental section is adapted to fair-value anchors rather than company cash flows, and the company-specific template items are flagged as justified N/A and replaced by their metal analogues. The method is deliberately transparent and open to scrutiny; the aim is education in how fundamental, technical and probabilistic lenses combine, not a recommendation.")
H1('Disclosure & Disclaimer')
for head, body in [
    ('Not investment advice. ', 'This document is an educational valuation exercise and an expression of personal analytical opinion. It is not investment advice, not a recommendation or solicitation to buy, sell or hold platinum or any related instrument, and not directed at the objectives, financial situation or needs of any reader.'),
    ('No licence, no clients. ', 'The preparer is not licensed or registered with any securities or commodities regulator in any jurisdiction, provides no financial consultancy, manages no money, solicits no funds and accepts no fees or clients.'),
    ('Model outputs, not certainties. ', 'All figures are model outputs presented as ranges and probability distributions. No single number is a forecast or a price target; the Monte Carlo describes a distribution of possible outcomes, not an expected path.'),
    ('Sources and estimates. ', 'Market data are drawn from the WPIC Platinum Quarterly (Metals Focus), the LBMA, CFTC positioning data, published analyst research and exchange data; forward-looking inputs are the preparer\'s own judgments and are flagged as estimates throughout. Errors and omissions are possible.'),
    ('Risk warning. ', 'Platinum is a highly volatile asset (forward volatility ~34% annualised; a −17.6% single day occurred in January 2026); prices can move sharply and past behaviour does not indicate future results. Anyone acting on their own view should consult a licensed financial advisor first.'),
    ('Position disclosure. ', 'The preparer may hold, and may in the future take or dispose of, a position in the instrument discussed in this report.'),
]:
    rich([(head, dict(bold=True)), (body, dict())], size=9.6, space_after=4)
P('Testahil · Independent Valuation Study — Educational Analysis · Platinum (XPT/USD) · combined 1/3/12-month edition · 20 Jul 2026', size=8.7, color=GREY, space_before=8)

doc.save('XPTUSD_Valuation_Study_20-07-2026_public.docx')
print('docx saved')
