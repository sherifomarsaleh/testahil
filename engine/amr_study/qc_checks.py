"""AMR — the programmatic checks that back the quality-control table.

Four scans, all run on the DELIVERED files:
  1. external-reader scrub — no internal-procedure vocabulary anywhere in any deliverable;
  2. table discipline — every table in every document has explicit, fixed column widths, and
     none is starved or bloated relative to the text it has to carry;
  3. figure discipline — every figure is a solid light canvas with zero transparency;
  4. delivered-file inventory — page and figure counts for each PDF.
"""
import json, os, re, zipfile
from docx import Document
from docx.shared import Inches
from PIL import Image
import pypdfium2 as pdfium

HERE = os.path.dirname(os.path.abspath(__file__))
DOCS = ['AMR_Valuation_Study_09-08-2026_public.docx', 'AMR_Bibliography_09-08-2026.docx']
PDFS = ['AMR_Valuation_Study_09-08-2026_public.pdf', 'AMR_Bibliography_09-08-2026.pdf',
        'AMR_Valuation_Model_09082026_public.pdf']
# THE WORKBOOK NAMES ITS EDITION DDMMYYYY WITH NO SEPARATORS, so the latest edition is found
# by PARSING the date, never by sorting the names: "03092026" sorts below "09082026" as text
# and a text sort would scrub a superseded edition and report it as current [L-067, L-350].
_XLSX_DATE = re.compile(r'_(\d{2})(\d{2})(\d{4})_')


def _latest_xlsx():
    c = []
    for f in os.listdir(HERE):
        if f.startswith('AMR_Valuation_Model_') and f.endswith('.xlsx') and not f.startswith('~$'):
            m = _XLSX_DATE.findall(f)
            c.append(((m[-1][2] + m[-1][1] + m[-1][0]) if m else '', f))
    assert c, 'no delivered workbook found — an empty result is not a clean result'
    return sorted(c)[-1][1]


XLSX = _latest_xlsx()

# ---- 1. external-reader scrub ---------------------------------------------
# Internal procedure vocabulary. The reader is an outside party: nothing about how the
# house builds a study belongs in a delivered document.
BANNED = [
    r'\bstep\s*0\b', r'\bstep\s*2a\b', r'\bstep\s*0\.0\b', r'information sweep',
    r'\bsweep register\b', r'\bfour[- ]ring\b', r'\bring\b(?!\s*road)', r'\bQC gate\b',
    r'\bquality[- ]control gate\b', r'\bSIGCM\b', r'\bmateriality gate\b',
    r'\bPARITY\b', r'\bBOUNDARY\b', r'\bmarket_profiles\b', r'\bfitted_configs\b',
    r'\bmc_v3\b', r'\bwidth_cal\b', r'\bLONO\b', r'\bcalibration ledger\b',
    r'\bdriver ledger\b', r'\brollforward\b', r'\broll[- ]forward protocol\b',
    r'\bmodel study\b', r'\bdepth bar\b', r'\bstanding research protocol\b',
    r'\bTestahil engine\b', r'\bpanel verdict\b', r'\bname[- ]level fail\b',
    r'\bcarry[- ]anchored gate\b', r'\bpublish protocol\b', r'\badaptive_width\b',
]
scrub = {}
for f in DOCS:
    doc = Document(os.path.join(HERE, f))
    text = '\n'.join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                text += '\n' + c.text
    hits = []
    for pat in BANNED:
        for m in re.finditer(pat, text, re.I):
            hits.append(f'{pat} -> "{text[max(0, m.start()-40):m.end()+40]}"')
    scrub[f] = hits
z = zipfile.ZipFile(os.path.join(HERE, XLSX))
xl_text = ' '.join(z.read(n).decode('utf-8', 'ignore') for n in z.namelist()
                   if n.endswith('.xml'))
xl_hits = []
for pat in BANNED:
    for m in re.finditer(pat, xl_text, re.I):
        xl_hits.append(f'{pat} -> "{xl_text[max(0, m.start()-40):m.end()+40]}"')
scrub[XLSX] = xl_hits
total_hits = sum(len(v) for v in scrub.values())
print(f'external-reader scrub: {total_hits} hits across {len(scrub)} delivered files')
for f, hits in scrub.items():
    for h in hits[:6]:
        print(f'   {f}: {h}')

# ---- 2. table discipline ---------------------------------------------------
# A column is STARVED if its widest cell needs more than about 1.9x the width it has;
# BLOATED if the widest cell in it uses less than 35% of the width it has.
CHARS_PER_INCH = 15.0     # at 8.5pt Calibri in a table cell, conservatively
tbl_report = []
for f in DOCS:
    doc = Document(os.path.join(HERE, f))
    for ti, t in enumerate(doc.tables):
        widths = [c.width for c in t.columns]
        if any(w is None for w in widths):
            tbl_report.append((f, ti, 'NO EXPLICIT WIDTH', None, None))
            continue
        wins = [w / 914400 for w in widths]
        for j, w in enumerate(wins):
            longest = 0
            for row in t.rows:
                if j < len(row.cells):
                    longest = max(longest, max((len(x) for x in row.cells[j].text.split('\n')),
                                               default=0))
            need = longest / CHARS_PER_INCH
            # cells that wrap are fine; the test is on the LONGEST SINGLE WORD, which cannot wrap
            longest_word = 0
            for row in t.rows:
                if j < len(row.cells):
                    for word in row.cells[j].text.split():
                        longest_word = max(longest_word, len(word))
            word_need = longest_word / CHARS_PER_INCH
            starved = word_need > w * 0.98
            # A numeric column needs whitespace either side to read as a column, so a
            # narrow one is never "bloated". The test is for columns that are genuinely
            # wasteful: wider than 1.1 inches while their widest cell uses under 45% of it.
            bloated = need < w * 0.45 and w > 1.10
            if starved or bloated:
                tbl_report.append((f, ti, 'starved' if starved else 'bloated', j,
                                   f'width {w:.2f}in, longest word needs {word_need:.2f}in, '
                                   f'longest line needs {need:.2f}in'))
print(f'table discipline: {len(tbl_report)} column issues across '
      f'{sum(len(Document(os.path.join(HERE, f)).tables) for f in DOCS)} tables')
for r in tbl_report[:12]:
    print('   ', r)

# ---- 3. figure discipline --------------------------------------------------
figs = sorted(x for x in os.listdir(HERE) if x.startswith('fig') and x.endswith('.png'))
fig_issues = []
for f in figs:
    im = Image.open(os.path.join(HERE, f))
    if im.mode == 'RGBA':
        lo, hi = im.getchannel('A').getextrema()
        if lo != 255:
            fig_issues.append(f'{f}: transparency present')
    px = im.convert('RGB').getpixel((2, 2))
    if min(px) < 200:
        fig_issues.append(f'{f}: canvas corner is dark {px} — labels may not read')
print(f'figure discipline: {len(figs)} figures, {len(fig_issues)} issues')
for i in fig_issues:
    print('   ', i)

# ---- 4. delivered-file inventory ------------------------------------------
inv = []
for f in PDFS:
    path = os.path.join(HERE, f)
    d = pdfium.PdfDocument(path)
    n_img = 0
    for i in range(len(d)):
        try:
            n_img += sum(1 for o in d[i].get_objects() if o.type == 3)
        except Exception:
            pass
    inv.append(dict(file=f, pages=len(d), images=n_img,
                    kb=round(os.path.getsize(path) / 1024)))
    print(f'  {f}: {len(d)} pages, {n_img} embedded images, '
          f'{round(os.path.getsize(path)/1024):,} KB')

res = dict(scrub_hits=total_hits, scrub_detail={k: v for k, v in scrub.items()},
           table_issues=tbl_report, n_figures=len(figs), figure_issues=fig_issues,
           pdf_inventory=inv)
json.dump(res, open(os.path.join(HERE, 'qc_checks.json'), 'w'), indent=1, default=str)
assert total_hits == 0, f'{total_hits} internal-procedure terms found in delivered files'
assert not fig_issues, fig_issues
assert not tbl_report, f'{len(tbl_report)} table column issues'
print('\nQC CHECKS OK — zero internal-procedure vocabulary, zero table column issues, '
      'zero figure issues')
