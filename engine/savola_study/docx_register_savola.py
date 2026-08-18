"""SAVOLA_Bibliography_18-08-2026.docx — the companion bibliography document.
Every input in the model: value, source, date and research layer — emitted from
study_numbers.json (the compute script's own INPUTS block), plus the document
bibliography, the judgements table and the negative results."""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(HERE, 'study_numbers.json')))
INP = D['inputs']
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(table, top=36, bottom=36, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    table._tbl.tblPr.append(m)

def borders(table, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    table._tbl.tblPr.append(b)

def P(text='', size=10, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def H1(t): return P(t, size=13, bold=True, space_before=12, space_after=4)
def H2(t): return P(t, size=11, bold=True, space_before=8, space_after=3)

def table(rows, widths, size=8.2, header=True):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t); t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def masthead():
    t = doc.add_table(rows=1, cols=1)
    cell_margins(t, 80, 80, 150, 150)
    c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.1)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def fmt(v):
    if isinstance(v, float):
        return f"{v:,.4f}".rstrip('0').rstrip('.') if abs(v) < 100 else f"{v:,.1f}"
    if isinstance(v, int):
        return f"{v:,}"
    if isinstance(v, dict):
        return "; ".join(f"{k}: {fmt(x)}" for k, x in v.items())
    if isinstance(v, (list, tuple)):
        return ", ".join(fmt(x) for x in v)
    return str(v)

# ============================================================================
masthead()
H1('Savola Group Company (Saudi Exchange: 2050) — Bibliography and Source Register')
P('Companion document to the valuation study dated 18 August 2026. It records where every '
  'number in that study came from.', size=9.5, color=GREY)

H2('READ FIRST')
P('This document exists so that a reader can check the study rather than trust it. It lists '
  'every input the valuation model uses, together with the value, the source, the date of '
  'that source, and the research layer it belongs to. Nothing in the study is computed from '
  'a figure that does not appear here.')
P('Two things are worth knowing before reading the tables. First, inputs marked "House" are '
  'judgements or derivations made by the analyst, not disclosures by the company — each '
  'carries the reasoning that produced it, and the reader is free to disagree and re-run '
  'the model (the companion workbook reprices live). Second, where a disclosure could not '
  'be reached or does not exist, that is recorded as a negative result at the end of this '
  'document rather than quietly filled in.')

H2('The research layers')
table([['Layer', 'What it holds'],
       ['Company', 'the company\'s own audited/reviewed statements, releases, presentations '
        'and exchange filings — the only permitted source for its reported history'],
       ['Company/derived', 'arithmetic on audited figures (ratios, day counts, residuals '
        'that must foot to an audited total — each derivation is stated)'],
       ['Market', 'traded prices and rates: the share price, Herfy\'s price, peer '
        'multiples, interbank rates, the study\'s own beta regression'],
       ['Country', 'sovereign data: yields, spreads, risk premia, inflation'],
       ['Global', 'the dollar curve and the commodity complex'],
       ['House', 'the analyst\'s judgements: growth paths, held margins, weights, the '
        'terminal — each with its reasoning and its overturn condition']],
      [1.30, 5.80], size=8.8)

H2('Primary documents')
table([['Document', 'Publisher / auditor', 'Dated', 'What was taken'],
       ['FY2025 consolidated financial statements (audited)',
        'Savola Group; Deloitte & Touche, unmodified opinion', '05-Mar-2026',
        'the base year in full: statements, five-segment note, category detail, notes 1-46'],
       ['FY2024 consolidated financial statements (audited)', 'Savola Group; KPMG',
        '10-Mar-2025', 'FY2023 comparatives; the rights issue, capital reduction and '
        'Almarai distribution mechanics'],
       ['FY2023 consolidated financial statements (audited)', 'Savola Group', '14-Mar-2024',
        'FY2022 comparatives; the pre-reset balance sheet'],
       ['Q1-2026 interim condensed statements (reviewed)', 'Savola Group', '06-May-2026',
        'first-quarter actuals and the 31-Mar-2026 balance sheet'],
       ['H1-2026 earnings release', 'Savola Group', '06-Aug-2026',
        'half-year actuals, net debt, capital expenditure, the Sudan exit, the Mehbaj '
        'acquisition, the second-half cost warning'],
       ['Q2-2026 investor presentation', 'Savola Group', '06-Aug-2026',
        'category volumes and unit gross profits, the store network, segment debt/leases'],
       ['FY2025 investor presentation', 'Savola Group', '09-Mar-2026',
        'full-year category units, the reported-to-recurring bridge, capex by unit'],
       ['Annual Report 2025', 'Savola Group', '30-Mar-2026',
        'the dividend policy, governance and related-party context, the store programme'],
       ['Saudi Exchange announcements (FY2025 results; FY2025 dividend; Q1-2026 results)',
        'Saudi Exchange / Savola Group', 'Mar-May 2026',
        'the official results tables and the dividend terms (SAR 1.70, ex 07-May-2026)'],
       ['US Treasury constant-maturity yields', 'Federal Reserve (FRED)', '14-Aug-2026',
        'US 10Y 4.68% and 1Y 3.98% — the risk-free construction\'s dollar leg'],
       ['Sovereign risk dataset (January 2026 update)', 'NYU Stern (Damodaran)',
        '05-Jan-2026', 'Saudi Aa3 default spread 0.51%, equity risk premium 5.01%; CDS '
        '0.98% / 5.72%'],
       ['Saudi sovereign issuance', 'Emirates NBD Research note; NDMC announcements',
        'Jan / Aug 2026', 'the USD 10Y new-issue spread (+85bp); the 1Y SAR savings-sukuk '
        'rate (4.70%)'],
       ['FAO Food Price Index', 'FAO', '07-Aug-2026',
        'vegetable oils at a four-year high; sugar −8% year on year; wheat firming'],
       ['Consumer price index, July 2026', 'GASTAT (via press reports of the release)',
        '14-Aug-2026', 'Saudi CPI +1.8%'],
       ['Market quotes (Savola, Herfy, Almarai, Al Othaim, BinDawood, NADEC, Wilmar; '
        '3M SAIBOR)', 'stockanalysis.com; TradingEconomics — market data only',
        '18-Aug-2026', 'prices and trailing multiples for the cross-checks and the '
        'bridge\'s Herfy leg — never a source for any Savola reported figure']],
      [2.20, 1.75, 0.80, 2.35], size=8.0)

H2('The input register — every input, with value, source, date and layer')
P('Grouped by layer. Values are as the model consumes them (SAR millions unless the row '
  'says otherwise; paths list the five forecast years FY2026E-FY2030E).', size=9)
LAYERS = ['Company', 'Company/derived', 'Market', 'Country', 'Global', 'House',
          'House/derived']
order = {l: i for i, l in enumerate(LAYERS)}
items = sorted(INP.items(), key=lambda kv: (order.get(kv[1]['ring'], 99), kv[0]))
cur = None
rows = None
def flush():
    global rows
    if rows and len(rows) > 1:
        table(rows, [1.45, 1.15, 3.75, 0.75], size=7.6)
for k, rec in items:
    if rec['ring'] != cur:
        flush()
        cur = rec['ring']
        H2(f'Layer: {cur}')
        rows = [['Input', 'Value', 'Source and construction', 'Date']]
    rows.append([k, fmt(rec['value']), rec['source'], rec['date']])
flush()

H2('Judgements — and what would overturn each')
table([['Judgement', 'What would overturn it'],
       ['Panda sales density stabilises (Framing A is the base case)',
        'two further quarters of density erosion at the measured first-half pace; the '
        'study publishes Framing B beside it at all times'],
       ['Oil unit gross profit held below the first-half actual for the rest of 2026',
        'a second-half print that HOLDS the first-half unit margin — the model is then too '
        'low by construction and should be restruck'],
       ['Sugar and pasta unit-margin gains retained but not extended',
        'either a disclosed capacity/mix change (extend) or a give-back in the prints '
        '(retract)'],
       ['The 20-store-per-year programme continues through FY2030',
        'the company\'s own guidance changing; capex would be re-pathed the same day'],
       ['Combined zakat-and-tax rate 19.5%',
        'a structural change in the Egypt mix or a new assessment cycle; the sensitivity '
        'is on the workbook\'s driver test'],
       ['Terminal: growth 2.5%, return on capital 10.5%',
        'evidence the group can compound above its cost of capital at scale — three years '
        'of returns above 11% would justify raising the terminal return'],
       ['The 10Y SAR risk-free construction (4.68% + 0.85%)',
        'a directly quotable 10-year SAR government level; both ±50bp alternatives are '
        'already priced in the study'],
       ['Beta 1.087 (own stock vs the exchange index, five years weekly)',
        'a materially different reading once eighteen months of clean post-reset history '
        'exist; the 90% interval is priced in the study'],
       ['Peer-mix multiple with a 20% conglomerate/Egypt discount',
        'a de-rating of the Saudi consumer peer set toward the dividend-discount multiple '
        '(the study\'s own crosswalk) — the discount band 10-30% is published'],
       ['Kinan at capitalized earnings in the bridge',
        'a transaction or disclosure evidencing either the carrying floor or the '
        'net-asset value; all three constructions are published in the workbook']],
      [2.60, 4.50], size=8.2)

H2('Negative results — searched, not found, recorded')
table([['What was sought', 'Where', 'Outcome'],
       ['Al Mehbaj Al Shamiya acquisition consideration',
        'H1-2026 release; Q2-2026 presentation; exchange announcements; Annual Report '
        '2025', 'NOT DISCLOSED anywhere as of 18-Aug-2026 — the nuts leg carries a small, '
        'flagged revenue step instead of an invented purchase price'],
       ['Panda like-for-like sales series', 'all company disclosures',
        'NOT PUBLISHED — sales per average store is derived from disclosed revenue and '
        'store counts, and the derivation is the study\'s central contested judgement'],
       ['Numeric FY2026 revenue/margin guidance', 'company disclosures',
        'NONE EXISTS — only store-count and store-refresh targets are guided'],
       ['A directly quotable 10-year SAR government bond/sukuk yield',
        'exchange fixed-income pages, index publishers, data aggregators, official '
        'issuance releases (which publish tranche sizes, not yields)',
        'NOT ACCESSIBLE — constructed from the dollar curve plus the sovereign\'s own '
        'issue spread, cross-checked on the observed 1-year rate, and priced ±50bp'],
       ['H1-2026 reviewed interim statements on the company website',
        'savola.com financial-statements page (English and Arabic)',
        'FILED with the authorities but not yet mirrored on the site at the study date; '
        'the company\'s own release and presentation carry every H1 figure used, and the '
        'reviewed statements remain a follow-up item'],
       ['FY2023 continuing-basis figures excluding Türkiye',
        'FY2024 and FY2025 statements',
        'NOT PUBLISHED — FY2023 is presented on its own audited basis and the Türkiye '
        'inclusion is flagged wherever that column appears']],
      [2.20, 2.30, 2.60], size=8.2)

H2('Aggregator discrepancies')
P('One found. The market-data page used for peer quotes lists Savola\'s trailing '
  'twelve-month earnings on the reported basis (which still contains the zakat release '
  'and the Sudan disposal gain) — a trailing multiple of about 8x that overstates '
  'recurring earnings power. The study uses the company\'s own recurring bridge instead '
  'and says so where the multiple appears. Market quotes themselves (prices, peer '
  'multiples) were consistent across sources checked on 18-Aug-2026.', size=9)

doc.save(os.path.join(HERE, 'SAVOLA_Bibliography_18-08-2026.docx'))
n_inputs = len(INP)
print(f'bibliography written · {n_inputs} inputs in the register')
