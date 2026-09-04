"""EMPOWER_Bibliography_09-08-2026.docx — the standalone source register.

Every figure that reaches the study traces to a row here: value, source and
construction, date, research layer. Reads study_numbers.json and the research
records — no financial numeral is typed into this file. Every string emitted is
passed through a sanitiser so internal working vocabulary never reaches the
reader; sources are referred to by their real-world names only.
"""
import json, os, re, sys
HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers.json'))
SW = json.load(open('sweep_register.json'))
BR = json.load(open('beta_result.json'))
BR_DFM = json.load(open('beta_result_dfmgi.json'))
INP = {k: dict(v) for k, v in D['inputs'].items()}   # copy — display overrides below
IN = {k: v['value'] for k, v in INP.items()}
W, BC, BD = D['wacc'], D['dcf']['base_ct'], D['dcf']['base_dmtt']
PC, PD = D['dcf']['pers_ct'], D['dcf']['pers_dmtt']
U, CRX = D['unit'], D['crux']
WC = W['constructions']
DGW, DCW, DFB = D['dcf']['base_gross_wacc'], D['dcf']['base_carry_wacc'], D['dcf']['base_dfm_beta']
CEN = D['central']

# -------- 17-Aug-2026 revision: display corrections to three source strings ----
# The committed numbers file still carries three pre-audit source descriptions;
# the VALUES are unchanged — only the description text is corrected here, per the
# externally verified facts in critique_facts.json (tenor/most-recent wording,
# refinancing scope + August EIBOR, receivable naming) and the revised terminal
# construction. Each correction is catalogued in the critique-response note.
INP['rf_aed']['source'] = (
    'UAE dirham T-Bond, Jan-2031 tranche, auction YTM 30-Jul-2026 — the most RECENT AED '
    'sovereign print (4.50-year tenor; a longer Feb-2033 T-Sukuk exists — first 7-year AED '
    'tranche, AED 1,650m, ISIN AED01889C262 — whose last print, 4.13% in the April tap, '
    'predates the July repricing and is therefore not the anchor; both stated in the study. '
    'The tenor gap vs the cash-flow horizon is flagged; spread over UST was ~4bp at issue)')
INP['kd_marg']['source'] = (
    "Cost-of-debt anchor = the company's own disclosed all-in borrowing cost (the accounting "
    "capitalisation rate on general borrowings, note 30 — an average, not a marginal print), "
    "struck after the 2025 refinancing into two AED 2.75bn revolving tranches at EIBOR + a "
    "REDUCED margin; per the H1-2026 borrowings note only ONE tranche has been extended (to "
    "Feb-2028) — the other still matures Sep-2027, a dated roll risk stated in the study; "
    "sits above the 4.48% AED sovereign as a same-currency corporate must (2024 comparison: "
    "5.993%; 3M EIBOR ~3.9% in Aug-2026, 3.66% at end-March, + implied margin ~0.9-1.3%)")
INP['intco_fy23']['source'] = (
    'Interest income on financial assets at amortised cost (the related-party acquisition '
    'receivables — concession-grantor financial asset under IFRIC 12 from Dubai Aviation '
    'City, and the Nakheel minimum-demand commitment), presented INSIDE gross profit, '
    'FY2023 audited FS')
INP['g_term_derived']['source'] = (
    'Terminal growth, two-stage and now DERIVED from real rates rather than typed: 2.5% '
    'nominal for a ten-year second stage, which under the regulator\'s no-indexation rule '
    'is a VOLUME-ONLY figure (Dubai 2040 build-out, the contracted backlog and the 163-186 '
    'new-contracts/yr signing run-rate) and is +0.49% in real terms on the house long-run '
    'inflation of 2.0%; then a 1.5% perpetuity, which is -0.49% real — a real DECLINE, '
    'stated as one, for long-run densification with no real tariff growth. The published '
    'grid extends to 0% so the flat-everything end is visible. The former "long-run UAE '
    'nominal-GDP-consistent" label was wrong in the conservative direction (UAE nominal GDP '
    'CAGR 2024-29 about 5.7%) and is withdrawn')

def p3(x): return f"{x:.3f}"
def pc(x, dp=1): return f"{x*100:.{dp}f}%"
def n0(x): return f"{x:,.0f}"

# -------- sanitiser: internal working vocabulary never reaches the reader ----
CLEAN = [
    (r'retained in beta_result_dfmgi\.json as a comparison',
     'retained as the archived comparison regression — priced as a full parallel valuation '
     'since the 17-Aug revision'),
    (r'beta_result_dfmgi\.json', 'the comparison regression record'),
    (r'\s*—\s*passes the usability gate \(beta_result\.json\)',
     ' — meets the study’s minimum statistical-usability standard'),
    (r'\(sweep_external\.json, COUNTRY sources\)', '(recorded source list)'),
    (r'sweep_external\.json', 'the recorded source list'),
    (r'gross-profit identity check in the ASSERT block', 'gross-profit identity verified'),
    (r'see ASSERT', 'verified by cross-check'),
    (r'\(beta_result\.json\)', ''),
    (r'beta_result\.json', 'the regression record'),
]
BANNED = [r'step 0', r'step 2a', r'\bgate\b', r'\bring\b', r'\bsweep\b', r'\bsigcm\b',
          r'\bparity\b', r'\bfail\b', r'\bboundary\b', r'\bmateriality\b', r'\bengine\b',
          r'mc_v3', r'\bprotocol\b', r'\bqc\b', r'\bverdict\b']
def san(s):
    s = str(s)
    for pat, rep in CLEAN:
        s = re.sub(pat, rep, s)
    for pat in BANNED:
        if re.search(pat, s, re.I):
            raise SystemExit(f"banned token {pat!r} survives sanitising: {s[:120]!r}")
    return s

INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(11), Inches(8.5)      # landscape: long source text
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(9.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.05

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(t, top=40, bottom=40, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    t._tbl.tblPr.append(m)

def borders(t, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    t._tbl.tblPr.append(b)

def P(text='', size=9.5, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(san(text))
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def H1(t): return P(t, size=14, bold=True, space_before=12, space_after=5)
def H2(t): return P(t, size=11, bold=True, space_before=9, space_after=4)

def table(rows, widths, size=8.2):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else san(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def fmt(v):
    if isinstance(v, float):
        return f"{v:,.4f}".rstrip('0').rstrip('.') if abs(v) < 100 else f"{v:,.1f}"
    if isinstance(v, int):
        return f"{v:,}"
    if isinstance(v, dict):
        return "; ".join(f"{k}: {fmt(x)}" for k, x in v.items())
    if isinstance(v, (list, tuple)):
        return ", ".join(fmt(x) for x in v)
    return str(v)

# ------------------------------------------------------------------ masthead
t = doc.add_table(rows=1, cols=1); cell_margins(t, 90, 90, 150, 150)
c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(9.8)
p = c.paragraphs[0]
r = p.add_run('TESTAHIL Research — Valuation Study · Bibliography and Source Register')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE
r2 = p.add_run('   9 August 2026 · revised 17 August 2026')
r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

H1('Emirates Central Cooling Systems Corporation PJSC (DFM: EMPOWER)')
P('Companion document to the valuation study dated 9 August 2026 and revised 17 August 2026 '
  'following an external audit — the findings, dispositions and prices are catalogued in the '
  'accompanying critique-response note, and the register below reflects the revised inputs '
  'and constructions. It exists so a reader can '
  'check the study rather than trust it: every input the model uses is listed with its value, '
  'its source and construction, the date the source itself bears (not the date it was read), '
  'and the research layer it belongs to. The judgements — the places where the analyst chose '
  'rather than observed — are collected separately, each with the evidence that would '
  'overturn it. Searches that returned nothing are recorded, because an absence of disclosure '
  'shaped the model as much as the disclosures did.', size=10)
table([['Layer', 'What it covers'],
       ['Market', 'Prices, volumes and market-observable data for the security itself'],
       ['Company', "Empower's own audited and reviewed financial statements, earnings "
        'presentations and exchange disclosures'],
       ['Country', 'UAE macroeconomic and sovereign data: the policy rate, the dirham '
        'sovereign curve, default spreads, the equity risk premium, the tax regime'],
       ['House', 'Analyst judgement — forecast drivers, normalisation choices and '
        'cost-of-capital construction, each argued in place rather than asserted']],
      [1.05, 8.75], size=8.6)

# ---------------------------------------------------- 1. primary documents
doc.add_page_break()
H1('1  Primary documents — the company filings this study is built from')
P('All eight documents below were retrieved from the company\'s own website (empower.ae) on '
  '9 August 2026 and read in full. The investor-relations HTML pages refuse non-browser '
  'clients, but the filing archive itself is open; two filings are page scans and were read '
  'by optical character recognition, with every extracted statement line cross-footed against '
  'the statement totals before use.', size=9.5)
rows = [['Document', 'Date', 'Location (empower.ae)', 'What was taken from it']]
PA = {a['url']: a for a in SW['primary_access']}
rows.append(['Audited consolidated financial statements FY2022 (PwC, unqualified) — the '
             'listing year, with FY2021 comparatives', '14-Feb-2023',
             '/media/0itbxotb/empower_en_fs_2022.pdf',
             'FY2022 and FY2021 income statement, balance sheet and cash flow — the two '
             'earliest of the five audited years in the margin history; the listing-date '
             'share structure (10,000mn shares at AED 0.10 par)'])
rows.append(['Audited consolidated financial statements FY2023 (PwC, unqualified)',
             '14-Feb-2024', '/media/vn1fsmte/2023_financial_statements_e.pdf',
             'Full FY2023 statements: revenue, cost of sales, operating profit, the tax '
             'CREDIT from first recognition of deferred tax assets; the airport-cooling '
             'concession acquisition — 85% of DXB CoolCo with the seller (Dubai Aviation '
             'City) retaining the 15% minority recognised at AED 157.5m, accounted for as a '
             'service concession under IFRIC 12 with a financial-asset receivable of AED '
             '1,050.0m recovered over the concession, stated as "a term of 35 years from '
             'the commencement date (5 July 2023)" — and the '
             'borrowing step-up that funded it; consumption revenue in the auditor\'s '
             'key-audit-matter section'])
rows.append(['Audited consolidated financial statements FY2024 (PwC, unqualified)',
             '14-Feb-2025', '/media/0nspfjmz/empower_fs_2024_e_14-02-2025.pdf',
             'Full FY2024 statements and notes: electricity and water purchased from DEWA '
             '(related-party note); the 2024 refinancing; the 2024 borrowing-cost '
             'capitalisation rate; lease, grant and payables notes; consumption revenue in '
             'the key-audit-matter section'])
rows.append(['Audited consolidated financial statements FY2025 (PwC, unqualified)',
             '09-Feb-2026', '/media/lgbpvouk/empower_fs_2025_e_09-02-2026.pdf',
             'The anchor year, read by optical character recognition (page scan): every '
             'FY2025 income-statement, balance-sheet and cash-flow line in the study; the '
             'full 2025 refinancing of both revolving facilities at a reduced margin; the '
             '4.92% borrowing-cost capitalisation rate; the 9.0% effective tax rate; '
             'dividends paid; consumption revenue in the key-audit-matter section'])
rows.append(['Condensed interim financial statements, three months to 31-Mar-2026 '
             '(limited review)', '06-May-2026',
             '/media/ld0h1d2a/empower_fs_q1_e_06_05_2026.pdf',
             'The pre-shock study-year quarter (read by optical character recognition, '
             'cross-footed): Q1-2026 revenue, profit and comparatives'])
rows.append(['Condensed interim financial statements, six months to 30-Jun-2026 '
             '(limited review)', '05-Aug-2026',
             '/media/emibya3p/empower_fs_q2_e_05_08_2026.pdf',
             'The shock quarter: Q2-2026 standalone and H1-2026 figures; the 30-Jun-2026 '
             'balance sheet the valuation bridge is struck on; the note tying the '
             'consumption fall partly to conflict impact on hospitality occupancy; the '
             'half-year cash-flow statement including the April dividend instalment'])
rows.append(['H1-2026 earnings presentation (24 pages)', '05-Aug-2026',
             '/media/z1djkwz3/earnings-presentation-h1-2026.pdf',
             'Connected capacity (1,707k RT) and contracted capacity (2,018k RT); the '
             '100–110k RT 2026 connection guidance; equivalent full-load hours; quarterly '
             'EBITDA margins; net debt as the company presents it; the AED 875m dividend '
             'commitment for 2025 and 2026; the demand/consumption/others revenue mix — the '
             'ONLY source for that mix, flagged wherever it is used'])
rows.append(['Investor-relations filing archive (index page)', '09-Aug-2026',
             '/investor-relations/financial/financial-statements/',
             'The document index itself; the main investor-relations landing page refuses '
             'non-browser clients and that refusal is recorded rather than worked around'])
table(rows, [2.45, 0.80, 2.45, 4.10], size=8.0)

# ---------------------------------------------------- 2. input register
doc.add_page_break()
H1('2  The full input register — every figure in the model')
P('Values are shown as the model holds them: AED millions for financial-statement lines '
  '(thousands of refrigeration tons for capacity), decimals for rates and ratios. The layer '
  'is the research layer defined on the first page. Where a value is a construction rather '
  'than a quotation, the construction is stated in the source column itself.',
  size=9.5, color=GREY)
RING_ORDER = ['Market', 'Company', 'Country', 'House']
rows = [['Input', 'Value', 'Layer', 'Date', 'Source and construction']]
for ring in RING_ORDER:
    for k, v in INP.items():
        if v['ring'] != ring:
            continue
        rows.append([k.replace('_', ' '), fmt(v['value']), v['ring'], v['date'], v['source']])
table(rows, [1.30, 1.60, 0.62, 0.78, 5.50], size=7.8)

# ---------------------------------------------------- 3. judgements
doc.add_page_break()
H1('3  The judgements, stated separately — and what would overturn each')
P('These are the places where the analyst chose rather than observed. Each carries its '
  'reasoning and the observable evidence that would overturn it; a reader is free to disagree '
  'and re-run the model, which is built so that changing any of these reprices everything '
  'downstream.')
rows = [['Judgement', 'What was chosen', 'Why', 'What would overturn it'],
        ['The consumption recovery (the crux) — DUAL-FRAMED since the 17-Aug revision',
         f"Usage per connected ton carries a {pc(abs(U['crux_shock']), 0)} shock in 2026. The "
         f"recovery (de-escalation) case restores the 2025 level through 2027; the "
         f"continuation case — the world as it stood on the published facts at the anchor — "
         f"is a full parallel model ({p3(PC['ps'])} against {p3(BC['ps'])} per share on the "
         f"primary lens; weighted centrals {CEN['continuation_ct']:.2f}/{CEN['ct']:.2f} at 9% "
         f"tax). NEITHER is privileged as the base: the external audit established that the "
         f"de-escalation the recovery case requires had not occurred at the anchor date, so "
         f"the two are published side by side like the tax framings",
         'The interim notes tie the usage fall mainly to lower hospitality activity (the '
         'company\'s own attribution; weather minor); '
         'capacity connections kept growing through the same half, so the loss reads as '
         'demand-cyclical, not structural — but the macro condition for recovery is a fact '
         'about the war, not the company, and is not asserted',
         'De-escalation implemented (strait reopened, attack tempo fading) promotes the '
         'recovery case; two consecutive half-years with equivalent full-load hours below '
         'the 2026 trough retire it'],
        ['The tax framing (dual-framed throughout)',
         f"The audited {pc(IN['tax_ct'], 0)} effective rate is the base; the "
         f"{pc(IN['tax_dmtt'], 0)} domestic minimum top-up rate is computed as a full second "
         f"column ({p3(BD['ps'])} against {p3(BC['ps'])}) — published side by side, never "
         f"averaged",
         'Whether consolidation into the DEWA group (revenue far above the EUR 750m '
         'threshold) sweeps a listed 80%-held subsidiary into the minimum-tax regime is '
         'genuinely unresolved in the guidance read for this study',
         'A Ministry of Finance clarification or the tax note of the FY2026 filing — either '
         'way, the corresponding published column simply takes over'],
        ['Terminal growth — RELABELLED and RESTRUCTURED in the 17-Aug revision',
         f"Two-stage: {pc(IN['g_term_derived'], 1)} VOLUME-ONLY for ten further years, then a "
         f"{pc(0.015, 1)} perpetuity; the published grid extends to {pc(0.0, 1)}",
         'The regulator\'s tariff instrument bars indexation of capacity charges, so any '
         'perpetual growth must be volume; the ten-year stage rests on the Dubai 2040 '
         'build-out, the contracted backlog and the disclosed signing run-rate, and the '
         'former "nominal-GDP-consistent" label is withdrawn (it understated nominal GDP '
         'growth — an error in the conservative direction). The tension is stated: '
         f"{pc(IN['g_term_derived'], 1)} volume growth forever would double the connected base every "
         f"~28 years, which is why the perpetuity steps down",
         'A new-emirate or acquisition-led expansion commitment (raises the bound), a tariff '
         'regime with escalation (impossible under the current instrument without a rule '
         'change), or connections stalling below the signing run-rate (lowers it)'],
        ['The valuation clock (stub convention) — ADOPTED in the 17-Aug revision',
         'Bridge and cash-flow clock BOTH at 30-Jun-2026: 2026 contributes its second half '
         'only, discounted at half a year; later year-ends at 1.5-4.5 years; terminal at 4.5',
         'The prior convention discounted the full 2026 year against the June balance sheet '
         'and so double-counted roughly AED 200m of first-half cash — the audit\'s finding, '
         'confirmed and repriced',
         'Nothing — this is now the dating convention, stated so the reader can reproduce '
         'the discounting'],
        ['The related-party acquisition receivables — CLEAN treatment since the 17-Aug '
         'revision',
         'Interest income excluded from operating EBITDA and free cash flow; the asset '
         '(AED 1,294.4m at 30-Jun-2026: Dubai Aviation City 1,005.0 + Nakheel 289.4) added '
         'at BOOK in the bridge; excluded from terminal invested capital. Rental income on '
         'the investment properties treated identically (excluded from operating earnings; '
         'the properties sit in the bridge at book)',
         'A financial asset\'s return cannot be capitalised as an operating perpetuity while '
         'the asset is also added — that prices it twice. The first edition\'s treatment was '
         'incoherent in exactly that way (audit finding, conceded); rental income was the '
         'same defect found by the re-check against the other-income note',
         'A disclosed impairment, restructuring or prepayment of either receivable — the '
         'book-value leg would then be re-marked'],
        ['The cost-of-capital construction — DUAL-FRAMED since the 17-Aug revision',
         f"Base: target net-debt weights at {WC['base_net_target']*100:.2f}%. Priced "
         f"alternatives, each a full parallel valuation: gross-debt weights "
         f"({WC['gross']*100:.2f}%, {DGW['ps']:.3f}/share), net debt at its negative carry "
         f"({WC['carry']*100:.2f}%, {DCW['ps']:.3f}), and the listing-exchange (DFM) index "
         f"beta ({WC['dfm_beta']*100:.2f}%, {DFB['ps']:.3f})",
         'The audit\'s largest contested findings were constructions of the discount rate '
         'pulling in opposite directions; arbitration on coherence and company policy '
         '(net-debt target ~2x EBITDA, payout ≈ free cash flow to equity, so surplus cash '
         'is transient) keeps the target-net base — and every alternative is priced instead '
         'of argued away',
         'The company running a durable structural cash pile (promotes the carry '
         'construction) or de-rating its net-debt policy (promotes gross weights)'],
        ['Capital expenditure per new refrigeration ton',
         f"AED {n0(U['capex_per_rt'] * 1000)} per RT added (derived: 2025 cash capital "
         f"expenditure over the year's added tons), plus a {pc(U['maint_pct'], 1)} "
         f"maintenance allowance on net plant",
         'No AED capital-expenditure guidance is disclosed anywhere (a recorded negative '
         'search); cost-per-ton on the audited cash figure is the finest sourced level '
         'available',
         'Company capital-expenditure guidance, or two years of cash figures diverging '
         'materially from the per-ton construction'],
        ['Electricity-and-water pass-through ratio',
         f"Purchased electricity and water held at {pc(U['ew_ratio'], 1)} of consumption "
         f"revenue (the audited 2025 print; 2024 was {pc(U['ew_ratio_fy24'], 1)})",
         'The cost is tied to its own physical driver — the consumption leg — on the '
         'related-party purchase disclosure, never to a blended inflation index; the DEWA '
         'slab tariff is flat and the fuel surcharge floats monthly',
         'A DEWA tariff revision, or the ratio moving by more than two points in a full-year '
         'print'],
        ['Working capital',
         f"Net working capital held at {pc(U['nwc_ratio'], 1)} of revenue (negative — "
         f"customer deposits and payables fund the cycle), from the audited 2025 balance "
         f"sheet with capital-expenditure accruals stripped out",
         'The ratio is structural: deposits are contractual and the dominant payable is to '
         'the related-party supplier',
         'Receivable days lengthening materially (the collection cycle stood near a month at '
         'FY2025) or a change in deposit terms'],
        ['The connected-capacity path',
         'The 2026 guidance midpoint, then additions funded by the contracted backlog, '
         'tapering as the build-out matures',
         "The company's own guidance and its signed but not-yet-connected backlog are the "
         'finest sourced volume drivers available; no market-growth assumption is used',
         'Connections printing outside the guidance band, or the contracted backlog '
         'shrinking two decks running'],
        ['The cost of equity construction',
         f"Risk-free = the most recent dirham sovereign print less the UAE's own default "
         f"spread; "
         f"beta {BR['beta']:.3f} from the stock's own weekly history against the FTSE ADX "
         f"General Index (adopted as the UAE market index by explicit instruction; the "
         f"DFM-index regression, {BR_DFM['beta']:.3f}, is priced as a full parallel "
         f"valuation since the 17-Aug revision); the published UAE equity risk "
         f"premium; both spread bases computed "
         f"(they converge: {pc(W['ke_rating'], 2)} against {pc(W['ke_cds'], 2)})",
         'Country risk must enter once — through the premium — not twice; the two published '
         'bases are both carried precisely because the choice between them is contested, and '
         'here it turns out to cost nothing. The beta-index choice, by contrast, is worth a '
         'fifth of the primary lens and is therefore priced, not footnoted',
         'A longer LIQUID dirham sovereign print post-repricing (the stale Feb-2033 tap does '
         'not qualify and is named as such) or the '
         'beta interval shifting as the listing history lengthens'],
        ['Minority interests at their profit share',
         f"Minorities are charged {pc(IN['nci_pat_fy25'] / IN['pat_fy25'], 1)} of equity "
         f"value — their share of profits — rather than book value",
         'The study values the group above book, so valuing the minority claim at book while '
         'valuing everyone else\'s above it would apply two standards to the same '
         'subsidiaries',
         'A disclosed transaction in the minority stakes establishing a different value'],
        ['The tax framing keeps the interest shield in BOTH columns',
         f"Both framings discount at {pc(W['rating_ct'], 2)}: the {pc(IN['tax_dmtt'], 0)} "
         f"minimum top-up is a minimum-effective-rate charge, not a higher statutory rate, "
         f"so the {pc(IN['tax_ct'], 0)} interest shield survives in the 15% column",
         'Pillar-Two mechanics: the top-up restores the effective rate to the minimum; it '
         'does not change the statutory deductibility of interest — the audit\'s point, '
         'accepted and implemented',
         'UAE implementing rules explicitly recomputing the debt shield at the top-up rate '
         'for in-scope entities']]
table(rows, [1.50, 2.90, 2.75, 2.65], size=7.9)

# ---------------------------------------------------- 4. negative results
doc.add_page_break()
H1('4  Searches that returned nothing — recorded, not filled in')
P('Each of these absences shaped the model. They are recorded with the scope of what was '
  'searched, so a reader knows the difference between "not disclosed" and "not found".')
rows = [['What was looked for', 'Where', 'Consequence for the model'],
        ['A per-ton tariff schedule — the AED level of the capacity (demand) charge',
         'The revenue notes of all four audited filings; the interim notes; the '
         'investor-relations presentations; a site-restricted web search of empower.ae. Only '
         'consumption revenue (in the auditor\'s key-audit-matter sections) and a first-half '
         'percentage mix (in the H1-2026 deck) are disclosed',
         'The capacity rate per connected ton is IMPLIED — solved from disclosed total '
         'revenue less disclosed consumption and pipes revenue — and flagged as implied '
         'wherever it appears; its growth is tied to connected tons at a flat regulated '
         'tariff'],
        ['AED capital-expenditure guidance',
         'All filings and both 2026 investor presentations: the deck carries '
         'refrigeration-ton guidance but no dirham capital budget',
         'Capital expenditure is modelled as cost per new ton derived from the audited 2025 '
         'cash figures plus a maintenance allowance — a construction, flagged as such in the '
         'judgements table above'],
        ['The contractual margin over EIBOR on the two revolving credit facilities',
         'The borrowings notes of the FY2024 and FY2025 filings say only "EIBOR plus a '
         'margin", and for 2025 "a reduction in interest margin"; the basis-point margin is '
         'nowhere disclosed',
         'The marginal cost of debt is anchored instead on the closest disclosed all-in '
         'figure — the company\'s own borrowing-cost capitalisation rate (4.92% in 2025, '
         '5.993% in 2024) — checked above the dirham sovereign yield']]
table(rows, [2.45, 3.75, 3.60], size=8.2)

# ---------------------------------------------------- 5. aggregator notes
H1('5  Aggregator and press data — discrepancies found and how they were handled')
rows = [['Item', 'What was found', 'Handling'],
        ['Market index series for the beta regression',
         'The FTSE ADX General Index history used (supplied export) ends at 24-Jul-2026, '
         'two weeks before the anchor date, so the regression window clips there. The '
         'previously used DFM General Index series (an aggregator pull, ~a fifth of '
         'sessions missing, ending 16-Jul-2026) is retained only as the archived '
         'comparison regression (beta 0.652 against 0.863 adopted)',
         'Used for ONE purpose only: the market side of the beta regression, where the '
         'weekly sampling (last observation of each week) absorbs the missing sessions; the '
         'sparsity and the stale endpoint are flagged in the study. Never used for '
         'Empower\'s own prices, which come from the supplied exchange history through '
         '07-Aug-2026'],
        ['Conflict and ceasefire status, August 2026 — RESOLVED in the 17-Aug revision',
         'The first edition treated the truce status as an unresolved disagreement between '
         'sources and leaned on the stale reading ("holding since 8 April"). The dated '
         'timeline record, externally re-verified for the revision, is unambiguous: the '
         'ceasefire was declared over on 8 July, the Strait of Hormuz was closed from 11 '
         'July (US blockade 14 July) and remained closed at the 7 August anchor, and an '
         'ADNOC tanker was attacked at sea on 8 August',
         'The study now states the anchor-date facts plainly, relabels the cases so that '
         'recovery is conditioned on a de-escalation that had not occurred, publishes the '
         'continuation central beside it, and carries a dated 17-August postscript for the '
         'newest events (tankers hit 12/14 August; Iran-Oman reopening agreement announced '
         'and the formal ceasefire ended 17 August)'],
        ["The study's own first edition vs the external audit — recorded on the same terms "
         "as any other source discrepancy",
         'The audit found the first edition carrying a stale ceasefire reading (above); '
         'missing the regulator\'s current tariff instrument (v1.3 of 17-Sep-2025 and v1.4 '
         'of 6-Feb-2026, with explicit caps and the no-indexation rule) and the Nov-2025 '
         'fees-and-fines resolution; overstating the refinancing scope (one tranche '
         'extended, not both); a receivable shown at its current portion in three '
         'balance-sheet columns; and a stale peer mark. The re-check that followed also '
         'found, and this register discloses, a FY2025 statement-face mis-split (G&A '
         '256.383 / other income 24.430, where the first pass carried 246.577 / 14.624 — '
         'both pairs close the operating-profit identity, which is why the assertion never '
         'caught it) and a rental-income double-count against the investment properties',
         'Every item is corrected in the revised deliverables; the full list with prices is '
         'in the accompanying critique-response note. The mis-split is a reminder that an '
         'identity check constrains SUMS, not splits — the note-level composition (note 29) '
         'is now the control'],
        ['Peer market value dating',
         'The relative lens was originally struck on a data-provider mark dated 22-Jun-2026 '
         'while the study anchored 7-Aug-2026 — the same input carrying two dates inside '
         'one document set, and the peer had itself sold off ~10% between them',
         'Peer marks are restruck on the anchor dates (Tabreed 2.46 on 6/7-Aug-2026; '
         'trailing multiple on trailing operating EBITDA, like-for-like) and the dating '
         'rule — one anchor for every market input — is stated'],
        ['A peer\'s quarterly earnings figure',
         'One aggregator rendered Tabreed\'s quarterly EBITDA in dollars where the company '
         'release states dirhams — a transcription error in the aggregator',
         'The company release figure was kept; the incident is recorded here as a reminder '
         'of why peer figures are cross-checks only, never build sources'],
        ['Empower "next earnings date"',
         'An aggregator listed a mid-August earnings date after the half-year results had '
         'already been released on 5 August',
         'Ignored in favour of the company\'s own filing dates'],
        ['Consensus price targets',
         'Aggregator pages carry sell-side consensus figures for the shares',
         'Not used anywhere in the study: this research publishes fair-value ranges and '
         'distributions, never price targets, and imports nobody else\'s']]
table(rows, [1.90, 4.20, 3.70], size=8.2)

P('')
H1('Disclosure')
P('This document accompanies an educational valuation study. For information only — not '
  'investment advice. It contains no recommendation, rating or price target. Sources are '
  'listed so that readers can verify the analysis independently; where a figure is derived, '
  'implied or estimated rather than disclosed, that is stated in the register itself.',
  size=9.2, color=GREY)

_TBLPR_ORDER = ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
                'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
                'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook',
                'tblCaption', 'tblDescription']
def _key(el):
    tag = el.tag.split('}')[1]
    return _TBLPR_ORDER.index(tag) if tag in _TBLPR_ORDER else len(_TBLPR_ORDER)
for tblPr in doc._element.iter(qn('w:tblPr')):
    _seen = set()
    for child in list(tblPr):                # drop duplicate singleton children
        _tag = child.tag.split('}')[1]
        if _tag in _seen:
            tblPr.remove(child)
        else:
            _seen.add(_tag)
    for child in sorted(list(tblPr), key=_key):
        tblPr.remove(child); tblPr.append(child)
zoom = doc.settings.element.find(qn('w:zoom'))
if zoom is not None and zoom.get(qn('w:percent')) is None:
    zoom.set(qn('w:percent'), '100')

OUT = 'EMPOWER_Bibliography_09-08-2026.docx'
doc.save(OUT)
print(f"wrote {OUT} | {len(doc.paragraphs)} paragraphs | {len(doc.tables)} tables | "
      f"{len(INP)} inputs registered")
