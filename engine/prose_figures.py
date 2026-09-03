"""Every percentage and multiple a reader sees, reconciled against the model's own numbers.

WHY THIS IS SHARED RATHER THAN COPIED
    engine/egch_study/prose_check.py was the only implementation of this check in the book,
    and porting it study by study would have produced twenty-four hand-maintained rendering
    sets with twenty-four different holes — which is precisely the failure L-084 was
    registered for this same day, on the sweeps, and [R-ENF-01 EXTENDED] the same day on the
    external-reader scrubs. A SHARED INSTRUMENT BEATS A GOOD LOCAL ONE EVEN WHEN THE LOCAL
    ONE IS BETTER WRITTEN, so the mechanism lives here and each study declares only what is
    genuinely its own: which documents a reader receives, and which figures may legitimately
    be quoted against something other than a model output.

WHAT IT DOES
    Flattens every number in the study's committed JSON into every plausible rendering — as
    a percentage, as a complement, as a multiple, at nought to three decimals — and then
    reads the DELIVERED documents and requires every figure carrying a per-cent or a
    multiple to match one. A figure with no computed counterpart FAILS.

    Bare integers are deliberately NOT matched. A page number, a note reference, a year and
    a count of sections are all bare integers, and a check that cries wolf is one everyone
    learns to ignore. The counts that bit AMOC this week ("four consecutive filed periods"
    against five) are a different instrument's job: they are caught by computing the count
    in the builder, which is what that study now does.

THE STANDING RULE FOR A FALSE POSITIVE, inherited verbatim from EGCH's version because it
is the whole discipline: A FALSE POSITIVE IS FIXED BY WIDENING THE RENDERING SET, NEVER BY
DELETING THE FIGURE FROM THE STUDY. If a figure is real and the model cannot produce it,
the model is what is missing.
"""
import glob
import json
import os
import re

# A figure a reader sees carries a unit. Anything bare is somebody else's check.
NUM = re.compile(r"(?<![\w.])(-?\d{1,3}(?:,\d{3})*(?:\.\d+)?|-?\d+(?:\.\d+)?)\s*"
                 r"(per cent|percent|%|x\b|times)")

# Constants a reader legitimately sees that no model produces: round shares, decile
# boundaries, the halves and quarters a sentence uses to describe a proportion.
STRUCTURAL = (0, 1, 2, 3, 4, 5, 10, 15, 20, 25, 30, 33, 40, 50, 60, 66, 67, 70, 75, 80,
              90, 95, 100, 0.5, 1.5, 2.5, 7.5, 12.5, 22.5)

# The input register printed verbatim is the SOURCE of the numbers, not a claim about them,
# so a table whose header is the register's own is skipped. Studies name their own header.
REGISTER_HEADERS = (
    ('Input', 'Value', 'Unit', 'Date', 'Source and construction'),
    ('Input', 'Value', 'Unit', 'Source, date and construction'),
)


def _flat(x, out):
    if isinstance(x, dict):
        for v in x.values():
            _flat(v, out)
    elif isinstance(x, (list, tuple)):
        for v in x:
            _flat(v, out)
    elif isinstance(x, (int, float)) and not isinstance(x, bool):
        out.append(float(x))


def numbers_from(study_dir, files=None, skip=()):
    """Every number in the study's committed JSON. `files` narrows it; `skip` excludes."""
    vals = []
    paths = ([os.path.join(study_dir, f) for f in files] if files
             else sorted(glob.glob(os.path.join(study_dir, '*.json'))))
    for p in paths:
        if os.path.basename(p) in skip or not os.path.exists(p):
            continue
        try:
            _flat(json.load(open(p, encoding='utf-8')), vals)
        except Exception:                                               # noqa: BLE001
            continue
    return vals


def rendering_set(values, extra=(), structural=STRUCTURAL):
    """Every plausible way a model number reaches a page."""
    r = set()
    # FOUR decimals, not three: AMOC prints a debt weight at pc(..., 4) — "0.0966% of the
    # capital structure" — and a three-decimal set rounds that to 0.097 and reports the
    # study's own correct figure as unmatched. The rendering set must reach as far as the
    # widest format any builder uses, and the honest way to find that out is to be told by
    # a false positive.
    for v in list(values) + list(extra):
        for x in (v, 1 - v, v - 1, -v, 100 * v, 100 * (1 - v), 100 * (v - 1), v / 100):
            for d in (0, 1, 2, 3, 4):
                r.add(round(x, d))
    for x in structural:
        for d in (0, 1, 2):
            r.add(round(x, d))
    return r


def relative_to(values, denominators):
    """A figure quoted as a distance — a level against a close, a lens against spot.

    EGCH's version learned this the hard way and the lesson is kept here: the TECHNICAL
    read is computed on the price library's last session while the study is struck on the
    latest known price, so a level's distance must be measured against the close the
    document itself states. A checker that models what a document OUGHT to divide by rather
    than reading what it does is checking a different document.
    """
    out = []
    for v in values:
        for d in denominators:
            if d:
                out.append(v / d - 1)
    return out


def texts_of(path):
    """Every paragraph and cell a reader sees, minus the input register printed verbatim."""
    import docx
    d = docx.Document(path)
    out = [p.text for p in d.paragraphs]
    for t in d.tables:
        hdr = tuple(c.text.strip() for c in t.rows[0].cells) if t.rows else ()
        if hdr in REGISTER_HEADERS:
            continue
        for row in t.rows:
            for c in row.cells:
                out.append(c.text)
    return out


def ratios_against(values, denominators):
    """A figure quoted as a GAP: "the lens reads X% against the price".

    A value and a price are both in the committed record; their ratio is not, and it is one
    of the commonest shapes on the page — every lens, every expert and every scenario is
    quoted as a distance from spot. The first run of this module against AMOC reported
    eleven of eighteen unmatched figures for exactly this reason, all of them correct.
    Distinct from relative_to() only in intent: that one is for a level against a close on
    the technical read's own clock, this one for a value against the price it is compared
    with.
    """
    out = []
    for v in values:
        for d in denominators:
            if d:
                out.append(v / d - 1)
                out.append(v / d)
    return out


def check(docs, render, label=''):
    """(checked, problems). A problem is a figure with no computed counterpart."""
    problems, checked = [], 0
    for f in docs:
        name = os.path.basename(f)
        for txt in texts_of(f):
            for m in NUM.finditer(txt):
                raw = m.group(1).replace(',', '')
                dec = len(raw.split('.')[1]) if '.' in raw else 0
                checked += 1
                if round(float(raw), dec) not in render:
                    lo = max(0, m.start() - 60)
                    problems.append('%s: %r in: ...%s...'
                                    % (name, m.group(0), txt[lo:m.end() + 40]))
    return checked, problems


def report(checked, problems, label=''):
    print('prose figures checked: %d; unmatched: %d%s'
          % (checked, len(problems), (' [%s]' % label) if label else ''))
    for p in problems:
        print('  !', p)
    return 1 if problems else 0

# ---------------------------------------------------------------------------------------
# A DIRECTION WORD IS A CLAIM AND IT IS CHECKED AGAINST THE SIGN BESIDE IT.
#
# The figures in this file are computed; the WORDS around them are typed, and a typed word
# does not look like a figure. ARCC shipped "+1.8% below the simple annualisation" — the
# ratio computed correctly and the direction word said the opposite, so the sentence
# understated the model's own forecast against the benchmark it was being challenged on.
# MODON shipped "AED 2.50, -12% above the market".
#
# Measured across the book at 41 delivered documents in their latest editions, this finds
# ONE contradiction and no false positives, once two innocent constructions are excluded:
# a temporal "over" ("fell about -2.4% over the same span") and a RANGE dash ("80-85%
# above the 2024 average"), both of which fired against the first draft. Bare "over" is
# therefore dropped from the upward set — above, more than, higher than, ahead of, exceeds
# and greater than cover it — and a sign preceded by a digit is a range, not a minus.
_UP = r'above|more than|ahead of|higher than|exceeds?|greater than'
_DOWN = r'below|under|less than|beneath|short of|lower than|behind'
_SIGN_UP = re.compile(r'(?<![\d])[+]\s?\d[\d,]*\.?\d*\s?%\s+(?:\w+\s+){0,2}(' + _DOWN + r')\b',
                      re.I)
_SIGN_DOWN = re.compile(r'(?<![\d])[-\u2212\u2013]\s?\d[\d,]*\.?\d*\s?%\s+(?:\w+\s+){0,2}('
                        + _UP + r')\b', re.I)


def sign_word_conflicts(texts):
    """Every place a signed percentage is contradicted by the direction word beside it."""
    out = []
    for t in texts:
        for rx in (_SIGN_UP, _SIGN_DOWN):
            for m in rx.finditer(t or ''):
                i = max(0, m.start() - 60)
                out.append((m.group(0).strip(), (t[i:m.end() + 25]).strip()))
    return out


def document_texts(path):
    """Every string a reader sees in a .docx — paragraphs and table cells alike."""
    import docx
    d = docx.Document(path)
    texts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                texts.append(c.text)
    return texts
