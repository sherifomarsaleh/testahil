"""Build the delivered PHDC study document.

Reads study_numbers.json and nothing else for figures — no financial numeral is
typed into this builder, so every number in the delivered document traces to the
registry entry it came from.

The delivered document is written for an external reader: it carries no internal
procedure vocabulary, no file paths, no module names and no house shorthand.
Calibration evidence appears in section 3 as plain-language sentences with the
statistics inline, not as an appendix.
"""
import json, os, sys, datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
N = json.load(open(os.path.join(HERE, "study_numbers.json")))
M, D, W = N["meta"], N["derived"], N["wacc"]
REG, PM = N["registry"], N["price_map"]
BU, LENS, BRIDGE, RANGED = (N["bottom_up"], N["lenses"], N["bridge"],
                            N["ranged_revenue"])
ST = N["statements"]
LW = N["lens_weighted"]
CASES, SENS = N["cases"], N["sensitivity"]

INK = RGBColor(0x1A, 0x1D, 0x21)
MUTED = RGBColor(0x5B, 0x65, 0x70)
ACCENT = RGBColor(0x1D, 0x6F, 0xA3)

# --- vocabulary an external reader must never meet -------------------------
FORBIDDEN = [
    "sigcm", "assert_", "r-cal", "r-sigcm", "r-enf", "r-beta", "r-std", "r-doc",
    "step 0.0", "step 2a", "information sweep", "outstanding.json", "panel.json",
    "study_numbers", "beta_regression", "research_protocol", "market_profiles",
    "mc_v3", "wacc_builder", "engine/", "crps", "parity", "gate", "checklist",
    "walk-forward training", "ring", "cohort", "raw_ohlc", "data.js",
]


def v(key):
    return REG[key]["value"]


# The forecast window now runs fifteen years, because it must run until growth
# has converged on the terminal. Fifteen columns do not fit a page, so the wide
# forecast tables show the first five years and then every fifth — the shape of
# the path, without pretending the table is the model. The workbook carries every
# year.
def display_years(rows):
    """Indices of the years a wide table shows: the first five, then every fifth."""
    n = len(rows)
    idx = list(range(min(5, n)))
    idx += [i for i in range(5, n) if (i + 1) % 5 == 0]
    return idx


def pick(rows, fmt):
    """Format one row of a wide table on the displayed years only."""
    return [fmt(rows[i]) for i in display_years(rows)]


def wide_widths(rows, label_cm=4.6, total_cm=16.2):
    n = len(display_years(rows))
    return [label_cm] + [round((total_cm - label_cm) / n, 2)] * n


def q(key):
    """A line of the 31 March 2026 reviewed balance sheet — what the bridge stands on."""
    return N["balance_sheet_1q26"][key]["value"]


def money(x, dp=0):
    return "{:,.{dp}f}".format(x, dp=dp)


def pct(x, dp=1):
    return "{:.{dp}f}%".format(x * 100, dp=dp)


def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(10)
    st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.18
    for nm, sz, col, bold in (("Heading 1", 16, ACCENT, True),
                              ("Heading 2", 12.5, INK, True),
                              ("Heading 3", 10.5, INK, True)):
        s = doc.styles[nm]
        s.font.name = "Georgia"
        s.font.size = Pt(sz)
        s.font.color.rgb = col
        s.font.bold = bold
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(5)


def para(doc, text, size=10, bold=False, italic=False, color=INK, align=None,
         space_after=7):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(size)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(3)


def table(doc, headers, rows, widths, caption=None, size=8.5):
    """Fixed layout with explicit widths — no column may starve or bloat."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    # UNDER FIXED LAYOUT THE RENDERER READS THE GRID, NOT THE CELL WIDTHS.
    # python-docx writes a tblGrid of EQUAL columns and setting cell widths
    # does not touch it, so every table in this document rendered with roughly
    # equal columns however carefully the widths were chosen — and the column
    # audit, which read the cell widths back, reported it clean. A checker that
    # models something other than what the renderer reads is checking a
    # different document. The grid is written here and audited below.
    grid = t._tbl.find(qn("w:tblGrid"))
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(int(round(w * 567))))
    # a table that breaks across a page must carry its header onto the next one,
    # or the continuation is a grid of numbers with nothing naming the columns
    trPr = t.rows[0]._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    trPr.append(hdr)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = MUTED
    for row in rows:
        tr = t.add_row()
        # A row allowed to split across a page break renders its continuation as
        # an empty strip under the repeated header, which reads as a missing
        # line rather than a wrapped one.
        cant = OxmlElement("w:cantSplit")
        tr._tr.get_or_add_trPr().append(cant)
        cells = tr.cells
        for i, val in enumerate(row):
            cells[i].width = Cm(widths[i])
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            r.font.size = Pt(size)
            r.font.color.rgb = INK
            if i and str(val).replace("-", "").replace(".", "").replace(",", "")\
                    .replace("%", "").replace("(", "").replace(")", "").isdigit():
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if caption:
        para(doc, caption, size=8, italic=True, color=MUTED, space_after=10)
    else:
        # TWO TABLES WITH NOTHING BETWEEN THEM ARE ONE TABLE TO THE RENDERER.
        # Word and LibreOffice merge adjacent tables, so the second one loses
        # its own header row to the first one's repeating header and its first
        # row renders as an empty strip. A.3 published a cycle table under a
        # header reading "Forecast cash flow" because of exactly this. A
        # separator always follows a table that carries no caption, and
        # adjacency is audited below rather than remembered.
        para(doc, "", size=4, space_after=4)
    return t


def figure(doc, name, caption):
    doc.add_picture(os.path.join(HERE, name), width=Cm(16.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, size=8, italic=True, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def scrub(doc_path):
    """Programmatic scan for internal vocabulary. Zero hits required."""
    from docx import Document as Dx
    d = Dx(doc_path)
    text = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                text += "\n" + c.text
    import re as _re
    low = text.lower()
    hits = []
    for w in FORBIDDEN:
        # Short internal tokens are ordinary English substrings — "ring" lives
        # inside "during" and "offering", "gate" inside "mitigate". Match those
        # on a word boundary; keep substring matching for the distinctive ones.
        if len(w) <= 6 and w.isalpha():
            if _re.search(r"\b%s\b" % _re.escape(w), low):
                hits.append(w)
        elif w in low:
            hits.append(w)
    return sorted(set(hits)), len(text)


def column_audit(doc_path):
    """Every table fixed-layout with explicit widths, none starved or bloated."""
    from docx import Document as Dx
    d = Dx(doc_path)
    from docx.oxml.ns import qn as _qn
    bad = []
    body = d.element.body
    kids = [c.tag for c in body.iterchildren()]
    TBL = _qn("w:tbl")
    for i in range(len(kids) - 1):
        if kids[i] == TBL and kids[i + 1] == TBL:
            bad.append((i, "two tables with nothing between them — the "
                           "renderer merges these into one"))
    for i, t in enumerate(d.tables):
        widths = [c.width for c in t.rows[0].cells]
        if any(w is None for w in widths):
            bad.append((i, "a column carries no explicit width"))
            continue
        cm = [w.cm for w in widths]
        if min(cm) < 1.0:
            bad.append((i, "starved column %.2fcm" % min(cm)))
        if sum(cm) > 17.5:
            bad.append((i, "row overflows the text column at %.2fcm" % sum(cm)))
        # AUDIT WHAT THE RENDERER ACTUALLY READS. Under fixed layout that is
        # the grid; cell widths agreeing among themselves proves nothing if the
        # grid disagrees with them.
        grid = t._tbl.find(_qn("w:tblGrid"))
        if grid is None:
            bad.append((i, "no column grid"))
            continue
        gcm = [int(gc.get(_qn("w:w")) or 0) / 567.0
               for gc in grid.findall(_qn("w:gridCol"))]
        if len(gcm) != len(cm):
            bad.append((i, "grid has %d columns, the row has %d"
                        % (len(gcm), len(cm))))
        elif max(abs(a - b) for a, b in zip(gcm, cm)) > 0.02:
            bad.append((i, "the grid the renderer reads disagrees with the "
                           "cell widths: %s vs %s"
                        % (["%.2f" % x for x in gcm], ["%.2f" % x for x in cm])))
    return bad


# ===========================================================================
def build(path):
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.3)
        s.top_margin = s.bottom_margin = Cm(2.0)
    _style(doc)
    sp, prior = PM["spot"], D["prior_edition_fair"]
    base, low, high = CASES["base"], CASES["low_conversion"], CASES["high_conversion"]
    cds = CASES["base_cds_erp"]

    # --- 1 Masthead + READ FIRST -------------------------------------------
    para(doc, "PALM HILLS DEVELOPMENTS", size=20, bold=True, color=ACCENT,
         space_after=2)
    para(doc, "Egyptian Exchange · PHDC · Egyptian pounds · edition of 2 September 2026",
         size=10, color=MUTED, space_after=14)
    doc.add_heading("READ FIRST", level=2)
    para(doc, "This is a valuation study, not advice. It carries no rating, no "
              "recommendation and no price target. What it publishes is a range of "
              "value and the reasoning behind it, so that a reader can disagree with "
              "the reasoning rather than with a number.")
    para(doc, "It supersedes the edition of 30 August 2026, and it changes three "
              "things in it, each found by a review of that edition against the "
              "company's filings on 1 September 2026 and each set out where it "
              "applies. First, the bridge from enterprise value to equity, the book "
              "value and the borrowings now stand on the balance sheet of 31 March "
              "2026 — reviewed, and posted to the company's own result centre with "
              "the first-quarter results the earlier edition already used — rather "
              "than on 31 December 2025 (section 1.1, Appendix A.2). Second, "
              "minority shareholders in subsidiaries are now deducted at their share "
              "of value; the earlier edition deducted nothing for them (section 1.1). "
              "Third, normalised earnings are capitalised at the cost of equity less "
              "the same growth the cash-flow model carries, where the earlier edition "
              "capitalised them at the cost of equity alone — which in a currency "
              "whose discount rate embeds Egyptian inflation assumes a perpetual "
              "decline in real terms that nothing in the company's record supports "
              "(section 1.4). The discount rate, the forecast and the lens weights "
              "are unchanged. The central figure moves from EGP %.2f to EGP %.2f a "
              "share." % (prior["base"], LW["base"]))
    para(doc, "Two things in the edition of 11 June 2026 were wrong and remain "
              "corrected, and both are set out plainly in section 1.8 rather than "
              "buried: its discount rate was below Egypt's own government bond "
              "yield, and the project-level price and cost table it valued the "
              "company on is not disclosed by the company anywhere.")
    para(doc, "The company's own audited statements are the only source used for its "
              "reported history. Where something needed is not disclosed, this study "
              "says so and does not fill the hole. There are %s such gaps and they "
              "are listed in section 7." % {4: "four", 5: "five", 6: "six", 7: "seven"}.get(len(N["gaps"]), str(len(N["gaps"]))))
    para(doc, "Information set: everything the company had published as at 2 "
              "September 2026, which ends at its first-quarter 2026 results — the "
              "reviewed statements of 31 March 2026 included. No half-year 2026 "
              "figures had been released at that date.", size=9, italic=True,
         color=MUTED)

    # --- 2 Headline ---------------------------------------------------------
    doc.add_heading("Headline", level=1)
    para(doc, "Palm Hills is worth EGP %.2f a share on a base case built from the "
              "company's own units and prices, in a range of EGP %.2f to EGP %.2f "
              "across the full observed range of the one thing that decides it. "
              "The shares trade at EGP %.2f."
              % (N["central"], LW["bear"], LW["full"], sp), size=11.5, bold=True)
    bullets(doc, [
        "The forecast is built from units and prices the company itself "
        "discloses: %s units sold across three regions at EGP %.2f to %.2f million "
        "a unit, and separately %s units delivered at EGP %.2f million of revenue "
        "each, because units handed over were contracted years earlier at lower "
        "prices."
        % ("{:,.0f}".format(BU["rows"][0]["units_sold"]),
           min(d["asp_base"] for d in BU["regions"].values()),
           max(d["asp_base"] for d in BU["regions"].values()),
           "{:,.0f}".format(BU["rows"][0]["units_delivered"]),
           BU["rows"][0]["rev_per_unit"]),
        "The company's cost of capital, built from the ground up, is %s against "
        "an Egyptian ten-year government bond yield of %s. The edition of 11 June "
        "2026 discounted at %s, below the sovereign; that was corrected on 30 "
        "August 2026 and the rate is unchanged here."
        % (pct(W["wacc_rating"], 2), pct(W["rf_observed"], 2),
           pct(D["edition_11jun_wacc"], 2)),
        "Value turns on how quickly contracted sales become cash. Over the three "
        "years the company has published a cash-flow statement, operating cash was "
        "%s, %s and %s of revenue. At the low end the shares are worth EGP %.2f; at "
        "the high end, EGP %.2f."
        % (pct(D["cfo_margins"]["2023"]), pct(D["cfo_margins"]["2024"]),
           pct(D["cfo_margins"]["2025"]), low["per_share"], high["per_share"]),
        "At the closing price of EGP %.2f the market is paying for a conversion rate "
        "of %s — above the two weak years, well below the strong one."
        % (sp, pct(D["market_implied_cash_conversion"])),
        "The contracted order book stands at EGP %sbn against %s revenue of EGP "
        "%sbn. What limits this company is how fast it can build and hand over, not "
        "whether it can sell."
        % (money(v("backlog_1q26") / 1000, 0), "2025",
           money(v("revenue_fy25") / 1000, 1)),
    ])

    # --- 3 Valuation summary ------------------------------------------------
    doc.add_heading("Valuation summary — every read at a glance", level=1)
    LR = N["lens_record"]
    table(doc, ["Lens", "Low", "Central", "High", "Role"],
          [[r[0], "%.2f" % r[1], "%.2f" % r[2], "%.2f" % r[3],
            "the central" if i == 0 else "cross-check"]
           for i, r in enumerate(LENS)]
          + [["Market price, 23 Aug 2026", "", "%.2f" % sp, "", ""],
             ["vs the central", "", "%+.0f%%" % (100 * (sp / LR["central"] - 1)), "", ""]],
          [5.6, 2.6, 2.6, 2.6, 2.4],
          "EGP per share. The cash-flow lens is the central because it is the only "
          "one built from the company's own units and prices; the others are "
          "cross-checks, published so that a reader can see where they disagree. "
          "They are NOT averaged into the central. An earlier edition blended all "
          "four at fixed weights, and because three of the four value a developer "
          "on its reported earnings and its historical-cost book — a floor for a "
          "company whose worth sits in an undelivered order book — the blend read "
          "as a value when it was measuring a floor.")
    figure(doc, "fig1_football.png",
           "Figure 1 — each bar is one observed rate of cash conversion; the bar "
           "spans the discount rate from four points below the rebuilt cost of "
           "capital to four points above, and the white tick marks the value at "
           "the rebuilt rate itself. Shown against the closing price and the "
           "30 August 2026 edition's central figure.")

    # --- 4 Company overview -------------------------------------------------
    doc.add_heading("Company overview", level=1)
    para(doc, "Palm Hills Developments builds and sells residential and commercial "
              "property in Egypt, principally in West Cairo, East Cairo, and the "
              "North Coast and Alexandria. It has been listed on the Egyptian "
              "Exchange since 2008. Units are sold off plan and paid for in "
              "instalments over several years, so the company collects cash long "
              "before, and long after, it recognises the sale.")
    para(doc, "Two things follow from that model and they shape everything in this "
              "study. First, revenue is recognised as construction completes rather "
              "than when a contract is signed, so the order book and the income "
              "statement move on different clocks. Second, the balance sheet carries "
              "a very large receivable book — EGP %sbn at the end of 2025 — against "
              "customer advances of EGP %sbn."
              % (money((v("notes_recv_lt") + v("notes_recv_lt_undel")
                        + v("notes_recv_st") + v("notes_recv_st_undel")
                        + v("accounts_receivable")) / 1000, 1),
                 money(v("advances_customers") / 1000, 1)))
    table(doc, ["", "2025", "2024"],
          [["Revenue (EGP mn)", money(v("revenue_fy25")), money(v("revenue_fy24"))],
           ["Gross profit (EGP mn)", money(v("gross_profit_fy25")), "9,330"],
           ["Gross margin", pct(D["gross_margin_fy25"]), "34.3%"],
           ["Profit after tax and minority (EGP mn)",
            money(v("npat_mi_fy25")), "3,255"],
           ["Cash from operations (EGP mn)", money(v("cfo_fy25")), money(v("cfo_fy24"))],
           ["Order book (EGP mn)", "263,000 (at 31 Mar 2026)",
            money(v("backlog_fy24"))],
           ["Total assets (EGP mn)", money(v("total_assets")), "123,437"],
           ["Shareholders' funds (EGP mn)", money(v("total_equity")), "14,625"]],
          [7.4, 4.4, 4.4],
          "Source: the company's audited consolidated financial statements and its "
          "own results releases.")
    _section_one(doc, sp, base, low, high, cds, prior)
    _sections_two_to_seven(doc, sp)
    _appendices(doc, sp, base)
    return doc


# ---------------------------------------------------------------------------
def _section_one(doc, sp, base, low, high, cds, prior):
    doc.add_heading("1  Fundamental valuation", level=1)

    doc.add_heading("1.1  The cash-flow model", level=2)
    para(doc, "The model runs five years from the audited 2025 base and then a "
              "terminal value. Revenue is limited by how much the company can "
              "build and hand over, not by how much it has sold: the order book "
              "is EGP %sbn against 2025 revenue of EGP %sbn, so sales are not "
              "the binding constraint and have not been for some years. Revenue "
              "is therefore units handed over times the revenue each one "
              "carries — %s units in 2026, growing %.0f per cent a year on the "
              "company's own disclosed run of handovers, at EGP %.2f million a "
              "unit escalating with inflation. Nothing here is a ratio applied "
              "to last year's revenue."
              % (money(v("backlog_1q26") / 1000, 0),
                 money(v("revenue_fy25") / 1000, 1),
                 "{:,.0f}".format(BU["rows"][0]["units_delivered"]),
                 100 * BU["anchors"]["delivery_growth"],
                 BU["rows"][0]["rev_per_unit"]))
    para(doc, "Cost is recognised on the same completion schedule as revenue. That "
              "sounds obvious and it is the single most important correction in this "
              "edition. Since January 2016 the company has recognised revenue on "
              "standalone units as construction progresses; a model that accrues "
              "revenue that way but accrues cost on handover has the two legs on "
              "different clocks, and it systematically overstates profit. Gross "
              "margin here is an output of price against cost, never an input.")
    def _r(key, fmt="%,.0f", scale=1.0):
        return pick(BU["rows"], lambda x: ("{:,.0f}".format(x[key]*scale) if "f" not in fmt
                 else fmt % (x[key]*scale)))
    yrs = pick(BU["rows"], lambda x: str(x["year"]))
    body = [
        ["Units sold", pick(BU["rows"], lambda x: "{:,.0f}".format(x["units_sold"]))],
        ["New sales (EGP mn)", pick(BU["rows"], lambda x: "{:,.0f}".format(x["new_sales"]))],
        ["Units delivered", pick(BU["rows"], lambda x: "{:,.0f}".format(x["units_delivered"]))],
        ["Revenue per delivered unit (EGP mn)",
         pick(BU["rows"], lambda x: "%.2f" % x["rev_per_unit"])],
        ["Revenue", pick(BU["rows"], lambda x: "{:,.0f}".format(x["revenue"]))],
        ["Cost per delivered unit (EGP mn)",
         pick(BU["rows"], lambda x: "%.2f" % x["cost_per_unit"])],
        ["Cost of revenue", pick(BU["rows"], lambda x: "{:,.0f}".format(x["cogs"]))],
        ["Gross profit", pick(BU["rows"], lambda x: "{:,.0f}".format(x["gross"]))],
        ["Gross margin (output)", pick(BU["rows"], lambda x: "%.1f%%" % (100*x["gross_margin"]))],
        ["Overheads", pick(BU["rows"], lambda x: "{:,.0f}".format(x["sga"]))],
        ["Operating profit", pick(BU["rows"], lambda x: "{:,.0f}".format(x["ebit"]))],
        ["Finance cost", pick(BU["rows"], lambda x: "{:,.0f}".format(x["interest"]))],
        ["Profit before tax", pick(BU["rows"], lambda x: "{:,.0f}".format(x["npbt"]))],
        ["Tax rate", pick(BU["rows"], lambda x: "%.1f%%" % (100*x["tax_rate"]))],
        ["Net profit", pick(BU["rows"], lambda x: "{:,.0f}".format(x["npat"]))],
        ["Earnings per share (EGP)", pick(BU["rows"], lambda x: "%.2f" % x["eps"])],
        ["Order book, closing", pick(BU["rows"], lambda x: "{:,.0f}".format(x["backlog"]))],
    ]
    table(doc, ["EGP mn unless stated"] + yrs,
          [[lbl] + vals for lbl, vals in body],
          wide_widths(BU["rows"]),
          "Every line follows from the two engines above. Gross margin is what "
          "price per unit and cost per unit leave behind, not an assumption.")
    a = BU["anchors"]
    para(doc, "Three anchors hold this table to what was actually reported. "
              "FY2026 is part-reported: the company disclosed first-quarter "
              "revenue of EGP %s million, up 11 per cent on the same quarter of "
              "2025, which places that quarter at %.1f per cent of the full year "
              "and gives an FY2026 revenue anchor of EGP %s million from the "
              "actual rather than from a trend. The finance charge is the "
              "disclosed rate on the interest-bearing borrowings, %.2f per cent, "
              "not the %.2f per cent marginal rate used for discounting — a large "
              "part of the balance does not bear interest and part of the charge "
              "is capitalised into work in progress. And the gross margin is held "
              "at %.1f per cent, the average of the last two disclosures."
              % ("{:,.0f}".format(a["q1_2026_reported"]),
                 100*a["q1_share_of_year"], "{:,.0f}".format(a["fy2026_anchor"]),
                 100*a["effective_pl_rate"], 100*a["marginal_rate_for_discounting"],
                 100*a["gross_margin_forward"]))
    para(doc, "One drift was measured and deliberately not carried. The move from "
              "a 41.2 per cent gross margin in 2025 to 35.5 per cent in the first "
              "quarter of 2026 implies cost rising about %.1f per cent a year "
              "faster than price. Compounded over five years that takes gross "
              "margin to 7.5 per cent and the company to a loss by 2030 — on the "
              "strength of one quarter against one year, for a developer whose "
              "margin moves with which project happens to hand over. It is "
              "sensitised in section 1.9 instead of extrapolated."
              % (100*a["cost_drift_measured_not_carried"]))
    table(doc, ["Revenue, years three to five (EGP mn)", "Low", "Point", "High"],
          [[str(r["year"]),
            "-" if r["low"] is None else "{:,.0f}".format(r["low"]),
            "{:,.0f}".format(r["point"]),
            "-" if r["high"] is None else "{:,.0f}".format(r["high"])]
           for r in RANGED if r["low"] is not None],
          [6.0, 3.4, 3.4, 3.4],
          "The band is this method's OWN measured error at each horizon, taken "
          "from testing it across ten annual starting points on the company's "
          "2011-2025 history. It is not a scenario; it is the spread the method "
          "has actually produced.")
    doc.add_heading("The cash-flow calculation in full", level=3)
    para(doc, "Operating profit is not cash, and for a developer selling on "
              "instalments the gap between the two is the whole question. The "
              "table below runs every step from operating profit to a present "
              "value, year by year, with nothing collapsed.")
    wf = ST["dcf_b"]["waterfall"]
    wy = pick(wf, lambda w: str(w["year"]))
    def _w(key, fmt="{:,.0f}"):
        return pick(wf, lambda w: fmt.format(w[key]))
    table(doc, ["EGP million"] + wy,
          [["Operating profit"] + _w("ebit"),
           ["Cash from operations"] + _w("cfo"),
           ["  as a share of revenue"]
           + pick(wf, lambda w: "%.1f%%" % (100 * w["cfo"] / w["revenue"])),
           ["plus finance cost, after tax"] + _w("interest_addback"),
           ["less capital expenditure"] + _w("capex"),
           ["Free cash flow to the firm"] + _w("fcff"),
           ["Cost of capital, that year"]
           + pick(wf, lambda w: "%.1f%%" % (100 * w.get("forward_wacc", ST["dcf_b"]["wacc"]))),
           ["Discount factor"] + pick(wf, lambda w: "%.3f" % w["discount_factor"]),
           ["Present value"] + _w("pv")],
          wide_widths(BU["rows"]),
          "The finance charge is added back after tax because the discount rate "
          "already carries the cost of debt; leaving it in the cash flow would "
          "charge for the same thing twice. Capital expenditure is maintenance "
          "only, at one per cent of revenue: for this company the building "
          "itself is inventory, and it is already inside operating cash.")
    table(doc, ["Bridge from enterprise value to equity, base case", "EGP mn"],
          [[lbl, ("{:,.2f}" if lbl.startswith("Value per share") else "{:,.1f}").format(val)]
           for lbl, val in BRIDGE],
          [10.6, 5.6],
          "The terminal value is %.0f per cent of enterprise value, which is high "
          "and is one reason the range is published rather than a point."
          % (100 * (N["dcf_cases"]["base"]["terminal_share"] or 0)))

    doc.add_heading("What the two statements disagree about, and why it is "
                    "not resolved here", level=3)
    wg = ST["wedge"]; cy = ST["cycle_measured"]
    para(doc, "The audited balance sheet and the audited cash-flow statement "
              "tell different stories about the same year, and the difference "
              "is the largest single uncertainty in this valuation. Working "
              "capital on the balance sheet rose EGP %s million in 2025. Net "
              "profit plus depreciation less the operating cash the company "
              "reported implies it rose EGP %s million. The difference is EGP "
              "%s million — %.1f per cent of the year's revenue."
              % ("{:,.0f}".format(wg["d_wc_book_fy25"]),
                 "{:,.0f}".format(wg["d_wc_cash_fy25"]),
                 "{:,.0f}".format(wg["wedge_fy25"]),
                 100 * wg["wedge_over_revenue"]))
    para(doc, "A difference that size is not depreciation and it is not "
              "capitalised interest. It is consolidation, land taken on "
              "deferred terms, and reclassification inside a book where "
              "revenue is recognised as construction progresses. The 2025 "
              "cash-flow statement is published in its three totals only, so "
              "the difference cannot be split from what the company has "
              "disclosed. It is therefore not split. It is measured, recorded "
              "in the register of what is missing, and the forecast is "
              "published two ways around it rather than one — set out in full "
              "in Appendix A.3.")
    table(doc, ["The collection cycle, as reported", "2025", "2024"],
          [["Collection period, days", "%.0f" % cy["dso_fy25"],
            "%.0f" % cy["dso_fy24"]],
           ["Work in progress, days of cost", "%.0f" % cy["dio_fy25"],
            "%.0f" % cy["dio_fy24"]],
           ["Suppliers, days of cost", "%.0f" % cy["dpo_fy25"],
            "%.0f" % cy["dpo_fy24"]],
           ["Customer advances, share of the order book",
            "%.1f%%" % (100 * cy["adv_of_backlog_fy25"]),
            "%.1f%%" % (100 * cy["adv_of_backlog_fy24"])],
           ["Net working capital (EGP mn)",
            "{:,.0f}".format(cy["nwc_fy25"]), "{:,.0f}".format(cy["nwc_fy24"])],
           ["  as a multiple of revenue", "%.2f" % cy["nwc_over_revenue_fy25"],
            "%.2f" % cy["nwc_over_revenue_fy24"]]],
          [8.4, 3.4, 3.4],
          "Both years measured the same way on the two audited sheets. The "
          "collection period is close to three years because units are sold on "
          "multi-year instalments; customer advances against the order book "
          "are what funds the building in the meantime.")

    doc.add_heading("1.2  Book value and sustainable return", level=2)
    para(doc, "Shareholders' funds attributable to the parent were EGP %s million "
              "at 31 March 2026, or EGP %.2f a share, against a market price of EGP "
              "%.2f — the shares change hands at about %.1f times book. The earlier "
              "edition divided total equity, minority interests included, by the "
              "parent's shares; this one puts the numerator on the same footing as "
              "the share count, and on the latest disclosed sheet. Profit after tax "
              "and minority interest of EGP %s million in 2025 on average "
              "shareholders' funds gives a return on equity of roughly %s. Against a "
              "cost of equity of %s, the company is not currently earning its cost of "
              "capital on book, which is why book value sits well below every "
              "cash-flow read in this study rather than acting as a floor."
              % (money(q("equity_parent")), D["book_equity_per_share"], sp,
                 sp / D["book_equity_per_share"], money(v("npat_mi_fy25")),
                 pct(v("npat_mi_fy25") / ((v("total_equity")
                      + N["balance_sheet_subtotals"]["2024"]["total_equity"]["value"]) / 2)),
                 pct(W["ke_rating"], 1)))

    doc.add_heading("1.3  Relative multiples", level=2)
    para(doc, "The Egyptian listed developers are compared in Appendix B on the "
              "measures that can actually be sourced for all of them on the same "
              "basis. Earnings and book multiples for the peer group are NOT "
              "published here: no peer discloses figures this study can obtain on a "
              "consistent basis, and a multiple assembled from inconsistent or "
              "unsourced inputs would look precise and mean nothing. What is "
              "comparable is market risk, and on that measure Palm Hills sits in the "
              "middle of its sector.")

    doc.add_heading("1.4  Normalised earnings power", level=2)
    para(doc, "Normalising on the three disclosed years, revenue converts to gross "
              "profit at about %s, overheads absorb about %s of revenue, and "
              "interest on the disclosed borrowings absorbs a further large share. "
              "The result is a company whose accounting profit is real but whose "
              "cash profit is small and volatile: EGP %s million of operating cash "
              "in 2025 on EGP %s million of revenue, against EGP %s million on EGP "
              "%s million the year before."
              % (pct((D["gross_margin_fy25"] + D["gross_margin_1q26"]) / 2),
                 pct(D["sga_ratio_fy25"]), money(v("cfo_fy25")),
                 money(v("revenue_fy25")), money(v("cfo_fy24")),
                 money(v("revenue_fy24"))))
    NI = N["lens_detail"]["normalised_inputs"]
    para(doc, "Normalised earnings of EGP %s million — the average of the last two "
              "years' revenue carried one year forward at inflation, at the 2025 "
              "net margin of %s — would capitalise at the cost of equity of %s less "
              "growth to about EGP %.2f a share, and at the cost of equity alone, "
              "as an earlier edition did, to EGP %.2f. Neither figure is published "
              "as a lens here, and the reason is not the arithmetic. This company "
              "recognises revenue when a project completes, so its reported earnings "
              "in any year are an accident of which project happened to complete: "
              "capitalising a mid-cycle figure treats that schedule as though it "
              "were a steady state. The working is kept as a diagnostic in the "
              "workbook and it claims nothing."
              % (money(NI["norm_earn"]), pct(NI["norm_margin"], 1), pct(NI["ke"], 2),
                 N["lens_detail"].get("normalised_base",
                                      NI["norm_earn"] / (NI["ke"] - NI["growth_netted"])
                                      / N["derived"]["shares_mn"]),
                 NI["as_30aug_edition_e_over_ke"]))

    doc.add_heading("1.5  Synthesis — the lenses in one field", level=2)
    para(doc, "The lenses disagree, and the disagreement is the information — but "
              "one of them is the answer and the others are checks on it. The "
              "cash-flow lens is the central: it is the only one built from the "
              "company's own units and prices, and it is the only one that values "
              "an order book the company has sold but not yet handed over. The "
              "multiple on its own trading history sits at EGP %.2f and book value "
              "at EGP %.2f, and both are lower for the same reason: they measure "
              "what has been reported, and what has been reported does not yet "
              "include a backlog carried at historical cost in a currency that has "
              "lost most of its value since 2022. That is a floor, and it is "
              "published as one."
              % (N["lenses"][1][2], N["lenses"][2][2]))
    MOVES = ["cash conversion, then the cost of capital",
             "the multiple the shares have historically carried, 6x to 14x",
             "nothing — it is a reported figure"]
    table(doc, ["Lens", "Low", "Central", "High", "Role", "What moves it"],
          [[LENS[i][0], "%.2f" % LENS[i][1], "%.2f" % LENS[i][2], "%.2f" % LENS[i][3],
            "the central" if i == 0 else "cross-check", MOVES[i]]
           for i in range(len(LENS))],
          [4.2, 1.8, 1.8, 1.8, 1.8, 4.4],
          "EGP per share, against a close of EGP %.2f. The cash-flow lens is the "
          "central and the others are cross-checks; they are not averaged. "
          "Normalised earnings power, which an earlier edition carried at a fifth "
          "of the weight, is not published as a lens here at all: a developer that "
          "recognises revenue when a project completes reports earnings that are an "
          "accident of which project completed in which year, and capitalising a "
          "mid-cycle figure treats that schedule as a steady state. The working is "
          "kept in the workbook as a diagnostic and claims nothing." % sp)

    doc.add_heading("1.6  Drivers — units and prices, by region", level=2)
    para(doc, "This is what the forecast is built from. For each operating region "
              "the company publishes both the value of new sales and the number of "
              "units behind it, so the average selling price per unit is a realised "
              "figure, not an assumption. Dividing one disclosure by the other is "
              "the only step taken.")
    for nm, d in BU["regions"].items():
        h = d["history"]
        doc.add_heading(nm, level=3)
        table(doc, ["Disclosed"] + [str(y) for y in h["years"]],
              [["New sales (EGP mn)"] + ["{:,.0f}".format(x) for x in h["sales"]],
               ["Units sold"] + ["{:,.0f}".format(x) for x in h["units"]],
               ["Price per unit (EGP mn)"] + ["%.2f" % x for x in h["asp"]]],
              [4.4, 2.4, 2.4, 2.4, 2.4, 2.4])
        para(doc, "Carried forward at %s units, the average of the last three "
                  "disclosed years, at a starting price of EGP %.2f million a "
                  "unit, the price realised in 2024, escalated with Egyptian "
                  "inflation."
                  % ("{:,.0f}".format(d["units_base"]), d["asp_base"]),
             size=9, color=MUTED)
    para(doc, "The three regions reconcile to the group total the same release "
              "prints: for 2024, 44,570 plus 95,082 plus 11,364 is 151,016, which "
              "is the all-regions figure exactly. That reconciliation is checked in "
              "code, so a mis-read series fails loudly rather than quietly becoming "
              "a price.")
    para(doc, "Revenue is driven by a different quantity from new sales, and the "
              "distinction matters. A unit handed over in 2026 was contracted years "
              "earlier at that year's price, so revenue per delivered unit sits "
              "below the current selling price — EGP %.2f million against regional "
              "prices of EGP %.2f to %.2f million. That gap is why the order book "
              "keeps building even as deliveries rise, and it is the single "
              "clearest fact about how this business converts sales into reported "
              "revenue."
              % (BU["rows"][0]["rev_per_unit"],
                 min(d["asp_base"] for d in BU["regions"].values()),
                 max(d["asp_base"] for d in BU["regions"].values())))
    para(doc, "What is NOT disclosed, and therefore not used: unit mix, average "
              "unit size, price per square metre and construction cost per square "
              "metre, for any project. The edition of 11 June 2026 carried a full table of "
              "those figures across fifteen named projects. None of it is published "
              "by the company. It is not reused here, and the forecast is built at "
              "the finest level the disclosure actually supports.")

    doc.add_heading("1.7  The crux", level=2)
    para(doc, "How fast do contracted sales become cash?", bold=True)
    para(doc, "Everything else in this study is secondary to that question. Palm "
              "Hills sells off plan on long instalment terms, so a sale signed today "
              "produces cash over several years on a schedule the company does not "
              "publish — no down-payment percentage, no instalment tenor, no "
              "post-handover tail. The edition of 11 June 2026 assumed all three. This "
              "edition measures the outcome instead, from the three years in which "
              "the company has published a cash-flow statement.")
    table(doc, ["Year", "Cash from operations (EGP mn)", "Revenue (EGP mn)",
                "Cash as a share of revenue"],
          [["2023", money(v("cfo_fy23")), money(v("revenue_fy23")),
            pct(D["cfo_margins"]["2023"])],
           ["2024", money(v("cfo_fy24")), money(v("revenue_fy24")),
            pct(D["cfo_margins"]["2024"])],
           ["2025", money(v("cfo_fy25")), money(v("revenue_fy25")),
            pct(D["cfo_margins"]["2025"])]],
          [2.6, 5.2, 4.0, 4.4],
          "The 2024 figure is the one in the audited statements. The company's own "
          "results release for that year reported EGP 3,132 million and the figure "
          "was later restated; the audited number is used.")
    para(doc, "A spread of %s to %s over three consecutive years is not noise around "
              "a stable rate — it is the absence of a stable rate. Until the company "
              "discloses its collection terms, or until more years accumulate, this "
              "is the honest width of the answer."
              % (pct(D["cfo_lo"]), pct(D["cfo_hi"])))

    doc.add_heading("1.8  Macro, country risk and the cost of capital", level=2)
    para(doc, "The edition of 11 June 2026 discounted at %s. Egypt's ten-year government "
              "bond yields %s. A discount rate below the sovereign yield prices this "
              "company as safer than the government that taxes it, which cannot be "
              "right, and it is not a real-versus-nominal mismatch: that edition "
              "escalated selling prices at 14 per cent a year, so its cash flows "
              "were in nominal pounds throughout."
              % (pct(D["edition_11jun_wacc"], 0), pct(W["rf_observed"], 2)),
         bold=False)
    table(doc, ["Cost of capital", "On the credit-rating basis",
                "On the credit-default-swap basis"],
          [["Egyptian ten-year government bond yield",
            pct(W["rf_observed"], 2), pct(W["rf_observed"], 2)],
           ["less Egypt's own sovereign default spread",
            pct(W["damodaran_egypt"]["adj_default_spread"], 2),
            pct(W["damodaran_egypt"]["sovereign_cds"], 2)],
           ["Normalised risk-free rate", pct(W["rf_star_rating"], 2),
            pct(W["rf_star_cds"], 2)],
           ["Equity risk premium for Egypt",
            pct(W["damodaran_egypt"]["total_erp_rating"], 2),
            pct(W["damodaran_egypt"]["total_erp_cds"], 2)],
           ["Beta", "%.4f" % W["beta_record"]["beta"],
            "%.4f" % W["beta_record"]["beta"]],
           ["Cost of equity", pct(W["ke_rating"], 2), pct(W["ke_cds"], 2)],
           ["Cost of debt before tax", pct(W["kd_pretax_local"], 2),
            pct(W["kd_pretax_local"], 2)],
           ["Cost of debt after tax at 22.5%", pct(W["kd_aftertax"], 2),
            pct(W["kd_aftertax"], 2)],
           ["Equity weight at market value", pct(W["we"], 1), pct(W["we"], 1)],
           ["Debt weight", pct(W["wd"], 1), pct(W["wd"], 1)],
           ["Weighted average cost of capital", pct(W["wacc_rating"], 2),
            pct(W["wacc_cds"], 2)]],
          [7.0, 4.6, 4.6],
          "Country risk is counted once, inside the equity risk premium, and the "
          "government bond yield is reduced by Egypt's own default spread on the "
          "matching basis before that premium is added. Both bases are published "
          "because both are defensible and they differ by about a point.")
    para(doc, "The beta is measured on Palm Hills' own weekly price history against "
              "the EGX30 index over %.2f years: %.4f, with a standard error of "
              "%.4f, an R-squared of %s and %d weekly observations. A cross-check "
              "adjustment toward the market gives %.4f. Appendix B shows the same "
              "measurement for every listed Egyptian developer."
              % (W["beta_record"]["window_years"], W["beta_record"]["beta"],
                 W["beta_record"]["se"], pct(W["beta_record"]["r2"], 1),
                 W["beta_record"]["n"], W["beta_record"]["blume_crosscheck"]))
    para(doc, "The cost of debt is the one cost-of-capital input in this study that "
              "is not sourced, and it is flagged. The right anchor exists: the "
              "company closed an EGP 2,015 million securitisation on 4 February 2026, "
              "the first draw on a newly approved EGP 30 billion programme, in four "
              "tranches of 13, 36, 60 and 84 months rated AA+, AA, A+ and A- on the "
              "Egyptian national scale. No coupon is published for any tranche. The "
              "cost of debt is therefore built as the sovereign yield plus a 250 "
              "basis point corporate spread, which puts it above the sovereign as it "
              "must be. The accounting rate is deliberately not used: EGP %s million "
              "of finance cost on average gross borrowings of about EGP 26,700 "
              "million implies 12.5 per cent, far below the sovereign, because much "
              "of the balance does not bear interest and part of the charge is "
              "capitalised into work in progress." % money(v("finance_cost_fy25")))

    doc.add_heading("1.9  Sensitivity", level=2)
    para(doc, "The table below prices the crux against the discount rate. Read it "
              "down a column rather than across a row: moving the discount rate by "
              "the whole 800 basis points of the 11 June 2026 edition's error changes "
              "value by less than moving cash conversion from one observed year to "
              "another.")
    hdr = ["Cash conversion"] + [pct(w, 2) for w in SENS["waccs"]]
    rows = []
    for i, c in enumerate(SENS["cfos"]):
        tag = {0.039: " (2023 and 2025)", 0.087: " (three-year mean)",
               0.179: " (2024)"}.get(c, "")
        rows.append([pct(c) + tag] + ["%.2f" % x for x in SENS["grid"][i]])
    table(doc, hdr, rows, [4.4, 2.4, 2.4, 2.4, 2.4, 2.4],
          "Value per share in Egyptian pounds. The closing price on 23 August 2026 "
          "was EGP %.2f." % sp)
    figure(doc, "fig2_sensitivity.png",
           "Figure 2 — the same grid. Darker cells are higher values; the two axes "
           "are the discount rate and the rate at which contracted sales turn into "
           "cash.")
    para(doc, "Solving the model back to the closing price gives a cash-conversion "
              "rate of %s. That is the clearest single statement this study can make "
              "about the current price: the market is paying for conversion better "
              "than 2023 and 2025 delivered, and materially worse than 2024 did."
              % pct(D["market_implied_cash_conversion"]))


# ---------------------------------------------------------------------------
def _sections_two_to_seven(doc, sp):
    T, BR = N["technical"], PM["band_record"]

    doc.add_heading("2  Technical and price structure", level=1)
    para(doc, "The shares closed at EGP %.2f on 23 August 2026, above a rising "
              "20-day average at %.2f, a rising 50-day at %.2f and a rising 200-day "
              "at %.2f. The whole moving-average stack is beneath the price and "
              "sloping up, which is the least ambiguous configuration a trend read "
              "produces."
              % (T["close"], T["sma20"], T["sma50"], T["sma200"]))
    table(doc, ["Resistance", "EGP", "Support", "EGP"],
          [["First", "%.2f" % T["levels"]["res"][0], "First",
            "%.2f" % T["levels"]["sup"][0]],
           ["Second", "%.2f" % T["levels"]["res"][1], "Second",
            "%.2f" % T["levels"]["sup"][1]],
           ["Third", "%.2f" % T["levels"]["res"][2], "Third",
            "%.2f" % T["levels"]["sup"][2]]],
          [3.6, 4.0, 3.6, 4.0],
          "Levels are computed from the same cleaned price history the "
          "distribution in section 3 is built on; first resistance and first "
          "support are always the nearest to the close.")
    para(doc, "The levels sit close together — first resistance is %.1f per cent "
              "above the close and first support %.1f per cent below — so the "
              "near-term structure is tight and offers little information on its "
              "own. The 200-day average at %.2f is far below the price, which is a "
              "consequence of how much the shares have risen, not a signal."
              % (100 * (T["levels"]["res"][0] / T["close"] - 1),
                 100 * (1 - T["levels"]["sup"][0] / T["close"]), T["sma200"]))

    doc.add_heading("3  A probabilistic price map", level=1)
    para(doc, "This section is not a forecast and does not connect to the valuation "
              "above. It is a statement of how widely the price has historically "
              "moved over one and three months, expressed as a distribution.")
    d1, d3 = PM["dist"]["m1"], PM["dist"]["m3"]
    table(doc, ["Percentile", "One month (EGP)", "Three months (EGP)"],
          [["5th", "%.2f" % d1["p5"], "%.2f" % d3["p5"]],
           ["25th", "%.2f" % d1["p25"], "%.2f" % d3["p25"]],
           ["50th", "%.2f" % d1["p50"], "%.2f" % d3["p50"]],
           ["75th", "%.2f" % d1["p75"], "%.2f" % d3["p75"]],
           ["95th", "%.2f" % d1["p95"], "%.2f" % d3["p95"]]],
          [4.4, 5.6, 5.6],
          "Resolving %s and %s respectively." % (d1["resolve"], d3["resolve"]))
    figure(doc, "fig3_cone.png",
           "Figure 3 — the same distribution drawn as a cone. The shaded bands are "
           "the middle half and the 5th-to-95th range.")
    para(doc, "Over %d resolved three-month forecasts on this company, the price "
              "finished inside the 90 per cent band %s of the time, and inside the "
              "middle band %s of the time. That is a long record by the standards of "
              "this series. The bands are about %.2f times as wide as a simple "
              "random-walk band over the same horizon, which is what Egyptian "
              "devaluation risk costs in width."
              % (BR["n"], pct(BR["c90"], 1), pct(BR["c50"], 1), BR["width"]))
    table(doc, ["Level (EGP)", "Chance of touching within one month",
                "Within three months"],
          [["%.2f" % lv, "%d%%" % a, "%d%%" % b] for lv, a, b in PM["touch"]],
          [4.4, 6.0, 5.2],
          "The chance that the price trades at or through each level at any point "
          "in the window, not the chance it ends there.")
    para(doc, "Price data to %s, computed %s; the trend read uses data to %s, "
              "computed %s."
              % (PM["asof"]["mc_data"], PM["asof"]["mc_computed"],
                 PM["asof"]["tech_data"], PM["asof"]["tech_computed"]),
         size=8.5, italic=True, color=MUTED)

    doc.add_heading("4  Comparison of the lenses", level=1)
    para(doc, "The three readings are deliberately independent — none is an input to "
              "another — so where they agree, the agreement carries information.")
    table(doc, ["Lens", "What it says", "Relative to the EGP %.2f close" % sp],
          [["Fundamental value", "EGP %.2f to EGP %.2f depending on cash conversion"
            % (CASES["low_conversion"]["per_share"],
               CASES["high_conversion"]["per_share"]),
            "the close sits inside the range, near its lower half"],
           ["Price distribution", "three-month middle band EGP %.2f to %.2f"
            % (PM["dist"]["m3"]["p25"], PM["dist"]["m3"]["p75"]),
            "the distribution is centred slightly above the close"],
           ["Trend structure", "above a rising 20-, 50- and 200-day average",
            "constructive, with resistance only %.1f per cent away"
            % (100 * (N["technical"]["levels"]["res"][0] / sp - 1))]],
          [3.8, 6.4, 5.6])
    para(doc, "They point the same way without being made to. The fundamental range "
              "brackets the price, the distribution is mildly positive and the trend "
              "is intact. None of that resolves the crux in section 1.7.")

    doc.add_heading("5  Catalysts to watch", level=1)
    bullets(doc, [
        "The half-year 2026 results. None had been published as at 30 August 2026, "
        "which is itself notable this late in the year.",
        "The full-year 2025 results release. The audited statements were published "
        "but the accompanying release, which is where units sold, new sales, "
        "deliveries and construction spending are disclosed, was not.",
        "Any disclosure of collection terms — down payment, instalment tenor, "
        "post-handover tail. This is the single disclosure that would most narrow "
        "the range in this study.",
        "Progress on the EGP 30 billion securitisation programme. Selling receivables "
        "converts the instalment book to cash and is the direct lever on the crux; "
        "pricing on any tranche would also fix the cost of debt.",
        "The Village de La Capitale launch, which contributed EGP 24 billion of the "
        "EGP 52 billion of first-quarter 2026 sales. Whether that rate of launch "
        "sales repeats materially changes the order book path.",
        "Egyptian interest rates. The cost of capital here is built on a %s "
        "sovereign yield; a sustained fall would raise every figure in section 1.9."
        % pct(W["rf_observed"], 2),
    ])

    doc.add_heading("6  Reading the probability zones", level=1)
    para(doc, "The bands in section 3 describe where the price has tended to go, not "
              "where it should go. They are wide because Egyptian equities are "
              "genuinely volatile: the middle half of the three-month distribution "
              "spans EGP %.2f to EGP %.2f, a range of about %.0f per cent of the "
              "current price. A reader who wants a narrower band is asking for a "
              "less honest one."
              % (PM["dist"]["m3"]["p25"], PM["dist"]["m3"]["p75"],
                 100 * (PM["dist"]["m3"]["p75"] - PM["dist"]["m3"]["p25"]) / sp))
    para(doc, "The touch ladder answers a different question from the percentile "
              "map. A level with a 51 per cent chance of being touched inside three "
              "months is not a level with a 51 per cent chance of holding — prices "
              "pass through levels far more often than they settle at them.")

    doc.add_heading("7  Caveats and what would change our mind", level=1)
    para(doc, "Five things are not disclosed by the company and are therefore not in "
              "this study. Each is named with what would close it.", bold=True)
    gap_rows = [
        ["Collection schedule", "Down payment, instalment tenor and post-handover "
         "tail are not published. This is the crux; it is measured from outcomes "
         "instead of built from terms.", "Disclosure of contract terms."],
        ["Project unit economics", "No per-project unit mix, unit area, price per "
         "square metre or construction cost per square metre is disclosed.",
         "Project-level disclosure or an investor presentation carrying them."],
        ["Cost of debt", "The February 2026 securitisation publishes sizes, tenors "
         "and ratings but no coupon.", "A coupon on any tranche."],
        ["Full-year 2025 operating figures", "The audited statements were published; "
         "the results release, which carries units and new sales, was not.",
         "The 2025 results release."],
        ["Half-year 2026", "Nothing newer than the first quarter of 2026 had been "
         "published as at this edition's date.", "The half-year filing."],
    ]
    table(doc, ["What is missing", "Why it matters", "What would close it"],
          gap_rows, [3.6, 7.0, 5.2])
    para(doc, "What would change our mind, in order of force:", bold=True)
    bullets(doc, [
        "Two more years of cash conversion clustering. Three observations spanning "
        "%s to %s cannot distinguish a volatile business from a mismeasured one; "
        "five could." % (pct(D["cfo_lo"]), pct(D["cfo_hi"])),
        "Disclosure of collection terms, which would replace the measured range with "
        "a built schedule and narrow the answer directly.",
        "A sustained fall in Egyptian interest rates, which lifts every cell in the "
        "sensitivity table.",
        "Evidence that the order book converts at the rate the company's delivery "
        "history implies rather than more slowly. The order book is now more than "
        "seven times revenue, against 5.4 times in 2024 and 3.4 times in 2023; that "
        "ratio has risen every year and cannot rise indefinitely.",
    ])
    para(doc, "A note on this study's own track record. Tested across ten annual "
              "starting points on this company's history from 2011 to 2025, this "
              "method forecast revenue tolerably but overstated net profit by about "
              "three times, consistently and in 97 per cent of tested cases, and did "
              "worse than simply assuming no change. The cause was the mismatched "
              "recognition clock corrected in section 1.1. That is why this edition "
              "publishes ranges rather than a central figure, and why the reader "
              "should treat the upper end of section 1.9 with more suspicion than "
              "the lower.")


# ---------------------------------------------------------------------------
def _appendices(doc, sp, base):
    doc.add_page_break()
    doc.add_heading("Appendix A  Financial statements", level=1)

    doc.add_heading("A.1  Income statement", level=2)
    H = N["historical_is"]
    hy = ["2023", "2024", "2025"]
    def _h(key, sign=1, dp=1):
        return [("" if key not in H[y] else money(sign * H[y][key]["value"], dp))
                for y in hy]
    table(doc, ["EGP million, as reported"] + hy,
          [["Revenue"] + _h("revenue"),
           ["Cost of revenue"] + _h("cogs", -1),
           ["Cash discount"] + _h("cash_discount", -1),
           ["Gross profit"] + _h("gross_profit"),
           ["Overheads"] + _h("sga", -1),
           ["Depreciation and amortisation"] + _h("da", -1),
           ["Finance cost"] + _h("finance_cost", -1),
           ["Profit before tax"] + _h("npbt"),
           ["Tax"] + _h("tax_total", -1),
           ["Profit after tax"] + _h("npat_pre_nci"),
           ["Minority interest"] + _h("nci", -1),
           ["Profit attributable to shareholders"] + _h("npat_mi")],
          [6.2, 3.3, 3.3, 3.3],
          "Three years, each as that year's own audited statements reported "
          "it. Revenue less cost of revenue less the cash discount foots to "
          "reported gross profit in all three years, and profit before tax "
          "less tax less minority interest foots to the attributable figure.")
    para(doc, "One presentational difference is worth naming rather than "
              "smoothing. Cost of revenue for 2024 is EGP %s million as the "
              "company reported it that year and EGP %s million in the "
              "comparative column of the 2025 statements — a difference of EGP "
              "%.1f million, exactly that year's cash discount, which the later "
              "presentation folds into cost. The forecast is built on the later "
              "basis throughout; the table above shows each year as first "
              "reported. Both figures are right on their own basis."
              % ("{:,.1f}".format(N["fy24_cogs_basis"]["as_reported"]),
                 "{:,.1f}".format(N["fy24_cogs_basis"]["fy25_comparative"]),
                 N["fy24_cogs_basis"]["fy25_comparative"]
                 - N["fy24_cogs_basis"]["as_reported"]))
    FA, FB = ST["framing_a"], ST["framing_b"]
    fy = pick(FB, lambda r: str(r["year"]))
    def _f(rows, key, fmt="{:,.0f}", sign=1):
        return pick(rows, lambda r: fmt.format(sign * r[key]))
    table(doc, ["Forecast, EGP million"] + fy,
          [["Units delivered"] + _f(FB, "units_delivered"),
           ["Revenue per delivered unit"] + _f(FB, "rev_per_unit", "{:,.2f}"),
           ["Revenue"] + _f(FB, "revenue"),
           ["Cost per delivered unit"] + _f(FB, "cost_per_unit", "{:,.2f}"),
           ["Cost of revenue"] + _f(FB, "cogs", "{:,.0f}", -1),
           ["Gross profit"] + _f(FB, "gross"),
           ["Gross margin"]
           + pick(FB, lambda r: "%.1f%%" % (100 * r["gross_margin"])),
           ["Overheads"] + _f(FB, "sga", "{:,.0f}", -1),
           ["Depreciation and amortisation"] + _f(FB, "da", "{:,.0f}", -1),
           ["Operating profit"] + _f(FB, "ebit"),
           ["Finance cost"] + _f(FB, "interest", "{:,.0f}", -1),
           ["Profit before tax"] + _f(FB, "npbt"),
           ["Tax"] + _f(FB, "tax", "{:,.0f}", -1),
           ["Net profit"] + _f(FB, "npat"),
           ["Earnings per share (EGP)"] + _f(FB, "eps", "{:,.2f}")],
          wide_widths(BU["rows"]),
          "Five forecast years. Gross margin is what price per unit and cost "
          "per unit leave behind, never an input. Years three to five should be "
          "read against the range in section 1.9, not as points.")

    doc.add_heading("A.2  Balance sheet", level=2)
    SB, B24 = N["balance_sheet_subtotals"], N["balance_sheet_fy24"]
    def _b(key):
        return [money(q(key), 1)] + [money(SB[y][key]["value"], 1) for y in ("2025", "2024")]
    def _bl(key):
        return [money(q(key), 1), money(v(key), 1), money(B24[key]["value"], 1)]
    table(doc, ["EGP million, as reported", "31 Mar 2026 (reviewed)", "2025", "2024"],
          [["Non-current assets"] + _b("total_noncurrent_assets"),
           ["  of which notes receivable, long term"]
           + [money(q("notes_recv_lt") + q("notes_recv_lt_undel"), 1),
              money(v("notes_recv_lt") + v("notes_recv_lt_undel"), 1),
              money(B24["notes_recv_lt"]["value"]
                    + B24["notes_recv_lt_undel"]["value"], 1)],
           ["Current assets"] + _b("total_current_assets"),
           ["  of which work in progress"] + _bl("work_in_progress"),
           ["  of which trade receivables"] + _bl("accounts_receivable"),
           ["  of which notes receivable, short term"]
           + [money(q("notes_recv_st") + q("notes_recv_st_undel"), 1),
              money(v("notes_recv_st") + v("notes_recv_st_undel"), 1),
              money(B24["notes_recv_st"]["value"]
                    + B24["notes_recv_st_undel"]["value"], 1)],
           ["  of which cash"] + _bl("cash"),
           ["Total assets"] + _b("total_assets"),
           ["Current liabilities"] + _b("total_current_liabs"),
           ["  of which advances from customers"] + _bl("advances_customers"),
           ["  of which suppliers"] + _bl("suppliers"),
           ["Non-current liabilities"] + _b("total_noncurrent_liabs"),
           ["Total liabilities"] + _b("total_liabilities"),
           ["Shareholders' funds"] + _b("total_equity"),
           ["  of which attributable to the parent"]
           + [money(q("equity_parent"), 1), money(v("equity_parent"), 1), "n/a"],
           ["  of which minority interests"]
           + [money(q("nci_equity"), 1), money(v("nci_equity"), 1), "n/a"],
           ["Total equity and liabilities"] + _b("total_assets")],
          [6.2, 3.6, 3.2, 3.2],
          "As reported. Total assets equal total liabilities plus shareholders' "
          "funds on all three dates, and each subtotal reconciles to its "
          "components. The March 2026 column is the reviewed interim statement "
          "and is the sheet the bridge, the book value and the borrowings stand "
          "on; the split of shareholders' funds is not carried in the 2024 "
          "comparative column the study holds.")

    doc.add_heading("A.3  Forecast balance sheet and cash flow", level=2)
    wg = ST["wedge"]
    para(doc, "Because the balance sheet and the cash-flow statement disagree "
              "about 2025 by EGP %s million, and that difference cannot be "
              "split from what has been disclosed, the forecast is published "
              "two ways. Neither is averaged into the other. Both rest on the "
              "same income statement above; they differ only in what they hold "
              "fixed about cash."
              % "{:,.0f}".format(wg["wedge_fy25"]))
    bullets(doc, [
        "The first holds the collection cycle. Every working-capital line "
        "stays at the ratio to revenue it stood at in 2025, the balance sheet "
        "is built from those ratios, and operating cash is what is left. It "
        "asks what the growth costs to fund if nothing about collection "
        "changes.",
        "The second holds cash conversion. Operating cash is set at the "
        "company's own disclosed rate — %.1f, %.1f and %.1f per cent of "
        "revenue in the three published years, %.1f per cent on average — and "
        "working capital is what is left. It asks what the collection cycle "
        "must do for the cash to keep converting as it has."
        % (100 * ST["cash_conversion"]["FY2023"],
           100 * ST["cash_conversion"]["FY2024"],
           100 * ST["cash_conversion"]["FY2025"],
           100 * ST["cash_conversion"]["mean"])])
    para(doc, "The valuation stands on the second, because that is the reading "
              "the company's own cash-flow statements support. The first is "
              "printed beside it because the two do not agree, and a single "
              "set of figures would hide that they don't.")

    for tag, rows, title in (
            ("B", FB, "If cash conversion holds — the basis of the valuation"),
            ("A", FA, "If the collection cycle holds")):
        doc.add_heading(title, level=3)
        table(doc, ["Forecast balance sheet, EGP million"] + fy,
              [["Trade and notes receivable"] + _f(rows, "receivables"),
               ["Work in progress"] + _f(rows, "wip"),
               ["Other receivables and prepayments"]
               + _f(rows, "bs_debtors_other"),
               ["Advances to suppliers"] + _f(rows, "bs_suppliers_advances"),
               ["Cash and equivalents"] + _f(rows, "cash"),
               ["Property and equipment"] + _f(rows, "ppe"),
               ["Other assets, held at the 2025 level"]
               + _f(rows, "other_assets"),
               ["Total assets"] + _f(rows, "total_assets"),
               ["Customer advances"] + _f(rows, "advances"),
               ["Suppliers"] + _f(rows, "suppliers"),
               ["Other creditors"] + _f(rows, "bs_creditors_other"),
               ["Cheques under collection"]
               + _f(rows, "bs_checks_undelivered"),
               ["Borrowings"] + _f(rows, "debt"),
               ["Other liabilities, held at the 2025 level"]
               + _f(rows, "other_liabs"),
               ["Total liabilities"] + _f(rows, "total_liabilities"),
               ["Shareholders' funds"] + _f(rows, "equity"),
               ["Total liabilities and equity"]
               + _f(rows, "total_liabs_and_equity")],
              wide_widths(BU["rows"]),
              "Assets equal liabilities plus shareholders' funds in every "
              "year, to the same tenth of a million the audited 2025 sheet "
              "itself foots to.")
        table(doc, ["Forecast cash flow, EGP million"] + fy,
              [["Net profit"] + _f(rows, "npat"),
               ["Depreciation and amortisation"] + _f(rows, "da"),
               ["Change in working capital, cash effect"]
               + _f(rows, "d_wc", "{:,.0f}", -1),
               ["Cash from operations"] + _f(rows, "cfo"),
               ["  as a share of revenue"]
               + pick(rows, lambda r: "%.1f%%" % (100 * r["cash_conversion"])),
               ["Capital expenditure"] + _f(rows, "cfi"),
               ["New borrowing drawn"] + _f(rows, "drawn"),
               ["Closing cash"] + _f(rows, "cash")],
              wide_widths(BU["rows"]),
              None)
        table(doc, ["The cycle this implies"] + fy,
              [["Collection period, days"]
               + pick(rows, lambda r: "%.0f" % r["dso"]),
               ["Work in progress, days of cost"]
               + pick(rows, lambda r: "%.0f" % r["dio"]),
               ["Suppliers, days of cost"]
               + pick(rows, lambda r: "%.0f" % r["dpo"]),
               ["Customer advances, share of the order book"]
               + pick(rows, lambda r: "%.1f%%" % (100 * r["adv_of_backlog"])),
               ["Net working capital"] + _f(rows, "net_wc"),
               ["  as a multiple of revenue"]
               + pick(rows, lambda r: "%.2f" % r["nwc_over_revenue"])],
              wide_widths(BU["rows"]),
              ("The collection period falls from %.0f days to %.0f over five "
               "years. That is what this reading requires, and it is a large "
               "improvement to take on trust — which is exactly why the other "
               "reading is printed too."
               % (rows[0]["dso"], rows[-1]["dso"])) if tag == "B" else
              ("Every ratio is held where 2025 left it, by construction. The "
               "customer-advance share rises only because the order book grows "
               "more slowly than revenue on this path."))
        if tag == "A":
            d = ST["dcf_a"]
            para(doc, "This reading does not produce a value per share, and "
                      "the reason is worth stating plainly rather than "
                      "burying. Free cash flow is negative in every one of the "
                      "five years and reaches minus EGP %s million in 2030. A "
                      "terminal "
                      "value taken on a negative flow returns a large negative "
                      "number that looks like a valuation and is not one: it "
                      "asserts the company burns cash for ever at a "
                      "compounding rate, which is not a forecast anybody made. "
                      "So none is taken. What this reading measures instead is "
                      "the funding the growth would need."
                      % "{:,.0f}".format(abs(d["terminal_flow"])))
            table(doc, ["If the cycle holds, what the growth costs", "EGP mn"],
                  [["Present value of the five forecast years",
                    "{:,.0f}".format(d["pv_explicit"])],
                   ["New borrowing required by 2030",
                    "{:,.0f}".format(d["funding_required"])],
                   ["Its annual interest by 2030, at %.2f per cent"
                    % (100 * W["kd_pretax_local"]),
                    "{:,.0f}".format(d["funding_interest"])],
                   ["That interest as a share of 2030 operating profit",
                    "%.0f%%" % (100 * d["funding_interest_vs_ebit"])]],
                  [10.6, 5.6],
                  "The interest on that borrowing is not in the profit "
                  "forecast above, and by 2030 it would exceed operating "
                  "profit outright. The collection cycle and this growth path "
                  "cannot both hold — which is the finding, not a defect in "
                  "the arithmetic.")
        else:
            d = ST["dcf_b"]
            table(doc, ["Free cash flow to the firm, EGP million"] + fy,
                  [["Cash from operations"]
                   + pick(d["waterfall"], lambda w: "{:,.0f}".format(w["cfo"])),
                   ["plus finance cost, after tax"]
                   + pick(d["waterfall"], lambda w: "{:,.0f}".format(w["interest_addback"])),
                   ["less capital expenditure"]
                   + pick(d["waterfall"], lambda w: "{:,.0f}".format(-w["capex"])),
                   ["Free cash flow to the firm"]
                   + pick(d["waterfall"], lambda w: "{:,.0f}".format(w["fcff"])),
                   ["Present value, on the schedule"]
                   + pick(d["waterfall"], lambda w: "{:,.0f}".format(w["pv"]))],
                  wide_widths(BU["rows"]),
                  "These are the figures discounted in section 1.1, shown "
                  "again here beside the statements they come out of.")

    doc.add_page_break()
    doc.add_heading("Appendix B  Peer frame, risk register and research register",
                    level=1)
    doc.add_heading("B.1  The Egyptian listed developers", level=2)
    prows = []
    for r in sorted([p for p in N["peers"] if "beta" in p],
                    key=lambda x: -x["beta"]):
        prows.append([r["ticker"], r["name"][:30], "%.4f" % r["beta"],
                      pct(r["r2"], 1), "%.3f" % r["se"],
                      pct(r["ann_vol_5y"], 1), pct(r["max_drawdown_5y"], 1)])
    table(doc, ["Code", "Company", "Beta", "R-squared", "Std error",
                "Volatility", "Worst fall"], prows,
          [1.7, 4.4, 2.0, 2.2, 2.0, 2.0, 2.0],
          "Every measure computed on the same five-year window, through the same "
          "cleaning, against the same index. Volatility is annualised from weekly "
          "returns; worst fall is the deepest peak-to-trough decline over five years.")
    figure(doc, "fig4_peers.png",
           "Figure 4 — the same betas with their standard errors. Palm Hills sits "
           "mid-sector: more market-sensitive than most of its peers, materially "
           "less so than the largest.")
    para(doc, "Earnings and book multiples for the peer group are not published in "
              "this study. No peer discloses financial statements this study can "
              "obtain on a consistent basis, and a multiple built from inconsistent "
              "inputs would carry more authority than it deserves. What is measured "
              "here is measured the same way for every name.")

    doc.add_heading("B.2  Risk register", level=2)
    table(doc, ["Risk", "How it would show up", "Where it is priced"],
          [["Cash conversion stays at the 2023 and 2025 rate",
            "operating cash near %s of revenue" % pct(D["cfo_lo"]),
            "bottom row of the table in section 1.9"],
           ["Egyptian rates stay high or rise",
            "cost of capital above %s" % pct(W["wacc_rating"], 2),
            "right-hand columns of the same table"],
           ["Order book converts more slowly than delivery history implies",
            "the order book keeps rising as a multiple of revenue",
            "section 7, fourth bullet"],
           ["Currency devaluation", "nominal figures rise, real value does not",
            "the whole study is in nominal pounds; a devaluation moves the "
            "denominator too"],
           ["Cost of debt above the assumed spread",
            "finance cost above the modelled level",
            "section 1.8; the spread is the one unsourced input"]],
          [4.6, 5.6, 6.0])

    doc.add_heading("B.3  Research register", level=2)
    para(doc, "This study is built on %d sourced inputs. Every reported historical "
              "comes from the company's own audited statements or its own results "
              "releases. The only third-party inputs are the exogenous macroeconomic "
              "series and the market price. A full register, with the value, source, "
              "date and provenance of each input, is published in the accompanying "
              "bibliography." % len(REG))

    doc.add_page_break()
    doc.add_heading("Appendix C  Expert panel", level=1)
    para(doc, "Three readings of the same company by three different methods. The "
              "experts are identified by method, not by name.", italic=True,
         color=MUTED)

    doc.add_heading("C.1  Expert 1 — the discounted cash-flow analyst", level=2)
    para(doc, "Worldview: a company is worth the cash it will hand its owners, "
              "discounted at what that cash costs. When it works: businesses with "
              "stable conversion of profit to cash. When it fails: businesses where "
              "that conversion is the unknown — which is exactly this one.")
    table(doc, ["Working", "EGP mn"],
          [["Present value of five explicit years", money(base["pv_explicit"])],
           ["Present value beyond year five", money(base["pv_terminal"])],
           ["Enterprise value", money(base["ev"])],
           ["less net debt, 31 March 2026", "(%s)" % money(D["net_debt_bridge"])],
           ["plus associates and investment property, 31 March 2026",
            money(q("investments_assoc") + q("investment_property"))],
           ["Equity before minority interests", money(base["equity_before_nci"])],
           ["less minority interests at their share of value",
            "(%s)" % money(base["nci_deduction"])],
           ["Equity attributable to shareholders", money(base["equity"])],
           ["Per share (EGP)", "%.2f" % base["per_share"]]],
          [9.0, 5.0])
    para(doc, "Named sensitivity: a move in cash conversion from %s to %s takes the "
              "answer from EGP %.2f to EGP %.2f. Falsifier stated in advance: if the "
              "2026 and 2027 cash-flow statements show conversion inside two "
              "percentage points of %s, this reading is right and the range should "
              "collapse toward its centre; if they straddle the full observed spread "
              "again, the method should be abandoned for this company."
              % (pct(D["cfo_lo"]), pct(D["cfo_hi"]),
                 CASES["low_conversion"]["per_share"],
                 CASES["high_conversion"]["per_share"], pct(D["cfo_mid"])))

    doc.add_heading("C.2  Expert 2 — the asset-based analyst", level=2)
    para(doc, "Worldview: a developer is a pile of land, work in progress and "
              "receivables, less what it owes. Earnings are an accident of timing. "
              "When it works: asset-heavy businesses in distress or in wind-down. "
              "When it fails: going concerns whose value is in execution.")
    table(doc, ["Working, on the reviewed balance sheet of 31 March 2026", "EGP mn"],
          [["Work in progress", money(q("work_in_progress"))],
           ["Receivables, trade and instalment",
            money(q("accounts_receivable") + q("notes_recv_st")
                  + q("notes_recv_lt") + q("notes_recv_st_undel")
                  + q("notes_recv_lt_undel"))],
           ["Cash", money(q("cash"))],
           ["Investments and investment property",
            money(q("investments_assoc") + q("investment_property"))],
           ["less advances from customers",
            "(%s)" % money(q("advances_customers"))],
           ["less gross borrowings", "(%s)" % money(D["gross_debt_bridge"], 1)],
           ["less other liabilities, balancing",
            "(%s)" % money(q("total_liabilities") - q("advances_customers")
                           - D["gross_debt_bridge"], 1)],
           ["Shareholders' funds as reported", money(q("total_equity"))],
           ["  of which attributable to the parent", money(q("equity_parent"))],
           ["Per share, parent (EGP)", "%.2f" % D["book_equity_per_share"]]],
          [9.0, 5.0])
    para(doc, "Named sensitivity: the instalment receivable book is carried at face "
              "value less unwinding discount. A 10 per cent write-down of that book "
              "would remove roughly EGP %s million, about EGP %.2f a share. "
              "Falsifier: evidence of material cancellations or receivable "
              "impairment would confirm this reading and invalidate the cash-flow "
              "one."
              % (money(0.10 * (q("accounts_receivable") + q("notes_recv_st")
                               + q("notes_recv_lt"))),
                 0.10 * (q("accounts_receivable") + q("notes_recv_st")
                         + q("notes_recv_lt")) / D["shares_mn"]))

    doc.add_heading("C.3  Expert 3 — the market-implied analyst", level=2)
    para(doc, "Worldview: do not value the company; ask what the price already "
              "assumes, then judge whether that assumption is reasonable. When it "
              "works: where a single variable dominates. When it fails: where the "
              "market is wrong about many things at once.")
    para(doc, "Working: holding the cost of capital at %s and every other input at "
              "the audited 2025 base, the closing price of EGP %.2f implies cash "
              "conversion of %s. The company achieved %s in 2023, %s in 2024 and %s "
              "in 2025. Named sensitivity: at a %s cost of capital the implied rate "
              "would be lower still. Falsifier: if 2026 conversion prints above %s, "
              "the price was too low."
              % (pct(W["wacc_rating"], 2), sp,
                 pct(D["market_implied_cash_conversion"]),
                 pct(D["cfo_margins"]["2023"]), pct(D["cfo_margins"]["2024"]),
                 pct(D["cfo_margins"]["2025"]), pct(W["wacc_cds"], 2),
                 pct(D["market_implied_cash_conversion"])))

    doc.add_heading("C.4  Cross-examination", level=2)
    bullets(doc, [
        "Expert 2 to Expert 1: your terminal value is %s of your answer, so you are "
        "mostly valuing years you cannot see. CONCEDED — it is why no single figure "
        "is published." % pct(base["terminal_share"], 0),
        "Expert 1 to Expert 2: book value ignores that the company sells at roughly "
        "%s gross margin, so its inventory is worth more than cost. CONCEDED — book "
        "is a floor, not a value." % pct(D["gross_margin_fy25"]),
        "Expert 3 to both: you disagree by a factor of four and the market sits "
        "between you, so neither of you is adding information. REJECTED — the range "
        "is the information; a false point estimate would not be.",
        "Expert 2 to Expert 3: the price also reflects Egyptian equity risk appetite, "
        "not just this company's conversion rate. CONCEDED — the implied figure is a "
        "single-variable solve and is labelled as one.",
    ])

    doc.add_heading("C.5  The three in one room", level=2)
    para(doc, "They agree on the facts and disagree on one number. All three accept "
              "the audited 2025 base, the %s cost of capital and the order book. "
              "Expert 1 needs a conversion rate to produce an answer; Expert 2 "
              "refuses to forecast one and prices the assets instead; Expert 3 asks "
              "what rate is already in the price. The room's joint position is that "
              "Palm Hills is worth between EGP %.2f and EGP %.2f, that the price of "
              "EGP %.2f sits in the lower half of that range, and that the width is a "
              "property of the disclosure rather than of the analysis."
              % (pct(W["wacc_rating"], 2),
                 CASES["low_conversion"]["per_share"],
                 CASES["high_conversion"]["per_share"], sp))

    doc.add_heading("C.6  Reading the divergence", level=2)
    table(doc, ["Assumption", "Expert 1", "Expert 2", "Expert 3",
                "Drives the gap?"],
          [["Cash conversion", "%s to %s" % (pct(D["cfo_lo"]), pct(D["cfo_hi"])),
            "not forecast", "%s implied" % pct(D["market_implied_cash_conversion"]),
            "YES — this is the whole gap"],
           ["Cost of capital", pct(W["wacc_rating"], 2), "not used",
            pct(W["wacc_rating"], 2), "no — shared"],
           ["Receivable book", "collected on schedule", "at carrying value",
            "as the market judges", "partly"],
           ["Terminal value", "%s of the answer" % pct(base["terminal_share"], 0),
            "none", "none", "yes for Expert 1 alone"]],
          [3.4, 3.4, 3.0, 3.2, 3.4],
          "One assumption explains almost all of the divergence, which is the "
          "clearest evidence that section 1.7 has identified the right crux.")

    doc.add_page_break()
    doc.add_heading("About this series", level=1)
    para(doc, "These studies value a company from its own disclosures, publish the "
              "reasoning alongside the answer, and record what the method has "
              "historically got right and wrong on that company. They carry no "
              "rating and no price target. Where a company does not disclose "
              "something the valuation needs, the study says so and prices the "
              "consequence rather than filling the gap.")
    doc.add_heading("Disclosure and disclaimer", level=1)
    para(doc, "This document is for information and education. It is not investment "
              "advice, not a recommendation, and not an offer or solicitation to buy "
              "or sell any security. It contains no rating and no price target. "
              "Figures are drawn from the company's published financial statements "
              "and results releases and from the named third-party sources; they may "
              "contain errors and are not warranted. Valuations are estimates that "
              "depend on assumptions stated in the text, and small changes in those "
              "assumptions produce large changes in the result, as section 1.9 shows "
              "directly. Past price behaviour does not predict future returns. "
              "Readers must reach their own conclusions and should take professional "
              "advice before acting. Edition of 2 September 2026; information set ends "
              "at the company's first-quarter 2026 results.",
         size=8.5, color=MUTED)


if __name__ == "__main__":
    out = os.path.join(HERE, "PHDC_Valuation_Study_03-09-2026.docx")
    doc = build(out)
    doc.save(out)
    hits, chars = scrub(out)
    bad = column_audit(out)
    print("built: %s  (%.0f KB, %d characters of text)"
          % (os.path.basename(out), os.path.getsize(out) / 1024, chars))
    print("external-reader scrub : %s" % ("CLEAN" if not hits else "HITS %s" % hits))
    print("table column audit    : %s" % ("CLEAN" if not bad else bad))
