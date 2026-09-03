"""Build the standalone PHDC bibliography document.

Five things the depth bar requires and this document carries: the primary
documents table, the FULL input register with value, date and source-and-
construction for every input grouped by research layer, a judgements table each
with what would overturn it, a negative-results table, and notes on every place
two sources disagree.
"""
import json, os, sys, collections

from docx import Document
from docx.shared import Pt, Cm

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from docx_phdc import _style, para, table, bullets, scrub, column_audit, MUTED, ACCENT

N = json.load(open(os.path.join(HERE, "study_numbers.json")))
REG, D, W = N["registry"], N["derived"], N["wacc"]

LAYER = {
    "revenue_": "Company — income statement", "cogs_": "Company — income statement",
    "gross_profit_": "Company — income statement", "sga_": "Company — income statement",
    "da_": "Company — income statement", "finance_cost_": "Company — income statement",
    "npbt_": "Company — income statement", "tax_": "Company — income statement",
    "npat_": "Company — income statement",
    "cfo_": "Company — cash flow", "cfi_": "Company — cash flow",
    "cff_": "Company — cash flow",
    "total_": "Company — balance sheet", "cash": "Company — balance sheet",
    "work_in_progress": "Company — balance sheet",
    "accounts_receivable": "Company — balance sheet",
    "notes_recv": "Company — balance sheet", "advances_": "Company — balance sheet",
    "suppliers": "Company — balance sheet", "investments_": "Company — balance sheet",
    "investment_property": "Company — balance sheet",
    "fixed_assets": "Company — balance sheet", "fin_inv_": "Company — balance sheet",
    "debtors_": "Company — balance sheet", "due_from_": "Company — balance sheet",
    "inv_fair_": "Company — balance sheet", "equity_": "Company — balance sheet",
    "nci_": "Company — balance sheet", "deferred_revenue": "Company — balance sheet",
    "checks_": "Company — balance sheet", "creditors_": "Company — balance sheet",
    "loans_": "Company — borrowings", "notes_payable": "Company — borrowings",
    "credit_facilities": "Company — borrowings", "banks_": "Company — borrowings",
    "current_portion": "Company — borrowings", "lease_": "Company — borrowings",
    "backlog_": "Company — operating", "new_sales": "Company — operating",
    "vdlc_": "Company — operating", "construction_": "Company — operating",
    "collections_": "Company — operating", "land_bank": "Company — operating",
    "units_": "Company — operating", "revenue_1q26": "Company — operating",
    "shares_": "Market", "spot": "Market",
}
# every line of the 31 March 2026 reviewed sheet is registered with a _1q26 suffix
LAYER_1Q26 = "Company — balance sheet, 31 March 2026 (reviewed)"


def layer_of(key):
    if key.endswith("_1q26") and not key.startswith(("revenue_", "gross_profit_", "npat_",
                                                    "backlog_", "new_sales_", "vdlc_")):
        return LAYER_1Q26
    for pre, lab in sorted(LAYER.items(), key=lambda kv: -len(kv[0])):
        if key.startswith(pre):
            return lab
    return "Other"


def build(path):
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.0)
        s.top_margin = s.bottom_margin = Cm(2.0)
    _style(doc)

    para(doc, "PALM HILLS DEVELOPMENTS", size=18, bold=True, color=ACCENT,
         space_after=2)
    para(doc, "Sources, inputs and judgements · edition of 2 September 2026",
         size=10.5, color=MUTED, space_after=14)
    para(doc, "This document accompanies the valuation study. It lists every "
              "document the study was built on, every input with its value, its "
              "date and how it was constructed, every judgement with what would "
              "overturn it, and everything that was tried and rejected. It exists "
              "so that a reader can check the study rather than trust it.")

    doc.add_heading("1  Primary documents", level=1)
    table(doc, ["Document", "Period", "Publisher", "Used for"],
          [["Consolidated financial statements", "FY2025",
            "Palm Hills Developments", "the entire reported base: income "
            "statement, balance sheet and cash-flow statement"],
           ["Consolidated financial statements", "FY2024 (comparative column)",
            "Palm Hills Developments", "prior-year balance sheet and cash flow"],
           ["Consolidated financial statements", "FY2023",
            "Palm Hills Developments", "revenue and gross profit"],
           ["Consolidated financial statements (reviewed)", "1Q2026 — 31 March 2026",
            "Palm Hills Developments", "the most recent reported quarter, and the "
            "balance sheet the bridge, the book value and the borrowings stand on "
            "(a scan; figures read off the rendered pages and held to the "
            "statement's own subtotals)"],
           ["Results release", "1Q2026", "Palm Hills Developments",
            "order book, new sales, the land-plot launch"],
           ["Results release", "FY2024", "Palm Hills Developments",
            "units, new sales, deliveries, construction spend, land bank"],
           ["Results release", "FY2023", "Palm Hills Developments",
            "units, cash from operations, construction spend"],
           ["Country risk premium file", "read 30 Aug 2026",
            "Damodaran, NYU Stern", "Egypt equity risk premium and sovereign "
            "default spread, both bases"],
           ["EGX30 index series", "to 22 Jul 2026", "Egyptian Exchange",
            "the regressor for the beta"],
           ["World Development Indicators", "read 30 Aug 2026", "World Bank",
            "Egyptian inflation, exchange rate and urban population"]],
          [4.4, 3.2, 3.6, 5.4],
          "All company documents were obtained from the company's own investor "
          "relations result centre.")

    doc.add_heading("2  Full input register", level=1)
    para(doc, "Every input in the study, grouped by research layer. Provenance A is "
              "the company's own audited statements or its own results releases; B "
              "is an exchange or regulator source; C is a credible third party. No "
              "reported historical in this study is below A.")
    groups = collections.OrderedDict()
    for k, rec in REG.items():
        groups.setdefault(layer_of(k), []).append((k, rec))
    for lab in sorted(groups):
        doc.add_heading(lab, level=3)
        rows = []
        for k, rec in sorted(groups[lab]):
            val = rec["value"]
            sval = ("{:,.4f}".format(val).rstrip("0").rstrip(".")
                    if isinstance(val, float) else "{:,}".format(val))
            rows.append([k.replace("_", " "), sval, rec.get("unit", ""),
                         rec["date"], rec["tier"], rec["source"][:200]])
        table(doc, ["Input", "Value", "Unit", "Date", "Tier",
                    "Source and construction"], rows,
              [3.2, 2.1, 1.6, 1.9, 1.4, 6.4], size=7.5)

    doc.add_heading("3  Judgements, and what would overturn each", level=1)
    table(doc, ["Judgement", "What was decided", "What would overturn it"],
          [["Discount rate",
            "Cost of capital rebuilt bottom-up at %s on the rating basis and %s on "
            "the swap basis, replacing the 11 June 2026 edition's %s; unchanged in "
            "this edition."
            % ("{:.2%}".format(W["wacc_rating"]), "{:.2%}".format(W["wacc_cds"]),
               "{:.0%}".format(D["edition_11jun_wacc"])),
            "A sustained fall in the Egyptian sovereign yield, or evidence that "
            "the company borrows materially below sovereign plus 250 basis points."],
           ["Cash conversion is the crux",
            "Value is driven by the share of revenue reaching operating cash, "
            "measured across its full observed range rather than assumed.",
            "Two further years clustering within about two points of the mean "
            "would collapse the range and make this an ordinary input."],
           ["Revenue is capacity-limited, not sales-limited",
            "The order book is more than seven times revenue, so delivery "
            "capacity binds and is what is grown.",
            "A collapse in new sales sustained long enough to bring the order book "
            "below about three years of revenue."],
           ["Cost accrues on the same clock as revenue",
            "Both legs accrue with completion, so gross margin is an output.",
            "Evidence that the company recognises cost on handover for a material "
            "part of its book."],
           ["Per-project economics are not used",
            "The 11 June 2026 edition's project price and cost table is not disclosed "
            "by the company and is not reused.",
            "Project-level disclosure of unit mix, area, price and cost."],
           ["The bridge stands on the latest disclosed balance sheet",
            "Net debt, associates, investment property and book equity are taken "
            "from the reviewed statement of 31 March 2026 (net debt EGP %s million "
            "against %s million at 31 December 2025); the projected statements keep "
            "the audited full year 2025 as their base."
            % ("{:,.1f}".format(D["net_debt_bridge"]), "{:,.1f}".format(D["net_debt"])),
            "A later balance sheet — the half-year 2026 statements, once published."],
           ["Minority interests deducted at their share of value",
            "The cash-flow model capitalises all of the subsidiaries' cash flow, so "
            "the minority's claim comes out at its share of the resulting value, "
            "proxied by its filed share of 2025 profit after tax (%s), applied to "
            "equity value. At book it would be EGP %s million (%s of equity); on the "
            "three-year mean profit share, %s. Both are shown as reference."
            % ("{:.2%}".format(D["nci_value_share"]), "{:,.1f}".format(D["nci_book_1q26"]),
               "{:.1%}".format(D["nci_book_share_1q26"]), "{:.2%}".format(D["nci_profit_share_3y"])),
            "Disclosure of the subsidiaries that carry the minority with their own "
            "economics, which would let the minority be valued directly."],
           ["Normalised earnings capitalised at cost of equity less growth",
            "Earnings power is E divided by (cost of equity less the %s terminal "
            "growth the cash-flow model carries), not E divided by the cost of "
            "equity alone: in a currency whose discount rate embeds Egyptian "
            "inflation, zero nominal growth is a perpetual real decline."
            % "{:.1%}".format(N["lens_detail"]["normalised_inputs"]["growth_netted"]),
            "Evidence that the company's real earnings are in secular decline, which "
            "would make the zero-growth form the right one."],
           ["Peer multiples are not published",
            "Only measures obtainable for every peer on the same basis are shown.",
            "Consistent, obtainable peer financial statements."]],
          [3.4, 6.6, 6.4])

    doc.add_heading("4  Negative results — what was tried and rejected", level=1)
    table(doc, ["Attempt", "Why it was rejected"],
          [["Recover the missing FY2024 unit count by summing the disclosed "
            "regional unit series",
            "On the four years where the company total IS disclosed, the sum "
            "overstates it by about a third because the regional charts overlap. "
            "FY2024 and FY2025 unit counts are therefore absent rather than "
            "inferred."],
           ["Use the accounting cost of debt",
            "Finance cost of EGP 3,347.5mn on average gross borrowings of about "
            "EGP 26,700mn implies 12.5%, far below the sovereign, because much of "
            "the balance does not bear interest and part of the charge is "
            "capitalised. A historical rate is not a marginal rate."],
           ["Read the FY2025 statements at a single scan resolution",
            "Neither resolution is reliable throughout: at one, notes receivable "
            "long term read 801.3 against a true 54,801.3; at the other, fixed "
            "assets read 5.0 against a true 4,522.0. Every component list is "
            "instead required to sum to the subtotal the company printed."],
           ["Source peer earnings and book multiples from an aggregator",
            "The pages carry no financial tables. An invented multiple would look "
            "authoritative and mean nothing."],
           ["Obtain a live Egyptian sovereign yield from the exchange or the "
            "central bank",
            "Both publish through interfaces this study could not read. The yield "
            "used is a dated market quote and is flagged as 24 days old."]],
          [5.6, 10.8])

    doc.add_heading("5  Where two sources disagree", level=1)
    table(doc, ["Item", "One source", "The other", "Which is used"],
          [["FY2024 cash from operations", "EGP 4,854.8mn — FY2025 audited "
            "statements, comparative column", "EGP 3,131.9mn — the company's own "
            "FY2024 results release", "the audited statements"],
           ["FY2023 revenue", "EGP 17,462.1mn — FY2023 audited statements",
            "EGP 17,454.6mn — FY2024 results release comparative",
            "the audited statements; the difference is 0.04%"],
           ["FY2016 new sales", "EGP 8,194mn as first reported",
            "EGP 8,467mn in later releases",
            "as first reported, for any measurement anchored at that date"],
           ["FY2015 revenue", "EGP 3,560.6mn as filed",
            "EGP 3,641.7mn restated after the 2016 accounting-policy change",
            "both, kept apart; the restatement is a change of basis, not an error"]],
          [3.6, 4.6, 4.6, 3.6])

    doc.add_heading("6  What is not disclosed", level=1)
    rows = [[k.replace("_", " ").title(), why] for k, why in N["gaps"].items()]
    table(doc, ["Gap", "Statement"], rows, [4.0, 12.4])
    return doc


if __name__ == "__main__":
    out = os.path.join(HERE, "PHDC_Bibliography_02-09-2026.docx")
    build(out).save(out)
    hits, chars = scrub(out)
    bad = column_audit(out)
    print("built: %s (%.0f KB, %d characters)"
          % (os.path.basename(out), os.path.getsize(out) / 1024, chars))
    print("external-reader scrub : %s" % ("CLEAN" if not hits else hits))
    print("table column audit    : %s" % ("CLEAN" if not bad else bad))
