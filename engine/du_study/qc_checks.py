"""Automated QC checks over the delivered documents:
  (k)/(m) no internal-procedure vocabulary reaches the external reader
  (o)      table column widths fit the text block and no column is starved or bloated
  (h)/(n)  figure canvases are opaque and light (numbers legible on any page)
  general  no placeholder or unformatted values leaked into the documents
"""
import os, re, sys, json
from docx import Document
from docx.shared import Inches
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
STUDY = os.path.join(HERE, 'DU_Valuation_Study_09-08-2026_public.docx')
BIB = os.path.join(HERE, 'DU_Bibliography_09-08-2026.docx')


def latest_ddmmyyyy(pat):
    """The workbook names its edition DDMMYYYY with no separators, so the date is PARSED
    rather than the filenames sorted as text — 03092026 sorts below 09082026 as a string and
    a text sort silently picks a superseded edition [L-067]."""
    c = []
    for f in os.listdir(HERE):
        if re.match(pat, f) and not f.startswith('~$'):
            m = re.findall(r'_(\d{2})(\d{2})(\d{4})_', f)
            c.append(((m[-1][2] + m[-1][1] + m[-1][0]) if m else '', f))
    return sorted(c)[-1][1] if c else None


# THE WORKBOOK IS A DELIVERED DOCUMENT AND THIS SCRUB EXCLUDED IT [L-350]. A reader receives
# three files and the list above named two, so the third was scanned by nothing. Only its
# STRING cells are read: a numeric cell is a model output the recalculation gate reconciles,
# and a formula is not a sentence.
XLSX = os.path.join(HERE, latest_ddmmyyyy(r'.*Valuation_Model_\d{8}.*\.xlsx$'))
TEXT_WIDTH = 7.0            # 8.5in page less 0.75in margins each side
BIB_WIDTH = 7.1             # bibliography uses 0.7in margins

BANNED = [
    r'\bstep\s*0(\.0)?\b', r'\bstep\s*2a\b', r'\bstep\s*[1-9]\b',
    r'\bfour[- ]ring\b', r'\bring\s+(classification|guide)\b', r'\bB/S/D/C\b',
    r'\binformation sweep\b', r'\bsweep register\b', r'\bQC gate\b', r'\bgate item\b',
    r'\bcalibration gate\b', r'\bmateriality gate\b', r'\bpromotion rule\b',
    r'\bstanding research protocol\b', r'\bresearch protocol\b',
    r'\bmc_v3\b', r'\bmarket_profiles\b', r'\bfitted_configs\b', r'\bstudy_numbers\b',
    r'\bcompute\.py\b', r'\bdata_quality\b', r'\bclean_ohlc\b', r'\bbacktest_v3\b',
    r'\bapply_breaks\b', r'\brobust_verdict\b', r'\bpanel_refresh\b', r'\bwacc_builder\b',
    r'\bLONO\b', r'\bCRPS\b', r'\bPIT\b', r'\bwidth_cal\b', r'\bkd_path\b',
    r'\bcalibration ledger\b', r'\bledger cohort\b',
    r'\bdevice A-\d\b', r'\bexpert persona library\b', r'\bpersona library\b',
    r'\bscale-normalis?zed\b', r'\bbootstrap block\b', r'\bengine reconciliation\b',
]

CASE_BANNED = [r'\bPARITY\b', r'\bBOUNDARY\b', r'\bFAIL\b(?! )']

def xlsx_strings(path):
    """Every string cell a reader sees in the workbook; formulas and numbers are skipped."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
    out = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if isinstance(v, str) and not v.startswith('='):
                    out.append(v)
    wb.close()
    return out


def doc_text(path):
    if path.lower().endswith(('.xlsx', '.xlsm')):
        return None, '\n'.join(xlsx_strings(path))
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return d, '\n'.join(parts)

fails = []

# ---- (k)/(m) procedure-reference scrub --------------------------------------
print('=' * 74)
print('(k)+(m)  internal-procedure vocabulary scrub')
for label, path in [('study', STUDY), ('bibliography', BIB), ('workbook', XLSX)]:
    _, txt = doc_text(path)
    low = txt.lower()
    hits = []
    for pat in BANNED + CASE_BANNED:
        flags = 0 if pat in CASE_BANNED else re.I
        hay = txt if pat in CASE_BANNED else low
        for m in re.finditer(pat, hay, flags):
            ctx = txt[max(0, m.start() - 45):m.end() + 45].replace('\n', ' ')
            hits.append(f'{pat} -> "...{ctx}..."')
    print(f'  {label}: {len(hits)} hits')
    for h in hits:
        print('     ', h)
    if hits:
        fails.append(f'procedure vocabulary in {label}')

# ---- (o) table column widths --------------------------------------------------
print('=' * 74)
print('(o)  table column widths')
for label, path, maxw in [('study', STUDY, TEXT_WIDTH), ('bibliography', BIB, BIB_WIDTH)]:
    d = Document(path)
    over, narrow, wide = [], [], []
    for i, t in enumerate(d.tables):
        ws = [c.width.inches if c.width else None for c in t.columns]
        if any(w is None for w in ws):
            over.append(f'table {i}: unset width')
            continue
        tot = sum(ws)
        if tot > maxw + 0.02:
            over.append(f'table {i}: total {tot:.2f}in > {maxw:.2f}in')
        # a text column narrower than 0.55in will wrap mid-word for ordinary labels
        for j, w in enumerate(ws):
            if w < 0.55:
                narrow.append(f'table {i} col {j}: {w:.2f}in')
        # a single column taking more than 75% of a multi-column table is bloated
        if len(ws) > 2 and max(ws) / tot > 0.75:
            wide.append(f'table {i}: col takes {max(ws)/tot:.0%} of width')
    print(f'  {label}: {len(d.tables)} tables | over-wide {len(over)} | starved {len(narrow)} | '
          f'bloated {len(wide)}')
    for x in over + narrow + wide:
        print('     ', x)
    if over or narrow or wide:
        fails.append(f'column widths in {label}')

# ---- (h)/(n) figure canvases ----------------------------------------------------
print('=' * 74)
print('(h)+(n)  figure canvases opaque and light')
for fn in sorted(f for f in os.listdir(HERE) if f.endswith('.png')):
    im = Image.open(os.path.join(HERE, fn)).convert('RGBA')
    alpha = im.getchannel('A')
    transparent = alpha.getextrema()[0] < 255
    rgb = im.convert('RGB')
    px = list(rgb.getdata())
    corners = [rgb.getpixel(p) for p in [(1, 1), (rgb.width - 2, 1), (1, rgb.height - 2),
                                         (rgb.width - 2, rgb.height - 2)]]
    light = all(sum(c) / 3 > 200 for c in corners)
    ok = (not transparent) and light
    print(f'  {fn}: {"OK " if ok else "FAIL"} '
          f'(transparent={transparent}, corner luminance={[round(sum(c)/3) for c in corners]})')
    if not ok:
        fails.append(f'figure canvas {fn}')

# ---- general: no leaked placeholders --------------------------------------------
print('=' * 74)
print('general  leaked placeholders / unformatted values')
for label, path in [('study', STUDY), ('bibliography', BIB), ('workbook', XLSX)]:
    d, txt = doc_text(path)
    bad = re.findall(r'\{[a-z_]+\}|\bnan\b|\d+e[+-]\d\d|\bTODO\b|\bXXX\b', txt)
    # a whole cell whose entire content is a Python repr is the real leakage mode
    cells = ([c.text for t in d.tables for row in t.rows for c in row.cells] if d is not None
             else xlsx_strings(path))
    for ctext in cells:
        if ctext.strip() in ('None', 'nan', 'inf', '-inf', '[]', '{}', '0.0%'):
            bad.append(f'cell="{ctext.strip()}"')
    print(f'  {label}: {len(bad)} hits {sorted(set(bad))[:8]}')
    if bad:
        fails.append(f'placeholders in {label}')

# ---- the result, written where the model can read it back ---------------------
# [R-ENF-01] A SELF-ATTESTED BOOLEAN IS NEVER A CHECK. This scan has always been real and
# nothing downstream could see it: compute.py could attest external_reader_scrub only on
# its own say-so. The result is now a file, it NAMES THE FILES IT SCANNED, and the model
# refuses to attest on a result covering an edition nobody receives.
json.dump({
    'files': [os.path.basename(STUDY), os.path.basename(BIB), os.path.basename(XLSX)],
    'clean': not fails,
    'hits': sorted(set(fails)),
    'patterns': len(BANNED) + len(CASE_BANNED),
    'chars': sum(len(doc_text(p)[1]) for p in (STUDY, BIB, XLSX)),
}, open(os.path.join(HERE, 'scrub_result.json'), 'w'), indent=1)

print('=' * 74)
if fails:
    print('QC CHECKS FAILED:', '; '.join(fails))
    sys.exit(1)
print('ALL AUTOMATED QC CHECKS PASSED')
