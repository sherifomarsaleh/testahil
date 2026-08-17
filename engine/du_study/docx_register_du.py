"""DU_Bibliography_09-08-2026.docx — the companion bibliography document.
Every input in the model: value, source, date and research layer — emitted from
study_numbers.json (the compute script's own INPUTS block), plus the document
bibliography and the negative results."""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(HERE, 'study_numbers.json')))
INP = D['inputs']
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(table, top=36, bottom=36, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    table._tbl.tblPr.append(m)

def borders(table, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    table._tbl.tblPr.append(b)

def P(text='', size=10, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def H1(t): return P(t, size=13, bold=True, space_before=12, space_after=4)
def H2(t): return P(t, size=11, bold=True, space_before=8, space_after=3)

def table(rows, widths, size=8.2, header=True):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t); t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def masthead():
    t = doc.add_table(rows=1, cols=1)
    cell_margins(t, 80, 80, 150, 150)
    c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.1)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def fmt(v):
    if isinstance(v, float):
        return f"{v:,.4f}".rstrip('0').rstrip('.') if abs(v) < 100 else f"{v:,.1f}"
    if isinstance(v, int):
        return f"{v:,}"
    if isinstance(v, dict):
        return "; ".join(f"{k}: {fmt(x)}" for k, x in v.items())
    if isinstance(v, (list, tuple)):
        return ", ".join(fmt(x) for x in v)
    return str(v)

# ============================================================================
masthead()
H1('Emirates Integrated Telecommunications Company PJSC (DFM: DU) — Bibliography and '
   'Source Register')
P('Companion document to the valuation study dated 9 August 2026. It records where every '
  'number in that study came from.', size=9.5, color=GREY)

H2('READ FIRST')
P('This document exists so that a reader can check the study rather than trust it. It lists '
  'every input the valuation model uses, together with the value, the source, the date of '
  'that source, and the research layer it belongs to. Nothing in the study is computed from '
  'a figure that does not appear here.')
P('Two things are worth knowing before reading the tables. First, inputs marked "House" are '
  'judgements or derivations made by the analyst, not disclosures by the company — each '
  'carries the reasoning that produced it, and the reader is free to disagree with it and '
  're-run the model. Second, where a disclosure could not be reached, that is recorded as a '
  'negative result at the end of this document rather than quietly filled in.')

H2('The research layers')
table([['Layer', 'What it covers'],
       ['Company', 'the company\'s own audited/reviewed statements, filings, releases and '
        'investor presentations — the only permitted source for its reported historicals'],
       ['Industry', 'peers, the regulator, sector pricing and multiples'],
       ['Country', 'UAE macro, rates, the sovereign, the fiscal regime'],
       ['Market', 'exchange data: prices, the index, the beta regression'],
       ['House', 'analyst judgements and constructions, each with its reasoning']],
      [1.10, 5.95])

H2('Primary documents relied upon')
table([['Document', 'What it provided', 'Date', 'Where read'],
       ['Integrated Annual Report 2025 (audited consolidated FS; KPMG Lower Gulf, '
        'unmodified opinion)', 'FY2025 and re-presented FY2024: every income-statement, '
        'balance-sheet, cash-flow and segment line; notes 4-39 incl. borrowings (none), '
        'leases, tax/royalty, segments', '09-Feb-2026', 'investors.du.ae'],
       ['Annual Report 2024 (audited; PwC, unqualified)', 'FY2024 original presentation and '
        'FY2023 comparatives; the FY2023 royalty-accrual disclosure that makes the '
        'working-capital series like-for-like', '10-Feb-2025', 'investors.du.ae'],
       ['Annual Report 2023 (audited; PwC, unqualified)', 'FY2023 and complete FY2022 '
        'comparatives; the pre-2024 royalty construction (Note 27) that defines Framing B',
        '13-Feb-2024', 'investors.du.ae'],
       ['H1-2026 condensed consolidated interim FS (KPMG ISRE 2410 review)', 'H1/Q2-2026 '
        'actuals; the 30-Jun-2026 balance sheet; the licence-term note; the royalty-regime '
        'note; dividend record', '22-Jul-2026', 'investors.du.ae'],
       ['Q1-2026 condensed consolidated interim FS (KPMG review)', 'Q1-2026 actuals',
        '22-Apr-2026', 'investors.du.ae'],
       ['Q2-2026 earnings release + analyst presentation', 'subscribers, ARPU, capex '
        'intensity, revised FY2026 guidance, interim dividend', '23-Jul-2026',
        'investors.du.ae'],
       ['Q4/FY2025 results presentation', 'FY2025 KPIs, subscriber base, original FY2026 '
        'guidance, dividend and payout record', '10-Feb-2026', 'investors.du.ae'],
       ['UAE MoF T-bond auction results; Emirates Islamic AED yield sheet', 'the AED '
        'risk-free anchor and curve context', 'Jul/Aug-2026', 'public market data'],
       ['Damodaran country risk dataset (ctryprem)', 'UAE rating-based default spread and '
        'equity risk premium; the live implied US premium', '05-Jan-2026 / 01-Aug-2026',
        'pages.stern.nyu.edu/~adamodar'],
       ['DFM official index API + cross-validated Yahoo history', 'the DFM General Index '
        'series behind the beta regression (identical closes on all 307 overlapping '
        'sessions)', '2021-2026', 'api2.dfm.ae / finance.yahoo.com'],
       ['stc USD sukuk curve quotes', 'the GCC telecom credit spread behind the marginal '
        'cost of debt', '06-Aug-2026', 'public market data'],
       ['Mobily FY2025 and H1-2026 results announcements (Saudi Exchange issuer filings) and '
        'FY2025 earnings release', 'the justified price/earnings multiple and the peer '
        'dividend-yield anchor, both re-derived from filed earnings after the first edition '
        'used a stale aggregator figure: trailing EPS SAR 4.76 = FY2025 4.51 + H1-2026 2.32 - '
        'H1-2025 2.07, giving 12.9x on the 06-Aug-2026 close of SAR 61.30 (the market-cap '
        'route ties independently)', '06-Aug-2026', 'saudiexchange.sa / mobily.com.sa'],
       ['e& FY2025, H1-2025 and H1-2026 results releases', 'the second filing-derived peer '
        'multiple: trailing EPS AED 1.33 reported-attributable, 15.7x on the 06-Aug-2026 '
        'close of AED 20.98. NOTE the basis trap: e&\'s own H1-2026 release quotes the '
        'H1-2025 comparative on an ADJUSTED basis, which if used fabricates a 12.6x',
        '06-Aug-2026', 'eand.com'],
       ['du dividend-distribution disclosure to DFM, Ref RME/14/2026', 'the interim dividend '
        'timetable that determines which dividends are still in the anchor price: ex-date '
        '31-Jul-2026, record 03-Aug-2026, payment 21-Aug-2026', '2026-07-31',
        'du filing via DFM'],
       ['du disclosure announcement, "Extension of Federal Royalty Scheme for the Period '
        '2027-2029"', 'the settled 2027-2029 fiscal regime, including the express retention '
        'of the AED 1.8bn combined annual floor', '2026-07-24', 'du company disclosure'],
       ['IMF World Economic Outlook (Apr-2026) and July-2026 Update', 'UAE and regional '
        'growth/inflation; the war\'s macro baseline', 'Apr/Jul-2026', 'imf.org']],
      [2.30, 2.95, 0.85, 0.90])

H1('The full input register')
import collections
groups = collections.OrderedDict()
for k, rec in INP.items():
    ring = rec['ring'].split('/')[0]
    groups.setdefault(ring, []).append((k, rec))
for ring, items in groups.items():
    H2(f'{ring} layer — {len(items)} inputs')
    rows = [['Input', 'Value', 'Source and construction', 'Date']]
    for k, rec in items:
        rows.append([k, fmt(rec['value']), rec['source'], rec['date']])
    table(rows, [1.05, 1.15, 4.10, 0.75], size=7.4)

H1('The judgements, stated separately')
P('Every forecast rests on judgements. Each is stated here with what would overturn it.')
table([['Judgement', 'Basis', 'What would overturn it'],
       ['THE CENTRAL JUDGEMENT — what required return the business deserves. Framing 1 uses '
        'du\'s own measured beta; Framing 2 declines the terminal re-rating that implies and '
        'holds du\'s current trailing EV/EBITDA in perpetuity', 'the beta is du\'s own '
        '5-year weekly regression, gate-passed; the market multiple is du\'s own audited '
        'figures at the anchor', 'Both are published side by side and never averaged. Either '
        'is overturned by the other being right: a sustained re-rating toward the implied exit '
        'multiple, or a beta that behaves like the sector\'s ~0.80 rather than the measured '
        'figure'],
       ['The 2027-2029 fiscal regime is settled at the current take (base case rests on '
        'disclosed fact)', 'du\'s own disclosure of 24-Jul-2026 extending the 38% royalty + '
        '9% corporate tax and retaining the AED 1.8bn combined floor', 'nothing — this is '
        'disclosed, not assumed. What IS assumed is the period after 2029, priced as a tail'],
       ['Mobile recovers to 9,450k subscribers by end-2026, then adds ~210-310k/yr',
        'the company\'s own Q2 commentary (recovery underway, gross adds below pre-war), '
        'Dubai population re-acceleration', 'a re-opened conflict, or two quarters of '
        'negative total net adds'],
       ['Blended ARPU held roughly flat — the study\'s most fragile revenue judgement',
        'the last three prints are 65.3 / 63.4 / 63.4, and the mix decomposition shows the '
        'flat headline is two offsetting forces: a postpaid mix shift worth about +2.6% '
        'against per-leg erosion of about -2.4%. Holding flat assumes both continue',
        'a quarter in which the postpaid mix share FALLS back toward 20% while blended ARPU '
        'follows it down — the mix-exhaustion case, priced in the sensitivity table'],
       ['Every direct-cost unit rate is anchored on the H1-2026 reviewed actual and held '
        'flat, except mobile interconnect (-1.5%/yr) and mobile commission (+3.0%/yr)',
        'the like-for-like half-year pair measures interconnect per subscriber falling 4.2% '
        'and commission per subscriber rising 3.1%; three of the four second-half 2025 rates '
        'came in cheaper than the first half, so carrying an H1 rate into H2 overstates cost',
        'a disclosed reset in any single line — wholesale repricing, an ICT mix collapse, or '
        'interconnect per subscriber turning UP in a half-year pair'],
       ['Capex peaks at 15.5% of revenue then glides to 13%', 'commitments up ~14% in six '
        'months; the disclosed data-centre programme; no numeric company guidance '
        '(flagged)', 'company capex guidance above the path, or the programme extending '
        'past FY2028'],
       ['Payout stays at 98%', 'FY2024 actual 98%, FY2025 ~100%, interim raised through the '
        'war quarter', 'any declared cut'],
       ['Beta 0.472 from the own-index regression', 'five years weekly, R² 0.20, gate '
        'passed, composite cross-check 0.394', 'a structural re-rating of DFM correlations; '
        'the 0.57-0.80 alternatives are priced'],
       ['Terminal growth 2.5%', 'below long-run UAE nominal GDP (~4%+); duopoly at '
        'population-plus-inflation minus price erosion', 'sustained sub-2% revenue growth '
        'after the recovery completes'],
       ['Lease replacement charged at right-of-use depreciation', 'conservative: actual '
        'FY2025 lease additions ran at a quarter of depreciation', 'a disclosed structural '
        'shift to leased infrastructure']],
      [1.85, 2.60, 2.60], size=7.8)

H1('Negative results — what could not be sourced')
table([['What was sought', 'Where', 'Outcome'],
       ['UAE sovereign 5-year credit-default-swap quote', 'cbonds (403), '
        'worldgovernmentbonds (script-only), Damodaran (prints NA for UAE)',
        'NOT SOURCED — the market-spread premium basis uses the traded Abu Dhabi USD '
        'curve (+25bp) instead, disclosed in section 1.8'],
       ['Peer EV/EBITDA multiples, clean', 'public aggregators', 'net-debt figures missing '
        'or stale — the relative lens runs on P/E and dividend yield only, stated in 1.3'],
       ['Concluded TDRA licence-renewal terms', 'TDRA site, du IR disclosure list, DFM '
        'announcements, press', 'NOT FOUND at the sweep date — the licence note in the '
        'H1-2026 interims is the latest primary word; treated as catalyst 2'],
       ['Post-2029 royalty/corporate-tax regime', 'du filings, e& filings, MoF releases, '
        'Cabinet decisions', 'NOT DISCLOSED by anyone — so a reversion after 2029 is carried '
        'as a priced tail, not forecast. NOTE: the FIRST EDITION of this study recorded the '
        'post-2026 regime as undisclosed by du. That was WRONG — du published its own '
        '"Extension of Federal Royalty Scheme for the Period 2027-2029" on 24-Jul-2026, '
        'sixteen days before the sweep date, retaining the AED 1.8bn floor. The negative '
        'search was stale, not the disclosure absent'],
       ['Licence FEE / revenue-share terms on the renewed licence', 'TDRA announcement of '
        '12-Aug-2026, du filings', 'NOT DISCLOSED — the 20-year renewal (effective '
        '09-Aug-2026) publishes obligations only, no economics. The study holds the fee ratio '
        'at its historical level and flags that as an assumption'],
       ['Peer EV/EBITDA from issuer filings', 'peer annual reports and results releases',
        'net-debt detail per peer not retrieved within scope — the relative lens runs on '
        'earnings and dividend yield only, both sourced, rather than on an estimated multiple'],
       ['Trailing P/E for Zain and a priced multiple for Omantel', 'aggregators and issuer '
        'filings', 'REFUSED rather than passed through: the provider returned 11,365x for '
        'Zain, and Omantel\'s quoted multiple and yield are mutually inconsistent. Neither '
        'enters the peer frame as a number'],
       ['Wholesale and ICT unit KPIs (minutes, racks, MW, backlog)', 'AR2023-AR2025, all '
        'interims, all decks', 'never disclosed — both segments built top-down, flagged'],
       ['Numeric FY2026 capex guidance', 'Feb/Apr/Jul-2026 releases and decks',
        'not guided — house path anchored on disclosed commitments, flagged'],
       ['Analyst consensus on du', 'public sources', 'not cleanly retrievable (searches '
        'polluted by the exchange\'s own ticker); noted, not used']],
      [2.20, 2.30, 2.55], size=7.8)

H1('Corrections made after external review')
P('Four input errors were found by external audit and by re-audit, and are corrected in this '
  'edition rather than left standing. First, the justified price/earnings multiple: the first '
  'edition used 15.5x, described as a peer median. It was neither a median of the stated peer set '
  'nor a current figure — it was a January-2026 aggregator reading built on FY2024 earnings, '
  'published with an August-2026 date. Re-derived from the peers\' own filings it is 12.9x, and '
  'no median is claimed because only two of six peers survive as clean observations. Second, the '
  'dividend roll: the first edition netted only the final dividend PAID before the anchor, on the '
  'reasoning that the interim was declared but unpaid. The correct test is the EX-date, which had '
  'passed six sessions earlier, so AED 0.66 is now netted, not 0.40. Third, the staff-cost input: '
  'its own source note said it applied the prior year\'s second-half seasonal ratio, but the '
  'figure used implied 1.09 against a true 1.196 — which pushed total forecast operating expenses '
  'BELOW the audited prior-year actual while revenue grew, and made the margin expansion an '
  'artefact of one input. Fourth, and most consequential for how this study reads: du had '
  'disclosed the 2027-2029 royalty extension itself, before the first edition\'s sweep date, so '
  'the "contested judgement" that edition was built around was not contested. It is recast here '
  'as the required return.')

H1('A note on source discrepancies')
P('Three discrepancies a checking reader will find, all explained rather than smoothed: '
  '(1) FY2024 EBITDA is 6,469.8 in this study (the IFRS 18 re-presented comparative on the '
  'face of the FY2025 statements) but derives as 6,472.2 from the original FY2024 '
  'presentation — the re-presented basis is used so FY2024 and FY2025 sit on one '
  'presentation. (2) FY2024 federal royalty is 1,571.6 on the re-presented face but 1,675.9 '
  'in the original FY2024 statements (which included prior-period adjustments); the '
  're-presented figure is carried, and the FY2024 effective-rate note (44.7%) belongs to '
  'the original basis. (3) The company\'s IR capex (cost-additions basis, 2,274 for FY2025) '
  'differs from the cash-flow-statement capex (2,353.2) used in the model; the audited cash '
  'basis is used and the IR basis is quoted only where the IR intensity ratios are cited.')

H1('Disclosure')
P('This bibliography is part of an educational analysis and is not investment advice. '
  'Sources are quoted for verification; all errors of transcription are the authors\' own. '
  '© Testahil, 2026.', size=9.3)

doc.save(os.path.join(HERE, 'DU_Bibliography_09-08-2026.docx'))
print('wrote DU_Bibliography_09-08-2026.docx')
