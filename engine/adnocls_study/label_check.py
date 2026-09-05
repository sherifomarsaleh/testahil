"""External-reader scrub for the delivered ADNOCLS study: scan every paragraph and every
table cell of the .docx for internal-procedure vocabulary. Zero hits required."""
import re, sys
from docx import Document

BANNED = [
    (r'\bstep\s*0\b', 'step 0'),
    (r'\bstep\s*2a\b', 'step 2A'),
    (r'\bstep\s*\d', 'step + number'),
    (r'\bfour[- ]ring\b', 'four-ring'),
    (r'\brings?\b', 'ring'),
    (r'\binformation sweep\b', 'information sweep'),
    (r'\bsweep\b', 'sweep'),
    (r'\bgates?\b', 'gate'),
    (r'\bpromotion rule\b', 'promotion rule'),
    (r'\bstanding research protocol\b', 'standing research protocol'),
    (r'\bmc_v3\b', 'mc_v3'),
    (r'\bmarket_profiles\b', 'market_profiles'),
    (r'\bfitted_configs\b', 'fitted_configs'),
    (r'\bstudy_numbers\b', 'study_numbers'),
    (r'compute\.py', 'compute.py'),
    (r'\bdata_quality\b', 'data_quality'),
    (r'\bwacc_builder\b', 'wacc_builder'),
    (r'\bLONO\b', 'LONO'),
    (r'\bCRPS\b', 'CRPS'),
    (r'\bPIT\b', 'PIT'),
    (r'\bwidth_cal\b', 'width_cal'),
    (r'\bbootstrap block\b', 'bootstrap block'),
    (r'\bscale[- ]normali[sz]ed\b', 'scale-normalised'),
    (r'\bPARITY\b', 'PARITY'),
    (r'\bBOUNDARY\b', 'BOUNDARY'),
    (r'\bFAIL\b', 'FAIL'),
    (r'\bpersonas?\b', 'persona'),
    (r'\bprice target\b', 'price target'),
    (r'\bexpert persona library\b', 'expert persona library'),
]
CASE_SENSITIVE = {'LONO', 'CRPS', 'PIT', 'PARITY', 'BOUNDARY', 'FAIL'}


def cells(doc):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            yield f'paragraph {i}', p.text
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, c in enumerate(row.cells):
                if c.text.strip():
                    yield f'table {ti} r{ri}c{ci}', c.text


def xlsx_cells(path):
    """String cells only, formulas skipped — modelled on engine/stc_study/scrub.py [L-350]."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
    for ws in wb.worksheets:
        for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
            for ci, v in enumerate(row, start=1):
                if isinstance(v, str) and not v.startswith('=') and v.strip():
                    yield f'{ws.title}!r{ri}c{ci}', v
    wb.close()


def main(path):
    if path.lower().endswith(('.xlsx', '.xlsm')):
        blocks = list(xlsx_cells(path))
        hits = []
        for where, text in blocks:
            for pat, label in BANNED:
                flags = 0 if label in CASE_SENSITIVE else re.IGNORECASE
                for mt in re.finditer(pat, text, flags):
                    s = max(0, mt.start() - 45)
                    hits.append((label, where, text[s:mt.end() + 45].replace('\n', ' ')))
        if hits:
            print(f'{len(hits)} HITS')
            for label, where, ctx in hits:
                print(f'  [{label}] {where}: ...{ctx}...')
            return 1
        print(f'scrub clean: 0 hits across {len(blocks)} workbook string cells')
        return 0
    doc = Document(path)
    hits = []
    for where, text in cells(doc):
        for pat, label in BANNED:
            flags = 0 if label in CASE_SENSITIVE else re.IGNORECASE
            for mt in re.finditer(pat, text, flags):
                s = max(0, mt.start() - 45)
                hits.append((label, where, text[s:mt.end() + 45].replace('\n', ' ')))
    # a bare "rating" is allowed only where it plainly means a sovereign credit rating
    for where, text in cells(doc):
        for mt in re.finditer(r'\brating\b', text, re.IGNORECASE):
            s = max(0, mt.start() - 60)
            ctx = text[s:mt.end() + 60].lower()
            if not any(k in ctx for k in ('sovereign', 'credit', 'agency', 'moody')):
                hits.append(('rating (unqualified)', where,
                             text[s:mt.end() + 60].replace('\n', ' ')))
    if hits:
        print(f'{len(hits)} HITS')
        for label, where, ctx in hits:
            print(f'  [{label}] {where}: ...{ctx}...')
        return 1
    print('scrub clean: 0 hits across '
          f'{len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables')
    return 0


if __name__ == '__main__':
    sys.exit(main(sys.argv[1]))
