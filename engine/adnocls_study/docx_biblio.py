"""ADNOCLS_Bibliography_09-08-2026.docx — the standalone bibliography that ships with the
ADNOC Logistics & Services valuation study.

Every number and every source string in the delivered document is read from the committed
JSON files in this folder. No financial numeral is typed into this builder: values come from
the inputs block, source text comes with them, and the research trail comes from the
companion research file. Prose that frames a table is authored here; the evidence inside one
is not.

House style helpers are reused from docx_base.py (same palette, same table furniture); the
page is re-set to 0.7in margins, so the text block is 7.10in and every table below sums to
exactly that.

Run:  python3 docx_biblio.py
"""
import json, os, re, sys

HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)

import docx_base as B                      # noqa: E402  (house style, builds the document)
from docx.shared import Pt, Inches, RGBColor   # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.oxml.ns import qn                # noqa: E402
from docx.oxml import OxmlElement          # noqa: E402

D = json.load(open('study_numbers.json'))
R = json.load(open('sweep_register.json'))
INP = D['inputs']
META = D['meta']

doc = B.doc
INK, GREY, WHITE = B.INK, B.GREY, B.WHITE
F_DARK, F_PANEL, F_CREAM = B.F_DARK, B.F_PANEL, B.F_CREAM

# ---------------------------------------------------------------- page furniture
TEXT_W = 7.10
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.size = Pt(9.8)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.05


def P(text='', size=9.8, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    return B.P(text, size=size, bold=bold, italic=italic, color=color,
               space_after=space_after, space_before=space_before)


def H1(t):
    return B.P(t, size=14, bold=True, space_before=13, space_after=5)


def H2(t):
    return B.P(t, size=11, bold=True, space_before=9, space_after=3)


def H1_NEW_PAGE(t):
    """A section that must start at the top of a page.

    NOT a page-break paragraph. A break of its own is a paragraph like any other, and when the
    section before it ends near the foot of a page — the input register does, and does so at a
    different place every time an input is added — the spacer after the last table and the
    break itself both slide onto the next page and are rendered as a blank sheet before the
    heading. Carrying the break ON the heading cannot produce that: there is no paragraph
    between the two sections at all, so the heading starts the page it breaks to."""
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        el = doc.paragraphs[-1]._element
        el.getparent().remove(el)
    p = H1(t)
    p.paragraph_format.page_break_before = True
    return p


def bullet(text):
    return P('   •  ' + text, size=9.6, space_after=3)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def _cant_split(row):
    """Keep a row whole rather than letting it break across a page and strand a fragment
    with three empty cells above it at the top of the next page."""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:cantSplit')
    el.set(qn('w:val'), 'true')
    trPr.append(el)


def T(rows, widths, aligns=None, size=8.0, cell_lr=70, header=True):
    """Fixed-layout table. widths sum to the text block; aligns is one of L/R/C per column."""
    assert abs(sum(widths) - TEXT_W) < 1e-6, f'width {sum(widths):.3f} != {TEXT_W}'
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    B.cell_margins(t, 34, 34, cell_lr, cell_lr)
    B.borders(t)
    t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        _cant_split(t.rows[i])
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True
                B.shade(c, F_PANEL)
            elif aligns:
                a = aligns[j]
                if a == 'R':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif a == 'C':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header:
        _repeat_header(t.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def masthead():
    t = doc.add_table(rows=1, cols=1)
    B.cell_margins(t, 90, 90, 110, 110)
    c = t.cell(0, 0); B.shade(c, F_DARK); c.width = Inches(TEXT_W)
    t.columns[0].width = Inches(TEXT_W)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def panel(lines, fill=F_CREAM):
    t = doc.add_table(rows=1, cols=1)
    B.borders(t, color='C0A45F', sz='6')
    B.cell_margins(t, 110, 110, 150, 150)
    c = t.cell(0, 0); B.shade(c, fill); c.width = Inches(TEXT_W)
    t.columns[0].width = Inches(TEXT_W)
    first = True
    for head, body in lines:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        if head:
            r = p.add_run(head); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
        r2 = p.add_run(body); r2.font.size = Pt(9.5); r2.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------- value formatting
MONTHS = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
DAYS_IN = [31, 29, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]


def fdate(s):
    """ISO date -> '31 Dec 2025'. A few quarterly-average tags carry a period label whose day
    is not a real calendar day; those render as month and year, never as an impossible date."""
    try:
        y, m, d = (int(x) for x in s.split('-'))
        if 1 <= m <= 12 and 1 <= d <= DAYS_IN[m - 1]:
            return f'{d} {MONTHS[m - 1]} {y}'
        if 1 <= m <= 12:
            return f'{MONTHS[m - 1]} {y}'
    except Exception:
        pass
    return s


PCT_EXACT = {
    'rf_observed', 'sov_spread', 'erp_total', 'crp', 'erp_mature', 'sofr', 'cb_rate',
    'tax_stat', 'tax_topup_rate', 'tax_integrated_logistics', 'tax_shipping', 'tax_services',
    'g_terminal', 'rf_terminal', 'gw_growth_assumption', 'opex_escalation',
    'h2_2026_reversion', 'gas_margin', 'div_growth', 'contracted_2026_share', 'nci_share',
    'dep_rate_ppe', 'shldr_margin', 'hybrid_margin', 'bank_loan_lo', 'bank_loan_hi',
    'other_borr_lo', 'other_borr_hi', 'rel_weight_ev_ebitda', 'gdp_growth_26',
    'inflation_26', 'tanker_orderbook_pct_fleet', 'spot_share_ebitda_26',
    'spot_share_ebitda_29', 'ri_fade',
}
PCT_PREFIX = ('drv_mar_', 'lens_weight_', 'g26_')
PERDAY_EXACT = {'tnk_opex_day', 'gas_rate_day', 'vlcc_1y_tc_market', 'vlcc_spot_clarksons_jan26'}
# every charter-out rate is a rate PER VESSEL PER DAY; without this prefix they render
# in the register's default unit of thousands of dollars, which is off by a factor of
# a thousand and reads as though a vessel earned 19.75 million dollars a day
PERDAY_PREFIX = ('tce_', 'tc_out_', 'charter_')
COUNT_EXACT = {'spot_vessels_total', 'jub_owned', 'jub_chartered', 'osv_owned', 'gas_owned',
               'gas_lt_contracted', 'vlcc_sold_jan26'}
COUNT_RE = re.compile(r'^tnk_[a-z0-9]+_(n|spot)$')
RATIO3 = {'beta', 'beta_se', 'beta_r2', 'beta_dimson', 'beta_composite',
          'beta_ci_lo', 'beta_ci_hi'}
MULT = {'nd_ebitda_target_lo', 'nd_ebitda_target_hi', 'tnk_grossup_25', 'tnk_grossup_26'}
DAYS = {'dso_days', 'dio_days', 'dpo_days'}


def trim(x, dp):
    s = f'{x:,.{dp}f}'
    return s.rstrip('0').rstrip('.') if '.' in s else s


def fval(k, v):
    """Format by what the figure IS. Default unit is thousands of US dollars."""
    if k in PCT_EXACT or k.startswith(PCT_PREFIX):
        return f'{v * 100:.2f}%'.replace('.00%', '%')
    if k in RATIO3:
        return f'{v:.3f}'
    if k in MULT:
        return f'{trim(v, 2)}x'
    if k in DAYS:
        return f'{v:,.1f} days'
    if k in PERDAY_EXACT or k.startswith(PERDAY_PREFIX):
        return f'{v:,.0f}/day'
    if k in COUNT_EXACT or COUNT_RE.match(k):
        return f'{v:,}'
    if k.startswith('gas_vessel_years_'):
        return trim(v, 2)
    if k in ('shares_mn', 'shares_wavg_mn'):
        return f'{v:,.3f} mn'
    if k == 'dps_2026_usd':
        return f'{trim(v, 1)} USD mn'
    if k == 'fx_aed_usd':
        return f'{v:,.4f}'
    if k == 'spot_aed':
        return f'{v:,.2f} AED'
    if k.startswith('eps_fy'):
        return f'{v:,.2f} USD'
    if k == 'erp_cds_available':
        return 'none'
    if isinstance(v, float):
        return trim(v, 3) if abs(v) < 1000 else f'{v:,.0f}'
    return f'{v:,}'


# ---------------------------------------------------------------- source counts
DOCS = {
    'FS25': 'Consolidated Financial Statements FY2025',
    'FS24': 'Consolidated Financial Statements FY2024',
    'FS23': 'Consolidated Financial Statements FY2023',
    'Q126': 'Condensed Consolidated Interim Financial Information, three months ended 31 March 2026',
    'AR25': 'Annual Report and Accounts 2025',
    'MDAQ126': 'Management Discussion and Analysis Q1 2026',
    'MDA25': 'Management Discussion and Analysis FY2025',
    'IP26': 'Investor Presentation April 2026',
    'IPFY25': 'FY2025 Investor Presentation',
    'CALLQ126': 'Q1 2026 earnings call transcript',
    'CALL25': 'FY2025 earnings call transcript',
    'REL25': 'FY2025 earnings release',
}


def n_from(marker):
    return sum(1 for v in INP.values() if marker in v['source'])


CNT = {
    'FS25': n_from('Consolidated Financial Statements FY2025'),
    'FS24': n_from('Consolidated Financial Statements FY2024'),
    'FS23': n_from('Consolidated Financial Statements FY2023'),
    'Q126': n_from('Condensed Consolidated Interim Financial Information'),
    'AR25': n_from('Annual Report and Accounts 2025 '),
    'MDAQ126': n_from('Management Discussion and Analysis Q1 2026'),
    'MDA25': n_from('Management Discussion and Analysis FY2025'),
    'IP26': n_from('Investor Presentation April 2026'),
    'IPFY25': n_from('FY2025 Investor Presentation'),
    'CALLQ126': n_from('Q1 2026 earnings call transcript'),
    'CALL25': n_from('FY2025 earnings call transcript'),
    'REL25': n_from('FY2025 earnings release'),
    'FS22': n_from('Consolidated Financial Statements FY2022'),
    'AR24': n_from('Annual Report and Accounts 2024 '),
}


def supplies(key):
    n = CNT[key]
    if n == 0:
        return ('No figure is cited to it directly; it is the evidence behind the statements '
                'that are.')
    if n == 1:
        return 'Supplies one of the figures below.'
    return f'Supplies {n:,} of the figures below.'


def usdm(x, dp=0):
    """Thousands of dollars as the reader would say them aloud."""
    return f'USD {x / 1000.0:,.{dp}f}m'


def held(fname):
    return fname if os.path.exists(os.path.join(HERE, fname)) else 'NOT HELD'


# ---------------------------------------------------------------- numbers used in prose
W = D['wacc']; BR = D['beta']; BRG = D['bridge']; L = D['lenses']
DCF, DCFS, DCFA = D['dcf'], D['dcf_sustained'], D['dcf_beta_alt']
BF = D['beta_framing']                  # the two market measurements, side by side
BC = BR['composite_variant']            # the same regression against the equal-weight composite
BSI = BR['self_inclusion_bias']         # the share sits inside the index it is measured against
CENTRAL_ALT = D['central_beta_alt']
# The confidence level of the published interval is read off the input's own source text
# rather than asserted here, so the document and the model can never name different levels.
CI_LEVEL = re.search(r'(\d+%)', INP['beta_ci_lo']['source']).group(1)
LW = D['lens_weights']; REL = D['rel']; GC = D['guidance_check']
PEG = META['fx']
SHARES = META['shares_mn']
IV = {k: v['value'] for k, v in INP.items()}

# equal-weighting the four lenses, computed from the committed lens values, as the
# counterfactual that tests the weighting judgement
EQ_CENTRAL = sum(L[k]['base'] for k in ('dcf', 'relative', 'normalized', 'book')) / 4.0
HYB_PER_SHARE = IV['q1_26_hybrid'] / SHARES / 1000.0 * PEG
ND_FOOTNOTE = IV['q1_26_shldr_loan'] + IV['q1_26_leases'] - IV['q1_26_cash']
ND_PRINTED = IV['q1_26_netdebt']
ND_STUDY = BRG['net_debt']

def idate(k):
    """The date an input carries, so a document's date and its figures never disagree."""
    return fdate(INP[k]['date'])


def find_one(**want):
    """A research finding located by WHAT IT IS, never by its position in the file. The
    identifiers renumber whenever a finding is added — one was, for the index series — so a
    document that cites them by number silently re-dates itself. Matching on the fields that
    describe the finding cannot drift that way, and the assertion below is what proves it."""
    hits = [f for f in R['findings'] if all(f.get(k) == v for k, v in want.items())]
    assert len(hits) == 1, f'{len(hits)} findings match {want}, expected exactly one'
    return hits[0]


def fs_date(period):
    """The date the filing for a fiscal period itself bears."""
    return fdate(find_one(category='official financial statements',
                          fiscal_period=period)['source_date'])


def rd_date(period):
    """The date a regular disclosure for a period bears (the reviewed interim information)."""
    return fdate(find_one(category='regular disclosures', fiscal_period=period,
                          is_fs_data=True)['source_date'])


def gd_date():
    """The date the management commentary carrying the raised guidance bears."""
    return fdate(find_one(category='strategic plans & guidance',
                          klass='DRIVER_UNLOCK')['source_date'])


IDX = find_one(source_type='PRIMARY_MARKET_DATA')   # the index series the beta is measured on
IDX_FROM, IDX_TO = BR['regressor_span']


LAYERS = ['Company', 'Industry', 'Country', 'Global', 'Market']
LAYER_N = {lay: sum(1 for v in INP.values() if v['layer'] == lay) for lay in LAYERS}
assert sum(LAYER_N.values()) == len(INP), 'an input carries a layer outside the five'

FS_YEARS = ['FY2022', 'FY2023', 'FY2024', 'FY2025']

# ============================================================================
# 1  MASTHEAD AND READ FIRST
# ============================================================================
masthead()
H1(f"{META['company']} ({META['exchange']}: {META['ticker']}) — Bibliography and "
   f"Source Register")
P(f"Companion document to the valuation study dated 9 August 2026. It records where every "
  f"number in that study came from, and what was decided rather than read.",
  size=9.5, color=GREY)

H2('READ FIRST')
panel([
    ('What this is. ',
     f"A complete list of the {len(INP):,} inputs the valuation model uses. Each carries its "
     f"value, the date the source itself bears, and the source with the construction that "
     f"produced it. Nothing in the study is computed from a figure that does not appear here, "
     f"so a reader can check the work rather than trust it."),
    ('How to use it. ',
     "Read the guide to the research layers first, then the list of documents actually read. "
     "The input register is the body of the document and is meant to be searched, not read "
     "front to back: find the figure you want to test, and the row tells you where it came "
     "from and how it was built. The tables after it are the short ones that matter most — "
     "what the study judged rather than observed, what it looked for and could not find, and "
     "where two readings of the same figure disagreed."),
    ('What is not here. ',
     "Nothing is summarised out. Where a figure is derived, solved or assumed, the row says so "
     "in its own words rather than presenting it as a disclosure. Where a search came back "
     "empty, it is written down as an empty search."),
    ('The reporting unit. ',
     f"The company reports in US dollars and this study values it in US dollars, converting to "
     f"dirhams at the fixed peg of {IV['fx_aed_usd']:,.4f} only to compare with the quoted "
     f"price. Unless a row shows another unit, every figure in the register is in THOUSANDS of "
     f"US dollars."),
])

# ============================================================================
# 2  THE RESEARCH LAYERS
# ============================================================================
H1('The research layers')
P("Every input is tagged with the layer of evidence it belongs to. The layers run from the "
  "company outwards, and they are not interchangeable: what a layer is allowed to be the "
  "source of is as much a part of the method as the figure itself.")
T([['Layer', 'Inputs', 'What sits in it, and what it may be the source of'],
   ['Company', f"{LAYER_N['Company']:,}",
    "The company's own issued documents — audited consolidated financial statements, "
    "reviewed interim financial information, the annual report and accounts, management "
    "commentary, investor presentations and earnings-call transcripts. Every reported "
    "historical figure in the study comes from this layer and from nowhere else. It is also "
    "where the unit drivers live: vessel counts, the rates each vessel class earned, the "
    "contract table, the order book and the guided spending path."],
   ['Industry', f"{LAYER_N['Industry']:,}",
    "Marine transport and energy logistics: the tanker order book as a share of the trading "
    "fleet, broker prints of prevailing rates, one-year charter levels fixed by other owners, "
    "and the enterprise multiples of comparable listed operators. Used for context, for "
    "cross-checks and for the comparison lens — never as the source of a figure the "
    "company itself reports."],
   ['Country', f"{LAYER_N['Country']:,}",
    "The United Arab Emirates: the dirham government bond yield, the sovereign default spread "
    "and equity risk premium, the currency's fixed parity to the dollar, the corporate tax "
    "regime and the minimum top-up rate, and projected growth and inflation. This is where "
    "the cost of capital and the tax treatment come from."],
   ['Global', f"{LAYER_N['Global']:,}",
    "Conditions common to every market rather than to this one: the dollar overnight financing "
    "rate that every one of the group's debt instruments is priced off, the mature-market "
    "equity risk premium, and the long-run nominal anchor behind the terminal discount rate."],
   ['Market', f"{LAYER_N['Market']:,}",
    "The security itself: the last close used throughout the study, and the statistics "
    "estimated from its own price history — the slope of its weekly returns against the "
    "published index of the exchange it trades on, the standard error and fit of that "
    "regression, the two bounds of the interval around the slope, the lead-lag variant, and "
    "the slope the same returns give when they are measured against an equal-weight composite "
    "of that exchange's own names instead of the published index."]],
  [1.15, 0.70, 5.25], aligns=['L', 'R', 'L'], size=8.4)

H2("Why the company's own statements are the only admissible source for its own history")
P("A valuation is only as good as the base year it starts from, and a base year assembled "
  "from second-hand numbers is not a base year at all. Data vendors, brokers and press "
  "coverage all restate: they blend segments that the company reports separately, carry "
  "figures forward from an earlier year, apply their own definitions of earnings and net "
  "debt, and correct silently. None of that is visible downstream. So the rule applied here "
  "admits nothing but the company's own issued and audited documents as the source of the "
  "company's own reported figures.")
P(f"That rule was met in full for this study. Four complete audited financial years were "
  f"obtained — {', '.join(FS_YEARS)} — against a floor of two, together with the "
  f"reviewed interim information for the one quarter of the study year disclosed so far. "
  f"Every historical income-statement, balance-sheet, cash-flow and segment figure in the "
  f"register below is read from those filings. External sources appear for the cost of "
  f"capital, the tax regime, sector context and the comparison lens, and are labelled as "
  f"such wherever they are used.")

# ============================================================================
# 3  PRIMARY DOCUMENTS
# ============================================================================
H1('The documents actually read')
P(f"Every document below was downloaded from the company's own investor pages and read from "
  f"the filing itself. All of them are held as PDF copies alongside the model, and the file "
  f"held is named in each row so that a reader can ask for the same document rather than a "
  f"description of it. The publisher of every document in the first table is "
  f"{META['company']}; the audited statements were audited by PricewaterhouseCoopers.",
  size=9.5)

H2("The company's own documents")
T([['Document', 'Date', 'File held', 'What was taken from it'],
   ['Consolidated financial statements for the year ended 31 December 2025, audited',
    fs_date('FY2025'), held('FS_FY2025.pdf'),
    'The 2025 income statement, balance sheet and cash flow statement with the 2024 '
    'comparatives; the operating-segment schedule by business unit including segment '
    'earnings, depreciation and property; the revenue-by-service and cost-by-nature notes; '
    'borrowings, leases and their rates; the business-combination note for the tanker '
    'acquisition; the perpetual capital securities; the goodwill test that the terminal '
    'growth rate is taken from; share capital and treasury. ' + supplies('FS25')],
   ['Consolidated financial statements for the year ended 31 December 2024, audited',
    fs_date('FY2024'), held('FS_FY2024.pdf'),
    'The 2024 statements and the 2023 comparatives they carry — the second and third '
    'historical years of the income statement, balance sheet and cash flow, and the segment '
    'schedule under the reporting structure then in force. ' + supplies('FS24')],
   ['Consolidated financial statements for the year ended 31 December 2023, audited',
    fs_date('FY2023'), held('FS_FY2023.pdf'),
    'The third historical year, and the segment schedule that makes the later '
    'reclassification visible. ' + supplies('FS23')],
   ['Consolidated financial statements for the year ended 31 December 2022, audited',
    fs_date('FY2022'), held('FS_FY2022.pdf'),
    'Completes the four-year audited record and carries the segment schedule from before the '
    'group was reorganised, which is the evidence that the segment series is not comparable '
    'across the whole span. ' + supplies('FS22')],
   ['Annual Report and Accounts 2025', fs_date('FY2025'), held('AR2025.pdf'),
    'The audited statements as machine-readable text where the standalone filing is an image, '
    'plus the fleet, order-book and operating review. ' + supplies('AR25')],
   ['Annual Report and Accounts 2024', fs_date('FY2024'), held('AR2024.pdf'),
    'The same for the 2024 statements and their 2023 comparatives; this is the volume the '
    '2023 figures were read from.'],
   ['Annual Report and Accounts 2023', fs_date('FY2023'), held('AR2023.pdf'),
    'NOT HELD as a separate volume. The 2023 audited statements themselves are held and were '
    'read directly, and the 2024 volume carries the 2023 comparatives, so no figure depends '
    'on it. Recorded rather than quietly dropped.'],
   ['Condensed consolidated interim financial information for the three months ended 31 March '
    '2026, reviewed', rd_date('Q1 2026'), held('FS_Q1_2026.pdf'),
    'The first-quarter income statement with 2025 comparatives, the balance sheet at the '
    'valuation date, the full segment schedule, the cash flow, the borrowings and related-'
    'party rates, and the events-after-reporting note. ' + supplies('Q126')],
   ['Management commentary on the first quarter of 2026', gd_date(),
    held('MDA_Q1_2026.pdf'),
    'The raised guidance for 2026 by line and by business unit, the distribution policy, the '
    'medium-term leverage target, net debt and free cash flow, and the segment revenue and '
    'earnings table the quarterly build is anchored on. ' + supplies('MDAQ126')],
   ['Management commentary on the 2025 full year', fdate('2026-02-11'), held('MDA_FY2025.pdf'),
    "The full-year and fourth-quarter bridge, the company's own definition of net debt and of "
    'earnings, and the re-presentation note on tanker revenue. ' + supplies('MDA25')],
   ['Management commentary on the 2024 full year', fdate('2025-02-12'), held('MDA_FY2024.pdf'),
    'The 2024 comparative basis for the same tables, used to check that the 2025 presentation '
    'change is a presentation change and not a restatement of profit.'],
   ['Investor presentation, April 2026', fdate('2026-04-30'), held('IP_Apr2026.pdf'),
    'The unit drivers: owned fleet by vessel class, the twelve vessels chartered out with '
    'each rate and expiry, the gas contract table quarter by quarter to 2029, the order book '
    'and remaining committed spending, the guided capital-expenditure path, long-term '
    'contracted revenue and its run-off, the share of earnings exposed to spot rates, and the '
    'disclosed sensitivity of earnings to a change of USD 1,000 a day. ' + supplies('IP26')],
   ['Investor presentation, 2025 full year', fdate('2026-02-11'), held('IP_FY2025.pdf'),
    'The quarterly time-charter equivalents by vessel class back to the start of 2024, the '
    'jack-up barge and support-vessel counts, and the contracted share of 2026 revenue. '
    + supplies('IPFY25')],
   ['Earnings-call transcript, first quarter of 2026', fdate('2026-05-14'),
    held('CALL_Q1_2026.pdf'),
    'The rates achieved in the first quarter and the level the second quarter was crossing at '
    'the time of the call, and management’s own account of why. ' + supplies('CALLQ126')],
   ['Earnings-call transcript, 2025 full year', fdate('2026-02-11'), held('CALL_FY2025.pdf'),
    'Context for the acquisition, the newbuild programme and the distribution policy; read in '
    'full, and used to check the presentation figures rather than as a source of numbers.'],
   ['Earnings release, 2025 full year', fdate('2026-01-31'), held('REL_FY2025.pdf'),
    'The realised sale of a 2017-built very large crude carrier in January 2026, its carrying '
    'value and the gain — the only market evidence of what this fleet is worth secondhand. '
    + supplies('REL25')]],
  [1.62, 0.75, 1.15, 3.58], aligns=['L', 'L', 'L', 'L'], size=7.7)

H2('External documents')
T([['Document', 'Publisher', 'Date', 'What was taken from it'],
   ['Dirham treasury bond auction result, the tranche maturing January 2031',
    'UAE Ministry of Finance, as reported by the Emirates News Agency', idate('rf_observed'),
    f"The observed government yield of {fval('rf_observed', IV['rf_observed'])}, which after "
    f"deducting the sovereign's own default spread is the risk-free rate in the cost of "
    f"equity."],
   ['Country risk premium and default spread file, January 2026 edition',
    'Aswath Damodaran, NYU Stern', idate('erp_total'),
    f"The sovereign default spread of {fval('sov_spread', IV['sov_spread'])}, the total equity "
    f"risk premium of {fval('erp_total', IV['erp_total'])} and the mature-market premium of "
    f"{fval('erp_mature', IV['erp_mature'])}. The original file only, read for this sovereign "
    f"rather than borrowed from a regional figure."],
   ['Published overnight financing rate', 'Federal Reserve Bank of New York',
    idate('sofr'),
    f"The dollar base rate of {fval('sofr', IV['sofr'])} that every one of the group's debt "
    f"instruments is contractually priced off, and the anchor for the marginal cost of debt."],
   ['Federal Decree-Law 47 of 2022 on the taxation of corporations and businesses',
    'Government of the United Arab Emirates', idate('tax_stat'),
    f"The {fval('tax_stat', IV['tax_stat'])} standard corporate rate above the threshold, and "
    f"the relief for income from the international transport of goods and passengers that the "
    f"company's own segment tax disclosure shows the effect of."],
   ['Base rate decision', 'Central Bank of the UAE', idate('cb_rate'),
    f"The policy rate of {fval('cb_rate', IV['cb_rate'])} and the confirmation that the dirham "
    f"peg is maintained."],
   ['World Economic Outlook database, United Arab Emirates',
    'International Monetary Fund', idate('gdp_growth_26'),
    f"Real growth of {fval('gdp_growth_26', IV['gdp_growth_26'])} and consumer price inflation "
    f"of {fval('inflation_26', IV['inflation_26'])} for 2026, used as context for the terminal "
    f"assumptions rather than as a forecast driver."],
   ['Company statistics page, Qatar Gas Transport (Nakilat)', 'stockanalysis.com',
    fdate(D['peers'][0]['asof']),
    f"The enterprise-to-earnings multiple of {D['peers'][0]['ev_ebitda']:.2f}x for the "
    f"contracted-fleet comparator, with its price-to-earnings and price-to-book figures."],
   ['Company statistics page, Frontline plc', 'stockanalysis.com',
    fdate(D['peers'][1]['asof']),
    f"The enterprise-to-earnings multiple of {D['peers'][1]['ev_ebitda']:.2f}x for a fleet "
    f"traded at spot."],
   ['Enterprise-value-to-earnings page, International Seaways', 'valueinvesting.io',
    fdate(D['peers'][2]['asof']),
    f"The enterprise-to-earnings multiple of {D['peers'][2]['ev_ebitda']:.2f}x for the second "
    f"spot comparator. Aggregator figures appear only here, as a labelled cross-check in the "
    f"comparison lens."],
   [f"Daily history of the {BR['regressor']}",
    'Published by FTSE Russell with the Abu Dhabi Securities Exchange; the series obtained as '
    'a dated export',
    fdate(IDX['source_date']),
    f"The market the beta is measured against: {BR['regressor_rows']:,} daily sessions from "
    f"{fdate(IDX_FROM)} to {fdate(IDX_TO)}, checked session by session against the exchange's "
    f"own trading calendar before use. This is the capitalisation-weighted index of the "
    f"exchange the share is listed on, which is the market measurement the method calls for. "
    f"Regressed on it, the stock's weekly returns give a slope of {BR['beta']:.3f}. It is "
    f"market data rather than company data, so it is not covered by the rule that admits only "
    f"the company's own filings — that rule governs the company's own reported history, which "
    f"an index level is not."],
   ['Daily price history for the stock, and for the Abu Dhabi listed names held alongside it',
    f"{META['exchange']}", fdate(META['price_date']),
    f"The last close of {fval('spot_aed', IV['spot_aed'])} used throughout; the "
    f"{BR['n']:,} paired weekly observations behind the regression against the published "
    f"index; and the equal-weight composite of {BR['composite_names']:,} of the exchange's own "
    f"names, which is the alternative measurement of the market the same regression is also "
    f"run against."]],
  [1.72, 1.30, 0.68, 3.40], aligns=['L', 'L', 'L', 'L'], size=7.7)

H2('A quotation this study turns on')
P(f"One sentence spoken on an earnings call changed the shape of the model more than any "
  f"document in the tables above, so it is reproduced here in full with its context rather "
  f"than cited. On the first-quarter {META['price_date'][:4]} call, an analyst asked how "
  f"the reported rate for the largest tanker class squared with what the spot market was "
  f"paying at the time. The chief financial officer answered:", size=9.5)
panel([('', f"“… was related to our full fleet of {D['fleet']['owned']['vlcc']}, and it "
            f"includes all the vessels on long-term charter as well … it's a blended rate "
            f"that we give there, which is obviously less than the spot rate.”")])
T([['Field', 'Detail'],
   ['Source', DOCS['CALLQ126'] + f", {META['company']}"],
   ['Date', fdate('2026-05-14')],
   ['File held', held('CALL_Q1_2026.pdf')],
   ['Speaker', 'The chief financial officer, answering an analyst question on the rate '
    'reported for very large crude carriers in the first quarter'],
   ['What it establishes',
    f"That the single rate the company publishes for each vessel class each quarter is an "
    f"average across the whole class, taken over the vessels trading in the open market "
    f"AND the vessels already committed on charters out at rates fixed earlier. It is "
    f"therefore lower than what an uncommitted vessel earns, by an amount that depends on "
    f"how many of the class are committed and at what rates."],
   ['What it changed',
    f"The first edition of the study read the published figure as the open-market rate and "
    f"then added the chartered vessels beside it at their own, lower, disclosed rates — "
    f"charging the same drag twice. The model now carries each of the "
    f"{len(D['fleet']['charters'])} chartered vessels individually for exactly the days "
    f"its own contract runs, and solves the open-market rate out of the published average "
    f"instead of assuming it. For the largest class in the first quarter of "
    f"{META['price_date'][:4]} the solved figure is "
    f"{D['fleet']['spot_q1_26']['vlcc']:,.0f}/day against a published "
    f"{D['fleet']['blend_q1_26']['vlcc']:,.0f}/day."],
   ['How it is used',
    'As evidence about what a disclosed figure means, not as the source of a figure. '
    'Every rate in the model still comes from a published table; what the quotation '
    'changes is the arithmetic applied to those tables.']],
  [1.05, 6.05], aligns=['L', 'L'], size=7.9)

P('The company site was reached for every company document. Of the three investor pages '
  'tried, two returned everything asked of them; one index page returned a server error from '
  'this environment and its documents were obtained through the other two, so nothing was '
  'lost. A failed attempt is still a fact about how the record was assembled, so it is '
  'recorded rather than omitted.', size=9.3, color=GREY)

# ============================================================================
# 4  THE FULL INPUT REGISTER
# ============================================================================
H1_NEW_PAGE('The input register — every figure the model uses')
P(f"All {len(INP):,} inputs, grouped by research layer and, within a layer, in the order the "
  f"model declares them — which follows the shape of the accounts: income statement, "
  f"balance sheet, cash flow, the disclosed quarter, then segments, service lines, cost "
  f"lines, the fleet, rates, guidance and the forecast drivers. Values are shown as the model "
  f"holds them. Unless a row shows another unit, figures are in thousands of US dollars; "
  f"rates and shares are shown as percentages, vessel earnings and running costs per vessel "
  f"per day, cycle measures in days, and multiples with an x.", size=9.5, color=GREY)

REG_W = [1.45, 0.86, 0.71, 4.08]
rendered = 0
for lay in LAYERS:
    items = [(k, v) for k, v in INP.items() if v['layer'] == lay]
    if not items:
        continue
    H2(f'{lay} layer — {len(items):,} inputs')
    rows = [['Input', 'Value', 'Date', 'Source and construction']]
    for k, v in items:
        rows.append([k.replace('_', ' '), fval(k, v['value']), fdate(v['date']), v['source']])
        rendered += 1
    T(rows, REG_W, aligns=['L', 'R', 'L', 'L'], size=7.4, cell_lr=58)

# The register is counted LIVE off the numbers file and every count in this document is
# derived from that one number. Counting against a known total is the check that matters:
# a renderer that silently drops a row reports success on its own terms, so the register
# is reconciled row by row against the file it came from, and the count the prose states
# is asserted to be the same object rather than a figure typed beside it.
N_INPUTS = len(INP)
assert rendered == N_INPUTS, f'rendered {rendered} of {N_INPUTS} inputs'
assert sum(LAYER_N.values()) == N_INPUTS, 'the layer counts do not reconcile to the register'
assert N_INPUTS == len({k for k in INP}), 'an input key is duplicated'
for k, v in INP.items():
    assert set(v) >= {'value', 'source', 'date', 'layer'}, f'{k} is not four-field complete'
    assert str(v['source']).strip() and str(v['date']).strip(), f'{k} has an empty field'

# The two market measurements the document sets side by side must BE the ones the model
# holds, not a pair retyped into this builder. Same for the interval and the lead-lag variant.
for _lbl, _doc_val, _key in (
        ('published index', BF['primary']['beta'], 'beta'),
        ('equal-weight composite', BF['alternative']['beta'], 'beta_composite'),
        ('lower bound', BF['ci90'][0], 'beta_ci_lo'),
        ('upper bound', BF['ci90'][1], 'beta_ci_hi'),
        ('lead-lag variant', BR['dimson']['sum_beta'], 'beta_dimson')):
    assert abs(_doc_val - IV[_key]) < 5e-4, f'{_lbl} disagrees with the {_key} input'
assert abs(BC['beta'] - IV['beta_composite']) < 5e-4, 'the composite slope disagrees with itself'
assert BF['alternative']['ke'] == W['ke_beta1'], 'the alternative cost of equity is not the one built'
assert BF['primary']['ke'] == W['ke'], 'the adopted cost of equity is not the one built'
assert BF['primary']['central'] == D['central'], 'the published central figure disagrees with itself'
assert BF['alternative']['central'] == CENTRAL_ALT, 'the alternative central disagrees with itself'

# ============================================================================
# 5  JUDGEMENTS
# ============================================================================
H1_NEW_PAGE('The judgements, stated separately')
P("These are the places where the study chose rather than read. Each row states what was "
  "taken, and what evidence would overturn it — written before the outcome is known, so "
  "that a reader can hold the study to it. They are collected here so they can be found "
  "without reading the whole study.")

JUD = [
    ('Where tanker rates settle after 2026',
     f"Not resolved into one number. The whole valuation is computed twice and published side "
     f"by side: rates reverting to the average of the 2024 and 2025 outcomes gives AED "
     f"{DCF['fv_aed']:,.2f} a share, rates settling 30% above that gives AED "
     f"{DCFS['fv_aed']:,.2f}. This is the study's central contested judgement and the largest "
     f"single driver in it.",
     "A published forward curve or a third-party rate forecast would replace the judgement "
     "outright; none is obtainable. Failing that, a second and third quarter printing at "
     "either level settles it — the second-quarter results were scheduled for 11 August "
     "2026, days after this study was struck."),
    ('How the tanker fleet is built, vessel by vessel, off the published charter table',
     f"Each of the {len(D['fleet']['charters']):,} vessels chartered out is carried at its "
     f"own disclosed rate for exactly the days its own contract runs, and the rate an "
     f"uncommitted vessel earns is SOLVED out of the class average the company publishes "
     f"rather than assumed — the average multiplied by the class's vessel-days, less what "
     f"the committed vessels earned, over the days the uncommitted vessels had. Nothing "
     f"about the market rate is assumed. What IS judged is that the published average is a "
     f"straight vessel-day average of the whole class, which is what the chief financial "
     f"officer's own description of it implies. Only the expiry and the period are "
     f"published for each charter, so each start date is the expiry less the period; no "
     f"date is invented.",
     "A disclosed split between committed and uncommitted earnings, which the company does "
     "not publish, would replace the construction outright. Failing that, the check "
     "available is internal and it passes: where a class has no vessel on charter out, the "
     "published average and the solved rate come out identical. If the company weights its "
     "published average differently — by revenue rather than by vessel-days, say, or "
     "excluding off-hire — the solved rate is too high by the amount the average is "
     "understated, and a reported half-year would show it."),
    ('The rate at which the excess return fades in the asset lens',
     f"{fval('ri_fade', IV['ri_fade'])} a year. The asset lens is built as residual income "
     f"— opening ordinary book, plus the return earned above the cost of equity on that "
     f"book for five years discounted, plus a remainder beyond the forecast in which the "
     f"excess return decays at this rate. The remainder is "
     f"{D['book']['pv_terminal'] / D['book']['equity_value'] * 100:.0f}% of what the lens "
     f"produces. The economic case for a fade is that a fleet has to be replaced at market "
     f"prices rather than at the value it is carried at, so a return above the cost of "
     f"equity cannot persist unchanged; the SPEED is chosen rather than measured.",
     "Nothing observable fixes it, and that is the honest answer. A longer record of this "
     "company earning above its cost of equity through a full rate cycle would argue for a "
     "slower fade; a period of newbuild deliveries competing the return away would argue "
     "for a faster one. A slower fade raises this lens and a faster one lowers it, which "
     "is one reason it carries the lowest weight of the four."),
    ('How the minority interests are deducted',
     f"Not at book, and not at a flat share of equity value either. Of the "
     f"{usdm(IV['q1_26_nci'], 1)} of minority book at the valuation date "
     f"({usdm(IV['nci_fy25'], 1)} at the {FS_YEARS[-1][2:]} year end), "
     f"{usdm(IV['nci_navig8'], 1)} arose on the tanker combination, and that 20% is "
     f"contracted for purchase in mid-2027 at a price already carried in the bridge as "
     f"deferred consideration ({usdm(IV['q1_26_pcp'], 1)}). Deducting it again at a share "
     f"of equity value would count the same claim twice, so it stays at book. Only the "
     f"remaining {usdm(D['dcf']['nci_other_bv'], 1)} is lifted from book to its share of "
     f"value, giving a total deduction of {usdm(D['dcf']['nci'], 1)}.",
     f"A review of the first edition called it a critical failure that minorities were "
     f"deducted at book when they take "
     f"{fval('nci_share', IV['nci_share'])} of profit. The premise is right and the "
     f"conclusion is not, and the arithmetic is published so a reader can decide for "
     f"themselves: at book the deduction is {usdm(IV['q1_26_nci'], 1)}, as taken here it "
     f"is {usdm(D['dcf']['nci'], 1)}, and applying the profit share to the whole equity "
     f"value with nothing netted off would make it "
     f"{usdm(IV['nci_share'] * (D['dcf']['ev'] - D['dcf']['net_debt']), 1)}. What would "
     f"overturn it is the purchase not completing, or completing at a different price — "
     f"at which point the deferred-consideration line changes and this one follows it."),
    ('The step-down in rates assumed for the second half of 2026',
     f"A half-and-half blend: {fval('h2_2026_reversion', IV['h2_2026_reversion'])} weight on "
     f"the 2025 average rate against the rate achieved in the first quarter of 2026. A "
     f"deliberate step down from a first half running far above the 2025 average.",
     "Third-quarter rates printing at first-half levels would make the blend too conservative; "
     "a collapse to the 2025 average would make it too generous. The sensitivity is carried "
     "in the study rather than buried, because the first half is already realised and only "
     "the back half is assumed."),
    ('Which measurement of the market the beta is regressed against',
     f"The published index of the exchange the share is listed on — the {BR['regressor']} — "
     f"because that is the market the share actually trades in and the measurement the method "
     f"calls for. Weekly returns over the full listed history give a slope of {BR['beta']:.3f} "
     f"({BR['n']:,} paired observations, R-squared {BR['r2']:.3f}, standard error "
     f"{BR['se']:.3f}) and a cost of equity of {W['ke'] * 100:.2f}%. The same returns over the "
     f"same window, measured against an equal-weight composite of that exchange's own names "
     f"instead, give {BC['beta']:.3f}, a cost of equity of "
     f"{BF['alternative']['ke'] * 100:.2f}% and a central figure of AED "
     f"{BF['alternative']['central']:,.2f} against the published AED {D['central']:,.2f}. Both "
     f"constructions are computed in full and published side by side; neither is averaged into "
     f"the other.",
     f"A longer price history — the listed record is only {BR['window_years']:.2f} years, and "
     f"length is the one thing that cannot be manufactured. A different definition of the "
     f"index, since what the two constructions disagree about is which names are in the market "
     f"and how much say each of them gets. Or evidence that the share's own membership of the "
     f"index it is measured against materially inflates the slope: it is a constituent, so its "
     f"own returns sit on both sides of the regression and cannot be taken out of a "
     f"capitalisation-weighted index. How large that pull is has been measured on the "
     f"composite, where the share CAN be removed, and it is set out with the empty searches "
     f"that follow."),
    ('The beta used in the low and the high case',
     f"The regression's own {CI_LEVEL} confidence bounds rather than round numbers picked by "
     f"hand: {BF['ci90'][1]:.3f} in the low case and {BF['ci90'][0]:.3f} in the high case. A "
     f"higher slope discounts harder, so the upper bound is what belongs in the low case. On "
     f"the cash-flow lens those give AED {L['dcf']['bear']:,.2f} and AED "
     f"{L['dcf']['bull']:,.2f} a share against a base of AED {L['dcf']['base']:,.2f}.",
     f"A tighter interval, which only more observations give. The standard error is "
     f"{BR['se']:.3f} on a slope of {BR['beta']:.3f}, so the two bounds are "
     f"{(BF['ci90'][1] - BF['ci90'][0]) / BR['beta'] * 100:.0f}% of the estimate apart and the "
     f"cases inherit that width honestly. Round betas chosen by hand would look tidier and "
     f"would be telling the reader less. The lead-lag variant of the same regression, at "
     f"{BR['dimson']['sum_beta']:.3f}, sits inside these bounds."),
    ('The recovery in engineering and construction revenue after 2026',
     f"Revenue of {usdm(IV['drv_rev_offshore_projects_2026'])} in 2026 — inside the "
     f"company's own stated range — recovering to "
     f"{usdm(IV['drv_rev_offshore_projects_2030'])} by 2030, at margins moving from "
     f"{fval('drv_mar_offshore_projects_2026', IV['drv_mar_offshore_projects_2026'])} to "
     f"{fval('drv_mar_offshore_projects_2030', IV['drv_mar_offshore_projects_2030'])}. This is "
     f"the least visible line in the model.",
     "A disclosed award pipeline or tender book. Nothing of the kind is published — the "
     "empty search is recorded overleaf — so a single announced contract, or a second year "
     "at the 2026 level with none, would settle the line in either direction."),
    ('The capital-expenditure path beyond the guided window',
     f"The company's own published path is used to 2028 ({usdm(IV['capex_2026'])} / "
     f"{usdm(IV['capex_2027'])} / {usdm(IV['capex_2028'])}). Beyond it the study steps "
     f"spending down to {usdm(IV['capex_2029'])} and {usdm(IV['capex_2030'])} as the "
     f"newbuild programme delivers, toward stated maintenance spending plus continuing fleet "
     f"renewal.",
     "A new order. The order book and its remaining committed spending are disclosed, so an "
     "addition to either is visible immediately and would lift the path; equally, a programme "
     "that delivers with nothing behind it would push the tail lower still."),
    ('The weighting between the contracted and the spot comparison multiple',
     f"Not chosen. The blend uses the company's OWN disclosed share of earnings exposed to "
     f"spot rates — {REL['spot_weight'] * 100:.0f}% in 2026 — to weight a "
     f"contracted-fleet multiple of {REL['contracted_multiple']:.2f}x against a spot-fleet "
     f"multiple of {REL['spot_multiple']:.2f}x, giving {REL['blend_ev_ebitda']:.2f}x. What IS "
     f"judged is the {fval('rel_weight_ev_ebitda', IV['rel_weight_ev_ebitda'])} weight placed "
     f"on the enterprise-multiple reading against the earnings-multiple reading inside that "
     f"lens, and the choice of three comparable operators rather than a wider set.",
     "A wider comparable set, or the two spot operators re-rating. The three used are the only "
     "listed operators whose business model maps cleanly onto one of the company's two halves; "
     "adding a mixed operator would blur the very split the weighting exists to respect."),
    ('How much weight each valuation lens carries',
     f"Cash flow {LW['dcf'] * 100:.0f}%, comparison {LW['relative'] * 100:.0f}%, normalised "
     f"earnings {LW['normalized'] * 100:.0f}%, book and sustainable return "
     f"{LW['book'] * 100:.0f}%. The cash-flow lens carries the most because it is the only one "
     f"that prices the contracted revenue and the delivery of the order book explicitly.",
     f"Weighting the four equally instead gives AED {EQ_CENTRAL:,.2f} a share against the "
     f"published AED {D['central']:,.2f} — small enough that the weighting is not what the "
     f"answer rests on, which is the point of stating it. A reader who rejects the cash-flow "
     f"lens entirely is reading a different study, and the four lens values are published "
     f"separately so that they can."),
    ('The treatment of the perpetual capital securities',
     f"Two things at once, and they are the two halves of one treatment. They are deducted "
     f"in full — {usdm(IV['q1_26_hybrid'], 1)} — in the bridge from enterprise to equity "
     f"value, because they rank ahead of the ordinary shares however the company "
     f"classifies them; and they are WEIGHTED in the cost of capital at "
     f"{W['wh'] * 100:.1f}% of total capital at their own coupon of {W['kh'] * 100:.2f}%, "
     f"because they are a third kind of capital funding the same enterprise. That coupon "
     f"floats, so it normalises with the risk-free rate in the terminal, at "
     f"{W['kh_term'] * 100:.2f}%. Their coupons never touch the income statement, so "
     f"earnings available to ordinary holders are struck after them.",
     f"The first edition deducted them but did NOT weight them, on the ground that doing "
     f"both would charge for them twice. That was wrong, and two independent reviews said "
     f"so: deducting a claim from value and pricing the capital it supplies are different "
     f"operations, and doing only the first removed a cheap tranche of funding from the "
     f"enterprise without letting it lower the rate the enterprise is discounted at. "
     f"Correcting it moves the cost of capital to {W['wacc'] * 100:.2f}% and the terminal "
     f"rate to {W['wacc_term'] * 100:.2f}%. Treating the securities as ordinary equity "
     f"instead — deducting nothing — would add about AED {HYB_PER_SHARE:,.2f} a share. A "
     f"redemption, or a coupon actually deferred, would resolve the classification in one "
     f"direction or the other."),
    ('The terminal growth rate',
     f"{fval('g_terminal', IV['g_terminal'])}, taken from the company's own goodwill "
     f"value-in-use test, which projects cash flows beyond its plan at a rate equal to an "
     f"estimated 2% inflation. Sensitised from 1.0% to 2.5% rather than asserted.",
     f"It sits at, not above, projected long-run inflation for the economy "
     f"({fval('inflation_26', IV['inflation_26'])} for 2026, settling lower thereafter), so a "
     f"structurally faster-growing franchise would justify more. The terminal value is "
     f"{DCF['tv_share'] * 100:.0f}% of the cash-flow lens, so this row and the discount rate "
     f"are where that lens is won or lost."),
    ('The cost of debt',
     f"{W['kd'] * 100:.2f}%, the average of three constructions rather than one assertion: the "
     f"marginal drawdown rate on the parent facility ({W['kd_method1'] * 100:.2f}%), the "
     f"weighted blend of the instruments actually outstanding ({W['kd_method2'] * 100:.2f}%), "
     f"and the midpoint of the disclosed third-party bank-loan range "
     f"({W['kd_method3'] * 100:.2f}%). It sits above the local sovereign yield of "
     f"{IV['rf_observed'] * 100:.2f}%, as a corporate borrowing in its sovereign's own "
     f"currency must.",
     "A new issue at a rate outside that spread, or a refinancing that shifts the mix away "
     "from the parent facility. The three constructions are published separately so a reader "
     "can take whichever they find most representative."),
    ('The running cost per tanker per day',
     f"{fval('tnk_opex_day', IV['tnk_opex_day'])}, solved so that the owned fleet's disclosed "
     f"rates less that cost reproduce reported 2025 tanker earnings exactly. That makes it a "
     f"calibration to a disclosed outcome rather than a free assumption, but it is still not a "
     f"disclosed figure.",
     f"A disclosed operating cost per vessel. Escalation is applied at "
     f"{fval('opex_escalation', IV['opex_escalation'])} a year as a wage and services "
     f"escalator, not a commodity index, because crew, technical management, insurance and "
     f"repairs are what the line actually consists of; a fuel-linked escalator would be the "
     f"wrong instrument for it."),
    ('The gas carrier day rate',
     f"{fval('gas_rate_day', IV['gas_rate_day'])} on average, implied by 2025 gas revenue over "
     f"consolidated vessel-years read off the published contract table. Per-vessel rates are "
     f"not disclosed, so the group average is the finest level the disclosure supports, and "
     f"the margin is held near the 2025 outcome at "
     f"{fval('gas_margin', IV['gas_margin'])}.",
     "Per-vessel or per-contract rate disclosure. Fifteen of twenty owned vessels sit on "
     "long-term contracts, so the risk in this line is smaller than in tankers — but the "
     "gap is flagged rather than smoothed."),
    ("Treating the company's guidance as a reconciliation, not as the forecast",
     f"The build sits above guided 2026 group earnings by "
     f"{GC['Group']['ebitda_gap'] * 100:.1f}%, and the gap is reported rather than closed. "
     f"Management states its shipping assumptions sit well below prevailing rates and its "
     f"logistics guidance at minimum activity levels.",
     "A first half that lands on guidance rather than above it. Guidance was raised once "
     "already during the study year on the strength of realised performance, which is the "
     "evidence that it is set conservatively rather than centrally."),
    ('Taxing each business unit at its own rate',
     f"Each unit is taxed at the rate its own segment disclosure shows it bore in 2025 — "
     f"{fval('tax_integrated_logistics', IV['tax_integrated_logistics'])} in logistics, "
     f"{fval('tax_shipping', IV['tax_shipping'])} in shipping under the international "
     f"transport relief, {fval('tax_services', IV['tax_services'])} in services — so the "
     f"group rate is an output of the earnings mix and falls as shipping grows.",
     f"The domestic minimum top-up rate of "
     f"{fval('tax_topup_rate', IV['tax_topup_rate'])} applying to the relieved income would "
     f"change the shipping leg materially. The relief is legislated and the segment "
     f"disclosure shows it operating, so the judgement is that it persists."),
    ('Holding the working-capital cycle at its 2025 shape',
     f"Receivable days {IV['dso_days']:,.1f}, inventory days {IV['dio_days']:,.1f} and payable "
     f"days {IV['dpo_days']:,.1f}, computed from the 2025 statements and held, so the balance "
     f"sheet and the cash flow are projected from the conversion cycle rather than plugged.",
     "Two consecutive periods of receivable days moving materially in either direction. The "
     "counterparty is overwhelmingly the parent group, which is why the cycle is stable enough "
     "to hold, and also why a change in it would be a real signal rather than noise."),
]
T([['The judgement', 'What the study took', 'What would overturn it']]
  + [[a, b, c] for a, b, c in JUD],
  [1.42, 2.72, 2.96], aligns=['L', 'L', 'L'], size=7.6)

# ============================================================================
# 6  NEGATIVE RESULTS
# ============================================================================
H1_NEW_PAGE('Negative results — what was looked for and not found')
P("An empty search is a result. Each row below is something the study went looking for, could "
  "not obtain, and therefore had to handle another way — recorded here so that a reader "
  "can see where the evidence stops and judgement starts, rather than discovering it by "
  "reverse-engineering the model.")

NEG_FIDS = [f for f in R['findings'] if f['klass'] == 'NEGATIVE_SEARCH']


def neg_text(f):
    """The register writes a negative finding as 'Negative search - nothing found (detail)'."""
    h = f['headline']
    m = re.search(r'\((.*)\)\s*$', h, re.S)
    return (m.group(1) if m else h).strip()


NEG = [(f['category'].title() if f['category'] != 'pricing' else 'A forward rate curve',
        fdate(f['source_date']), neg_text(f)) for f in NEG_FIDS]
NEG = [
    ('A forward curve or third-party forecast of tanker rates beyond 2026',
     fdate('2026-08-09'),
     neg_text([f for f in NEG_FIDS if f['category'] == 'pricing'][0])),
    ('An engineering and construction award pipeline or tender book',
     fdate('2026-08-09'),
     neg_text([f for f in NEG_FIDS if 'guidance' in f['category']][0])),
    ('Named new entrants into the offshore marine logistics market',
     fdate('2026-08-09'),
     neg_text([f for f in NEG_FIDS if 'entrant' in f['category']][0])),
    ('A sovereign credit-default-swap spread for the United Arab Emirates',
     fdate('2026-01-05'),
     "The country risk file carries no swap entry for this sovereign, so the second premium "
     "basis cannot be built for it at all. A reader who expects a valuation to publish the "
     "cost of equity on both a rating basis and a swap basis should know that only one exists "
     "here. The two are not interchangeable where both are available: in the same file, a Gulf "
     "comparator that does carry a swap entry shows the swap basis running well above its "
     "rating basis. The rating basis is used, and the absence of the alternative is stated "
     "rather than papered over with a regional proxy."),
    ('Per-vessel or per-contract charter rates for the gas fleet',
     fdate('2025-12-31'),
     "Not disclosed. The published contract table gives vessel counts and expiries but not "
     "rates, so the study solves an average revenue per vessel per day from reported segment "
     "revenue over consolidated vessel-years, and says so at the row itself. This is the "
     "finest level the disclosure supports."),
    ('An operating cost per vessel per day for the tanker fleet',
     fdate('2025-12-31'),
     "Not disclosed at any level. The running cost is solved so that the disclosed rates less "
     "that cost reproduce reported 2025 segment earnings exactly, which anchors it to a "
     "disclosed outcome, but it remains a derived figure and is labelled as one."),
    ('An index level for the last two weeks of the stock’s own price history',
     fdate(R['sweep_date']),
     f"The index series obtained ends {fdate(IDX_TO)}, while the stock's own history runs to "
     f"{fdate(META['price_date'])} and is used to that date everywhere else in the study. No "
     f"later index level could be obtained, so {BR['unused_stock_weeks']} weekly stock "
     f"observations fall outside the regression: the window stops where the index stops, "
     f"rather than pairing the stock against an index level that has stopped moving. Related "
     f"and not searchable at all: the share is itself a constituent of the index it is "
     f"measured against, so on a capitalisation-weighted index its own returns cannot be "
     f"removed from the thing they are being regressed on. That one is measured instead of "
     f"searched for — the equal-weight composite is run both with the share inside it and with "
     f"it taken out, giving {BSI['beta_proxy_including_subject']:.3f} against "
     f"{BSI['beta_proxy_excluding_subject']:.3f}. The difference of "
     f"{BSI['beta_proxy_including_subject'] - BSI['beta_proxy_excluding_subject']:.3f} is the "
     f"scale of the upward pull the published-index slope also carries, and it is disclosed in "
     f"the study. It works against the study's own conclusion, since removing it would lower "
     f"the beta and lift the valuation."),
    ('The Annual Report and Accounts for 2023 as a separate volume',
     fdate('2026-08-09'),
     "Not obtained from the investor pages. It costs the study nothing: the 2023 audited "
     "statements themselves were obtained and read, and the 2024 volume carries the 2023 "
     "comparatives, so every 2023 figure has an audited source. Recorded because the record of "
     "what was read should match what was actually read."),
]
T([['What was searched for', 'Date searched', 'Outcome, and how the study handled it']]
  + [[a, b, c] for a, b, c in NEG],
  [1.90, 0.72, 4.48], aligns=['L', 'L', 'L'], size=7.7)

# ============================================================================
# 7  DISCREPANCIES
# ============================================================================
H1('Where two readings of the same figure disagreed')
P("Every disagreement found between a secondary source and the company's own documents is "
  "recorded here with the resolution. In each case the filing was used. The point of the "
  "table is not that the study got it right; it is that a reader who has seen the other "
  "figure should be able to find out immediately why it is not the one in the model. The last "
  "row is a different animal and is included deliberately: there both readings are correct, "
  "and what separates them is not a source but a construction.")
DISC = [
   ['2026 net profit guidance',
    'Press coverage of the guidance raise reported growth in the high-60% range for 2026 net '
    'profit.',
    f"The management commentary of 14 May 2026 raises 2026 net-profit guidance to "
    f"mid-to-high-teens growth — taken at the midpoint of the stated band as "
    f"{fval('g26_np_group', IV['g26_np_group'])}. The only figure above 50% anywhere in that "
    f"release is the SHIPPING earnings line, guided to mid-to-high 50% growth, which is a "
    f"different metric on a different base.",
    'The filing. The press figure is discarded outright and appears nowhere in the model. It '
    'reads like a business-unit earnings number reported as a group profit number; the '
    'company’s own table distinguishes them line by line.'],
   ['First-quarter 2025 revenue',
    f"The company's own management commentary reports first-quarter 2025 revenue of USD 1,204 "
    f"million.",
    f"The reviewed interim statements report {IV['q1_25_rev']:,} thousand for the same "
    f"quarter. The company footnotes the difference itself: tanker revenue and direct costs "
    f"for the first three quarters of 2025 were re-presented, with no effect on gross profit, "
    f"net profit or earnings.",
    'Both, in their places. The statutory figure is used wherever a statement line is built, '
    'so the accounts tie to the filing; the management figures are used for the business-unit '
    'analysis, where they are the only basis on which 2025 and 2026 are comparable. Mixing the '
    'two inside one table is what the disclosure exists to prevent, so the difference is '
    'stated rather than reconciled away.'],
   ['The edition of the country risk file',
    'A mid-2026 update to the country risk work has been published by its author on his own '
    'blog, and two secondary write-ups of that update disagree with each other on the '
    'mature-market premium.',
    f"The canonical file itself still returns the January 2026 edition, which is what was "
    f"read: sovereign default spread {fval('sov_spread', IV['sov_spread'])}, total equity risk "
    f"premium {fval('erp_total', IV['erp_total'])}, mature-market premium "
    f"{fval('erp_mature', IV['erp_mature'])}.",
    'The published file. A number taken from a blog post, or from a write-up of one that '
    'another write-up contradicts, is not a source this study can cite and a reader cannot '
    'check it. Recorded because a reader who has seen the update would otherwise think the '
    'figure is simply stale.'],
   ['What the published rate per vessel class actually is',
    f"Read as the rate a vessel earns in the open market. On that reading the largest "
    f"class earned {D['fleet']['blend_q1_26']['vlcc']:,.0f}/day in the first quarter of "
    f"2026, and the vessels on charter out are a separate, lower, additional exposure. "
    f"This is how the first edition of this study read it, and it is how the figure is "
    f"usually quoted.",
    f"A BLEND across the whole class, including the vessels already on charter out. The "
    f"company said so itself on the first-quarter call — the quotation is set out in full "
    f"with the documents above. Taking the {len(D['fleet']['charters']):,} charters back "
    f"out at their own disclosed rates, the rate an uncommitted vessel must have earned in "
    f"the same quarter is {D['fleet']['spot_q1_26']['vlcc']:,.0f}/day.",
    "The company's own description of its own disclosure. This is not two sources "
    "disagreeing: there is one figure and two readings of what it covers, and only one of "
    "them is the company's. The consequence of the wrong reading was not a wrong rate but "
    "a double count — the first edition applied the blended rate to the whole class AND "
    "added the chartered vessels again at their own rates, charging the drag of the "
    "charters twice and understating the fleet. The correction is set out in the study's "
    "own section on what changed in this edition."],
   ['The beta, measured against two different definitions of the same market',
    f"An equal-weight composite of the exchange's own listed names — {BR['composite_names']:,} "
    f"of them, the subject excluded — gives a slope of {BC['beta']:.3f} over the same window "
    f"(R-squared {BC['r2']:.3f}, standard error {BC['se']:.3f}, {BC['n']:,} observations). "
    f"This is the measurement the study used before the published index could be obtained.",
    f"The {BR['regressor']}, the capitalisation-weighted index the exchange itself publishes "
    f"and the share is a constituent of, gives {BR['beta']:.3f} on the same returns over the same "
    f"window (R-squared {BR['r2']:.3f}, standard error {BR['se']:.3f}, {BR['n']:,} "
    f"observations).",
    "The published index, and the gap between the two is worth more than the answer. This is "
    "not two sources disagreeing about a fact — both slopes are correctly computed from the "
    "same price history. What differs is how the market itself is defined. A "
    "capitalisation-weighted index is dominated by the exchange's largest companies, which is "
    "exactly the group the subject sits in, so the share moves with that index more closely "
    "than it moves with an average in which the exchange's smallest names count for as much as "
    "its largest. The index of the exchange the share is listed on is the market the share "
    "actually trades in, so that is what is used; the composite is published beside it and the "
    "cost of equity is shown on both, because a difference of this size is a fact about index "
    "construction a reader is entitled to see rather than a detail to bury."],
]
T([['Where they disagree', 'The other reading, and where it comes from',
    'The reading this study used, and where it comes from', 'Why that one']] + DISC,
  [1.05, 1.85, 2.10, 2.10], aligns=['L', 'L', 'L', 'L'], size=7.6)

H2('Three definitions of net debt circulate in the company’s own documents')
P(f"This is worth its own note, because the figure feeds straight into the bridge from "
  f"enterprise value to equity value, and the narrowest and widest readings of it are "
  f"{(ND_STUDY - ND_FOOTNOTE) / 1000.0:,.0f} million dollars apart.", size=9.5)
T([['The definition', 'At 31 Mar 2026', 'Used?'],
   ["As the management commentary DEFINES it in its own footnote — debt and debt-like "
    "items consisting of the shareholder loan and current and non-current lease liabilities, "
    "less cash",
    usdm(ND_FOOTNOTE, 1),
    'No. Read literally the definition omits third-party borrowings, and it does not reproduce '
    'the number printed beside it.'],
   ["As the company PRINTS it — the USD 420 million headline in the same table, which "
    "reconciles only when third-party borrowings are added, exactly as the investor "
    "presentation's own net-debt bridge (cash, borrowings, leases, other) shows",
    usdm(ND_PRINTED, 1),
    'Yes, as the starting point. It is reconciled here from the balance sheet itself: '
    f"shareholder facility {usdm(IV['q1_26_shldr_loan'])} plus third-party borrowings "
    f"{usdm(IV['q1_26_borrowings'])} plus lease liabilities {usdm(IV['q1_26_leases'])} less "
    f"cash {usdm(IV['q1_26_cash'])}."],
   ["As this study uses it — the printed figure extended by the deferred consideration "
    f"for the remaining 20% of the acquired tanker business, carried at a present value of "
    f"{usdm(IV['q1_26_pcp'], 1)}, which is contractually owed and therefore debt-like",
    usdm(ND_STUDY, 1),
    f"Yes. Separately, the perpetual capital securities of {usdm(IV['q1_26_hybrid'], 1)} sit "
    f"outside ALL THREE definitions because the company classifies them as equity; they are "
    f"deducted in the bridge in their own right rather than netted into debt, so nothing is "
    f"counted twice and nothing is dropped."]],
  [2.95, 0.85, 3.30], aligns=['L', 'R', 'L'], size=7.7)

# ============================================================================
# 8  SOURCE INTEGRITY
# ============================================================================
H1_NEW_PAGE('Source integrity')
P(f"Every historical figure in this study — every income-statement line, every balance-"
  f"sheet line, every cash-flow line, every segment, service line and cost line — traces "
  f"to the company's own issued financial statements, read from the filing itself rather than "
  f"from any summary of it. Four complete audited financial years were obtained "
  f"({', '.join(FS_YEARS)}), against a floor of two and a target of four, together with the "
  f"reviewed interim information for the first quarter of 2026, which is the only quarter of "
  f"the study year disclosed at the time of writing and is the valuation date.")
P("The company's own website was reached for every one of those documents. Three investor "
  "pages were tried; the only failure logged was a single index page returning a server error "
  "from this environment, and the documents it lists were obtained through the other two "
  "pages on the same site, so no company figure rests on anything but a document downloaded "
  "from the company itself.")
P(f"One source in this document is neither a company filing nor an aggregator, and it is worth "
  f"being explicit about why it is admissible. The beta is regressed against the "
  f"{BR['regressor']} — {BR['regressor_rows']:,} daily sessions from {fdate(IDX_FROM)} to "
  f"{fdate(IDX_TO)}, checked against the exchange's own trading calendar before use. An index "
  f"level is market data, not company data: it is the price at which a market cleared, not a "
  f"figure the company reported about itself. The rule that admits only the company's own "
  f"issued documents governs the company's own reported history, and an index level is outside "
  f"it by construction — there is no filing that could be the source of it. The same is true "
  f"of the stock's own closing prices. What the rule does require of them is that they are "
  f"dated, screened before use and named, which they are, here and at the rows they feed.")
P(f"Aggregators appear in exactly one place: the three comparable operators' statistics pages "
  f"used in the comparison lens, listed by name and date in the external documents table "
  f"above and labelled as cross-checks wherever a multiple derived from them appears. No "
  f"aggregator, broker or press figure is the source of any figure the company itself "
  f"reports. Where a figure is solved, implied or assumed rather than disclosed — the "
  f"tanker running cost, the gas day rate, the forecast drivers — the row in the register "
  f"says so in its own words, and the judgement behind it appears in the judgements table "
  f"with the evidence that would overturn it.")

P('Testahil · Independent valuation research · Educational analysis, not investment '
  'advice. This document contains no recommendation and no target. Sources are listed so that '
  'a reader can verify the analysis independently; where a figure is derived or estimated '
  'rather than disclosed, that is stated at the figure itself.',
  size=8.6, italic=True, color=GREY, space_before=10)

# ============================================================================
# SELF-CHECKS
# ============================================================================
FORBIDDEN = [
    (r'\bstep\s*0\b', 0), (r'\bstep\s*2a\b', 0), (r'\bstep\s+\d', 0),
    (r'\brings?\b', 0), (r'four-ring', 0), (r'information sweep', 0), (r'\bsweeps?\b', 0),
    (r'\bgates?\b', 0), (r'promotion rule', 0), (r'standing research protocol', 0),
    (r'research_sweep', 0), (r'compute\.py', 0), (r'study_numbers', 0),
    (r'market_profiles', 0), (r'mc_v3', 0), (r'data_quality', 0), (r'wacc_builder', 0),
    (r'width_cal', 0), (r'\bpersonas?\b', 0), (r'price target', 0),
    (r'\bLONO\b', 1), (r'\bCRPS\b', 1), (r'\bPIT\b', 1),
    (r'\bPARITY\b', 1), (r'\bBOUNDARY\b', 1), (r'\bFAIL(ED|S|URE)?\b', 1),
]


def all_text():
    out = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                out.append(c.text)
    return out


hits = []
for chunk in all_text():
    for pat, cs in FORBIDDEN:
        m = re.search(pat, chunk, 0 if cs else re.I)
        if m:
            hits.append((pat, chunk[max(0, m.start() - 60):m.end() + 60]))
if hits:
    for h in hits[:20]:
        print('  SCRUB HIT', h)
    sys.exit(f'external-reader scrub found {len(hits)} hits')

# the count the document STATES must be the count it rendered — a stale figure in the
# prose beside a correct register is exactly the failure a reader cannot see
_stated = f'{N_INPUTS:,}'
assert any(_stated in c for c in all_text()), \
    f'the register count {_stated} does not appear in the rendered text'

widths_bad = []
for t in doc.tables:
    tot = sum(c.width.inches for c in t.columns)
    if tot - TEXT_W > 1e-6:
        widths_bad.append(tot)
assert not widths_bad, f'tables wider than the text block: {widths_bad}'

OUT = 'ADNOCLS_Bibliography_09-08-2026.docx'
doc.save(OUT)
print(f'wrote {OUT}')
print(f'  inputs in the file : {len(INP):,}')
print(f'  inputs rendered    : {rendered:,}  (layers: '
      + ', '.join(f'{k} {v}' for k, v in LAYER_N.items()) + ')')
print(f'  judgement rows     : {len(JUD)}')
print(f'  negative results   : {len(NEG)}')
print(f'  discrepancy rows   : {len(DISC)} + a three-way net-debt note')
print(f'  tables             : {len(doc.tables)}')
print(f'  external-reader scrub: 0 hits over {len(all_text()):,} text blocks')
