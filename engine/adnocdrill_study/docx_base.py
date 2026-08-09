"""Shared document furniture for the ADNOC Drilling deliverables — house style.

Palette: canvas 1C3A36 · panel EAF0EE/EFF3F1 · cream F6F1E6 · gold C0A45F ·
brass 896F36 · grey 6E7B77.

TABLE DISCIPLINE. Every table is fixed-layout with explicit column widths, so a
renderer cannot re-flow it into starved or bloated columns. The widths are
checked programmatically in qc_checks.py against the text they actually carry.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1C, 0x3A, 0x36)
GREY = RGBColor(0x6E, 0x7B, 0x77)
BRASS = RGBColor(0x89, 0x6F, 0x36)
GOLD = RGBColor(0xC0, 0xA4, 0x5F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_PANEL2, F_CREAM = '1C3A36', 'EAF0EE', 'EFF3F1', 'F6F1E6'


class Doc:
    def __init__(self, landscape=False):
        self.doc = Document()
        sec = self.doc.sections[0]
        if landscape:
            sec.orientation = WD_SECTION.NEW_PAGE
            sec.page_width, sec.page_height = Inches(11), Inches(8.5)
        else:
            sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.75)
        sec.top_margin = sec.bottom_margin = Inches(0.65)
        st = self.doc.styles['Normal']
        st.font.name = 'Calibri'
        st.font.size = Pt(10.5)
        st.font.color.rgb = INK
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.line_spacing = 1.06
        # The usable text block: page width less both margins. Every table is
        # scaled to fit inside it, so a width list that sums to more than the
        # page can never reach the reader as an overflowing table.
        self.text_width = Emu(sec.page_width - sec.left_margin - sec.right_margin).inches
        self.tables = []

    # ---------------------------------------------------------------- utils --
    @staticmethod
    def shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hexcolor)
        tcPr.append(shd)

    @staticmethod
    def cell_margins(table, top=40, bottom=40, left=90, right=90):
        m = OxmlElement('w:tblCellMar')
        for tag, v in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
            e = OxmlElement(f'w:{tag}')
            e.set(qn('w:w'), str(v))
            e.set(qn('w:type'), 'dxa')
            m.append(e)
        table._tbl.tblPr.append(m)

    @staticmethod
    def borders(table, color='C9D4D1', sz='4'):
        b = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), sz)
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), color)
            b.append(e)
        table._tbl.tblPr.append(b)

    # ------------------------------------------------------------ paragraphs --
    def P(self, text='', size=10.5, bold=False, italic=False, color=INK, align=None,
          space_after=6, space_before=0):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        return p

    def rich(self, parts, size=10.5, space_after=6, space_before=0, align=None):
        p = self.doc.add_paragraph()
        for txt, kw in parts:
            r = p.add_run(txt)
            r.font.size = Pt(kw.get('size', size))
            r.bold = kw.get('bold', False)
            r.italic = kw.get('italic', False)
            r.font.color.rgb = kw.get('color', INK)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def H1(self, text):
        return self.P(text, size=15, bold=True, space_before=14, space_after=6)

    def H2(self, text):
        return self.P(text, size=12, bold=True, space_before=10, space_after=4)

    def H3(self, text):
        return self.P(text, size=11, bold=True, color=BRASS, space_before=8, space_after=3)

    def caption(self, text):
        return self.P(text, size=8.7, italic=True, color=GREY, space_after=10)

    def bullet(self, text, bold_head=None):
        p = self.doc.add_paragraph(style='List Bullet')
        if bold_head:
            r = p.add_run(bold_head)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = INK
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(3)
        return p

    def page_break(self):
        self.doc.add_page_break()

    # ---------------------------------------------------------------- tables --
    def table(self, rows, widths, header=True, first_col_bold=False, size=9.3,
              header_fill=F_PANEL, align_right_from=1, band_rows=None):
        total = sum(widths)
        if total > self.text_width:
            widths = [w * self.text_width / total for w in widths]
        t = self.doc.add_table(rows=len(rows), cols=len(widths))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.cell_margins(t)
        self.borders(t)
        t.autofit = False
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        t._tbl.tblPr.append(layout)
        for j, w in enumerate(widths):
            t.columns[j].width = Inches(w)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                c = t.cell(i, j)
                c.width = Inches(widths[j])
                p = c.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run('' if val is None else str(val))
                r.font.size = Pt(size)
                r.font.color.rgb = INK
                if i == 0 and header:
                    r.bold = True
                    self.shade(c, header_fill)
                if band_rows and i in band_rows:
                    r.bold = True
                    self.shade(c, F_CREAM)
                if first_col_bold and j == 0 and i > 0:
                    r.bold = True
                if j >= align_right_from and i > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        self.tables.append((rows, widths))
        return t

    def figure(self, path, width, caption_text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
        p.paragraph_format.space_after = Pt(2)
        self.caption(caption_text)

    def box(self, lines, fill=F_CREAM, width=7.0):
        t = self.doc.add_table(rows=1, cols=1)
        self.borders(t, color='C0A45F', sz='6')
        self.cell_margins(t, 120, 120, 160, 160)
        c = t.cell(0, 0)
        self.shade(c, fill)
        c.width = Inches(width)
        first = True
        for head, body in lines:
            p = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            if head:
                r = p.add_run(head)
                r.bold = True
                r.font.size = Pt(9.8)
                r.font.color.rgb = INK
            r2 = p.add_run(body)
            r2.font.size = Pt(9.8)
            r2.font.color.rgb = INK
            p.paragraph_format.space_after = Pt(5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def masthead(self, width=7.0):
        t = self.doc.add_table(rows=1, cols=1)
        self.cell_margins(t, 90, 90, 160, 160)
        c = t.cell(0, 0)
        self.shade(c, F_DARK)
        c.width = Inches(width)
        p = c.paragraphs[0]
        r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = WHITE
        r2 = p.add_run('   Not investment advice')
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def save(self, path):
        self.doc.save(path)
