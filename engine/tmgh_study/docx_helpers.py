"""Shared document mechanics for the TMGH deliverables."""
import os
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1D, 0x21)
MUTED = RGBColor(0x5B, 0x65, 0x70)
ACCENT = RGBColor(0x2E, 0x5E, 0x4E)

# vocabulary an external reader must never meet
FORBIDDEN = [
    "sigcm", "assert_", "r-cal", "r-sigcm", "r-enf", "r-beta", "r-std", "r-doc",
    "r-fcal", "r-lesson", "r-tcal", "r-lens", "r-width", "r-grade", "r-shape",
    "step 0.0", "step 2a", "information sweep", "outstanding.json", "panel.json",
    "study_numbers", "beta_regression", "research_protocol", "market_profiles",
    "mc_v3", "wacc_builder", "engine/", "crps", "parity", "gate", "checklist",
    "walk-forward", "ring", "cohort", "raw_ohlc", "data.js", "pre-registration",
    "basis-break", "provenance tier", "lessons register", "watch flag",
]


def money(x, dp=0):
    return "{:,.{dp}f}".format(x, dp=dp)


def pct(x, dp=1):
    return "{:.{dp}f}%".format(x * 100, dp=dp)


def style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Georgia"
    st.font.size = Pt(10)
    st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.18
    for nm, sz, col, bold in (("Heading 1", 16, ACCENT, True),
                              ("Heading 2", 12.5, INK, True),
                              ("Heading 3", 10.5, INK, True)):
        s = doc.styles[nm]
        s.font.name = "Georgia"
        s.font.size = Pt(sz)
        s.font.color.rgb = col
        s.font.bold = bold
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(5)


def para(doc, text, size=10, bold=False, italic=False, color=INK, align=None,
         space_after=7):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(size)
        r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(3)


TABLES = []          # every table built, for the column-width audit below


def table(doc, headers, rows, widths, caption=None, size=8.5):
    """Fixed layout with an explicit GRID — the thing the renderer reads.

    EVERY TABLE IS RECORDED SO ITS COLUMNS CAN BE AUDITED AGAINST THEIR OWN CELLS. A column
    a character too narrow wraps, Word breaks after a hyphen, and a negative number prints
    as a bare dash with its digits on the line beneath — the sign of a printed figure
    changing from a typographic cause. The widths are chosen by hand here, so the check is
    that the hand was right: engine/col_width.py holds the measured per-character widths
    and assert_columns_fit() below runs it over everything this document built.
    """
    TABLES.append((list(headers), [list(r) for r in rows], list(widths)))
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    t._tbl.tblPr.append(lay)
    # Under fixed layout the renderer reads the GRID, not the cell widths.
    # python-docx writes a grid of equal columns and setting cell widths does
    # not touch it, so a table renders with equal columns however carefully the
    # widths were chosen — and a checker that reads the cell widths back reports
    # it clean. The grid is written here and audited from the delivered file.
    grid = t._tbl.find(qn("w:tblGrid"))
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(int(round(w * 567))))
    trPr = t.rows[0]._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    trPr.append(hdr)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Cm(widths[i])
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = MUTED
    for row in rows:
        tr = t.add_row()
        cant = OxmlElement("w:cantSplit")
        tr._tr.get_or_add_trPr().append(cant)
        for i, val in enumerate(row):
            cell = tr.cells[i]
            cell.width = Cm(widths[i])
            cell.text = ""
            r = cell.paragraphs[0].add_run("" if val is None else str(val))
            r.font.size = Pt(size)
            r.font.color.rgb = INK
            plain = str(val).replace("-", "").replace(".", "").replace(",", "") \
                .replace("%", "").replace("(", "").replace(")", "").replace("+", "")
            if i and plain.isdigit():
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if caption:
        para(doc, caption, size=8, italic=True, color=MUTED, space_after=10)
    else:
        # Two tables with nothing between them are ONE table to the renderer:
        # Word and LibreOffice merge adjacent tables and the second loses its
        # header row to the first one's repeating header.
        para(doc, "", size=4, space_after=4)
    return t


def figure(doc, path, caption):
    doc.add_picture(path, width=Cm(16.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, size=8, italic=True, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def scrub(doc_path):
    """Programmatic scan for internal vocabulary. Zero hits required."""
    import re
    from docx import Document as Dx
    d = Dx(doc_path)
    text = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                text += "\n" + c.text
    low = text.lower()
    hits = []
    for w in FORBIDDEN:
        if len(w) <= 6 and w.isalpha():
            if re.search(r"\b%s\b" % re.escape(w), low):
                hits.append(w)
        elif w in low:
            hits.append(w)
    return sorted(set(hits)), len(text)


def column_audit(doc_path):
    """Every table fixed-layout, gridded, none starved, bloated or adjacent."""
    from docx import Document as Dx
    d = Dx(doc_path)
    bad = []
    kids = [c.tag for c in d.element.body.iterchildren()]
    TBL = qn("w:tbl")
    for i in range(len(kids) - 1):
        if kids[i] == TBL and kids[i + 1] == TBL:
            bad.append((i, "two tables with nothing between them"))
    for i, t in enumerate(d.tables):
        widths = [c.width for c in t.rows[0].cells]
        if any(w is None for w in widths):
            bad.append((i, "a column carries no explicit width"))
            continue
        cm = [w.cm for w in widths]
        if min(cm) < 1.0:
            bad.append((i, "starved column %.2fcm" % min(cm)))
        if sum(cm) > 17.5:
            bad.append((i, "row overflows the text column at %.2fcm" % sum(cm)))
        grid = t._tbl.find(qn("w:tblGrid"))
        gw = [int(gc.get(qn("w:w"))) / 567.0 for gc in grid.findall(qn("w:gridCol"))]
        if len(gw) != len(cm) or any(abs(a - b) > 0.05 for a, b in zip(gw, cm)):
            bad.append((i, "the grid the renderer reads disagrees with the "
                           "cell widths: grid %s vs cells %s"
                        % ([round(x, 2) for x in gw], [round(x, 2) for x in cm])))
    return bad


def assert_columns_fit():
    """Every column clears its own widest unbreakable token, or the build stops.

    Run at the END of the build, over every table the document actually produced — not over
    a list somebody maintains, which is the shape that lets one table drift.
    """
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..'))
    import col_width
    bad = []
    for headers, rows, widths in TABLES:
        for col, declared, needed in col_width.audit(headers, rows, widths):
            bad.append('  %-42s declared %.2fcm, needs %.2fcm'
                       % (str(col)[:42], declared, needed))
    assert not bad, ('columns too narrow for their own cells — a cell that wraps '
                     'mid-number changes the figure a reader sees:\n' + '\n'.join(bad))
    return len(TABLES)
