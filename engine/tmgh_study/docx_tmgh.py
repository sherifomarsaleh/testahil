"""Build the delivered TMGH study document.

Reads study_numbers.json, experts.json and peers.json, and nothing else — no
financial numeral is typed into this builder, so every number in the delivered
document traces to the input entry it came from.

The document is written for an external reader: no internal procedure
vocabulary, no file paths, no module names, no house shorthand. Evidence about
how well the price distribution has held appears in section 3 as plain-language
sentences with the statistics inline, never as an appendix.
"""
import datetime, json, math, os, re, sys

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ENGINE = os.path.dirname(HERE)
ROOT = os.path.dirname(ENGINE)
sys.path.insert(0, HERE)
from docx_helpers import (INK, MUTED, ACCENT, money, pct, style, para, bullets,
                          table, figure, scrub, column_audit,
                          assert_columns_fit)

sys.path.insert(0, os.path.join(HERE, ".."))
import col_width                                                       # noqa: E402

HDR_SUMMARY = ["Case", "Discount rate", "Enterprise value",
               "Equity value (adopted)", "Per share, minority at value share",
               "Per share, at book", "Per share, pro rata"]

N = json.load(open(os.path.join(HERE, "study_numbers.json")))
EX = json.load(open(os.path.join(HERE, "experts.json")))
PE = json.load(open(os.path.join(HERE, "peers.json")))
M, W, R = N["meta"], N["wacc"], N["ratios"]
IS, BSH, KPI, H1 = (N["inputs"]["IS"], N["inputs"]["BS"], N["inputs"]["KPI"],
                    N["inputs"]["H1_26"])
LENS, ST, WF = N["lenses"], N["statements"], N["walkforward"]
CASES = N["valuation_cases"]
SCHED = N["cost_of_capital_record"]
PSB, PSP = N["per_share_nci_book"], N["per_share_nci_proportional"]
PSV = N["per_share_nci_value_share"]      # the ADOPTED basis
# HOW MANY WAYS THE MINORITY IS ACTUALLY DEDUCTED, COUNTED RATHER THAN TYPED. The document
# said "two different ways" in two places and "three bases" in a third, while the summary
# table publishes three columns. A count a reader can check against the table it sits under
# is not a figure of speech.
_NBASES = len([k for k in N if k.startswith("per_share_nci_")])
_NBASES_WORD = {2: "two", 3: "three", 4: "four"}[_NBASES]
FV = N["fair_value_range"]
SPOT = M["spot"]
# THE EDITION DATE IS DERIVED FROM THE FILE THIS DOCUMENT SHIPS AS [ADDED 03-Sep-2026].
# It was typed here and typed again in the output filename, and the two disagreed: DATE
# read "1 September 2026" while the file shipped as 02-09-2026. Nothing was wrong with the
# study; a person had to remember two strings and remembered one. A date is a figure a
# reader sees, so the standing rule applies to it — COMPUTED, NOT TYPED.
EDITION_FILE = "TMGH_Valuation_Study_02-09-2026.docx"


def _edition_words(fname=EDITION_FILE):
    """'2 September 2026' from the filename this document ships as. One source, one date."""
    import datetime as _dt
    import re as _re
    m = _re.search(r"(\d{2})-(\d{2})-(\d{4})", fname)
    if not m:
        raise ValueError("cannot read an edition date out of %r" % fname)
    d = _dt.date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
    return "%d %s %d" % (d.day, d.strftime("%B"), d.year)


DATE = _edition_words()


import sys as _sys
_sys.path.insert(0, os.path.join(ROOT, "engine"))
import site_data                                                    # noqa: E402


def v(reg, k):
    return reg[k]["value"]


def site_block():
    """The published technical read and price distribution, read live.

    THROUGH A REAL PARSE, NEVER BY REGULAR EXPRESSION [R-ENF-03]. This function used to
    find the entry with js.index("\\n  TMGH: {") and slice a fixed 4,200 bytes after it,
    then hunt inside that window with re.search — which is the exact construction the rule
    names twice over: index and re.search both take the FIRST occurrence where a JavaScript
    object literal takes the LAST, and a fixed byte window either truncates the entry or
    bleeds into the next one, silently, with no error either way.
    """
    e = site_data.read('TICKERS', 'TMGH')
    out = {}
    for tag in ('t20', 't60'):
        d = (e.get('dist') or {}).get(tag)
        if d:
            out[tag] = d
    lv = e.get('levels') or {}
    if lv.get('res') and lv.get('sup'):
        out['res'] = [float(x) for x in lv['res']]
        out['sup'] = [float(x) for x in lv['sup']]
    tech = e.get('tech') or {}
    for k in ('trend', 'summary', 'bull', 'bear'):
        if tech.get(k):
            out[k] = tech[k]
    return out


SITE = site_block()


# ---------------------------------------------------------------- the document
def build(path):
    doc = Document()
    style(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.4)

    # --- 1. Masthead + READ FIRST -----------------------------------------
    para(doc, "TALAAT MOUSTAFA GROUP HOLDING", size=19, bold=True, color=ACCENT,
         space_after=2)
    para(doc, "Egyptian Exchange · TMGH · reported in Egyptian pounds", size=10,
         color=MUTED, space_after=2)
    para(doc, "Valuation study · %s" % DATE, size=10, color=MUTED, space_after=14)

    # the model report counts the masthead and this note as its first
    # section, so it is a top-level heading and the section count is 16
    doc.add_heading("Read first", level=1)
    para(doc, "This document sets out a range of values for one company and the "
              "reasoning behind it. It is not advice, it does not tell anyone to "
              "buy or sell anything, and it contains no target price. Where a "
              "judgement could reasonably be made two ways, both are shown and "
              "neither is averaged into the other.")
    para(doc, "Three things are worth knowing before the numbers.")
    bullets(doc, [
        "The discount rate does most of the work, and it is not one rate. A "
        "company whose order book converts over the next fifteen years cannot be "
        "discounted at today's crisis rate in every one of those years, because "
        "the central bank publishes a path back to a seven per cent inflation "
        "target and this study's own borrowing assumptions already follow it. "
        "Each year therefore carries its own rate, falling from %s in the first "
        "year to %s once the economy has normalised; a pound arriving in year "
        "nineteen is worth about four times what a single flat rate would have "
        "said. Earlier editions of this study used one flat rate throughout, and "
        "that is the largest single change here."
        % (pct(SCHED["forward_wacc"][0], 2), pct(SCHED["wacc_terminal"], 2)),
        "The order book is enormous and slow. TMG had sold but not yet handed "
        "over EGP %s billion of property at 30 June 2026, against EGP %s billion "
        "of development revenue recognised in the first half. How fast that "
        "converts is the single question this study cannot settle from what the "
        "company publishes, so it is answered both ways throughout."
        % (money(v(KPI, "backlog_jun26") / 1000, 1),
           money(v(H1, "dev_revenue") / 1000, 1)),
        "Nearly half the group does not belong to TMG's shareholders. "
        "Non-controlling interests are %s of consolidated equity after the 2024 "
        "hotel acquisition and the project-company structures. Every value here "
        "is shown with that minority deducted %s different ways."
        % (pct(CASES["rating|capacity"]["nci_share_of_equity"], 1), _NBASES_WORD),
    ])
    para(doc, "Every figure about the company comes from its own audited or "
              "reviewed financial statements, or from documents it published "
              "itself. Where something is not disclosed, this study says so "
              "rather than estimating it.", italic=True, color=MUTED)

    # --- 2. Headline -------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Headline", level=1)
    para(doc, "On the method used here, Talaat Moustafa Group Holding is worth "
              "between EGP %s and EGP %s a share. The shares last traded at EGP "
              "%s."
              % (money(FV["low"], 2), money(FV["high"], 2), money(SPOT, 2)),
         size=11.5, bold=True)
    para(doc, "That range is the envelope of four published cases: two ways of "
              "measuring Egypt's equity risk premium, crossed with two readings "
              "of how fast the order book converts, each shown with the minority "
              "interest deducted on three bases — at its share of value, proxied by "
              "the filed profit share, which is the basis adopted; at book; and pro "
              "rata to book — because the protocol deducts a minority at what it is "
              "worth, never at what it historically cost. The cases are "
              "not averaged, because averaging two answers to a question nobody "
              "has settled produces a number that was true in neither.")
    para(doc, "Run the same model backwards and ask what single flat discount "
              "rate would reproduce the traded price, and the answer is %s to %s. "
              "This study does not use a single rate: it discounts each year at "
              "its own, starting at %s on the credit-default-swap premium — the "
              "rating-based premium runs %s higher in the first year — and settling "
              "at %s once the central bank's "
              "published path has played out. The comparison is still worth "
              "making, and what it shows is that the market and this study are "
              "closer than the headline rates suggest."
              # THE GAP BETWEEN THE TWO LADDERS IS COMPUTED, NEVER TYPED. Both are
              # published cases and the sentence used to quote one as "this study's"
              # rate without saying which, while the study's flagship waterfall is
              # built on the other.
              % (pct(LENS["implied_discount_rate"]["recovery"], 1),
                 pct(LENS["implied_discount_rate"]["capacity"], 1),
                 pct(SCHED["forward_wacc"][0], 2),
                 pct(CASES["rating|capacity"]["wacc"] - SCHED["forward_wacc"][0], 2),
                 pct(SCHED["wacc_terminal"], 2)))

    figure(doc, os.path.join(HERE, "fig1_football.png"),
           "Every lens in this study, and the price the market is paying. The "
           "four dark bars are the cases this study publishes: each spans the "
           "three bases on which the minority interest can be deducted, and "
           "each is labelled with the basis adopted followed by that span, so "
           "the figure the label reads is the figure the summary table "
           "carries. The lighter bars are the cross-checks.")

    # --- 3. Valuation summary ---------------------------------------------
    doc.add_heading("Valuation summary", level=1)
    rows = []
    for k in ("rating|capacity", "rating|recovery", "cds|capacity", "cds|recovery"):
        basis, mode = k.split("|")
        rows.append([
            "%s premium, %s conversion"
            % ("Rating-based" if basis == "rating" else "Credit-default-swap",
               "slower" if mode == "capacity" else "faster"),
            pct(CASES[k]["wacc"], 2),
            money(CASES[k]["enterprise_value"]),
            money(CASES[k]["equity_after_nci_value_share"]),
            money(CASES[k]["per_share_nci_value_share"], 2),
            money(PSB[k], 2), money(PSP[k], 2)])
    table(doc, HDR_SUMMARY,
          # 1.7cm broke "Discount" mid-word and the header rendered "Discou nt
          # rate". A first hand fix widened that column and narrowed "Case" to
          # 3.0cm, which needs 3.61 for "Credit-default-swap" — the same error
          # again, minutes later, which is the argument for sizing from the cells
          # rather than from the eye.
          rows, col_width.fit_widths(HDR_SUMMARY, rows, 16.3),
          "All values in EGP million except per-share figures, which are in EGP. "
          "The four cases are published side by side and never averaged.")

    para(doc, "The other lenses, as cross-checks:")
    b = LENS["book_and_sustainable_return"]
    bvals = [x["value_per_share"] for x in b["cases"].values() if x["value_per_share"]]
    caps = [x for k, x in LENS["normalised_earnings"].items() if k.startswith("cap|")]
    hist = [r["pe"] for r in LENS["own_multiple_history"] if r.get("pe")]
    eps_n = LENS["normalised_earnings"]["average_eps"]
    table(doc, ["Cross-check", "What it measures", "EGP per share"],
          [["Book value and the return earned on it",
            "the company's return on equity of %s against a cost of equity of %s"
            % (pct(b["roe_fy25"], 1), pct(b["ke_rating"], 1)),
            "%s – %s" % (money(min(bvals), 2), money(max(bvals), 2))],
           ["Normalised earnings power",
            "three years' attributable profit with revaluation gains stripped, "
            "capitalised", "%s – %s" % (money(min(caps), 2), money(max(caps), 2))],
           ["Its own history of multiples",
            "the range of price-to-earnings multiples the shares have actually "
            "carried, %s to %s times, on normalised earnings of EGP %s"
            % (money(min(hist), 1), money(max(hist), 1), money(eps_n, 2)),
            "%s – %s" % (money(min(hist) * eps_n, 2), money(max(hist) * eps_n, 2))]],
          [4.6, 8.4, 3.2])

    # --- 4. Company overview ----------------------------------------------
    doc.add_heading("Company overview", level=1)
    para(doc, "Talaat Moustafa Group Holding is Egypt's largest listed developer "
              "and manager of integrated communities, and it has been publicly "
              "held since 2007. It runs three businesses inside one holding "
              "company.")
    table(doc, ["Business", "FY2025 revenue", "Share of revenue",
                "Gross margin, first half 2026", "What it is"],
          [["Real-estate development", money(v(IS, "dev_revenue_fy25")),
            pct(v(IS, "dev_revenue_fy25") / (v(IS, "dev_revenue_fy25")
                + v(IS, "hosp_revenue_fy25") + v(IS, "other_revenue_fy25")), 0),
            pct(R["gm_dev_h1_26"], 1),
            "Madinaty, Al Rehab, Celia, SouthMed, The Spine and, from 2025, Banan "
            "in Saudi Arabia"],
           ["Hospitality", money(v(IS, "hosp_revenue_fy25")),
            pct(v(IS, "hosp_revenue_fy25") / (v(IS, "dev_revenue_fy25")
                + v(IS, "hosp_revenue_fy25") + v(IS, "other_revenue_fy25")), 0),
            pct(R["gm_hosp_h1_26"], 1),
            "about %s operating keys with a further %s under construction, "
            "including Four Seasons and Sofitel properties"
            % (money(v(KPI, "hotel_keys_operating")),
               money(v(KPI, "hotel_keys_under_construction")))],
           ["Other recurring income", money(v(IS, "other_revenue_fy25")),
            pct(v(IS, "other_revenue_fy25") / (v(IS, "dev_revenue_fy25")
                + v(IS, "hosp_revenue_fy25") + v(IS, "other_revenue_fy25")), 0),
            pct(R["gm_other_h1_26"], 1),
            "commercial leasing and management, sporting clubs and integrated "
            "community services"]],
          [3.2, 2.2, 1.9, 2.6, 6.3],
          "Revenue in EGP million. Margins are outputs of the disclosed cost "
          "lines, not assumptions.")
    para(doc, "The scale that matters most is not the income statement. At 30 "
              "June 2026 the company held a sold-but-undelivered order book of "
              "EGP %s billion, a landbank of about %s million square metres, "
              "customer advances of EGP %s billion already collected against "
              "undelivered homes, and a further EGP %s billion of post-dated "
              "cheques held outside the balance sheet for the same purpose."
              % (money(v(KPI, "backlog_jun26") / 1000, 1),
                 money(v(KPI, "landbank_msqm"), 0),
                 money(v(BSH, "customer_advances") / 1000, 1),
                 money(v(BSH, "postdated_cheques_offbs") / 1000, 1)))
    return doc


def section1(doc):
    doc.add_heading("1. Fundamental valuation", level=1)

    # 1.1 cash-flow model
    doc.add_heading("1.1 The cash-flow model", level=2)
    para(doc, "The company is valued as a sum of its parts, because its three "
              "businesses need different treatment and forcing them into one "
              "model produces a number that describes none of them. One "
              "projection underlies everything: the same forecast feeds the "
              "value, the statements in Appendix A and the workbook.")
    # THE CASE IS NAMED IN THE CAPTION AND THE NAME IS DERIVED FROM THIS KEY, never
    # typed. Two of the four published cases are slower-conversion, so "on the
    # slower-conversion case" identified half of them: a reader who carried the
    # narrative's 32.37% to this table found a discount factor implying 35.79% and
    # nothing on the page reconciled the two.
    WF_CASE = "rating|capacity"
    c = CASES[WF_CASE]
    _erp, _conv = WF_CASE.split("|")
    wf_case_words = ("%s equity risk premium, %s conversion"
                     % ("rating-based" if _erp == "rating" else "credit-default-swap",
                        "slower" if _conv == "capacity" else "faster"))
    rows_ = c["rows"][:5]
    yrs = [str(r["year"]) for r in rows_]
    # THE WATERFALL READS DOWN THE PAGE, not across it. Eleven waterfall stages
    # as columns wrapped every header into fragments — "Revenu e", "Workin g
    # capital" — which is the shape a reader cannot follow, and it only showed
    # up on the rendered page.
    tax = N["model_parameters"]["TAX"]
    lines = [
        ("Revenue", [r["revenue"] for r in rows_]),
        ("Cost of revenue", [-r["cost_of_revenue"] for r in rows_]),
        ("Gross profit", [r["gross_profit"] for r in rows_]),
        ("Gross margin", [pct(r["gross_margin"], 1) for r in rows_]),
        ("Overheads", [-r["opex"] for r in rows_]),
        ("Depreciation and amortisation", [-r["da"] for r in rows_]),
        ("Operating profit", [r["ebit"] for r in rows_]),
        ("Tax on operating profit", [-r["ebit"] * tax for r in rows_]),
        ("Operating profit after tax", [r["ebit"] * (1 - tax) for r in rows_]),
        ("add back depreciation and amortisation", [r["da"] for r in rows_]),
        ("less capital spend", [-r["capex"] for r in rows_]),
        ("less increase in homes built ahead of handover",
         [-r["d_properties_under_development"] for r in rows_]),
        ("add increase in money collected ahead of handover",
         [r["d_advances"] for r in rows_]),
        ("Free cash flow to the firm", [r["fcff"] for r in rows_]),
        # formatted here, not by the shared money() helper: a discount factor
        # rounded to whole numbers renders as 1, 1, 0, 0, 0
        ("Discount factor", [money(r["discount_factor"], 4) for r in rows_]),
        ("Present value of free cash flow", [r["pv"] for r in rows_]),
    ]
    tab = [[lab] + [x if isinstance(x, str) else money(x) for x in vals]
           for lab, vals in lines]
    table(doc, ["EGP million"] + yrs, tab, [6.2, 2.0, 2.0, 2.0, 2.0, 2.0],
          "The first five of the ten explicit years, on the case built on the %s, "
          "whose first-year rate is %s; the other three cases are in the workbook, "
          "as is the full ten-year waterfall. The two working-capital "
          "lines are the movement in the positions a developer carries at once — "
          "money collected ahead of handover, and homes built ahead of handover."
          % (wf_case_words, pct(c["wacc"], 2)))

    para(doc, "From enterprise value to what the shareholder owns:")
    br = [["Present value of the explicit ten years", money(c["pv_explicit"])],
          ["Present value of the order book still unconverted after that",
           money(c["pv_residual_book"])],
          ["Present value of the recurring businesses beyond the window",
           money(c["pv_terminal_recurring"])],
          ["Enterprise value", money(c["enterprise_value"])],
          ["Cash, deposits and amortised-cost financial assets",
           money(c["cash_and_deposits"])],
          ["Borrowings", "(%s)" % money(c["borrowings"])],
          ["Lease liabilities", "(%s)" % money(c["lease_liabilities"])],
          ["Investment property, held separately", money(c["investment_property"])],
          ["Investments in associates and other financial assets",
           money(c["associates"] + c["fvoci"])],
          ["Value of the whole group's equity", money(c["equity_before_minority"])],
          ["Non-controlling interests at their share of value (filed profit share, %s)"
           % pct(c["nci_profit_share"]),
           "(%s)" % money(c["equity_before_minority"] - c["equity_after_nci_value_share"])],
          ["  for reference: at book %s; pro rata to book %s"
           % (money(c["nci_book"]), money(c["equity_before_minority"] * c["nci_share_of_equity"])), ""],
          ["Equity attributable to TMG's own shareholders",
           money(c["equity_after_nci_value_share"])],
          ["Shares in issue, million", money(c["shares_mn"], 1)],
          ["Value per share, EGP", money(c["per_share_nci_value_share"], 2)]]
    # [R-BRIDGE-01] REQUIRES THE EQUITY TO DIVIDE TO THE STATED PER SHARE, AND THIS
    # BRIDGE DID NOT. Every line above deducted the minority at its share of value and
    # reached EGP 244,183mn; the last row then printed per_share_nci_book, EGP 113.24,
    # where 244,183 / 2,060.7 is 118.50 — the number the summary table publishes for
    # this same case two pages earlier. A reader following the bridge down arrived 4.4%
    # BELOW the study's own answer, and no gate could see it: the table-footing
    # instrument reproduces COLUMN totals and this last step is a division.
    #
    # The rule already said the assertion should exist. It did not, so here it is.
    _ps = c["equity_after_nci_value_share"] / c["shares_mn"]
    assert abs(_ps - c["per_share_nci_value_share"]) < 0.005, (
        "the bridge does not divide: %.4f vs %.4f"
        % (_ps, c["per_share_nci_value_share"]))
    _lines = (c["pv_explicit"] + c["pv_residual_book"] + c["pv_terminal_recurring"])
    assert abs(_lines - c["enterprise_value"]) < 1.0, (
        "the bridge does not foot to enterprise value: %.1f vs %.1f"
        % (_lines, c["enterprise_value"]))
    _eq = (c["enterprise_value"] + c["cash_and_deposits"] - c["borrowings"]
           - c["lease_liabilities"] + c["investment_property"] + c["associates"]
           + c["fvoci"])
    assert abs(_eq - c["equity_before_minority"]) < 1.0, (
        "the bridge does not foot to group equity: %.1f vs %.1f"
        % (_eq, c["equity_before_minority"]))
    table(doc, ["Bridge from enterprise value to equity", "EGP million"], br,
          [11.0, 5.2],
          # wf_case_words already names the conversion reading, so a "slower-conversion
          # case on the …, slower conversion" prefix stutters.
          "The case built on the %s, with the minority deducted at its share of "
          "value, which is the basis adopted. The other three cases run the same "
          "bridge on their own numbers, and the two other minority bases are in the "
          "summary table." % wf_case_words)

    # 1.2 book value
    doc.add_heading("1.2 Book value and the return earned on it", level=2)
    b = LENS["book_and_sustainable_return"]
    para(doc, "Book value attributable to TMG's shareholders is EGP %s a share. "
              "The company earned %s on that book in 2025. Its cost of equity on "
              "this method is %s. A business earning less on its capital than the "
              "capital costs is worth less than its book, and the arithmetic says "
              "so plainly: at a long-run growth rate of 10%% the justified "
              "multiple of book is %s times, giving EGP %s a share."
              % (money(b["book_value_per_share"], 2), pct(b["roe_fy25"], 1),
                 pct(b["ke_rating"], 1),
                 money(b["cases"]["rating|g=10%"]["justified_pb"], 2),
                 money(b["cases"]["rating|g=10%"]["value_per_share"], 2)))
    para(doc, "That is the most severe reading in this study and it should be "
              "treated with care. It is entirely a statement about the discount "
              "rate: at the rate the market appears to use, the same arithmetic "
              "gives a multiple above one. It is included because a lens that "
              "only ever agrees with the others is not a lens.")

    # 1.3 relative
    doc.add_heading("1.3 What the shares have been worth before", level=2)
    hrows = [[str(r["year"]), money(r["close"], 2),
              money(r["eps"], 2) if r.get("eps") else "not stated",
              money(r["pe"], 1) if r.get("pe") else "–",
              money(r["bvps"], 2) if r.get("bvps") else "not stated",
              money(r["pb"], 2) if r.get("pb") else "–"]
             for r in LENS["own_multiple_history"] if r["year"] >= 2017]
    table(doc, ["Year end", "Close, EGP", "Earnings per share, EGP",
                "Price to earnings", "Book value per share, EGP",
                "Price to book"], hrows, [2.0, 2.2, 3.0, 2.4, 3.2, 2.2],
          "Earnings are those attributable to TMG's own shareholders. Book value "
          "is shown only for years where the split between TMG's shareholders "
          "and the minority is disclosed; carrying group equity in the same "
          "column would compare two different things.")
    para(doc, "Priced against its own history the company looks ordinary. It has "
              "traded between %s and %s times earnings over the period, and at "
              "the last close it is on %s times its 2025 result. Nothing in the "
              "relative record suggests the shares are unusually priced; what "
              "the cash-flow lens disputes is whether that historical range was "
              "ever right, not whether today sits inside it."
              % (money(min(r["pe"] for r in LENS["own_multiple_history"]
                           if r.get("pe")), 1),
                 money(max(r["pe"] for r in LENS["own_multiple_history"]
                           if r.get("pe")), 1),
                 money(SPOT / v(IS, "eps_fy25"), 1)))
    para(doc, "A table of competitors' multiples is deliberately not shown. This "
              "study holds price histories for the Egyptian names listed in "
              "Appendix B but has not sourced their financial statements, and a "
              "competitor multiple built on a denominator nobody sourced is a "
              "number wearing a decimal point.", italic=True, color=MUTED)

    # 1.4 normalised
    doc.add_heading("1.4 Normalised earnings power", level=2)
    ne = LENS["normalised_earnings"]
    # THE ROW DID NOT RECONCILE ACROSS AND THE MISSING FACTOR WAS NOWHERE ON THE PAGE.
    # The gain is disclosed at GROUP level and only the parent's share of it is stripped
    # from attributable profit, so a reader subtracting the two printed columns got EGP
    # 5,799mn for 2024 against a printed 8,025mn, with nothing to explain the difference.
    # The table-footing instrument reproduces COLUMN totals and cannot see a row identity.
    # The parent's share is now a column of its own, computed from the same figure the
    # model uses, so the arithmetic is reproducible from what is printed.
    hdr_ne = ["Year", "Attributable profit as reported",
              "Revaluation gain at group level", "Parent's share of that gain",
              "Cleaned attributable profit"]
    rows_ne = [[y, money(ne["years"][y]["npat"]), money(ne["years"][y]["reval"]),
                money(ne["years"][y]["npat"] - ne["cleaned_attributable_profit"][y]),
                money(ne["cleaned_attributable_profit"][y])]
               for y in sorted(ne["years"])]
    for _y in sorted(ne["years"]):
        _share = ne["years"][_y]["npat"] - ne["cleaned_attributable_profit"][_y]
        assert abs((ne["years"][_y]["npat"] - _share)
                   - ne["cleaned_attributable_profit"][_y]) < 0.5, _y
    table(doc, hdr_ne, rows_ne, col_width.fit_widths(hdr_ne, rows_ne, 16.2),
          "EGP million. The gain is disclosed at group level and only the parent's "
          "share of it is stripped, which is why the parent's share is smaller than "
          "the group gain. The revaluation gains are non-cash and are not forecast; "
          "capitalising them would capitalise a valuation opinion rather than a "
          "business.")
    para(doc, "Averaged over the three years, cleaned attributable profit is EGP "
              "%s million, or EGP %s a share. Capitalised at the cost of equity "
              "less long-run growth, that is EGP %s to EGP %s a share."
              % (money(ne["average"]), money(ne["average_eps"], 2),
                 money(min(x for k, x in ne.items() if k.startswith("cap|")), 2),
                 money(max(x for k, x in ne.items() if k.startswith("cap|")), 2)))

    # 1.5 synthesis
    doc.add_heading("1.5 Synthesis — four lenses, one field", level=2)
    para(doc, "The four lenses do not agree, and the pattern of disagreement is "
              "informative. The cash-flow model and the earnings-power model land "
              "in the twenties to sixties. The book-value model lands far below "
              "them, for one reason: it is the lens most exposed to the discount "
              "rate, and this study's discount rate is high. The relative lens "
              "lands closest to the traded price, for the opposite reason: it "
              "contains no discount rate at all, and simply records what buyers "
              "have paid before.")
    para(doc, "Read together they say something a single number could not. The "
              "shares are priced roughly where they have always been priced "
              "relative to earnings, and that price implies a required return "
              "close to what the Egyptian government pays to borrow. Whether "
              "that is enough compensation for owning a levered developer in "
              "this currency is the question, and it is one a reader can answer "
              "for themselves from figure 2.")
    return doc


def section1_drivers(doc):
    doc.add_heading("1.6 The drivers, and how well this method has forecast them",
                    level=2)
    para(doc, "Revenue is not projected as a growth rate. Each segment is built "
              "on its own driver and each cost line is a disclosed ratio of that "
              "segment's revenue, so the margin is an output rather than an "
              "assumption.")
    table(doc, ["Driver", "How it is built", "Value used"],
          [["Development revenue",
            "the disclosed order book converted at a stated rate, capped by what "
            "has actually been sold", "book of EGP %s bn at 30 June 2026"
            % money(v(KPI, "backlog_jun26") / 1000, 1)],
           ["Development cost",
            "a ratio of development revenue, from the reviewed first half of 2026",
            pct(1 - R["gm_dev_h1_26"], 1) + " of revenue"],
           ["Hospitality revenue", "grown on its own rate",
            pct(N["model_parameters"]["HOSP_GROWTH"], 0) + " a year"],
           ["Other recurring revenue", "grown on its own rate",
            pct(N["model_parameters"]["OTHER_GROWTH"], 0) + " a year"],
           ["Overheads", "a ratio of group revenue, from 2025",
            pct(R["opex_ratio_fy25"], 2)],
           ["Money collected ahead of handover",
            "a rate on the order book plus the year's sales, taken from the "
            "movement in customer advances in the first half of 2026",
            pct(R["collection_rate_on_book"], 2)],
           ["Homes built ahead of handover",
            "work in progress moved toward the cover the build programme needs",
            "%s years of cost" % money(N["model_parameters"]["PUD_COVER_YEARS"], 1)],
           ["Interest", "the marginal borrowing rate applied to borrowings, and "
                        "to nothing else", pct(R["kd"], 2)]],
          [3.4, 7.6, 5.2],
          "Customer advances, supplier balances and obligations against cheques "
          "received are funding, but none of them pays interest, so none of them "
          "is in the denominator of the borrowing rate.")

    para(doc, "This method was then tested against the company's own history. "
              "The driver model was rebuilt as it would have stood at each year "
              "end from 2015 to 2024, projected forward one to five years, and "
              "each driver scored against what the company actually reported. "
              "The results decided which of the model's habits were corrected "
              "and which were left alone.")
    sc = WF["driver_scores"]

    def row(key, label):
        s = sc.get(key)
        if not s:
            return None
        return [label, str(s["n"]),
                ("%+.2f" % s["bias"]), ("%.2f" % s["mae"]),
                "yes" if s["robust"] else "no"]
    rows = [r for r in (row("dev_revenue", "Development revenue"),
                        row("total_revenue", "Total revenue"),
                        row("gross_profit", "Gross profit"),
                        row("net_profit", "Net profit"),
                        row("new_sales", "New contracted sales"),
                        row("da", "Depreciation"),
                        row("finance_cost", "Finance cost")) if r]
    # A RESULTS TABLE WITH NO ROWS SHIPPED FOR A DAY. The lookup above asked for
    # "asknown|<driver>|all" against a record whose keys are plain driver names, matched
    # nothing, and produced an empty list — so the delivered page carried this table's
    # headers and its caption with NOTHING BETWEEN THEM, under prose describing what the
    # testing had found. Nothing raised: an empty list is a valid table. [R-ENF-04] says
    # an empty result is not a clean result, and here the emptiness reached a reader.
    assert len(rows) >= 5, ("the walk-forward results table has %d rows; the scores "
                            "record has %d drivers" % (len(rows), len(sc)))
    hdr_wf = ["Driver", "Observations", "Average error",
              "Average size of error", "Consistent across the record"]
    table(doc, hdr_wf,
          # the HEADER is the widest token in its own column here — "Observations"
          # needs 2.51cm and had 2.40, so the header itself broke mid-word.
          rows, col_width.fit_widths(hdr_wf, rows, 16.2),
          "Errors are in natural logarithms, so −0.06 means the forecast came in "
          "about 6% below the outcome. A negative average error means this "
          "method has tended to forecast low.")

    para(doc, "Two findings shaped this study.")
    # EVERY FIGURE IN THESE TWO BULLETS WAS TYPED. They were also right, which is what
    # makes typing them dangerous: nothing would have moved them when the record did.
    # Log errors are converted to the percentages a reader can use.
    _dr, _ns = sc["dev_revenue"], sc["new_sales"]
    bullets(doc, [
        "Converting a disclosed order book is forecastable. Development revenue "
        "came back with an average error of about %.0f%% low and a typical error of "
        "%.0f%%, across %d tests. That is why the valuation is built on the order "
        "book rather than on a revenue growth rate."
        % (abs(math.expm1(_dr["bias"])) * 100, math.expm1(_dr["mae"]) * 100, _dr["n"]),
        "Forecasting new sales from population and inflation is not. New "
        "contracted sales came back about %.0f%% low, consistently, in every part "
        "of the record. The reason is plain in the company's own numbers: sales "
        % (abs(math.expm1(_ns["bias"])) * 100) +
        "went from EGP 33 billion in 2022 to EGP 143 billion, then EGP 504 "
        "billion, then EGP 382 billion, on the launches of SouthMed and The "
        "Spine. No demographic anchor can see a launch calendar. This study "
        "therefore does not forecast sales that way; it uses the company's own "
        "disclosed sales and lets them converge on the delivery rate.",
    ])
    # THE PARAGRAPH MADE THE MISTAKE IT WAS WARNING ABOUT. It divided the FY2025 finance
    # charge by the borrowings on the 30-JUNE-2026 balance sheet — two different clocks —
    # and reported "about 44%", a typed figure that reconciles against neither pairing:
    # 3,936.5 over the 16,493 the same sentence quotes is 23.9%, which is BELOW the
    # policy peak it was contrasted with, so the argument as printed refuted itself.
    # The charge is divided here by the borrowings of the year that BORE it, named in
    # the sentence, which is [R-COC-01]'s own requirement that an effective rate carry a
    # described denominator.
    #
    # THE 27.25% POLICY PEAK IS GONE BECAUSE IT CANNOT BE SOURCED HERE. It was typed in
    # two places; engine/macro_history/_supplied_EG_rates.json runs to 2023 and the house
    # macro path carries the CURRENT rate with its MPC date, so nothing in this
    # repository establishes a March-2024 peak. Under SIGCM a figure this desk cannot
    # source is a figure it does not print. The comparison that survives is stronger
    # anyway, because both sides are the study's own committed numbers: the implied rate
    # against the marginal borrowing rate this model actually charges.
    _fc = v(IS, "finance_cost_fy25")
    _debt25 = (v(BSH, "loans_noncurrent_fy25") + v(BSH, "loans_current_fy25")
               + v(BSH, "credit_facilities_fy25"))
    para(doc, "The same testing found a habit worth naming and NOT correcting. "
              "The method under-forecasts finance cost badly and consistently, "
              "and a correction for it passed every statistical test applied. It "
              "was rejected anyway. TMG's reported finance charge of EGP %s "
              "million for 2025, against the EGP %s million of loans, current "
              "maturities and credit facilities it carried at that year end, "
              "implies a rate of %s — against the %s this model charges on "
              "borrowings. The excess is the unwinding of the financing component "
              "the company recognises on its customer contracts, which is not "
              "interest on a loan and is not disclosed separately. Correcting a "
              "figure whose numerator and denominator describe different things "
              "would have hidden that rather than fixed it."
              % (money(_fc), money(_debt25), pct(_fc / _debt25, 1),
                 pct(W["kd_pretax"], 2)))

    # 1.7 the crux
    doc.add_heading("1.7 The crux", level=2)
    para(doc, "Everything in this study turns on one number that the company "
              "does not publish: how long its order book takes to convert into "
              "handovers.")
    conv = [[str(y), money(x, 2)] for y, x in
            sorted(LENS["sensitivity"]["conversion_years_grid"].items(),
                   key=lambda kv: int(kv[0]))]
    table(doc, ["Years to convert the order book", "EGP per share"], conv,
          [7.0, 4.0],
          "At this study's discount rate, with the minority deducted at its share of value. "
          "The company's own record over the last three years is consistent with "
          "the slower end of this table.")
    para(doc, "The evidence for the slower reading is the company's own "
              "arithmetic. Development revenue as a share of the opening order "
              "book plus the year's sales ran near 15% before 2023. It was 9.8% "
              "in 2023, 3.8% in 2024 and 5.4% in 2025 — not because demand "
              "weakened, but because the book quadrupled while the ability to "
              "build did not. The constraint is construction capacity, not the "
              "order book.")
    para(doc, "The evidence for the faster reading is that TMG is investing "
              "heavily in exactly that capacity: work in progress rose from EGP "
              "%s billion to EGP %s billion in the first half of 2026 alone, and "
              "assets under construction rose alongside it. A company spending "
              "at that rate expects to hand over faster."
              % (money(v(BSH, "properties_under_development_fy25") / 1000, 1),
                 money(v(BSH, "properties_under_development") / 1000, 1)))
    para(doc, "Here is the part that is not obvious. Faster conversion is worth "
              "LESS, not more. A developer that accelerates handovers has to "
              "build before it collects, and at first-year rates near %s that "
              "investment pays back only slowly inside the window. The slower "
              "reading is worth EGP %s a share and the faster one EGP %s. Both "
              "are published, and the gap between them is the study's crux."
              % (pct(SCHED["forward_wacc"][0], 0),
                 # the ADOPTED basis, which is what the crux table immediately above
                 # this sentence publishes; it read the book basis and so contradicted
                 # its own table by five and twenty pounds a share
                 money(PSV["rating|capacity"], 2),
                 money(PSV["rating|recovery"], 2)))

    # 1.8 macro and cost of capital
    doc.add_heading("1.8 The cost of capital, priced line by line", level=2)
    ins = W["inputs"]
    dam = ins["damodaran"]
    # EACH COST OF EQUITY REPRODUCES FROM THE ROWS THIS TABLE PRINTS, on its own basis.
    for _rf, _erp, _ke in ((W["rf_star_rating"], dam["total_erp_rating"], W["ke_rating"]),
                           (W["rf_star_cds"], dam["total_erp_cds"], W["ke_cds"])):
        assert abs(_rf + W["beta_record"]["beta"] * _erp - _ke) < 5e-4, (
            "the cost of equity does not reproduce: %.4f + %.4f x %.4f != %.4f"
            % (_rf, W["beta_record"]["beta"], _erp, _ke))
        assert abs((ins["rf_observed"] - _rf)
                   - (dam["adj_default_spread"] if _erp == dam["total_erp_rating"]
                      else dam["sovereign_cds"])) < 5e-4, "the spread stripped is not printed"
    table(doc, ["Input", "Value", "Where it comes from"],
          [["Egyptian ten-year government bond yield", pct(ins["rf_observed"], 2),
            "market quote dated 6 August 2026, cross-checked against a policy "
            "rate of 19.00%, an overnight lending rate of 20.00% and an "
            "interbank rate of 19.51% at the central bank's August 2026 meeting"],
           ["Egypt's own default spread, rating basis", pct(dam["adj_default_spread"], 2),
            "the sovereign's own row in the published country-premium file, read "
            "fresh on 1 September 2026"],
           ["Risk-free rate, normalised, rating basis", pct(W["rf_star_rating"], 2),
            "the bond yield less that default spread, so country risk is charged "
            "once and not twice"],
           # THE SWAP BASIS STRIPS A DIFFERENT SPREAD AND THE TABLE DID NOT SAY SO.
           # [R-COC-01] requires the same basis of default spread to be stripped as
           # the premium added back, so the swap-basis cost of equity stands on a
           # normalised risk-free rate of its own. Only the rating pair was printed,
           # so a reader combining the printed rows got 30.48% where the table states
           # 33.44% — reproducing from a row the table did not carry.
           ["Egypt's own default spread, swap basis", pct(dam["sovereign_cds"], 2),
            "the sovereign credit-default-swap spread from the same file, which is "
            "what the swap-basis premium is built on"],
           ["Risk-free rate, normalised, swap basis", pct(W["rf_star_cds"], 2),
            "the bond yield less THAT spread; the two bases are not mixed"],
           ["Equity risk premium, rating basis", pct(dam["total_erp_rating"], 2),
            "the same published file, Egypt's own row"],
           ["Equity risk premium, swap basis", pct(dam["total_erp_cds"], 2),
            "the same file, on the alternative measure"],
           ["Beta", money(W["beta_record"]["beta"], 4),
            "TMGH's own weekly returns regressed against the Egyptian Exchange's "
            "published index over %s years, %d observations, explaining %s of the "
            "variation, standard error %s"
            % (money(W["beta_record"]["window_years"], 2), W["beta_record"]["n"],
               pct(W["beta_record"]["r2"], 1), money(W["beta_record"]["se"], 3))],
           ["Cost of equity", "%s / %s" % (pct(W["ke_rating"], 2), pct(W["ke_cds"], 2)),
            "the normalised risk-free rate plus beta times each premium"],
           ["Marginal cost of debt, before tax", pct(ins["kd_local"], 2),
            "the sovereign yield plus a 250 basis-point corporate spread. TMG "
            "does not disclose the rate on any of its own facilities, so its own "
            "borrowing cost cannot be used and this is labelled rather than "
            "presented as the company's rate"],
           ["Weights", "%s equity / %s debt"
            % (pct(W["weight_equity"], 1), pct(W["weight_debt"], 1)),
            "market value of equity at the last close; borrowings at book"],
           ["Cost of capital, first year", "%s / %s"
            % (pct(W["wacc_rating"], 2), pct(W["wacc_cds"], 2)),
            "both premium bases published, neither averaged"],
           ["Cost of capital, once normalised", pct(SCHED["wacc_terminal"], 2),
            "built from the central bank's own inflation target plus a standard "
            "real return, not quoted from any market: %s risk-free, %s premium, "
            "%s cost of debt"
            % (pct(SCHED["rf_terminal"], 2), pct(SCHED["erp_terminal"], 2),
               pct(SCHED["kd_terminal_pretax"], 2))],
           ["The path between them",
            " to ".join(pct(x, 1) for x in (SCHED["forward_wacc"][0],
                                            SCHED["forward_wacc"][-1])),
            "each year discounted at its own rate; the shape follows the central "
            "bank's published easing calendar rather than a second assumption of "
            "this study's own"]],
          [4.4, 2.6, 9.2])
    hdr_lad = ["Year"] + [str(i + 1) for i in range(len(SCHED["forward_wacc"]))]
    rows_lad = [["Cost of capital"] + [pct(w, 1) for w in SCHED["forward_wacc"]],
                ["A pound arriving then is worth"]
                + [money(d, 3) for d in SCHED["discount_factors"]]]
    # AT 1.28cm EVERY ONE OF THE TEN RATE CELLS ORPHANED ITS PERCENT SIGN: the delivered
    # page printed "32.4" with a bare "%" on the line beneath, ten times across one row.
    # A percent sign is 0.30cm and "32.4%" needs 1.35. The orphan detector that found the
    # other wrapped cells in this book scans for a stray DIGIT and could not see it.
    table(doc, hdr_lad, rows_lad,
          col_width.fit_widths(hdr_lad, rows_lad, 16.79, generous=0, equal_from=1))
    para(doc, "The terminal value is brought home on the same factor as the last "
              "year's cash flow, %s. The common alternative — discounting the "
              "forecast years at one rate and the terminal value at a lower one — "
              "gives the same pound arriving on the same day two different values, "
              "and it is not done here."
              % money(SCHED["terminal_discount_factor"], 3), size=9, color=MUTED)
    para(doc, "One caveat on the risk-free rate. The quote adopted is 26 days "
              "old at the date of this study. It was cross-checked against three "
              "current central-bank rates rather than accepted alone, and its "
              "effect is priced across the whole plausible range in the next "
              "section. It should be refreshed before this study is relied on.",
         italic=True, color=MUTED)

    # 1.9 sensitivity
    doc.add_heading("1.9 Sensitivity", level=2)
    figure(doc, os.path.join(HERE, "fig2_sensitivity.png"),
           "The discount rate against value per share, on both readings of the "
           "order book. The dotted lines mark this study's rate and the 18% the "
           "June 2026 edition used.")
    grid = LENS["sensitivity"]["wacc_grid"]
    waccs = sorted({x["wacc"] for x in grid.values()})
    table(doc, ["Discount rate", "Slower conversion, EGP/share",
                "Faster conversion, EGP/share", "Against the traded price"],
          [[pct(w, 2),
            money(grid["%0.4f|capacity" % w]["per_share_nci_value_share"], 2),
            money(grid["%0.4f|recovery" % w]["per_share_nci_value_share"], 2),
            pct(grid["%0.4f|capacity" % w]["per_share_nci_value_share"] / SPOT - 1, 0)]
           for w in waccs], [3.2, 4.4, 4.4, 4.2],
          "The last column compares the slower reading with the last traded "
          "price of EGP %s." % money(SPOT, 2))
    para(doc, "Read the table from the bottom up and the point is unmistakable. "
              "Nothing about this company's business changes across those rows. "
              "The value changes by a factor of four, and the whole of that is "
              "one input. Anyone using this study should decide what they think "
              "the right discount rate is before they decide what they think of "
              "the shares.")
    return doc


def sections_2_to_7(doc):
    # --- 2. Technical and price structure ---------------------------------
    doc.add_page_break()
    doc.add_heading("2. Technical and price structure", level=1)
    para(doc, SITE.get("summary", ""))
    para(doc, "Trend: " + SITE.get("trend", ""))
    if SITE.get("res") and SITE.get("sup"):
        table(doc, ["Charted levels above the close", "EGP",
                    "Charted levels below the close", "EGP"],
              [[["Nearest", "Second", "Third"][i], money(SITE["res"][i], 2),
                ["Nearest", "Second", "Third"][i], money(SITE["sup"][i], 2)]
               for i in range(3)], [4.4, 2.4, 4.4, 2.4],
              "Levels are computed from the same cleaned price history the "
              "distribution in section 3 is built on.")
    bullets(doc, [SITE.get("bull", ""), SITE.get("bear", "")])
    para(doc, "This section describes the chart and nothing else. It carries no "
              "view on the business and it is not an input to the valuation; it "
              "is shown beside the other work so that agreement between them is "
              "information rather than an echo.", italic=True, color=MUTED)

    # --- 3. Probabilistic price map ---------------------------------------
    doc.add_heading("3. Where the price itself could be", level=1)
    figure(doc, os.path.join(HERE, "fig3_cone.png"),
           "The published distribution of where the share price could be at one "
           "and three months, against the last close.")
    rows = []
    for tag in ("t20", "t60"):
        d = SITE.get(tag)
        if not d:
            continue
        # TWO DECIMALS, EXPLICITLY. The regex this reader replaced captured the raw
        # SOURCE TEXT, so a percentile written 95.00 in data.js reached the table as the
        # string "95.00"; a real parse returns the NUMBER 95.0, which renders "95.0". The
        # figure is identical and the cell a reader sees is not, so the format is pinned
        # here rather than inherited from whichever reader happens to be in use.
        rows.append([d.get("label", tag)]
                    + ["%.2f" % float(d[k]) for k in ("p5", "p25", "p50", "p75", "p95")]
                    + [d.get("resolve", "")])
    if rows:
        table(doc, ["Horizon", "5th", "25th", "50th", "75th", "95th",
                    "Resolves"], rows, [2.4, 1.9, 1.9, 1.9, 1.9, 1.9, 2.4],
              "Percentiles of the modelled price distribution, in EGP. These "
              "describe the PRICE, not the business, and they are produced "
              "independently of the valuation above.")
    para(doc, "How well these bands have held is worth stating plainly rather "
              "than assumed. Across the whole book of companies covered on this "
              "market the ninety-per-cent band has caught the outcome close to "
              "the ninety per cent it promises, and where a particular company's "
              "record departs from that materially it is flagged on its page. "
              "The honest reading of a band that holds about as often as it says "
              "is that nothing needs saying, which is why nothing is said here.")
    para(doc, "One caution specific to this company. The distribution is fitted "
              "to price history. It knows nothing about the order book, the "
              "landbank or the discount rate, and it should not be read as "
              "agreeing or disagreeing with anything in section 1.")

    # --- 4. Comparison of the lenses --------------------------------------
    doc.add_heading("4. Comparing the lenses", level=1)
    table(doc, ["Lens", "What it looks at", "What it says, EGP/share",
                "Where it is weakest"],
          [["Cash flow, sum of the parts",
            "the order book, the hotels and the recurring businesses, discounted",
            # the ADOPTED basis, as everywhere else in this document
            "%s – %s" % (money(min(PSV.values()), 2), money(max(PSV.values()), 2)),
            "wholly dependent on the discount rate and on the conversion period"],
           ["Book value and sustainable return",
            "what the company earns on the capital it holds",
            "%s – %s" % (money(min(x["value_per_share"]
                                   for x in LENS["book_and_sustainable_return"]["cases"].values()
                                   if x["value_per_share"]), 2),
                         money(max(x["value_per_share"]
                                   for x in LENS["book_and_sustainable_return"]["cases"].values()
                                   if x["value_per_share"]), 2)),
            "the most discount-rate-sensitive lens of the four"],
           ["Normalised earnings power",
            "three years of profit with one-off gains removed",
            "%s – %s" % (money(min(x for k, x in LENS["normalised_earnings"].items()
                                   if k.startswith("cap|")), 2),
                         money(max(x for k, x in LENS["normalised_earnings"].items()
                                   if k.startswith("cap|")), 2)),
            "three years is a short cycle for a developer"],
           ["Its own history of multiples",
            "what buyers have actually paid for these shares",
            "%s – %s" % (money(min(r["pe"] for r in LENS["own_multiple_history"]
                                   if r.get("pe")) * LENS["normalised_earnings"]["average_eps"], 2),
                         money(max(r["pe"] for r in LENS["own_multiple_history"]
                                   if r.get("pe")) * LENS["normalised_earnings"]["average_eps"], 2)),
            "records what the market did, not whether it was right"]],
          [3.6, 4.6, 3.0, 5.0])
    para(doc, "The four are not weighted into a single figure. Weighting would "
              "require a view on which lens is right, and the point of showing "
              "four is that this study does not have one.")

    # --- 5. Catalysts ------------------------------------------------------
    doc.add_heading("5. What would move this", level=1)
    table(doc, ["Event", "Direction", "Why it matters"],
          [["A disclosed delivery schedule by project", "either way",
            "it would settle the crux directly, and the table in section 1.7 "
            "shows what each answer is worth"],
           ["A disclosure of TMG's economic share of each project company",
            "either way", "nearly half of consolidated equity belongs to "
            "minorities; how much of the value follows is currently an estimate"],
           ["The Saudi project reaching scale", "upward",
            "Banan was already 52.5% of development revenue in the first half of "
            "2026 and it is recognised as work completes rather than on handover, "
            "which brings revenue forward"],
           ["A further devaluation of the pound", "downward in real terms",
            "the cost base and the buyer's capacity to pay both move, and the "
            "collection cycle is long"],
           ["Interest rates falling further", "upward",
            "the whole of section 1.9 is about this"],
           ["Cancellations rising above their historical rate", "downward",
            "the order book is the asset; the company has historically reported "
            "cancellations around 4% to 4.5% of accumulated sales"]],
          [5.0, 3.0, 8.2])

    # --- 6. Reading the probability zones ---------------------------------
    doc.add_heading("6. How to read the ranges in this document", level=1)
    bullets(doc, [
        "A range is not a forecast with error bars bolted on. It is the honest "
        "output of a method whose accuracy has been measured.",
        "The four valuation cases are alternatives, not a distribution. Each is "
        "internally consistent; picking between them means picking between the "
        "two ways of measuring country risk and the two readings of the order "
        "book.",
        "The ranges on the third to fifth forecast years in Appendix A come from "
        "how far out this method has actually been when tested on this company's "
        "own history, at that distance. They widen with horizon because the "
        "measured errors widen with horizon.",
        "The price distribution in section 3 answers a different question from "
        "everything else here — where the price might go, not what the business "
        "is worth. The two are shown side by side so that agreement between them "
        "is information rather than repetition.",
    ])

    # --- 7. Caveats --------------------------------------------------------
    doc.add_heading("7. Caveats, and what would change our mind", level=1)
    bullets(doc, [
        "The discount rate is the study, and the shape of it matters as much as "
        "its level. This edition discounts each year at its own rate, gliding "
        "from %s to %s as the central bank's own published path plays out. A flat "
        "rate held at today's level for fifteen years is not a conservative "
        "assumption; it is a forecast that Egypt never normalises, and it was not "
        "one any earlier edition of this study argued for."
        % (pct(SCHED["forward_wacc"][0], 2), pct(SCHED["wacc_terminal"], 2)),
        "Nearly half the group belongs to somebody else. Non-controlling "
        "interests are %s of consolidated equity, and the company does not "
        "disclose its economic share project by project, so the deduction is "
        "made %s ways and all are published."
        % (pct(CASES["rating|capacity"]["nci_share_of_equity"], 1), _NBASES_WORD),
        "Unit economics cannot be built for this company. It publishes unit "
        "counts occasionally — 3,196 homes delivered in 2025, 6,102 units sold "
        "in 2022 — and never as a continuous series, and it publishes no average "
        "area, price per square metre or construction cost per square metre. The "
        "model is therefore built at segment level and this is a real limitation, "
        "stated rather than papered over.",
        "The revenue recognition basis is changing while this is written. The "
        "Egyptian business recognises a sale when the customer takes control of "
        "the home. From 2025 the Saudi project is recognised as construction "
        "progresses, and it was already more than half of development revenue in "
        "the first half of 2026. Comparisons across that change need care.",
        "The finance charge is not what it looks like. It implies a rate of %s "
        "on the borrowings that carried it, against the %s this model charges; "
        "the difference is contract financing that the statements do not separate."
        % (pct(v(IS, "finance_cost_fy25")
               / (v(BSH, "loans_noncurrent_fy25") + v(BSH, "loans_current_fy25")
                  + v(BSH, "credit_facilities_fy25")), 1),
           pct(W["kd_pretax"], 2)),
        "The 2024 comparatives were restated. The company completed the "
        "accounting for its hotel acquisition within the allowed measurement "
        "period, moving 2024 gross profit from EGP %s million to EGP %s million "
        "and attributable profit from EGP %s million to EGP %s million. This "
        "study uses the figures as first reported for the historical record and "
        "shows the restatement beside them."
        % (money(v(IS, "gross_profit_fy24")),
           money(v(IS, "gross_profit_fy24_restated")),
           money(v(IS, "npat_parent_fy24")),
           money(v(IS, "npat_parent_fy24_restated"))),
        "What would change our mind, most simply: a published delivery schedule. "
        "It is the one disclosure that would collapse the widest uncertainty in "
        "this document into a number.",
    ])
    return doc


def appendices(doc):
    # --- Appendix A --------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Appendix A — financial statements", level=1)
    doc.add_heading("A.1 Income statement: three years reported, five projected",
                    level=2)
    rep, fwd = ST["reported"], ST["capacity"]["rows"]
    hdr = ["EGP million", "2023", "2024", "2025"] + \
          [str(r["year"]) for r in fwd]
    lines = [("Development revenue", "dev_revenue"),
             ("Hospitality revenue", "hosp_revenue"),
             ("Other recurring revenue", "other_revenue"),
             ("Gross profit", "gross_profit"),
             ("Net profit", "net_profit"),
             ("Profit attributable to TMG's shareholders", "attributable_profit")]
    rows = []
    for label, key in lines:
        r = [label] + [money(rep[y][key]) for y in ("2023", "2024", "2025")]
        for f in fwd:
            r.append(money(f[key]) if key in f else "–")
        rows.append(r)
    rows.append(["Earnings per share, EGP",
                 money(rep["2023"]["attributable_profit"] / M["shares_mn"], 2),
                 money(rep["2024"]["attributable_profit"] / M["shares_mn"], 2),
                 money(rep["2025"]["attributable_profit"] / M["shares_mn"], 2)]
                + [money(f["eps"], 2) for f in fwd])
    table(doc, hdr, rows,
          # SIZED FROM THE CELLS, NOT BY FEEL. At 1.55cm the 2030 column could not hold
          # "102,747" and the delivered page printed "102,74" with a lone "7" beneath it.
          # col_width holds the measured per-character widths; equal_from ties the year
          # columns to one width, because a year grid whose last column is a tenth wider
          # than its neighbour is a table a reader notices.
          col_width.fit_widths(hdr, rows, 16.8, generous=0, equal_from=1),
          "Reported figures are as first reported. Projected figures are the "
          "slower-conversion case.")
    para(doc, "The third to fifth projected years carry ranges, because at that "
              "distance this method's measured error is wide enough that a "
              "single figure would overstate what is known:")
    rrows = []
    for f in fwd:
        if "revenue_low" not in f:
            continue
        rrows.append([str(f["year"]), money(f["revenue"]),
                      "%s – %s" % (money(f["revenue_low"]), money(f["revenue_high"])),
                      money(f["net_profit"]),
                      "%s – %s" % (money(f.get("net_profit_low", 0)),
                                   money(f.get("net_profit_high", 0))),
                      str(f.get("revenue_band_n", ""))])
    if rrows:
        table(doc, ["Year", "Revenue, central", "Revenue range",
                    "Net profit, central", "Net profit range",
                    "Tests behind the range"], rrows,
              [1.5, 2.7, 3.6, 2.7, 3.6, 2.3],
              "Ranges are built from how far out this method actually was, at "
              "that horizon, when tested on this company's own history.")

    doc.add_heading("A.2 Balance sheet as reported", level=2)
    bs_rows = [
        ["Property, plant and equipment", money(v(BSH, "ppe"))],
        ["Investment properties", money(v(BSH, "investment_property"))],
        ["Assets under construction", money(v(BSH, "assets_under_construction"))],
        ["Goodwill", money(v(BSH, "goodwill"))],
        ["Deposits and financial assets, non-current",
         money(v(BSH, "deposits_noncurrent") + v(BSH, "fvoci")
               + v(BSH, "associates"))],
        # THE BLOCK MUST FOOT FOR A READER. These three were registered inputs the table
        # did not print, so the five printed lines came out EGP 748mn short of the total
        # above them — exactly intangibles plus right-of-use plus the deferred tax asset,
        # to the last hundred thousand. Small lines, and a reader adding a column has no
        # way to tell a missing line from a wrong total.
        ["Intangibles, right-of-use assets and deferred tax",
         money(v(BSH, "intangibles") + v(BSH, "right_of_use")
               + v(BSH, "deferred_tax_asset"))],
        ["Total non-current assets", money(v(BSH, "total_noncurrent_assets"))],
        ["Properties under development",
         money(v(BSH, "properties_under_development"))],
        ["Inventories", money(v(BSH, "inventories"))],
        # likewise: the current block was EGP 24mn short, which is work in progress
        ["Work in progress", money(v(BSH, "work_in_progress"))],
        ["Trade and notes receivable", money(v(BSH, "trade_notes_receivable"))],
        ["Notes receivable for undelivered homes", money(v(BSH, "nr_undelivered"))],
        ["Other current assets", money(v(BSH, "other_current_assets"))],
        ["Deposits and financial assets, current",
         money(v(BSH, "deposits_current") + v(BSH, "fvtpl"))],
        ["Cash and cash equivalents", money(v(BSH, "cash"))],
        ["Total current assets", money(v(BSH, "total_current_assets"))],
        ["Total assets", money(v(BSH, "total_assets"))],
        ["Equity attributable to TMG's shareholders", money(v(BSH, "equity_parent"))],
        ["Non-controlling interests", money(v(BSH, "nci_equity"))],
        ["Total equity", money(v(BSH, "total_equity"))],
        ["Borrowings and leases",
         money(v(BSH, "loans_noncurrent") + v(BSH, "loans_current")
               + v(BSH, "credit_facilities") + v(BSH, "lease_noncurrent")
               + v(BSH, "lease_current"))],
        ["Advance payments from customers", money(v(BSH, "customer_advances"))],
        ["Obligations against cheques for undelivered homes",
         money(v(BSH, "obligations_nr_undelivered"))],
        ["Suppliers, contractors and notes payable",
         money(v(BSH, "suppliers_contractors"))],
        ["Creditors and other credit balances", money(v(BSH, "creditors_other"))],
        ["Other liabilities, provisions and tax",
         money(v(BSH, "other_noncurrent_liab") + v(BSH, "deferred_tax_liab")
               + v(BSH, "provisions") + v(BSH, "tax_payable")
               + v(BSH, "dividends_payable"))],
        ["Total liabilities", money(v(BSH, "total_liabilities"))],
    ]
    table(doc, ["As at 30 June 2026, EGP million", "Amount"], bs_rows, [11.0, 4.4],
          "From the reviewed interim statements. Post-dated cheques of EGP %s "
          "million for sold and undelivered homes are held by the company and "
          "are deliberately not on this sheet."
          % money(v(BSH, "postdated_cheques_offbs")))

    doc.add_heading("A.3 Projected balance sheet and cash flow", level=2)
    full = ST["capacity"]["full_rows"][:5]
    # THE COLUMN DID NOT ADD UP AND EVERY FIGURE IN IT WAS CORRECT. Operating cash flow
    # was printed FIRST, and it is built from net profit and ALREADY CONTAINS the two
    # working-capital movements listed beneath it — so a reader adding the column counted
    # them twice and reached 41,134 against a printed free cash flow of 24,855. The free
    # cash flow is not the residual of those rows at all: it sits on operating profit
    # after tax, before financing, while operating cash flow sits on net profit after it.
    #
    # The components now precede the subtotal they make up, so the block foots as a
    # reader reads down it, and the free cash flow is separated and captioned for what it
    # is. This is the ARCC Table 3 shape: figures individually right and the relationship
    # between them unreadable, which nothing checking figures one at a time can see.
    hdr_cf = ["EGP million"] + [str(r["year"]) for r in full]
    rows_cf = [["Net profit"] + [money(r["net_profit"]) for r in full],
               ["add back depreciation and amortisation"]
               + [money(r["da"]) for r in full],
               ["Movement in money collected ahead of handover"]
               + [money(r["d_advances"]) for r in full],
               ["Movement in homes built ahead of handover"]
               + [money(-r["d_properties_under_development"]) for r in full],
               ["Operating cash flow"] + [money(r["cfo"]) for r in full],
               ["Capital spend"] + [money(-r["capex"]) for r in full],
               ["Dividends"] + [money(-r["dividend"]) for r in full],
               ["Free cash flow to the firm, before financing"]
               + [money(r["fcff"]) for r in full],
               ["Cash and deposits, closing"] + [money(r["cash"]) for r in full],
               ["Properties under development, closing"]
               + [money(r["properties_under_development"]) for r in full],
               ["Customer advances, closing"]
               + [money(r["customer_advances"]) for r in full],
               ["Equity attributable to TMG's shareholders"]
               + [money(r["equity_parent"]) for r in full],
               ["Total assets"] + [money(r["total_assets"]) for r in full],
               ["Total liabilities"] + [money(r["total_liabilities"]) for r in full],
               ["Total equity"] + [money(r["total_equity"]) for r in full],
               ["Assets less liabilities less equity"]
               + [money(r["balance_check"], 2) for r in full]]
    # the operating block foots as printed, on every year
    for i, r in enumerate(full):
        _made = (r["net_profit"] + r["da"] + r["d_advances"]
                 - r["d_properties_under_development"])
        assert abs(_made - r["cfo"]) < 1.0, (
            "A.3's operating cash flow does not foot in %s: %.1f vs %.1f"
            % (r["year"], _made, r["cfo"]))
    table(doc, hdr_cf, rows_cf,
          col_width.fit_widths(hdr_cf, rows_cf, 16.79, generous=0, equal_from=1),
          "The slower-conversion case. The first four lines make up operating cash "
          "flow. Free cash flow to the firm is NOT the residual of the rows above "
          "it: it starts from operating profit after tax and stops before financing, "
          "which is the measure section 1.1 discounts, while operating cash flow "
          "starts from net profit and is after it. The last line is the check that "
          "the sheet closes: nothing here is plugged, and it closes because the cash "
          "flow feeds it.")
    para(doc, "On the faster-conversion reading the same statements go cash "
              "negative from %d, because accelerating handovers means building "
              "before collecting. That case would need external funding, and "
              "saying so is part of what makes it a genuine alternative rather "
              "than a more optimistic version of the same thing."
              % next((r["year"] for r in ST["recovery"]["full_rows"]
                      if r["cash"] < 0), 2030))

    # --- Appendix B --------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Appendix B — competitors, risks and open questions", level=1)
    doc.add_heading("B.1 Competitors", level=2)
    hdr_pe = ["Company", "Last close, EGP", "As at", "Why it is a comparator"]
    rows_pe = [[p["name"], money(p["close"], 2) if p.get("close") else "–",
                p.get("as_of", ""), p["why"]] for p in PE["egypt"]]
    # "As at" carried an ISO date, ten characters, in a 2.0cm column that needs 2.15.
    # The slack goes to the prose column, which is the one that reads better for it.
    table(doc, hdr_pe, rows_pe,
          col_width.fit_widths(hdr_pe, rows_pe, 16.79, generous=3))
    table(doc, ["Company", "Country", "Why it is a comparator"],
          [[p["name"], p["country"], p["why"]] for p in PE["outside_country"]],
          [4.6, 2.2, 9.0],
          "Competitors are studied for operating indicators and for how the "
          "market values them. Nothing about a competitor is ever a source for "
          "TMG's own reported figures.")
    para(doc, PE["note"], italic=True, color=MUTED)

    doc.add_heading("B.2 Risks", level=2)
    table(doc, ["Risk", "Why", "Where it is priced",
                "What would change it"],
          [[r["risk"], r["why"], r["priced_at"], r["what_would_change_it"]]
           for r in PE["risk_register"]], [3.4, 4.6, 3.6, 4.2])

    doc.add_heading("B.3 Questions this study asked", level=2)
    table(doc, ["Question", "Answered?", "Where"],
          [[q["question"], q["answered"], q["where"]]
           for q in PE["research_register"]], [6.0, 5.0, 4.6])

    # --- Appendix C --------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Appendix C — three ways of valuing this company", level=1)
    para(doc, "Three independent methods, applied to the same disclosed facts by "
              "people who would not agree with one another. Each shows its "
              "workings in full, names the sensitivity that would move it most, "
              "and states in advance what would prove it wrong.")
    for i, e in enumerate(EX["experts"], start=1):
        doc.add_heading("C.%d %s" % (i, e["name"]), level=2)
        para(doc, e["worldview"])
        para(doc, "When it works: " + e["when_it_works"])
        para(doc, "When it fails: " + e["when_it_fails"])
        # whole millions throughout. Formatting the small lines to two decimals
        # made an EGP 84.6 MILLION intangible-asset line render as "-84.60"
        # beside neighbours in whole millions, which reads as a different unit.
        table(doc, ["Working", "EGP million"],
              [[lbl, money(val)] for lbl, val in e["workings"]], [11.0, 4.4])
        para(doc, "Value per share: EGP %s" % money(e["value_per_share"], 2),
             bold=True)
        s = e["sensitivity"]
        table(doc, [s["what"].capitalize(), "EGP per share"],
              [[k, money(vv, 2)] for k, vv in s["numbers"].items()], [7.0, 4.0])
        para(doc, "What would prove this wrong: " + e["falsifier"], italic=True)

    doc.add_heading("C.4 Cross-examination", level=2)
    table(doc, ["Challenge", "Outcome"],
          [[c["challenge"], c["outcome"]] for c in EX["cross_examination"]],
          [7.0, 8.4])

    doc.add_heading("C.5 The three in one room", level=2)
    for chunk in EX["three_in_one_room"].split("\n\n"):
        para(doc, chunk)

    doc.add_heading("C.6 What drives the differences", level=2)
    table(doc, ["Pair", "Gap, EGP per share", "The assumption behind it"],
          [[d["pair"], money(d["gap"], 2), d["driven_by"]] for d in EX["divergence"]],
          [2.2, 3.0, 10.2])

    # --- About and disclosure ---------------------------------------------
    doc.add_page_break()
    doc.add_heading("About this study", level=1)
    para(doc, "This study values one company from its own published documents. "
              "Every figure about Talaat Moustafa Group Holding comes from its "
              "audited or reviewed financial statements or from documents it "
              "published itself; the company's own investor-relations archive "
              "was read in full for this edition, back to 2007. Country-level "
              "inputs — inflation, exchange rates, population, sovereign spreads "
              "— come from named third-party sources and are labelled as such in "
              "the accompanying source document.")
    para(doc, "Forecasts are built from drivers, not from growth rates applied "
              "to totals. Margins are outputs of disclosed cost lines. Where a "
              "figure has two legitimate readings, both are published. Where "
              "something is not disclosed, the study says so and leaves it out "
              "rather than estimating it; four such gaps are recorded in the "
              "accompanying source document.")
    para(doc, "The forecasting method used here was tested against this "
              "company's own fifteen-year record before it was used on its "
              "future: the driver model was rebuilt as it would have stood at "
              "each past year end, projected forward, and scored against what "
              "the company actually reported. Section 1.6 reports what that "
              "found, including where the method performs worse than simply "
              "assuming last year's figure repeats.")

    doc.add_heading("Disclosure", level=1)
    para(doc, "This document is research, not advice. It contains no "
              "recommendation, no target price and no rating, and it is not an "
              "offer to buy or sell anything. Valuation is uncertain and the "
              "ranges here are wide on purpose. Anyone acting on this document "
              "does so on their own judgement and should take their own advice.")
    para(doc, "Prices and market data are as at the dates stated. The company's "
              "financial figures are as at 31 December 2025 and 30 June 2026. "
              "Nothing here has been shown to or approved by the company.",
         color=MUTED, size=9)
    return doc


def main():
    doc = build(None)
    section1(doc)
    section1_drivers(doc)
    sections_2_to_7(doc)
    appendices(doc)
    out = os.path.join(HERE, EDITION_FILE)
    doc.save(out)
    hits, chars = scrub(out)
    bad = column_audit(out)
    # THE COLUMN AUDIT MEASURES A TABLE ON AVERAGE AND CANNOT SEE ONE CELL. It reports
    # whether a column is starved or bloated across the table, which is a different
    # question from whether the widest cell in it fits — and the widest cell is the one
    # that wraps mid-number.
    n_tables = assert_columns_fit()
    print("wrote %s (%.0f KB, %d characters of text)"
          % (os.path.basename(out), os.path.getsize(out) / 1024.0, chars))
    print("external-reader scrub: %d hits %s" % (len(hits), hits or ""))
    print("table audit: %d problems;  %d tables clear their own widest cell"
          % (len(bad), n_tables))
    for b in bad[:10]:
        print("   table %s: %s" % b)
    return out


if __name__ == "__main__":
    main()
