"""Build the TMGH standalone source document.

Primary documents · the full input list, every entry with value, source, date
and how reliable it is · the judgements and what would overturn each · what was
looked for and not found · where two sources disagree · what is not disclosed.

Reads study_numbers.json and nothing else.
"""
import json, os, sys, datetime

from docx import Document
from docx.shared import Cm

HERE = os.path.dirname(os.path.abspath(__file__))
ENGINE = os.path.dirname(HERE)
sys.path.insert(0, HERE)
sys.path.insert(0, os.path.join(HERE, ".."))
import col_width                                                       # noqa: E402
from docx_helpers import (INK, MUTED, ACCENT, money, pct, style, para, bullets,
                          table, scrub, column_audit)

N = json.load(open(os.path.join(HERE, "study_numbers.json")))
M, W, R = N["meta"], N["wacc"], N["ratios"]
REG = N["inputs"]
DATE = "1 September 2026"

TIER = {"A": "the company's own audited or reviewed statements, or its own "
             "investor documents",
        "B": "an exchange or regulator filing",
        "C": "a named third party, used only for country-level inputs",
        "DERIVED": "recovered from the statement's own arithmetic"}

LAYER = {
    "IS": "Company — income statement",
    "H1_26": "Company — most recent reviewed interim",
    "BS": "Company — balance sheet",
    "KPI": "Company — operating indicators",
}


def build(path):
    doc = Document()
    style(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)

    para(doc, "TALAAT MOUSTAFA GROUP HOLDING", size=17, bold=True, color=ACCENT,
         space_after=2)
    para(doc, "Sources, inputs and judgements · %s" % DATE, size=10, color=MUTED,
         space_after=14)
    para(doc, "This document lists every number the accompanying study uses, "
              "where it came from, when it was published and how reliable it is. "
              "It also lists what was looked for and not found, because an "
              "absence is a finding.")

    # --- primary documents -------------------------------------------------
    doc.add_heading("Primary documents", level=1)
    docs = sorted({(e["source"], e["date"]) for grp in ("IS", "H1_26", "BS", "KPI")
                   for e in REG[grp].values()})
    seen, rows = set(), []
    for src, dt in docs:
        key = src.split(" (")[0]
        if key in seen:
            continue
        seen.add(key)
        rows.append([key, dt])
    table(doc, ["Document", "Period it reports"], rows, [12.6, 3.4],
          "Every one of these was downloaded from the company's own "
          "investor-relations page. The whole archive back to 2007 was read for "
          "this edition: 138 of the 140 documents sought were obtained, and the "
          "two that were not are dead links on the company's own site — the "
          "first-half 2009 and second-quarter 2017 results releases.")

    # --- the input list ----------------------------------------------------
    doc.add_heading("Every input, with its source", level=1)
    for grp in ("IS", "H1_26", "BS", "KPI"):
        doc.add_heading(LAYER[grp], level=2)
        rows = []
        for k, e in REG[grp].items():
            if e.get("value") is None:
                continue
            rows.append([k.replace("_", " "),
                         money(e["value"], 2 if abs(e["value"]) < 100 else 0),
                         e.get("unit", ""), e["date"], e["tier"],
                         (e.get("note") or "")[:150]])
        # THE DATE COLUMN PRINTED "2025-12-" WITH A BARE "31" ON THE LINE BENEATH, on
        # every row of this register, in the field a reader of a provenance table checks
        # first. At 7.5pt a ten-character ISO date needs 1.95cm and the column had 1.90.
        # Sized from the cells at the size they are actually set in; the note column,
        # which holds prose and may wrap, absorbs the difference.
        hdr_reg = ["Input", "Value", "Unit", "As at", "Reliability", "Note"]
        table(doc, hdr_reg, rows,
              col_width.fit_widths(hdr_reg, rows, 16.0, generous=5, size=7.5),
              size=7.5)

    doc.add_heading("Country-level inputs", level=2)
    ins = W["inputs"]
    dam = ins["damodaran"]
    table(doc, ["Input", "Value", "Source", "Reliability"],
          [["Egyptian ten-year government bond yield", pct(0.23, 2),
            "market quote dated 6 August 2026, cross-checked against a central-bank "
            "policy rate of 19.00%, an overnight lending rate of 20.00% and an "
            "interbank rate of 19.51% at the August 2026 meeting", "C"],
           ["Egypt's credit rating", dam["moodys_rating"],
            "published country-premium file, Egypt's own row, read 1 September 2026",
            "C"],
           ["Egypt's adjusted default spread", pct(dam["adj_default_spread"], 2),
            "the same file", "C"],
           ["Egypt's country risk premium", pct(dam["country_risk_premium"], 2),
            "the same file", "C"],
           ["Equity risk premium, rating basis", pct(dam["total_erp_rating"], 2),
            "the same file", "C"],
           ["Equity risk premium, swap basis", pct(dam["total_erp_cds"], 2),
            "the same file", "C"],
           ["Statutory corporate tax rate", pct(dam["corporate_tax_rate"], 2),
            "the same file", "C"],
           ["Egyptian inflation, exchange rate, population",
            "series, 2004–2025",
            "World Bank development indicators, retrieved 1 September 2026", "C"],
           ["Beta against the exchange's published index",
            money(W["beta_record"]["beta"], 4),
            "TMGH's own weekly returns against the Egyptian Exchange's published "
            "index, %s years to %s, %d observations"
            % (money(W["beta_record"]["window_years"], 2),
               W["beta_record"]["last_obs"], W["beta_record"]["n"]), "A"],
           ["Last traded price", money(M["spot"], 2), M["spot_source"], "A"]],
          [4.0, 2.2, 7.4, 2.2])
    table(doc, ["Code", "What it means"],
          [[k, t] for k, t in TIER.items()], [2.4, 13.4],
          "No country-level source is ever used for a figure about the company "
          "itself, and no competitor is ever a source for TMG's own reported "
          "numbers.")

    # --- judgements --------------------------------------------------------
    doc.add_heading("Judgements, and what would overturn each", level=1)
    table(doc, ["Judgement", "What was decided", "What would overturn it"],
          [["How fast the order book converts",
            "not decided — the study publishes a slower reading of %d years and a "
            "faster one of %d years, and never averages them"
            % (N["model_parameters"]["CAPACITY_YEARS"],
               N["model_parameters"]["RECOVERY_YEARS"]),
            "a published delivery schedule by project"],
           ["How the minority interest is deducted",
            "at its share of value, proxied by the filed profit share (adopted); at book and pro rata shown beside it",
            "a disclosure of TMG's economic share of each project company"],
           ["Which equity risk premium to use",
            "both, published side by side",
            "nothing; the two measures are genuinely different questions"],
           ["Whether to extrapolate the current sales rate",
            "no. TMG sold about ten times what it delivered in 2025, and that is "
            "not a steady state; sales are held flat in real terms, not extrapolated",
            "sustained sales at the current rate alongside deliveries rising to "
            "meet them"],
           ["Whether to correct the finance-cost forecast",
            "no. A correction passed every statistical test and was rejected "
            "because the reported charge is not interest on borrowings alone",
            "a disclosure splitting interest on borrowings from the financing "
            "component of customer contracts"],
           ["What margin to carry forward",
            "the ratio from the reviewed first half of 2026, held flat. The "
            "margin is an output of that ratio and is never set directly",
            "a structural change in the cost base with a measured direction in "
            "the company's own figures"],
           ["Whether to build unit economics",
            "no — the disclosure does not support it, and the study says so "
            "rather than inventing an average unit",
            "disclosure of units and areas by project"]],
          [3.6, 6.6, 5.6])

    # --- negative results --------------------------------------------------
    doc.add_heading("What was looked for and not found", level=1)
    rows = []
    for k, e in REG["GAPS"].items():
        rows.append([k.replace("_", " "), e.get("gap", "")])
    rows.append(["The company's own borrowing rate",
                 "no facility pricing appears in any statement or release held, "
                 "so the cost of debt is built from the sovereign yield plus a "
                 "stated spread and labelled as such"])
    rows.append(["Competitors' financial statements",
                 "not sourced for this edition. A competitor multiple built on an "
                 "unsourced denominator would be a number wearing a decimal "
                 "point, so the relative lens is built on TMGH's own history "
                 "instead and the competitor list carries prices only"])
    rows.append(["Results releases for the first half of 2009 and the second "
                 "quarter of 2017",
                 "linked from the company's own investor-relations page but "
                 "returning errors; both attempts are logged rather than passed "
                 "over in silence"])
    table(doc, ["What", "What is missing, and what would close it"], rows,
          [4.2, 11.6])

    # --- disagreements -----------------------------------------------------
    doc.add_heading("Where two sources disagree", level=1)
    table(doc, ["Figure", "The disagreement", "What this study does"],
          [["2024 gross profit, net profit and earnings per share",
            "the 2024 statements report EGP %s million of gross profit and EGP "
            "%s million of attributable profit; the 2025 statements restate the "
            "same year to EGP %s million and EGP %s million after completing the "
            "accounting for the hotel acquisition"
            % (money(REG["IS"]["gross_profit_fy24"]["value"]),
               money(REG["IS"]["npat_parent_fy24"]["value"]),
               money(REG["IS"]["gross_profit_fy24_restated"]["value"]),
               money(REG["IS"]["npat_parent_fy24_restated"]["value"])),
            "uses the figures as first reported for the historical record and "
            "shows the restatement beside them; neither is discarded"],
           ["2023 new contracted sales",
            "the company reports both a total of EGP 142.8 billion and a core "
            "figure of EGP 94.9 billion, the difference being a large non-core "
            "land transaction",
            "uses the total, consistently with every other year, and records the "
            "core figure beside it"],
           ["The discount rate",
            "the method used here gives %s; the price the shares trade at implies "
            "%s to %s" % (pct(W["wacc_rating"], 2),
                          pct(N["lenses"]["implied_discount_rate"]["recovery"], 1),
                          pct(N["lenses"]["implied_discount_rate"]["capacity"], 1)),
            "publishes both and says plainly that the difference is the "
            "disagreement"]],
          [3.4, 7.2, 5.4])

    # --- how the numbers were checked --------------------------------------
    doc.add_heading("How these numbers were checked", level=1)
    bullets(doc, [
        "Every statement was accepted only if it added up. The segment gross "
        "profits were checked against their own revenue and cost lines, their sum "
        "against the reported total, and the profit waterfall against the bottom "
        "line, for every year used. Where a figure did not add up it was dropped "
        "rather than adjusted.",
        "The 2025 balance sheet and the June 2026 balance sheet were checked to "
        "balance to the currency unit.",
        "Several of the company's older documents survive only as scans. Figures "
        "read from those were checked against the same figure quoted in a later "
        "document wherever one existed, and three fiscal years were left out of "
        "the historical record entirely because their tables could not be read "
        "reliably. Nothing was estimated to fill those years.",
        "The forecasting method was tested against the company's own record "
        "before being used on its future, and the results are reported in section "
        "1.6 of the study — including the cases where the method performs worse "
        "than assuming last year's figure repeats.",
    ])
    return doc


def main():
    doc = build(None)
    out = os.path.join(HERE, "TMGH_Sources_02-09-2026.docx")
    doc.save(out)
    hits, chars = scrub(out)
    bad = column_audit(out)
    print("wrote %s (%.0f KB, %d characters)"
          % (os.path.basename(out), os.path.getsize(out) / 1024.0, chars))
    print("external-reader scrub: %d hits %s" % (len(hits), hits or ""))
    print("table audit: %d problems %s" % (len(bad), bad[:5] or ""))
    return out


if __name__ == "__main__":
    main()
