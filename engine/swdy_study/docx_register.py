"""SWDY_Bibliography_05-08-2026.docx — the companion bibliography document.
Every input in the model: value, source, date and research layer — emitted from
study_numbers.json (the compute script's own INPUTS block), plus the document
bibliography and the negative results."""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
D = json.load(open(os.path.join(HERE, 'study_numbers.json')))
INP = D['inputs']
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(table, top=36, bottom=36, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    table._tbl.tblPr.append(m)

def borders(table, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    table._tbl.tblPr.append(b)

def P(text='', size=10, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def H1(t): return P(t, size=13, bold=True, space_before=12, space_after=4)
def H2(t): return P(t, size=11, bold=True, space_before=8, space_after=3)

def table(rows, widths, size=8.2, header=True):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t); t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def masthead():
    t = doc.add_table(rows=1, cols=1)
    cell_margins(t, 80, 80, 150, 150)
    c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.1)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def fmt(v):
    if isinstance(v, float):
        return f"{v:,.4f}".rstrip('0').rstrip('.') if abs(v) < 100 else f"{v:,.1f}"
    if isinstance(v, int):
        return f"{v:,}"
    if isinstance(v, dict):
        return "; ".join(f"{k}: {fmt(x)}" for k, x in v.items())
    if isinstance(v, (list, tuple)):
        return ", ".join(fmt(x) for x in v)
    return str(v)

# ============================================================================
masthead()
H1('Elsewedy Electric Company S.A.E. (EGX: SWDY) — Bibliography and Source Register')
P('Companion document to the valuation study dated 5 August 2026. It records where every number '
  'in that study came from.', size=9.5, color=GREY)

H2('READ FIRST')
P('This document exists so that a reader can check the study rather than trust it. It lists every '
  'input the valuation model uses, together with the value, the source, the date of that source, '
  'and the research layer it belongs to. Nothing in the study is computed from a figure that does '
  'not appear here.')
P('Two things are worth knowing before reading the tables. First, inputs marked "House" are '
  'judgements or derivations made by the analyst, not disclosures by the company — each carries '
  'the reasoning that produced it, and the reader is free to disagree with it and re-run the '
  'model. Second, where a company disclosure could not be reached, that is recorded as a negative '
  'result at the end of this document rather than quietly filled in.')

H2('The research layers')
table([['Layer', 'What it covers'],
       ['Market', 'Prices, volumes and market-observable data for the security itself'],
       ['Country', 'Egyptian macroeconomic and sovereign data: policy rates, government bond '
        'yields, sovereign spreads, the exchange rate, the tax regime, the country equity risk '
        'premium'],
       ['Industry', 'The wires and cables sector, engineering and construction, input costs and '
        'the competitive frame'],
       ['Company', 'Elsewedy Electric\'s own audited and interim financial statements, earnings '
        'releases, exchange filings and disclosures'],
       ['House', 'Analyst judgement — forecast drivers, normalisation choices and cost-of-capital '
        'construction. Each is argued in place rather than asserted']],
      [1.05, 5.95], size=8.6)

H2('Primary documents relied upon')
table([['Document', 'Publisher', 'Date', 'What was taken from it'],
       ['Consolidated financial statements for the year ended 31 December 2023 (audited, '
        'translated from Arabic), KPMG Hazem Hassan, unqualified opinion', 'Elsewedy Electric '
        'Company', '13 March 2024',
        'Full consolidated income statement, balance sheet and cash flow statement for FY2023 '
        'with FY2022 comparatives; three-segment revenue and profit note; loans and borrowings '
        'note with average rates by currency; shareholder structure; exchange-rate disclosure '
        'note and the March-2024 devaluation subsequent event'],
       ['Consolidated financial statements for the year ended 31 December 2024 (audited, '
        'translated from Arabic), KPMG Hazem Hassan, unqualified opinion', 'Elsewedy Electric '
        'Company', '13 March 2025',
        'Full consolidated income statement, balance sheet and cash flow statement for FY2024 '
        'with FY2023 comparatives; three-segment revenue and profit note; loans and borrowings '
        'note including the financing-liability reconciliation; interest-rate risk note with '
        'average rates by currency; currency risk note; shareholder structure'],
       ['Consolidated financial statements for the year ended 31 December 2025 (audited, '
        'translated from Arabic), KPMG Hazem Hassan, unqualified opinion', 'Elsewedy Electric '
        'Company', '15 March 2026',
        'Full consolidated income statement, balance sheet and cash flow statement for FY2025 '
        'with FY2024 comparatives; three-segment revenue (Note 5-3) and profit (Note 16) note, '
        'each reconciling exactly to consolidated revenue and to operating profit through the '
        'corporate cost load; capital-management note (Note 29-1) with the company\'s own '
        '"net debt" definition; loans and borrowings note including the financing-liability '
        'reconciliation and average rates by currency; equity-accounted investees note; '
        'shareholder structure'],
       ['Condensed interim consolidated financial statements for the three months ended 31 March '
        '2026 (limited review), KPMG Hazem Hassan', 'Elsewedy Electric Company', '13 May 2026',
        'Q1-2026 income statement with Q1-2025 comparatives; balance sheet at 31 March 2026; '
        'three-segment revenue note reconciling exactly to Q1-2026 revenue; loans and borrowings '
        'note with average rates by currency; cash-flow statement; confirmation that no FY2025 '
        'dividend had been declared, proposed or approved as at the report date'],
       ['Country risk premium and default spread file', 'Damodaran, NYU Stern',
        '5 January 2026', 'Egypt equity risk premium and sovereign default spread, credit-default-'
        'swap basis and rating basis'],
       ['Egypt 10-year local-currency government bond yield', 'Market data, house cost-of-capital '
        'reference', '21 July 2026, re-verified 5 August 2026', 'The risk-free rate anchor'],
       ['Worldwide Tax Summaries — Egypt', 'PwC', '2026', 'Corporate income tax rate'],
       ['Daily price history for SWDY on the Egyptian Exchange', 'Supplied price series',
        'to 5 August 2026',
        'The anchor price, the volatility estimate, the moving-average structure, the beta '
        'regression and the price distributions'],
       ['Daily price history for the covered Egyptian equity library', 'House data library',
        'to August 2026', 'The 31-name equal-weight composite used as the market proxy in the '
        'beta regression']],
      [1.55, 1.25, 0.95, 3.25], size=8.0)

# ---- the four-field input register ------------------------------------------
H1('The full input register')
P('Every input to the valuation model, in the order the model declares them. The "Layer" column '
  'is the research layer defined above. Values are shown as the model holds them: EGP millions '
  'for financial-statement lines, decimals for rates and shares.', size=9.5, color=GREY)

for ring in ['Market', 'Company', 'Country', 'House']:
    items = [(k, v) for k, v in INP.items() if v['ring'] == ring]
    if not items:
        continue
    H2(f'{ring} layer — {len(items)} inputs')
    rows = [['Input', 'Value', 'Date', 'Source and construction']]
    for k, v in items:
        rows.append([k.replace('_', ' '), fmt(v['value']), v['date'], v['source']])
    table(rows, [1.15, 0.95, 0.72, 4.18], size=7.6)

# ---- judgements ---------------------------------------------------------------
H1('The judgements, stated separately')
P('These are the places where the analyst chose rather than observed. They are collected here so '
  'a reader can find them without reading the whole study.')
table([['Judgement', 'What was chosen', 'Why', 'What would overturn it'],
       ['Currency of discounting',
        'The full Egyptian cost of capital is charged to the whole company as the primary reading; '
        'the hard-currency alternative is computed and shown but not averaged in',
        'The shares trade, and dividends are paid, in Egyptian pounds on an Egyptian exchange, and '
        'realising value depends on Egyptian capital-account conditions',
        'Evidence that convertibility is not a binding constraint for this issuer would shift the '
        'primary reading materially higher'],
       ['Exchange-rate path',
        'About 6% a year of depreciation, far below what the interest-rate differential implies',
        'The base case assumes the central bank\'s disinflation path closes most of the gap rather '
        'than the currency absorbing it',
        'A disorderly move in the pound; the sensitivity table carries the parity case'],
       ['Working capital held near the FY2025 share of revenue',
        'Net working capital stays near 19.9% of revenue, the FY2025 audited level',
        'All three audited years show working capital absorbing cash rather than converting, '
        'though FY2025 improved (24.1% -> 23.1% -> 19.9% of revenue)',
        'Two consecutive years of operating cash flow above 60% of EBITDA'],
       ['Terminal growth of 5%',
        'The standing centre for established names in this market, sensitised 3–7%',
        'It is below the blended long-run nominal growth ceiling of the economies the company '
        'operates in, and is reconciled to the return on capital and reinvestment rate',
        'A demonstrated structural change in the export franchise\'s long-run growth'],
       ['Cost of debt at the blended rather than the domestic rate',
        'A currency-blended rate near the independently computed effective rate',
        'The audited notes disclose materially cheaper hard-currency borrowing, and most of the '
        'book is hard currency',
        'A shift of the debt book back into Egyptian pounds'],
       ['Minority interests charged at their share of group profit',
        'Rather than at book value',
        'Minorities take a larger share of profit than of book equity, so the profit share is the '
        'conservative charge',
        'A reader preferring the book convention can add the difference back; the amount is stated '
        'in the study'],
       ['Effective tax rate of 24.5% for NOPAT',
        'Above the Egyptian statutory rate, between the FY2025 print and the historical average',
        'Audited effective rates ran 31.3% (FY2023), 30.1% (FY2024) and 22.6% (FY2025); no '
        'statutory-vs-effective reconciliation is disclosed, and the group pays tax in 15+ '
        'jurisdictions plus Free-Zone entities on a revenue basis',
        'A sustained repeat of the FY2025 low or the Q1-2026 print (25.75%) in either direction'],
       ['Segment margins compressed rather than held at their FY2025 level',
        'Cables and Constructions margins recover PARTIALLY over the forecast; Electrical '
        'products holds closest to its FY2025 level',
        'All three disclosed segments compressed in FY2025 versus FY2023; the causes (copper '
        'pass-through, competitive pricing) are not disclosed as transitory or structural',
        'Two consecutive years of segment margin at or above the FY2023-24 average']],
      [1.35, 1.75, 2.05, 1.85], size=7.8)

# ---- negative results ---------------------------------------------------------
H1('Negative results — what could not be sourced')
P('This study was rebuilt once the company\'s own audited FY2023, FY2024 and FY2025 consolidated '
  'financial statements and its Q1-2026 condensed interim statements became available. An earlier '
  'version of this study, built before those filings were in hand, had to triangulate and derive '
  'several FY2025 figures from press coverage and company commentary; every one of those '
  'derivations has now been replaced by the audited figure itself, and the items below record what '
  'is still not disclosed anywhere in the four primary filings, rather than what could not be '
  'reached.')
table([['What was sought', 'Outcome', 'How the study handled it'],
       ['An order book, backlog or unit-volume (tonnage, MVA, meter-count) disclosure for any '
        'segment, in any of the four filings including the Q1-2026 interim',
        'Not disclosed. The company reports only segment revenue (Note 5-3) and segment profit '
        '(Note 16) — no volumes, prices or backlog',
        'The forecast is built as a taper on each segment\'s own recent revenue CAGR and, for '
        'Cables, a copper-price and FX-translation driver, rather than a reconstructed unit model'],
       ['A facility-by-facility or currency-by-currency breakdown of the debt book finer than the '
        'two-way EGP / hard-currency split disclosed in the FY2025 and Q1-2026 borrowings notes',
        'Not disclosed at finer granularity; the FY2024 filing\'s own three-way EGP/USD/EUR split '
        'was itself replaced by the simpler two-way format in FY2025',
        'The Egyptian-pound share of the book is back-solved from the independently computed '
        'effective interest rate against the two disclosed currency-bucket rates, and is labelled '
        'as inferred'],
       ['A declared, proposed or approved FY2025 dividend',
        'Not disclosed in either the FY2025 annual filing (board-approved 12 March 2026) or the '
        'Q1-2026 interim (board-approved 12 May 2026, the most recent primary source available)',
        'No FY2025 dividend per share is carried in the model; the forecast payout ratio is struck '
        'near the FY2025 disclosed cash-dividend rate relative to FY2024 profit instead'],
       ['An explanation for the sharp single-session price move on 4 August 2026',
        'No corresponding company disclosure or news item was found',
        'Not used. The study\'s anchor is the closing price on 5 August 2026 and no narrative is '
        'attached to the move']],
      [1.55, 2.35, 3.10], size=7.8)

H1('A note on aggregator and press data')
P('An earlier version of this study relied on press coverage and company commentary for FY2025, '
  'because the audited filings were not reachable from the research environment at the time. One '
  'material discrepancy was found in that process and is recorded here for the audit trail: a '
  'widely syndicated "FY2025" balance sheet — total assets, total equity, total debt and cash — '
  'reproduced the company\'s audited 31 December 2024 figures exactly, one year stale, and was '
  'discarded. All balance-sheet, income-statement and segment figures in the current study are '
  'taken directly from the audited FY2023-25 consolidated statements and the Q1-2026 condensed '
  'interim statements; no aggregator or press figure remains in the model.')

H1('Disclosure')
P('This document accompanies an educational valuation study. It is not investment advice and '
  'contains no recommendation, rating or price target. Sources are listed so that readers can '
  'verify the analysis independently. Where a figure is derived or estimated rather than '
  'disclosed, that is stated.', size=9.2, color=GREY)

out = os.path.join(HERE, 'SWDY_Bibliography_05-08-2026.docx')
doc.save(out)
print(f'wrote {out} | {len(doc.paragraphs)} paragraphs | {len(doc.tables)} tables | '
      f'{len(INP)} inputs registered')
