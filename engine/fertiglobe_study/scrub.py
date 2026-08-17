"""Checks on the delivered bibliography document, run after it is built.

(1) EXTERNAL-READER SCRUB. The document goes to an outside reader, so it must not contain
    any of the internal working vocabulary — the names of internal procedures, checks,
    scoring machinery or house tooling. Every paragraph and every table cell is scanned
    for a list of banned tokens and any hit is a FAILURE.

(2) COLUMN WIDTHS. Every table must use a fixed layout, must fit the usable page width,
    and must have no starved column (one too narrow to show its longest unbreakable word)
    and no bloated column (one carrying far more width than its content can ever use).

    python3 scrub.py [file.docx]
"""
import os
import re
import sys

from docx import Document
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT = os.path.join(HERE, 'Fertiglobe_Bibliography_09-08-2026.docx')

# --- (1) the banned vocabulary --------------------------------------------------------
BANNED = [
    r'step\s*0', r'step\s*2a', r'information sweep', r'\bsweeps?\b', r'\bswept\b',
    r'\brings?\b', r'\bgates?\b', r'\bgated\b', r'\bregisters?\b', r'\bregistered\b',
    r'\bengines?\b', r'\bsigcm\b', r'\bparity\b', r'\bcones?\b',
    r'\bstrikes?\b', r'\bstruck\b', r'\bqc\b', r'\bmateriality\b', r'\bledgers?\b',
    r'\bprotocols?\b', r'\bverdicts?\b', r'walk[- ]forward', r'\blono\b',
    r'width_cal', r'\bpanels?\b', r'roll[- ]?forward', r'\bbacktest\w*\b',
    r'\bmodel study\b', r'\breference set\b', r'\bdepth bar\b', r'\bhouse style\b',
    r'\bdriver ledger\b', r'\bcalibration appendix\b', r'\bfitted_configs\b',
    r'\bmarket_profiles\b', r'\bdata quality gate\b', r'\bpipeline\b',
]
PATS = [(p, re.compile(p, re.I)) for p in BANNED]

# --- (2) column-width geometry --------------------------------------------------------
USABLE_IN = 7.0          # 8.5in page less the two 0.75in house margins
CELL_PAD_IN = 0.125      # 90 + 90 dxa of cell margin
STARVED_TOL_IN = 0.01
BLOAT_SLACK_IN = 0.80    # dead space, in every row, before a column counts as bloated

NARROW = set("iljt.,;:'!|()[]{}fr ")
WIDE = set('mwMW')
BOLD_FACTOR = 1.06

# Measure against the font the renderer will ACTUALLY use. The document asks for Calibri;
# if no metric-compatible face is installed the renderer substitutes something wider, and
# a column sized against Calibri metrics then wraps or clips in the delivered PDF. Ask
# fontconfig which file it resolves Calibri to and measure with that.
_FONT = None
try:
    import subprocess

    from PIL import ImageFont

    _path = subprocess.run(['fc-match', '--format=%{file}', 'Calibri'],
                           capture_output=True, text=True).stdout.strip()
    if _path and os.path.exists(_path):
        _FONT = ImageFont.truetype(_path, 200)
        _FONT_NAME = os.path.basename(_path)
except Exception:
    _FONT = None


def em(ch):
    if ch in NARROW:
        return 0.30
    if ch in WIDE:
        return 0.85
    if ch.isupper():
        return 0.62
    if ch.isdigit():
        return 0.51
    return 0.50


def text_in(s, pt, bold=False):
    """Rendered width, in inches, of a string at pt in the substituted face."""
    if _FONT is not None:
        w = _FONT.getlength(s) / 200.0
    else:
        w = sum(em(c) for c in s)
    return w * pt / 72.0 * (BOLD_FACTOR if bold else 1.0)


def cell_text(cell):
    return '\n'.join(p.text for p in cell.paragraphs)


def cell_font_pt(cell, default=8.0):
    """(size in points, bold) of the first sized run in the cell."""
    for p in cell.paragraphs:
        for r in p.runs:
            if r.font.size is not None:
                return r.font.size.pt, bool(r.bold)
    return default, False


def tokens(s):
    """Break points a renderer can actually use: whitespace, and the zero-width spaces
    deliberately inserted inside long web addresses."""
    return [t for t in re.split(r'[\s​]+', s) if t]


def scrub(path):
    doc = Document(path)
    chunks = []
    for i, p in enumerate(doc.paragraphs):
        chunks.append((f'paragraph {i}', p.text))
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                chunks.append((f'table {ti + 1} r{ri} c{ci}', cell_text(cell)))

    hits = []
    for where, txt in chunks:
        clean = txt.replace('​', '')
        for pat, rx in PATS:
            for m in rx.finditer(clean):
                lo, hi = max(0, m.start() - 40), min(len(clean), m.end() + 40)
                hits.append((where, pat, m.group(0), clean[lo:hi].replace('\n', ' ')))
    return doc, chunks, hits


def columns(doc):
    problems, report = [], []
    for ti, t in enumerate(doc.tables, 1):
        widths = [c.width.inches if c.width else None for c in t.columns]
        if any(w is None for w in widths):
            problems.append(f'table {ti}: a column carries no explicit width')
            continue
        total = sum(widths)
        layout = t._tbl.tblPr.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblLayout')
        fixed = any(l.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    == 'fixed' for l in layout)
        if total > USABLE_IN + 0.02:
            problems.append(f'table {ti}: total width {total:.2f}in exceeds the usable '
                            f'{USABLE_IN:.2f}in')
        if len(widths) == 1:
            # a single-cell full-width panel (masthead, highlight box) has no columns to
            # starve or bloat and no layout to fix
            report.append(f'table {ti}: 1 col, {total:.2f}in (full-width panel) OK')
            continue
        if not fixed:
            problems.append(f'table {ti}: layout is not fixed')
        natural = [0.0] * len(widths)
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if ci >= len(widths):
                    continue
                txt = cell_text(cell)
                if not txt.strip():
                    continue
                pt, bold = cell_font_pt(cell)
                avail = widths[ci] - CELL_PAD_IN
                longest = max((text_in(w, pt, bold) for w in tokens(txt)), default=0.0)
                if longest > avail + STARVED_TOL_IN:
                    problems.append(
                        f'table {ti} col {ci} row {ri}: starved — longest word needs '
                        f'{longest:.2f}in, column allows {avail:.2f}in ('
                        f'{max(tokens(txt), key=lambda w: text_in(w, pt, bold))!r})')
                natural[ci] = max(natural[ci], text_in(txt, pt, bold))
        for ci, w in enumerate(widths):
            if natural[ci] and natural[ci] < (w - CELL_PAD_IN) - BLOAT_SLACK_IN:
                problems.append(
                    f'table {ti} col {ci}: bloated — widest content is {natural[ci]:.2f}in '
                    f'but the column is {w:.2f}in')
        report.append(f'table {ti}: {len(widths)} cols, '
                      + ' + '.join(f'{w:.2f}' for w in widths)
                      + f' = {total:.2f}in')
    return problems, report


def main():
    path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT
    doc, chunks, hits = scrub(path)
    print(f'scanned {os.path.basename(path)} — {len(chunks)} text blocks, '
          f'{len(doc.tables)} tables, {len(BANNED)} banned tokens')
    print(f'column widths measured against {_FONT_NAME if _FONT else "an estimate table"} '
          '(the face the renderer substitutes for Calibri here)')

    ok = True
    if hits:
        ok = False
        print(f'FAIL — external-reader scrub: {len(hits)} hit(s)')
        for where, pat, tok, ctx in hits[:40]:
            print(f'  {where}: /{pat}/ matched {tok!r} ... {ctx}')
    else:
        print('PASS — external-reader scrub: 0 hits across all banned tokens')

    problems, report = columns(doc)
    for line in report:
        print('  ' + line)
    if problems:
        ok = False
        print(f'FAIL — column widths: {len(problems)} problem(s)')
        for p in problems[:40]:
            print('  ' + p)
    else:
        print('PASS — column widths: every table fixed-layout, within the page, '
              'no starved or bloated column')

    sys.exit(0 if ok else 1)


if __name__ == '__main__':
    main()
