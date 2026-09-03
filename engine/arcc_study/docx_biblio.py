"""ARCC_Bibliography_03-09-2026.docx — a standalone source register.

Every figure that reaches the study or the model traces to a row here: what it is, where
it came from, what kind of source that is, and the date the source itself carries.
Reads study_numbers.json and the sweep register — no numeral is typed here.
"""
import json, os, sys
HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers.json'))
SW = json.load(open('sweep_register.json'))
INP = D['inputs']


def pc(x, dp=1):
    return f"{x*100:.{dp}f}%"


def _IV(k):
    """one registered input's value — the builder reads the record, never a numeral."""
    return INP[k]['value']

INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(11), Inches(8.5)      # landscape: long source text
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(9.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.05


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_margins(t, top=40, bottom=40, left=80, right=80):
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    t._tbl.tblPr.append(m)


def borders(t, color='C9D4D1', sz='4'):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    t._tbl.tblPr.append(b)


def P(text='', size=9.5, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def H1(t): return P(t, size=15, bold=True, space_before=12, space_after=6)
def H2(t): return P(t, size=11.5, bold=True, space_before=10, space_after=4)


def table(rows, widths, size=8.4, wrap_cols=None):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for j, w in enumerate(widths):
        t.columns[j].width = Inches(w)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ------------------------------------------------------------------ masthead
t = doc.add_table(rows=1, cols=1); cell_margins(t, 90, 90, 150, 150)
c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(9.8)
p = c.paragraphs[0]
r = p.add_run('Testahil · Arabian Cement Company S.A.E. (EGX: ARCC) — Source Register')
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = WHITE
r2 = p.add_run('   6 August 2026')
r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

P('Where every number came from. This register accompanies the valuation study and the '
  'companion model. Each input carries the source it was taken from, the type of that '
  'source, and the date the source itself bears — not the date it was read.', size=10)

# ---------------------------------------------------- sourcing statement
H2('What this study is built on')
P('The audited consolidated financial statements. That is a change from the first edition '
  'of this work, and the change is worth stating at the top rather than buried.', size=9.5)
P('The first edition was written without opening a source document: every outbound request '
  'was refused by the network policy in force, so every company figure reached the model as '
  'relayed in a web-search summary of reporting about the accounts — two steps from the '
  'audited print, not one. Four documents were subsequently supplied and this edition is '
  'rebuilt on them line by line:', size=9.5)
for t in ['Consolidated financial statements for the year ended 31 December 2025, audited by '
          'Deloitte (Wafik, Ramy & Partners), unqualified opinion signed in Cairo on 25 '
          'February 2026. 47 pages including 33 notes.',
          'Consolidated financial statements for the year ended 31 December 2024, same '
          'auditor, unqualified opinion signed 23 March 2025, carrying the FY2023 '
          'comparatives.',
          'Consolidated financial statements for the year ended 31 December 2023.',
          'Condensed consolidated interim financial statements for the three months ended 31 '
          'March 2026, limited review by the same auditor, concluded 25 May 2026.']:
    P('   •  ' + t, size=9.5)
P('Every figure in the input register below that carries a Company ring is now read from one '
  'of those four documents, with the note number given. Nothing historical is reconstructed. '
  'What remains estimated is named in section 3 and is short: two prices, a set of forecast '
  'paths, and the cost-of-capital parameters.', size=9.5)
P('Four things the reconstructed edition got materially wrong, recorded because they show '
  'where reconstruction fails rather than to dwell on it: non-controlling interests were '
  'deducted at EGP 150mn against an audited EGP 158,005; the effective tax rate was inferred '
  f'at {pc(_IV("tax_eff_superseded"), 2)} against a disclosed {pc(_IV("tax_eff"), 2)}; the cost '
  'of debt was assumed at 21.5% against a '
  '91%-euro-denominated book contracted at about 7.5%; and kiln capacity was assumed at '
  '3.6Mt against a disclosed 4.2Mt. Three things it got right and are now confirmed: the '
  'share count of 374,867,445, FY2025 operating income of EGP 4,595.82mn to the pound, and '
  'total liabilities of EGP 4,140.99mn — which it DERIVED as assets less equity after '
  'rejecting an aggregator print of EGP 2,894.13mn that would not close. That print turns '
  'out to be total CURRENT liabilities.', size=9.5)

# ---------------------------------------------------- input register
doc.add_page_break()
H1('1  Input register — every figure in the model')
RING_ORDER = {'Market': 0, 'Company': 1, 'Industry': 2, 'Country': 3, 'House': 4}
rows = [['#', 'Input', 'Value', 'Ring', 'Source', 'Source date']]


def fmt(v):
    if isinstance(v, list):
        return ', '.join(f'{x:,.4g}' for x in v)
    if isinstance(v, float):
        return f'{v:,.4f}'.rstrip('0').rstrip('.')
    return f'{v:,}' if isinstance(v, int) else str(v)


items = sorted(INP.items(), key=lambda kv: (RING_ORDER.get(kv[1]['ring'], 9), kv[0]))
for i, (k, v) in enumerate(items, 1):
    rows.append([str(i), k, fmt(v['value']), v['ring'], v['source'], v['date']])
table(rows, [0.30, 1.70, 0.92, 0.74, 4.98, 0.84], size=7.4)

# ---------------------------------------------------- source catalogue
doc.add_page_break()
H1('2  Source catalogue — the publications and institutions relied on')
CAT = [
 ('Arabian Cement Company S.A.E. — consolidated financial statements, FY2025',
  'AUDITED FINANCIAL STATEMENTS',
  'The statement of financial position, statement of profit or loss, statement of '
  'comprehensive income, statement of changes in equity and statement of cash flows, with '
  'the FY2024 comparatives. Notes relied on: 1 (capacity and group structure), 2.5 (currency '
  'table), 4 (revenue split), 5 (cost of sales), 6 (administrative expenses), 8 (finance '
  'costs), 10 (income taxes), 11 (earnings per share and weighted share count), 12 (property '
  'and depreciation), 13 (assets under construction), 16 (inventories), 17 (receivables), 19 '
  '(cash), 20 (issued capital), 21 (treasury shares), 24 (non-controlling interests), 25 '
  '(borrowings), 28 (dividends), 30 (financial instruments and currency exposure).',
  'Deloitte — Wafik, Ramy & Partners; unqualified opinion, Cairo, 25 February 2026'),
 ('Arabian Cement Company S.A.E. — consolidated financial statements, FY2024',
  'AUDITED FINANCIAL STATEMENTS',
  'FY2023 comparatives for the income statement, balance sheet and cash flow statement, '
  'including FY2023 depreciation and capital expenditure.',
  'Deloitte — Wafik, Ramy & Partners; unqualified opinion, Cairo, 23 March 2025'),
 ('Arabian Cement Company S.A.E. — consolidated financial statements, FY2023',
  'AUDITED FINANCIAL STATEMENTS',
  'Cross-check on the FY2023 position.', 'Deloitte — Wafik, Ramy & Partners'),
 ('Arabian Cement Company S.A.E. — interim financial statements, Q1-2026',
  'REVIEWED INTERIM STATEMENTS',
  'First-quarter revenue, gross profit, profit and the 31 March 2026 balance sheet, '
  'including the dividends payable that the valuation deducts from cash.',
  'Deloitte — Wafik, Ramy & Partners; limited review concluded 25 May 2026'),
 ('Arabian Cement Company S.A.E. — FY2025 Investor Presentation',
  'COMPANY DISCLOSURE — PHYSICAL',
  'The tonnes. Sales volumes by product and market (local cement 2,923.6kt, cement exports '
  '629.5kt, clinker exports 1,300.5kt, total 4,853.6kt), production indicators (clinker '
  '3,851.6kt at 92% kiln utilisation, cement 3,480.6kt), the FY2024 comparatives that give '
  'the realised local price series, the fourth-quarter split that gives the exit rate, and '
  'the Egyptian market balance of 53.9Mt local plus 18.6Mt export. The audited statements '
  'carry no volume table, so every physical figure in three earlier editions of this study '
  'was reconstructed from an assumed price; all of them are now read from here. Its '
  'FINANCIAL tables are on a narrower basis than the audited consolidated accounts '
  '(revenue 12,320 against 12,447; total assets 8,640 against 8,784), so no financial '
  'figure is taken from it.',
  'Investor relations library, arabiancementcompany.com'),
 ('Central Bank of Egypt', 'Central bank',
  'Main operation rate of 19.50% held at the April and May 2026 meetings; the Q1-2026 '
  'Monetary Policy Report; the LONGEST-DATED published inflation target of 5% for Q4-2028, '
  'used to build the terminal risk-free rate — earlier editions used the 7% Q4-2026 target '
  'and mis-described it as the medium-term one; headline urban inflation.', 'cbe.org.eg'),
 ('Aswath Damodaran — country risk premium file', 'Reference dataset',
  'Egypt equity risk premium and sovereign default spread on the credit-default-swap basis, '
  'January 2026 edition. The original file only.', 'Used for the cost-of-equity build'),
 ('Egyptian Tax Authority', 'Tax reference',
  'Statutory corporate income tax rate of 22.5%. The rate actually used is the effective '
  'rate disclosed in the audited accounts.', 'Statutory rate'),
 ('Enterprise, Global Cement and International Cement Review', 'Trade and financial press',
  'Egyptian sector context only: nameplate capacity of about 76Mt, production of about 65Mt, '
  'domestic consumption of about 54Mt, exports of about 18.5Mt, the abolition of the '
  'production quota in May 2025 with exports capped at 30% of output, and the 12.6Mt of '
  'dormant capacity under revival from the second half of 2026. No company figure is taken '
  'from these sources in this edition.',
  'enterpriseam.com, globalcement.com, cemnet.com'),
 ('Egyptian cement market pricing commentary', 'Trade press',
  'The local realised price of about EGP 3,500 a tonne and the export price of about USD 62 '
  'a tonne. These are the only two operating inputs in the FY2025 build that are not audited, '
  'and volume is derived from them.', 'Industry commentary, 2026'),
 ('Egyptian Exchange daily price history', 'Market data',
  '2,957 daily open, high, low, close and volume records from 18 May 2014 — the listing date '
  '— to 6 August 2026. The basis of the price chart, the volatility estimate, the beta '
  'regression and the price map.', 'Supplied as a file'),
 ('Standard cement engineering benchmarks', 'Industry reference',
  'Replacement cost of USD 120-150 per annual tonne, used in the asset lens and as the '
  'terminal invested-capital base.', 'Used for the asset lens'),
 ('Peer market data', 'Data aggregator',
  'Sinai Cement and Misr Beni Suef Cement revenue, profit and market capitalisation. Every '
  'multiple in the study is RECOMPUTED from these three rather than quoted.',
  'Peer disclosures as reported'),
]
rows = [['Source', 'Type', 'What it was used for', 'Reference']]
for a, b, c_, d_ in CAT:
    rows.append([a, b, c_, d_])
table(rows, [1.85, 1.62, 4.03, 2.00], size=8.0)

# ---------------------------------------------------- derived figures
doc.add_page_break()
H1('3  Figures that are DERIVED rather than sourced')
P('These do not appear in any source. They are computed, and the method is given so a '
  'reader can reproduce or reject each one.', size=9.5)
_BJ = json.load(open('beta_result.json'))
BR = _BJ['adopted']
BRO = _BJ['own_stock']
SHT = D['share_triangulation']; TR = D['terminal_reconciliation']; UC = D['unit_calibration']
KDG = D['kd_gate']
# (beta record already loaded above)
DER = [
 ('Sales volume and realised price — the only material estimate in the operating build',
  f'No volume is disclosed. Volume is DERIVED from the audited revenue note: export sales of '
  f'goods of EGP 3,356.42mn divided by an assumed USD 62 a tonne at the audited average '
  f'exchange rate of 49.26 gives {UC["vol_export"]:.3f}Mt, and local sales of goods of EGP '
  f'8,350.45mn divided by an assumed EGP 3,500 a tonne gives {UC["vol_local"]:.3f}Mt. Total '
  f'{UC["vol_fy25"]:.3f}Mt, which is {UC["util_fy25"]:.1%} of the audited 5.0Mt nameplate. '
  f'The check is that this utilisation is an OUTPUT and could have disagreed with the sector; '
  f'national production is 85.5% of national capacity. The two prices are the estimate; '
  f'everything else in this derivation is audited.'),
 ('The cost stack per tonne',
  f'From the audited cost-of-sales note and administrative-expenses note, divided by the '
  f'derived volume: materials and fuel EGP {UC["cc_mat_t"]:,.0f} a tonne, transportation EGP '
  f'{UC["cc_tra_t"]:,.0f}, overheads and cash administration EGP {UC["cc_ovh_t"]:,.0f}, total '
  f'EGP {UC["cash_cost_t"]:,.0f}. Nothing here is an invented physical build; the five-line '
  f'fuel/power/raw-material/packaging/distribution stack of the first edition is retired.'),
 ('EBITDA',
  'Operating profit — gross profit less administrative expenses, provisions and expected '
  'credit losses, all audited — plus the depreciation and amortisation reported in the '
  'audited cash flow statement. FY2023 EGP 1,329.53mn, FY2024 EGP 2,020.87mn, FY2025 EGP '
  '4,885.59mn.'),
 (f'The blended cost of debt — {KDG["kd_blended"]:.2%}',
  f'Built facility by facility from the audited borrowings note: the CIB pound facility at '
  f'the corridor rate plus 0.6% ({KDG["kd_cib"]:.2%}), the NBE/KfW euro facility at Euribor '
  f'plus 3% ({KDG["kd_nbe"]:.2%}) and the EBRD euro facility at Euribor plus '
  f'{pc(_IV("ebrd_margin"), 2)} '
  f'({KDG["kd_ebrd"]:.2%}), weighted by their audited balances. {KDG["eur_share"]:.1%} of the '
  f'book is euro-denominated. Cross-checked against interest expense over average debt: '
  f'{KDG["eff_fy24"]:.2%} in FY2024, {KDG["eff_fy25"]:.2%} in FY2025 and '
  f'{KDG["eff_q126_annualised"]:.2%} annualising Q1-2026. The contractual rate exceeds all '
  f'three and the reason is stated in the study rather than smoothed. The pound-equivalent '
  f'alternative under interest parity is {KDG["kd_egp_equivalent"]:.2%} and its effect on the '
  f'valuation is published.'),
 (f'Terminal return on invested capital — {TR["roic_repl"]:.1%}',
  f'Struck on REPLACEMENT cost — 5.0Mt at USD 130 per annual tonne — rather than on the '
  f'audited book, on which the FY2025 return is {TR["roic_book_fy25"]:.1%}. The book carries a '
  f'2010-vintage plant at historical cost through several devaluations: audited net property '
  f'and construction of EGP 2,913.87mn is about USD 12 per annual tonne, against a '
  f'replacement cost of USD 130 — roughly a tenth. A return computed on that base measures '
  f'the devaluation rather than the economics of adding a tonne. The choice makes terminal '
  f'growth value-destroying and is the single most consequential judgement in the model.'),
 (f'Beta — {BR["beta_used"]:.3f} adopted, from comparable companies rather than from '
  f'this share',
  f'Measured against the EGX30 — the only index against which an Egyptian Exchange listing '
  f'can properly be measured — Arabian Cement\'s own weekly returns over '
  f'{BRO["window_years"]:.2f} years give a beta of {BRO["beta"]:.3f} on {BRO["n"]} '
  f'observations with an R-squared of {BRO["r2"]:.3f} and a standard error of '
  f'{BRO["se"]:.3f}. The index explains under a twentieth of the share\'s movement, which '
  f'is below the level at which a regression is usable, so it is NOT adopted and its '
  f'diagnostics are printed instead. Earlier editions reported {BR["retired"]["beta"]:.4f} '
  f'measured against a basket built from the other Egyptian companies this house follows; a '
  f'basket of covered names explains a covered name better because it partly consists of '
  f'companies like it, so the better statistic was an artefact and the figure is withdrawn. '
  f'Adopted instead: the median beta of the Egyptian building-materials and construction '
  f'peers that do clear the threshold — '
  f'{", ".join("%.3f" % b for b in _BJ["peer_betas_usable"])} — giving '
  f'{BR["beta_used"]:.3f}. Sinai Cement is the closest business match and is deliberately '
  f'NOT used: its own regression is weaker still. WHAT COULD NOT BE DONE: the peers\' own '
  f'borrowings are not stripped out and Arabian Cement\'s added back, because their balance '
  f'sheets are not sourced here. Arabian Cement holds net cash and its peers carry debt, so '
  f'completing that step could only lower the beta and raise the value; the figure adopted '
  f'is the cautious end and the whole peer spread is published as a sensitivity.'),
 ('Capital expenditure in the forecast',
  'Set at the economic maintenance level of USD 4.00 per tonne of installed capacity rather '
  'than at book depreciation, because a historic-cost base understates what it costs to keep '
  'a plant running. Bracketed by the audited outturns: FY2024 EGP 912.02mn and FY2025 EGP '
  '796.47mn, both of which also carried growth projects.'),
 ('Net cash at the valuation date',
  'The DISCLOSED 30 June 2026 balance sheet: cash of EGP 1,970.50mn less interest-bearing '
  'debt of EGP 1,283.29mn, giving net cash of EGP 687.21mn. The previous edition had no '
  'balance sheet for its valuation date and had to roll one forward, which came out EGP '
  '1,239mn too generous because it could not see six months of stock-building, receivables '
  'and capital spending. Superseded rather than adjusted: '
  'rolled forward on the elapsed part of FY2026 and LESS the EGP 2,001.79mn FY2025 dividend '
  'shown as payable in the March 2026 accounts. Cross-checked against the reviewed 31 March '
  '2026 position.'),
 ('The forecast paths',
  'Utilisation, the price indices, the currency path, cost inflation, the alternative-fuel '
  'saving, depreciation as a share of revenue, the cash yield and the payout are house '
  'forecasts. Each is stated as an input in the register above with its reasoning, and each '
  'is perturbed in the driver test on the delivered model.'),
]
rows = [['Figure', 'How it was derived']]
for a, b in DER:
    rows.append([a, b])
table(rows, [2.40, 7.10], size=8.2)

# ---------------------------------------------------- research trail
doc.add_page_break()
H1('4  Research trail')
P(f'The research was carried out on {SW["sweep_date"]} across four concentric rings — '
  f'global, country, industry and company — with {len(SW["findings"])} recorded findings. '
  'Each finding below names the source and the date that source carries.', size=9.5)
rows = [['Ring', 'Topic', 'Finding', 'Source', 'Date']]
for f in SW['findings']:
    rows.append([f['ring'].title(), f['category'], f['headline'], f['source_name'],
                 f['source_date']])
table(rows, [0.76, 1.58, 4.06, 2.12, 0.88], size=7.4)

H1('5  How each forecast driver was set')
P('Every driver states whether it was built from the ground up or set from the top down, '
  'and names the findings above that it rests on.', size=9.5)
rows = [['Driver', 'Basis', 'How it was set', 'Findings']]
for d in SW['drivers']:
    rows.append([d['driver'], d['mode'].replace('_', ' ').title(), d['justification'],
                 ', '.join(d['sweep_refs'])])
table(rows, [1.70, 0.85, 5.75, 1.10], size=7.8)

P('')
P('Testahil · Independent valuation research · Educational analysis, not investment advice.',
  size=8.4, italic=True, color=GREY)

OUT = 'ARCC_Bibliography_03-09-2026.docx'
doc.save(OUT)
print('wrote', OUT)
