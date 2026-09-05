"""The external-reader scrub and the table-column audit for ARCC's delivered files.

WHY THIS FILE EXISTS. Until the edition of 2 September 2026 this study attested
`external_reader_scrub=True` in its own model-report checklist and NOTHING RAN.
The boolean was typed. That is the exact failure [R-ENF-01] names — a rule that
can be checked must be checked from outside the thing it governs, and a
self-attested boolean is never a check — and it survived here because the
attestation looked like the same word every other study uses while the study
next door was running a real scan and reporting a real count.

The scrub reads the DELIVERED documents, not the builders that wrote them: a
check that opens the source and infers the output is checking a different file
from the one that ships. It fails, it does not warn.

    python3 scrub_gate.py           scan and write scrub_result.json
    python3 scrub_gate.py --check   scan and exit non-zero on any hit
"""
from __future__ import annotations
import glob
import json
import os
import re
import sys

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))

# Vocabulary an external reader must never meet: the machinery's own words.
# Deliberately the same list the other studies carry, so a term caught on one
# name is caught on all of them, plus the two this study's own drafts used.
FORBIDDEN = [
    "sigcm", "assert_", "r-cal", "r-sigcm", "r-enf", "r-beta", "r-std", "r-doc",
    "r-macro", "r-lens", "r-coc", "r-bridge", "r-gap", "r-fcal",
    "step 0.0", "step 2a", "information sweep", "outstanding.json", "panel.json",
    "study_numbers", "beta_regression", "research_protocol", "market_profiles",
    "mc_v3", "wacc_builder", "engine/", "crps", "gate", "checklist",
    "walk-forward training", "raw_ohlc", "data.js", "macro_path",
    "cost_of_capital.py", "lens_record", "bridge_record", "compute.py",
]

# "parity" is matched case-sensitively in capitals only: lower-case parity is an
# ordinary word in a study that discusses purchasing-power parity in every
# section about the currency, and a check that cries wolf is one everyone learns
# to ignore.
FORBIDDEN_CASED = ["PARITY", "PASS/PARITY/FAIL"]


def _texts(path):
    """Every string a reader can see: body paragraphs, tables, headers, footers.

    A WORKBOOK IS A DELIVERED DOCUMENT AND THIS SCRUB NAMED TWO OF THE THREE FILES A READER
    RECEIVES [L-350]. The .xlsx branch reads every STRING cell: a numeric cell is a model
    output the recalculation gate reconciles, a formula is skipped because data_only=False
    hands back its text and a formula is not a sentence, and a numeral inside a label or a
    note is prose that happens to live in a spreadsheet — the shape this scrub exists for.
    """
    if path.lower().endswith((".xlsx", ".xlsm")):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
        out = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for v in row:
                    if isinstance(v, str) and not v.startswith("="):
                        out.append(v)
        wb.close()
        return [x for x in out if x and x.strip()]
    d = Document(path)
    out = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                out.extend(p.text for p in c.paragraphs)
    for sec in d.sections:
        for part in (sec.header, sec.footer):
            out.extend(p.text for p in part.paragraphs)
    return [x for x in out if x and x.strip()]


# A term made only of letters, digits and spaces is matched as a WHOLE WORD; a
# term carrying punctuation (a path, a module name) is matched as a substring
# because that is how it appears. The first run of this scrub matched a bare
# "gate" inside "aggregated" and reported a clean sentence about peer data as
# internal vocabulary — a check that cries wolf is one everybody learns to
# ignore, and the fix is to make the term mean what it is meant to mean rather
# than to delete it.
_PATTERNS = []
for _t in FORBIDDEN:
    if re.fullmatch(r"[a-z0-9 .]+", _t) and not _t.endswith(".py"):
        _PATTERNS.append((_t, re.compile(r"\b" + re.escape(_t) + r"\b")))
    else:
        _PATTERNS.append((_t, re.compile(re.escape(_t))))


def scrub(path):
    hits = []
    chars = 0
    for text in _texts(path):
        chars += len(text)
        low = text.lower()
        for term, rx in _PATTERNS:
            for m in rx.finditer(low):
                i = m.start()
                hits.append({"file": os.path.basename(path), "term": term,
                             "context": text[max(0, i - 60):i + 60].strip()})
        for term in FORBIDDEN_CASED:
            if term in text:
                hits.append({"file": os.path.basename(path), "term": term,
                             "context": text[:140].strip()})
    return hits, chars


def column_audit(path):
    """Starved and bloated columns, on the delivered file's own grid.

    A column narrower than 0.7cm cannot hold a formatted number and a table
    wider than the text frame runs off the page; both are invisible to every
    other check and obvious to a reader.
    """
    d = Document(path)
    # THE FRAME IS READ OFF THE DOCUMENT, NEVER ASSUMED. The first draft of this
    # audit carried a typed 16.6cm and reported twenty-two tables as too wide;
    # the study is set on US Letter with 1.91cm margins (17.78cm of text) and the
    # bibliography on a wider sheet (24.89cm), so most of those "defects" were
    # tables exactly filling their own frame and the check was describing a
    # document that does not exist. A checker that models the layout instead of
    # reading it is checking a different file from the one that ships.
    sec = d.sections[0]
    frame_cm = (sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm)
    bad = []
    for n, t in enumerate(d.tables, 1):
        widths = []
        for gc in t._tbl.tblGrid.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridCol"):
            w = gc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")
            widths.append(int(w) / 567.0 if w else None)
        if not widths or any(w is None for w in widths):
            bad.append({"table": n, "problem": "no explicit column grid — the layout "
                                               "is whatever the renderer decides"})
            continue
        if min(widths) < 0.7:
            bad.append({"table": n, "problem": "starved column %.2fcm" % min(widths)})
        if sum(widths) > frame_cm + 0.05:
            bad.append({"table": n, "problem": "table %.2fcm against a %.2fcm text frame"
                                              % (sum(widths), frame_cm)})
    return bad


def delivered():
    """The CURRENT edition's files, by the latest date on disk — never a name
    typed here. A check that opens a superseded file by name reports on
    something nobody receives, which this repository has already paid for once."""
    out = []
    for pat in ("ARCC_Valuation_Study_*.docx", "ARCC_Bibliography_*.docx"):
        files = sorted(glob.glob(os.path.join(HERE, pat)),
                       key=lambda f: re.sub(r".*_(\d\d)-(\d\d)-(\d{4})",
                                            r"\3\2\1", f))
        if files:
            out.append(files[-1])
    # THE WORKBOOK NAMES ITS EDITION DDMMYYYY WITH NO SEPARATORS, so its date is PARSED
    # rather than the filenames sorted as text: as strings 03092026 sorts BELOW 06082026 and
    # a text sort hands this scrub the superseded 6 August file [L-067, L-350].
    xl = []
    for f in glob.glob(os.path.join(HERE, "ARCC_Valuation_Model_*.xlsx")):
        m = re.search(r"_(\d{2})(\d{2})(\d{4})_", os.path.basename(f))
        if m and not os.path.basename(f).startswith("~$"):
            xl.append((m.group(3) + m.group(2) + m.group(1), f))
    if xl:
        out.append(sorted(xl)[-1][1])
    return out


def main(argv):
    files = delivered()
    if not files:
        print("FAIL: no delivered document found to scrub. An empty result is not "
              "a clean result.")
        return 1
    result = {"files": [os.path.basename(f) for f in files], "hits": [],
              "chars": 0, "column_problems": []}
    for f in files:
        h, c = scrub(f)
        result["hits"].extend(h)
        result["chars"] += c
        if f.lower().endswith(".docx"):        # the column grid is a Word property
            result["column_problems"].extend(
                dict(file=os.path.basename(f), **b) for b in column_audit(f))
    result["clean"] = not result["hits"] and not result["column_problems"]

    with open(os.path.join(HERE, "scrub_result.json"), "w") as fh:
        json.dump(result, fh, indent=1)

    print("scrubbed %d delivered document(s), %s characters"
          % (len(files), "{:,}".format(result["chars"])))
    for f in result["files"]:
        print("   %s" % f)
    if result["hits"]:
        print("\n%d FORBIDDEN TERM(S):" % len(result["hits"]))
        for h in result["hits"][:40]:
            print("   %-22s %-16s %s" % (h["file"], h["term"], h["context"][:90]))
    else:
        print("external-reader scrub: CLEAN — 0 hits across %d terms"
              % (len(FORBIDDEN) + len(FORBIDDEN_CASED)))
    if result["column_problems"]:
        print("\n%d TABLE PROBLEM(S):" % len(result["column_problems"]))
        for b in result["column_problems"][:20]:
            print("   %-40s table %s: %s" % (b["file"], b["table"], b["problem"]))
    else:
        print("table column audit: CLEAN")
    return 0 if result["clean"] else 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
