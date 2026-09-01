"""Build Lessons_Register.docx from the same source as the Markdown.

Three renderings, one source: `lessons_register.py` -> the module a study calls,
the Markdown a reader browses, and this Word document. None of the three is
hand-edited, so none can drift from the others — the failure this project has
already had twice, once between its two protocol documents and once between a
page and the fact it was supposed to remember.
"""
import os, sys, datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

import lessons_register as LR
import build_lessons_register as MD

OUT = os.path.join(HERE, "Lessons_Register.docx")

INK = RGBColor(0x1A, 0x1D, 0x21)
MUTED = RGBColor(0x5B, 0x65, 0x70)
ACCENT = RGBColor(0x1D, 0x6F, 0xA3)
RULE = RGBColor(0xB0, 0x3A, 0x2E)

SCOPE_WORD = {
    "ALL": ("Every study", ACCENT),
    "CLASS": ("A class of company", RGBColor(0x7E, 0x4A, 0x94)),
    "STOCK": ("One company only", RULE),
}
ORIGIN_WORD = {"walk_forward_fundamental": "fundamental walk-forward test",
               "walk_forward_price": "price-engine walk-forward test",
               "critique": "outside critique", "self_audit": "self-audit",
               "build": "found while building"}


def _style(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.font.color.rgb = INK
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.12
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def para(doc, text, size=10.5, bold=False, italic=False, color=INK,
         space_after=6, align=None, left=0.0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if left:
        p.paragraph_format.left_indent = Cm(left)
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    t._tbl.tblPr.append(lay)
    # the renderer reads the GRID under fixed layout, not the cell widths
    grid = t._tbl.find(qn("w:tblGrid"))
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(int(round(w * 567))))
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = Cm(widths[i]); c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = MUTED
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].width = Cm(widths[i]); cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9); r.font.color.rgb = INK
    doc.add_paragraph()
    return t


def lesson(doc, x):
    word, colour = SCOPE_WORD[x["scope"]]
    h = doc.add_paragraph()
    r = h.add_run("%s  %s" % (x["id"], x["headline"]))
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = INK
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.keep_with_next = True

    who = ("every study" if x["scope"] == "ALL"
           else "every %s" % x["applies_to"] if x["scope"] == "CLASS"
           else "%s only" % x["applies_to"])
    tag = doc.add_paragraph()
    a = tag.add_run(word.upper())
    a.bold = True; a.font.size = Pt(8); a.font.color.rgb = colour
    b = tag.add_run("   applies to %s   ·   learned from a %s, %s"
                    % (who, ORIGIN_WORD[x["origin"]], x["source"]))
    b.font.size = Pt(8); b.font.color.rgb = MUTED
    if x["status"] != "adopted":
        c = tag.add_run("   ·   STATUS: %s" % x["status"].upper())
        c.bold = True; c.font.size = Pt(8); c.font.color.rgb = RULE
    tag.paragraph_format.space_after = Pt(5)
    tag.paragraph_format.keep_with_next = True

    para(doc, x["plain"], space_after=5)
    para(doc, "How we know.  " + x["evidence"], size=9.5, color=MUTED,
         space_after=3, left=0.5)
    para(doc, "What would overturn it.  " + x["overturned_by"], size=9.5,
         color=MUTED, space_after=8, left=0.5)


def status_section(doc):
    """Where the campaign stands — done, blocked, still to go. At the front.

    Every figure comes from the SAME functions the Markdown calls, so the two
    documents cannot disagree about what has been calibrated, what stopped, or
    what is left. Three renderings, one source, as everywhere else here.
    """
    cal, blk, rem = MD.calibrated(), MD.blocked(), MD.remaining()

    doc.add_heading("Where the calibration stands", level=1)
    para(doc, "%d calibrated  ·  %d attempted and blocked  ·  %d still to go, "
              "of %d names in the book."
         % (len(cal), len(blk), len(rem["todo"]), rem["total"]),
         size=11, bold=True, space_after=10)

    doc.add_heading("Stocks that have been calibrated", level=2)
    para(doc, "Generated from the fair-value store, never typed — the same "
              "figures the fair-value register renders, so the two cannot "
              "disagree.", italic=True, color=MUTED, size=9)
    table(doc,
          ["Stock", "Market", "Class", "Fair price before",
           "Fair price after", "Move"],
          [[r["ticker"], r["market"], r["klass"],
            MD.money(r["before"], r["ccy"]), MD.money(r["after"], r["ccy"]),
            "—" if r["move"] is None else "%+.1f%%" % r["move"]]
           for r in cal],
          [2.0, 2.2, 5.4, 2.9, 2.9, 1.8])
    for r in [x for x in cal if x["note"]]:
        para(doc, "%s — %s." % (r["ticker"], r["note"]), size=9, color=MUTED)
    para(doc, "The fair price is the central case. Where a study publishes "
              "several cases and no single point, this column carries the "
              "median of them so the before-and-after can be compared at all; "
              "the study's own range is the thing to quote, never this cell.",
         size=9, color=MUTED)

    doc.add_heading("Stocks we tried to calibrate and could not", level=2)
    para(doc, "A run that stopped is evidence too. Leaving these out would "
              "make the list above read as though nothing else was attempted.",
         italic=True, color=MUTED, size=9)
    if not blk:
        para(doc, "None. Every name attempted so far has completed.")
    for x in blk:
        para(doc, "%s — %s" % (x["ticker"], x.get("company", "")),
             bold=True, size=11, space_after=3)
        para(doc, "Blocked at %s. %s."
             % (x.get("blocked_at", "an unrecorded step"),
                x.get("reason_short", "no reason recorded")),
             bold=True, color=RULE, size=10, space_after=4)
        if x.get("reason_full"):
            para(doc, x["reason_full"], size=10)
        if x.get("documents_needed"):
            para(doc, "What would unblock it:", bold=True, size=9.5,
                 space_after=2)
            for d in x["documents_needed"]:
                para(doc, "•  " + d, size=9.5, left=0.5, space_after=2)
        if x.get("would_still_be_short"):
            para(doc, "And what would still be missing.  "
                 + x["would_still_be_short"], size=9.5, color=MUTED)
        if x.get("what_was_still_learned"):
            para(doc, "What was learned anyway.  " + x["what_was_still_learned"]
                 + (("  (%s)" % ", ".join(x["lessons"]))
                    if x.get("lessons") else ""), size=9.5, color=MUTED)
        if x.get("not_filed"):
            para(doc, "What could not be recorded.  " + x["not_filed"],
                 size=9.5, color=MUTED)

    doc.add_heading("Stocks still to be calibrated", level=2)
    para(doc, "Resolved from the campaign queue at the moment this document "
              "was built, in the order the campaign runs them — never written "
              "down and kept by hand. A name leaves this list the moment its "
              "run exists.", italic=True, color=MUTED, size=9)
    table(doc, ["Market", "To go", "Names, in queue order"],
          [[label, str(len(rows)), ", ".join(r["ticker"] for r in rows)]
           for label, rows in rem["by_market"].items()],
          [3.6, 1.4, 12.2])
    if rem["excluded"]:
        ex = ", ".join(e[0] if isinstance(e, (tuple, list)) else str(e)
                       for e in rem["excluded"])
        para(doc, "%d names excluded from the campaign by construction (%s): "
                  "no issuer, no financial statements and no drivers, so there "
                  "is no forecasting method to test."
             % (len(rem["excluded"]), ex), size=9, color=MUTED)
    para(doc, "The campaign stops after the first market to review whether the "
              "method generalises before the next begins, so this list is an "
              "order rather than a schedule.", size=9, color=MUTED)


def build(path=OUT):
    LR.assert_lessons_register()
    c = LR.counts()
    doc = Document()
    _style(doc)

    para(doc, "TESTAHIL", size=9, bold=True, color=ACCENT, space_after=2)
    para(doc, "The Lessons Register", size=22, bold=True, space_after=4)
    para(doc, "Every lesson this research has learned, in plain language, and "
              "how far each one travels.", size=12, color=MUTED, space_after=4)
    para(doc, "Generated %s from the register's own source. Not hand-written, "
              "and not hand-editable — the Word file, the Markdown and the code "
              "a study calls are three renderings of one thing."
         % datetime.date.today().strftime("%d %B %Y"),
         size=9, italic=True, color=MUTED, space_after=14)

    p = doc.add_paragraph()
    r = p.add_run("Nothing in this register binds any study.")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RULE
    r2 = p.add_run(
        "  It is a record, not a gate. No standing rule refers to it and no "
        "quality gate consults it, deliberately \u2014 the walk-forward method "
        "that produced its strongest lessons has not itself been validated. "
        "The house rule is that nothing enters the method without surviving "
        "the same out-of-sample test the forecasts must survive. Read this to "
        "think with; do not cite it as authority.")
    r2.font.size = Pt(10.5); r2.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(14)

    status_section(doc)

    doc.add_page_break()
    doc.add_heading("How to read this", level=1)
    para(doc, "A lesson is useless until you know how far it carries. Every "
              "entry is therefore tagged with one of three scopes, and the "
              "difference is not a matter of taste.")
    table(doc, ["Scope", "Means", "Who must read it"],
          [["ALL", "About method, arithmetic, or how work gets checked",
            "Every study, every time"],
           ["CLASS", "True of every company that works the same way",
            "The next study of that class only"],
           ["STOCK", "True of this one company and nothing else",
            "That company's next update only"]],
          [2.4, 7.0, 6.4])
    para(doc, "Applying a STOCK lesson to another company is superstition. "
              "Applying an ALL lesson is mandatory. The middle category is "
              "where most of the value sits, and it is the one that takes "
              "longest to fill.")
    para(doc, "Every lesson also records how it was learned, because the "
              "strength of the evidence differs enormously. A walk-forward "
              "test measures the method against real outcomes across many "
              "starting points and is the strongest evidence this research "
              "produces; an outside critique, a self-audit and a defect found "
              "while building are each weaker in turn. Where two lessons "
              "disagree, the walk-forward one wins.")
    para(doc, "And every lesson carries what would overturn it. A lesson with "
              "no falsifier is a habit, not a finding, and habits are how a "
              "method quietly stops being tested.")

    doc.add_heading("What is in here, and what is honestly missing", level=1)
    table(doc, ["", "Lessons"],
          [["Bind on every study (ALL)", c["ALL"]],
           ["Bind on a class of company (CLASS)", c["CLASS"]],
           ["Bind on a single company (STOCK)", c["STOCK"]],
           ["Total", c["total"]],
           ["", ""],
           ["From fundamental walk-forward testing",
            c["by_origin"]["walk_forward_fundamental"]],
           ["From price-engine walk-forward testing",
            c["by_origin"]["walk_forward_price"]],
           ["From outside critiques", c["by_origin"]["critique"]],
           ["From self-audits", c["by_origin"]["self_audit"]],
           ["Found while building", c["by_origin"]["build"]]],
          [11.0, 4.8])

    doc.add_heading("Two different tests are both called a walk-forward",
                    level=2)
    para(doc, "They test different machinery on different evidence, and the "
              "first edition of this register conflated them — which "
              "understated the evidence base badly and is recorded here rather "
              "than quietly corrected.")
    pop = MD.population()
    table(doc, ["", "What it tests", "Names", "Resolved forecasts"],
          [["Fundamental",
            "The forecasting method — project each driver from a past origin, "
            "score revenue, cost and profit against what happened",
            "%d (%s)" % (len(pop["fundamental"]), ", ".join(pop["fundamental"])),
            "10 origins x 5 horizons"],
           ["Price engine",
            "The probability cone — strike it at a past origin and score band "
            "coverage and a proper score against a naive rule",
            str(len(pop["price"])), str(pop["price_origins"])]],
          [3.0, 7.4, 2.6, 2.8])
    para(doc, "The price engine is well tested; the fundamental method is not. "
              "%d names carry price-engine evidence, DU (%s forecasts) and "
              "GBCO (%s) among them. The fundamental method has been through a "
              "full training run on %s alone, and that run's own record states "
              "its corrections rest on two starting points, its intervals are "
              "wide with several straddling zero, and its observations are not "
              "independent. Every lesson from the fundamental method is "
              "therefore marked PROVISIONAL; price-engine lessons are not, "
              "because they rest on %d forecasts across %d names."
         % (len(pop["price"]), dict(pop["price"]).get("DU", "no"),
            dict(pop["price"]).get("GBCO", "no"),
            " and ".join(pop["fundamental"]), pop["price_origins"],
            len(pop["price"])),
         bold=True)
    outstanding = [x for x in LR.LESSONS if x["status"] != "adopted"]
    if outstanding:
        para(doc, "Not yet acted on (%d): %s. Recorded as open rather than "
                  "quietly carried as done."
             % (len(outstanding),
                "; ".join("%s — %s" % (x["id"], x["headline"].rstrip("."))
                          for x in outstanding)))

    doc.add_page_break()
    doc.add_heading("Lessons that bind on EVERY study", level=1)
    para(doc, "Read these before starting any study, of any company, in any "
              "market.", italic=True, color=MUTED)
    for x in sorted([l for l in LR.LESSONS if l["scope"] == "ALL"],
                    key=lambda y: y["id"]):
        lesson(doc, x)

    doc.add_page_break()
    doc.add_heading("Lessons that bind on a CLASS of company", level=1)
    para(doc, "Read the section for the class being studied. Do not read "
              "across classes — a rule true of developers is not evidence "
              "about airlines.", italic=True, color=MUTED)
    for k in LR.CLASSES:
        rows = sorted([l for l in LR.LESSONS
                       if l["scope"] == "CLASS" and l["applies_to"] == k],
                      key=lambda y: y["id"])
        if not rows:
            continue
        doc.add_heading(k[0].upper() + k[1:], level=2)
        for x in rows:
            lesson(doc, x)

    doc.add_page_break()
    doc.add_heading("Lessons that bind on ONE company", level=1)
    para(doc, "Read only the section for the company being updated. These do "
              "not generalise and must not be applied to another name.",
         italic=True, color=MUTED)
    for t in sorted({l["applies_to"] for l in LR.LESSONS
                     if l["scope"] == "STOCK"}):
        doc.add_heading(t, level=2)
        for x in sorted([l for l in LR.LESSONS
                         if l["scope"] == "STOCK" and l["applies_to"] == t],
                        key=lambda y: y["id"]):
            lesson(doc, x)

    doc.add_page_break()
    doc.add_heading("How a lesson gets added", level=1)
    para(doc, "A walk-forward run's own results are read mechanically and the "
              "candidate findings are drafted with their evidence already "
              "filled in from what the run measured — so a lesson's 'how we "
              "know' clause cannot drift from the numbers behind it.")
    para(doc, "Choosing the scope is not automated, deliberately. It is a "
              "judgement, and getting it wrong is costly in both directions: "
              "too narrow and the next study repeats the mistake, too broad "
              "and one company's quirk becomes a house rule nobody can "
              "dislodge. Every draft therefore arrives unscoped and nothing "
              "enters this register until someone decides. When genuinely "
              "unsure, file at the narrower scope and widen when a second "
              "company shows the same thing — one observation is not a "
              "pattern.")
    para(doc, "Every drafted finding ends either registered or declined with a "
              "reason. A finding nobody ruled on fails the check, because an "
              "unanswered question must not be able to pass as a clean result.")
    doc.save(path)
    return path


if __name__ == "__main__":
    p = build()
    c = LR.counts()
    d = Document(p)
    print("wrote %s (%d KB)" % (os.path.basename(p),
                                os.path.getsize(p) // 1024))
    print("  %d lessons — ALL %d · CLASS %d · STOCK %d"
          % (c["total"], c["ALL"], c["CLASS"], c["STOCK"]))
    print("  %d headings, %d tables, %d paragraphs"
          % (sum(1 for x in d.paragraphs if x.style.name.startswith("Heading")),
             len(d.tables), len(d.paragraphs)))
