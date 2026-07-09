"""GBCO_Valuation_Study_08-07-2026_public.docx — python-docx builder, TMPV house style.
Palette: canvas 1C3A36 · panel EAF0EE/EFF3F1 · cream F6F1E6 · gold C0A45F · brass 896F36 · grey 6E7B77."""
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = json.load(open('study_numbers.json'))
INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
BRASS = RGBColor(0x89, 0x6F, 0x36); GOLD = RGBColor(0xC0, 0xA4, 0x5F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_PANEL2, F_CREAM = '1C3A36', 'EAF0EE', 'EFF3F1', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.left_margin = sec.right_margin = Inches(0.75)
sec.top_margin, sec.bottom_margin = Inches(0.65), Inches(0.65)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.06

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def cell_margins(table, top=40, bottom=40, left=90, right=90):
    tblPr = table._tbl.tblPr
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tblPr.append(m)

def borders(table, color='C9D4D1', sz='4'):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    tblPr.append(b)

def P(text='', size=10.5, bold=False, italic=False, color=INK, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.color.rgb = color
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def rich(parts, size=10.5, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    for txt, kw in parts:
        r = p.add_run(txt)
        r.font.size = Pt(kw.get('size', size)); r.bold = kw.get('bold', False)
        r.italic = kw.get('italic', False); r.font.color.rgb = kw.get('color', INK)
    p.paragraph_format.space_after = Pt(space_after); p.paragraph_format.space_before = Pt(space_before)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def H1(text):
    p = P(text, size=15, bold=True, space_before=14, space_after=6)
    return p

def H2(text):
    return P(text, size=12, bold=True, color=INK, space_before=10, space_after=4)

def caption(text):
    return P(text, size=8.7, italic=True, color=GREY, space_after=10)

def bullet(text, bold_head=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_head:
        r = p.add_run(bold_head); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = INK
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(3)
    return p

def table(rows, widths, header=True, first_col_bold=False, size=9.3, header_fill=F_PANEL,
          align_right_from=1, band_rows=None):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(t); borders(t)
    t.autofit = False
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, header_fill)
            if band_rows and i in band_rows:
                r.bold = True; shade(c, F_CREAM)
            if first_col_bold and j == 0 and i > 0:
                r.bold = True
            if j >= align_right_from and i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def figure(path, width, caption_text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    p.paragraph_format.space_after = Pt(2)
    caption(caption_text)

def box(lines, fill=F_CREAM):
    t = doc.add_table(rows=1, cols=1)
    borders(t, color='C0A45F', sz='6'); cell_margins(t, 120, 120, 160, 160)
    c = t.cell(0, 0); shade(c, fill); c.width = Inches(7.0)
    first = True
    for head, body in lines:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        if head:
            r = p.add_run(head); r.bold = True; r.font.size = Pt(9.8); r.font.color.rgb = INK
        r2 = p.add_run(body); r2.font.size = Pt(9.8); r2.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def masthead():
    t = doc.add_table(rows=1, cols=1)
    cell_margins(t, 90, 90, 160, 160)
    c = t.cell(0, 0); shade(c, F_DARK); c.width = Inches(7.0)
    p = c.paragraphs[0]
    r = p.add_run('Testahil · Independent Valuation Study — Educational Analysis')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p.add_run('   Not investment advice')
    r2.font.size = Pt(9.5); r2.font.color.rgb = RGBColor(0x9F, 0xB0, 0xAC)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# expose everything for the content scripts
G = dict(globals())
