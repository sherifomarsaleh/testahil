#!/usr/bin/env python3
"""STC — the external-reader scrub. Depth-bar standard 4, and it did not exist.

A DELIVERED DOCUMENT IS WRITTEN FOR SOMEBODY OUTSIDE THIS HOUSE, and the vocabulary this
house uses to run itself must not reach them. There are two instruments and they catch
different things. The shared one matches BY SHAPE — a standing-rule identifier, a
repository path — neither of which can occur innocently in a document written for a
reader, which is what makes shape-matching safe there where a word list is not. This one
matches PROCEDURE NOUNS, which need judgement about their ordinary senses: "gate" is a
house term and also an ordinary English word, so the pattern has to look at how it is used.

WHAT ITS ABSENCE COST. STC shipped with no scrub of its own and an outside audit found
eight hits across the two delivered documents: the research trail described the sweep as
running "in four rings", a table column was headed "Ring", four bullet heads read
"Global ring —", "Country ring —", "Industry ring —", and a table note said a regression
had "failed its usability gate". Every one of them is this house talking to itself.

A FALSE POSITIVE IS FIXED BY NARROWING THE PATTERN TO THE HOUSE SENSE, never by deleting
a sentence a reader needs — "the balance sheet is the gate on the bridge" would be an
ordinary use and must not fire.
"""
import glob
import os
import re
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
DOCS = ['STC_Valuation_Study_05-09-2026_public.docx',
        'STC_Bibliography_05-09-2026.docx']

#: (name, pattern, why). Each is a noun this house uses for a step of its own process.
PATTERNS = [
    ('the sweep ring', r'\b(?:four |the )?rings?\b(?=[\s,.;:—-]|$)',
     'a "ring" is this house\'s name for a research layer'),
    ('a house gate', r'\b(?:usability|hard|QC|depth|promotion|materiality|publication)\s+gate\b'
                     r'|\bgate\s+(?:passed|failed|cleared|refuses|goes red)\b'
                     r'|\bpassed (?:its|the) \w+ gate\b',
     'a "gate" is this house\'s name for a check it runs on itself'),
    ('a house register', r'\b(?:sweep|driver|lessons|input|band|beta|bridge|artefact)\s+register\b',
     'these registers are internal records, not things a reader receives'),
    ('a house step number', r'\bStep\s*0(?:\.0)?\b|\bStep\s*2A\b',
     'the numbered steps are this house\'s own procedure'),
    ('an engine module', r'\b[a-z_]+\.py\b|\bmc_v[23]\b|\bstrike_cohorts\b|\bresearch_protocol\b',
     'a module name is machinery, not a source'),
    ('a scoring verdict', r'\bPARITY\b|\bROBUST FAIL\b|\bCRPS\b',
     'the retired skill verdict may not reach a reader in any form'),
    ('a ratchet', r'\bratchet(?:ed|s)?\b|\boutstanding\.json\b',
     'a ratchet is how this house schedules its own debt'),
    ('a rule identifier', r'\[R-[A-Z]+-\d+\]|\bL-\d{3}\b|\bT-\d{2,3}\b',
     'a standing-rule identifier is internal machinery on the page'),
    ('an internal device code', r'\bdevice [A-Z]-\d+\b',
     'the device codes are this house\'s own shorthand for a table shape'),
]


def texts(path):
    import docx
    d = docx.Document(path)
    out = [(i, p.text) for i, p in enumerate(d.paragraphs)]
    for ti, t in enumerate(d.tables):
        for ri, r in enumerate(t.rows):
            for ci, c in enumerate(r.cells):
                out.append(('t%d.%d.%d' % (ti, ri, ci), c.text))
    return out


def main():
    hits, read = [], 0
    for name in DOCS:
        path = os.path.join(HERE, name)
        if not os.path.exists(path):
            print('MISSING: %s' % name)
            return 1
        read += 1
        for where, txt in texts(path):
            for label, pat, why in PATTERNS:
                for m in re.finditer(pat, txt):
                    a = max(0, m.start() - 60)
                    hits.append((name, where, label, m.group(0),
                                 txt[a:m.end() + 60].replace('\n', ' '), why))
    print('external-reader scrub: %d document(s), %d pattern class(es), %d hit(s)'
          % (read, len(PATTERNS), len(hits)))
    for name, where, label, tok, ctx, why in hits:
        print('  ! %s [%s] %s: %r — %s' % (name, where, label, tok, why))
        print('      ...%s...' % ctx)
    if hits:
        print('\nFAIL — internal-procedure vocabulary reaches a reader. Rewrite the '
              'sentence for the reader; narrow a pattern only if it fired on an ORDINARY '
              'sense of the word, never to let a house term through.')
        return 1
    print('\nOK — zero hits.')
    return 0


if __name__ == '__main__':
    sys.exit(main())
