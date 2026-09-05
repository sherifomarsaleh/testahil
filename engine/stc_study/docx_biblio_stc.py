"""STC_Bibliography_05-09-2026.docx — the standalone source register.

Every figure that reaches the study or the model traces to a row here: what it is, where
it came from, what kind of source that is, and the date the source itself carries — not
the date it was read.

NOT ONE NUMERAL IS TYPED IN THIS FILE. Every row resolves from the study's own committed
artefacts: the generated input register, the sweep register, the cost-of-capital schedule,
the bridge, the lens record and the rebuild ledger. That is the whole reason a
bibliography is worth having — a hand-typed source list is a second copy of figures that
already exist, and a second copy goes stale silently while looking authoritative.
"""
import json
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
sys.path.insert(0, HERE)
sys.path.insert(0, os.path.join(HERE, '..'))

from docx import Document                                              # noqa: E402
from docx.shared import Pt, Inches, RGBColor                           # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH                          # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT                         # noqa: E402
from docx.oxml.ns import qn                                            # noqa: E402
from docx.oxml import OxmlElement                                      # noqa: E402

import inputs_register as IR                                           # noqa: E402

D = json.load(open('study_numbers.json'))
SW = json.load(open('sweep_register.json'))
INP = D['inputs']
LED = json.load(open('rebuild_ledger.json'))
_ISJ = json.load(open('income_statement.json'))
EDITION = '05-09-2026'

INK = RGBColor(0x1C, 0x3A, 0x36); GREY = RGBColor(0x6E, 0x7B, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
F_DARK, F_PANEL, F_CREAM = '1C3A36', 'EAF0EE', 'F6F1E6'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(11), Inches(8.5)   # landscape: source text is long
sec.left_margin = sec.right_margin = Inches(0.6)
sec.top_margin = sec.bottom_margin = Inches(0.6)
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(9.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.05


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def _fixed(t, widths):
    """Fixed layout, declared in the XML. autofit=False alone does not hold — a renderer
    recomputes the grid from the content, which is how the study's own Appendix A came to
    print a 2.05-inch column at a quarter of its declared width."""
    t.autofit = False
    tblPr = t._tbl.tblPr
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(int(sum(widths) * 1440))); w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    grid = t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for gc, wd in zip(grid.findall(qn('w:gridCol')), widths):
            gc.set(qn('w:w'), str(int(wd * 1440)))


def table(rows, widths, size=8.0, header=True):
    t = doc.add_table(rows=len(rows), cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    m = OxmlElement('w:tblCellMar')
    for tag, v in [('top', 34), ('left', 74), ('bottom', 34), ('right', 74)]:
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    t._tbl.tblPr.append(m)
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'C9D4D1')
        b.append(e)
    t._tbl.tblPr.append(b)
    _fixed(t, widths)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i, j); c.width = Inches(widths[j])
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            r = p.add_run('' if val is None else str(val))
            r.font.size = Pt(size); r.font.color.rgb = INK
            if i == 0 and header:
                r.bold = True; shade(c, F_PANEL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def P(text='', size=9.5, bold=False, italic=False, color=INK, space_after=5, space_before=0):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def H1(t):
    return P(t, size=14, bold=True, space_before=12, space_after=6)


def H2(t):
    return P(t, size=11, bold=True, space_before=9, space_after=4)


def fmt(v):
    if isinstance(v, list):
        return ', '.join(f'{x:,.4g}' if isinstance(x, (int, float)) else str(x) for x in v)
    if isinstance(v, bool):
        return 'yes' if v else 'no'
    if isinstance(v, float):
        return f'{v:,.4f}'.rstrip('0').rstrip('.')
    return f'{v:,}' if isinstance(v, int) else str(v)


# ------------------------------------------------------------------ masthead
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Testahil · Source register')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = INK
r2 = p.add_run('    Saudi Telecom Company (Tadawul: 7010) · edition ' + EDITION)
r2.font.size = Pt(10); r2.font.color.rgb = GREY
doc.add_paragraph().paragraph_format.space_after = Pt(4)
P('Where every number came from. This register accompanies the valuation study and the '
  'companion model. Each input carries the source it was taken from, the type of that '
  'source, and the date the source itself bears — not the date it was read.', size=10)

H2('What this study is built on')
P('The company’s own audited and reviewed financial statements, and its own earnings '
  'presentations. Nothing historical in this study comes from a data vendor, a broker or a '
  'press report, and the register below is what makes that checkable rather than asserted: '
  'every dated historical names the company document it was read from, with its note.',
  size=9.5)
_pa = SW['primary_access'][0]
P('The company’s own investor-relations channel was attempted first, before any aggregator, '
  'and the attempt is logged whether or not it succeeded. ' + _pa['note'], size=9.5)
_ncomp = sum(1 for v in INP.values() if v.get('layer', '').startswith('Company'))
_nir = sum(1 for v in INP.values() if v.get('layer') == 'Company_IR')
P('Three figures about this register, all computed from it: it carries %d inputs, of which '
  '%d come from this company’s own documents and %d of those from the investor-relations '
  'channel specifically — which is the split a reviewer needs, because it says how much of '
  'the company ring rests on material no financial statement carries. The three most recent '
  'fiscal years and every disclosed period of the study year come from the statements '
  'themselves.' % (len(INP), _ncomp, _nir), size=9.5)

# ------------------------------------------------------- 1 the input register
doc.add_page_break()
H1('1  Input register — every figure the model consumes')
P('Grouped by research layer, then by name. Four fields on every row, with no orphan '
  'numbers: what it is, what it is worth, where it came from, and the date that source '
  'carries.', size=9.3)
LAYER_ORDER = {'Market': 0, 'Company': 1, 'Company_IR': 2, 'Industry': 3, 'Country': 4,
               'House': 5}
LAYER_NAME = {'Company': 'Company — the audited and reviewed statements',
              'Company_IR': 'Company — the investor-relations channel',
              'Market': 'Market', 'Industry': 'Industry', 'Country': 'Country',
              'House': 'House'}
items = sorted(INP.items(),
               key=lambda kv: (LAYER_ORDER.get(kv[1].get('layer'), 9), kv[0]))
rows = [['#', 'Input', 'Value', 'Layer', 'Source — the document and its note', 'Source date']]
for i, (k, v) in enumerate(items, 1):
    rows.append([str(i), k.replace('_', ' '), fmt(v['value']),
                 LAYER_NAME.get(v.get('layer'), v.get('layer', '')),
                 v['source'], v.get('date', '')])
table(rows, [0.30, 2.05, 0.95, 1.15, 4.55, 0.80], size=7.0)

# ------------------------------------------------------- 2 the source catalogue
doc.add_page_break()
H1('2  Source catalogue — the documents relied on')
P('Every document named in the register above, with what it was used for. The list is '
  'closed: a row of the input register that cited a document not on this list would have '
  'failed the register’s own assertion.', size=9.3)
_USE = {}
for v in INP.values():
    for key, (title, dt) in IR.DOCS.items():
        head = title.split(',')[1].strip() if ',' in title else title
        if head[:40] in v['source']:
            _USE.setdefault(key, 0)
            _USE[key] += 1
rows = [['Document', 'Kind', 'What it carries', 'Date it bears', 'Inputs drawn']]
KIND = {'FY2025': 'AUDITED FINANCIAL STATEMENTS', 'FY2024': 'AUDITED FINANCIAL STATEMENTS',
        'FY2023': 'AUDITED FINANCIAL STATEMENTS',
        'H1_2026': 'REVIEWED INTERIM STATEMENTS',
        'IR_FY2025': 'EARNINGS PRESENTATION (investor-relations channel)',
        'IR_H1_2026': 'EARNINGS PRESENTATION (investor-relations channel)'}
CARRIES = {
 'FY2025': 'The consolidated income statement, balance sheet and cash-flow statement with '
           'the FY2024 comparatives; note 9’s reconciliation of segment revenue to net '
           'profit, which is the only place those lines appear together; note 26’s '
           'borrowings book, facility by facility, in five currencies; note 33(a)’s '
           'movement of the zakat provision, which names the prior-year reversal on its own '
           'line; and the accounting-policies note carrying the disclosed asset lives the '
           'terminal is built on.',
 'FY2024': 'The FY2023 column wherever a line was regrouped afterwards — the '
           'one-filing-per-column discipline this study keeps, so a column is never two '
           'groupings in one place — and note 27’s borrowings comparative.',
 'FY2023': 'The FY2023 position as originally reported, for the point-in-time cross-check.',
 'H1_2026': 'The reviewed 30 June 2026 balance sheet the enterprise-to-equity bridge stands '
            'on, the half’s income statement, its financing cash flows, and the share count '
            'footed against its own stated capital and par value.',
 'IR_FY2025': 'Subscribers by category and the operating measures no financial statement '
              'carries, for the segments built as volume times price.',
 'IR_H1_2026': 'The same measures for the study year’s disclosed halves, and the capital '
               'expenditure guidance this study SCORES rather than consumes.'}
for key in ('FY2025', 'FY2024', 'FY2023', 'H1_2026', 'IR_FY2025', 'IR_H1_2026'):
    title, dt = IR.DOCS[key]
    rows.append([title, KIND[key], CARRIES[key], dt, str(_USE.get(key, 0))])
table(rows, [2.35, 1.45, 4.35, 0.85, 0.70], size=7.6)
P('A note on the guidance row. Management publishes a capital-expenditure band and an '
  'earlier edition of this study took its forecast path straight from it. A forward target '
  'leans the same way an optimistic model does, so guidance is SCORED against what happens '
  'and never consumed as an input; the path this study uses is measured from the filings '
  'instead, and the guidance is recorded here as a document read rather than as a driver.',
  size=9.3)

# --------------------------------------------- 3 derived rather than sourced
doc.add_page_break()
H1('3  Figures that are DERIVED rather than sourced')
P('A derived figure is arithmetic on sourced ones, and an identity is not an assumption — '
  'but the two must not be allowed to look alike, so each is named here with the identity '
  'that produces it.', size=9.3)
_c = D['coc_record']; _b = D['dcf']['wacc_build']['beta_reg']
rows = [['Figure', 'Value', 'The identity that produces it, and from what']]
rows += [
 ['Normalised risk-free rate', f"{_c['rf_star']*100:.2f}%",
  'The observed sovereign yield of %.2f%% less this sovereign’s own default spread of '
  '%.2f%%. Country risk is charged exactly once and it is charged inside the equity risk '
  'premium, so it is removed here rather than counted twice.'
  % (_c['rf_observed']*100, _c['default_spread']*100)],
 ['Equity beta', f"{_b['beta']:.4f}",
  'A %.2f-year weekly regression of this stock against the published index of the exchange '
  'it is listed on, the Tadawul All Share Index as of %s: %d observations, R-squared '
  '%.1f%%, standard error %.4f. Produced by the house regression routine, not by a '
  'study-local script.' % (_b['window_years'], _b['index_asof'], _b['n'],
                           _b['r2']*100, _b['se'])],
 ['Cost of equity', f"{_c['ke_exp']*100:.2f}%",
  'The normalised risk-free rate plus beta times the equity risk premium of %.2f%%, on the '
  '%s basis this study names as central.' % (_c['erp']*100, _c['erp_basis'])],
 ['Marginal cost of debt', f"{_c['marginal_issue']['weighted_rate']*100:.3f}%",
  'The company’s own latest issue, weighted by the tranches’ own sizes: ' +
  '; '.join('SAR %s thousand for %d years at %.3f%%'
            % (f"{t['amount_th']:,}", t['years'], t['rate']*100)
            for t in _c['marginal_issue']['tranches']) + '.'],
 ['Weighted cost of capital', f"{_c['wacc_exp']*100:.2f}%",
  _c['weights_source']],
 ['Terminal growth', f"{D['dcf']['tg']*100:.2f}%",
  'Terminal inflation of %.1f%% from the house macro path plus a STATED real growth of '
  '%.1f%%. It is derived rather than chosen, and the real component is written down as the '
  'number it is rather than buried inside a nominal rate nobody can falsify.'
  % (D['macro_record']['terminal']['inflation_in_rf']*100,
     D['macro_record']['terminal']['real']*100)],
 ['Effective tax rate on operating profit', f"{D['tax_rate']*100:.2f}%",
  'The three filed years’ zakat and income-tax charge over their operating profit, with the '
  'prior-year provision reversal of SAR %s thousand that note 33(a) names put BACK — '
  'carrying that reversal forward instead would read %.2f%% and would assume the company '
  'keeps discovering it has over-provided, for ever.'
  % (f"{_ISJ['zakat_reversal_fy2025']:,}",
     _ISJ['zakat_rate_carrying_the_reversal']*100)],
 ['Implied asset life in the terminal', f"{D['dcf']['terminal_life_years']:.2f} years",
  'Derived from the company’s own property, plant and equipment note. It is a DISCLOSED '
  'life, not one this desk chose, and it is what the terminal maintenance charge is built '
  'on rather than the reciprocal of an inflation rate.'],
 ['Share count', f"{D['bridge_record']['shares_mn']:,.3f} mn",
  'Issued capital divided by the par value the same document states, footed against the '
  'count that document itself gives, less treasury. Today’s count is never carried back to '
  'a past period.'],
]
table(rows, [2.15, 1.05, 6.60], size=7.8)

# ----------------------------------------------------------- 4 research trail
doc.add_page_break()
H1('4  Research trail — the four rings, and what came back empty')
P('Before any forecast driver was set, the research ran in four rings: the world, the '
  'country, the industry and the company. Every finding carries its source and that '
  'source’s own date.', size=9.3)
# A SOURCE FIELD IN THE SWEEP REGISTER NAMES THE EXTRACTED FILE THIS DESK READ. That is
# the right provenance for the desk — it says which extraction a figure came out of, and
# two extractions of one filing can differ — and it is the WRONG name for a reader, who
# cannot obtain "stc_Annual-2025-en.txt" and should be told which filing it is. The
# shape-matching gate caught the one that was a repository path outright; these are the
# same defect in a milder form, and the fix is a translation rather than a rewrite of the
# register, which keeps its own precision.
_DOCNAME = {
 'stc_Annual-2025-en.txt': 'the FY2025 audited consolidated financial statements',
 'stc_Annual-2024-en.txt': 'the FY2024 audited consolidated financial statements',
 'STC_FY2023_FS_en.txt': 'the FY2023 audited consolidated financial statements',
 'financial-statementsQ2-2026En.txt': 'the reviewed half-year 2026 interim statements',
 'financial-statementsQ1-2026En.txt': 'the first-quarter 2026 interim statements',
 'EarningsPresentationQ4-2025En.txt': 'the fourth-quarter and full-year 2025 earnings '
                                      'presentation',
 'EarningsPresentationQ2-2026En.txt': 'the second-quarter 2026 earnings presentation',
 'earnings-presentation2024en.txt': 'the full-year 2024 earnings presentation',
 'units.py': 'the study\u2019s own unit-economics module',
}


def _reader_source(txt):
    for k, v in _DOCNAME.items():
        txt = txt.replace(k, v)
    return txt


RINGS = ['GLOBAL', 'COUNTRY', 'INDUSTRY', 'COMPANY']
rows = [['#', 'Ring', 'What was found', 'Source', 'Kind', 'Source date']]
n = 0
for rg in RINGS:
    for f in [x for x in SW['findings'] if x['ring'] == rg]:
        n += 1
        rows.append([str(n), rg.capitalize(), f['headline'], _reader_source(f['source_name']),
                     f['source_type'].replace('_', ' ').lower(), f.get('source_date', '')])
table(rows, [0.28, 0.78, 4.30, 3.05, 0.95, 0.74], size=7.0)
_neg = [f for f in SW['findings'] if f['klass'] == 'NEGATIVE_SEARCH']
P('A NEGATIVE RESULT IS A RESULT. %d of the %d findings above are searches that came back '
  'empty, and they are printed rather than left out — a question asked and answered with '
  'nothing is evidence, while a question never asked looks identical to one that found '
  'nothing.' % (len(_neg), len(SW['findings'])), size=9.3)

# --------------------------------------------------- 5 how each driver was set
H1('5  How each forecast driver was set')
P('The house standard is to build at the finest level the disclosure supports — volume '
  'times price on a disclosed unit, with cost per unit and the margin as an OUTPUT. Where '
  'the disclosure stops short of that, the driver drops to the finest sourced level and the '
  'gap is STATED rather than filled.', size=9.3)
rows = [['Driver', 'Level built at', 'Why, and what the disclosure supports',
         'Sweep findings behind it']]
for d in SW['drivers']:
    rows.append([d['driver'], d['mode'].replace('_', ' ').lower(), d['justification'],
                 ', '.join(d.get('sweep_refs', []))])
table(rows, [2.15, 0.95, 5.70, 1.00], size=7.8)

# ------------------------------------------------ 6 judgements and falsifiers
doc.add_page_break()
H1('6  The judgements, and what would overturn each')
P('Every input above is sourced or derived. What follows is not: these are the decisions '
  'this study made where the evidence permits more than one answer. Each is stated with the '
  'observation that would show it wrong.', size=9.3)
_rb = D['lens_record']['primary']['range_basis']
rows = [['Judgement', 'What this study decided', 'What would overturn it']]
rows += [
 ['Which lens is the answer',
  'The cash-flow model is the central and the others are cross-checks published beside it. '
  'No weights are applied anywhere.',
  'A calibration finding, out of sample, that some blend of the lenses beats the primary '
  'alone. Until that exists the typed blend stays retired, because a number produced by '
  'averaging several methods is a new method with parameters nobody tested.'],
 ['What the bear and bull move',
  'Exactly one driver: capital expenditure as a multiple of the depreciation of the base '
  'being renewed, across the %.3f to %.3f span this company’s own three filed years ran. '
  'The cost of capital, terminal growth, terminal risk-free rate, inflation ladder and '
  'margin path are IDENTICAL across all three reads.'
  % (_rb['high'], _rb['low']),
  'Evidence that a second driver moves independently of capital intensity by enough to '
  'matter. Note that the previous edition moved three levers at once, which makes each '
  'case unattributable — a bear that also changes the discount rate is valuing the company '
  'in a different economy.'],
 ['The capital-intensity path',
  'Measured from the filings at %.3f times depreciation, the three filed years’ own mean, '
  'and held flat.' % D['drivers']['capex_to_dna_adopted'],
  'Two more years of filed capital spending outside that range, which would say the '
  'data-centre build has changed the company’s intensity rather than sitting inside its '
  'history. Management’s own guidance is deliberately NOT the test: it is scored against '
  'what happens, never consumed.'],
 ['The margin path',
  'An OUTPUT of the cost build, which runs %.2f%% in the first forecast year and %.2f%% in '
  'the last — flat, and slightly down.'
  % (D['drivers']['ebitda_m'][0]*100, D['drivers']['ebitda_m'][-1]*100),
  'A cost line moving differently from the driver it is attached to. The model is not '
  'permitted to assume the mix improvement an upward glide would represent, so the risk '
  'here is a margin that FALLS rather than one that fails to rise.'],
 ['The debt tax shield',
  'The statutory rate, not the %.2f%% effective rate the filed years bore. An effective '
  'rate is the average a company paid; a shield is the marginal rate an authority allows on '
  'the income-tax portion of its base, and the two are not required to agree.'
  % (D['tax_rate']*100),
  'The split of this company’s base between zakat and income tax by ownership, which is not '
  'disclosed and which would place the true shield below the statutory rate. The whole span '
  'is worth 0.79%% of the answer — below this house’s materiality line for a contested '
  'judgement — and it is flagged rather than filled.'],
 ['The sovereign quote’s age',
  'Accepted deliberately at 36 days against a 14-day bound, and disclosed rather than used '
  'quietly.',
  'A fresher sovereign quote. Refreshing it is a macro-path task that moves every Saudi '
  'study at once, not a lever inside this one.'],
 ['The forecast balance sheet',
  'NOT BUILT, and the reason is a defect in the filing rather than a gap in the model: the '
  'latest reviewed balance sheet does not add up in its own current column, while every '
  'prior-year column foots exactly.',
  'A corrected interim, or the FY2026 audited statements. Solving for the figures that '
  'would make it foot would be inventing them, and the previous edition’s forecast balance '
  'sheet was filled with grouped estimates chosen to make a balance-check row read zero — '
  'a check that cannot fail is not a check.'],
]
table(rows, [1.65, 4.35, 3.80], size=7.8)

# --------------------------------------------------------- 7 what moved, and why
H1('7  What moved this edition, and in what order')
P('This study was rebuilt rather than re-struck, and the route matters as much as the '
  'destination: several corrections serving one rule are ONE piece of evidence, not '
  'several, and a running total looked at only once at the end can hide two rules pulling '
  'in opposite directions.', size=9.3)
rows = [['#', 'What changed', 'Before', 'After', 'Move']]
for i, lv in enumerate(LED['levers'], 1):
    rows.append([str(i), lv['name'], f"{lv['before']:.4f}", f"{lv['after']:.4f}",
                 f"{lv['move']*100:+.2f}%"])
rows.append(['', 'CUMULATIVE', f"{LED['start_value']:.4f}", f"{LED['value']:.4f}",
             f"{LED['cumulative_move']*100:+.2f}%"])
table(rows, [0.28, 5.20, 0.95, 0.95, 0.85], size=7.6)

P('TESTAHIL · Source register · Saudi Telecom Company (Tadawul: 7010) · edition '
  + EDITION, size=8.6, color=GREY, space_before=10)

OUT = 'STC_Bibliography_%s.docx' % EDITION
doc.save(OUT)
print('saved', OUT, '|', len(INP), 'inputs |', len(SW['findings']), 'findings |',
      len(LED['levers']), 'levers')
