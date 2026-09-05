"""The directional-claim sweep, run book-wide 05-09-2026. EVIDENCE, NOT A GATE.

It is committed because its MEASUREMENT is the finding [L-349], and a false-positive rate
quoted from memory is worth nothing. Two drafts, both re-pointed rather than widened:
9.2% then 28.1%, the second WORSE because it tested more sentences and the residue is
concentrated in exactly the ones a naive detector reads wrongly.

WHY IT IS NOT WIRED INTO CI. A gate firing on one claim in four is the permanently-red
check the enforcement rules forbid, and widening its tolerance is the free parameter the
promotion rule forbids. The third option is the one the waterfall gate took: the subject
of a directional sentence is not recoverable from the page, so the test belongs with
whoever knows which value the sentence is about, which is the builder.

WHAT IT FOUND ANYWAY, and no gate in this repository can see either: ARCC's section 4
says two lenses sit ABOVE the market price when all three of its committed lens values
sit below it, listing three names under "two" and quoting two figures its record does not
hold — a sentence written for the values that study carried before it was re-struck. And
GBCO's football-field caption inverts the direction of its own central. Both are recorded
for those studies' next re-issue; neither is fixed here, because a delivered study is
re-issued on an explicit request and not in passing.

Run: python3 engine/method_reassessment/directional_sweep_05-09-2026.py
"""

import glob, json, os, re
from docx import Document

ROOT = '/home/user/testahil'
CCY = r'(?:SAR|AED|EGP|USD|QAR|KRW|INR)'
# (d) the strong form: "<value>/share, N% below spot"
PCT = re.compile(r'(' + CCY + r'\s*)?([\d,]+\.?\d*)\s*(?:/share|a share|per share)?\s*,?\s*'
                 r'(?:which is\s*)?([\d.]+)\s*%\s+(above|below)\s+(?:the\s+)?'
                 r'(?:market|spot|traded price|market price|latest known price)', re.I)
# (c) the weak form: a figure, then within 90 chars a direction word
NEAR = re.compile(CCY + r'\s*([\d,]+\.?\d*)((?:(?!' + CCY + r').){0,90}?)\b(above|below)\b'
                  r'\s+(?:the\s+)?(?:market|spot|traded price|market price)', re.I | re.S)
PROB = re.compile(r'P\s*\(|probability|odds|chance', re.I)


def spot_of(sn):
    for k in ('spot',):
        if k in sn:
            return float(sn[k])
    for path in (('meta', 'spot'), ('coc_record', 'spot')):
        cur = sn
        try:
            for p in path:
                cur = cur[p]
            return float(cur)
        except Exception:
            pass
    return None


rows, allbad = [], []
for sd in sorted(glob.glob(os.path.join(ROOT, 'engine', '*_study'))):
    tk = os.path.basename(sd)[:-6].upper()
    snp = os.path.join(sd, 'study_numbers.json')
    if not os.path.exists(snp):
        continue
    try:
        sn = json.load(open(snp))
    except Exception:
        continue
    spot = spot_of(sn)
    docs = sorted(f for f in glob.glob(os.path.join(sd, '*.docx'))
                  if 'Bibliograph' not in f and not os.path.basename(f).startswith('~'))
    if not spot or not docs:
        continue
    d = Document(docs[-1])
    txt = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                txt.append(c.text)
    blob = '\n'.join(txt)
    n_claims, bad = 0, []
    for m in PCT.finditer(blob):
        val = float(m.group(2).replace(',', ''))
        pct = float(m.group(3)); word = m.group(4).lower()
        if not (0.1 * spot < val < 10 * spot):
            continue
        n_claims += 1
        真 = 'above' if val > spot else 'below'
        gap = abs(val / spot - 1) * 100
        if word != 真 or abs(gap - pct) > 1.5:
            bad.append(('pct', val, pct, word, 真, round(gap, 1),
                        blob[max(0, m.start()-70):m.end()+30].replace('\n', ' ')))
    for m in NEAR.finditer(blob):
        val = float(m.group(1).replace(',', '')); mid = m.group(2); word = m.group(3).lower()
        if PROB.search(mid) or abs(val - spot) < 0.005 * spot:
            continue
        if not (0.1 * spot < val < 10 * spot):
            continue
        n_claims += 1
        真 = 'above' if val > spot else 'below'
        if word != 真:
            bad.append(('near', val, None, word, 真, None,
                        blob[max(0, m.start()-60):m.end()+30].replace('\n', ' ')))
    rows.append((tk, spot, n_claims, len(bad)))
    for b in bad:
        allbad.append((tk, spot) + b)

print('%-12s %9s %8s %8s' % ('study', 'spot', 'claims', 'wrong'))
tc = tb = 0
for tk, spot, n, b in rows:
    tc += n; tb += b
    print('%-12s %9.2f %8d %8d' % (tk, spot, n, b))
print('\n%d studies, %d testable directional claims, %d contradicting the record (%.1f%%)'
      % (len(rows), tc, tb, 100.0 * tb / max(tc, 1)))
print()
for r in allbad:
    print('  %s spot %.2f | %s says %s, is %s | %s' % (r[0], r[1], r[3], r[5], r[6], r[8][-155:]))
