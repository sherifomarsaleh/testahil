"""Build the plan as a Word document from PLAN_02-09-2026.md — the same source as the
page, so neither can drift from the markdown. The Gantt is not read from the markdown
at all: the fenced block under the Gantt heading is replaced by gantt.docx_table(),
which renders the schedule from gantt.ROWS as a real table (the first edition typed
it as block characters and it rendered wrong in Word).

Run from the repository root:  python3 engine/method_reassessment/build_plan_docx.py
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
import gantt  # noqa: E402

SRC = HERE / "PLAN_02-09-2026.md"
OUT = HERE / "Fundamental_Method_Reassessment_02-09-2026.docx"

INK = RGBColor(0x1C, 0x25, 0x30); MUTED = RGBColor(0x5B, 0x66, 0x72); BRASS = RGBColor(0x8F, 0x73, 0x26)
TOK = re.compile(r'(\*\*.+?\*\*|`[^`]+`|(?<![\w*])\*(?!\s).+?(?<!\s)\*(?![\w*])|\[R-[A-Z]+-\d+(?: amended)?\])')


def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_); tcPr.append(shd)


def runs(par, text, size=None, color=None):
    for piece in TOK.split(text):
        if not piece:
            continue
        if piece.startswith('**') and piece.endswith('**'):
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith('`') and piece.endswith('`'):
            r = par.add_run(piece[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt((size or 10.5) - 1)
        elif piece.startswith('[R-'):
            r = par.add_run(piece); r.font.name = 'Consolas'; r.font.color.rgb = BRASS; r.font.size = Pt((size or 10.5) - 1)
        elif piece.startswith('*') and piece.endswith('*') and len(piece) > 2:
            r = par.add_run(piece[1:-1]); r.italic = True
        else:
            r = par.add_run(piece)
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color


def restart_numbering(doc) -> int:
    """Every numbered list in the plan starts at 1: python-docx's 'List Number' style shares one
    w:num across the document, so a later list would otherwise continue the earlier one's count
    (the first edition printed Part E's list starting at 8). A fresh w:num per list, pointing at
    the same abstract definition with a start override, is the Word mechanism for a restart."""
    numbering = doc.part.numbering_part.element
    style = doc.styles['List Number']
    base_num_id = style.element.pPr.numPr.numId.val
    base = numbering.num_having_numId(base_num_id)
    new = numbering.add_num(base.abstractNumId.val)
    ovr = OxmlElement('w:lvlOverride'); ovr.set(qn('w:ilvl'), '0')
    so = OxmlElement('w:startOverride'); so.set(qn('w:val'), '1'); ovr.append(so); new.append(ovr)
    return new.numId


def build(src: Path = SRC, out: Path = OUT) -> Path:
    md = src.read_text(encoding='utf-8')
    doc = Document()
    sec = doc.sections[0]; sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2); sec.top_margin = sec.bottom_margin = Cm(2.0)
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
    for lvl, size in ((1, 20), (2, 15), (3, 12)):
        h = doc.styles['Heading %d' % lvl]; h.font.name = 'Cambria'; h.font.size = Pt(size)
        h.font.bold = (lvl != 1); h.font.color.rgb = INK
        h.paragraph_format.space_before = Pt(18 if lvl < 3 else 12); h.paragraph_format.space_after = Pt(6)
        rpr = h.element.rPr; rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rf.set(qn(a), 'Cambria')

    lines = md.split('\n'); i = 0; buf = []
    state = {'first_h1': True, 'in_gantt': False}

    def flush():
        nonlocal buf
        if buf:
            runs(doc.add_paragraph(), ' '.join(buf)); buf = []

    while i < len(lines):
        ln = lines[i]
        if ln.startswith('# '):
            flush()
            if state['first_h1']:
                p = doc.add_paragraph(); r = p.add_run('TESTAHIL · Standing Research Protocol · 02 September 2026')
                r.font.size = Pt(9); r.font.color.rgb = BRASS; r.font.name = 'Consolas'
                doc.add_heading('Fundamental Method Reassessment', level=1)
                p = doc.add_paragraph()
                r = p.add_run('Why the house is systematically pessimistic, measured cause by cause; the programme '
                              'that ends it without deferring to the price; and the timeline through the calibration '
                              'of every market.')
                r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = MUTED
                state['first_h1'] = False
            i += 1; continue
        # the Gantt lives between markers in the markdown; the docx renders it from data
        if ln.strip() == gantt.MD_BEGIN:
            flush(); j = i + 1
            while j < len(lines) and lines[j].strip() != gantt.MD_END:
                j += 1
            gantt.docx_table(doc)
            p = doc.add_paragraph()
            runs(p, 'Dark cells are work, brass cells are milestones (a hard stop or a batch publish); each week '
                    'is two half-week cells. Weeks are relative to the start of Phase 1 at the base case of five '
                    'half-windows a week; Part G re-cuts the dates from the cap measured on the first two nights.',
                 size=9, color=MUTED)
            i = j + 1; continue
        if ln.startswith('```'):
            flush(); j = i + 1; block = []
            while j < len(lines) and not lines[j].startswith('```'):
                block.append(lines[j]); j += 1
            for b in block:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
                r = p.add_run(b if b else ' '); r.font.name = 'Consolas'; r.font.size = Pt(7.5)
            doc.add_paragraph(); i = j + 1; continue
        if ln.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[-| :]+\|$', lines[i + 1].strip()):
            flush(); hdr = [c.strip() for c in ln.strip().strip('|').split('|')]; rows = []; j = i + 2
            while j < len(lines) and lines[j].startswith('|'):
                rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')]); j += 1
            t = doc.add_table(rows=1 + len(rows), cols=len(hdr)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            wide = len(hdr) >= 5
            for k, h in enumerate(hdr):
                c = t.rows[0].cells[k]; c.text = ''; runs(c.paragraphs[0], h, size=8 if wide else 9)
                for r in c.paragraphs[0].runs:
                    r.bold = True
                shade(c, 'EDF0F3')
            for ri, row in enumerate(rows, 1):
                for k, cval in enumerate(row[:len(hdr)]):
                    c = t.rows[ri].cells[k]; c.text = ''; runs(c.paragraphs[0], cval, size=8 if wide else 9)
            fixed = {5: (Cm(0.8), Cm(6.2), Cm(2.0), Cm(3.9), Cm(3.7)),          # the cause table
                     6: (Cm(2.6), Cm(1.5), Cm(2.2), Cm(1.8), Cm(1.7), Cm(6.8))}  # the Phase 2 market table
            if len(hdr) in fixed:
                t.autofit = False
                for gc, w in zip(t._tbl.tblGrid.findall(qn('w:gridCol')), fixed[len(hdr)]):
                    gc.set(qn('w:w'), str(int(w.twips)))
                for row in t.rows:
                    for k, w in enumerate(fixed[len(hdr)]):
                        row.cells[k].width = w
            doc.add_paragraph(); i = j; continue
        if ln.startswith('## '):
            flush(); doc.add_heading(ln[3:].replace('**', ''), level=2); i += 1; continue
        if ln.startswith('### '):
            flush(); doc.add_heading(ln[4:].replace('**', '').replace('`', ''), level=3); i += 1; continue
        if ln.startswith('#### '):
            flush(); doc.add_heading(ln[5:].replace('**', ''), level=4); i += 1; continue
        if ln.startswith('> '):
            flush(); p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); runs(p, ln[2:])
            pPr = p._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr'); l = OxmlElement('w:left')
            l.set(qn('w:val'), 'single'); l.set(qn('w:sz'), '12'); l.set(qn('w:color'), '8F7326'); l.set(qn('w:space'), '8')
            bdr.append(l); pPr.append(bdr)
            i += 1; continue
        if re.match(r'^\d+\. ', ln):
            flush(); j = i; num_id = restart_numbering(doc)
            while j < len(lines) and re.match(r'^\d+\. ', lines[j]):
                p = doc.add_paragraph(style='List Number'); runs(p, re.sub(r'^\d+\. ', '', lines[j])); j += 1
                pPr = p._p.get_or_add_pPr(); numPr = OxmlElement('w:numPr')
                ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); numPr.append(ilvl)
                nid = OxmlElement('w:numId'); nid.set(qn('w:val'), str(num_id)); numPr.append(nid); pPr.append(numPr)
            i = j; continue
        if ln.startswith('- '):
            flush(); j = i
            while j < len(lines) and lines[j].startswith('- '):
                p = doc.add_paragraph(style='List Bullet'); runs(p, lines[j][2:]); j += 1
            i = j; continue
        if ln.strip() == '---':
            i += 1; continue
        if ln.strip() == '':
            flush(); i += 1; continue
        if ln.startswith('*(') and ln.endswith(')*'):
            flush(); p = doc.add_paragraph(); runs(p, ln[2:-2], size=9, color=MUTED); i += 1; continue
        buf.append(ln); i += 1
    flush()

    f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = f.add_run('Fundamental Method Reassessment · 02-Sep-2026 · page '); r.font.size = Pt(8); r.font.color.rgb = MUTED
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); rr = OxmlElement('w:r'); t = OxmlElement('w:t')
    t.text = '1'; rr.append(t); fld.append(rr); f._p.append(fld)
    doc.save(out)
    return out


if __name__ == '__main__':
    print('wrote', build())
