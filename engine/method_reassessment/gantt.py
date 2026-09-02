"""The plan's Gantt — one data table, three renderings.

The first edition drew the Gantt as monospaced block characters. That failed twice
over: the week header and the bars were positioned independently, so the columns
never corresponded to the weeks even in a terminal, and Word substituted a fallback
font for the block glyphs, so the bars compressed on the page. A chart whose
positions are typed by eye is the same defect as a number typed in prose.

Here the schedule is DATA — task, phase, first and last half-week — and every
rendering (markdown table for the plan file, HTML table for the page, Word table
for the docx) is generated from it, so the three cannot disagree and a re-dating
(Part G: the Gantt is re-cut from the cap measured on the first two nights) is one
edit to the rows below.

Units: half-weeks, 1..2*WEEKS inclusive. Base case as in Part D: five half-windows
a week. Phase 1 = weeks 1-3; Phase 2 = weeks 4-13 in the standing market order.
"""
from __future__ import annotations
import argparse
import re
import sys
from pathlib import Path

WEEKS = 13
SLOTS = 2 * WEEKS

# (phase, label, first_half_week, last_half_week, kind)   kind: "work" | "milestone"
ROWS = [
    ("PHASE 1 — the method (weeks 1–3)", None, None, None, "phase"),
    ("Day 0: merge · refetch filings · stamps",            1, 1, "work"),
    ("Day 1–2: apply reviewed corrections (class A)",      1, 1, "work"),
    ("WS2 macro path · WS3 lenses · WS4 bridge",           1, 2, "work"),
    ("WS1 cost of capital",                                1, 2, "work"),
    ("WS6 macro-history archive (sourcing)",               1, 4, "work"),
    ("WS7 output gates + negative controls",               2, 3, "work"),
    ("WS5 actuator (decision rule, asserts)",              3, 4, "work"),
    ("WS8 re-issue TMGH → PHDC → ARCC",                    3, 4, "work"),
    ("WS6 valuation walk-forward, five names",             4, 5, "work"),
    ("WS10 governance, lessons, digest",                   4, 6, "work"),
    ("WS8 EGCH → AMOC · WS9 sweep report",                 5, 6, "work"),
    ("promote (LONO + era) · re-issue · publish 5",        6, 6, "milestone"),
    ("PHASE 2 — the other 85, standing market order (weeks 4–13)", None, None, None, "phase"),
    ("EGX 32 names (5 re-issue, 27 first build)",          7, 12, "work"),
    ("EGX hard stop: does the method generalise?",        13, 13, "milestone"),
    ("publish EGX batch",                                 13, 13, "milestone"),
    ("UAE 28 names (10 / 18)",                            14, 18, "work"),
    ("publish UAE batch",                                 19, 19, "milestone"),
    ("KSA 13 names (3 / 10)",                             19, 21, "work"),
    ("publish KSA batch",                                 22, 22, "milestone"),
    ("Qatar · India · Korea · USA, 12 names",             22, 24, "work"),
    ("walk-forward record refresh · promotion re-check",  25, 25, "work"),
    ("final batch publish · governance close",            26, 26, "milestone"),
]

# The rows above are tuples of two shapes; normalise once.
def rows():
    out = []
    for r in ROWS:
        if r[-1] == "phase":
            out.append({"kind": "phase", "label": r[0]})
        else:
            label, a, b, kind = r
            assert 1 <= a <= b <= SLOTS, (label, a, b)
            out.append({"kind": kind, "label": label, "a": a, "b": b})
    return out


def _check():
    """Every task sits inside the grid and the phases are contiguous in time."""
    rs = rows()
    assert rs[0]["kind"] == "phase"
    phase = 0
    last_end = {1: 0, 2: 0}
    for r in rs:
        if r["kind"] == "phase":
            phase += 1
            continue
        last_end[phase] = max(last_end[phase], r["b"])
    assert last_end[1] <= 6, "Phase 1 runs past week 3"
    assert last_end[2] == SLOTS, "Phase 2 must end on the last half-week of the grid"
    return True


# ── markdown ────────────────────────────────────────────────────────────────
def markdown() -> str:
    """One column per week; a filled week is ●●, a half-week ●, empty is blank.
    Milestones are ◆ so they read differently in plain text too."""
    hdr = "| Task | " + " | ".join(f"W{w}" for w in range(1, WEEKS + 1)) + " |"
    sep = "|---|" + "|".join([":-:"] * WEEKS) + "|"
    lines = [hdr, sep]
    for r in rows():
        if r["kind"] == "phase":
            lines.append(f"| **{r['label']}** |" + " |" * WEEKS)
            continue
        cells = []
        for w in range(1, WEEKS + 1):
            s1, s2 = 2 * w - 1, 2 * w
            h1 = r["a"] <= s1 <= r["b"]
            h2 = r["a"] <= s2 <= r["b"]
            if r["kind"] == "milestone":
                cells.append("◆" if (h1 or h2) else "")
            else:
                cells.append("●●" if (h1 and h2) else ("●" if (h1 or h2) else ""))
        lines.append(f"| {r['label']} | " + " | ".join(cells) + " |")
    lines.append("")
    lines.append("●● a full week, ● a half-week, ◆ a milestone (a hard stop or a batch publish). "
                 "Weeks are relative to the start of Phase 1 at the base case of five half-windows "
                 "a week; Part G re-cuts the dates from the cap measured on the first two nights. "
                 "Generated from `engine/method_reassessment/gantt.py`, never typed.")
    return "\n".join(lines)


# ── html ────────────────────────────────────────────────────────────────────
HTML_CSS = """.gantt{border-collapse:collapse;font-family:var(--sans);font-size:12px;width:100%;min-width:760px}
.gantt th,.gantt td{border:1px solid var(--hair);padding:0;height:22px}
.gantt th{font-weight:600;color:var(--muted);font-family:var(--mono);font-size:11px;text-align:center}
.gantt td.lbl,.gantt th.lbl{text-align:left;padding:2px 10px;white-space:nowrap;width:1%}
.gantt tr.phase td{background:var(--panel);font-weight:600;padding:3px 10px;text-align:left}
.gantt td.w{background:var(--ink)}
.gantt td.m{background:var(--brass)}
.gantt td.wk{border-left:1px solid var(--muted)}
.gantt-key{font-size:12px;color:var(--muted);margin:8px 0 0}
.gantt-key i{display:inline-block;width:14px;height:10px;vertical-align:middle;margin:0 4px 0 10px}
"""


def html() -> str:
    out = ['<table class="gantt">']
    out.append('<tr><th class="lbl">Task</th>' +
               "".join(f'<th colspan="2" class="wk">W{w}</th>' for w in range(1, WEEKS + 1)) + "</tr>")
    for r in rows():
        if r["kind"] == "phase":
            out.append(f'<tr class="phase"><td colspan="{SLOTS + 1}">{r["label"]}</td></tr>')
            continue
        cls = "m" if r["kind"] == "milestone" else "w"
        cells = []
        for s in range(1, SLOTS + 1):
            c = [cls] if r["a"] <= s <= r["b"] else []
            if s % 2 == 1:
                c.append("wk")
            cells.append(f'<td class="{" ".join(c)}"></td>' if c else "<td></td>")
        out.append(f'<tr><td class="lbl">{r["label"]}</td>' + "".join(cells) + "</tr>")
    out.append("</table>")
    out.append('<p class="gantt-key"><i style="background:var(--ink)"></i> work'
               '<i style="background:var(--brass)"></i> milestone (hard stop, batch publish) — '
               "each week is two half-week cells; weeks are relative to the start of Phase 1 at "
               "five half-windows a week; Part G re-cuts the dates from the measured cap.</p>")
    return "\n".join(out)


# ── docx ────────────────────────────────────────────────────────────────────
def docx_table(doc, label_cm=6.4, slot_cm=0.39, font_pt=7.5):
    """Append the Gantt to a python-docx Document as a fixed-layout table.
    Label column + 2*WEEKS half-week columns; week header merged across each pair."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    INK, BRASS, PANEL, HAIR, MUTED = "1C2530", "8F7326", "EDF0F3", "BFC5CC", "5B6672"
    rs = rows()
    t = doc.add_table(rows=1 + len(rs), cols=1 + SLOTS)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), HAIR); borders.append(e)
    tblPr.append(borders)

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill); tcPr.append(shd)

    def tight(cell, text="", bold=False, color=None, size=font_pt, center=False):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_with_next = True   # the chart stays on one page
        if center:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
            if color: r.font.color.rgb = RGBColor.from_string(color)
        # keep every cell the same height so bars read as bars
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        side_w = 15 if center else 60
        for side, w in (("top", 20), ("bottom", 20), ("left", side_w), ("right", side_w)):
            m = OxmlElement(f"w:{side}"); m.set(qn("w:w"), str(w)); m.set(qn("w:type"), "dxa"); mar.append(m)
        tcPr.append(mar)

    widths = [Cm(label_cm)] + [Cm(slot_cm)] * SLOTS
    # a fixed-layout table takes its geometry from tblGrid, not from the cells alone
    grid = t._tbl.tblGrid
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(int(w.twips)))
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(int(sum(w.twips for w in widths)))); tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for row in t.rows:
        row.height = Cm(0.42)
        trPr = row._tr.get_or_add_trPr(); trPr.append(OxmlElement("w:cantSplit"))
        for k, w in enumerate(widths):
            row.cells[k].width = w

    # header: label + merged week pairs
    hdr = t.rows[0]
    tight(hdr.cells[0], "Task  ·  columns are weeks 1–13", bold=True, color=MUTED); shade(hdr.cells[0], PANEL)
    for w in range(1, WEEKS + 1):
        a, b = hdr.cells[2 * w - 1], hdr.cells[2 * w]
        m = a.merge(b)
        tight(m, str(w), bold=True, color=MUTED, size=font_pt - 1, center=True); shade(m, PANEL)

    for i, r in enumerate(rs, 1):
        row = t.rows[i]
        if r["kind"] == "phase":
            m = row.cells[0].merge(row.cells[SLOTS])
            tight(m, r["label"], bold=True); shade(m, PANEL)
            continue
        tight(row.cells[0], r["label"])
        fill = BRASS if r["kind"] == "milestone" else INK
        for s in range(1, SLOTS + 1):
            tight(row.cells[s])
            if r["a"] <= s <= r["b"]:
                shade(row.cells[s], fill)
    return t


# ── writers ─────────────────────────────────────────────────────────────────
MD_BEGIN, MD_END = "<!-- gantt:begin -->", "<!-- gantt:end -->"
HTML_BEGIN, HTML_END = "<!-- gantt:begin -->", "<!-- gantt:end -->"


def _splice(text, begin, end, body, what):
    if begin in text and end in text:
        pre, rest = text.split(begin, 1)
        _, post = rest.split(end, 1)
        return pre + begin + "\n" + body + "\n" + end + post
    raise SystemExit(f"{what}: gantt markers not found")


def write_markdown(path: Path):
    text = path.read_text(encoding="utf-8")
    path.write_text(_splice(text, MD_BEGIN, MD_END, markdown(), str(path)), encoding="utf-8")


def write_html(path: Path):
    text = path.read_text(encoding="utf-8")
    if ".gantt{" not in text:
        text = text.replace("</style>", HTML_CSS + "</style>", 1)
    path.write_text(_splice(text, HTML_BEGIN, HTML_END, html(), str(path)), encoding="utf-8")


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--md", type=Path, help="plan markdown to splice the table into (between gantt markers)")
    ap.add_argument("--html", type=Path, help="page HTML to splice the table into (between gantt markers)")
    ap.add_argument("--print", choices=("md", "html"), help="print one rendering to stdout")
    a = ap.parse_args(argv)
    _check()
    if a.print == "md":
        print(markdown())
    elif a.print == "html":
        print(html())
    if a.md:
        write_markdown(a.md); print("wrote gantt into", a.md)
    if a.html:
        write_html(a.html); print("wrote gantt into", a.html)
    if not (a.print or a.md or a.html):
        print(f"gantt: {len([r for r in rows() if r['kind'] != 'phase'])} tasks over {WEEKS} weeks — ok")


if __name__ == "__main__":
    main()
