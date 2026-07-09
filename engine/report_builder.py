"""
Testahil — Word report builder (canonical study skeleton).

Emulates the TMPV study's structure and house furniture: masthead, READ FIRST box, §1–§7,
Appendices A/B/C, About, and the standing six-clause Disclosure (including the holdings clause).
The quantitative tables that come straight from the engine — the §3 probability-read table, the
§2 technical indicator table, the §3 MC percentile table, and the Appendix-B Step 0 result — are
written in full; the narrative prose and the three-expert appendix are left as clearly marked
placeholders for the analyst / model to author on top. Figures produced by house_figures are
embedded where they belong.

build_report(path, ctx) -> path
  ctx = { ticker, exch, company, spot, close_date, shares, mkt_cap, thesis, asset_class,
          prob_read (list of (label, value)), technicals (list of (indicator, reading, signal)),
          mc_percentiles (dict), step0 (Step0Result), figures {fan, dist, ladder, pit, ma} }
"""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x0F, 0x76, 0x6E)
GREY = RGBColor(0x77, 0x77, 0x77)

DISCLAIMER = [
    "1. This report is an educational valuation study published under the Testahil brand. It is not "
    "investment advice, an offer, or a solicitation to buy or sell any security.",
    "2. Testahil is not a licensed financial adviser and is not regulated by the FRA. Nothing here "
    "should be read as a personal recommendation.",
    "3. All figures are estimates derived from public information believed reliable but not "
    "guaranteed; sources and as-of dates are stated where used.",
    "4. Valuation is expressed as fair-value ranges and probability distributions — never a rating, "
    "price target, or buy/sell/hold call.",
    "5. Forward-looking statements are subject to risks and uncertainties; realized outcomes may "
    "differ materially.",
    "6. The preparer may hold, and may in the future take or dispose of, a position in the security "
    "discussed in this report.",
]


def _shade(cell, hex_fill):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(sh)


def _boxrow(doc, text, fill, color=None, italic=False, bold=False, size=None):
    t = doc.add_table(rows=1, cols=1); t.autofit = True
    c = t.cell(0, 0); _shade(c, fill)
    p = c.paragraphs[0]; r = p.add_run(text)
    if color: r.font.color.rgb = color
    r.italic = italic; r.bold = bold
    if size: r.font.size = Pt(size)
    return t


def _kvtable(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        run = t.cell(0, j).paragraphs[0].add_run(str(h)); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].paragraphs[0].add_run("" if v is None else str(v))
    return t


def _h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = TEAL
    return h


def _placeholder(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("[ANALYST NARRATIVE] " + text); r.italic = True; r.font.color.rgb = GREY
    return p


def build_report(path, ctx):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    # ---- masthead ----
    _boxrow(doc, "Testahil · Independent Valuation Study — Educational Analysis · Not investment advice",
            "0F766E", color=RGBColor(0xFF, 0xFF, 0xFF), bold=True, size=10)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run(f"{ctx['company']} — Valuation Study"); tr.bold = True; tr.font.size = Pt(18)
    sub = doc.add_paragraph(); sr = sub.add_run(f"{ctx['company']} ({ctx['exch']}: {ctx['ticker']})")
    sr.font.size = Pt(12); sr.font.color.rgb = TEAL
    line = doc.add_paragraph(); lr = line.add_run(
        "Fundamental analysis · Technical analysis · Monte Carlo simulation — one integrated read")
    lr.italic = True; lr.font.color.rgb = GREY
    anchor = doc.add_paragraph()
    anchor.add_run(f"Spot {ctx['spot']:.2f} · close {ctx['close_date']} · shares {ctx['shares']:.1f}mn · "
                   f"mkt cap {ctx['mkt_cap']} · {ctx['thesis']}").font.size = Pt(9.5)

    # ---- READ FIRST ----
    _boxrow(doc,
            "READ FIRST — This is an educational study. Testahil publishes fair-value ranges and "
            "probability distributions, never a rating or price target. The fundamental value gap "
            "(§1) is kept out of the Monte Carlo drift (§3): the paths map where price could go from "
            "today, not where value sits.", "FFF7E6", italic=True, size=9.5)

    _h(doc, "Headline", 1)
    _placeholder(doc, "3–4 sentence thesis: what the name is, the fair-value range vs spot, and the "
                      "single crux that decides it.")

    _h(doc, "Company overview", 1)
    _kvtable(doc, ["Item", "Value"], [
        ["Ticker", f"{ctx['exch']}: {ctx['ticker']}"], ["Spot", f"{ctx['spot']:.2f}"],
        ["Shares (mn)", f"{ctx['shares']:.1f}"], ["Market cap", ctx["mkt_cap"]],
        ["Asset class", ctx["asset_class"]], ["Close date", ctx["close_date"]],
    ])
    _placeholder(doc, "Business description, segments, ownership, and where earnings sit in the cycle.")

    # ---- §1 fundamental ----
    _h(doc, "1  Fundamental valuation", 1)
    for s, t in [("1.1", "Primary lens"), ("1.2", "Cash-flow cross-check"),
                 ("1.3", "Relative multiples"), ("1.4", "Normalized earnings power / cycle framing"),
                 ("1.5", "Synthesis — five-lens weights & football field"),
                 ("1.6", "Deeper asset / segment look with scorecard"),
                 ("1.7", "The crux / swing-factor (sensitized in real observable units)"),
                 ("1.8", "Macro / country"), ("1.9", "Primary-lens sensitivity grid")]:
        _h(doc, f"{s}  {t}", 2)
        _placeholder(doc, f"{t}: workings, sourced driver table where bottom-up, else top-down "
                          f"normalized margins. Cost of equity published as rf + ERP, beta 0.8–1.3, "
                          f"sensitized on a WACC × terminal-g grid.")

    # ---- §2 technical ----
    _h(doc, "2  Technical and price structure", 1)
    _kvtable(doc, ["Indicator", "Reading", "Signal"], ctx["technicals"])
    if ctx["figures"].get("ma"):
        doc.add_picture(ctx["figures"]["ma"], width=Inches(6.2))
    _placeholder(doc, "Read the tape: trend vs the MA stack, RSI/MACD, and how an over/under-stretched "
                      "tape widens the near-term tail.")

    # ---- §3 Monte Carlo ----
    _h(doc, "3  Monte Carlo — a probabilistic price map", 1)
    p = doc.add_paragraph(); pr = p.add_run(
        "These paths diffuse from spot and deliberately do not embed the fundamental NAV. "
        "50,000 paths, seed 42; YZ-HAR width; unit-variance Student-t(5) innovations; "
        "asset-class-conditional drift.")
    pr.italic = True; pr.font.color.rgb = GREY
    _h(doc, "The probability read", 2)
    _kvtable(doc, ["Read (T+60)", "Value"], [[k, v] for k, v in ctx["prob_read"]])
    _h(doc, "Percentile table", 2)
    hdr = ["Horizon", "P5", "P25", "P50", "P75", "P95"]
    rows = [[name] + [f"{d[k]:.2f}" for k in (5, 25, 50, 75, 95)]
            for name, d in ctx["mc_percentiles"].items()]
    _kvtable(doc, hdr, rows)
    for key in ("fan", "dist", "ladder"):
        if ctx["figures"].get(key):
            doc.add_picture(ctx["figures"][key], width=Inches(6.2))

    # ---- §4–§7 ----
    for s, t in [("4", "Comparison of the lenses, and a verdict"), ("5", "Catalysts to watch"),
                 ("6", "Reading the probability zones"), ("7", "Caveats and what would change our mind")]:
        _h(doc, f"{s}  {t}", 1)
        _placeholder(doc, f"{t}.")

    # ---- Appendix A ----
    _h(doc, "Appendix A  Financial statements", 1)
    for s, t in [("A.1", "Income statement"), ("A.2", "Balance sheet"), ("A.3", "Cash flow")]:
        _h(doc, f"{s}  {t}", 2)
        _kvtable(doc, ["Line", "FY-3", "FY-2", "FY-1", "F+1", "F+2"],
                 [["…", "", "", "", "", ""]])
        _placeholder(doc, f"{t}: 3 years historical + 5-year forecast from primary filings.")

    # ---- Appendix B  Step 0 ----
    _h(doc, "Appendix B  Step 0 calibration backtest", 1)
    s0 = ctx["step0"]
    _kvtable(doc, ["Metric", "Value"], [
        ["Horizon", f"{s0.horizon} trading days"],
        ["Non-overlapping origins", str(s0.n_origins)],
        ["CRPS skill vs RW-lognormal", f"{s0.crps_skill:+.3f} ({'PASS' if s0.passed else 'FAIL'})"],
        ["Winkler skill", f"{s0.winkler_skill:+.3f}"],
        ["CRPS model / benchmark", f"{s0.crps_model:.3f} / {s0.crps_bench:.3f}"],
    ])
    _kvtable(doc, ["Interval", "Empirical coverage (diagnostic)"],
             [[f"{k}%", f"{v:.0%}"] for k, v in s0.coverage.items()])
    if ctx["figures"].get("pit"):
        doc.add_picture(ctx["figures"]["pit"], width=Inches(4.8))

    _h(doc, "Appendix C  Peer set, sector structure, and risks", 1)
    _kvtable(doc, ["Peer", "Note"], [["…", ""]])
    _placeholder(doc, "Peer set and sector structure.")

    # ---- Expert appendix placeholder (authored on top) ----
    _h(doc, "Appendix D  Three-expert valuation appendix", 1)
    _placeholder(doc, "Cast three personas (Expert 1 / 2 / 3) by method from the Expert Persona "
                      "Library; each: worldview → when it works/fails → workings with tables → worked "
                      "example → sensitivity → cross-examination → verdict + falsification. Close with "
                      "'three in one room' and 'reading the divergence'.")

    _h(doc, "About this series", 1)
    doc.add_paragraph("Testahil publishes educational valuation studies — distributions, not tips.")

    _h(doc, "Disclosure & Disclaimer", 1)
    for clause in DISCLAIMER:
        doc.add_paragraph(clause).runs[0].font.size = Pt(8.5)

    foot = doc.add_paragraph(); fr = foot.add_run(
        f"Testahil · {ctx['company']} ({ctx['exch']}: {ctx['ticker']}) · educational valuation study")
    fr.font.size = Pt(8); fr.font.color.rgb = GREY

    doc.save(path)
    return path
