"""Programmatic quality-control checks on the DELIVERED files.

Four gates, all run on what actually ships:
  1. EXTERNAL-READER SCRUB — a scan of every delivered document for internal-procedure
     vocabulary and verdict tokens. Zero hits required.
  2. TABLE DISCIPLINE — every table in every document must use fixed layout with explicit
     column widths, and no column may be starved or bloated relative to its content.
  3. FIGURE DISCIPLINE — every figure must be a solid opaque canvas with no transparency.
  4. NUMERIC TRACEABILITY — no financial numeral may be typed into a builder; every builder
     must read the committed numbers file.
"""
import json, os, re, sys, glob
HERE = os.path.dirname(os.path.abspath(__file__))
import docx
import openpyxl
from PIL import Image
import numpy as np

STUDY = os.path.join(HERE, 'ADNOCDIST_Valuation_Study_09-08-2026.docx')
BIB = os.path.join(HERE, 'ADNOCDIST_Bibliography_09-08-2026.docx')
XLSX = os.path.join(HERE, 'ADNOCDIST_Valuation_Model_09082026.xlsx')

# ---- 1. external-reader scrub -------------------------------------------------
BANNED = [
    r'\bstep\s*0(\.0)?\b', r'\bstep\s*2a\b', r'\bsweep\b', r'\bsweep register\b',
    r'\binformation sweep\b', r'\bfour[- ]ring\b', r'\bring\s*[1-4]\b',
    r'\bregister\b(?!ed|ing|s\b)', r'\bQC gate\b', r'\bhard gate\b', r'\bgate\b',
    r'\bSIGCM\b', r'\bmc_v3\b', r'\bmarket_profiles\b', r'\bfitted_configs\b',
    r'\bstudy_numbers\b', r'\bwidth_cal\b', r'\bLONO\b', r'\bwalk-forward\b',
    r'\bPARITY\b', r'\bVERDICT\b', r'\bPASS\b', r'\bFAIL\b', r'\bBOUNDARY\b',
    r'\bmodel study\b', r'\bTestahil Standing\b', r'\bprotocol\b', r'\bdriver ledger\b',
    r'\bcalibration ledger\b', r'\bmateriality gate\b', r'\bpanel\b(?!ed)',
    r'\bCRPS\b', r'\bPIT\b', r'\bnu=', r'\bblock bootstrap\b',
]
# Ordinary English that collides with an internal token. 'pass-through' is the correct
# term for a fuel retailer's cost base and must not be blocked by the verdict token PASS.
ALLOW = {'expert panel', 'panel of', 'the panel does',
         'pass-through', 'passes through', 'a pass through'}


def doc_text(path):
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return '\n'.join(parts)


scrub = {}
for name, path in (('study', STUDY), ('bibliography', BIB)):
    txt = doc_text(path)
    low = txt.lower()
    hits = []
    for pat in BANNED:
        for m in re.finditer(pat, txt, re.IGNORECASE):
            ctx = txt[max(0, m.start() - 45):m.end() + 45].replace('\n', ' ')
            if any(a in ctx.lower() for a in ALLOW):
                continue
            hits.append((pat, m.group(0), ctx))
    scrub[name] = hits
    print(f'external-reader scrub, {name}: {len(hits)} hit(s)')
    for pat, g, ctx in hits[:20]:
        print(f'   {g!r:24s} ... {ctx}')

# also scan the spreadsheet's own text
wb = openpyxl.load_workbook(XLSX)
xl_hits = []
for ws in wb.worksheets:
    for row in ws.iter_rows():
        for c in row:
            if isinstance(c.value, str) and not c.value.startswith('='):
                for pat in BANNED:
                    m = re.search(pat, c.value, re.IGNORECASE)
                    if m and not any(a in c.value.lower() for a in ALLOW):
                        xl_hits.append((ws.title, c.coordinate, m.group(0), c.value[:90]))
print(f'external-reader scrub, workbook: {len(xl_hits)} hit(s)')
for sh, co, g, v in xl_hits[:20]:
    print(f'   {sh}!{co} {g!r}: {v}')

# ---- 2. table discipline -------------------------------------------------------
CHARS_PER_INCH = 15.5      # 8-9pt Calibri, conservative
tbl_problems = []
ntbl = 0
for name, path in (('study', STUDY), ('bibliography', BIB)):
    d = docx.Document(path)
    for ti, t in enumerate(d.tables):
        ntbl += 1
        layout = t._tbl.tblPr.findall(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblLayout')
        if not layout or layout[0].get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') != 'fixed':
            tbl_problems.append(f'{name} table {ti}: layout is not fixed')
        widths = []
        for col in t.columns:
            widths.append(col.width.inches if col.width else None)
        if any(w is None for w in widths):
            tbl_problems.append(f'{name} table {ti}: a column has no explicit width')
            continue
        # starved: the longest unbroken word in a column cannot fit on one line
        for j, w in enumerate(widths):
            longest_word = 0
            longest_cell = 0
            for row in t.rows:
                txt = row.cells[j].text
                longest_cell = max(longest_cell, len(txt))
                for word in txt.split():
                    longest_word = max(longest_word, len(word))
            cap = w * CHARS_PER_INCH
            if longest_word > cap:
                tbl_problems.append(
                    f'{name} table {ti} column {j}: STARVED — a {longest_word}-character word '
                    f'against {cap:.0f} characters of width ({w:.2f}in)')
            if longest_cell < cap * 0.30 and w > 1.2:
                tbl_problems.append(
                    f'{name} table {ti} column {j}: BLOATED — longest content {longest_cell} '
                    f'characters against {cap:.0f} of width ({w:.2f}in)')
print(f'\ntable discipline: {ntbl} tables checked, {len(tbl_problems)} problem(s)')
for p in tbl_problems[:25]:
    print('   ', p)

# ---- 3. figure discipline --------------------------------------------------------
fig_problems = []
figs = sorted(glob.glob(os.path.join(HERE, 'fig*.png')))
for fpath in figs:
    im = Image.open(fpath)
    if im.mode == 'RGBA' and np.array(im)[:, :, 3].min() < 255:
        fig_problems.append(f'{os.path.basename(fpath)}: contains transparent pixels')
    px = np.array(im.convert('RGB'))
    corner = px[0, 0]
    if corner.mean() < 200:
        fig_problems.append(f'{os.path.basename(fpath)}: canvas is not light '
                            f'(corner {tuple(int(x) for x in corner)})')
print(f'figure discipline: {len(figs)} figures checked, {len(fig_problems)} problem(s)')
for p in fig_problems:
    print('   ', p)

# ---- 4. numeric traceability ------------------------------------------------------
# No builder may contain a typed financial numeral. Structural constants (row/column
# offsets, font sizes, column widths, percentages of a page) are permitted; anything that
# looks like a monetary figure is not.
MONEY = re.compile(r'(?<![\w.])\d{1,3}(?:,\d{3})+(?:\.\d+)?(?![\w])'      # 1,234,567
                   r'|(?<![\w.])\d{4,}\.\d{3,}(?![\w])')                  # 1234.567890
trace_problems = []
for builder in ('docx_adnocdist.py', 'docx_biblio.py', 'build_xlsx_adnocdist.py', 'figures.py'):
    src = open(os.path.join(HERE, builder)).read()
    for i, line in enumerate(src.split('\n'), 1):
        if line.strip().startswith('#') or "'''" in line:
            continue
        for m in MONEY.finditer(line):
            trace_problems.append(f'{builder}:{i} typed numeral {m.group(0)!r}: '
                                  f'{line.strip()[:100]}')
print(f'\nnumeric traceability: {len(trace_problems)} typed financial numeral(s) in builders')
for p in trace_problems[:25]:
    print('   ', p)

res = dict(scrub_study=len(scrub['study']), scrub_bib=len(scrub['bibliography']),
           scrub_workbook=len(xl_hits), tables_checked=ntbl,
           table_problems=tbl_problems, figures_checked=len(figs),
           figure_problems=fig_problems, traceability_problems=trace_problems)
json.dump(res, open(os.path.join(HERE, 'qc_result.json'), 'w'), indent=1)

fails = []
if scrub['study'] or scrub['bibliography'] or xl_hits:
    fails.append('external-reader scrub')
if tbl_problems:
    fails.append('table discipline')
if fig_problems:
    fails.append('figure discipline')
if trace_problems:
    fails.append('numeric traceability')
assert not fails, f'quality-control checks failed: {fails}'
print('\nALL PROGRAMMATIC CHECKS PASSED')
